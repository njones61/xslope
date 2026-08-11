# Copyright 2025 Norman L. Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Checks for the Analysis Report.

What is being defended:

  A. THE CONTENT TREE — a solved model produces the sections the report is
     specified to have, in order, with the tables holding the model's own
     numbers and every figure rendered to a real file.

  B. THE TOGGLES ARE REAL — turning a section off removes it from the tree, and
     turning a parent off takes its children with it. A checkbox that changed
     nothing would pass a "the report has sections" test.

  C. THE METHOD PICKER — the picked method drives the critical-surface figure
     and the slice table, and NOT the factor-of-safety summary: EVERY method the
     solver offers is reported whichever one the detail follows.

  H. NOTHING IS SAID THAT IS NOT SO — the report describes the features the model
     actually has. A dry section gets one plain statement about water, not a
     key-value block explaining how its water loads are derived; a model with no
     piles gets no Piles section; the model checks are filtered to the analyses
     the report contains; and an empty title-page field prints no row.

  D. THE DOCUMENT — the .docx is a real OOXML package: the title is in it, the
     template's styles are referenced, the slice table sits in a landscape
     section, the table of contents is a TOC field whose cached result is the
     report's own heading list, and the running head and foot carry live
     fields — the head naming the section the page is in, over the body and not
     over the front matter. Its tables are fitted to their CONTENT — fixed columns, measured,
     ending where the content ends rather than ruled across the page, cut down
     only where the content would overrun the text width, and indented so their
     borders line up with the body text — and generating one writes one file.

  E. THE COLUMN REGISTRY — every column the registry marks report-worthy exists
     in a solved slice_df, so the registry cannot drift away from the solver.

  F. THE SHARED PLOT — ``mode="shared"`` suppresses the trial surfaces and draws
     the water line a reservoir boundary states, which is where the derived
     water loads on the same figure come from.

  J. EVERY FIGURE AND TABLE IS CITED — the convention a technical report is read
     under. Every numbered block is named by number in prose that prints
     whenever the block does, from its own section or one above it; nothing is
     cited that the options left out; and each citation is a live
     cross-reference to the bookmark on the block's caption.

  K. THE OTHER TWO ENGINES — a seepage run is documented with what it solved,
     the mesh and conductivities it solved on, and the flow that came out of it;
     a strength reduction run states its factor of safety in bold and a single
     trial states none. Neither section is built without its engine's solution.

  G. THE DIALOG — it constructs offscreen, its toggles move the options it hands
     over, and what it remembers survives a second construction.

One small LEM model is solved once and shared. A second model (a dam with a
reservoir head boundary) is loaded, not solved: the shared-plot check reads its
inputs only. The seepage and finite element models are loaded and their SHIPPED
solutions read back, so no check here spends a seepage iteration or an SSRM
bisection.
"""

import contextlib
import copy
import io
import itertools
import json
import math
import os
import re
import shutil
import sys
import tempfile
import zipfile

_HERE = os.path.dirname(os.path.abspath(__file__))
_REPO = os.path.dirname(_HERE)
if _REPO not in sys.path:
    sys.path.insert(0, _REPO)

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

REINF_XLSX = os.path.join(_REPO, "docs", "inputs", "slope", "xslope_reinf.xlsx")
DAM_XLSX = os.path.join(_REPO, "docs", "inputs", "slope", "xslope_dam.xlsx")
RFACE_XLSX = os.path.join(_REPO, "docs", "inputs", "slope", "xslope_rface.xlsx")
SIMPLE1_XLSX = os.path.join(_REPO, "docs", "inputs", "slope", "xslope_simple1.xlsx")

# Models carrying one feature each, for the checks that can only be made where
# the feature is present: a right-facing slope with real horizontal forces on its
# slices (xslope_rface carries none, so its mirror rule reads correct however it
# is written), a slope reinforced entirely by passive geosynthetic, one reinforced
# axially, and one with water in a tension crack.
MIRROR_XLSX = os.path.join(_REPO, "docs", "verification", "files", "rocscience",
                           "vp039b.xlsx")
PASSIVE_XLSX = os.path.join(_REPO, "docs", "verification", "files", "rocscience",
                            "vp088.xlsx")
AXIAL_XLSX = os.path.join(_REPO, "docs", "inputs", "slope",
                          "xslope_nail_axial.xlsx")
TENSION_XLSX = os.path.join(_REPO, "docs", "lem", "files",
                            "xslope_tension_KEY.xlsx")
#: A model whose piezometric line has a shape: a level line is stated as its one
#: elevation, so only this kind gets a table of its points to introduce.
PIEZO_XLSX = os.path.join(_REPO, "docs", "lem", "files",
                          "xslope_gsat_piezo.xlsx")

#: The width a report figure prints at: the text column of a US Letter portrait
#: page inside the report's own margins. A figure is rendered wider than this and
#: placed at it, so type in a figure reaches the page in that ratio.
_TEXT_WIDTH_IN = 6.5

#: The smallest a printed number may be and still be read. Below this the digits
#: of a slice key close up on a laser printer and the key stops keying anything.
_LEGIBLE_PT = 5.0

_SOLVED = {}
_MODELS = {}


def load_slope_data_cached(xlsx):
    """One model, loaded once however many checks read it."""
    if xlsx not in _MODELS:
        from xslope.fileio import load_slope_data
        with contextlib.redirect_stdout(io.StringIO()):
            _MODELS[xlsx] = load_slope_data(xlsx)
    return _MODELS[xlsx]


def _solved():
    """The sample model with one circular surface solved by two methods, plus a
    small synthetic search bundle. Solved once and cached: nothing here reads a
    converged value, only the report's treatment of one.

    Each bundle records the options it was solved under, the way Studio's do. A
    report asked for a method these did not solve RUNS it, under those options
    (:func:`~xslope.report.run_requested_methods`), so the recorded options are
    what keeps that run the same cheap single solve these were rather than a
    search this fixture never made."""
    if "it" in _SOLVED:
        return _SOLVED["it"]
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.slice import generate_slices
    from xslope.solve import solve_selected

    slope_data = load_slope_data(REINF_XLSX)
    ok, out = generate_slices(slope_data, circle=slope_data["circles"][0],
                              num_slices=15)
    if not ok:
        raise RuntimeError(f"the sample model produced no slices: {out}")
    slice_df, surface = out

    bundles = []
    for name in ("spencer", "bishop"):
        df = slice_df.copy()
        results = solve_selected(name, df)
        bundles.append({"slice_df": df, "failure_surface": surface,
                        "results": results, "search": None, "method": name,
                        "options": {"analysis": "single_surface",
                                    "surface": "circular", "num_slices": 15}})
    # A search bundle built from the surfaces that were solved: the search
    # section reads counts and factors of safety, not the search's internals, so
    # this exercises every line of it without spending a real search.
    bundles[0]["search"] = {
        "kind": "circular",
        "fs_cache": [{"Xo": 10.0, "Yo": 50.0, "Depth": 0.0,
                      "FS": b["results"]["FS"], "slices": b["slice_df"],
                      "failure_surface": surface,
                      "solver_result": b["results"]} for b in bundles],
        "search_path": [{"x": 10.0, "y": 50.0, "FS": None},
                        {"x": 9.0, "y": 49.0, "FS": bundles[0]["results"]["FS"]}],
        "circle_cache": None,
    }
    _SOLVED["it"] = (slope_data, {"lem": bundles})
    return _SOLVED["it"]


def _build(options=None, figure_dir=None, fast=True):
    """A content tree from the sample model.

    ``fast`` builds it with the figures switched off: a tree check reads
    sections, tables and numbers, and rendering three 300-dpi plots for each of
    the twenty trees these checks build would be most of the run time for
    nothing. The checks that are ABOUT the figures pass ``fast=False`` and get
    the shipped defaults, resolution included.
    """
    from xslope.report import build_report
    slope_data, solutions = _solved()
    opts = {"input_path": REINF_XLSX, "title": "Sample Levee"}
    if fast:
        opts.update({"pd_figure": False, "lem_inputs_figure": False,
                     "lem_search_figure": False, "lem_solution_figure": False})
    opts.update(options or {})
    if figure_dir is None:
        figure_dir = tempfile.mkdtemp(prefix="xslope_report_")
    return build_report(slope_data, solutions, opts, figure_dir)


# --------------------------------------------------------------------------
# A. the content tree
# --------------------------------------------------------------------------

#: The sections a report of this solved LEM model has, in order, with the shipped
#: defaults. The sample carries reinforcement lines and no piles, so Reinforcement
#: is here and Piles is not — the two are separate sections, each present only
#: where the model has that feature. Model Checks is opt-in and is absent.
#:
#: Project Definition holds only what is true of the model whatever is run on it
#: — the geometry and the units. Everything an engine READS stands where that
#: engine is documented: a material's shear strength, the pore pressures a
#: stability analysis takes, the loads it applies, and the reinforcement it
#: resolves onto the slice bases are all limit equilibrium inputs and are all in
#: the limit equilibrium section.
#:
#: There is no factor of safety summary: the shipped default documents ONE
#: method, and a one-row comparison table restates the number that method's own
#: section states. Everything after Loads is that method's block, headed by its
#: name. The SEARCH is inside the block: it found the critical surface for that
#: method, and two methods searched separately settle on different surfaces.
EXPECTED_SECTIONS = [
    (1, "Traceability"),
    (1, "Project Definition"),
    (1, "Limit Equilibrium Analysis"),
    (2, "Analysis Inputs"),
    (2, "Materials"),
    (2, "Loads"),
    # The reinforcement stands with the engine that reads it, below the inputs
    # that engine reads and above the answers it produced.
    (2, "Reinforcement"),
    (2, "Spencer's Method"),
    (3, "Search for the Critical Surface"),
    (3, "Results"),
    (3, "Slice Table"),
    (3, "Calculations"),
]


def test_tree():
    """The tree has the specified sections, in order, and the units statement
    leads the Project Definition rather than trailing it."""
    fails = []
    report = _build()
    got = report.section_titles()
    if got != EXPECTED_SECTIONS:
        fails.append(f"section order is {got}, expected {EXPECTED_SECTIONS}")
    if not report.meta.get("title"):
        fails.append("the report carries no title")
    if not report.meta.get("date"):
        fails.append("the report carries no date")

    # The units statement is a lead block of Project Definition: a reader meets
    # the section's numbers already knowing what they are in.
    pd_sec = next((s for s in report.sections if s.title == "Project Definition"),
                  None)
    if pd_sec is None:
        return fails + ["there is no Project Definition section"]
    kinds = [b.kind for b in pd_sec.blocks]
    prose = [b.text for b in pd_sec.blocks if b.kind == "prose"]
    units = [i for i, t in enumerate(prose) if "units" in t or "unit system" in t]
    if not units:
        fails.append(f"no units statement in the Project Definition: {prose}")
    elif units[0] != 1:
        fails.append(f"the units statement is prose {units[0]} of the section; it "
                     f"belongs immediately after the intro")
    with_fig = _build({"pd_figure": True})
    fsec = next(s for s in with_fig.sections if s.title == "Project Definition")
    kinds = [b.kind for b in fsec.blocks]
    if "figure" not in kinds:
        fails.append("the model figure is not a block of Project Definition")
    else:
        texts = [b.text for b in fsec.blocks[:kinds.index("figure")]
                 if b.kind == "prose"]
        if not any("units" in t or "unit system" in t for t in texts):
            fails.append("the units statement comes after the model figure")
    return fails


def test_tables_carry_the_model():
    """The tables hold the model's own numbers, not placeholders."""
    fails = []
    slope_data, solutions = _solved()
    report = _build()
    tables = {t.caption: t for t in report.tables()}

    mats = tables.get("Material properties")
    if mats is None:
        return ["there is no materials table"]
    if len(mats.rows) != len(slope_data["materials"]):
        fails.append(f"the materials table has {len(mats.rows)} rows for "
                     f"{len(slope_data['materials'])} materials")
    names = {r[1] for r in mats.rows}
    for m in slope_data["materials"]:
        if m["name"] not in names:
            fails.append(f"material {m['name']!r} is missing from the table")
    # Only populated columns: this model states no gamma_sat and no r_u.
    heads = " ".join(mats.headers)
    for absent in ("γ_sat", "r_u"):
        if absent in heads:
            fails.append(f"the materials table prints {absent}, which no "
                         f"material populates")
    for present in ("γ", "c", "φ"):
        if present not in heads:
            fails.append(f"the materials table omits {present}")

    # The comparison table exists only where there is a comparison, so it is
    # asked for from a report that documents both methods.
    both = _build({"method": ["spencer", "bishop"]})
    fs = {t.caption: t for t in both.tables()}.get("Computed factors of safety")
    if fs is None:
        fails.append("a two-method report carries no factor-of-safety table")
    else:
        for b in solutions["lem"]:
            want = f"{b['results']['FS']:.3f}"
            if not any(want in row for row in fs.rows):
                fails.append(f"FS {want} ({b['method']}) is not in the summary table")

    reinf = tables.get("Reinforcement lines")
    if reinf is None:
        fails.append("there is no reinforcement table")
    elif len(reinf.rows) != len(slope_data["reinforcement_lines"]):
        fails.append(f"the reinforcement table has {len(reinf.rows)} rows for "
                     f"{len(slope_data['reinforcement_lines'])} lines")

    # No header word wider than its column: Word does not wrap a word that will
    # not fit, it breaks it, and "Applicatio n" is what that looks like in print.
    from xslope.report import max_header_word
    for t in report.tables():
        if t.landscape:
            continue                        # the slice table gets a whole page
        limit = max_header_word(len(t.headers))
        for head in t.headers:
            for word in str(head).split():
                if len(word) > limit:
                    fails.append(f"{t.caption!r}: the header word {word!r} is "
                                 f"{len(word)} characters, over the {limit} a "
                                 f"{len(t.headers)}-column table allows, and will "
                                 f"break mid-word")

    slices = [t for t in report.tables() if t.landscape]
    if len(slices) != 1:
        fails.append(f"{len(slices)} landscape tables; the slice table should be "
                     f"the only one")
    elif len(slices[0].rows) != len(solutions["lem"][0]["slice_df"]):
        fails.append(f"the slice table has {len(slices[0].rows)} rows for "
                     f"{len(solutions['lem'][0]['slice_df'])} slices")
    elif not slices[0].legend:
        fails.append("the slice table carries no column legend")
    return fails


def test_figures_rendered():
    """Every figure block points at a real, non-trivial PNG."""
    fails = []
    from xslope.report import DEFAULT_OPTIONS, FIGURE_DPI

    if DEFAULT_OPTIONS["dpi"] != FIGURE_DPI or FIGURE_DPI < 300:
        fails.append(f"figures default to {DEFAULT_OPTIONS['dpi']} dpi; a report "
                     f"figure is a print figure")
    tmp = tempfile.mkdtemp(prefix="xslope_report_fig_")
    report = _build(figure_dir=tmp, fast=False)
    figures = report.figures()
    if len(figures) < 3:
        fails.append(f"only {len(figures)} figures were built (model, search, "
                     f"solution expected)")
    numbers = [f.number for f in figures]
    if numbers != sorted(set(numbers)) or (numbers and numbers[0] != 1):
        fails.append(f"figures are not numbered 1..n in order: {numbers}")
    for f in figures:
        if not os.path.exists(f.path):
            fails.append(f"figure {f.number} was not written to {f.path}")
        elif os.path.getsize(f.path) < 20000:
            fails.append(f"figure {f.number} is {os.path.getsize(f.path)} bytes — "
                         f"too small to be a rendered plot")
        if not f.caption:
            fails.append(f"figure {f.number} has no caption")
    return fails


def test_traceability():
    """The stamp identifies the inputs the answers came from."""
    from xslope.report import file_digest
    from xslope._version import __version__

    report = _build()
    kv = [b for b in report.blocks("keyvalues")]
    if not kv:
        return ["the traceability block is missing"]
    items = dict(kv[0].items)
    fails = []
    if items.get("xslope version") != __version__:
        fails.append(f"the stamp reports version {items.get('xslope version')!r}")
    if items.get("Input file") != os.path.basename(REINF_XLSX):
        fails.append(f"the stamp names {items.get('Input file')!r} as the input")
    digest = file_digest(REINF_XLSX)
    if items.get("Input file SHA-256") != digest:
        fails.append("the stamp's digest is not the input file's SHA-256")
    if len(digest) != 64:
        fails.append(f"the digest is {len(digest)} characters, not a SHA-256")

    # When the analysis was run, which is on record or is not. It used to
    # default to the moment the report was made — so a run restored from
    # companions saved months earlier was stamped with today, twice, once as the
    # analysis and once as the generation.
    from datetime import datetime, timedelta
    if "Analysis run" in items:
        fails.append(f"a run whose record carries no solve time is stamped "
                     f"{items['Analysis run']!r}")
    if "Report generated" not in items:
        fails.append("the stamp does not say when the report was generated")

    when = datetime(2019, 3, 4, 5, 6)
    told = dict(_build({"solved_at": when}).blocks("keyvalues")[0].items)
    if told.get("Analysis run") != when.strftime("%Y-%m-%d %H:%M"):
        fails.append(f"the caller gave the solve time and the stamp reads "
                     f"{told.get('Analysis run')!r}")
    if told.get("Report generated") == told.get("Analysis run"):
        fails.append("the stamp reports the solve and the generation as one "
                     "moment")

    # And where the RUN's own record carries it — the only source a restored run
    # has — the stamp reads it off the record rather than off the clock.
    from xslope.report import SOLVED_AT_KEY
    slope_data, bundle = _fem_1d_bundle(FEM_REINF_XLSX)
    recorded = datetime.now() - timedelta(days=97)
    on_record = dict(bundle, meta={SOLVED_AT_KEY: recorded.isoformat()})
    read = dict(_built_report(
        slope_data, {"fem": on_record},
        {"input_path": FEM_REINF_XLSX, "lem": False, "pd_figure": False}
    ).blocks("keyvalues")[0].items)
    if read.get("Analysis run") != recorded.strftime("%Y-%m-%d %H:%M"):
        fails.append(f"the run recorded its solve time and the stamp reads "
                     f"{read.get('Analysis run')!r}")
    return fails


# --------------------------------------------------------------------------
# B. the toggles
# --------------------------------------------------------------------------

def test_toggles():
    """Every content toggle removes what it names, and a parent takes its
    children with it."""
    fails = []
    full = [t for _lvl, t in _build().section_titles()]

    cases = [
        ({"traceability": False}, "Traceability"),
        ({"lem_search": False}, "Search for the Critical Surface"),
        ({"lem_slice_table": False}, "Slice Table"),
        ({"lem_calculations": False}, "Calculations"),
        ({"lem_materials": False}, "Materials"),
        ({"lem_loads": False}, "Loads"),
        ({"lem_members": False}, "Reinforcement"),
    ]
    for opts, gone in cases:
        titles = [t for _lvl, t in _build(opts).section_titles()]
        if gone in titles:
            fails.append(f"{opts} left {gone!r} in the report")
        if len(titles) != len(full) - 1:
            fails.append(f"{opts} changed the section count by "
                         f"{len(full) - len(titles)}, not 1")

    # The model checks are not part of a report at all: they are an interface
    # surface, and a warning about the model belongs where the model is being
    # built rather than in the submittal it produced (the owner's ruling,
    # fem_reinforce review). Asking for the section by its retired option name
    # changes nothing.
    if "Model Checks" in full:
        fails.append("Model Checks is in the report; it belongs to the "
                     "interface")
    on = [t for _lvl, t in _build({"model_checks": True}).section_titles()]
    if "Model Checks" in on:
        fails.append("the retired model_checks option still builds a section")
    if on != full:
        fails.append(f"the retired model_checks option changed the report: "
                     f"{set(on) ^ set(full)}")

    # The units statement is a block, not a section: its toggle takes the
    # paragraph and leaves the section it leads.
    def units_prose(report):
        return [b.text for b in report.blocks("prose")
                if "units" in b.text or "unit system" in b.text]

    if not units_prose(_build()):
        fails.append("the default report states no units")
    off = _build({"pd_units": False})
    if units_prose(off):
        fails.append("pd_units=False still stated the units")
    if "Project Definition" not in [t for _l, t in off.section_titles()]:
        fails.append("pd_units=False removed the Project Definition section")

    # A parent off takes the whole branch.
    titles = [t for _lvl, t in _build({"project_definition": False}).section_titles()]
    if "Project Definition" in titles:
        fails.append("project_definition=False left 'Project Definition' in the "
                     "report")
    # And takes nothing that is not its own. The reinforcement is read by the
    # stability analysis and stands with it: a report built without the general
    # description of the model still documents what that analysis was run on.
    if "Reinforcement" not in titles:
        fails.append("project_definition=False took the Reinforcement section "
                     "out of the limit equilibrium analysis with it")
    titles = [t for _lvl, t in _build({"lem": False}).section_titles()]
    for gone in ("Limit Equilibrium Analysis", "Slice Table", "Results",
                 "Calculations", "Materials", "Loads", "Reinforcement"):
        if gone in titles:
            fails.append(f"lem=False left {gone!r} in the report")

    fails += _member_switch_stands_under_its_engine()
    return fails


def _member_switch_stands_under_its_engine():
    """The member tables are switched on and off per engine, and each switch is
    a child of the engine that prints them.

    The dialog forces a sub-item false whenever its parent section is off, so
    where a switch sits IS what it can mean. Under Project Definition — where it
    was — a report built without the general description of the model lost the
    properties its stability analysis was run on. Under one engine and shared
    with the other it would be the same defect turned sideways: the finite
    element tables would go when the limit equilibrium section did.

    So there is one switch per engine, each under its own engine, and neither
    reaches across.
    """
    fails = []
    from studio.report_dialog import CONTENT_TREE
    from xslope.report import DEFAULT_OPTIONS

    parents = {}
    for key, _label, _tip, children in CONTENT_TREE:
        for child_key, _l, _t in children:
            parents[child_key] = key
    for key, engine in (("lem_members", "lem"), ("fem_members", "fem")):
        if key not in DEFAULT_OPTIONS:
            fails.append(f"{key} is not an option the builder reads")
            continue
        if DEFAULT_OPTIONS[key] is not True:
            fails.append(f"{key} is off by default; the properties an analysis "
                         f"was run on are part of documenting it")
        if key not in parents:
            fails.append(f"the dialog offers no {key} row")
        elif parents[key] != engine:
            fails.append(f"the {key} row is a child of {parents[key]!r}; it "
                         f"belongs under {engine!r}, because the dialog turns a "
                         f"child off with its parent and this switch may only "
                         f"be turned off by the engine that prints it")
    # The retired key is gone from both the tree and the builder: a row nothing
    # reads is a box that does nothing.
    if "pd_reinforcement" in parents:
        fails.append("the dialog still offers pd_reinforcement, which the "
                     "builder no longer reads")
    if "pd_reinforcement" in DEFAULT_OPTIONS:
        fails.append("pd_reinforcement is still a builder option")

    # And the switches really are independent: one engine's tables survive the
    # other engine's section being switched off entirely.
    slope_data, bundle = _fem_1d_bundle(FEM_REINF_XLSX)
    with tempfile.TemporaryDirectory() as tmp:
        from xslope.report import build_report
        report = build_report(slope_data, {"fem": [bundle]},
                              dict(FAST_FIGURES, pd_figure=False,
                                   lem=False, lem_members=False), tmp)
    if "Reinforcement" not in [t for _l, t in report.section_titles()]:
        fails.append("switching the limit equilibrium members off took the "
                     "finite element section's reinforcement with them")
    with tempfile.TemporaryDirectory() as tmp:
        off = build_report(slope_data, {"fem": [bundle]},
                           dict(FAST_FIGURES, pd_figure=False,
                                fem_members=False), tmp)
    if "Reinforcement" in [t for _l, t in off.section_titles()]:
        fails.append("fem_members=False left the reinforcement section in the "
                     "finite element report")
    return fails

    # A figure toggle removes the figure, not the section that holds it.
    on = _build({"pd_figure": True})
    off = _build({"pd_figure": False})
    if not any(f.source == "shared model" for f in on.figures()):
        fails.append("pd_figure=True drew no model figure, so its 'off' case "
                     "proves nothing")
    if any(f.source == "shared model" for f in off.figures()):
        fails.append("pd_figure=False still drew the model figure")
    if "Project Definition" not in [t for _lvl, t in off.section_titles()]:
        fails.append("pd_figure=False removed the Project Definition section")
    return fails


# --------------------------------------------------------------------------
# C. the method picker
# --------------------------------------------------------------------------

def test_method_picker():
    """The picked method drives the detail and leaves the summary alone."""
    fails = []
    from xslope.report import method_label, solved_methods

    slope_data, solutions = _solved()
    if solved_methods(solutions) != ["spencer", "bishop"]:
        fails.append(f"the sample offers {solved_methods(solutions)}")

    reports = {}
    for name in ("spencer", "bishop"):
        reports[name] = _build({"method": name,
                                "lem_search_figure": False,
                                "lem_solution_figure": True})

    a, b = reports["spencer"], reports["bishop"]

    # The detail follows the pick: figure provenance and slice-table caption.
    for name, report in reports.items():
        sources = [f.source for f in report.figures()]
        if f"{name} solution surface" not in sources:
            fails.append(f"method={name}: no solution-surface figure for it "
                         f"(sources {sources})")
        captions = [t.caption for t in report.tables() if t.landscape]
        if not captions or method_label(name) not in captions[0]:
            fails.append(f"method={name}: the slice table is captioned "
                         f"{captions}")
    if ([f.source for f in a.figures()] == [f.source for f in b.figures()]):
        fails.append("the two methods produced the same figure provenance — the "
                     "picker does not reach the figures")
    if ([f.path for f in a.figures()] == [f.path for f in b.figures()]):
        fails.append("the two methods wrote the same figure files")

    # The summary follows the pick, because the summary IS the pick: a report
    # of one method carries none, and a report of two lists those two.
    for report, name in ((a, "spencer"), (b, "bishop")):
        if _fs_rows(report) is not None:
            fails.append(f"method={name}: a single-method report was given a "
                         f"comparison table")

    # An unknown pick falls back rather than failing.
    report = _build({"method": "no_such_method"})
    if "Limit Equilibrium Analysis" not in [t for _l, t in report.section_titles()]:
        fails.append("an unrecognised method dropped the LEM section")
    return fails


def _fs_rows(report):
    """The factor-of-safety summary as ``(headers, rows)``, or None."""
    for t in report.tables():
        if t.caption == "Computed factors of safety":
            return (t.headers, t.rows)
    return None


def _fs_table_block(report):
    for t in report.tables():
        if t.caption == "Computed factors of safety":
            return t
    return None


def _fs_prose(report):
    """The paragraphs of the Factors of Safety section, as one string."""
    for section in report.sections:
        for _l, sub in section.walk():
            if sub.title == "Factors of Safety":
                return " ".join(b.text for b in sub.blocks if b.kind == "prose")
    return ""


#: The two things a summary paragraph can say about where its numbers came from,
#: and they are alternatives. A report that says both says that every method
#: found its own surface AND that no surface was searched for.
_SEARCHED_CLAIM = "searched for"
_UNSEARCHED_CLAIM = "No search was performed"


def _searched_bundles(bundles, which):
    """The same bundles with a search attached to the named methods and taken off
    the rest — the three provenances, off one solved model."""
    out = []
    for bundle in bundles:
        copy = dict(bundle)
        if copy.get("method") in which:
            copy["search"] = copy.get("search") or {
                "kind": "circular", "fs_cache": [], "search_path": []}
        else:
            copy["search"] = None
        out.append(copy)
    return out


def test_the_fs_summary_says_where_its_numbers_came_from():
    """The summary paragraph states one provenance for its rows, and it is the
    true one.

    Three cases, and exactly one sentence is true of any table. Every method
    searched: each row is that method's own critical factor of safety, on its own
    critical surface. Some searched: those rows are minima and the rest are the
    specified surface. None searched: every row is the specified surface, and no
    row is a minimum over anything.

    The paragraph used to open by saying every method finds its own surface,
    whatever the bundles held, and then append the no-search sentence under it —
    so a report solved entirely on one specified circle told the reader both at
    once. The two claims are alternatives, and no paragraph may carry both.
    """
    fails = []
    from xslope.report import build_report, method_label

    slope_data, solutions = _solved()
    bundles = solutions.get("lem") or []
    methods = [b.get("method") for b in bundles]
    if len(bundles) < 2:
        return ["the sample solves fewer than two methods; the summary is untested"]

    def built(which):
        with tempfile.TemporaryDirectory() as tmp:
            return build_report(
                slope_data, {"lem": _searched_bundles(bundles, which)},
                {"pd_figure": False, "lem_search_figure": False,
                 "lem_solution_figure": False, "lem_slice_key": False,
                 "lem_slice_table": False, "lem_calculations": False,
                 "method": methods}, tmp)

    cases = [("every method searched", set(methods), True, False),
             ("one method searched", {methods[0]}, True, True),
             ("no method searched", set(), False, True)]
    for where, which, wants_searched, wants_specified in cases:
        said = _fs_prose(built(which))
        if not said:
            fails.append(f"{where}: the report carries no summary paragraph")
            continue
        # Never both. This is the defect: the two claims contradict each other.
        if _SEARCHED_CLAIM in said and _UNSEARCHED_CLAIM in said:
            fails.append(f"{where}: the paragraph says a search was made and "
                         f"that none was: {said!r}")
        if ("finds its own surface" in said
                and _UNSEARCHED_CLAIM in said):
            fails.append(f"{where}: the paragraph says every method finds its "
                         f"own surface and that no search was performed: "
                         f"{said!r}")
        if wants_searched and _SEARCHED_CLAIM not in said:
            fails.append(f"{where}: the paragraph does not say a search was "
                         f"made: {said!r}")
        if not wants_searched and _SEARCHED_CLAIM in said:
            fails.append(f"{where}: no bundle carries a search and the "
                         f"paragraph says one was made: {said!r}")
        if wants_specified and "specified in the input" not in said:
            fails.append(f"{where}: rows stand on the specified surface and the "
                         f"paragraph does not say so: {said!r}")
        if not wants_specified and "specified in the input" in said:
            fails.append(f"{where}: every method searched and the paragraph "
                         f"still sends the reader to the specified surface: "
                         f"{said!r}")

    # Where every method searched, the paragraph says what the rows ARE: each
    # method's own critical value. That is the summary the owner asked for, and
    # the wording is what makes the column readable as one.
    said = _fs_prose(built(set(methods)))
    for phrase in ("critical factor of safety", "its own critical surface",
                   "the surfaces differ from method to method"):
        if phrase not in said:
            fails.append(f"every method searched: the paragraph does not say "
                         f"{phrase!r}: {said!r}")

    # And where only some did, the ones that did are named.
    said = _fs_prose(built({methods[0]}))
    if method_label(methods[0]) not in said:
        fails.append(f"one method searched: the paragraph does not name "
                     f"{method_label(methods[0])}: {said!r}")
    return fails


def test_multi_method_detail():
    """Several methods, several full detail blocks — in the order asked for, each
    headed and captioned with its own method's name.

    The failure this defends against is the one that makes a multi-method report
    worse than a single-method one: two blocks of numbers a reader cannot tell
    apart. Every heading, figure caption and table caption in a block has to name
    its method, and the blocks have to come in a stated order.

    A block is headed by the method's bare name, so the blocks are found by
    matching the level-2 headings against the labels of every method the solver
    offers — which is what makes a block for a method nobody asked for a failure
    as well.
    """
    fails = []
    from xslope.report import detail_bundle, method_label, supported_methods

    slope_data, solutions = _solved()
    labels = {method_label(m) for m in supported_methods()}

    wanted = ["bishop", "spencer"]
    report = _build({"method": wanted, "lem_solution_figure": True})
    titles = [t for lvl, t in report.section_titles() if lvl == 2]

    heads = [t for t in titles if t in labels]
    expected = [method_label(m) for m in wanted]
    if heads != expected:
        fails.append(f"the detail blocks are {heads}, expected {expected}")

    # One of everything per method, each naming its own.
    for m in wanted:
        label = method_label(m)
        sources = [f.source for f in report.figures()]
        if f"{m} solution surface" not in sources:
            fails.append(f"{m}: no solution-surface figure ({sources})")
        captions = [f.caption for f in report.figures()
                    if f.source == f"{m} solution surface"]
        if not captions or label not in captions[0]:
            fails.append(f"{m}: the solution-surface caption is {captions}")
        slices = [t for t in report.tables() if t.landscape and label in t.caption]
        if len(slices) != 1:
            fails.append(f"{m}: {len(slices)} slice tables captioned for it")

    # Each block carries the whole subtree, not just a heading — and where the
    # method searched, its own search stands at the head of that subtree.
    blocks = {}
    for sec in report.sections:
        for _lvl, node in sec.walk():
            if node.title in labels:
                blocks[node.title] = node
    for m in wanted:
        node = blocks.get(method_label(m))
        if node is None:
            continue
        detail, _note = detail_bundle(slope_data, solutions, m)
        want_kids = (["Search for the Critical Surface"]
                     if (detail or {}).get("search") else [])
        want_kids += ["Results", "Slice Table", "Calculations"]
        kids = [c.title for c in node.children]
        if kids != want_kids:
            fails.append(f"{node.title!r} carries {kids}, expected {want_kids}")

    # Figures and tables are numbered straight through, whatever the block.
    numbers = [f.number for f in report.figures()]
    if numbers != list(range(1, len(numbers) + 1)):
        fails.append(f"the figures are not numbered 1..n across methods: {numbers}")
    numbers = [t.number for t in report.tables()]
    if numbers != list(range(1, len(numbers) + 1)):
        fails.append(f"the tables are not numbered 1..n across methods: {numbers}")

    # The summary comes ONCE, and before the first detail block: it lists every
    # method, so it belongs to the analysis. The search does NOT — it belongs to
    # the block whose surface it found, so there is one per method that searched
    # and none above the blocks.
    if sum(1 for t in report.tables()
           if t.caption == "Computed factors of safety") != 1:
        fails.append("the factor-of-safety summary is not printed exactly once")
    searched = [m for m in wanted
                if (detail_bundle(slope_data, solutions, m)[0] or {}).get("search")]
    levels = [lvl for lvl, t in report.section_titles()
              if t == "Search for the Critical Surface"]
    if len(levels) != len(searched):
        fails.append(f"the report documents {len(levels)} searches for the "
                     f"{len(searched)} method(s) that searched: {searched}")
    if any(lvl != 3 for lvl in levels):
        fails.append(f"a search is written at heading level {levels}; it is a "
                     f"child of the method block whose surface it found, not a "
                     f"section standing over all of them")
    if "Factors of Safety" not in titles:
        fails.append(f"there is no factor-of-safety section: {titles}")
    elif heads and titles.index("Factors of Safety") > titles.index(heads[0]):
        fails.append("the summary comes after the detail blocks")

    # Order is the caller's, not the solver's.
    other = _build({"method": ["spencer", "bishop"]})
    heads2 = [t for lvl, t in other.section_titles() if lvl == 2 and t in labels]
    if heads2 != list(reversed(expected)):
        fails.append(f"reversing the request gave {heads2}")

    # A bare string is still one method — every caller written before the list
    # option existed keeps working.
    one = _build({"method": "bishop"})
    heads3 = [t for lvl, t in one.section_titles() if lvl == 2 and t in labels]
    if heads3 != [method_label("bishop")]:
        fails.append(f"method='bishop' produced {heads3}")

    # A method the analysis did not run is RUN for the report, and documented
    # from that run like any other (the owner's ruling). Asking for one
    # produces its detail block, its own factor of safety, and no note anywhere
    # about where the number came from — because it came from the same place
    # every other number in the report did.
    from xslope.report import solved_methods
    run = solved_methods(solutions)
    unrun = next(m for m in ("janbu", "lowe", "corps") if m not in run)
    extra = _build({"method": [unrun]})
    prose = " ".join(b.text for b in extra.blocks("prose"))
    heads4 = [t for _l, t in extra.section_titles() if t in labels]
    if heads4 != [method_label(unrun)]:
        fails.append(f"asking for {unrun}, which the analysis did not run, "
                     f"documented {heads4}")
    if f"{method_label(unrun)} gives a factor of safety" not in prose:
        fails.append(f"{unrun} was run for the report and its block reports no "
                     f"factor of safety: {prose[-400:]!r}")
    for banned in ("It was not run in the analysis",
                   "the report solved it on the same",
                   "was not run in this analysis"):
        if banned in prose:
            fails.append(f"a method the report ran is described as unrun, or "
                         f"as solved on another method's surface: {banned!r}")
    # Asked for a run method AND an unrun one, BOTH are documented, and both
    # rows of the summary carry a number of their own.
    mixed = _build({"method": [run[0], unrun]})
    heads5 = [t for _l, t in mixed.section_titles() if t in labels]
    if heads5 != [method_label(run[0]), method_label(unrun)]:
        fails.append(f"a mixed request documented {heads5}")
    summary = [t for t in mixed.tables()
               if t.caption == "Computed factors of safety"]
    rows = {r[0]: r[1] for t in summary for r in t.rows}
    if method_label(unrun) not in rows:
        fails.append(f"the summary table carries no row for {unrun}, which the "
                     f"report ran: {sorted(rows)}")
    elif rows[method_label(unrun)] == "did not converge":
        fails.append(f"{unrun} was run for the report and its row says it did "
                     f"not converge: {rows}")
    return fails


def test_each_engine_presents_the_loads_it_applies():
    """The loads are an analysis input, and the second engine cites the first
    rather than printing the table twice.

    A distributed load is not a property of the section; it is something an
    analysis puts on it. So it stands with the engine that applies it, not in a
    general description of the model. Where both engines are documented the
    blocks are identical — same points, same pressures — so the finite element
    section points at the limit equilibrium table: two copies of one table are
    two things to keep in step.
    """
    fails = []
    from xslope.report import build_report

    slope_data, solutions = _load_bearing_model()
    if slope_data is None:
        return ["no loaded model with both engines; the citation is untested"]

    with tempfile.TemporaryDirectory() as tmp:
        report = build_report(slope_data, solutions,
                              dict(FAST_FIGURES, pd_figure=False), tmp)

    # Not in Project Definition.
    pd_sec = next((s for s in report.sections
                   if s.title == "Project Definition"), None)
    if pd_sec is not None:
        under = [t for _l, t in pd_sec.walk()]
        if "Loads" in under:
            fails.append(f"Project Definition still carries a Loads section: "
                         f"{under}")

    # One table, in the limit equilibrium section; the other engine cites it.
    loads = [t for t in report.tables() if t.caption == "Distributed loads"]
    if len(loads) != 1:
        fails.append(f"{len(loads)} distributed-load tables were printed; the "
                     f"second engine must cite the first")
    for head in ("Limit Equilibrium Analysis",
                 "Deformation and Strength Reduction"):
        sec = next((s for s in report.sections if s.title == head), None)
        if sec is None:
            fails.append(f"the report has no {head} section")
            continue
        sub = next((s for _l, s in sec.walk() if s.title == "Loads"), None)
        if sub is None:
            fails.append(f"{head} presents no loads")
            continue
        if head.startswith("Deformation"):
            if any(b.kind == "table" for b in sub.blocks):
                fails.append("the finite element section printed its own copy "
                             "of the loads table")
            said = " ".join(b.text for b in sub.blocks if b.kind == "prose")
            if loads and f"Table {loads[0].number}" not in said:
                fails.append(f"the finite element section does not cite the "
                             f"loads table: {said!r}")

    fails += _water_load_mechanism_checks()
    return fails


def _loads_prose(xlsx):
    """The Loads subsection's prose for a model, as one string."""
    from xslope.report import _loads_section, water_features, _Counter
    slope_data = load_slope_data_cached(xlsx)
    section = _loads_section(slope_data, water_features(slope_data), _Counter())
    return " ".join(b.text for b in section.blocks if b.kind == "prose")


#: What the Loads prose has to say about a load the user never entered: how big
#: it is, and which way it points. Anything less leaves an engineer reading a
#: factor of safety that stands partly on a force the report never described.
_WATER_MECHANISM = ("depth of water", "unit weight of water",
                    "normal to the ground surface")


def _water_load_mechanism_checks():
    """A load the engine derived is a load the report states the derivation of.

    Where the water surface stands above the ground surface, the engine applies
    the standing water to the ground: pressure equal to the depth of water above
    each point times the unit weight of water, acting normal to the surface and
    tapering to zero at the shoreline. That force enters the factor of safety and
    the user never typed it, so the analysis is not documented until the report
    says what it is.

    Said only where there is one. A section with no water standing on it carries
    no such load, and describing one would describe an analysis that did not
    happen.
    """
    fails = []
    from xslope.water import with_water_loads, has_derived_loads
    import xslope.report as report_mod

    ponding, dry = DAM_XLSX, SIMPLE1_XLSX
    if not has_derived_loads(with_water_loads(load_slope_data_cached(ponding))):
        fails.append(f"{os.path.basename(ponding)} derives no water load, so "
                     f"the sentence describing one is untested")
    if has_derived_loads(with_water_loads(load_slope_data_cached(dry))):
        fails.append(f"{os.path.basename(dry)} derives a water load, so it "
                     f"cannot show the sentence staying away")

    said = _loads_prose(ponding)
    for named in _WATER_MECHANISM:
        if named not in said:
            fails.append(f"the Loads prose of a ponded section does not say "
                         f"{named!r}: {said!r}")

    # The unit weight is named, not printed again. The units paragraph declares
    # its value; a second copy is a number that can disagree with itself.
    gamma = load_slope_data_cached(ponding).get("gamma_water")
    if gamma and f"{float(gamma):g}" in said:
        fails.append(f"the Loads prose restates the unit weight of water "
                     f"({float(gamma):g}), which the units paragraph declares")

    quiet = _loads_prose(dry)
    for named in _WATER_MECHANISM:
        if named in quiet:
            fails.append(f"a section with no water standing on it is told how "
                         f"standing water is loaded ({named!r}): {quiet!r}")

    # Mutation: the check has to be able to fail. With the mechanism withheld
    # the ponded section's Loads prose goes back to naming a derived load
    # without saying what it is, which is the report the defect was found in.
    real = report_mod._water_load_mechanism
    report_mod._water_load_mechanism = lambda slope_data, **kw: None
    try:
        muted = _loads_prose(ponding)
    finally:
        report_mod._water_load_mechanism = real
    if any(named in muted for named in _WATER_MECHANISM):
        fails.append(f"the mechanism sentence is written somewhere other than "
                     f"the one place that decides whether there is one: "
                     f"{muted!r}")
    return fails


def _load_bearing_model():
    """A loaded model solved by BOTH engines, or ``(None, None)``.

    The finite element sample carries distributed loads and a shipped strength
    reduction solution; one limit equilibrium trial on its own starting circle
    puts the other engine beside it, which is the only shape in which the
    citation exists to be checked.
    """
    from xslope.slice import generate_slices
    from xslope.solve import solve_selected

    slope_data, fem = _fem_bundle()
    if not slope_data.get("dloads") or not slope_data.get("circles"):
        return None, None
    with contextlib.redirect_stdout(io.StringIO()):
        ok, out = generate_slices(slope_data, circle=slope_data["circles"][0])
        if not ok:
            return None, None
        df, surface = out
        results = solve_selected("spencer", df)
    lem = [{"slice_df": df, "failure_surface": surface, "results": results,
            "search": None, "method": "spencer"}]
    return slope_data, {"lem": lem, "fem": fem}


def test_critical_is_a_word_a_search_earns():
    """A surface is called critical only where a search found it.

    "Critical" means least of a family. A factor of safety computed on a surface
    the user entered is the factor of safety of THAT surface, and printing it
    against the word critical invites an engineer to read a specified-surface
    answer as a searched minimum — the one misreading a stability report must not
    permit. Where no search ran the report says so outright rather than softening
    the adjective.

    Both shapes are built from the same sample: its Spencer bundle carries a
    search and its Bishop bundle does not.
    """
    fails = []
    _slope_data, solutions = _solved()

    searched = {b["method"]: bool(b.get("search")) for b in solutions["lem"]}
    if set(searched.values()) != {True, False}:
        return [f"the sample's bundles are all the same shape ({searched}); "
                f"one of the two wordings is untested"]

    report = _build({"method": ["spencer", "bishop"]})
    from xslope.report import method_label
    for name, ran in searched.items():
        block = next((s for _l, s in _sections(report)
                      if s.title == method_label(name)), None)
        if block is None:
            fails.append(f"{name}: no method block")
            continue
        said = " ".join(b.text for _l, node in block.walk()
                        for b in node.blocks if b.kind == "prose")
        captions = " ".join(f.caption for _l, node in block.walk()
                            for f in node.blocks if f.kind == "figure")
        if ran:
            if "critical surface" not in said:
                fails.append(f"{name} searched and its block never says critical "
                             f"surface: {said!r}")
            if "no search" in said.lower():
                fails.append(f"{name} searched and its block says no search ran")
        else:
            if "critical surface" in captions:
                fails.append(f"{name} did not search and its figure is captioned "
                             f"{captions!r}")
            if "specified surface" not in said:
                fails.append(f"{name} did not search and its block never calls "
                             f"the surface specified: {said!r}")
            if "no search" not in said.lower():
                fails.append(f"{name} did not search and the block never says so "
                             f"outright: {said!r}")
            # …and it says what that means for the number.
            if "not a minimum" not in said:
                fails.append(f"{name}: nothing says the reported factor is not a "
                             f"minimum over a family of surfaces: {said!r}")

    # The engine's own inputs say which it was, too.
    inputs = next((s for _l, s in _sections(report)
                   if s.title == "Analysis Inputs"), None)
    rows = dict(item for b in (inputs.blocks if inputs else [])
                if b.kind == "keyvalues" for item in b.items)
    if "Surface" not in rows:
        fails.append(f"the analysis inputs do not say how the surface was "
                     f"arrived at: {list(rows)}")

    # A report with no search anywhere never uses the word.
    nosearch = {"lem": [dict(b, search=None) for b in solutions["lem"]]}
    plain = _build({"method": ["spencer", "bishop"]})
    with tempfile.TemporaryDirectory() as tmp:
        from xslope.report import build_report
        plain = build_report(_slope_data, nosearch,
                             {"pd_figure": False, "lem_search_figure": False,
                              "lem_solution_figure": False,
                              "method": ["spencer", "bishop"]}, tmp)
    for block in plain.blocks("prose"):
        if "critical surface" in block.text and "No search" not in block.text \
                and "no search" not in block.text:
            fails.append(f"a report with no search calls a surface critical: "
                         f"{block.text!r}")
    if "Search for the Critical Surface" in [t for _l, t in plain.section_titles()]:
        fails.append("a report with no search carries a search section")
    return fails


def test_lem_inputs_figure():
    """The limit equilibrium section opens on the model the method of slices
    reads, as the other two engines' sections do.

    Every engine section states its own inputs, and the first of them is the
    section drawn the way that engine sees it. The shared figure of the Project
    Definition is not that figure: it deliberately suppresses the surface family,
    which is exactly what a limit equilibrium analysis is run over.
    """
    fails = []
    from xslope.report import DEFAULT_OPTIONS, build_report, planned_figures, \
        resolve_options

    if DEFAULT_OPTIONS.get("lem_inputs_figure") is not True:
        fails.append("the limit equilibrium model figure is not on by default")

    slope_data, solutions = _solved()

    def built(**over):
        opts = {"input_path": REINF_XLSX, "method": ["spencer", "bishop"],
                "pd_figure": False, "lem_search_figure": False,
                "lem_solution_figure": False}
        opts.update(FAST_FIGURES)
        opts.update(over)
        tmp = tempfile.mkdtemp(prefix="xslope_leminputs_")
        with contextlib.redirect_stdout(io.StringIO()):
            return opts, build_report(slope_data, solutions, opts, tmp)

    opts, report = built()
    figures = [f for f in report.figures() if f.source == "lem model"]
    if not figures:
        fails.append(f"the limit equilibrium section drew no model figure: "
                     f"{[f.source for f in report.figures()]}")
        return fails

    # One per section. Two methods are documented, and both are run on this one
    # model; a copy under each method's heading is the same picture twice.
    if len(figures) != 1:
        fails.append(f"the report drew {len(figures)} limit equilibrium model "
                     f"figures; it is an input of the section, not of a method")
    figure = figures[0]
    if "limit equilibrium" not in figure.caption.lower():
        fails.append(f"the figure is captioned {figure.caption!r}")

    # It stands in the section's own Analysis Inputs, ahead of the rows.
    lem = next((s for s in report.sections
                if s.title == "Limit Equilibrium Analysis"), None)
    inputs = next((c for c in (lem.children if lem else [])
                   if c.title == "Analysis Inputs"), None)
    if inputs is None:
        fails.append("the limit equilibrium section has no Analysis Inputs")
    else:
        kinds = [b.source if b.kind == "figure" else b.kind
                 for b in inputs.blocks if b.kind in ("figure", "keyvalues")]
        if kinds != ["prose", "lem model", "keyvalues"][1:]:
            if kinds[:2] != ["lem model", "keyvalues"]:
                fails.append(f"the limit equilibrium inputs are ordered {kinds}, "
                             f"not the model figure and then the rows")
        cites = [b for b in inputs.blocks
                 if b.kind == "prose" and f"Figure {figure.number}" in b.text]
        if not cites:
            fails.append(f"nothing cites Figure {figure.number}: "
                         f"{[b.text for b in inputs.blocks if b.kind == 'prose']}")
        else:
            said = cites[0].text
            if "limit equilibrium model" not in said:
                fails.append(f"the sentence does not say whose view of the model "
                             f"the figure is: {said!r}")
            # What it actually draws, and only that: the sample carries a
            # searched circular family and reinforcement.
            for named in ("the starting circle for the search",
                          "the reinforcement lines"):
                if named not in said:
                    fails.append(f"the sentence does not name {named!r}: {said!r}")

    # It is the LEM view, not the shared one: the surface family the shared
    # figure suppresses is drawn here.
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.figure as mplfig
    from xslope.plot import plot_inputs
    gids = {}
    for mode in ("lem", "shared"):
        f = mplfig.Figure()
        plot_inputs(slope_data, fig=f, mode=mode, show_title=False,
                    frame="content", show_mesh=False)
        gids[mode] = {a.get_gid() for a in f.axes[0].get_children()
                      if a.get_gid()}
    if not (gids["lem"] - gids["shared"]):
        fails.append(f"the limit equilibrium view draws nothing the shared view "
                     f"does not, so it is the same figure twice: {gids['lem']}")

    # …and it is the view the SECTION asks for. Reading the two modes apart says
    # nothing about which one the report drew, and the report ships a PNG the
    # check cannot interrogate, so the request itself is watched.
    import xslope.plot as xp
    real = xp.plot_inputs
    asked = []

    def spy(sd, **kw):
        asked.append(kw.get("mode"))
        return real(sd, **kw)

    xp.plot_inputs = spy
    try:
        built()
    finally:
        xp.plot_inputs = real
    if "lem" not in asked:
        fails.append(f"the limit equilibrium model figure was drawn in mode(s) "
                     f"{asked}, none of them the engine's own view")

    # The sentence claims what this view draws, and every part of it: a model
    # whose pool a head boundary states has that pool drawn on the limit
    # equilibrium figure, beside the load derived from it, and both are named.
    # A piezometric line it does not carry is not.
    dam, dam_solutions = _restored(SEEP_XLSX)
    from xslope.report import water_features
    dam_feats = water_features(dam)
    if dam_feats["piezo"]:
        fails.append("the seepage sample now carries a piezometric line; the "
                     "no-piezometric-line wording is untested")
    else:
        f = mplfig.Figure()
        plot_inputs(dam, fig=f, mode="lem", show_title=False, frame="content",
                    show_mesh=False)
        drawn = {a.get_gid() for a in f.axes[0].get_children() if a.get_gid()}
        if any("PIEZO" in (g or "") for g in drawn):
            fails.append(f"the seepage sample does draw a piezometric line "
                         f"after all: {sorted(drawn)}")
        if not any(g == "WATER_LINE" for g in drawn):
            fails.append(f"the limit equilibrium view draws no pool for a model "
                         f"whose head boundaries state one: {sorted(drawn)}")
        report2 = _cite_report(SEEP_XLSX, ("spencer",), engines=("seep",))
        lem2 = next(s for s in report2.sections
                    if s.title == "Limit Equilibrium Analysis")
        inputs2 = next(c for c in lem2.children
                       if c.title == "Analysis Inputs")
        said2 = " ".join(b.text for b in inputs2.blocks if b.kind == "prose")
        if "piezometric" in said2.lower():
            fails.append(f"the limit equilibrium figure's sentence promises a "
                         f"piezometric line the view does not draw: {said2!r}")
        for named in ("the water surface", "the distributed loads"):
            if named not in said2:
                fails.append(f"the figure draws {named} and the sentence does "
                             f"not name it: {said2!r}")

    # The toggle takes the figure and leaves the section.
    off_opts, off = built(lem_inputs_figure=False)
    if any(f.source == "lem model" for f in off.figures()):
        fails.append("lem_inputs_figure=False still drew the model figure")
    titles = [t for _l, t in off.section_titles()]
    if "Limit Equilibrium Analysis" not in titles or \
            "Analysis Inputs" not in titles:
        fails.append("lem_inputs_figure=False removed the section, not only the "
                     "figure")
    if len(off.figures()) != len(report.figures()) - 1:
        fails.append(f"lem_inputs_figure=False changed the figure count by "
                     f"{len(report.figures()) - len(off.figures())}, not 1")

    # And the count a caller is promised matches what was built, both ways.
    for these, made in ((opts, report), (off_opts, off)):
        planned = planned_figures(slope_data, solutions, resolve_options(these))
        if planned != len(made.figures()):
            fails.append(f"lem_inputs_figure="
                         f"{these.get('lem_inputs_figure', True)}: planned "
                         f"{planned} figures and built {len(made.figures())}")
    return fails


def test_search_figure_is_read_for_the_engineer():
    """The search subsection names what the search plot draws.

    Several thousand overlaid trials are unreadable to anyone who has not been
    told which mark is which, and the marks are not self-explanatory: a gray arc,
    a black dot and a green arrow are three different statements about the
    search. The prose names each, and names it as the artists actually draw it.
    """
    fails = []
    report = _build(dict({"method": ["spencer"], "lem_search_figure": True},
                         **FAST_FIGURES))
    search = next((s for _l, s in _sections(report)
                   if s.title == "Search for the Critical Surface"), None)
    if search is None:
        return ["the sample's report carries no search section"]
    figure = next((b for b in search.blocks if b.kind == "figure"), None)
    if figure is None:
        return ["the search section drew no figure, so its prose proves nothing"]
    said = " ".join(b.text for b in search.blocks if b.kind == "prose")
    if f"Figure {figure.number}" not in said:
        fails.append(f"the search figure is not cited: {said!r}")

    # Every element the circular search plot draws is named. The colours are the
    # ones plot.py hardcodes — nothing in the style names them — so the prose
    # can and must say which mark is which.
    for element in ("black dot", "red", "green arrow", "center"):
        if element not in said:
            fails.append(f"the search prose does not name the {element!r} the "
                         f"figure draws: {said!r}")
    if "refine" not in said:
        fails.append(f"the search prose never says the grid is refined: {said!r}")

    # The claim is checked against the artists, not against the wording: a plot
    # that stopped drawing centers would leave the sentence describing a figure
    # that is not there.
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.figure as mplfig
    from xslope.plot import plot_circular_search_results
    slope_data, solutions = _solved()
    bundle = next(b for b in solutions["lem"] if b.get("search"))
    run = bundle["search"]
    f = mplfig.Figure()
    plot_circular_search_results(slope_data, run.get("fs_cache") or [],
                                 search_path=run.get("search_path"),
                                 circle_cache=run.get("circle_cache"), fig=f,
                                 show_title=False, show_legend=False)
    ax = f.axes[0]
    drawn = {a.get_gid() for a in ax.get_children() if a.get_gid()}
    if "CIRCLE_CENTERS" not in drawn:
        fails.append(f"the search plot draws no circle centers, and the prose "
                     f"says it does: {sorted(drawn)}")
    if "CRITICAL_SURFACE" not in drawn:
        fails.append(f"the search plot highlights no critical surface: "
                     f"{sorted(drawn)}")
    centers = [a for a in ax.get_children()
               if a.get_gid() == "CIRCLE_CENTERS" and hasattr(a, "get_color")]
    colors = {str(a.get_color()) for a in centers}
    if not {"k", "black"} & colors:
        fails.append(f"the trial centers are drawn {colors}, and the prose calls "
                     f"them black")
    if not {"r", "red"} & colors:
        fails.append(f"no center is drawn red, and the prose says the critical "
                     f"one is: {colors}")
    if run.get("search_path") and len(run["search_path"]) > 1:
        arrows = [a for a in ax.get_children() if a.get_gid() == "SEARCH_PATH"]
        if not arrows:
            fails.append("the search plot draws no path, and the prose describes "
                         "arrows running from stage to stage")
        elif not any("green" in str(a.get_facecolor()) or
                     a.get_facecolor()[:3] == (0.0, 0.5019607843137255, 0.0)
                     for a in arrows):
            fails.append(f"the search path is not green, and the prose calls it "
                         f"green: {[a.get_facecolor() for a in arrows]}")

    # A model with no circle cache says only what it drew: the centers, not the
    # arcs it did not keep.
    from xslope.report import _search_section, _Counter, resolve_options
    thin = dict(bundle, search=dict(run, circle_cache=None))
    tmp = tempfile.mkdtemp(prefix="xslope_searchprose_")
    with contextlib.redirect_stdout(io.StringIO()):
        sub = _search_section(slope_data, thin,
                              resolve_options({"input_path": REINF_XLSX}),
                              _Counter(), tmp, "spencer")
    thin_said = " ".join(b.text for b in sub.blocks if b.kind == "prose")
    if "drawn in gray" in thin_said:
        fails.append(f"a search that kept no trial circles still says they are "
                     f"drawn: {thin_said!r}")
    if "black dot" not in thin_said:
        fails.append(f"a search that kept no trial circles no longer names its "
                     f"centers: {thin_said!r}")
    return fails


def _two_searches():
    """The sample model with BOTH methods searched, each over its own grid.

    Two searches settle on two surfaces, and the grids are pulled apart here so
    that a section showing the wrong one shows a visibly different picture rather
    than a coincidence.
    """
    import copy

    slope_data, solutions = _solved()
    bundles = copy.deepcopy(solutions["lem"])
    first = bundles[0]["search"]
    bundles[1]["search"] = {
        "kind": "circular",
        "fs_cache": [dict(c, Xo=c["Xo"] + 40.0, Yo=c["Yo"] + 25.0)
                     for c in first["fs_cache"]],
        "search_path": [{"x": 50.0, "y": 75.0, "FS": None},
                        {"x": 47.0, "y": 71.0,
                         "FS": bundles[1]["results"]["FS"]}],
        "circle_cache": None,
    }
    return slope_data, {"lem": bundles}


def test_each_searched_method_draws_its_own_search():
    """A method's search section draws that method's search.

    Two methods that each searched settled on two different surfaces over two
    different grids. Every one of them rendered into one ``search.png``, so the
    last to draw owned the file and both sections embedded it: Bishop's section
    printed Spencer's trial grid under a caption naming Bishop. The plots are
    named for their method, the way the critical-surface and slice-key plots
    are, and the figure count a caller is promised counts one per searching
    method rather than one per section.
    """
    fails = []
    from xslope.report import (build_report, file_digest, method_searched,
                               planned_figures, resolve_options)

    slope_data, solutions = _two_searches()
    opts = {"input_path": REINF_XLSX, "title": "Two Searches",
            "method": ["spencer", "bishop"], "pd_figure": False,
            "lem_inputs_figure": False, "lem_solution_figure": False,
            "lem_slice_key": False, "figure_dpi": 60}
    tmp = tempfile.mkdtemp(prefix="xslope_twosearch_")
    with contextlib.redirect_stdout(io.StringIO()):
        report = build_report(slope_data, solutions, opts, tmp)

    searches = [(label, s) for label, s in _sections(report)
                if s.title == "Search for the Critical Surface"]
    if len(searches) != 2:
        return [f"two methods searched and {len(searches)} search sections were "
                f"built"]
    figs = [next((b for b in s.blocks if b.kind == "figure"), None)
            for _l, s in searches]
    if any(f is None for f in figs):
        return ["a search section drew no figure, so nothing here is proven"]
    if figs[0].path == figs[1].path:
        fails.append(f"both searches were drawn into {figs[0].path}, so one "
                     f"method's section shows the other's grid")
    else:
        digests = [file_digest(f.path) for f in figs]
        if digests[0] == digests[1]:
            fails.append("the two searches drew the same picture, so the two "
                         "grids this check pulled apart are not being drawn")

    # And the count a caller is promised is the count that was built.
    planned = planned_figures(slope_data, solutions, resolve_options(opts))
    drawn = len(report.figures())
    if planned != drawn:
        fails.append(f"{planned} figures were planned and {drawn} were built")

    # A method that did not search is not credited with one another method ran.
    _sd, one = _solved()
    if method_searched(one, "bishop"):
        fails.append("a method that never searched is reported as having "
                     "searched")
    if not method_searched(one, "spencer"):
        fails.append("the method that searched is not reported as having "
                     "searched")
    return fails


def test_a_method_the_run_did_not_solve_is_run_for_the_report():
    """A method the report is asked for is RUN for it, and finds its own surface.

    The owner's rule: every method selected for a report is solved, the way it
    would have been had it been ticked in the Run dialog — a full search for ITS
    critical surface where the analysis searched, the specified surface where it
    did not. Two failures this stands against, and they are opposite ones.

    Dropping it silently: a report asked for five methods and documenting the two
    that had been run is a table missing the rows a reader asked to see, with
    nothing anywhere saying they are missing.

    Borrowing a surface: evaluating the method on the surface another method's
    search found produces a factor of safety no analysis made. Two searches over
    the same model settle on two different circles, so the surfaces here are
    checked to be different ones — a borrowed surface would be the same object as
    its lender's, and the same numbers down the slice table.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.report import (build_report, method_label, methods_to_run,
                               resolve_options, run_requested_methods)
    from xslope.search import run_lem_analysis

    slope_data = load_slope_data(REINF_XLSX)
    # The analysis: ONE method, searched, on a slice count that keeps this check
    # to a few seconds. What is under test is the surface a second method gets,
    # which is a question about the machinery and not about the mesh.
    how = {"analysis": "auto_search", "surface": "circular", "num_slices": 12}
    with contextlib.redirect_stdout(io.StringIO()):
        first = run_lem_analysis(slope_data, "bishop", **how)
    solutions = {"lem": [dict(first, options=how)]}

    opts = resolve_options({"input_path": REINF_XLSX, "title": "Two Methods",
                            "method": ["bishop", "spencer"], "pd_figure": False,
                            "lem_inputs_figure": False,
                            "lem_search_figure": False,
                            "lem_solution_figure": False,
                            "lem_slice_key": False})
    if methods_to_run(slope_data, solutions, opts) != ["spencer"]:
        return ["the fixture already carries Spencer; nothing is run for the "
                "report and the run it makes is untested"]
    with contextlib.redirect_stdout(io.StringIO()):
        after = run_requested_methods(slope_data, solutions, opts)

    if len(solutions["lem"]) != 1:
        fails.append("the caller's own solutions were added to; the extra run "
                     "belongs to the report, not to the session")
    second = next((b for b in after["lem"] if b.get("method") == "spencer"), None)
    if second is None:
        return fails + ["Spencer was asked for and not run"]

    # It ran, and it ran the analysis this model's own run was: its own search.
    if not second.get("search"):
        fails.append("the method the report ran carries no search of its own, "
                     "though the analysis it was run beside searched")
    if (second.get("results") or {}).get("FS") is None:
        fails.append("the method the report ran reports no factor of safety")

    # And what it searched is ITS surface. A borrowed one is the lender's own
    # object; a found one is a different circle, with different slices on it.
    for key in ("failure_surface", "slice_df"):
        if second.get(key) is first.get(key):
            fails.append(f"the method the report ran was handed the other "
                         f"method's {key}")
    centers = []
    for bundle in (first, second):
        df = bundle.get("slice_df")
        centers.append(None if df is None or "xo" not in df.columns
                       else (round(float(df["xo"].iloc[0]), 6),
                             round(float(df["yo"].iloc[0]), 6)))
    if None in centers:
        fails.append("the slice tables carry no circle center; the surfaces "
                     "cannot be told apart")
    elif centers[0] == centers[1]:
        fails.append(f"both methods report the same critical circle {centers[0]}; "
                     f"each search finds its own")

    # No bundle is a borrowed surface: one with no search of its own standing on
    # another method's geometry is the courtesy solve, and it is gone.
    for i, a in enumerate(after["lem"]):
        for b in after["lem"][i + 1:]:
            if a.get("search") is None and b.get("search") is None:
                continue
            if a.get("failure_surface") is b.get("failure_surface") \
                    and (a.get("search") is None or b.get("search") is None):
                fails.append(f"{a.get('method')} and {b.get('method')} stand on "
                             f"one surface, and only one of them searched for it")

    # The report built from those runs documents both, in full.
    with contextlib.redirect_stdout(io.StringIO()):
        report = build_report(slope_data, after, dict(opts),
                              tempfile.mkdtemp(prefix="xslope_ranfor_"))
    titles = [t for _l, t in report.section_titles()]
    if method_label("spencer") not in titles:
        fails.append(f"the method the report ran got no detail block: {titles}")
    block = next((s for sec in report.sections for _l, s in sec.walk()
                  if s.title == method_label("spencer")), None)
    if block is not None:
        kids = [c.title for c in block.children]
        if "Search for the Critical Surface" not in kids:
            fails.append(f"the method the report searched for documents no "
                         f"search: {kids}")
    rows = {r[0]: r[1] for t in report.tables()
            if t.caption == "Computed factors of safety" for r in t.rows}
    fs = rows.get(method_label("spencer"))
    if fs is None or fs == "did not converge":
        fails.append(f"the summary does not carry Spencer's own answer: {rows}")
    if fs is not None and fs == rows.get(method_label("bishop")):
        fails.append(f"both methods report the same factor of safety on "
                     f"different surfaces: {rows}")

    # Stopping the report stops the SEARCH it is making, at the search's next
    # iteration boundary. A cancel that is only read between figures leaves the
    # user holding a Cancel button through the longest stretch of the build —
    # minutes, on a model worth searching. The flag is asked here from the start,
    # so the run unwinds at the first boundary it reaches rather than at the end.
    from xslope.search import AnalysisCancelled
    stopped = "the build ran to the end"
    try:
        with contextlib.redirect_stdout(io.StringIO()):
            build_report(slope_data, solutions,
                         dict(opts, cancel_check=lambda: True),
                         tempfile.mkdtemp(prefix="xslope_cancel_"))
    except AnalysisCancelled:
        stopped = "cancelled"
    except Exception as exc:                                # noqa: BLE001
        stopped = f"raised {exc!r}"
    if stopped != "cancelled":
        fails.append(f"a report cancelled while it was solving did not stop: "
                     f"{stopped}")
    # And the cancellation reaches the caller of generate_report rather than
    # being turned into a failure message, which is what a build wrapped in
    # "except Exception" would have done with it.
    from xslope.report import generate_report
    stopped = "generate_report returned"
    try:
        with tempfile.TemporaryDirectory() as tmp:
            with contextlib.redirect_stdout(io.StringIO()):
                generate_report(slope_data, solutions,
                                dict(opts, cancel_check=lambda: True),
                                os.path.join(tmp, "cancelled.docx"))
    except AnalysisCancelled:
        stopped = "cancelled"
    if stopped != "cancelled":
        fails.append(f"generate_report swallowed the cancellation: {stopped}")
    return fails


def test_a_report_run_method_is_run_at_the_analysis_discretization():
    """A method the report runs is cut into the same slices the analysis was.

    The count a run was made with cannot be read back off its slice table: the
    slicer splits at material and water boundaries, so a run of 20 comes back as
    21 rows. Inferring the option from the table ran the second method at 21,
    which is a different discretization of the same surface — and the section's
    Analysis Inputs, which prints the slice count only where the methods share
    one, silently lost the row.

    So the run records what it was made under and the record is what a second
    method is run from. Pinned end to end: a scripted 20 stays a scripted 20,
    both tables come out the same size, and the shared row is printed.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.report import (analysis_options, build_report, resolve_options,
                               run_requested_methods)
    from xslope.search import run_lem_analysis

    slope_data = load_slope_data(REINF_XLSX)
    asked = 20
    with contextlib.redirect_stdout(io.StringIO()):
        first = run_lem_analysis(slope_data, "bishop", analysis="single_surface",
                                 num_slices=asked)
    if len(first["slice_df"]) == asked:
        return [f"this model cuts {asked} slices into {asked} rows, so the "
                f"count cannot be told from the table and nothing is pinned"]
    if (first.get("options") or {}).get("num_slices") != asked:
        fails.append(f"the run does not record the slice count it was made "
                     f"with: {first.get('options')}")

    solutions = {"lem": [first]}
    opts = resolve_options({"input_path": REINF_XLSX,
                            "method": ["bishop", "spencer"], "pd_figure": False,
                            "lem_inputs_figure": False,
                            "lem_search_figure": False,
                            "lem_solution_figure": False,
                            "lem_slice_key": False})
    if analysis_options(slope_data, solutions)["num_slices"] != asked:
        fails.append(f"the analysis is read as "
                     f"{analysis_options(slope_data, solutions)['num_slices']} "
                     f"slices, not the {asked} it was run with")
    with contextlib.redirect_stdout(io.StringIO()):
        after = run_requested_methods(slope_data, solutions, opts)
    second = next((b for b in after["lem"] if b.get("method") == "spencer"), None)
    if second is None:
        return fails + ["Spencer was asked for and not run"]
    if (second.get("options") or {}).get("num_slices") != asked:
        fails.append(f"the method the report ran was asked for "
                     f"{(second.get('options') or {}).get('num_slices')} slices, "
                     f"not the {asked} the analysis was run with")
    if len(second["slice_df"]) != len(first["slice_df"]):
        fails.append(f"the analysis was cut into {len(first['slice_df'])} slices "
                     f"and the method the report ran into "
                     f"{len(second['slice_df'])}")

    # …so the section can still state one slice count for the analysis.
    with contextlib.redirect_stdout(io.StringIO()):
        report = build_report(slope_data, after, dict(opts),
                              tempfile.mkdtemp(prefix="xslope_slices_"))
    shared = _shared_slice_counts(_analysis_inputs(report))
    if shared != [str(len(first["slice_df"]))]:
        fails.append(f"the two methods are cut alike and the engine inputs "
                     f"print {shared} for the slice count")
    return fails


def test_a_method_that_is_not_run_says_why():
    """A method the report is asked for and does not run says so, and says why.

    Two names are declined rather than run: one this surface family cannot take,
    and one the solver does not offer. Both used to disappear — no section, no
    row, nothing on the console — which is the same silence this whole
    arrangement was built to end, arriving by a different door. The reason
    printed is the rule's own, so it is the sentence the Run dialog dims that
    method with.

    Read off BUILD_REPORT, which is the path a report takes. Asking the solving
    function directly proves nothing about the report: the builder used to call
    it only when there was something to run, so a report whose every requested
    method was declined — the one case where the sentence is the whole of what
    the user gets — went through a build that never reached it and said nothing
    at all.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.preflight import method_surface_reason
    from xslope.report import build_report, method_label

    slope_data = load_slope_data(NONCIRC_XLSX)
    from xslope.search import run_lem_analysis
    with contextlib.redirect_stdout(io.StringIO()):
        first = run_lem_analysis(slope_data, "spencer",
                                 analysis="single_surface",
                                 surface="noncircular", num_slices=20)
    solutions = {"lem": [first]}
    base = {"input_path": NONCIRC_XLSX, "pd_figure": False,
            "lem_inputs_figure": False, "lem_search_figure": False,
            "lem_solution_figure": False, "lem_slice_key": False,
            "lem_slice_table": False, "lem_calculations": False}

    def built(methods):
        said = io.StringIO()
        with tempfile.TemporaryDirectory() as tmp:
            with contextlib.redirect_stdout(said):
                report = build_report(slope_data, solutions,
                                      dict(base, method=methods), tmp)
        return report, said.getvalue()

    reason = method_surface_reason("bishop", "noncircular")
    cases = [
        # A method that IS run stands beside the declines.
        ("with a method to run", ["spencer", "bishop", "no_such_method"]),
        # …and the case the builder used to skip entirely: nothing to run, so
        # the declines are the whole of what there is to say.
        ("with nothing to run", ["bishop", "no_such_method"]),
    ]
    for where, methods in cases:
        report, said = built(methods)
        if method_label("bishop") not in said:
            fails.append(f"{where}: Bishop was asked for and declined without a "
                         f"word: {said!r}")
        elif reason.split(",")[0] not in said:
            fails.append(f"{where}: Bishop is declined without the rule's own "
                         f"reason: {said!r}")
        if "not a method this solver offers" not in said:
            fails.append(f"{where}: a method the solver does not have was "
                         f"declined without saying so: {said!r}")
        # And no declined method is documented as though it had run.
        titles = {t for _l, t in report.section_titles()}
        for name in ("bishop", "no_such_method"):
            if method_label(name) in titles:
                fails.append(f"{where}: {name} was declined and documented "
                             f"anyway")
    return fails


def test_rapid_drawdown_names_the_governing_stage():
    """The rapid drawdown block says which stage the reported factor came from.

    Three stage factors stood above one unattributed number, and which stage
    governs IS the engineering answer: stage 2 governing means the undrained
    strengths control the drawn-down section, stage 3 means the drained ones do.
    A factor that is neither stage's is left unattributed rather than assigned to
    the nearer one.
    """
    fails = []
    from xslope.report import _Counter, _rapid_section

    def _governing(s2, s3, fs=None):
        res = {"stage1_FS": 1.900, "stage2_FS": s2, "stage3_FS": s3,
               "FS": min(s2, s3) if fs is None else fs}
        sub = _rapid_section(res, _Counter())
        kv = next(b for b in sub.blocks if b.kind == "keyvalues")
        return next(label for label, _v in kv.items if "Governing" in label)

    # Undrained controls.
    label = _governing(1.204, 1.611)
    if "stage 2" not in label:
        fails.append(f"stage 2 gave the lower factor and the governing line "
                     f"reads {label!r}")
    # Drained controls.
    label = _governing(1.611, 1.204)
    if "stage 3" not in label:
        fails.append(f"stage 3 gave the lower factor and the governing line "
                     f"reads {label!r}")
    # A reported factor that is neither stage's names no stage.
    label = _governing(1.611, 1.204, fs=0.900)
    if "stage" in label:
        fails.append(f"a factor matching neither stage was attributed to one: "
                     f"{label!r}")
    return fails


def test_fs_summary_reaches_the_document():
    """The summary the tree holds is the summary the document prints: the
    documented methods, and no others.

    The rows used to be set in bold to pick the documented methods out of a table
    that also listed every method that had not been run. There is nothing left to
    pick out — every row is a method the report documents — so what the document
    owes the reader is that the table it prints is that table and not a longer
    one it inherited.
    """
    fails = []
    from docx import Document

    from xslope.report import generate_report, method_label, supported_methods

    slope_data, solutions = _solved()
    featured = ["bishop", "spencer"]
    expect = [method_label(m) for m in featured]

    block = _fs_table_block(_build({"method": featured}))
    if block is None:
        return ["there is no factor-of-safety summary"]
    if [r[0] for r in block.rows] != expect:
        fails.append(f"the tree's summary lists {[r[0] for r in block.rows]}, "
                     f"not {expect}")

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "summary.docx")
        ok, out = generate_report(
            slope_data, solutions,
            {"input_path": REINF_XLSX, "method": featured, "pd_figure": False,
             "lem_search_figure": False, "lem_solution_figure": False,
             "lem_slice_key": False}, path)
        if not ok:
            return fails + [f"generate_report failed: {out}"]

        doc = Document(path)
        table = next((t for t in doc.tables
                      if [c.text for c in t.rows[0].cells][:2]
                      == ["Method", "Factor of safety"]), None)
        if table is None:
            return fails + ["the summary table is not in the document"]
        printed = [row.cells[0].text for row in table.rows[1:]]
        if printed != expect:
            fails.append(f"the document's summary lists {printed}, not {expect}")
        if len(printed) >= len(supported_methods()):
            fails.append(f"the summary still enumerates the solver "
                         f"({len(printed)} rows for {len(featured)} documented "
                         f"methods)")
    return fails


#: Which slice numbers a model's key leaves touching, by model.
#:
#: The numbers stand at the mid-height of their own slices and come down in size
#: together until they clear, so on any ordinary section none of them touch. The
#: dam is not an ordinary section: cut into forty, its crest carries slice 38 at
#: a foot and a half wide between two slices eight and a half feet wide, and no
#: size a reader could read separates a number in that slice from the numbers
#: either side of it. That pair, and only that pair, is what the key gives up —
#: pinned here, so a change that starts printing numbers over each other anywhere
#: else is a failure rather than one more entry on a list.
_SLICE_KEY_TOUCHING = {
    "the sample": [],
    "the dam": [("37", "38"), ("38", "39")],
}


def test_slice_key_figure():
    """Every slice table is preceded by a figure of the slices it lists, numbered.

    The table's first column is a slice number and nothing else in the report
    says where slice 9 is. The key is a figure of the sliced mass with the numbers
    on it, and it has to sit immediately BEFORE the table it keys — a key printed
    after the table is a key the reader has already needed.
    """
    fails = []
    from xslope.report import method_label

    tmp = tempfile.mkdtemp(prefix="xslope_key_")
    methods = ["spencer", "bishop"]
    report = _build({"method": methods, "lem_solution_figure": False},
                    figure_dir=tmp, fast=False)

    for m in methods:
        label = method_label(m)
        keys = [f for f in report.figures() if f.source == f"{m} slice key"]
        if len(keys) != 1:
            fails.append(f"{m}: {len(keys)} slice-key figures")
            continue
        key = keys[0]
        if label not in key.caption:
            fails.append(f"{m}: the slice key is captioned {key.caption!r}")
        if not os.path.exists(key.path):
            fails.append(f"{m}: the slice key was not written to {key.path}")
        elif os.path.getsize(key.path) < 20000:
            fails.append(f"{m}: the slice key is {os.path.getsize(key.path)} "
                         f"bytes — too small to be a rendered plot")
        # A portrait figure, printed at a stated width. The landscape page
        # belongs to the table's twenty columns; a key that took one too spent a
        # sheet on a picture that reads at a sixth of one.
        #
        # The width itself is the width the key's own numbers ask for, between
        # the narrowest a key is printed and the text width. The text width is a
        # limit, not a size the key is kept away from: a crowded surface asks for
        # more width than the page has, and what it gets is all of it. Printing
        # such a key narrower than the page allows would be shrinking the one
        # key whose numbers are already smallest — the nail model's fifteen
        # slices come out at 5.5 pt at the full width, and a forty-slice surface
        # at 5.2 pt, both of which read and neither of which has anything to
        # spare.
        if key.landscape or key.width_in <= 0:
            fails.append(f"{m}: the slice key takes a landscape page instead of "
                         f"printing at a stated width (landscape={key.landscape}, "
                         f"width_in={key.width_in})")
        from xslope.report import Figure as ReportFigure, SLICE_KEY_MIN_IN
        text_width = ReportFigure("", "").width_in
        if not SLICE_KEY_MIN_IN <= key.width_in <= text_width:
            fails.append(f"{m}: the slice key is {key.width_in} in wide, outside "
                         f"the {SLICE_KEY_MIN_IN} to {text_width} in it is "
                         f"printed between")

    # Immediately before the table, in the same section, with nothing between.
    for sec in report.sections:
        for _lvl, node in sec.walk():
            if node.title != "Slice Table":
                continue
            kinds = [b.kind for b in node.blocks]
            if kinds[-2:] != ["figure", "table"]:
                fails.append(f"the Slice Table section runs {kinds}; the key must "
                             f"be the block immediately before the table")

    # And it is the slice-key option that puts it there.
    off = _build({"method": methods, "lem_slice_key": False})
    if [f for f in off.figures() if "slice key" in f.source]:
        fails.append("lem_slice_key=False still drew a slice key")
    if not [t for t in off.tables() if t.landscape]:
        fails.append("lem_slice_key=False took the slice table with it")

    # The frame is the SLICED MASS, not the model: the key is tighter than the
    # solution plot of the same surface, which is the whole point of it.
    from xslope.plot import sliced_mass_bounds
    import matplotlib
    matplotlib.use("Agg")
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    from matplotlib.figure import Figure as MplFigure
    from xslope.plot import plot_solution

    slope_data, solutions = _solved()
    bundle = solutions["lem"][0]
    spans = {}
    for frame in ("fill", "slices"):
        fig = MplFigure(figsize=(9.0, 5.5))
        FigureCanvasAgg(fig)
        plot_solution(slope_data, bundle["slice_df"], bundle["failure_surface"],
                      bundle["results"], fig=fig, show_title=False,
                      slice_numbers=(frame == "slices"), frame=frame)
        ax = fig.axes[0]
        spans[frame] = (ax.get_xlim(), ax.get_ylim(), sliced_mass_bounds(ax),
                        [t for t in ax.texts if t.get_gid() == "SLICE_NUMBER"])
    (fx0, fx1), _fy, _fb, fill_labels = spans["fill"]
    (sx0, sx1), (sy0, sy1), bounds, labels = spans["slices"]
    if fill_labels:
        fails.append("frame='fill' drew slice numbers without being asked")
    if len(labels) != len(bundle["slice_df"]):
        fails.append(f"{len(labels)} slice labels for "
                     f"{len(bundle['slice_df'])} slices")
    if not (sx1 - sx0) < (fx1 - fx0):
        fails.append(f"the slice frame ({sx0:.1f}, {sx1:.1f}) is no tighter than "
                     f"the model frame ({fx0:.1f}, {fx1:.1f})")
    if bounds is None:
        fails.append("the sliced mass has no measurable bounds")
    else:
        bx0, bx1, by0, by1 = bounds
        # Everything the sliced mass draws is inside the frame — the base-stress
        # bars stand off the base and are the reason a frame on the slice corners
        # alone would clip.
        if not (sx0 <= bx0 and bx1 <= sx1 and sy0 <= by0 and by1 <= sy1):
            fails.append(f"the frame ({sx0:.1f}..{sx1:.1f}, {sy0:.1f}..{sy1:.1f}) "
                         f"clips the sliced mass ({bx0:.1f}..{bx1:.1f}, "
                         f"{by0:.1f}..{by1:.1f})")
        # And the cushion is uniform and modest, not a slab of white.
        pads = (bx0 - sx0, sx1 - bx1, by0 - sy0, sy1 - by1)
        size = max(bx1 - bx0, by1 - by0)
        if max(pads) > 0.10 * size:
            fails.append(f"the cushion is {max(pads) / size:.0%} of the mass; the "
                         f"frame is not tight")
        if max(pads) - min(pads) > 1e-6 * size:
            fails.append(f"the cushion is not uniform: {pads}")

    # Every number stands at the mid-height of its own slice, they are all one
    # size, that size reads ON THE PAGE, and none is printed over another except
    # where a slice is too narrow to hold a legible number at any size.
    #
    # The key is drawn at its own figure size and printed at the width its
    # smallest number asks for, so the size on the page is that number through
    # the same function the builder prints it by — which is what makes a change
    # to either constant fail here rather than on the page.
    #
    # Three models, because they land in three different places against the
    # width rule. The sample's fifteen slices are sparse and are printed well
    # inside the text width. The nail model's fifteen stand on a mass as tall as
    # it is wide, which leaves the numbers a narrow band and asks for more width
    # than the page has. The dam's forty ask for more still, and its crest
    # carries a sliver a sixth the width of the slices either side of it. The
    # last two are the case the width rule CLAMPS, which is where the size on
    # the page is smallest and the only place the legibility floor is really
    # tested.
    import numpy as np

    from xslope.plot import SLICE_LABEL_MIN_PT
    from xslope.report import (Figure as ReportFigure, SLICE_KEY_MIN_IN,
                               SLICE_KEY_SIZE, slice_key_width)
    from xslope.slice import generate_slices
    from xslope.solve import solve_selected

    def _keyed(path, n):
        """A solved model to key, from its own first circle."""
        with contextlib.redirect_stdout(io.StringIO()):
            sd = load_slope_data_cached(path)
            ok, out = generate_slices(sd, circle=sd["circles"][0], num_slices=n)
            if not ok:
                return None
            return sd, out[0], out[1], solve_selected("spencer", out[0])

    cases = [("the sample", slope_data, bundle["slice_df"],
              bundle["failure_surface"], bundle["results"])]
    for name, path, n in (("the nail model", AXIAL_XLSX, 15),
                          ("the dam", SEEP_XLSX, 40)):
        keyed = _keyed(path, n)
        if keyed is None:
            fails.append(f"{name} produced no slices; the clamped key is "
                         f"untested")
        else:
            cases.append((name,) + keyed)

    for name, sd, df, surface, results in cases:
        fig = MplFigure(figsize=SLICE_KEY_SIZE)
        FigureCanvasAgg(fig)
        with contextlib.redirect_stdout(io.StringIO()):
            plot_solution(sd, df, surface, results, fig=fig, show_title=False,
                          slice_numbers=True, frame="slices")
        fig.canvas.draw()
        ax = fig.axes[0]
        renderer = fig.canvas.get_renderer()
        marks = [t for t in ax.texts if t.get_gid() == "SLICE_NUMBER"]

        # --- one band, at mid-height ---
        #
        # Every number at the middle of its own slice: halfway up at the slice's
        # own center line, and nowhere else. Alternate numbers used to stand a
        # line higher to buy the key a larger size, and a key whose numbers zigzag
        # reads as saying something about the slices.
        #
        # The mid-height is taken from the slice frame — the base and the top of
        # each slice, which is what "mid-height" means — and NOT from the
        # function the plotter places the numbers with. Asking that function
        # where the numbers should be is asking the plotter whether it agrees
        # with itself, which it does whatever it is made to draw: a stagger
        # reintroduced there would move every number and every expectation with
        # it, and pass.
        want = (df['y_cb'].values.astype(float)
                + df['y_ct'].values.astype(float)) / 2.0
        astray = [(m.get_text(), round(m.get_position()[1] - want[i], 4))
                  for i, m in enumerate(marks)
                  if abs(m.get_position()[1] - want[i]) > 1e-9]
        if astray:
            fails.append(f"{name}: slice numbers stand off the mid-height of "
                         f"their own slice: {astray[:4]}")
        off_center = [m.get_text() for i, m in enumerate(marks)
                      if abs(m.get_position()[0] - float(df['x_c'].values[i]))
                      > 1e-9]
        if off_center:
            fails.append(f"{name}: slice numbers stand off the center of their "
                         f"own slice: {off_center[:4]}")

        # --- one size ---
        sizes = {round(m.get_fontsize(), 6) for m in marks}
        if len(sizes) > 1:
            fails.append(f"{name}: the slice numbers are set at {sorted(sizes)}, "
                         f"not one size")
        smallest = min((t.get_fontsize() for t in marks), default=0.0)
        width_in = slice_key_width(smallest)
        printed = width_in / SLICE_KEY_SIZE[0]
        if smallest * printed < _LEGIBLE_PT:
            fails.append(f"{name}: the slice numbers are set at "
                         f"{smallest:.1f} pt, which is "
                         f"{smallest * printed:.1f} pt on the page — under the "
                         f"{_LEGIBLE_PT} pt a number has to be to be read")
        # The width this model's key would be printed at is a width the page can
        # take. A model whose numbers ask for more than the text width is
        # printed at the text width, and the legibility floor above is what says
        # whether that is enough — so a rule that stopped short of the page, or
        # ran past it, is caught on the models that reach the limit rather than
        # on the one that never approaches it.
        if not SLICE_KEY_MIN_IN <= width_in <= ReportFigure("", "").width_in:
            fails.append(f"{name}: its key would be printed {width_in} in wide, "
                         f"outside the {SLICE_KEY_MIN_IN} to "
                         f"{ReportFigure('', '').width_in} in a key is printed "
                         f"between")

        # --- and they clear each other, unless the slice cannot hold them ---
        #
        # The size comes down until they clear, and it comes down for all of them
        # at once. Where even the floor is not small enough the floor is what is
        # kept: a slice a sixth the width of the ones either side has no room at
        # its own mid-height for a number anybody could read, and a key set at
        # four points to separate that one pair keys nothing at all. So an
        # overlap is allowed only at the floor, and only against a slice narrower
        # than the label printed on it.
        boxes = [t.get_window_extent(renderer) for t in marks]
        widths = ax.transData.transform(np.column_stack(
            [df['x_r'].values.astype(float), want])) - ax.transData.transform(
            np.column_stack([df['x_l'].values.astype(float), want]))
        touching = [(i, j) for i in range(len(boxes))
                    for j in range(i + 1, len(boxes))
                    if boxes[i].overlaps(boxes[j])]
        for i, j in touching:
            at_floor = abs(smallest - SLICE_LABEL_MIN_PT) < 1e-9
            narrow = min(widths[i][0], widths[j][0]) < max(boxes[i].width,
                                                           boxes[j].width)
            if not (at_floor and narrow):
                fails.append(
                    f"{name}: slice numbers {marks[i].get_text()} and "
                    f"{marks[j].get_text()} are printed over each other at "
                    f"{smallest:.2f} pt, and neither slice is too narrow to "
                    f"hold one ({widths[i][0]:.1f} and {widths[j][0]:.1f} px "
                    f"wide against a {max(boxes[i].width, boxes[j].width):.1f} "
                    f"px label)")
        pinned = _SLICE_KEY_TOUCHING.get(name)
        got = sorted((marks[i].get_text(), marks[j].get_text())
                     for i, j in touching)
        if pinned is not None and got != pinned:
            fails.append(f"{name}: the numbers that touch are {got}; "
                         f"{pinned} is what this model leaves touching")
    return fails


def test_figure_progress_counts():
    """The figure count a caller is shown is the number of figures there are.

    N methods mean N critical-surface plots and N slice keys; a progress line
    that said "3" while eleven were rendering would be worse than none.
    """
    fails = []
    from xslope.report import planned_figures, resolve_options

    seen = []
    tmp = tempfile.mkdtemp(prefix="xslope_prog_")
    opts = {"method": ["spencer", "bishop"],
            "progress": lambda done, total, label: seen.append((done, total, label))}
    report = _build(opts, figure_dir=tmp, fast=False)

    slope_data, solutions = _solved()
    planned = planned_figures(slope_data, solutions,
                              resolve_options({"input_path": REINF_XLSX,
                                               "title": "Sample Levee", **opts}))
    drawn = len(report.figures())
    if planned != drawn:
        fails.append(f"{planned} figures were planned and {drawn} were built")
    if [d for d, _t, _l in seen] != list(range(1, len(seen) + 1)):
        fails.append(f"the progress steps are {[d for d, _t, _l in seen]}")
    if seen and seen[-1][0] != seen[-1][1]:
        fails.append(f"the last step is {seen[-1][0]} of {seen[-1][1]}")
    if {t for _d, t, _l in seen} != {planned}:
        fails.append(f"the total reported was {[t for _d, t, _l in seen]}, not "
                     f"{planned}")
    labels = [l for _d, _t, l in seen]
    if len([l for l in labels if "slice key" in l]) != 2:
        fails.append(f"the two slice keys are not both counted: {labels}")
    return fails


def test_fs_table_compares_only_what_was_documented():
    """The summary lists the methods the report documents, each with its own
    answer, and only where there is something to compare.

    It used to list every method xslope offers, filling in the ones that had not
    been run by solving them on the surface the featured method had found. Each
    method has its own critical surface, so those rows were not that method's
    factor of safety — a column of numbers that read as a comparison and was not
    one. What the reader wants compared is what the reader asked the report to
    document, so those are the rows, and a report of ONE method carries no table
    at all: one row is the number that method's own section already states, and a
    second statement of a number is a second number to keep in step with it.
    """
    fails = []
    from xslope.report import build_report, method_label

    slope_data, solutions = _solved()

    # One method: no table, and the factor of safety is still stated.
    single = _build({"method": "spencer"})
    if _fs_rows(single) is not None:
        fails.append("a single-method report still carries a comparison table")
    if "Factors of Safety" in [t for _l, t in single.section_titles()]:
        fails.append("a single-method report still has a Factors of Safety "
                     "section")
    fs = solutions["lem"][0]["results"]["FS"]
    if not any(f"{fs:.3f}" in b.text for b in single.blocks("prose")):
        fails.append(f"a single-method report states its factor of safety "
                     f"{fs:.3f} nowhere")

    # Two methods: a table of exactly those two, each carrying its own answer.
    pair = _build({"method": ["spencer", "bishop"]})
    got = _fs_rows(pair)
    if got is None:
        return fails + ["a two-method report carries no summary table"]
    _headers, rows = got
    want = {method_label(b["method"]): b["results"]["FS"]
            for b in solutions["lem"]}
    if [r[0] for r in rows] != list(want):
        fails.append(f"the summary lists {[r[0] for r in rows]}, not the "
                     f"documented methods {list(want)}")
    for row in rows:
        expect = want.get(row[0])
        if expect is None:
            fails.append(f"{row[0]} is in the summary and is not documented")
        elif row[1] != f"{expect:.3f}":
            fails.append(f"{row[0]}: the summary says {row[1]!r}, its own run "
                         f"gave {expect:.3f}")

    # And the surface every row's number belongs to is in the PARAGRAPH, not in a
    # column of the table. Every row of one table is arrived at the same way, so
    # the column was one sentence repeated down the rows, standing between the
    # reader and the numbers they came for
    # (test_the_fs_summary_says_where_its_numbers_came_from holds the sentence).
    headers, _rows = got
    for banned in ("Surface", "surface"):
        if banned in headers:
            fails.append(f"the summary carries a {banned!r} column: {headers}")
    if any("critical" in cell.lower() or "specified" in cell.lower()
           for row in rows for cell in row):
        fails.append(f"a row of the summary states its own provenance: {rows}")

    # A method that does not converge is stated, not dropped.
    bundle = dict(solutions["lem"][0])
    df = bundle["slice_df"].copy()
    df["w"] = float("nan")             # no method can produce an answer from this
    bad = {"lem": [{"slice_df": df, "failure_surface": bundle["failure_surface"],
                    "results": {}, "search": None, "method": "spencer"},
                   {"slice_df": df, "failure_surface": bundle["failure_surface"],
                    "results": {}, "search": None, "method": "bishop"}]}
    with tempfile.TemporaryDirectory() as tmp:
        broken = build_report(slope_data, bad,
                              {"pd_figure": False, "lem_solution_figure": False,
                               "lem_search_figure": False,
                               "method": ["spencer", "bishop"]}, tmp)
    got = _fs_rows(broken)
    if got is None:
        fails.append("a report whose methods all failed lost its summary table")
    else:
        _h, rows = got
        if not [r for r in rows if "did not converge" in r[1]]:
            fails.append(f"no method reported 'did not converge' on an "
                         f"unsolvable slice table: {rows}")
        if len(rows) != 2:
            fails.append(f"the unsolvable run dropped rows: {len(rows)} of 2")
    return fails


# --------------------------------------------------------------------------
# D. the document
# --------------------------------------------------------------------------

def _docx_parts(path):
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        return names, {n: z.read(n).decode("utf-8", "replace")
                       for n in names if n.endswith(".xml")}


def test_docx():
    """The written .docx is a real OOXML package with the structure the report
    depends on."""
    fails = []
    from xslope.report import generate_report

    slope_data, solutions = _solved()
    with tempfile.TemporaryDirectory() as tmp:
        out_path = os.path.join(tmp, "report.docx")
        ok, out = generate_report(
            slope_data, solutions,
            {"input_path": REINF_XLSX, "title": "Sample Levee Report",
             "project_number": "26-001", "organization": "Example Engineering",
             "author": "A. Engineer", "method": "spencer",
             "signature_lines": True, "lem_search_figure": False},
            out_path)
        if not ok:
            return [f"generate_report failed: {out}"]
        if not os.path.exists(out_path):
            return ["generate_report reported success but wrote no file"]

        names, xml = _docx_parts(out_path)
        doc = xml.get("word/document.xml", "")
        if not doc:
            return ["the package has no word/document.xml"]

        if "Sample Levee Report" not in doc:
            fails.append("the document does not contain the project title")
        if "A. Engineer" not in doc:
            fails.append("the document does not contain the author")
        if 'DOCPROPERTY "Title"' not in doc:
            fails.append("the title page does not use a document-property field")
        if "TOC \\o" not in doc:
            fails.append("the document has no table-of-contents field")
        if 'w:dirty="true"' in doc:
            fails.append("a field is marked dirty; Word would prompt about "
                         "updating fields on every open")
        if "Prepared by" not in doc or "Checked by" not in doc:
            fails.append("signature_lines=True produced no signature lines")

        # The template's styles are the ones referenced.
        for style in ("Heading1", "Caption", "TableGrid"):
            if style not in doc:
                fails.append(f"the document references no {style} style")

        # The slice table gets its own landscape section, and the report returns
        # to portrait after it.
        orients = [chunk.split('w:orient="')[1].split('"')[0]
                   if 'w:orient="' in chunk else "portrait"
                   for chunk in doc.split("<w:pgSz ")[1:]]
        if orients.count("landscape") != 1:
            fails.append(f"page orientations are {orients}; exactly one landscape "
                         f"section is expected")
        if not orients or orients[-1] != "portrait":
            fails.append(f"the report ends on a {orients[-1:]} page")

        # Running head and foot, with live fields.
        header = next((v for k, v in xml.items() if k.startswith("word/header")), "")
        footer = next((v for k, v in xml.items() if k.startswith("word/footer")), "")
        if 'DOCPROPERTY "Title"' not in header:
            fails.append("the running head carries no title field")
        if "PAGE" not in footer or "NUMPAGES" not in footer:
            fails.append("the footer has no page-of-pages fields")

        # The figures actually went in.
        media = [n for n in names if n.startswith("word/media/")]
        if len(media) < 2:
            fails.append(f"only {len(media)} images were embedded")

        # A sentence that introduces a figure is bound to it. Four lead sentences
        # per transient report were left at the foot of one page with the figure
        # they name at the top of the next, one of them over a third of a blank
        # page. keepNext is what moves the sentence to the figure.
        import re
        paras = re.findall(r"<w:p[ >].*?</w:p>", doc, re.S)
        leads = 0
        for prev, para in zip(paras, paras[1:]):
            if "<w:drawing" not in para or 'w:val="BodyText"' not in prev:
                continue
            leads += 1
            if "<w:keepNext/>" not in prev:
                text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", prev))
                fails.append(f"the sentence a figure stands under is not kept "
                             f"with it: {text[:80]!r}")
        if not leads:
            fails.append("no figure in the report stands under a body sentence, "
                         "so nothing here checks that the two stay together")

        # And a figure's caption is bound to nothing after it. The template's
        # Caption style carries keep-with-next — right for a table, whose caption
        # stands above its first row — and it chained each figure to the heading,
        # sentence and figure that followed. With a lead sentence bound to its own
        # figure as well, the chain ran longer than a page: the reinforcement
        # report grew six pages of half-blank ones, each holding a sentence whose
        # figure had been pushed to the next.
        captions = [p for p in paras
                    if 'w:val="Caption"' in p
                    and re.search(r"<w:t[^>]*>Figure \d", p)]
        if not captions:
            fails.append("the report has no figure captions to check")
        for p in captions:
            if '<w:keepNext w:val="0"/>' not in p:
                text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", p))
                fails.append(f"a figure caption is kept with whatever follows it: "
                             f"{text[:60]!r}")
        # The table caption keeps its own binding: it stands above its first row.
        table_caps = [p for p in paras
                      if 'w:val="Caption"' in p
                      and re.search(r"<w:t[^>]*>Table \d", p)]
        for p in table_caps:
            if "<w:keepNext/>" not in p:
                text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", p))
                fails.append(f"a table caption is not kept with its table: "
                             f"{text[:60]!r}")

        # And the returned bundle names what was produced.
        if out["path"] != out_path:
            fails.append(f"generate_report returned {out['path']!r}")
        if not out["figures"]:
            fails.append("generate_report reported no figures")

    # An unavailable format is refused, in words, rather than half-written.
    with tempfile.TemporaryDirectory() as tmp:
        ok, msg = generate_report(slope_data, solutions, {},
                                  os.path.join(tmp, "r.pdf"))
        if ok:
            fails.append("a PDF report was reported as written; PDF is S3")
        elif "not available" not in str(msg):
            fails.append(f"the PDF refusal reads {msg!r}")
        if os.listdir(tmp):
            fails.append(f"the refused format left files behind: {os.listdir(tmp)}")
    return fails


def _sections_and_headers(path):
    """Every Word section of ``path``, in order, with the header part it prints.

    Returns ``(sections, headers)`` — one dict per section, carrying its page
    geometry in twips, whether it is a "different first page" section, and the
    name of its default header part; and the header parts themselves, by name.
    """
    import re
    with zipfile.ZipFile(path) as z:
        doc = z.read("word/document.xml").decode()
        rels = z.read("word/_rels/document.xml.rels").decode()
        headers = {n.rsplit("/", 1)[-1]: z.read(n).decode()
                   for n in z.namelist() if re.match(r"word/header\d+\.xml", n)}
    part = dict(re.findall(r'Id="([^"]+)"[^>]*Target="([^"]+)"', rels))

    sections = []
    for sect in re.findall(r"<w:sectPr\b.*?</w:sectPr>", doc, re.S):
        size = re.search(r"<w:pgSz [^>]*/>", sect)
        mar = re.search(r"<w:pgMar [^>]*/>", sect)

        def attr(el, name, default=0):
            m = el and re.search(r'w:%s="(-?\d+)"' % name, el.group(0))
            return int(m.group(1)) if m else default

        ref = re.search(r'<w:headerReference w:type="default" r:id="(\w+)"/>',
                        sect)
        sections.append({
            "landscape": 'w:orient="landscape"' in sect,
            "width": attr(size, "w"), "left": attr(mar, "left"),
            "right": attr(mar, "right"),
            "title_page": "<w:titlePg" in sect,
            "header": part.get(ref.group(1), "") if ref else "",
        })
    return sections, headers


def test_running_head_names_the_section():
    """The head over a body page names the section the page is in, as fields.

    The section name is two STYLEREF fields — the heading's number, then its
    text — so the head reads what the heading reads and follows the page into
    whatever section it lands in. STYLEREF is given the style's UI name; the
    "heading 1" that styles.xml stores is the internal name and resolves in
    neither Word nor LibreOffice.

    The front matter is excluded at the section break, not by hoping the field
    finds nothing: a STYLEREF over the contents page has no heading behind it
    and reaches FORWARD, which would label the contents with the first section
    of the report. So the title page and the contents keep a head that names the
    report alone, and the body opens a Word section of its own.

    The section name rides a tab to the right margin, and a landscape section's
    margin is three inches further out than a portrait one's. Each section's
    head is therefore its own — the stop is checked against the width of the
    page it prints on.
    """
    import re
    fails = []
    from xslope.report import generate_report
    from xslope.report_docx import RUNNING_HEAD_STYLE, STYLE

    if RUNNING_HEAD_STYLE != STYLE["heading"] % 1:
        fails.append(f"the running head reads {RUNNING_HEAD_STYLE!r}, which is "
                     f"not the style the top-level headings are written in")

    styleref = re.compile(r'STYLEREF "([^"]+)"(\s*\\n)?\s*<')

    slope_data, solutions = _solved()
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "report.docx")
        ok, out = generate_report(
            slope_data, solutions,
            {"input_path": REINF_XLSX, "title": "Sample Levee Report",
             "method": "spencer"}, path)
        if not ok:
            return fails + [f"generate_report failed: {out}"]

        sections, headers = _sections_and_headers(path)
        if len(sections) < 3:
            return fails + [f"the report has {len(sections)} Word sections; the "
                            f"front matter, the body and its landscape table "
                            f"are three"]

        front, body = sections[0], sections[1:]
        if not front["title_page"]:
            fails.append("the first section is not a different-first-page "
                         "section; the title page would carry a running head")
        head = headers.get(front["header"], "")
        if 'DOCPROPERTY "Title"' not in head:
            fails.append("the front matter's head carries no title field")
        if "STYLEREF" in head:
            fails.append("the front matter carries a STYLEREF head; over the "
                         "contents page it would name the first section of the "
                         "report")

        if len({s["header"] for s in body}) != len(body):
            fails.append(f"{len(body)} body sections share "
                         f"{len({s['header'] for s in body})} header parts; a "
                         f"shared head cannot put the tab at both margins")

        for i, sect in enumerate(body, start=1):
            where = f"section {i} ({'landscape' if sect['landscape'] else 'portrait'})"
            head = headers.get(sect["header"], "")
            if not head:
                fails.append(f"{where} prints no header part of its own")
                continue
            if sect["title_page"]:
                fails.append(f"{where} opens on a page with no running head")
            if 'DOCPROPERTY "Title"' not in head:
                fails.append(f"{where}'s head carries no title field")

            named = styleref.findall(head)
            if len(named) != 2:
                fails.append(f"{where}'s head has {len(named)} STYLEREF fields, "
                             f"not the heading's number and its text")
            for style, number in named:
                if style != RUNNING_HEAD_STYLE:
                    fails.append(f"{where}'s head reads the {style!r} style, "
                                 f"not {RUNNING_HEAD_STYLE!r}")
            if [bool(n) for _s, n in named] != [True, False]:
                fails.append(f"{where}'s head is not the heading's number "
                             f"followed by its text: {named}")
            if "<w:tab/>" not in head:
                fails.append(f"{where}'s head does not tab the section name "
                             f"away from the title")

            # The stop the tab lands on is this section's own right margin.
            want = sect["width"] - sect["left"] - sect["right"]
            stops = re.findall(r'<w:tab w:pos="(\d+)" w:val="(\w+)"/>', head)
            right = [int(pos) for pos, val in stops if val == "right"]
            if right != [want]:
                fails.append(f"{where}'s head tabs to {right}, not to its right "
                             f"margin at {want} twips")

        # Every heading the head can name is written in the style it reads.
        with zipfile.ZipFile(path) as z:
            doc = z.read("word/document.xml").decode()
        used = STYLE["heading"] % 1
        if f'<w:pStyle w:val="{used.replace(" ", "")}"/>' not in doc:
            fails.append(f"no paragraph is written in {used!r}; the head would "
                         f"name nothing")
    return fails


def _toc_result(doc_xml):
    """The text of the contents field's cached result, entry by entry.

    Everything between the field's ``separate`` and its ``end`` is what a reader
    sees before Word ever updates the field — so this is exactly what is on the
    contents page of a document nobody has pressed F9 in.
    """
    import re
    instr = doc_xml.index("TOC \\o")
    start = doc_xml.index('w:fldCharType="separate"', instr)
    end = doc_xml.index('w:fldCharType="end"', start)
    return re.findall(r"<w:t[^>]*>([^<]*)</w:t>", doc_xml[start:end])


def _sections_usable(doc_xml):
    """Every section's usable text width in twips, with the position in the XML
    where the section ends.

    A table belongs to the first section that ends after it — that is how Word
    reads a document, and it is what says whether the slice table had a landscape
    page's width to fill.
    """
    import re
    out = []
    for m in re.finditer(r"<w:sectPr[ >].*?</w:sectPr>", doc_xml, re.S):
        sect = m.group(0)
        page = re.search(r'<w:pgSz[^>]*\bw:w="(\d+)"', sect)
        margins = re.search(r"<w:pgMar[^>]*/>", sect)
        if page is None or margins is None:
            continue
        left = int(re.search(r'w:left="(\d+)"', margins.group(0)).group(1))
        right = int(re.search(r'w:right="(\d+)"', margins.group(0)).group(1))
        out.append((m.start(), int(page.group(1)) - left - right))
    return out


def _cell_margin(tbl_xml, styles_xml, style_id):
    """The leading cell margin a table carries, in twips.

    Its own declaration where it makes one — the report states its padding on
    every table it writes, because the template's is a body table's and starves a
    twenty-two column slice table of the width its numbers need — and the table
    style's otherwise. Whichever it is, it is the number the indent and the
    column measurement both have to be built on.
    """
    import re
    own = re.search(r"<w:tblCellMar>.*?<w:left [^>]*w:w=\"(\d+)\"",
                    tbl_xml, re.S)
    if own:
        return int(own.group(1))
    style = re.search(rf'<w:style [^>]*w:styleId="{style_id}".*?</w:style>',
                      styles_xml, re.S)
    if style is None:
        return None
    left = re.search(r"<w:tblCellMar>.*?<w:left [^>]*w:w=\"(\d+)\"",
                     style.group(0), re.S)
    return int(left.group(1)) if left else None


def _table_cell_texts(tbl_xml):
    """Everything that prints in each column of a table, one list per column.

    The header row, the body rows and the totals row — the same strings the
    renderer measured its columns from, read back out of the document.

    Word math counts too: a cell holding one would carry its characters in
    ``m:t`` runs rather than ``w:t`` ones, and read only for text runs a column
    of symbols would measure as empty and every width in it look stretched.
    """
    import re
    columns = []
    for row in re.findall(r"<w:tr[ >].*?</w:tr>", tbl_xml, re.S):
        for j, cell in enumerate(re.findall(r"<w:tc>.*?</w:tc>", row, re.S)):
            while len(columns) <= j:
                columns.append([])
            columns[j].append("".join(
                re.findall(r"<[wm]:t[^>]*>([^<]*)</[wm]:t>", cell)))
    return columns


def _table_point_size(tbl_xml):
    """The point size the table's cells are set in, as the document states it.

    Word stores a run's size in half-points; the size that appears most often in
    the table is the one its cells carry.
    """
    import re
    sizes = [int(v) for v in re.findall(r'<w:sz w:val="(\d+)"/>', tbl_xml)]
    return max(set(sizes), key=sizes.count) / 2.0 if sizes else None


def _content_widths(columns, family, size_pt, usable, margin):
    """What each column's content needs, in twips.

    Its widest line plus the cell margins either side, floored at its longest
    single word — Word breaks a word that will not fit rather than widening the
    column — and that floor capped at an equal share of the page, past which
    refusing to break would starve every other column. An empty column still
    reads as a column, so it floors at one em.
    """
    from xslope.report_docx import TWIPS_PER_PT, _text_width

    pad = 2 * margin
    fair_share = usable / max(1, len(columns))
    out = []
    for texts in columns:
        widest = max((_text_width(t, family, size_pt) for t in texts), default=0.0)
        longest_word = max((_text_width(w, family, size_pt)
                            for t in texts for w in t.split()), default=0.0)
        out.append(max(widest + pad,
                       min(max(longest_word, size_pt * TWIPS_PER_PT) + pad,
                           fair_share)))
    return out


def test_table_geometry():
    """Every table is laid out to its content, on the page it sits on.

    Three defects this defends against, none of which Word fixes on its own: a
    table with no indent hangs its left border out into the margin; an
    autofitting table gives "#" the same width as "Material"; and a table ruled
    across the full text width when it holds three columns of factors of safety
    is a table pretending to be a page. The cure is a fixed layout with measured
    columns, an indent of exactly one cell margin, and a width that is the
    content's — cut down to the text width only where the content would overrun
    it.
    """
    import re
    fails = []
    from xslope.report import generate_report
    from xslope.report_docx import DEFAULT_CELL_MARGIN

    slope_data, solutions = _solved()
    with tempfile.TemporaryDirectory() as tmp:
        out_path = os.path.join(tmp, "report.docx")
        # The model checks are on: their Finding column is the long-text column
        # the widths have to wrap rather than let it starve its neighbours.
        ok, out = generate_report(
            slope_data, solutions,
            {"input_path": REINF_XLSX, "title": "Sample Levee Report",
             "project_number": "26-001", "author": "A. Engineer",
             "method": "spencer", "signature_lines": True, "model_checks": True,
             "pd_figure": False, "lem_search_figure": False,
             "lem_solution_figure": False},
            out_path)
        if not ok:
            return [f"generate_report failed: {out}"]
        _names, xml = _docx_parts(out_path)
        doc = xml.get("word/document.xml", "")
        styles = xml.get("word/styles.xml", "")
        # The face the columns were measured in, read off the document's own
        # Normal style — the widths mean nothing without it.
        from docx import Document

        from xslope.report_docx import _table_font
        family = _table_font(Document(out_path))

    sections = _sections_usable(doc)
    if not sections:
        return ["the document declares no page size to fit the tables to"]
    tables = list(re.finditer(r"<w:tbl>.*?</w:tbl>", doc, re.S))
    if len(tables) < 5:
        return [f"only {len(tables)} tables were written; the report has more"]

    for i, m in enumerate(tables):
        tbl = m.group(0)
        tbl_pr = re.search(r"<w:tblPr>.*?</w:tblPr>", tbl, re.S).group(0)
        style = re.search(r'<w:tblStyle w:val="([^"]+)"', tbl_pr)
        where = f"table {i + 1} ({style.group(1) if style else 'unstyled'})"
        usable = next((u for pos, u in sections if pos > m.start()),
                      sections[-1][1])

        if '<w:tblLayout w:type="fixed"/>' not in tbl_pr:
            fails.append(f"{where} is not fixed-layout; Word will autofit it and "
                         f"give every column the same width")

        # The indent belongs to a BORDERED table, whose border is the edge the
        # reader sees and belongs on the text margin. A borderless table's edge is
        # its text, and that belongs where a line of prose begins — so it carries
        # no indent, which is what leaves its text flush with the body.
        margin = _cell_margin(tbl_pr, styles,
                              style.group(1) if style else "TableNormal")
        indent = re.search(r"<w:tblInd [^>]*/>", tbl_pr)
        if style is None:
            if indent is not None and int(
                    re.search(r'w:w="(-?\d+)"', indent.group(0)).group(1)):
                fails.append(f"{where} is indented {indent.group(0)}; a "
                             f"borderless table's text would then start inside "
                             f"the body text edge")
        elif indent is None:
            fails.append(f"{where} carries no indent; its left border will sit in "
                         f"the margin, left of the body text")
        elif margin is None:
            fails.append(f"{where} names a style with no cell margin to match")
        else:
            got = re.search(r'w:w="(-?\d+)"', indent.group(0))
            kind = re.search(r'w:type="(\w+)"', indent.group(0))
            if got is None or int(got.group(1)) != margin or (
                    kind is None or kind.group(1) != "dxa"):
                fails.append(f"{where} is indented {indent.group(0)}, not "
                             f"{margin} dxa — the cell margin its border "
                             f"overhangs by")

        grid = [int(w) for w in re.findall(r'<w:gridCol w:w="(\d+)"/>', tbl)]
        if not grid:
            fails.append(f"{where} declares no column grid")
            continue
        if sum(grid) > usable:
            fails.append(f"{where} spans {sum(grid)} twips of a {usable}-twip "
                         f"text width; it is wider than the page")
        if min(grid) <= 0:
            fails.append(f"{where} has a column of {min(grid)} twips")

        # The table declares the width its own columns come to, not the page's:
        # that is the number Word stops the last column at.
        declared = re.search(r'<w:tblW [^>]*w:w="(-?\d+)"[^>]*w:type="(\w+)"',
                             tbl_pr) or re.search(
            r'<w:tblW [^>]*w:type="(\w+)"[^>]*w:w="(-?\d+)"', tbl_pr)
        if declared is None:
            fails.append(f"{where} declares no width of its own")
        else:
            width, kind = declared.groups()
            if not width.lstrip("-").isdigit():
                width, kind = kind, width
            if int(width) != sum(grid) or kind != "dxa":
                fails.append(f"{where} declares a width of {width} {kind}, not "
                             f"the {sum(grid)} dxa its grid comes to")

        # The width IS the content's. Each column is measured in the document's
        # own face at the size the document states, so a column stretched past
        # what it holds shows up as a column wider than its text, and a table
        # padded out to the margin shows up as every column at once.
        #
        # The signature blocks are the one table that asks for the whole page —
        # two of them belong one at each margin, not side by side in the middle —
        # so that one is required to span it and the rest are required not to.
        columns = _table_cell_texts(tbl)
        size = _table_point_size(tbl)
        if "Prepared by" in tbl:
            if abs(sum(grid) - usable) > 1:
                fails.append(f"the signature blocks span {sum(grid)} twips of a "
                             f"{usable}-twip text width; they belong one at each "
                             f"margin")
        elif len(columns) != len(grid) or size is None:
            fails.append(f"{where}: {len(columns)} columns of text for "
                         f"{len(grid)} grid columns at {size} pt")
        else:
            needs = _content_widths(
                columns, family, size, usable,
                DEFAULT_CELL_MARGIN if margin is None else margin)
            for j, width in enumerate(grid):
                if width > needs[j] + 1:
                    fails.append(f"{where} gives column {j} {width} twips for "
                                 f"content {needs[j]:.0f} wide; the table is "
                                 f"being stretched rather than fitted")
            want = min(int(round(sum(needs))), usable)
            if sum(grid) < want - 1:
                fails.append(f"{where} spans {sum(grid)} twips where its content "
                             f"needs {want}; it is being starved")

            # A NUMBER is never printed on two lines. Word breaks a line after a
            # hyphen-minus, so a column a twip too narrow prints "-" over
            # "1416.5" and a total as "78357." over "0" — which is what the
            # slice table did. Two things have to hold at once: the column is at
            # least as wide as the widest value in it, and every cell holding a
            # number says it may not be broken. The first alone is a measurement
            # that a layout is free to disagree with; the second alone would
            # overflow a column too narrow to hold what it forbids breaking.
            from xslope.report_docx import _text_width as _tw
            from xslope.columns import is_number
            for j, texts in enumerate(columns):
                values = [t for t in texts if is_number(t)]
                if not values or any(t.strip() and not is_number(t)
                                     for t in texts[1:]):
                    continue
                widest = max(_tw(t, family, size) for t in values)
                if grid[j] < widest + 2 * margin - 1:
                    fails.append(f"{where} gives column {j} {grid[j]} twips for "
                                 f"a value {widest + 2 * margin:.0f} wide; the "
                                 f"number will be broken across two lines")
            numbers = sum(1 for texts in columns for t in texts if is_number(t))
            if tbl.count("<w:noWrap/>") != numbers:
                fails.append(f"{where} holds {numbers} numbers but marks "
                             f"{tbl.count('<w:noWrap/>')} cells unbreakable")

        # Word wants the widths in the cells as well as in the grid.
        for row in re.findall(r"<w:tr>.*?</w:tr>", tbl, re.S):
            cells = [int(w) for w in
                     re.findall(r'<w:tcW w:type="dxa" w:w="(\d+)"/>', row)]
            if cells != grid:
                fails.append(f"{where} has a row whose cell widths {cells} do "
                             f"not match its grid {grid}")
                break

        # The header row is a header: bold, and repeated on every page it spans.
        if style is not None:
            head = re.search(r"<w:tr>.*?</w:tr>", tbl, re.S).group(0)
            if "<w:tblHeader" not in head:
                fails.append(f"{where} does not repeat its header row")
            if "<w:b/>" not in head:
                fails.append(f"{where} has no bold in its header row")

    # The columns are measured, not equal: "#" is narrower than "Material", and
    # the long Finding column is wide without taking everything.
    materials = next((t.group(0) for t in tables if "Material" in t.group(0)
                      and "Strength" in t.group(0)), "")
    grid = [int(w) for w in re.findall(r'<w:gridCol w:w="(\d+)"/>', materials)]
    if len(grid) > 2 and not grid[0] < grid[1]:
        fails.append(f"the materials table gives '#' {grid[0]} twips and "
                     f"'Material' {grid[1]}: the columns are not measured")
    # A long text column is wide by wrapping and not by starving the short one
    # beside it. Measured on the reinforcement properties table, whose Label
    # column is short and whose coordinate columns are not.
    lines = next((t.group(0) for t in tables if "Spacing" in t.group(0)
                  and "Label" in t.group(0)), "")
    grid = [int(w) for w in re.findall(r'<w:gridCol w:w="(\d+)"/>', lines)]
    if len(grid) > 2:
        from xslope.report_docx import CELL_MARGIN, TABLE_PT, _text_width
        needed = _text_width("Label", "Calibri", TABLE_PT) + 2 * CELL_MARGIN[0]
        if grid[0] < needed:
            fails.append(f"the Label column is {grid[0]} twips, under the "
                         f"{needed:.0f} its own header needs — a column beside "
                         f"it starved it")
    elif not lines:
        fails.append("the reinforcement properties table was not written")
    return fails


def test_section_breaks_take_no_room():
    """A section break can never manufacture a page of its own.

    Word writes one as a paragraph, and that paragraph is a real empty line at
    the end of the section it closes. It is invisible until the section's content
    ends near the foot of a page — which a table spanning pages does sooner or
    later — and then it has nowhere to go but a fresh page, and the report grows
    a blank landscape sheet between the slice table and the section after it.

    So every paragraph that carries one is made to vanish: the smallest size the
    format expresses, an exact line height to match, no space above or below, and
    a hidden paragraph mark. Checked on all of them, not on the one that showed.
    """
    import re
    fails = []
    from xslope.report import generate_report
    from xslope.report_docx import SECT_BREAK_PT

    slope_data, solutions = _solved()
    with tempfile.TemporaryDirectory() as tmp:
        out_path = os.path.join(tmp, "report.docx")
        ok, out = generate_report(
            slope_data, solutions,
            {"input_path": REINF_XLSX, "method": "spencer", "pd_figure": False,
             "lem_search_figure": False, "lem_solution_figure": False},
            out_path)
        if not ok:
            return [f"generate_report failed: {out}"]
        _names, xml = _docx_parts(out_path)
    doc = xml.get("word/document.xml", "")

    carriers = [p for p in re.findall(r"<w:p\b[^>]*>.*?</w:p>", doc, re.S)
                if "<w:sectPr" in p]
    if not carriers:
        return ["the report opens no section of its own; there is no section "
                "break to check and the landscape slice table is missing"]
    half_points = SECT_BREAK_PT * 2
    for i, p in enumerate(carriers):
        where = f"section break {i + 1} of {len(carriers)}"
        p_pr = re.search(r"<w:pPr>.*?</w:pPr>", p, re.S)
        r_pr = re.search(r"<w:rPr>.*?</w:rPr>", p_pr.group(0) if p_pr else "",
                         re.S)
        if r_pr is None or "<w:vanish/>" not in r_pr.group(0):
            fails.append(f"{where} does not hide its paragraph mark; an empty "
                         f"line will print where it sits")
            continue
        size = re.search(r'<w:sz w:val="(\d+)"/>', r_pr.group(0))
        if size is None or int(size.group(1)) > half_points:
            fails.append(f"{where} sets its paragraph mark at "
                         f"{size.group(1) if size else 'the body size'} "
                         f"half-points, over the {half_points} it may be")
        spacing = re.search(r"<w:spacing [^>]*/>", p_pr.group(0))
        if spacing is None:
            fails.append(f"{where} names no spacing; it keeps the style's")
            continue
        for attr in ("before", "after"):
            got = re.search(rf'w:{attr}="(\d+)"', spacing.group(0))
            if got is None or int(got.group(1)) != 0:
                fails.append(f"{where} leaves {attr}-spacing "
                             f"{got.group(1) if got else 'unset'} on a line "
                             f"that is meant to take no room")
        line = re.search(r'w:line="(\d+)"', spacing.group(0))
        rule = re.search(r'w:lineRule="(\w+)"', spacing.group(0))
        if line is None or rule is None or rule.group(1) != "exact" or \
                int(line.group(1)) > half_points * 10:
            fails.append(f"{where} sets no exact line height; a hidden mark in "
                         f"a body-height line is still a body-height line")
    return fails


#: How far two rows of one table may differ in height and still be one table,
#: in points, and how far the rendered heights are read to. A row is about eleven
#: points tall; a fifth of a point is the rounding on the rules the renderer draws
#: them with, and a mark left at the body size added two and a half.
ROW_HEIGHT_TOLERANCE_PT = 0.5


def _table_rows(doc):
    """``[(table index, row index, row xml), ...]`` for every row of every table
    in a ``document.xml``."""
    out = []
    for t, table in enumerate(re.findall(r"<w:tbl>.*?</w:tbl>", doc, re.S)):
        for r, row in enumerate(re.findall(r"<w:tr\b[^>]*>.*?</w:tr>", table,
                                           re.S)):
            out.append((t + 1, r + 1, row))
    return out


def test_table_rows_are_one_height():
    """Every row of a table is as tall as the next, and its text sits on one line
    across.

    Two things Word does that show the moment a table has an empty cell in it.
    The paragraph MARK — the invisible character at the end of a paragraph —
    carries a size of its own, and a cell written at eight and a half points
    whose mark is left at the document's eleven is laid out for eleven: a row of
    factors of safety with nothing in its Solution parameters column came out a
    fifth taller than the rows either side of it. And a cell hangs its text from
    the TOP of its row, so the shorter cells in a row that one cell has made tall
    float above the line the row is read on.

    So every cell sets its mark at its own size and is vertically centered. Both
    are read back out of the written document, and where the machine can render
    one, the rows of a table with empty cells in it are measured on the page.
    """
    import re as _re
    fails = []
    from xslope.report import generate_report

    slope_data, solutions = _solved()
    with tempfile.TemporaryDirectory() as tmp:
        out_path = os.path.join(tmp, "report.docx")
        with contextlib.redirect_stdout(io.StringIO()):
            ok, out = generate_report(
                slope_data, solutions,
                {"input_path": REINF_XLSX, "title": "Row Heights",
                 "method": ["oms", "bishop", "janbu", "spencer"],
                 "pd_figure": False, "lem_search_figure": False,
                 "lem_slice_key": False, "lem_solution_figure": False},
                out_path)
        if not ok:
            return [f"generate_report failed: {out}"]
        _names, xml = _docx_parts(out_path)
        doc = xml.get("word/document.xml", "")
        rows = _table_rows(doc)
        if len(rows) < 20:
            return [f"only {len(rows)} table rows were written; the report has "
                    f"more"]
        for table, r, row in rows:
            where = f"table {table} row {r}"
            for c, cell in enumerate(_re.findall(r"<w:tc>.*?</w:tc>", row,
                                                 _re.S), start=1):
                if '<w:vAlign w:val="center"/>' not in cell:
                    fails.append(f"{where} cell {c} is not vertically centered; "
                                 f"its text hangs from the top of the row")
                p_pr = _re.search(r"<w:pPr>.*?</w:pPr>", cell, _re.S)
                mark = _re.search(r'<w:rPr>.*?<w:sz w:val="(\d+)"/>',
                                  p_pr.group(0) if p_pr else "", _re.S)
                if mark is None:
                    fails.append(f"{where} cell {c} leaves its paragraph mark at "
                                 f"the document's size; an empty cell will lay "
                                 f"the row out for the body text")
                    continue
                runs = {int(s) for s in _re.findall(
                    r'<w:r>\s*<w:rPr>.*?<w:sz w:val="(\d+)"/>', cell, _re.S)}
                if runs and int(mark.group(1)) not in runs:
                    fails.append(f"{where} cell {c} sets its mark at "
                                 f"{mark.group(1)} half-points and its text at "
                                 f"{sorted(runs)}")

        # --- and on the page ---
        heights, why = _rendered_row_heights(out_path)
        if why:
            print(f"Report: {why} — the rendered row heights were not measured")
        for caption, tall in (heights or {}).items():
            if len(tall) < 3:
                continue
            if max(tall) - min(tall) > ROW_HEIGHT_TOLERANCE_PT:
                fails.append(
                    f"{caption}: its rows come out {min(tall):.2f} to "
                    f"{max(tall):.2f} pt tall on the page, a spread of "
                    f"{max(tall) - min(tall):.2f} pt")
    return fails


def _rendered_row_heights(path):
    """``({caption: [row height, ...]}, why not)`` — the height of every row of
    every single-line table on the rendered page.

    The rules a table is drawn with are the row boundaries, so the heights are
    read off the page itself rather than out of the document. Only tables whose
    rows all hold one line are measured: a Finding that wraps to two lines is
    meant to be taller than the row above it.
    """
    import subprocess
    soffice = _soffice()
    if soffice is None:
        return None, "LibreOffice is not installed"
    try:
        import fitz
    except Exception:
        return None, "PyMuPDF is not installed"
    try:
        subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                        "--outdir", os.path.dirname(path), path],
                       capture_output=True, timeout=300)
    except Exception as exc:
        return None, f"LibreOffice did not render it ({exc})"
    pdf = os.path.splitext(path)[0] + ".pdf"
    if not os.path.exists(pdf):
        return None, "LibreOffice wrote no PDF"
    out = {}
    for page in fitz.open(pdf):
        ys = set()
        for drawing in page.get_drawings():
            for item in drawing["items"]:
                if item[0] == "l" and abs(item[1].y - item[2].y) < 0.3 \
                        and abs(item[1].x - item[2].x) > 50:
                    ys.add(round(item[1].y, 2))
                elif item[0] == "re" and item[1].height < 0.6 \
                        and item[1].width > 50:
                    ys.add(round(item[1].y0, 2))
        ys = sorted(ys)
        spans = [(s["bbox"], s["text"]) for block in page.get_text("dict")["blocks"]
                 for line in block.get("lines", []) for s in line["spans"]]
        captions = [(bbox[1], text) for bbox, text in spans
                    if text.startswith("Table ") and ". " in text]
        for top, bottom in zip(ys, ys[1:]):
            if not 4 < bottom - top < 40:
                continue
            inside = [t for bbox, t in spans
                      if top < (bbox[1] + bbox[3]) / 2 < bottom]
            if not inside:
                continue
            named = [t for y, t in captions if y < top]
            caption = named[-1] if named else f"a table on page {page.number + 1}"
            out.setdefault(caption, []).append(round(bottom - top, 2))
    # A row of two lines is twice a row of one, and its table is not measured.
    return {c: h for c, h in out.items()
            if max(h) < 1.5 * min(h)}, ""


#: Where LibreOffice is looked for. It renders the written document to PDF, and
#: the text drawn on those pages is what a reader can see — the only evidence
#: that a paragraph of the tree reached a page. Without it that leg is skipped
#: and the structural one still runs.
SOFFICE = (os.environ.get("XSLOPE_SOFFICE"), "soffice",
           "/Applications/LibreOffice.app/Contents/MacOS/soffice")

NONCIRC_XLSX = os.path.join(_REPO, "docs", "lem", "files",
                            "xslope_noncircular.xlsx")


def _soffice():
    """LibreOffice, or None on a machine without it."""
    import shutil
    for cand in SOFFICE:
        if cand and (shutil.which(cand) or os.path.isfile(cand)):
            return shutil.which(cand) or cand
    return None


def _rendered_text(path, title):
    """The text of a written report as it comes out on the page, or None when
    this machine cannot render one.

    The document is converted to PDF and the text read off the pages, so a block
    the renderer dropped is missing from the result. The running head and foot go
    with it, and the wrapped lines are rejoined, which puts a paragraph back
    together across a page break it was broken over.
    """
    import subprocess
    soffice = _soffice()
    if soffice is None:
        return None
    try:
        import pypdf
    except Exception:
        return None
    try:
        subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                        "--outdir", os.path.dirname(path), path],
                       capture_output=True, timeout=300)
    except Exception:
        return None
    pdf = os.path.splitext(path)[0] + ".pdf"
    if not os.path.exists(pdf):
        return None
    # The running head is the title and the numbered section the page is in, on
    # one line. It has to go with the foot: extracted, it lands between the two
    # halves of any paragraph that runs over a page break, and the paragraph is
    # then nowhere in the text as it was written.
    head = re.compile(r"^%s\s+\d+(\.\d+)*\s+\S" % re.escape(title)) if title else None
    lines = []
    for page in pypdf.PdfReader(pdf).pages:
        for line in page.extract_text().splitlines():
            line = line.strip()
            if (not line or line == title or re.match(r"^Page \d+ of \d+", line)
                    or (head is not None and head.match(line))):
                continue
            lines.append(line)
    # A cross-reference is a Word FIELD, and the extractor puts a space at every
    # field boundary — "given in Section 3 ." where the page reads "Section 3."
    # No sentence the report writes has a space before its punctuation, so one
    # here is the extraction and not the page.
    return re.sub(r"\s+([.,;:)])", r"\1", re.sub(r"\s+", " ", " ".join(lines)))


def _refused_solutions():
    """A model, and a solution whose working the report has to refuse.

    Janbu's method is solved on a copy of the slices that is then discarded, so
    the report is handed the unsolved frame beside the converged factor of
    safety. The equation evaluated on those values does not return the solution,
    the calculation is refused, and the refusal sentence stands as the last block
    of the landscape section the slice table opens — where a section break
    written onto a paragraph of content swallows it.
    """
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.slice import generate_slices
    from xslope.solve import solve_selected

    slope_data = load_slope_data(NONCIRC_XLSX)
    ok, out = generate_slices(slope_data, non_circ=slope_data["non_circ"],
                              num_slices=15)
    if not ok:
        raise RuntimeError(f"the non-circular model produced no slices: {out}")
    slice_df, surface = out
    solved = slice_df.copy()
    with contextlib.redirect_stdout(io.StringIO()):
        results = solve_selected("janbu", solved)
    bundle = {"slice_df": slice_df.copy(), "failure_surface": surface,
              "results": results, "search": None, "method": "janbu"}
    return slope_data, {"lem": [bundle]}


def _paragraph_text(p):
    """What a paragraph of the document puts on the page — the text of its runs,
    and a marker for a picture, which a text scan would otherwise miss."""
    text = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", p, re.S))
    if "<w:drawing>" in p:
        text += "[picture]"
    return text.strip()


def _collapsed_with_content(doc, where):
    """Failures for every paragraph that is collapsed to a hidden mark and still
    holds something — a section break's carrier, or any other hidden mark."""
    fails = []
    for p in re.findall(r"<w:p\b[^>]*>.*?</w:p>", doc, re.S):
        p_pr = re.search(r"<w:pPr>.*?</w:pPr>", p, re.S)
        if p_pr is None:
            continue
        if "<w:sectPr" in p_pr.group(0):
            why = "carries a section break"
        elif "<w:vanish/>" in p_pr.group(0):
            why = "is marked hidden"
        else:
            continue
        text = _paragraph_text(p)
        if text:
            fails.append(f"{where}: a paragraph that {why} holds {text[:90]!r}. "
                         f"Collapsed to a hidden mark at an exact line height of "
                         f"a point, that never prints")
    return fails


def _page_pattern(text):
    """One paragraph as a pattern for the text of the page it is set on.

    A symbol the prose names is SET as a symbol — N_S reaches the page as an N
    with a subscript S, which is the whole point of it — so the underscore and the
    braces are not on the page and the extractor may or may not put a space where
    the script starts. A hyphenated word broken over a line end comes back with a
    space after the hyphen ("out-of- plane"), which is the line break and not the
    page. Everything else has to be there, character for character.
    """
    from xslope.report_docx import INLINE_MATH

    out, at = [], 0
    for found in INLINE_MATH.finditer(text):
        out.append(re.escape(text[at:found.start()]))
        base, _sep, script = found.group(0).partition("_")
        out.append(re.escape(base) + r"\s*" + re.escape(script.strip("{}")))
        at = found.end()
    out.append(re.escape(text[at:]))
    return "".join(out).replace(r"\-", r"\-\s?")


def _missing_from_the_page(report, text, where):
    """Failures for every prose block of the tree that is not in the rendered
    text."""
    fails = []
    for block in report.blocks("prose"):
        want = re.sub(r"\s+", " ", block.text).strip()
        if want and not re.search(_page_pattern(want), text):
            fails.append(f"{where}: a paragraph of the report is not on a page "
                         f"of it: {block.text!r}")
    return fails


def _written(where, slope_data, solutions, opts, tmp):
    """``(report, document.xml, path)`` for one report written to ``tmp``."""
    from xslope.report import generate_report
    path = os.path.join(tmp, "report.docx")
    with contextlib.redirect_stdout(io.StringIO()):
        ok, out = generate_report(slope_data, solutions, dict(opts), path)
    if not ok:
        raise RuntimeError(f"{where}: generate_report failed: {out}")
    _names, xml = _docx_parts(path)
    return out["report"], xml.get("word/document.xml", ""), path


#: The reports this check writes and reads back: the sample model, and the one
#: whose landscape section ends in a sentence.
_PAGE_REPORTS = (
    ("the sample report", REINF_XLSX,
     {"title": "Sample Levee", "method": "spencer"}),
    ("a refused calculation", NONCIRC_XLSX,
     {"title": "Non-circular Surface", "method": "janbu"}),
)


def test_every_block_reaches_the_page():
    """Every paragraph of the report is on a page of the report.

    A block can be in the tree, be written into the document, and still not reach
    a reader. The paragraph that carries a closing section break is collapsed to
    a hidden mark at an exact line height of a point, and a sentence set in a
    paragraph like that is cropped to nothing — which is what becomes of the
    sentence a refused calculation prints when it is the last block of the
    landscape section the slice table opens.

    Two things are required of every report written here. Nothing is collapsed
    but an empty paragraph: a section break rides its own, never a sentence's.
    And every prose block of the tree is found in the text of the rendered pages,
    so a block dropped or hidden by any path through the renderer is caught by
    what came out rather than by what was intended.
    """
    fails = []
    rendered = 0
    for where, xlsx, extra in _PAGE_REPORTS:
        slope_data, solutions = (_solved() if xlsx is REINF_XLSX
                                 else _refused_solutions())
        opts = {"input_path": xlsx, "pd_figure": False,
                "lem_search_figure": False, "lem_solution_figure": False}
        opts.update(extra)
        with tempfile.TemporaryDirectory() as tmp:
            report, doc, path = _written(where, slope_data, solutions, opts, tmp)
            fails += _collapsed_with_content(doc, where)
            text = _rendered_text(path, opts["title"])
            if text is None:
                continue
            rendered += 1
            fails += _missing_from_the_page(report, text, where)
    if not rendered:
        print("Report: LibreOffice not installed — the rendered pages were not "
              "read back.")

    # Mutation. The defect this check is built on is the section break's
    # treatment reaching a paragraph of content: put it back, on the report whose
    # landscape section ends in a sentence, and the sentence has to be reported
    # missing.
    import xslope.report_docx as report_docx
    saved = report_docx._sect_break_carrier
    report_docx._sect_break_carrier = lambda doc: next(
        (p for p in reversed(doc.paragraphs) if p.text.strip()), None)
    try:
        slope_data, solutions = _refused_solutions()
        opts = {"input_path": NONCIRC_XLSX, "title": "Non-circular Surface",
                "method": "janbu", "pd_figure": False,
                "lem_search_figure": False, "lem_solution_figure": False}
        with tempfile.TemporaryDirectory() as tmp:
            _report, doc, _path = _written("the mutation", slope_data,
                                           solutions, opts, tmp)
        if not _collapsed_with_content(doc, "the mutation"):
            fails.append("a sentence was collapsed to a hidden mark and this "
                         "check passed the document")
    finally:
        report_docx._sect_break_carrier = saved
    return fails


def test_contents_page():
    """The contents page lists the report's own headings, from generation.

    The field is never marked dirty — that flag is what makes Word for Mac ask
    about fields on every open — so nothing updates it before a reader sees it,
    and what is cached in it has to be the contents themselves. They come from
    the same tree the document is written from, which is what makes a section
    that was toggled off absent from both.
    """
    import re
    fails = []
    from xslope.report import generate_report
    from xslope.report_docx import TOC_LEVELS

    slope_data, solutions = _solved()
    opts = {"input_path": REINF_XLSX, "title": "Sample Levee",
            "method": "spencer", "pd_figure": False,
            "lem_search_figure": False, "lem_solution_figure": False}
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "report.docx")
        ok, out = generate_report(slope_data, solutions, dict(opts), path)
        if not ok:
            return [f"generate_report failed: {out}"]
        _names, xml = _docx_parts(path)
        doc = xml.get("word/document.xml", "")
        entries = _toc_result(doc)
        expected = [f"{number} {sec.title}"
                    for number, lvl, sec in out["report"].section_numbers()
                    if lvl <= TOC_LEVELS]

        if entries[:-1] != expected:
            fails.append(f"the contents list {entries[:-1]}, not the report's "
                         f"headings {expected}")
        hint = entries[-1] if entries else ""
        if "Update Field" not in hint or "F9" not in hint:
            fails.append(f"the last thing in the field result is {hint!r}, not "
                         f"the line about updating it")
        if hint in doc.split('w:fldCharType="end"')[-1]:
            fails.append("the update hint sits outside the field result; Word "
                         "would leave it behind when the table is updated")
        if "Page numbers" not in hint:
            fails.append("the hint does not say what an update adds")
        # Levels are levels: a sub-section is indented, in the template's style.
        styles = re.findall(r'<w:pStyle w:val="(TOC\d)"/>', doc)
        levels = [lvl for lvl, _t in out["report"].section_titles()
                  if lvl <= TOC_LEVELS]
        if styles != [f"TOC{lvl}" for lvl in levels]:
            fails.append(f"the contents are styled {styles}, not one TOC style "
                         f"per heading level {levels}")
        if "<w:t>1</w:t>" in doc[doc.index("TOC \\o"):doc.index("Traceability")]:
            fails.append("the cached contents carry a page number; the pagination "
                         "is Word's to compute")

    # A section switched off is in neither the document nor its contents.
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "no_slices.docx")
        ok, out = generate_report(slope_data, solutions,
                                  dict(opts, lem_slice_table=False), path)
        _names, xml = _docx_parts(path)
        entries = _toc_result(xml.get("word/document.xml", ""))
        if "Slice Table" in entries:
            fails.append("the contents list a Slice Table section the report was "
                         "told not to write")
    return fails


def _headings_in(doc_xml):
    """``[(style_id, text), ...]`` for every heading paragraph of a document."""
    import re
    out = []
    for para in re.findall(r"<w:p[ >].*?</w:p>", doc_xml, re.S):
        style = re.search(r'<w:pStyle w:val="(Heading\d)"/>', para)
        if style:
            text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", para))
            out.append((style.group(1), text))
    return out


def test_heading_numbers():
    """The headings are numbered 1, 1.1, 1.1.1 by Word's own multilevel list.

    Numbering that is typed into the heading strings is numbering that goes wrong
    the first time a section is switched off. The document instead carries the
    numbering definition, bound to the heading styles, that Word writes when a
    user applies a multilevel list from the ribbon — so Word computes every
    number, the contents field picks them up, and a cross-reference to a section
    resolves to whatever its number has become.
    """
    import re
    fails = []
    from xslope.report import HEADING_LEVELS, generate_report
    from xslope.report_docx import HEADING_NUM_ID

    slope_data, solutions = _solved()
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "numbered.docx")
        ok, out = generate_report(
            slope_data, solutions,
            {"input_path": REINF_XLSX, "method": "spencer", "pd_figure": False,
             "lem_search_figure": False, "lem_solution_figure": False}, path)
        if not ok:
            return [f"generate_report failed: {out}"]
        _names, xml = _docx_parts(path)
        doc = xml.get("word/document.xml", "")
        numbering = xml.get("word/numbering.xml", "")
        styles = xml.get("word/styles.xml", "")

        # The definition: one level per heading depth, each naming its style and
        # printing the whole path to it.
        if not numbering:
            return ["the package carries no word/numbering.xml, so nothing "
                    "numbers the headings"]
        block = re.search(rf'<w:abstractNum w:abstractNumId="{HEADING_NUM_ID}".*?'
                          rf'</w:abstractNum>', numbering, re.S)
        if block is None:
            return [f"the numbering part defines no list "
                    f"{HEADING_NUM_ID}: {numbering[:400]!r}"]
        text = block.group(0)
        want = [".".join(f"%{i + 1}" for i in range(lvl + 1))
                for lvl in range(HEADING_LEVELS)]
        got = re.findall(r'<w:lvlText w:val="([^"]*)"/>', text)
        if got != want:
            fails.append(f"the heading list prints {got}, not {want}")
        named = re.findall(r'<w:pStyle w:val="([^"]*)"/>', text)
        if named != [f"Heading{i + 1}" for i in range(HEADING_LEVELS)]:
            fails.append(f"the heading list's levels name {named}, not the "
                         f"heading styles")
        if f'<w:num w:numId="{HEADING_NUM_ID}"' not in numbering:
            fails.append(f"list {HEADING_NUM_ID} is defined but never made a "
                         f"numbering a style can point at")

        # And the binding: every heading style carries it, at its own level.
        for level in range(1, HEADING_LEVELS + 1):
            style = re.search(rf'<w:style [^>]*w:styleId="Heading{level}".*?'
                              rf'</w:style>', styles, re.S)
            if style is None:
                fails.append(f"the template defines no Heading{level} style")
                continue
            num_pr = re.search(r"<w:numPr>.*?</w:numPr>", style.group(0), re.S)
            if num_pr is None:
                fails.append(f"the Heading{level} style carries no numbering, so "
                             f"its headings print unnumbered")
                continue
            ids = re.findall(r'<w:numId w:val="(\d+)"/>', num_pr.group(0))
            ilvl = re.findall(r'<w:ilvl w:val="(\d+)"/>', num_pr.group(0))
            if ids != [str(HEADING_NUM_ID)]:
                fails.append(f"the Heading{level} style is numbered by list "
                             f"{ids}, not the heading list")
            if ilvl != [str(level - 1)]:
                fails.append(f"the Heading{level} style is bound to level {ilvl}, "
                             f"not level {level - 1} of the heading list")

        # Nothing is numbered twice: no heading string starts with a number.
        headings = _headings_in(doc)
        if not headings:
            fails.append("the document carries no heading paragraphs")
        for style_id, text in headings:
            if re.match(r"^\s*\d+(\.\d+)*[.\s)]", text):
                fails.append(f"the {style_id} {text!r} carries its own number; "
                             f"Word puts one in front of it as well")

        # The headings in the document are the report's own titles, and the
        # numbers the report caches for them are what this tree numbers to.
        report = out["report"]
        titles = [t for lvl, t in report.section_titles()
                  if lvl <= HEADING_LEVELS]
        if [t for _s, t in headings] != titles:
            fails.append(f"the document's headings are {[t for _s, t in headings]}, "
                         f"not the report's {titles}")
        numbers = report.section_numbers()
        top = [n for n, lvl, _s in numbers if lvl == 1]
        if top != [str(i + 1) for i in range(len(top))]:
            fails.append(f"the top-level sections number {top}")
        seen = {}
        for number, _lvl, _sec in numbers:
            seen[number] = seen.get(number, 0) + 1
            parent = number.rsplit(".", 1)[0]
            if "." in number and parent not in seen:
                fails.append(f"section {number} numbers under {parent}, which "
                             f"the report does not have")
        repeated = [n for n, count in seen.items() if count > 1]
        if repeated:
            fails.append(f"two sections are numbered the same: {repeated}")

        # The contents page carries them too, so a document nobody has pressed
        # F9 in already reads as a numbered report.
        entries = _toc_result(doc)
        if not entries or not entries[0].startswith("1 "):
            fails.append(f"the contents page opens on {entries[:1]!r}, which "
                         f"carries no heading number")
    return fails


def test_report_writes_one_file():
    """A report is a document, not a document and a folder.

    The figures are embedded in the .docx, so the PNGs are rendered into a
    temporary directory and go with it; a report generated with the figures
    switched off creates no directory at all.
    """
    fails = []
    from xslope.report import generate_report

    slope_data, solutions = _solved()
    opts = {"input_path": REINF_XLSX, "title": "Sample Levee",
            "method": "spencer"}
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "report.docx")
        ok, out = generate_report(slope_data, solutions, dict(opts), path)
        if not ok:
            return [f"generate_report failed: {out}"]
        left = sorted(os.listdir(tmp))
        if left != ["report.docx"]:
            fails.append(f"a default report left {left} behind, not the document "
                         f"alone")
        if not out["figures"]:
            fails.append("the result names none of the figures the document "
                         "carries")
        captions = [f.caption for f in out["report"].figures()]
        if list(out["figures"]) != captions:
            fails.append(f"the result reports {out['figures']}, not the "
                         f"captions {captions}")

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "nofig.docx")
        ok, _out = generate_report(
            slope_data, solutions,
            dict(opts, pd_figure=False, lem_search_figure=False,
                 lem_solution_figure=False), path)
        left = sorted(os.listdir(tmp))
        if not ok or left != ["nofig.docx"]:
            fails.append(f"a report with the figures off left {left} behind")

    # A caller that asks for the PNGs still gets them, in the directory it named.
    with tempfile.TemporaryDirectory() as tmp:
        figures = os.path.join(tmp, "figures")
        ok, _out = generate_report(slope_data, solutions, dict(opts),
                                   os.path.join(tmp, "kept.docx"),
                                   figure_dir=figures)
        kept = sorted(os.listdir(figures)) if os.path.isdir(figures) else []
        if not ok or not [f for f in kept if f.endswith(".png")]:
            fails.append(f"an explicit figure_dir kept {kept}")
    return fails


def test_docx_template():
    """The shipped template exists, and the document is built on it."""
    fails = []
    from xslope.report_docx import DEFAULT_TEMPLATE

    if not os.path.exists(DEFAULT_TEMPLATE):
        return [f"the default template is missing: {DEFAULT_TEMPLATE}"]
    _names, xml = _docx_parts(DEFAULT_TEMPLATE)
    styles = xml.get("word/styles.xml", "")
    for style in ("Heading1", "Title", "Caption", "TableGrid"):
        if style not in styles:
            fails.append(f"the template defines no {style} style")

    # The title page's look is the TEMPLATE's: ranged left, and closed by the rule
    # the Title style carries. A company template put in its place decides both,
    # which is only true while the renderer does not draw them itself.
    import re
    title_style = re.search(r'<w:style [^>]*w:styleId="Title".*?</w:style>',
                            styles, re.S)
    if title_style is None:
        fails.append("the template has no Title style to read")
    else:
        block = title_style.group(0)
        if "w:pBdr" not in block:
            fails.append("the Title style carries no rule under the title")
        if 'w:jc w:val="left"' not in block:
            fails.append("the Title style is not ranged left")
    body = xml.get("word/document.xml", "")
    if "<w:p " in body or "<w:p>" in body:
        fails.append("the template carries body content; it must ship empty")

    # The template is reproducible from its build script.
    import importlib.util
    script = os.path.join(_REPO, "tools", "build_report_template.py")
    if not os.path.exists(script):
        fails.append("tools/build_report_template.py is missing — the template "
                     "would be an unreproducible binary")
    else:
        spec = importlib.util.spec_from_file_location("build_report_template", script)
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
        with tempfile.TemporaryDirectory() as tmp:
            rebuilt = mod.build(os.path.join(tmp, "t.docx"))
            _n, rebuilt_xml = _docx_parts(rebuilt)
            if rebuilt_xml.get("word/styles.xml") != styles:
                fails.append("rebuilding the template produces different styles "
                             "from the committed one")
    return fails


# --------------------------------------------------------------------------
# E. the column registry
# --------------------------------------------------------------------------

def test_column_registry():
    """Every report-worthy column is a column a solved model actually has."""
    fails = []
    from xslope import columns as cols

    _slope_data, solutions = _solved()
    slice_df = solutions["lem"][0]["slice_df"]
    have = set(slice_df.columns)
    for column in cols.solver_columns():
        if column.key not in have:
            fails.append(f"report column {column.key!r} is not in slice_df")
    for column in cols.report_columns():
        if not column.description.strip().endswith("."):
            fails.append(f"column {column.key!r} has no complete description")
        if not column.label.strip():
            fails.append(f"column {column.key!r} has no label")
    keys = [c.key for c in cols.SLICE_COLUMNS]
    if len(keys) != len(set(keys)):
        fails.append("the registry declares a column twice")

    # The units reach the headers, and only where they belong.
    from xslope.units import labels
    lbl = labels("imperial")
    headers, rows, legend = cols.slice_table(slice_df, lbl)
    if "W (lb/ft)" not in headers:
        fails.append(f"the weight column is not unit-labeled: {headers}")
    if "α (deg)" not in headers:
        fails.append(f"the base-angle column is not in degrees: {headers}")
    # A moment per unit width of section is spelled one way in this report, and
    # it is the way the pile capacities are already spelled
    # (:func:`xslope.fem_details.unit_labels`): force x length per length, which
    # reads as a moment. Cancelled to "lb/ft·ft" it reads as force per area.
    from xslope.fem_details import unit_labels as member_unit_labels
    for system in ("imperial", "si"):
        want = member_unit_labels({"unit_system": system})["moment"]
        for key in ("m_res", "m_drv"):
            got = cols.unit_label(cols.BY_KEY[key], labels(system))
            if got != want:
                fails.append(f"{system}: the slice table's {key} is in {got!r} "
                             f"and the member tables' moments in {want!r}")
    if "Slice" not in headers:
        fails.append(f"the slice number column is missing: {headers}")
    if any(h.startswith("Slice (") for h in headers):
        fails.append("the slice number carries a unit")
    if len(rows) != len(slice_df):
        fails.append(f"the table has {len(rows)} rows for {len(slice_df)} slices")
    if len(legend) != len(headers):
        fails.append(f"{len(legend)} legend entries for {len(headers)} columns")

    # An undeclared unit system leaves the headers bare, exactly as the plots do.
    bare, _rows, _legend = cols.slice_table(slice_df, None)
    if any("(lb/ft)" in h for h in bare):
        fails.append("an undeclared model got unit labels anyway")
    if "α (deg)" not in bare:
        fails.append("degrees are not a declared unit and must always be shown")

    # drop_empty removes a column of zeros, and never a geometry one.
    kept = {h.split(" (")[0] for h in headers}
    if "kW" in kept:
        fails.append("a seismic column was printed for a model with no seismic load")
    for needed in ("Δx", "α", "W", "c", "φ"):
        if needed not in kept:
            fails.append(f"drop_empty removed {needed}, which is never optional")
    fails += _the_load_and_its_angle_travel_together(slice_df)
    return fails


def _the_load_and_its_angle_travel_together(slice_df):
    """A distributed load and the angle it acts at are printed together or not
    at all.

    β is the inclination of the slice's top edge. The slicer computes it from the
    geometry for every slice, whether or not a load is applied there, so it is
    NOT zero on a model that carries no distributed load — and a column dropped
    for being all zeros is the wrong test for it. A model loaded by a line load
    alone printed β = 84.55° in a table whose D column had been dropped: an
    inclination of a force the model does not have, with nothing on the page to
    read it against.

    Run on the columns rather than on a sample model, so the pair is checked in
    both states no matter which distributed loads the corpus happens to carry.
    """
    import numpy as np
    from xslope import columns as cols

    fails = []
    if cols.BY_KEY["beta"].gated_by != "dload":
        fails.append(f"β is kept on its own values, and the angle is stored "
                     f"whether or not a load acts: "
                     f"gated_by={cols.BY_KEY['beta'].gated_by!r}")

    df = slice_df.copy()
    df["dload"] = 0.0
    df["beta"] = np.linspace(10.0, 84.55, len(df))
    kept = {c.key for c in cols.selected_columns(df)}
    if "beta" in kept:
        fails.append("β is printed on a model whose D column is zero on every "
                     "slice")
    if "dload" in kept:
        fails.append("a D column of zeros was printed")

    df["dload"] = np.linspace(1.0, 50.0, len(df))
    kept = {c.key for c in cols.selected_columns(df)}
    for key in ("dload", "beta"):
        if key not in kept:
            fails.append(f"{key} is missing from a model that carries a "
                         f"distributed load")

    # And the gate never reaches a column that has none: every other column is
    # still judged on its own values.
    df["kw"] = 0.0
    if "kw" in {c.key for c in cols.selected_columns(df)}:
        fails.append("an ungated column of zeros survived the gate")
    return fails


# --------------------------------------------------------------------------
# I. the calculations
#
# The section's whole claim is that the printed numbers ARE the solution. So the
# checks do what a reviewer would: read the numbers off the page, divide them,
# and see whether the factor of safety comes back. Nothing here reads an
# internal float — every operand is parsed out of the string the document
# carries, and the reproduction is asserted at that printed precision.
# --------------------------------------------------------------------------

#: The methods a calculation is checked on. The three the ruling names, plus the
#: rest of the solver's methods, because each writes a different equation and a
#: silent failure on one of them is a section that quietly disappears.
CALC_METHODS = ("oms", "bishop", "spencer", "janbu", "corps", "lowe", "mprice")

#: The methods whose factor of safety IS a quotient of two sums, and which are
#: therefore checked by dividing the printed operands and getting F back.
#:
#: Spencer's is not one of them. Its F and θ are the pair at which two
#: equilibrium equations both vanish, reached by iterating the two together, and
#: the ratio of two force sums over that mass is a general horizontal balance
#: that any solution satisfies — printing it as "the equation" said something
#: true about the answer and nothing at all about the method. Spencer is verified
#: instead on its two equations balancing
#: (:func:`test_calculation_residuals`) and on reproducing each row's Q
#: (:func:`test_spencer_force_sums`).
QUOTIENT_METHODS = tuple(m for m in CALC_METHODS if m != "spencer")

#: How Spencer's section writes an equilibrium equation, transcribed from the
#: derivation — ``R_1 = sum{Q}``. The number the derivation gives it is in the
#: sentence above it, not on the equation.
SPENCER_EQUATION = r"^R_([12]) = (sum\{.+\})$"

#: And how it writes that same equation EVALUATED at the pair the iteration
#: reached: the equation restated, then the number it comes out at. Restated,
#: because the close is the section's demonstration that the equations balance —
#: a bare ``R_2 = 9.136e-05`` a page below the equation it belongs to is a digit
#: the reader has to go back and pair up for themselves.
SPENCER_EVALUATION = r"^R_([12]) = (sum\{.+\}) = (-?\d[^ ]*)$"

#: And the third thing the section prints: equation (1) or (2) reduced to the
#: terms this model carries. No page publishes it — the transcribed pair above it
#: is the derivation's, this is what is left of them here.
SPENCER_REDUCED_SUM = r"^F_[hv] = (0|[^=]+)$"

_CALC = {}


def _calc_report(method, options=None, xlsx=None):
    """A report of the sample model solved by ``method``, and its solved bundle.

    Cached per method: each of these is a full solve, and several checks read the
    same one. ``xlsx`` runs another model through the same path, for a check that
    needs a section built on a different shape of slope.
    """
    xlsx = xlsx or REINF_XLSX
    key = (method, tuple(sorted((options or {}).items())), xlsx)
    if key in _CALC:
        return _CALC[key]
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.report import build_report
    from xslope.slice import generate_slices
    from xslope.solve import solve_selected

    slope_data = load_slope_data(xlsx)
    # A model carries circles or a non-circular surface, and the corpus models
    # the tolerance check reads carry both kinds.
    circles = slope_data.get("circles") or []
    surface_kw = ({"circle": circles[0]} if circles
                  else {"non_circ": slope_data.get("non_circ")})
    ok, out = generate_slices(slope_data, num_slices=15, **surface_kw)
    if not ok:
        raise RuntimeError(f"the sample model produced no slices: {out}")
    df, surface = out[0].copy(), out[1]
    with contextlib.redirect_stdout(io.StringIO()):
        results = solve_selected(method, df)
    if not isinstance(results, dict):
        _CALC[key] = (None, None)
        return _CALC[key]
    bundle = {"slice_df": df, "failure_surface": surface, "results": results,
              "search": None, "method": method}
    opts = {"method": method, "pd_figure": False, "lem_search_figure": False,
            "lem_solution_figure": False}
    opts.update(options or {})
    with tempfile.TemporaryDirectory() as tmp:
        report = build_report(slope_data, {"lem": [bundle]}, opts, tmp)
    _CALC[key] = (report, bundle)
    return _CALC[key]


def _calc_section(report):
    """The Calculations section of a report, or None."""
    for section in report.sections:
        for _lvl, node in section.walk():
            if node.title == "Calculations":
                return node
    return None


def _per_method_report(specs, options=None):
    """``(report, [(method, slice count), ...])`` for a report of several
    methods, each solved on its own slice frame.

    A search per method finds each its own critical surface, and each surface is
    cut into its own number of slices; ``specs`` is ``[(method, num_slices),
    ...]`` and reproduces that difference without running a search per method.
    The counts come back as the frames actually came out, not as they were asked
    for.
    """
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.report import build_report
    from xslope.slice import generate_slices
    from xslope.solve import solve_selected

    slope_data = load_slope_data(REINF_XLSX)
    circle = (slope_data.get("circles") or [None])[0]
    bundles, counts = [], []
    for method, num in specs:
        ok, out = generate_slices(slope_data, num_slices=num, circle=circle)
        if not ok:
            raise RuntimeError(f"the sample model produced no slices: {out}")
        df, surface = out[0].copy(), out[1]
        with contextlib.redirect_stdout(io.StringIO()):
            results = solve_selected(method, df)
        if not isinstance(results, dict):
            raise RuntimeError(f"the sample model did not solve with {method}")
        bundles.append({"slice_df": df, "failure_surface": surface,
                        "results": results, "search": None, "method": method})
        counts.append((method, len(df)))
    opts = {"method": [m for m, _n in specs], "pd_figure": False,
            "lem_search_figure": False, "lem_solution_figure": False,
            "lem_inputs_figure": False}
    opts.update(options or {})
    with tempfile.TemporaryDirectory() as tmp:
        return build_report(slope_data, {"lem": bundles}, opts, tmp), counts


def _analysis_inputs(report):
    """The engine-input rows of the limit equilibrium section's Analysis Inputs,
    as ``[(label, value), ...]``."""
    for section in report.sections:
        if section.title != "Limit Equilibrium Analysis":
            continue
        for _lvl, node in section.walk():
            if node.title == "Analysis Inputs":
                for block in node.blocks:
                    if block.kind == "keyvalues":
                        return list(block.items)
    return []


def _method_section(report, method):
    """One method's detail section, by the heading it is written under."""
    from xslope.report import method_label
    label = method_label(method)
    for section in report.sections:
        for _lvl, node in section.walk():
            if node.title == label:
                return node
    return None


def _numbers(text):
    """Every number in a string, in order, as they are printed."""
    import re
    return [float(m) for m in re.findall(r"-?\d+\.?\d*(?:e[+-]?\d+)?", text)]


def _operands(section):
    """The printed operands of the factor of safety, read off the equations.

    Returns ``(numerator, denominator, quotient, corrected)`` as STRINGS, exactly
    as the document prints them — the corrected pair is Janbu's f_o line and is
    ``(None, None)`` for every other method.

    Two shapes carry the arithmetic. The force methods print the division and the
    answer on lines of their own. The two moment methods print their page's own
    quotient with the numbers substituted into it and the answer on the end of the
    same line — ``F = N_S/(D_W + D_D) = 56676.4/31859 = 1.779`` — because the
    letters and the numbers are the same statement and separating them made the
    reader carry the letters down the page.
    """
    import re
    num = den = quotient = corrected = factor = None
    for block in section.blocks:
        if block.kind != "math":
            continue
        text = block.notation
        chain = re.match(r"^F = frac\{[^{}]+\}\{[^{}]+\} = "
                         r"frac\{([\d.eE+-]+)\}\{([\d.eE+-]+)\} = ([\d.]+)$",
                         text)
        if chain:
            num, den, quotient = chain.group(1), chain.group(2), chain.group(3)
            continue
        frac = re.match(r"^F = frac\{([^{}]+)\}\{([^{}]+)\}$", text)
        if frac:
            num, den = frac.group(1), frac.group(2)
            continue
        plain = re.match(r"^F = ([\d.]+)$", text)
        if plain:
            quotient = plain.group(1)
            continue
        corr = re.match(r"^F_corr = f_o·F = ([\d.]+)·([\d.]+) = ([\d.]+)$", text)
        if corr:
            factor, quotient, corrected = corr.group(1), corr.group(2), corr.group(3)
    return num, den, quotient, factor, corrected


def _reproduces(num, den, quotient, factor, corrected, fs):
    """Does the factor of safety come back out of the printed operands?

    The tolerance is one unit in the last digit the factor of safety is printed
    to — the strictest statement that can be made about a rounded number.
    """
    from xslope.columns import format_fs

    tolerance = 10 ** -len(format_fs(1.0).split(".")[-1])
    if num is None or den is None or quotient is None:
        return False, "the equations do not print a quotient"
    computed = float(num) / float(den)
    if abs(computed - float(quotient)) > tolerance:
        return False, (f"{num}/{den} = {computed:.6f}, printed as {quotient}")
    if corrected is not None:
        product = float(factor) * float(quotient)
        if abs(product - float(corrected)) > tolerance:
            return False, (f"{factor}x{quotient} = {product:.6f}, printed as "
                           f"{corrected}")
        if abs(product - fs) > tolerance:
            return False, (f"the printed operands give {product:.6f}, the solver "
                           f"{fs:.6f}")
        return True, ""
    if abs(computed - fs) > tolerance:
        return False, (f"the printed operands give {computed:.6f}, the solver "
                       f"{fs:.6f}")
    return True, ""


def test_force_term_registry():
    """Every force is in every equation, or the registry says why it is not.

    :data:`xslope.report.FORCE_TERMS` is the one declaration of the forces a
    slice can carry, and the equations the section prints — the two moment sums,
    the horizontal balance, Spencer's two force sums and the base-normal
    equation — are assembled from it. A force carried into some of them and
    forgotten in the rest is what prints a term in one equation and denies it in
    the next, so every entry has to hold either a contribution or a stated reason
    for each of them.

    And the contributions are all evaluated on a solved model: a term that names
    an array or a column nothing writes would never appear, and would never be
    missed, because a term with no values is a term that reads as absent. That
    includes the terms a stated reason carries as PUBLISHED — the shear on the
    top of the slice, which Spencer's section transcribes and no model exercises.
    """
    fails = []
    import numpy as np
    from xslope.report import (CONSUMERS, EQUATION_SYMBOLS, FORCE_TERMS,
                               ForceTerm, NotApplicable, PASSIVE_COLUMNS,
                               SECTION_SYMBOL_GROUPS, SYMBOL_GROUPS, Term,
                               _Calc, _calc_arrays)

    _slope_data, solutions = _solved()
    df = solutions["lem"][0]["slice_df"]
    A = _calc_arrays(df)
    C = _Calc(df, A, right_facing=False)
    have = set(df.columns)

    keys = [t.key for t in FORCE_TERMS]
    if len(set(keys)) != len(keys):
        fails.append(f"two entries share a key: {keys}")

    for term in FORCE_TERMS:
        for column in term.columns + tuple(c for _n, c, _d in term.arrays):
            if column not in have:
                fails.append(f"{term.key}: names column {column!r}, which a "
                             f"solved slice table does not carry")
        for consumer in CONSUMERS:
            got = getattr(term, consumer)
            if isinstance(got, NotApplicable):
                if not got.reason.strip():
                    fails.append(f"{term.key}: {consumer} is not applicable for "
                                 f"no stated reason")
                # A published term of a force the solver does not carry is held
                # against the same rules as any other, and has to be zero on
                # every slice: it is printed in the transcription and summed
                # nowhere, and a nonzero one is a term gone missing from the
                # arithmetic.
                for contribution in got.published:
                    values = np.asarray(contribution.values(C), dtype=float)
                    if values.shape != (len(df),):
                        fails.append(f"{term.key}: {consumer} publishes "
                                     f"{contribution.symbol!r}, which gives "
                                     f"{values.shape}, not one value per slice")
                    elif values.any():
                        fails.append(f"{term.key}: {consumer} publishes "
                                     f"{contribution.symbol!r} as unexercised, "
                                     f"and it is not zero on every slice")
                continue
            if not got:
                fails.append(f"{term.key}: {consumer} carries no term and no "
                             f"reason for carrying none")
                continue
            for contribution in got:
                if not isinstance(contribution, Term):
                    fails.append(f"{term.key}: {consumer} holds "
                                 f"{contribution!r}, which is not a term")
                    continue
                if contribution.sign not in (+1, -1):
                    fails.append(f"{term.key}: {consumer} term "
                                 f"{contribution.symbol!r} has sign "
                                 f"{contribution.sign}")
                try:
                    values = np.asarray(contribution.values(C), dtype=float)
                except Exception as exc:
                    fails.append(f"{term.key}: {consumer} term "
                                 f"{contribution.symbol!r} does not evaluate: "
                                 f"{exc!r}")
                    continue
                if values.shape != (len(df),):
                    fails.append(f"{term.key}: {consumer} term "
                                 f"{contribution.symbol!r} gives "
                                 f"{values.shape}, not one value per slice")
        for symbol in term.symbols:
            if symbol.group in SECTION_SYMBOL_GROUPS:
                # A letter one derivation gives a force the others give another.
                # It is defined for its own section and must NOT be the report's
                # definition, which is the whole reason it is held apart: a
                # section symbol that agrees with the nomenclature is a row the
                # nomenclature already had.
                if EQUATION_SYMBOLS.get(symbol.name) == symbol.meaning:
                    fails.append(f"{term.key}: symbol {symbol.name!r} is held "
                                 f"apart for one section and is what the "
                                 f"nomenclature already defines")
                continue
            if symbol.group not in SYMBOL_GROUPS:
                fails.append(f"{term.key}: symbol {symbol.name!r} is in group "
                             f"{symbol.group!r}, which the nomenclature has no "
                             f"place for")
            if EQUATION_SYMBOLS.get(symbol.name) != symbol.meaning:
                fails.append(f"{term.key}: symbol {symbol.name!r} is not the "
                             f"one the nomenclature defines")

    passive = tuple(c for t in FORCE_TERMS if t.passive for c in t.columns)
    if passive != PASSIVE_COLUMNS:
        fails.append(f"the passive gate is {PASSIVE_COLUMNS}, and the registry's "
                     f"passive entries carry {passive}")
    if not passive:
        fails.append("no entry is marked passive, so the gate tests nothing")

    # Mutation: a force added to some equations and not the rest has to be
    # impossible to write. Every consumer is a required field, so the half-added
    # entry does not construct.
    whole = {"key": "x", "columns": (), "arrays": (), "symbols": (),
             "feature": "", "passive": False}
    whole.update({c: NotApplicable("nothing") for c in CONSUMERS})
    for consumer in CONSUMERS:
        half = {k: v for k, v in whole.items() if k != consumer}
        try:
            ForceTerm(**half)
        except TypeError:
            continue
        fails.append(f"a force declared for every equation but {consumer} was "
                     f"accepted; nothing stops a term being added half-way")
    return fails


#: Models whose converged solution the report used to refuse, with the method
#: that solved them: the quotient reproduced the factor of safety to a few parts
#: in a million, or Spencer's two imbalances vanished to the tolerance the Newton
#: pair was driven to, and a relative 1e-6 turned each of them into a method
#: section with no working in it. The last two are the collapsed-scale case,
#: where Q acts through the coordinate origin on every slice and the moment terms
#: sum to two parts in a billion of one force-length unit.
_CONVERGED_BUT_REFUSED = (
    ("geostudio", "gs2_46", "spencer"),
    ("rocscience", "vp037", "spencer"),
    ("rocscience", "vp040", "janbu"),
    ("rocscience", "vp061a", "bishop"),
    ("rocscience", "vp061a", "janbu"),
    ("geostudio", "gs2_26", "spencer"),
    ("rocscience", "vp043", "spencer"),
)


def test_calculation_tolerance_follows_the_solver():
    """A solution that converged as far as it was asked to gets its working.

    The report evaluates its equation and compares the answer with the solver's
    before printing anything, which is what keeps a wrong calculation off the
    page. The comparison has to be made against what the solver delivers:
    Bishop's and Janbu's iterations stop at a step in F of 1e-6 and Spencer's
    Newton pair at imbalances of 1e-4, and a gate of a relative 1e-6 demanded
    more than either and refused seven converged model-method pairs outright.

    Every one of them is required to print here, and the gate is required to
    still refuse an answer that is wrong rather than rounded.
    """
    fails = []
    from xslope.report import (CALC_SAFETY_FACTOR, CALC_TOLERANCE, _closes,
                               _solver_tolerance)

    for vendor, model, method in _CONVERGED_BUT_REFUSED:
        xlsx = os.path.join(_REPO, "docs", "verification", "files", vendor,
                            f"{model}.xlsx")
        report, _bundle = _calc_report(method, xlsx=xlsx)
        if report is None:
            fails.append(f"{model} did not solve with {method}")
            continue
        if _calc_section(report) is None:
            fails.append(f"{model} under {method} converged and still prints no "
                         f"calculation")

    # The tolerances are read off the solvers themselves, so retuning a solver
    # moves the gate that judges its answers with it.
    for method, wanted in (("bishop", 1e-6), ("janbu", 1e-6), ("spencer", 1e-4),
                           ("corps", 1e-6), ("lowe", 1e-6), ("mprice", 1e-6)):
        got = _solver_tolerance(method)
        if got != wanted:
            fails.append(f"{method} converges to {wanted:.0e} and the gate "
                         f"reads {got:.0e}")
    if _solver_tolerance("oms") != 0.0:
        fails.append("the Ordinary Method of Slices is closed form; its gate "
                     "should allow no iteration tolerance at all")

    # Mutation, the quotient: a factor of safety out by one percent is not a
    # converged solution rounding, and no allowance may let it through.
    for method, scale in (("oms", 2.0), ("bishop", 2.0), ("janbu", 2.0),
                          ("mprice", 2.0), ("spencer", 144.0)):
        if _closes(0.01 * scale, scale, method):
            fails.append(f"{method}: a residual one percent of {scale} was "
                         f"accepted as a converged solution")
    # Mutation, the residual: an imbalance a thousand times what the solver
    # itself stops at, against a scale at which the relative test is no looser
    # than the absolute one, so neither statement can excuse it.
    for method in ("spencer", "bishop"):
        allowance = CALC_SAFETY_FACTOR * _solver_tolerance(method)
        if _closes(1000 * allowance, allowance / CALC_TOLERANCE, method):
            fails.append(f"{method}: an imbalance a thousand times the "
                         f"allowance was accepted")
    # And the allowance is a real widening: what the solver delivers has to be
    # more than the relative test alone would take.
    if not _closes(15 * _solver_tolerance("bishop"), 1.94, "bishop"):
        fails.append("a Bishop solution fifteen tolerances from its own fixed "
                     "point is still refused")
    return fails


def _method_block(report, method):
    """The whole of one method's block — its section and every descendant."""
    from xslope.report import method_label

    label = method_label(method)
    for section in report.sections:
        for _lvl, node in section.walk():
            if node.title == label:
                return node
    return None


def _subtree_prose(section):
    """Every paragraph in a section and its children, as one list."""
    return [b.text for _lvl, node in section.walk()
            for b in node.blocks if b.kind == "prose"]


def test_a_method_block_never_goes_quiet():
    """A method whose equilibrium cannot be worked through says what is true of
    it instead.

    The passive-support model is the case: its capacity mobilizes with the soil,
    so it enters divided by the factor of safety and stands on both sides of the
    balance the force methods solve. There is no quotient behind that factor of
    safety, and for five methods the block simply stopped after the slice table
    — a factor of safety with nothing after it, which reads as an omission.

    The moment methods CAN show passive support: it makes a resisting moment of
    its own. So the same model is required to carry the working under Bishop and
    the Ordinary Method of Slices, and the sentence under neither.
    """
    fails = []
    import xslope.report as report

    for method in ("janbu", "corps", "lowe", "mprice", "spencer"):
        built, _bundle = _calc_report(method, xlsx=PASSIVE_XLSX)
        if built is None:
            fails.append(f"the passive model did not solve with {method}")
            continue
        block = _method_block(built, method)
        if block is None:
            fails.append(f"{method}: the report has no block for the method")
            continue
        if any(node.title == "Calculations" for _lvl, node in block.walk()):
            fails.append(f"{method}: a passive model has no quotient to work "
                         f"through and the block carries a Calculations section")
        said = [p for p in _subtree_prose(block) if p == report.PASSIVE_NOTE]
        if len(said) != 1:
            fails.append(f"{method}: the block says why no working is printed "
                         f"{len(said)} times, not once")

    for method in ("bishop", "oms"):
        built, _bundle = _calc_report(method, xlsx=PASSIVE_XLSX)
        if built is None:
            fails.append(f"the passive model did not solve with {method}")
            continue
        block = _method_block(built, method)
        if not any(node.title == "Calculations" for _lvl, node in block.walk()):
            fails.append(f"{method}: passive support makes a resisting moment "
                         f"and the block prints no calculation")
        if report.PASSIVE_NOTE in _subtree_prose(block):
            fails.append(f"{method}: the block prints its working and says it "
                         f"cannot")

    # A tolerance refusal states the two numbers. The gate is widened to the
    # solver's own, so this is exercised on a residual constructed to fail it
    # rather than on a model: what is required is that the sentence carries the
    # quotient and the solution it does not return.
    denied = report._calculation(None, {"slice_df": None, "results": {}}, "bishop")
    if denied != (None, ""):
        fails.append(f"a model with no slices produced {denied!r}")

    # Mutation: with nothing to say, the block goes quiet again, which is what
    # this check exists to catch.
    saved = report.PASSIVE_NOTE
    report.PASSIVE_NOTE = ""
    try:
        quiet, _bundle = _calc_report("janbu", options={"title": "quiet"},
                                      xlsx=PASSIVE_XLSX)
        block = _method_block(quiet, "janbu") if quiet is not None else None
        if block is not None and saved in _subtree_prose(block):
            fails.append("the passive sentence is printed from somewhere other "
                         "than the refusal that withheld the calculation, so "
                         "the two can disagree")
    finally:
        report.PASSIVE_NOTE = saved
    return fails


def _quoted_numbers(sentence, fs, bounds):
    """The numbers a refusal sentence prints that are not factors of safety.

    The reported factor of safety is one whatever it is; every other number in
    the sentence has to lie in the range one can take.
    """
    low, high = bounds
    return [n for n in _numbers(sentence)
            if abs(n - fs) > 5e-7 and not low <= n <= high]


def test_a_refusal_prints_no_number_it_cannot_stand_behind():
    """A refused calculation states the mismatch; it does not typeset a
    degenerate quotient as a measurement.

    The equation is evaluated on the values the bundle carries, and a bundle
    whose slices were never solved carries initial ones: the driving sum is near
    zero and the quotient comes out at 1e14. Printed to six decimals beside the
    factor of safety the solution reports, that reads as a computed result and is
    an artifact of arithmetic on an unsolved frame. A mismatch between two
    factors of safety still states both.
    """
    fails = []
    import xslope.report as report
    from xslope.report import CREDIBLE_FS, _mismatch_note

    fs = 1.759948
    near = _mismatch_note(1.752341, fs)
    for want in ("1.752341", f"{fs:.6f}"):
        if want not in near:
            fails.append(f"a mismatch between two factors of safety does not "
                         f"state {want}: {near!r}")

    for computed in (2.7206093492393162e14, 1e-9, -3.4, CREDIBLE_FS[1] * 1.5):
        said = _mismatch_note(computed, fs)
        loose = _quoted_numbers(said, fs, CREDIBLE_FS)
        if loose:
            fails.append(f"the equation evaluated at {computed:g} is reported "
                         f"as {loose}, which no slope has: {said!r}")
        if f"{fs:.6f}" not in said:
            fails.append(f"the refusal at {computed:g} does not state the "
                         f"factor of safety the solution reports: {said!r}")

    # And on the model it was measured on: a non-circular slope whose slices the
    # report is handed unsolved.
    slope_data, solutions = _refused_solutions()
    from xslope.report import build_report
    with tempfile.TemporaryDirectory() as tmp:
        with contextlib.redirect_stdout(io.StringIO()):
            built = build_report(
                slope_data, solutions,
                {"input_path": NONCIRC_XLSX, "method": "janbu",
                 "pd_figure": False, "lem_search_figure": False,
                 "lem_solution_figure": False}, tmp)
    refusals = [p for p in _prose(built)
                if "does not return the solution" in p]
    if len(refusals) != 1:
        fails.append(f"the unsolved frame produced {len(refusals)} refusals, "
                     f"not one")
    fs_reported = solutions["lem"][0]["results"]["FS"]
    for said in refusals:
        loose = _quoted_numbers(said, fs_reported, CREDIBLE_FS)
        if loose:
            fails.append(f"the refusal prints {loose}: {said!r}")

    # Mutation: with the range opened, the degenerate quotient is typeset again,
    # which is what this check exists to catch.
    saved = report.CREDIBLE_FS
    report.CREDIBLE_FS = (0.0, float("inf"))
    try:
        said = report._mismatch_note(2.7206093492393162e14, fs)
        if not _quoted_numbers(said, fs, saved):
            fails.append("the raw quotient was printed and this check passed "
                         "the sentence")
    finally:
        report.CREDIBLE_FS = saved
    return fails


def test_calculation_reproduces_fs():
    """The factor of safety is re-derived from the operands as PRINTED, for
    every method, and matches the solver to the last digit it prints.

    This is the check that keeps the section from drifting: if the equation ever
    stops being the one the solver evaluates, or a sum is printed to too few
    digits to divide, the number will not come back.
    """
    fails = []
    for method in QUOTIENT_METHODS:
        report, bundle = _calc_report(method)
        if report is None:
            fails.append(f"{method}: the sample model did not solve")
            continue
        section = _calc_section(report)
        if section is None:
            fails.append(f"{method}: the report carries no Calculations section")
            continue
        fs = float(bundle["results"]["FS"])
        num, den, quotient, factor, corrected = _operands(section)
        ok, why = _reproduces(num, den, quotient, factor, corrected, fs)
        if not ok:
            fails.append(f"{method}: {why}")

    # And Spencer prints no quotient at all. A ratio of two force sums is true of
    # its answer and true of everybody else's, which is exactly why printing it
    # as Spencer's equation was wrong: it is not the method.
    report, _bundle = _calc_report("spencer")
    section = _calc_section(report) if report is not None else None
    if section is None:
        fails.append("spencer: the report carries no Calculations section")
    else:
        num, den, quotient, _factor, _corrected = _operands(section)
        if num is not None or den is not None:
            fails.append(f"spencer prints a quotient of two sums, "
                         f"{num}/{den}, as though it were the method")
        if quotient is not None:
            fails.append(f"spencer closes on the arithmetic F = {quotient}; its "
                         f"solution is the root of two equilibrium equations, "
                         f"not the value of an expression")

    # Mutation: the check has to be able to fail. A sum wrong in its third
    # significant digit moves the factor of safety in its third decimal, which is
    # the smallest error the printed precision can be asked to catch.
    report, bundle = _calc_report("bishop")
    section = _calc_section(report)
    num, den, quotient, factor, corrected = _operands(section)
    fs = float(bundle["results"]["FS"])
    for scale in (1.01, 0.999):
        nudged = f"{float(num) * scale:.6g}"
        ok, _why = _reproduces(nudged, den, quotient, factor, corrected, fs)
        if ok:
            fails.append(f"a numerator moved from {num} to {nudged} still "
                         f"reproduced the factor of safety; the check cannot "
                         f"fail")
    return fails


def test_calculation_sums_are_printed_precisely_enough():
    """The sums carry enough digits to divide, and that number is derived rather
    than picked: :data:`xslope.columns.SUM_DIGITS` follows from the format the
    factor of safety is printed in."""
    fails = []
    from xslope.columns import FS_FMT, SUM_DIGITS, format_sum

    decimals = len(FS_FMT.format(1.0).split(".")[-1])
    # The quotient's relative error is at most 10^(1-s); the factor of safety
    # needs it under 10^-decimals / F. Six digits must clear that for F = 100.
    if 10 ** (1 - SUM_DIGITS) > 10 ** -decimals / 100:
        fails.append(f"{SUM_DIGITS} significant digits cannot reproduce a "
                     f"factor of safety printed to {decimals} decimals")
    for value, want in ((1234567.89, 7), (0.000123456789, 6), (1.5, 2)):
        text = format_sum(value)
        got = len(text.replace("-", "").replace(".", "").lstrip("0"))
        if got < min(want, SUM_DIGITS):
            fails.append(f"format_sum({value}) = {text!r} keeps {got} digits")
        if abs(float(text) - value) > abs(value) * 10 ** (1 - SUM_DIGITS):
            fails.append(f"format_sum({value}) = {text!r} loses too much")
    return fails


def test_calculation_terms_follow_the_model():
    """The equation carries the terms the model exercises and no others.

    Two models of the same shape: one with reinforcement crossing the failure
    surface and one without. The reinforcement terms are in the first equation
    and absent from the second — the section describes THIS analysis, not the
    general case.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.report import build_report
    from xslope.slice import generate_slices
    from xslope.solve import solve_selected

    def equation(xlsx, method="bishop"):
        slope_data = load_slope_data(xlsx)
        ok, out = generate_slices(slope_data, circle=slope_data["circles"][0],
                                  num_slices=15)
        if not ok:
            return None, None
        df, surface = out[0].copy(), out[1]
        with contextlib.redirect_stdout(io.StringIO()):
            results = solve_selected(method, df)
        if not isinstance(results, dict):
            return None, None
        with tempfile.TemporaryDirectory() as tmp:
            report = build_report(
                slope_data,
                {"lem": [{"slice_df": df, "failure_surface": surface,
                          "results": results, "search": None, "method": method}]},
                {"method": method, "pd_figure": False,
                 "lem_search_figure": False, "lem_solution_figure": False}, tmp)
        section = _calc_section(report)
        if section is None:
            return None, report
        maths = [b.notation for b in section.blocks if b.kind == "math"]
        prose = " ".join(b.text for b in section.blocks if b.kind == "prose")
        return (maths, prose), report

    nailed = os.path.join(_REPO, "docs", "inputs", "slope",
                          "xslope_nail_axial.xlsx")
    got, _report = equation(nailed)
    if got is None:
        fails.append("the nailed sample produced no calculation to compare")
    else:
        maths, prose = got
        equation_line = next((m for m in maths if m.startswith("F = frac{sum")), "")
        if "P" not in equation_line:
            fails.append(f"a model with reinforcement crossing the surface "
                         f"prints no P term: {equation_line}")
        # The absent list is the CLAUSE that follows, and a semicolon ends a
        # clause. The same sentence goes on to say what the model does carry —
        # passive support, which mobilizes with the soil and joins the numerator
        # — and reading a fixed 200 characters past "carries no" read that as
        # part of the list, calling a reinforced model unreinforced on the
        # strength of the sentence that says it is reinforced.
        absent = prose.split("The model carries no")[-1].split(";")[0][:200]
        if "reinforcement" in absent:
            fails.append("a reinforced model is described as carrying no "
                         "reinforcement")
        # And the reading still catches the thing it is for.
        named = ("The model carries no reinforcement crossing the failure "
                 "surface; the piles in this model mobilize with the soil")
        if "reinforcement" not in \
                named.split("The model carries no")[-1].split(";")[0][:200]:
            fails.append("the absent-force clause is read in a way that would "
                         "miss a reinforced model called unreinforced")

    got, _report = equation(REINF_XLSX)
    if got is None:
        fails.append("the sample model produced no calculation")
    else:
        maths, prose = got
        equation_line = next((m for m in maths if m.startswith("F = frac{sum")), "")
        for absent in ("P ", "P·", "kW", "H cos", "T·a_t"):
            if absent in equation_line:
                fails.append(f"the equation prints a {absent!r} term for a model "
                             f"with none: {equation_line}")
        if "carries no" not in prose:
            fails.append("the section does not say what was left out")
        for named in ("seismic load", "reinforcement", "pile force"):
            if named not in prose:
                fails.append(f"the omission sentence does not name {named!r}")
    return fails


def test_calculation_columns():
    """The per-slice terms of the sums are columns of the slice table.

    That is what lets the section be four lines instead of a walk through every
    slice, so the columns have to be really there — in the table the report
    prints, under the labels the section names.
    """
    fails = []
    from xslope import columns as cols

    # Spencer's F_h and F_v are in the list for the same reason as Q_s: the
    # preamble prints the equation that turns them into Q, and a reader can only
    # check a row against it if the row carries them. F_R and F_D are NOT: they
    # are the two halves of a quotient, and Spencer's section prints none.
    for method, wanted in (("bishop", ("M_R", "M_D")),
                           ("spencer", ("F_h", "F_v", "Q_s", "y_Q"))):
        report, _bundle = _calc_report(method)
        if report is None:
            fails.append(f"{method}: the sample model did not solve")
            continue
        table = next((t for t in report.tables() if t.landscape), None)
        if table is None:
            fails.append(f"{method}: there is no slice table")
            continue
        labels = [h.split(" (")[0] for h in table.headers]
        for label in wanted:
            if label not in labels:
                fails.append(f"{method}: the slice table has no {label} column "
                             f"({labels})")
        legend = dict(table.legend)
        for head in table.headers:
            if head not in legend:
                fails.append(f"{method}: {head} has no legend entry")
        # And the section points at every one of them by the label the table
        # prints — a column nothing in the prose sends the reader to is a column
        # they have no reason to read.
        section = _calc_section(report)
        prose = " ".join(b.text for b in section.blocks if b.kind == "prose")
        for label in wanted:
            if label not in prose:
                fails.append(f"{method}: the calculation never names column "
                             f"{label}")
        if method == "spencer":
            for unwanted in ("F_R", "F_D"):
                if unwanted in labels:
                    fails.append(f"spencer's slice table carries a {unwanted} "
                                 f"column; it exists to be divided, and there "
                                 f"is no quotient in the section to divide it")

    # The computed columns are declared as the report's own, so a solved
    # slice_df is not expected to carry them.
    _slope_data, solutions = _solved()
    have = set(solutions["lem"][0]["slice_df"].columns)
    for column in cols.computed_columns():
        if column.key in have:
            fails.append(f"{column.key!r} is declared as computed by the report "
                         f"but the solver already writes it")
    if not cols.computed_columns():
        fails.append("no column is declared as the report's own; the split is "
                     "not being tested")
    return fails


def test_calculation_residuals():
    """Spencer's two equilibrium equations, rebuilt from the values as PRINTED.

    This is Spencer's answer to the check the other methods get from dividing
    two printed sums: its solution is the pair (F, θ) at which R_1 and R_2 both
    vanish, so what a reviewer verifies is that they DO vanish — at the printed
    θ, from the printed columns, against the size of the sums they cancel within.
    Both are re-formed here the way the section says a reader can re-form them:
    R_1 is the Q_s column added up, R_2 is that column weighted by x_c and y_Q.

    The solver's own residuals are vanishing. A reader adding up the printed
    column gets something larger — the rounding of every row — and the section
    says how large that is allowed to be. This checks both statements.
    """
    import math
    import re
    fails = []
    from xslope.report import CALC_TOLERANCE, PRINTED_RESIDUAL_TOLERANCE

    report, bundle = _calc_report("spencer")
    if report is None:
        return ["the sample model did not solve with Spencer's method"]
    section = _calc_section(report)
    table = next((t for t in report.tables() if t.landscape), None)
    if section is None or table is None:
        return ["there is no calculation or no slice table to read"]

    labels = [h.split(" (")[0] for h in table.headers]
    needed = ("Q_s", "x_c", "y_Q")
    if not set(needed) <= set(labels):
        return [f"the slice table carries no {sorted(set(needed) - set(labels))} "
                f"column, so neither equilibrium sum can be re-formed: {labels}"]
    at = {name: labels.index(name) for name in needed}
    values = [float(row[at["Q_s"]]) for row in table.rows
              if row[at["Q_s"]].strip()]
    if len(values) != len(table.rows):
        fails.append(f"{len(values)} of {len(table.rows)} Q_s cells hold a number")
    total = sum(values)
    scale = sum(abs(v) for v in values) or 1.0
    if abs(total) > PRINTED_RESIDUAL_TOLERANCE * scale:
        fails.append(f"the printed Q_s column sums to {total:.4g}, which is "
                     f"{abs(total) / scale:.1e} of the {scale:.6g} its "
                     f"magnitudes come to — past the stated "
                     f"{PRINTED_RESIDUAL_TOLERANCE:.0e}")

    prose = " ".join(b.text for b in section.blocks if b.kind == "prose")
    if "thousandth" not in prose:
        fails.append("the section does not state, in words, how closely the "
                     "printed values close")

    # The converged F is the answer a reader is looking for, so the sentence
    # that states it sets it in bold — findable without reading the sentence.
    from xslope.columns import format_fs
    printed_fs = format_fs(bundle["results"]["FS"])
    said = [b for b in section.blocks
            if b.kind == "prose" and f"F = {printed_fs}" in b.text]
    if not said:
        fails.append(f"the section never states the converged F = {printed_fs}")
    elif not any(printed_fs in b.bold for b in said):
        fails.append("the converged factor of safety is not set in bold where "
                     "the section states it")

    # θ comes off the page too: the section states it, and without it neither
    # the moment sum nor a single row of Q can be reproduced by a reader.
    stated = re.search(r"θ = (-?\d+\.\d+) degrees", prose)
    if stated is None:
        fails.append("the section never prints the interslice inclination it "
                     "converged at, which the moment sum cannot be re-formed "
                     "without")
    else:
        theta = math.radians(float(stated.group(1)))
        moments = [float(row[at["Q_s"]]) * (float(row[at["x_c"]]) * math.sin(theta)
                                            - float(row[at["y_Q"]]) * math.cos(theta))
                   for row in table.rows]
        m_total, m_scale = sum(moments), sum(abs(v) for v in moments) or 1.0
        if abs(m_total) > PRINTED_RESIDUAL_TOLERANCE * m_scale:
            fails.append(f"the moment sum re-formed from the printed columns is "
                         f"{m_total:.4g}, which is {abs(m_total) / m_scale:.1e} "
                         f"of the {m_scale:.6g} its terms come to — past the "
                         f"stated {PRINTED_RESIDUAL_TOLERANCE:.0e}")

    # The two equilibrium equations, and each of them evaluated at the pair the
    # iteration reached. Their numbers are in the sentence above them — no
    # equation on the page carries one — and that sentence is what
    # :func:`test_equation_numbers_are_in_the_prose` holds to the derivation.
    maths = [b.notation for b in section.blocks if b.kind == "math"]
    transcribed = {}
    for notation in maths:
        found = re.match(SPENCER_EQUATION, notation)
        if found:
            transcribed[found.group(1)] = found.group(2)
    if sorted(transcribed) != ["1", "2"]:
        fails.append(f"the section does not print the two equilibrium equations "
                     f"of the derivation: {maths}")
    if "equations (27) and (28)" not in " ".join(
            b.text for b in section.blocks if b.kind == "prose"):
        fails.append("the section prints the two equilibrium equations without "
                     "saying which equations of the derivation they are")
    evaluated = [re.match(SPENCER_EVALUATION, n) for n in maths
                 if re.match(SPENCER_EVALUATION, n)]
    if len(evaluated) != 2:
        fails.append(f"the section evaluates {len(evaluated)} of its two "
                     f"equilibrium equations at the converged pair: {maths}")
    printed = {}
    for found in evaluated:
        which, notation, value = found.groups()
        name = f"R_{which}"
        # The evaluated line has to BE the equation it is evaluating. An
        # evaluation whose left side drifted from the transcription is a number
        # a reader cannot tie to anything.
        stated_eq = transcribed.get(which, "")
        if notation != stated_eq:
            fails.append(f"{name} is evaluated as {notation!r} but transcribed "
                         f"as {stated_eq!r}; the two do not read as one equation")
        if "e-" not in value and "e+" not in value:
            fails.append(f"a residual is not in scientific notation: {value}")
        printed[name] = float(value)
    # And they really are zero, to the tolerance the section is gated on — the
    # same 1e-6 the other methods' quotients must reproduce F to.
    for name, against in (("R_1", scale), ("R_2", m_scale if stated else None)):
        if name not in printed or against is None:
            continue
        if abs(printed[name]) > CALC_TOLERANCE * against:
            fails.append(f"{name} = {printed[name]:.4g} against a scale of "
                         f"{against:.6g} is {abs(printed[name]) / against:.1e} — "
                         f"past the {CALC_TOLERANCE:.0e} a converged solution "
                         f"has to close to")

    # Both equations close AGAINST SOMETHING, and the section says what. A
    # residual printed on its own is a digit a reader has no way to judge, so
    # each is required to be stated against the sum of the magnitudes it cancels
    # within — the solver's own totals, in the report's own format — and the
    # closure each is claimed to reach is required to be true and not overstated.
    fails += _spencer_closures_are_true(bundle, prose, printed)

    # Morgenstern-Price prints its two residuals as well.
    report, _bundle = _calc_report("mprice")
    section = _calc_section(report) if report is not None else None
    if section is None:
        fails.append("Morgenstern-Price produced no calculation section")
    else:
        maths = [b.notation for b in section.blocks if b.kind == "math"]
        if not any(m.startswith("Z_n = ") for m in maths):
            fails.append(f"Morgenstern-Price prints no force residual: {maths}")
        if not any(m.startswith("sum{M_O} = ") for m in maths):
            fails.append(f"Morgenstern-Price prints no moment residual: {maths}")
    return fails


#: The counts and magnitudes a closure is spoken in — "one part in eight
#: billion". Read back off the page here and turned into a number, so the phrase
#: is a claim the check can measure and not decoration.
_CLOSURE_COUNTS = {"a": 1, "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
                   "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10}
_CLOSURE_MAGNITUDES = {"thousand": 1e3, "million": 1e6,
                       "billion": 1e9, "trillion": 1e12}


def _spencer_closures_are_true(bundle, prose, printed):
    """Each equilibrium equation is stated against the size of the sums it
    cancels within, and the closure claimed for it is the closure it has.

    Zero is not a number a residual can be judged against — 9.136e-05 is tiny in
    a moment sum of ninety thousand and enormous in one of a tenth. So the
    section states both totals and what fraction of each the residual is, and
    both statements are checked here: the totals against the solver's own, the
    claimed closure against the arithmetic. The phrase carries one figure, so it
    is allowed the rounding of one figure either way; past that — a closure
    overstated by more than that rounding, or understated by a factor of ten —
    the sentence has stopped describing this solution.
    """
    import re
    fails = []
    from xslope.columns import format_sum
    from xslope.report import _spencer_state

    state = _spencer_state(bundle["slice_df"])
    if state is None:
        return ["the Spencer state the section is written from cannot be rebuilt"]
    for name, key, what in (("R_1", "scale", "the magnitudes of Q"),
                            ("R_2", "m_scale", "its moment terms")):
        total = format_sum(state[key])
        if total not in prose:
            fails.append(f"{name} is printed against nothing: the section never "
                         f"states the {total} that {what} come to, so a reader "
                         f"has no size to judge the residual by")

    claims = re.findall(r"one part in ([a-z]+|[\d,]+) "
                        r"(thousand|million|billion|trillion)", prose)
    if len(claims) != 2:
        return fails + [f"the section states {len(claims)} closures for its two "
                        f"equations; each has to say how completely it closes"]
    for (count, word), (name, key) in zip(claims, (("R_1", "scale"),
                                                   ("R_2", "m_scale"))):
        if count in _CLOSURE_COUNTS:
            claimed = _CLOSURE_COUNTS[count] * _CLOSURE_MAGNITUDES[word]
        elif count.replace(",", "").isdigit():
            claimed = int(count.replace(",", "")) * _CLOSURE_MAGNITUDES[word]
        else:
            fails.append(f"{name}'s closure is stated as 'one part in {count} "
                         f"{word}', which is not a number")
            continue
        residual = abs(printed.get(name, state[name.replace("R_", "R")]))
        if not residual:
            continue
        actual = abs(state[key]) / residual
        if claimed > actual * 1.1:
            fails.append(f"{name} is claimed to close to one part in {count} "
                         f"{word}, but {residual:.4g} against {state[key]:.6g} "
                         f"is one part in {actual:.3g} — the section overstates "
                         f"its own closure")
        if claimed < actual / 10:
            fails.append(f"{name} is claimed to close to one part in {count} "
                         f"{word} when it closes to one part in {actual:.3g}; "
                         f"the phrase is more than a factor of ten out")
    return fails


#: Spencer's page in LaTeX, in the notation the report prints. Everything the
#: two published force sums are written with, and nothing else: a macro the map
#: does not carry survives the translation as a backslash and fails, rather than
#: quietly comparing equal to something.
_LATEX_TO_NOTATION = {
    r"\sin": "sin", r"\cos": "cos", r"\beta": "β", r"\psi": "ψ",
    r"\theta_p": "θ_p", r"\delta": "δ", "-": "−",
}


def _page_equation(page, number):
    """One numbered equation of a documentation page, in the report's notation.

    The pin behind the transcription: the report assembles equations (1) and (2)
    from its own registry, and this reads them off the page they came from, so
    that a registry entry renamed, resigned or reordered stops matching the
    published form instead of quietly publishing a new one.
    """
    import re

    found = re.search(r"\$([^$]*?)\\qquad \(%s\)\$" % number, page)
    if not found:
        return None, f"{METHOD_DOC_PAGES['spencer']} has no equation ({number})"
    text = found.group(1)
    for latex, plain in _LATEX_TO_NOTATION.items():
        text = text.replace(latex, plain)
    # A leading unary minus is set against its term. LaTeX spaces it or not as
    # the author typed it — equation (2) is written "$F_v = - W ...$" — and that
    # is typesetting, not notation.
    text = " ".join(text.split()).replace("= − ", "= −")
    if "\\" in text:
        return None, (f"equation ({number}) of "
                      f"{METHOD_DOC_PAGES['spencer']} is written with a macro "
                      f"the notation map does not carry: {text!r}")
    return text, ""


#: The base normal, as the two iterative methods' sections write it. Their
#: sections print it transcribed-then-reduced exactly as they print the equation
#: they solve, so a section carries TWO such pairs and every check that reads one
#: has to say which. The left-hand sides tell them apart: no method's own
#: transcription writes a line beginning ``N' =`` or ``m_α =``, and the march
#: writes ``N'·(tan φ_m …``, which is neither.
_NORMAL_FORCE_PREFIXES = ("N' = ", "m_α = ")


def _transcription_pairs(section):
    """``[(sentence, published, reduced), ...]`` — every published-then-reduced
    pair a section prints, in the order it prints them.

    What marks a pair is what marks it for a reader: the sentence between the two
    forms, which names the forces this model does not carry and the equation that
    loses them. The published form is the run of equations directly above that
    sentence and this model's is the run directly below it.
    """
    blocks = list(section.blocks)
    out = []
    for i, block in enumerate(blocks):
        if block.kind != "prose" or not _REDUCTION_LEAD.search(block.text):
            continue
        published, j = [], i - 1
        while j >= 0 and blocks[j].kind == "math":
            published.insert(0, blocks[j].notation)
            j -= 1
        reduced, j = [], i + 1
        while j < len(blocks) and blocks[j].kind == "math":
            reduced.append(blocks[j].notation)
            j += 1
        out.append((block.text, published, reduced))
    return out


def _is_normal_force(run):
    """Is this run of equations the base normal rather than a method's own?"""
    return bool(run) and all(
        any(n.startswith(p) for p in _NORMAL_FORCE_PREFIXES) for n in run)


def _method_pair(section):
    """``(sentence, published, reduced)`` for the equation the METHOD solves.

    Empty where nothing was dropped: the published form is then this model's and
    stands alone, with no sentence between.
    """
    for pair in _transcription_pairs(section):
        if not _is_normal_force(pair[1]):
            return pair
    return "", [], []


def _normal_force_pair(section):
    """``(sentence, published, reduced)`` for the base normal, or empties."""
    for pair in _transcription_pairs(section):
        if _is_normal_force(pair[1]):
            return pair
    return "", [], []


def _transcription_split(section, prefixes=("",)):
    """``(transcribed, this model's)`` — the equations a section prints as its
    derivation publishes the equation it SOLVES, and every other equation it
    prints.

    Everything else — the base normal reduced to this model, the reduction of the
    method's own equation, the arithmetic — carries only the terms this model
    has, and is what a claim about this model is checked against. Every published
    form is kept out of it, the base normal's included: a transcription is the
    derivation speaking and says nothing about this model.
    """
    _sentence, full, _reduced = _method_pair(section)
    published = {n for _s, pub, _r in _transcription_pairs(section) for n in pub}
    keep = [n for n in full if any(n.startswith(p) for p in prefixes)]
    others = [b.notation for b in section.blocks
              if b.kind == "math" and b.notation not in published
              and any(b.notation.startswith(p) for p in prefixes)]
    return keep, others


#: The sentence between a transcription and this model's reduction of it: it
#: names what the model does not carry and then the equation that loses it.
#: Everything printed below it is the model's own.
_REDUCTION_LEAD = re.compile(r", so equations? \(")


#: The consumers whose equations are printed as a quotient of named sums rather
#: than as one expression. A term of those is a piece of one of several lines,
#: not a whole term of one, so which terms a form carries is read by looking each
#: registry symbol up in the lines rather than by cutting the lines into terms.
_PART_CONSUMERS = ("oms_num", "bishop_num", "page_drv")


def _both_forms(full, reduced, consumers):
    """``(published, reduced)`` — the registry terms each printed form carries.

    Every method but the two moment ones prints one equation per form, and a
    term is what stands between its top-level operators. The moment methods
    print theirs as a quotient of named sums spread over several lines, where
    the same reading would return the part letters; those are read by asking,
    of each symbol the registry declares, whether the lines print it.
    """
    if any(c in _PART_CONSUMERS for c in consumers):
        from xslope.report import FORCE_TERMS, NotApplicable
        declared = []
        for term in FORCE_TERMS:
            for consumer in consumers:
                got = getattr(term, consumer)
                published = (got.published if isinstance(got, NotApplicable)
                             else got)
                declared += [c.symbol for c in published]
        return ([s for s in declared if any(s in n for n in full)],
                [s for s in declared if any(s in n for n in reduced)])
    return ([s for n in full for s in _terms_printed(n)],
            [s for n in reduced for s in _terms_printed(n)])


def _reduction_is_true(section, full, reduced, consumers, where):
    """The sentence between the published equations and this model's reduction of
    them says what actually went, and nothing else.

    Two ways for it to lie, and both are tested. A term the published form
    carries and the reduced form does not, with nothing in the sentence to
    account for it, is a term that vanished off the page. A force the sentence
    says the model does not carry, still standing in the reduced form, is the
    denial that the omission sentence has already been caught making once.

    The forces and the terms they are printed as come from the registry the
    section assembles both forms from, so a force added to it is covered here
    without being named here.
    """
    from xslope.report import FORCE_TERMS, NotApplicable

    fails = []
    owner = {}
    for term in FORCE_TERMS:
        for consumer in consumers:
            got = getattr(term, consumer)
            published = (got.published if isinstance(got, NotApplicable) else got)
            for contribution in published:
                owner[contribution.symbol] = term

    published_terms, reduced_terms = _both_forms(full, reduced, consumers)
    # The sentence that reduces THIS pair. A section that also prints the base
    # normal transcribed-then-reduced carries two, and reading the first one
    # judged the factor-of-safety equation against the base normal's sentence.
    sentence = _method_pair(section)[0]
    dropped = [s for s in published_terms if s not in reduced_terms]

    if not reduced:
        # Nothing was dropped, so the published equations ARE this model's and
        # stand alone. A sentence reducing them would be reducing nothing.
        if sentence:
            fails.append(f"{where}: the section prints no reduced equation and "
                         f"still says {sentence!r}")
        return fails
    if dropped and not sentence:
        fails.append(f"{where}: {dropped} are dropped from the published "
                     f"equations with no sentence saying so")

    for symbol in dropped:
        term = owner.get(symbol)
        if term is None:
            fails.append(f"{where}: {symbol!r} is dropped and belongs to no "
                         f"force in the registry")
        elif term.key == "T_top":
            if "shear on the top of the slice" not in sentence:
                fails.append(f"{where}: {symbol!r} is dropped and the sentence "
                             f"does not say that no shear on the top of the "
                             f"slice is simulated: {sentence!r}")
        elif term.feature not in sentence and symbol not in sentence:
            fails.append(f"{where}: {symbol!r} is dropped from the published "
                         f"equations and the sentence accounts for neither it "
                         f"nor the {term.feature}: {sentence!r}")

    for term in FORCE_TERMS:
        if not term.feature or term.feature not in sentence:
            continue
        carried = [s for s, t in owner.items()
                   if t is term and s in reduced_terms]
        if carried:
            fails.append(f"{where}: the sentence says the model carries no "
                         f"{term.feature}, and the reduced equations print "
                         f"{carried}")
    return fails


def test_spencer_force_sums():
    """Spencer's preamble prints equations (1) and (2) as published, reduces them
    to this model, and a row of the slice table reproduces its own Q.

    The section's claim is that the printed numbers ARE the solution, and Q is
    where a reader would have had to take that on trust: the force sums behind it
    were named in words and shown nowhere. So this does what a reviewer does —
    reads F_h, F_v, α, c, Δl, u and φ off one row, puts them through the Q
    equation as printed at the converged F and θ, and requires the row's own Q_s
    back, to the precision the columns are printed at.

    Above that row-level check sits the transcription. The full forms are pinned
    against the derivation page symbol for symbol, in that page's own order, so
    that the registry the report assembles them from cannot drift from the
    equations it claims to be printing. And the reduction below them is held to
    what actually went: every term the full form carries and the reduced form
    does not has to be accounted for in the sentence between them, and nothing
    the sentence names may still be standing in the reduced form.
    """
    import math
    fails = []

    report, bundle = _calc_report("spencer")
    if report is None:
        return ["the sample model did not solve with Spencer's method"]
    section = _calc_section(report)
    table = next((t for t in report.tables() if t.landscape), None)
    if section is None or table is None:
        return ["there is no calculation or no slice table to read"]

    maths = [b.notation for b in section.blocks if b.kind == "math"]
    full, reduced = _transcription_split(section, ("F_h = ", "F_v = "))
    if len(full) != 2:
        fails.append(f"the preamble does not print equations (1) and (2) of the "
                     f"derivation: {maths}")

    from xslope.report import METHOD_DOC_PAGES
    with open(os.path.join(_REPO, "docs", METHOD_DOC_PAGES["spencer"]),
              encoding="utf-8") as f:
        page = f.read()

    # --- the transcription, against the page it is transcribed from ---
    #
    # Symbol for symbol, in that page's own order: this is what keeps the
    # registry the two forms are assembled from from drifting away from the
    # equations the section says it is printing.
    for notation, number in zip(full, ("1", "2")):
        published, why = _page_equation(page, number)
        if why:
            fails.append(why)
        elif notation != published:
            fails.append(f"the section prints equation ({number}) as "
                         f"{notation!r}; {METHOD_DOC_PAGES['spencer']} publishes "
                         f"{published!r}")

    # --- the reduction, against what it says went ---
    fails += _reduction_is_true(section, full, reduced,
                                ("spencer_h", "spencer_v"), "the sample model")
    if not any("W" in n for n in reduced):
        fails.append(f"neither reduced force sum carries the slice weight: "
                     f"{reduced}")
    # This model has a distributed load and nothing else, so no seismic,
    # reinforcement, pile, line load or tension-crack term survives the
    # reduction.
    for notation in reduced:
        for absent in ("kW", "R cos", "R sin", "H cos", "H sin", "L cos",
                       "L sin", "V"):
            if absent in notation:
                fails.append(f"{notation!r} prints a {absent!r} term for a model "
                             f"with none")

    # And the row-level reproduction, on this model and on a right-facing one —
    # which the section is solved as the mirror image of, and which the preamble
    # therefore has to warn the reader about before they check a row.
    fails += _spencer_rows_reproduce(table, bundle["results"], mirror=1)

    report, bundle = _calc_report("spencer", xlsx=RFACE_XLSX)
    if report is None:
        return fails + ["the right-facing model did not solve with Spencer's "
                        "method"]
    table = next((t for t in report.tables() if t.landscape), None)
    section = _calc_section(report)
    prose = " ".join(b.text for b in section.blocks if b.kind == "prose")
    if "mirror image" not in prose:
        fails.append("a right-facing section does not say that it is solved as "
                     "the mirror image of the slope, which is the only way its "
                     "Q column can be checked against the equations")
    fails += _spencer_rows_reproduce(table, bundle["results"], mirror=-1)

    # The transcription and its reduction on every shape of model there is a
    # Spencer section for: one carrying a distributed load, one carrying water in
    # a tension crack, one right-facing and carrying neither. Each drops a
    # different set of terms, and each sentence has to be true of its own.
    for label, xlsx in (("the right-facing model", RFACE_XLSX),
                        ("the tension-crack model", TENSION_XLSX)):
        report, _bundle = _calc_report("spencer", xlsx=xlsx)
        section = _calc_section(report) if report is not None else None
        if section is None:
            fails.append(f"{label} produced no Spencer calculation to reduce")
            continue
        full, reduced = _transcription_split(section, ("F_h = ", "F_v = "))
        for notation, number in zip(full, ("1", "2")):
            published, why = _page_equation(page, number)
            if why:
                fails.append(f"{label}: {why}")
            elif notation != published:
                fails.append(f"{label}: the section prints equation ({number}) "
                             f"as {notation!r}; {METHOD_DOC_PAGES['spencer']} "
                             f"publishes {published!r}")
        fails += _reduction_is_true(section, full, reduced,
                                    ("spencer_h", "spencer_v"), label)
    return fails


def _spencer_rows_reproduce(table, results, mirror):
    """Every row of a Spencer slice table put back through the printed Q
    equation, and the mutation that proves the arithmetic is being done.

    ``mirror`` is -1 on a right-facing surface, where the derivation is applied
    to the mirrored section and α, c and tan φ enter with reversed signs — as
    the preamble states.
    """
    import math
    fails = []

    labels = [h.split(" (")[0] for h in table.headers]
    needed = ("α", "Δl", "u", "c", "φ", "F_h", "F_v", "Q_s")
    if not set(needed) <= set(labels):
        return [f"the slice table is missing {sorted(set(needed) - set(labels))}"]
    at = {name: labels.index(name) for name in needed}
    F = float(results["FS"])
    theta = math.radians(float(results["theta"]))

    def Q_of(cell, F_v):
        a = mirror * math.radians(cell["α"])
        tan_p = mirror * math.tan(math.radians(cell["φ"]))
        c = mirror * cell["c"]
        m_a = 1.0 / (math.cos(a - theta) + math.sin(a - theta) * tan_p / F)
        return (- F_v * math.sin(a) - cell["F_h"] * math.cos(a)
                - c * cell["Δl"] / F
                + (F_v * math.cos(a) - cell["F_h"] * math.sin(a)
                   + cell["u"] * cell["Δl"]) * tan_p / F) * m_a

    # Every operand is printed rounded, and the reproduction carries that
    # rounding: a thousandth of the magnitudes on the row covers the hundredth
    # of a degree α is printed to — the term that dominates — on a slice of any
    # size, with an order of magnitude to spare. The floor is for a sliver at
    # the toe with almost nothing on it.
    def tolerance(cell):
        return 0.2 + 1e-3 * (abs(cell["F_v"]) + abs(cell["F_h"])
                             + (abs(cell["c"]) + abs(cell["u"])) * cell["Δl"])

    caught = False
    for row in table.rows:
        cell = {name: float(row[at[name]]) for name in needed}
        Q = Q_of(cell, cell["F_v"])
        if abs(Q - cell["Q_s"]) > tolerance(cell):
            fails.append(f"slice {row[0]}: the printed F_h, F_v and base "
                         f"properties give Q = {Q:.3f}, the table {cell['Q_s']}")
        # Mutation: an F_v that did not come from the solve would not reproduce.
        # It has to be caught somewhere rather than on a nominated row — a slice
        # whose Q passes through zero is insensitive to F_v on its own, and
        # requiring the failure there would be requiring luck.
        caught = caught or abs(Q_of(cell, cell["F_v"] * 1.01)
                               - cell["Q_s"]) > tolerance(cell)
    if not caught:
        fails.append("an F_v column wrong by a percent still reproduced every "
                     "Q; the check cannot fail")
    return fails


#: The quantities a row's Q is rebuilt from that the mirror sentence has to
#: place, and the words it can name each of them in. Every one of them is either
#: reversed by the reader or already reversed in the column, and the sentence is
#: only true if it says which. "the horizontal forces" is F_h — the phrase the
#: rule was wrong in, which named F_h among the reversals when the column is
#: printed reversed already, so a reader following it reversed F_h twice.
_MIRROR_NAMES = {
    "α": (r"α",),
    "c": (r"c",),
    "tan φ": (r"tan φ",),
    "F_h": (r"F_h", r"horizontal forces?"),
    "F_v": (r"F_v", r"vertical forces?"),
}


def _mirror_rule(prose):
    """The sign rule a right-facing section STATES, as ``{quantity: reversed}``.

    Read out of the sentence itself rather than assumed, so that the arithmetic
    below is the arithmetic the words describe: a quantity named in the sentence
    that reverses signs is reversed, one named in the sentence that says the
    columns enter as printed is not. Returns ``(rule, complaints)``.
    """
    import re

    fails = []
    para = next((s for s in re.split(r"(?<=\.)\s+", prose)
                 if "mirror image" in s), None)
    if para is None:
        return {}, ["no sentence says the section is solved as a mirror image"]
    where = prose[prose.index(para):]
    sentences = [s for s in re.split(r"(?<=\.)\s+", where) if s]
    rule = {}
    for sentence in sentences:
        if "reversed" in sentence:
            reversed_by = True
        elif "as printed" in sentence:
            reversed_by = False
        else:
            continue
        for quantity, spellings in _MIRROR_NAMES.items():
            if quantity in rule:
                continue
            for spelling in spellings:
                if re.search(rf"(?<![A-Za-zα-ωΑ-Ω_]){spelling}"
                             rf"(?![A-Za-zα-ωΑ-Ω_])", sentence):
                    rule[quantity] = reversed_by
                    break
    # A quantity the sentence never places is a quantity a reader would take as
    # it is printed. That is a defect in the sentence and it is reported, and the
    # arithmetic is still run under that reading, so a wording that misplaces one
    # quantity and omits another fails on both counts.
    missing = [q for q in _MIRROR_NAMES if q not in rule]
    if missing:
        fails.append(f"the mirror sentence never says whether {missing} enter "
                     f"the equations reversed or as printed: {para!r}")
    return rule, fails


def test_spencer_mirror_rule():
    """A right-facing section's mirror sentence IS the arithmetic behind its Q_s
    column.

    The rule is parsed out of the sentence and applied to the printed columns,
    and the row's own Q has to come back. The model is one with real horizontal
    forces on its slices: where F_h is zero on every slice — the sample
    right-facing slope — a rule that reverses F_h and a rule that does not give
    the same Q, and the sentence can say either and read correct.
    """
    import math
    fails = []

    report, bundle = _calc_report("spencer", xlsx=MIRROR_XLSX)
    if report is None:
        return ["the right-facing model did not solve with Spencer's method"]
    section = _calc_section(report)
    table = next((t for t in report.tables() if t.landscape), None)
    if section is None or table is None:
        return ["there is no calculation or no slice table to read"]

    prose = " ".join(b.text for b in section.blocks if b.kind == "prose")
    rule, fails = _mirror_rule(prose)
    if not rule:
        return fails

    labels = [h.split(" (")[0] for h in table.headers]
    needed = ("α", "Δl", "u", "c", "φ", "F_h", "F_v", "Q_s")
    if not set(needed) <= set(labels):
        return [f"the slice table is missing {sorted(set(needed) - set(labels))}"]
    at = {name: labels.index(name) for name in needed}
    F = float(bundle["results"]["FS"])
    theta = math.radians(float(bundle["results"]["theta"]))

    def Q_of(cell, rule):
        def signed(quantity, value):
            return -value if rule.get(quantity) else value
        a = signed("α", math.radians(cell["α"]))
        tan_p = signed("tan φ", math.tan(math.radians(cell["φ"])))
        c = signed("c", cell["c"])
        F_h = signed("F_h", cell["F_h"])
        F_v = signed("F_v", cell["F_v"])
        m_a = 1.0 / (math.cos(a - theta) + math.sin(a - theta) * tan_p / F)
        return (- F_v * math.sin(a) - F_h * math.cos(a) - c * cell["Δl"] / F
                + (F_v * math.cos(a) - F_h * math.sin(a)
                   + cell["u"] * cell["Δl"]) * tan_p / F) * m_a

    # The tolerance of the sibling check: the operands are printed rounded and
    # the reproduction carries that rounding.
    def tolerance(cell):
        return 0.2 + 1e-3 * (abs(cell["F_v"]) + abs(cell["F_h"])
                             + (abs(cell["c"]) + abs(cell["u"])) * cell["Δl"])

    cells = [{name: float(row[at[name]]) for name in needed} for row in table.rows]
    if not any(abs(cell["F_h"]) > 1e-6 for cell in cells):
        return ["every printed F_h is zero on this model, so the mirror rule "
                "cannot be told from a rule that reverses the horizontal forces"]
    reproduced = True
    for row, cell in zip(table.rows, cells):
        Q = Q_of(cell, rule)
        if abs(Q - cell["Q_s"]) > tolerance(cell):
            reproduced = False
            fails.append(f"slice {row[0]}: the rule the mirror sentence states "
                         f"gives Q = {Q:.4f} from the printed columns, the "
                         f"table {cell['Q_s']}")
    # Mutation: the same rule with F_h reversed the other way — which is the
    # rule the sentence used to state — must NOT reproduce, or the sentence's
    # F_h clause is a clause nothing here tests.
    other = dict(rule, F_h=not rule.get("F_h"))
    if reproduced and all(abs(Q_of(cell, other) - cell["Q_s"]) <= tolerance(cell)
                          for cell in cells):
        fails.append("the same Q comes back with F_h reversed the other way, so "
                     "what the sentence says about the horizontal forces is "
                     "not tested")
    return fails


#: The letters each force of the general equation is printed as, so that a
#: sentence saying the model carries none of one can be tested against the
#: equations under it. Spencer's page follows UTEXAS and gives three of them
#: other letters, which is why the two are listed apart: P is the distributed
#: load there and the reinforcement everywhere else.
_FEATURE_SYMBOLS = {
    "distributed load": ("D",),
    "seismic load": ("kW",),
    "tension-crack water force": ("T",),
    "reinforcement crossing the failure surface": ("P", "P_p"),
    "pile force": ("H", "H_p"),
    "line load": ("L",),
}
_SPENCER_FEATURE_SYMBOLS = dict(
    _FEATURE_SYMBOLS,
    **{"distributed load": ("P",),
       "tension-crack water force": ("V",),
       "reinforcement crossing the failure surface": ("R",)})


def test_absent_features_are_really_absent():
    """No section says the model carries a force its own equations print.

    The omission sentence is what tells a reader that a missing term is a term
    the model does not have. A model reinforced entirely by passive capacity has
    nothing in the tangent and axial columns and its P_p terms in the quotient
    directly below, and the sentence denied the reinforcement anyway. So every
    force the sentence names as absent is required absent from every equation the
    section prints.
    """
    fails = []
    from xslope.report import EQUATION_SYMBOLS, _present_symbols

    models = [("the passive-reinforcement model", PASSIVE_XLSX,
               ("bishop", "oms")),
              ("the right-facing reinforced model",
               os.path.join(_REPO, "docs", "lem", "files",
                            "xslope_reinforce_rface.xlsx"), ("bishop", "oms")),
              ("the axially reinforced model", AXIAL_XLSX, ("bishop", "oms")),
              ("the tension-crack model", TENSION_XLSX, CALC_METHODS),
              ("the sample model", REINF_XLSX, CALC_METHODS)]

    claimed = set()
    for where, xlsx, methods in models:
        for method in methods:
            report, _bundle = _calc_report(method, xlsx=xlsx)
            if report is None:
                continue
            section = _calc_section(report)
            if section is None:
                continue
            intro = next((b.text for b in section.blocks
                          if b.kind == "prose" and "carries no" in b.text), "")
            names = (_SPENCER_FEATURE_SYMBOLS if method == "spencer"
                     else _FEATURE_SYMBOLS)
            absent = [name for name in names if name in intro]
            claimed.update(absent)
            # Every section transcribes its equations in full before reducing
            # them, and the transcription is the derivation's — it carries every
            # force there is, and says nothing about this model. What the
            # sentence is a claim about is the reduction below it.
            _full, model_own = _transcription_split(section)
            text = " ".join(model_own)
            # Every symbol the report knows goes in as a candidate, not just the
            # ones being looked for: the longest match wins, and R_1 has to claim
            # its characters before the reinforcement force R can be read out of
            # them.
            present = _present_symbols(
                text, [s for syms in names.values() for s in syms]
                + list(EQUATION_SYMBOLS))
            for name in absent:
                carried = [s for s in names[name] if s in present]
                if carried:
                    fails.append(f"{where} under {method}: the section says it "
                                 f"carries no {name}, and its equations print "
                                 f"{carried}")
    # A passive-reinforced model has to be one of the models this ran on, and it
    # has to print its reinforcement: without that the check above passes on a
    # sentence that names nothing.
    report, _bundle = _calc_report("bishop", xlsx=PASSIVE_XLSX)
    section = _calc_section(report) if report is not None else None
    if section is None:
        fails.append("the passive-reinforcement model produced no calculation, "
                     "so nothing here tests a passive force against its prose")
    else:
        text = " ".join(b.notation for b in section.blocks if b.kind == "math")
        if "P_p" not in text:
            fails.append("the passive-reinforcement model prints no P_p term")
    if not claimed:
        fails.append("no section named an absent force, so nothing was tested")
    return fails


#: Every equation number a method's Calculations section names, and the
#: documentation page each one is published on. Written out here rather than read
#: from the report: a section that starts citing a different equation has to
#: fail, not pass on its own new claim.
#:
#: Corps of Engineers, Lowe & Karafiath and Morgenstern-Price transcribe the
#: per-slice equilibrium of the force-equilibrium derivation, which is the page
#: the first two are documented on and the page Morgenstern-Price's own sends the
#: reader to for its march.
#:
#: The two moment methods each print their OWN page's factor-of-safety equation
#: — the Ordinary Method's (8), Bishop's (10) — and then the sums the arithmetic
#: is formed from, which are that equation in the general moment arms. The
#: Ordinary Method numbers that form (8a); Bishop's page gives it under
#: Composite Surfaces without a number, so Bishop's section names no number for
#: it. Neither section may cite the other's page: an equation of Bishop's method
#: assembled out of the Ordinary Method's numbering is what this is here to stop.
#: (8a) is not among the Ordinary Method's: it is the general moment arms, which
#: only a COMPOSITE surface is shown — on a true circle the base shear's arm is
#: the radius, the named sums of equation (8) are the equilibrium the solver
#: evaluated, and the section evaluates them. The models here are true circles.
_EQUATION_NUMBERS = {
    "oms": (("lem/oms.md", ("4", "8")),),
    "bishop": (("lem/bishop.md", ("8", "10")),),
    "janbu": (("lem/janbu.md", ("1", "4", "5", "6", "7")),),
    "corps": (("lem/force_eq.md", ("6", "7", "12")),),
    "lowe": (("lem/force_eq.md", ("6", "7", "12")),),
    "mprice": (("lem/force_eq.md", ("6", "7", "12")),
               ("lem/mprice.md", ("2", "4", "8"))),
    "spencer": (("lem/spencer.md", ("1", "2", "23", "24", "27", "28")),),
}


def _terms_printed(notation):
    """Every term one printed equation carries, as the registry writes it.

    A term is what stands between the top-level operators: inside a Σ where the
    equation sums over the slices, bare where it is written for one slice, and
    inside a mobilized fraction where the equation divides it by F. Whichever it
    is, what comes back is the symbol the registry declares, so a term can be
    followed from the published form into the model's own however the two are
    arranged.
    """
    def bare(piece):
        piece = piece.strip().lstrip("+−").strip()
        if piece.startswith("sum{") and piece.endswith("}"):
            piece = piece[4:-1]
        if piece.startswith("frac{") and piece.endswith("}{F}"):
            piece = piece[5:-4]
        return piece

    if "sum{" in notation:
        out, at = [], notation.find("sum{")
        while at >= 0:
            depth, j = 0, at + 3
            while j < len(notation):
                if notation[j] == "{":
                    depth += 1
                elif notation[j] == "}":
                    depth -= 1
                    if not depth:
                        break
                j += 1
            out.append(notation[at:j + 1])
            at = notation.find("sum{", j + 1)
        return [bare(p) for p in out]
    body = notation.split(" = ", 1)[-1].strip()
    if body == "0":
        return []
    return [bare(p) for p in re.split(r" [−+] ", body.lstrip("−"))]


#: How the prose names a documentation equation. One form, so that a number
#: reaching the page any other way is a number this cannot check.
#: A reference to a numbered equation of a derivation, written either way round:
#: "equation (4) of the derivation", and "the normal force on the base of the
#: slice (equation 4)", which is the shape a reference takes where it is an aside
#: to the sentence rather than its subject.
_EQUATION_REFERENCE = re.compile(
    r"equations? \(([0-9a-z]+)\)(?: and \(([0-9a-z]+)\))?"
    r"|\(equations? ([0-9a-z]+)\)", re.I)

#: And how a page numbers one. The pages set the number with \qquad or \quad, so
#: the spacing is not what a reference is looked up by.
_PAGE_NUMBER = r"\\q?quad ?\(%s\)"


def _cited_equations(section, method):
    """``[(page, number), ...]`` — every equation number the section names, each
    against the page the sentence that names it resolves on.

    A section that transcribes from two derivations cites numbers from two
    documents, and both documents number an equation (6). What decides which is
    meant is the sentence: one that links a page cites THAT page's numbering, and
    one that links none continues the page the sentence before it established —
    a reduction reads on from the equations introduced above it — starting from
    the method's own derivation. Reading the numbers without their page lets a
    citation drift to the other document and still validate, because the number
    exists on both.
    """
    from xslope.report import METHOD_DOC_PAGES, docs_url

    known = set(METHOD_DOC_PAGES.values())
    for entries in _EQUATION_NUMBERS.values():
        known |= {path for path, _numbers in entries}
    by_url = {docs_url(path).rstrip("/"): path for path in known}
    own = METHOD_DOC_PAGES.get(method, "")

    out, page = [], own
    for block in section.blocks:
        if block.kind != "prose":
            continue
        for _text, target in (getattr(block, "links", None) or ()):
            linked = by_url.get(str(target).rstrip("/"))
            if linked:
                page = linked
                break
        out += [(page, n) for match in _EQUATION_REFERENCE.finditer(block.text)
                for n in match.groups() if n]
    return out


def _citations_are_declared(method, cited, wanted, page_of):
    """The pages and numbers the section cites are the ones declared for it, and
    each number is published on the page it is cited against."""
    fails = []
    if cited != wanted:
        fails.append(f"{method}: the prose cites {sorted(cited)}; the section "
                     f"transcribes {sorted(wanted)}")
    for path, number in sorted(cited | wanted):
        if not re.search(_PAGE_NUMBER % re.escape(number), page_of(path)):
            fails.append(f"{method}: the prose names equation ({number}) of "
                         f"{path}, which publishes no such number")
    return fails


def test_equation_numbers_are_in_the_prose():
    """No equation is printed with a number beside it, and every number the prose
    names is published where it says.

    A report numbers its own equations consecutively. These come from seven
    derivations that number theirs independently, so printing those numbers on
    the page runs (1), (2), (7), (1) down a section with two different equations
    under the same number. The number belongs in the sentence that introduces the
    equation — "equation (7) of the derivation" — where it reads as the reference
    it is, and the section's opening statement links the page it resolves on.

    So: nothing beside the equations, every number named in the prose really on
    the page it is attributed to, and every number a section is expected to name
    really named.
    """
    fails = []
    from xslope.report import METHOD_DOC_PAGES, Math

    # The mechanism itself: an equation has nothing to print a number with.
    if "label" in getattr(Math, "__dataclass_fields__", {}):
        fails.append("Math still carries a label; an equation number can be set "
                     "beside the equation again")

    pages = {}

    def page_of(path):
        if path not in pages:
            with open(os.path.join(_REPO, "docs", path), encoding="utf-8") as f:
                pages[path] = f.read()
        return pages[path]

    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            fails.append(f"{method}: no calculation to read the numbering of")
            continue
        for block in section.blocks:
            if block.kind == "math" and getattr(block, "label", ""):
                fails.append(f"{method}: {block.notation!r} is printed with "
                             f"{block.label!r} beside it")
        declared = dict(_EQUATION_NUMBERS.get(method, ()))
        wanted = {(path, n) for path, numbers in declared.items()
                  for n in numbers}
        cited = set(_cited_equations(section, method))
        fails += _citations_are_declared(method, cited, wanted, page_of)
        # And an equation the section prints has to be introduced by a sentence
        # that says which one it is: a number named nowhere near the equation is
        # a reference the reader cannot follow.
        if wanted and not any(
                b.kind == "prose" and _EQUATION_REFERENCE.search(b.text)
                and b.text.rstrip().endswith(":")
                for b in section.blocks):
            fails.append(f"{method}: no sentence introduces an equation by the "
                         f"number the derivation gives it")
        if method not in _EQUATION_NUMBERS:
            continue
        if not METHOD_DOC_PAGES.get(method):
            fails.append(f"{method}: names equation numbers with no page mapped")

    # The mutation: a citation that drifted to the OTHER page the section
    # transcribes from. Morgenstern-Price cites (8) of its own derivation and
    # (6) and (7) of the force-equilibrium one, and both pages number an
    # equation (6), (7) and (8) — so a number checked without its page passes
    # wherever it is pointed. This moves (8) onto the force-equilibrium page,
    # which publishes an equation (8) of its own, and requires the drift to be
    # caught anyway.
    declared = dict(_EQUATION_NUMBERS["mprice"])
    wanted = {(path, n) for path, numbers in declared.items() for n in numbers}
    drifted = {("lem/force_eq.md", n) if (path, n) == ("lem/mprice.md", "8")
               else (path, n) for path, n in wanted}
    if drifted == wanted:
        fails.append("the mutation moved no citation, so it tests nothing")
    elif not _citations_are_declared("mprice", drifted, wanted, page_of):
        fails.append("a citation moved to the other page the section "
                     "transcribes from still validated")
    return fails


def test_the_published_equation_comes_first():
    """Every method prints the equation its derivation publishes, in full, before
    it prints this model's.

    A section that opens on the reduced form shows a reader an equation with the
    seismic term missing and nothing to say a seismic term exists. So each of the
    seven transcribes the published equation first — every force a slice can
    take, assembled from the registry the model's own equation is assembled from
    — then states in one sentence what this model does not carry, then prints
    what is left.

    Three things are required of that transcription. It carries every
    contribution the registry declares for the equation, so it is the published
    form and not a longer reduction. The forces it carries beyond the ones the
    page's own equation writes are only those the registry marks passive, which
    the pages describe in prose and put on the mobilized side. And the sentence
    below it accounts for every term that then goes.
    """
    fails = []
    from xslope.report import FORCE_TERMS, NotApplicable, TRANSCRIPTIONS

    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            fails.append(f"{method}: no calculation to read the transcription of")
            continue
        spec = TRANSCRIPTIONS[method]
        full, model_own = _transcription_split(section)
        reduced = bool(full)
        if not reduced:
            # Nothing was dropped: the published equation IS this model's, and
            # the section prints it once with no sentence between.
            full = model_own
        printed, _model = _both_forms(full, [], spec.consumers)
        for term in FORCE_TERMS:
            for consumer in spec.consumers:
                got = getattr(term, consumer)
                published = (got.published if isinstance(got, NotApplicable)
                             else got)
                for contribution in published:
                    if contribution.symbol not in printed:
                        fails.append(
                            f"{method}: the published {consumer} carries "
                            f"{contribution.symbol!r} and the transcription "
                            f"does not: {full}")
        # And nothing in it that the registry does not declare for those
        # equations — a term from another equation would be a claim the page
        # never made.
        declared = {c.symbol for t in FORCE_TERMS for consumer in spec.consumers
                    for c in (getattr(t, consumer).published
                              if isinstance(getattr(t, consumer), NotApplicable)
                              else getattr(t, consumer))}
        # The interslice resultant carried in from the slice before is not a
        # force the model applies but what the march computes, so it is written
        # by the equation rather than declared in the registry. Named here so
        # that it is the only thing the registry does not have to account for.
        declared |= {"Z_i·cos θ_i", "Z_i·sin θ_i"}
        if reduced and not any(c in _PART_CONSUMERS for c in spec.consumers):
            # The two moment methods print theirs as a quotient of named sums,
            # where a line carries part of a term rather than a term; what keeps
            # anything the registry does not declare out of THEIR transcription
            # is the recomposition, which has to give the published equation
            # term for term (test_the_moment_quotient_recomposes).
            for symbol in printed:
                if symbol not in declared:
                    fails.append(f"{method}: the transcription prints "
                                 f"{symbol!r}, which the registry declares for "
                                 f"no equation it transcribes")
        if reduced:
            fails += _reduction_is_true(section, full, model_own,
                                        spec.consumers, method)
    return fails


#: Every LaTeX the derivation pages write their transcribed equations with, in
#: the notation the report prints. Applied longest key first. Anything left with
#: a backslash after the pass has not been translated, and the comparisons below
#: refuse rather than quietly matching something else.
_PAGE_LATEX = {
    r"\Delta \ell": "Δl", r"\Delta\ell": "Δl", r"\ell": "l",
    r"\theta_p": "θ_p", r"\theta_{i+1}": "θ_{i+1}", r"\theta_{i}": "θ_i",
    r"\theta_j": "θ_j",
    r"\theta_i": "θ_i", r"\alpha": "α", r"\beta": "β", r"\phi'": "φ",
    r"\phi_m": "φ_m", r"\phi": "φ", r"\psi": "ψ", r"\delta": "δ",
    r"\lambda": "λ", r"\pi": "π",
    r"\dfrac": "frac", r"\frac": "frac", r"\sum": "sum",
    r"\sin": "sin", r"\cos": "cos", r"\tan": "tan",
    r"\left[": "[", r"\right]": "]", r"\left(": "(", r"\right)": ")",
    r"\,": " ", r"\;": " ", r"\ ": " ",
    "-": "−",
}

#: What a term is compared on: the symbols and operators, with the grouping and
#: the spacing taken out. The report writes ``frac{1}{R}·sum{D cos β·a_dx}`` and
#: the page ``\frac{1}{R}\sum D \cos \beta \, a_{dx}``; they are the same term,
#: and the two ways of setting it are not what this is checking.
_GROUPING = "{}[]() ·\t\n"


def _page_latex(text):
    """One page's equation in the report's notation, or ``(None, why)``."""
    out = text
    for latex in sorted(_PAGE_LATEX, key=len, reverse=True):
        out = out.replace(latex, _PAGE_LATEX[latex])
    out = re.sub(r"_\{([A-Za-z0-9+]+)\}", r"_\1", out)
    if "\\" in out:
        return None, f"a macro the notation map does not carry: {out!r}"
    return " ".join(out.split()), ""


def _split_signed(text):
    """``[(sign, text), ...]`` — one side of an equation cut at its top-level
    operators, with everything inside a fraction, a sum or a bracket left whole
    and each piece kept exactly as it was written.
    """
    out, depth, start, sign = [], 0, 0, +1
    for i, ch in enumerate(text):
        if ch in "{[(":
            depth += 1
        elif ch in "}])":
            depth -= 1
        elif depth == 0 and ch in "+−" and i > start:
            out.append((sign, text[start:i]))
            sign, start = (+1 if ch == "+" else -1), i + 1
    out.append((sign, text[start:]))
    return out


def _signed_terms(text):
    """``[(sign, term), ...]`` — :func:`_split_signed` with each term reduced to
    what it says."""
    return [(s, t) for s, t in
            [(s, "".join(c for c in t if c not in _GROUPING))
             for s, t in _split_signed(text)]
            if t]


def _sum_body(term):
    """What a term that is one Σ over a bracketed group sums, or None.

    ``sum{A + B}`` and ``\\sum [A + B]`` give back the group; ``sum ( A + B )
    cos α``, where the bracket closes before the term does, is not one sum over a
    group but a product, and gives back nothing.
    """
    text = term.strip()
    if not text.startswith("sum"):
        return None
    rest = text[3:].strip()
    if not rest or rest[0] not in "{[(":
        return None
    depth = 0
    for i, ch in enumerate(rest):
        if ch in "{[(":
            depth += 1
        elif ch in "}])":
            depth -= 1
            if not depth:
                return rest[1:i] if i == len(rest) - 1 else None
    return None


def _balance_atoms(side):
    """One side of a balance as a sorted ``[(sign, term), ...]``, with every Σ
    distributed over the terms it carries.

    A sum written once over many terms and a sum written over each of them are
    the same statement — ``Σ[A + B]`` and ``ΣA + ΣB`` — and the two pages that
    publish the whole-mass balance write it both ways and in different orders.
    Distributing the sum and comparing the terms as a set is what holds a
    transcription to a page across that: every term of the published equation, on
    the side the page puts it, with the sign the page gives it.
    """
    out = []
    for sign, term in _split_signed(side):
        body = _sum_body(term)
        pieces = _split_signed(body) if body is not None else [(+1, term)]
        for inner, piece in pieces:
            bare = "".join(c for c in piece.replace("sum", "")
                           if c not in _GROUPING)
            if bare:
                out.append((sign * inner, bare))
    return sorted(out)


def _as_quotient(text, where):
    """``(numerator, denominator)`` of ``F = frac{A}{B}``."""
    body = text.split("=", 1)[-1].strip()
    if not body.startswith("frac{"):
        return None, None, f"{where} is not a quotient: {text!r}"
    pieces, at = [], len("frac")
    for _side in (0, 1):
        if at >= len(body) or body[at] != "{":
            return None, None, f"{where} is not a quotient: {text!r}"
        depth, j = 0, at
        while j < len(body):
            if body[j] == "{":
                depth += 1
            elif body[j] == "}":
                depth -= 1
                if not depth:
                    break
            j += 1
        pieces.append(body[at + 1:j])
        at = j + 1
    return pieces[0], pieces[1], ""


def _recompose(lines):
    """The named parts of one printed quotient substituted back into it.

    ``lines`` are the equations the section prints for the published form: the
    quotient itself, the numerator's sum, the group the base normal is formed
    from, and one part per force. Returns ``(numerator, denominator, why)`` with
    every part letter replaced by what the section defines it as — which is what
    is then held against the page.

    A part standing as a whole term of the quotient is substituted bare, so that
    the terms inside it become terms of the quotient, as they are on the page.
    The group inside the numerator is multiplied by tan φ, so it is bracketed —
    also as the page writes it.
    """
    defined = {}
    top = ""
    for line in lines:
        name, _sep, body = line.partition(" = ")
        if name == "F":
            top = line
        else:
            defined[name] = body
    if not top:
        return None, None, f"no quotient among the printed equations: {lines}"
    num, den, why = _as_quotient(top, "the printed quotient")
    if why:
        return None, None, why

    def substitute(text, bracket):
        for name in sorted(defined, key=len, reverse=True):
            if name in text:
                body = defined[name]
                text = text.replace(name, f"({body})" if bracket else body)
        return text

    # The quotient's own parts first, bare; then the group inside the numerator,
    # bracketed, which is the only substitution that lands inside a product.
    num, den = substitute(num, False), substitute(den, False)
    return substitute(num, True), substitute(den, True), ""


def _page_quotient_sides(page, number, where):
    """One numbered equation of a documentation page as ``(numerator terms,
    denominator terms)``, in the report's notation."""
    found = re.search(r"\$([^$]*?)\\q?quad ?\(%s\)\$" % re.escape(number), page)
    if not found:
        return None, None, f"{where} publishes no equation ({number})"
    text, why = _page_latex(found.group(1))
    if why:
        return None, None, f"equation ({number}) of {where} is written with {why}"
    num, den, why = _as_quotient(text, f"equation ({number}) of {where}")
    if why:
        return None, None, why
    return _signed_terms(num), _signed_terms(den), ""


def _recomposes_to(lines, page, number, where):
    """The printed parts, put back together, against the published equation."""
    fails = []
    num, den, why = _recompose(lines)
    if why:
        return [f"{where}: {why}"]
    want_num, want_den, why = _page_quotient_sides(page, number, where)
    if why:
        return [f"{where}: {why}"]
    for side, got, want in (("numerator", _signed_terms(num), want_num),
                            ("denominator", _signed_terms(den), want_den)):
        if sorted(got) != sorted(want):
            missing = [t for t in want if t not in got]
            extra = [t for t in got if t not in want]
            fails.append(
                f"{where}: the named parts recompose to a {side} of {got}; "
                f"equation ({number}) publishes {want}"
                + (f" — missing {missing}" if missing else "")
                + (f", printed instead {extra}" if extra else ""))
    return fails


#: Which equation each moment method's Calculations section prints in named
#: parts, on which page. The pin: the parts substituted back have to give this
#: equation, term for term and sign for sign, on the page's own side of the bar.
_RECOMPOSED = {"oms": ("lem/oms.md", "8"), "bishop": ("lem/bishop.md", "10")}


def _printed_parts(section):
    """The FIRST quotient-in-named-parts a section prints — the published form.

    Read by shape rather than by the sentence below it: the quotient opens the
    run, and every line after it that defines one of its letters belongs to it.
    A section that drops nothing prints the run once and has no sentence, and
    this finds it either way.
    """
    # The part letters come from the module that prints them, so a part renamed
    # there is a part this still finds.
    from xslope.report import MOMENT_PART_SYMBOLS

    maths = [b.notation for b in section.blocks if b.kind == "math"]
    at = next((i for i, n in enumerate(maths) if n.startswith("F = frac{N_S}")),
              None)
    if at is None:
        return []
    out = [maths[at]]
    for notation in maths[at + 1:]:
        name = notation.partition(" = ")[0]
        if name not in ("N_S", "N_v", "N'") and name not in MOMENT_PART_SYMBOLS:
            break
        out.append(notation)
    return out


#: The arithmetic the two moment methods close on: their page's own quotient in
#: named sums, the same quotient with the numbers substituted, and the answer.
_PART_ARITHMETIC = re.compile(
    r"^F = frac\{(?P<top>[^{}]+)\}\{(?P<bottom>[^{}]+)\} = "
    r"frac\{(?P<num>[\d.eE+-]+)\}\{(?P<den>[\d.eE+-]+)\} = (?P<fs>[\d.]+)$")


def _part_arithmetic(section):
    """``(match, {name: value})`` for a moment method's closing arithmetic, or
    ``(None, {})`` — the quotient in named sums with its numbers, and every named
    sum the section printed a value for."""
    values = {}
    found = None
    for block in section.blocks:
        if block.kind != "math":
            continue
        got = _PART_ARITHMETIC.match(block.notation)
        if got:
            found = got
            continue
        name, sep, body = block.notation.partition(" = ")
        if sep and re.fullmatch(r"-?[\d.]+(?:e[+-]?\d+)?", body):
            values[name] = float(body)
    return found, values


def _named_sums(side):
    """``[(sign, name), ...]`` for one side of a quotient written in named sums."""
    out = []
    for piece in re.split(r"\s+([+−])\s+", " " + side.strip()):
        piece = piece.strip()
        if not piece:
            continue
        if piece in "+−":
            out.append(piece)
        elif out and out[-1] in "+−":
            out[-1] = (-1 if out[-1] == "−" else +1, piece)
        else:
            out.append((+1, piece))
    return [p for p in out if isinstance(p, tuple)]


def test_the_evaluated_equation_introduces_its_terms():
    """The two moment methods evaluate the equation they printed, and every named
    sum they divide is one the equations above it define and print a value for.

    Each prints its page's own quotient in named sums, this model's reduction of
    it, and then those same sums evaluated. Three things have to hold between the
    equation and its arithmetic, and each of them was really broken once.

    Nothing enters the quotient between the equation and its evaluation, and
    nothing leaves: the passive reinforcement moment used to arrive in the sums
    without appearing in any equation above them, standing for about 1700 of a
    20810 numerator while the sentence directly over it said the reinforcement
    terms were zero on every slice. Every sum the quotient names carries a printed
    value, so a letter is never divided without a number. And the values divide to
    the factor of safety the solver reported — which is what makes them the
    working behind it rather than a second calculation beside it.
    """
    fails = []
    from xslope.columns import format_fs
    from xslope.report import _closes

    models = (("the sample model", REINF_XLSX),
              ("the passive-reinforcement model", PASSIVE_XLSX),
              ("the tension-crack model", TENSION_XLSX))
    exercised = False
    tolerance = 10 ** -len(format_fs(1.0).split(".")[-1])

    for label, xlsx in models:
        for method in _RECOMPOSED:
            report, bundle = _calc_report(method, xlsx=xlsx)
            section = _calc_section(report) if report is not None else None
            if section is None:
                fails.append(f"{label} under {method}: no calculation")
                continue
            where = f"{label} under {method}"
            maths = [b.notation for b in section.blocks if b.kind == "math"]
            got, values = _part_arithmetic(section)
            if got is None:
                fails.append(f"{where}: the section closes on no quotient in "
                             f"named sums: {maths}")
                continue
            reduced = [n for n in maths if n.startswith("F = frac{N_S")]
            used = _named_sums(got["top"]) + _named_sums(got["bottom"])
            if "N_P" in got["top"] or "N_H" in got["top"]:
                exercised = True
            # The equation and its arithmetic name the same sums, in the same
            # places: the reduced form directly above is the one evaluated.
            want = (_named_sums(_as_quotient(reduced[-1], where)[0])
                    + _named_sums(_as_quotient(reduced[-1], where)[1])
                    if reduced else [])
            if sorted(used) != sorted(want):
                fails.append(f"{where}: the arithmetic divides {used} and the "
                             f"equation above it reads {want}")
            for _sign, name in used:
                if not any(n.startswith(f"{name} = ") for n in maths):
                    fails.append(f"{where}: the arithmetic uses {name!r} and no "
                                 f"equation above it defines it: {maths}")
                if name not in values:
                    fails.append(f"{where}: {name!r} is divided and never "
                                 f"evaluated: {sorted(values)}")
            # And the printed values are the printed operands, which are the
            # solver's factor of safety.
            top = sum(sign * values.get(name, 0.0) for sign, name in
                      _named_sums(got["top"]))
            bottom = sum(sign * values.get(name, 0.0) for sign, name in
                         _named_sums(got["bottom"]))
            fs = float(bundle["results"]["FS"])
            for said, computed in (("numerator", (got["num"], top)),
                                   ("denominator", (got["den"], bottom))):
                printed, formed = float(computed[0]), computed[1]
                if abs(printed - formed) > abs(formed) * 1e-5:
                    fails.append(f"{where}: the {said} is printed as {printed} "
                                 f"and its parts come to {formed}")
            if not bottom or abs(top / bottom - fs) > tolerance:
                fails.append(f"{where}: the printed sums give "
                             f"{top / bottom if bottom else float('nan')}, the "
                             f"solver {fs}")

    if not exercised:
        fails.append("no model put passive support in an evaluated equation, "
                     "so the case this is here for was not reached")

    # The mutations, on the section this was written against. A part dropped from
    # the arithmetic, and a value nudged in its last printed digit, both have to
    # be caught — the first by the equation above it, the second by the division.
    report, bundle = _calc_report("bishop", xlsx=PASSIVE_XLSX)
    got, values = _part_arithmetic(_calc_section(report))
    if got is None or "N_P" not in got["top"]:
        return fails + ["the passive model prints no N_P part, so the mutations "
                        "test nothing"]
    dropped = _named_sums(got["top"].replace(" + N_P", ""))
    if dropped == _named_sums(got["top"]):
        fails.append("the mutation dropped no part, so it tests nothing")
    elif sorted(dropped + _named_sums(got["bottom"])) == sorted(
            _named_sums(got["top"]) + _named_sums(got["bottom"])):
        fails.append("a part dropped from the arithmetic still matched the "
                     "equation above it")
    nudged = dict(values, N_S=values["N_S"] * 1.01)
    top = sum(sign * nudged[name] for sign, name in _named_sums(got["top"]))
    bottom = sum(sign * nudged[name] for sign, name in _named_sums(got["bottom"]))
    if abs(top / bottom - float(bundle["results"]["FS"])) <= tolerance:
        fails.append("a named sum moved by a hundredth still reproduced the "
                     "factor of safety; the division cannot fail")
    return fails


def _calc_of(method, xlsx):
    """``(slope_data, calc)`` — one method's worked calculation on one model,
    off the same solved frame the section is built from."""
    from xslope.fileio import load_slope_data
    from xslope.report import _calculation

    report, bundle = _calc_report(method, xlsx=xlsx)
    if report is None:
        return None, None
    slope_data = load_slope_data(xlsx)
    calc, _note = _calculation(slope_data, bundle, method)
    return slope_data, calc


def _names_an_equation(blocks, number):
    """Every sentence of ``blocks`` that turns the equation above into numbers,
    and whether it names equation ``number``."""
    return [(b.text, f"equation ({number})" in b.text) for b in blocks
            if b.kind == "prose" and "can be evaluated with the solved "
                                     "values" in b.text]


def test_the_general_arms_name_no_equation():
    """Where a section prints the general moment arms, its arithmetic names no
    equation.

    A moment method whose page's own named sums do not return the solution is
    shown the same equilibrium in the general moment arms, under a sentence that
    says exactly that: the named sums of equation (10) are not the working behind
    this factor of safety. Two pages further on, the arithmetic said "equation
    (10) can be evaluated with the solved values" — the equation the section had
    just set aside — and the reader was left with one number and two
    contradictory claims about where it came from. The arms are numbered (8a) on
    one page and unnumbered on the other, so the sentence there names no equation
    at all.

    The axially reinforced model reaches that branch under Bishop: its passive
    support mobilizes with the soil and stands inside the base normal, in a group
    of vertical forces equation (8) writes for the active case alone.
    """
    fails = []
    from xslope.report import (EVALUATED_EQUATION, _quotient_close,
                               _unit_labels)

    number = EVALUATED_EQUATION["bishop"]
    report, _bundle = _calc_report("bishop", xlsx=AXIAL_XLSX)
    section = _calc_section(report) if report is not None else None
    if section is None:
        return ["the axially reinforced model produced no Bishop calculation, "
                "so the arms branch was not reached"]
    if not any(b.kind == "prose" and "do not return the solution" in b.text
               for b in section.blocks):
        return ["the axially reinforced model no longer prints the general "
                "moment arms under Bishop, so nothing here tests them"]
    said = _names_an_equation(section.blocks, number)
    if not said:
        fails.append("the arithmetic is introduced by no sentence at all")
    for text, names_it in said:
        if names_it:
            fails.append(f"the section sets equation ({number}) aside and then "
                         f"names it as the one evaluated: {text!r}")
        if "the balance can be evaluated" not in text:
            fails.append(f"the arithmetic names neither an equation nor the "
                         f"balance it evaluates: {text!r}")

    # The mutation: the close named its method's equation whatever the section
    # above it printed. Everything else about the calculation is held fixed —
    # only the flag that says the arms are being printed is cleared.
    slope_data, calc = _calc_of("bishop", AXIAL_XLSX)
    if calc is None or not calc.get("arms"):
        return fails + ["the Bishop calculation of the axially reinforced model "
                        "no longer falls to the moment arms; the mutation tests "
                        "nothing"]
    blocks = _quotient_close(dict(calc, arms=""), 9, "mark",
                             _unit_labels(slope_data))
    mutated = _names_an_equation(blocks, number)
    if not mutated or not any(names_it for _text, names_it in mutated):
        fails.append("the close named no equation even with the arms flag "
                     "cleared, so nothing here would catch it naming one")
    return fails


def test_a_line_loads_moment_is_the_pages():
    """A line load's driving moment is the same number in the page's symbols and
    in the arms the solver formed it in.

    The load is applied as a magnitude L at an angle δ from the horizontal, and
    every equation that carries it is written in L cos δ and L sin δ. It is not
    stored that way: it is folded into the distributed-load convention, and the
    sign of its vertical component goes into the stored angle — so LL·cos(ll_β)
    is the NEGATIVE of the page's L sin δ. Read off the stored arrays without
    that mapping, the page's moment came out with one component reversed, the
    named sums of equation (8) returned 1.851 where the solver returned 1.834,
    and every model with a line load was refused its own page's working and shown
    the general moment arms instead.

    Two things are pinned. The registry's page term and the solver's arms give
    one number on the model that carries a line load, which is what makes them
    the same equilibrium; and the section built on it prints the page's named
    sums and divides them to the solver's factor of safety.
    """
    import numpy as np
    fails = []
    from xslope.columns import format_fs
    from xslope.report import FORCE_TERMS, _Calc, _calc_arrays, _calculation

    for where, xlsx in (("the axially reinforced model", AXIAL_XLSX),
                        ("its right-facing twin",
                         os.path.join(_REPO, "docs", "inputs", "slope",
                                      "xslope_nail_axial_rface.xlsx"))):
        report, bundle = _calc_report("oms", xlsx=xlsx)
        if report is None:
            fails.append(f"{where}: no solution to read a line load out of")
            continue
        df = bundle["slice_df"]
        if not np.any(np.abs(df["lload"].values.astype(float)) > 0):
            fails.append(f"{where} carries no line load; this tests nothing")
            continue
        A = _calc_arrays(df)
        C = _Calc(df, A, bool(df.attrs.get(
            "right_facing", df["y_lb"].iat[0] > df["y_rb"].iat[-1])))
        line = next(t for t in FORCE_TERMS if t.key == "L")
        page = sum(t.sign * np.asarray(t.values(C), dtype=float)
                   for t in line.page_drv)
        arms = sum(t.sign * np.asarray(t.values(C), dtype=float)
                   for t in line.moment_drv)
        # Each is the moment the load makes about the center of rotation, signed
        # as it enters the driving side, so the two agree slice by slice.
        if not np.allclose(page, arms, rtol=1e-9, atol=1e-9):
            fails.append(f"{where}: the page's line-load moment comes to "
                         f"{float(np.sum(page)):.4f} and the arms the solver "
                         f"formed it in to {float(np.sum(arms)):.4f}")
        # Both are the transcription of terms nothing else evaluates, so the
        # symbols they print are checked to be the two components of the load.
        printed = " ".join(t.symbol for t in line.page_drv)
        for symbol in ("L cos δ", "L sin δ"):
            if symbol not in printed:
                fails.append(f"{where}: the page's driving term prints "
                             f"{printed!r}, without {symbol}")

        section = _calc_section(report)
        if section is None:
            fails.append(f"{where}: the Ordinary Method's calculation is not "
                         f"printed on a model with a line load")
            continue
        if any(b.kind == "prose" and "do not return the solution" in b.text
               for b in section.blocks):
            fails.append(f"{where}: the named sums of equation (8) are refused "
                         f"on a model whose only feature they cannot carry is "
                         f"the line load")
        got, values = _part_arithmetic(section)
        if got is None:
            fails.append(f"{where}: the section closes on no quotient in named "
                         f"sums")
            continue
        if not any(name == "D_L" for _sign, name in _named_sums(got["bottom"])):
            fails.append(f"{where}: the driving side is {got['bottom']!r} and "
                         f"carries no line-load sum")
        top = sum(sign * values.get(name, 0.0)
                  for sign, name in _named_sums(got["top"]))
        bottom = sum(sign * values.get(name, 0.0)
                     for sign, name in _named_sums(got["bottom"]))
        fs = float(bundle["results"]["FS"])
        tolerance = 10 ** -len(format_fs(1.0).split(".")[-1])
        if not bottom or abs(top / bottom - fs) > tolerance:
            fails.append(f"{where}: the printed sums give "
                         f"{top / bottom if bottom else float('nan')}, the "
                         f"solver {fs}")

    # The mutation: the term mis-signed as it was, and the gate that refused the
    # page's sums has to refuse them again.
    import dataclasses

    import xslope.report as report_mod
    from xslope.fileio import load_slope_data

    line = next(t for t in FORCE_TERMS if t.key == "L")
    flipped = tuple(dataclasses.replace(t, sign=-t.sign) for t in line.page_drv)
    saved = report_mod.FORCE_TERMS
    report_mod.FORCE_TERMS = tuple(
        dataclasses.replace(t, page_drv=flipped) if t.key == "L" else t
        for t in saved)
    try:
        report, bundle = _calc_report("oms", xlsx=AXIAL_XLSX)
        calc, _note = _calculation(load_slope_data(AXIAL_XLSX), bundle, "oms")
    finally:
        report_mod.FORCE_TERMS = saved
    if calc is not None and calc.get("parts") and not calc.get("arms"):
        fails.append("a mis-signed line-load moment still returned the "
                     "solution; the gate cannot fail")
    return fails


def test_the_moment_quotient_recomposes():
    """The two moment methods print their own page's factor-of-safety equation.

    Equation (8) of the Ordinary Method of Slices derivation and equation (10)
    of Bishop's each run several times the width of a page, and what went onto
    the page instead was a reconstruction: Bishop's section printed a moment
    balance rewritten in another derivation's moment arms, which is not Bishop's
    equation and does not carry its number.

    Each now prints its own equation as a quotient of named sums, one part per
    force. This is what holds those parts to the page: they are substituted back
    into the quotient and the result compared with the published equation, term
    for term and sign for sign, numerator against numerator. A term dropped from
    a part, or a part moved from one side of the bar to the other, cannot
    survive it.
    """
    fails = []

    for method, (path, number) in _RECOMPOSED.items():
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            fails.append(f"{method}: no calculation to recompose")
            continue
        full = _printed_parts(section)
        if len(full) < 3:
            printed = [b.notation for b in section.blocks if b.kind == "math"]
            fails.append(f"{method}: the section prints no quotient in named "
                         f"parts: {printed}")
            continue
        with open(os.path.join(_REPO, "docs", path), encoding="utf-8") as f:
            page = f.read()
        fails += _recomposes_to(full, page, number, path)

    # The mutations: the comparison has to catch a part that lost a term and a
    # part that changed sides. Run on Bishop's own printed parts, so what is
    # mutated is the thing the check reads.
    report, _bundle = _calc_report("bishop")
    section = _calc_section(report) if report is not None else None
    if section is None:
        return fails + ["there is no Bishop calculation to mutate"]
    full = _printed_parts(section)
    with open(os.path.join(_REPO, "docs", "lem", "bishop.md"),
              encoding="utf-8") as f:
        page = f.read()

    dropped = [re.sub(r" − frac\{1\}\{R\}·sum\{D sin β·a_dy\}", "", line)
               for line in full]
    if dropped == full:
        fails.append("the mutation removed nothing, so it tests nothing")
    elif not _recomposes_to(dropped, page, "10", "the mutation"):
        fails.append("a term dropped from a named part still recomposed to "
                     "equation (10)")

    moved = [line.replace("F = frac{N_S}{D_W", "F = frac{N_S − D_P}{D_W")
             .replace(" − D_P − D_H", " − D_H") for line in full]
    if moved == full:
        fails.append("the mutation moved nothing, so it tests nothing")
    elif not _recomposes_to(moved, page, "10", "the mutation"):
        fails.append("a part moved across the fraction bar still recomposed to "
                     "equation (10)")
    return fails


#: The four methods whose transcription is one printed equation per published
#: one, and where those are published: the page, the numbers in the order the
#: section prints them, and the identities the page itself states between its
#: letters and the report's. Each identity is checked to be ON the page before it
#: is used, so a substitution here cannot excuse a notation the page never
#: introduced.
#:
#: The other three are pinned elsewhere and for the same purpose: the two moment
#: methods by recomposing their named parts (test_the_moment_quotient_recomposes)
#: and Spencer's two force sums against its own page (test_spencer_force_sums).
_FULL_FORMS = {
    "janbu": ("lem/janbu.md", ("7",),
              (("sumNsinα", "sumN'sinα+sumuΔlsinα",
                r"N = N' + u\,\Delta\ell"),)),
    "corps": ("lem/force_eq.md", ("6", "7"), ()),
    "lowe": ("lem/force_eq.md", ("6", "7"), ()),
    "mprice": ("lem/force_eq.md", ("6", "7"), ()),
}


def _canonical(text):
    """One equation reduced to what it says: the symbols and the operators, with
    the grouping and the spacing taken out."""
    out, why = _page_latex(text)
    if why:
        return None, why
    return "".join(c for c in out if c not in _GROUPING), ""


def test_the_full_forms_match_their_pages():
    """Every method prints its own page's equation, symbol for symbol.

    A section that transcribes the published equation is making a claim about
    another document, and the claim is checkable. Each of the four methods that
    print one equation per published one has its printed form reduced to its
    symbols and operators and held against the same reduction of the equation on
    the page it names — the whole equation, not a list of the letters in it.

    Janbu's page writes the total base normal N and states, in the sentence under
    equation (7), that N = N' + u·Δl; the report writes it out, as the two sums
    the one sum over N distributes to — which is how the force-equilibrium page
    writes the same thrust in its own equation (12), and what lets the pore-water
    term leave the reduced form on a model with no pore pressure. That identity
    is the one substitution made here, and the page is required to carry it.
    """
    fails = []

    for method, (path, numbers, identities) in _FULL_FORMS.items():
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            fails.append(f"{method}: no calculation to read the transcription of")
            continue
        full, _model_own = _transcription_split(section)
        if len(full) != len(numbers):
            fails.append(f"{method}: the section transcribes {full}, and "
                         f"{path} is cited for {len(numbers)} equation(s)")
            continue
        with open(os.path.join(_REPO, "docs", path), encoding="utf-8") as f:
            page = f.read()
        for number, printed in zip(numbers, full):
            found = re.search(r"\$([^$]*?)\\q?quad ?\(%s\)\$" % re.escape(number),
                              page)
            if not found:
                fails.append(f"{method}: {path} publishes no equation ({number})")
                continue
            want, why = _canonical(found.group(1))
            if why:
                fails.append(f"{method}: equation ({number}) of {path} is "
                             f"written with {why}")
                continue
            for was, becomes, stated in identities:
                if stated not in page:
                    fails.append(f"{method}: {path} does not state {stated!r}, "
                                 f"so {was!r} may not be read as {becomes!r}")
                    continue
                want = want.replace(was, becomes)
            got, why = _canonical(printed)
            if why:
                fails.append(f"{method}: the section prints {printed!r}, which "
                             f"carries {why}")
            elif got != want:
                fails.append(f"{method}: the section prints equation ({number}) "
                             f"as {got!r}; {path} publishes {want!r}")

    # The mutation: a printed form that has drifted from its page has to be
    # caught. One term dropped from the horizontal balance Janbu publishes.
    report, _bundle = _calc_report("janbu")
    section = _calc_section(report)
    full, _model_own = _transcription_split(section)
    with open(os.path.join(_REPO, "docs", "lem", "janbu.md"),
              encoding="utf-8") as f:
        page = f.read()
    found = re.search(r"\$([^$]*?)\\q?quad ?\(7\)\$", page)
    want, _why = _canonical(found.group(1))
    want = want.replace("sumNsinα", "sumN'sinα+sumuΔlsinα")
    drifted = full[0].replace(" + sum{kW}", "")
    if drifted == full[0]:
        fails.append("the mutation dropped nothing, so it tests nothing")
    elif _canonical(drifted)[0] == want:
        fails.append("a term dropped from the printed equation still matched "
                     "the page")
    return fails


#: The whole-mass balance the three marching methods print under their march, and
#: where it is published: equation (12) of the force-equilibrium derivation.
_WHOLE_MASS_PAGE, _WHOLE_MASS_NUMBER = "lem/force_eq.md", "12"

#: How the sentence below that equation names it, which is what marks the pair.
_WHOLE_MASS_REDUCTION = "so equation (12) reduces to"


def _whole_mass_pair(section):
    """``(published, sentence, reduced)`` for the whole-mass balance a marching
    method's section prints, or empties.

    A marching section carries two published-then-reduced pairs — equations (6)
    and (7) in its preamble and this one below them — so the pair is found by the
    sentence that names the equation, not by taking the first.
    """
    for sentence, published, reduced in _transcription_pairs(section):
        if _WHOLE_MASS_REDUCTION in sentence:
            return published, sentence, reduced
    return [], "", []


def _whole_mass_page_form():
    """Equation (12) of the force-equilibrium page as ``(numerator, denominator)``
    atoms, or ``(None, None, why)``."""
    with open(os.path.join(_REPO, "docs", _WHOLE_MASS_PAGE),
              encoding="utf-8") as f:
        page = f.read()
    found = re.search(r"\$([^$]*?)\\q?quad ?\(%s\)\$" % _WHOLE_MASS_NUMBER, page)
    if not found:
        return None, None, (f"{_WHOLE_MASS_PAGE} publishes no equation "
                            f"({_WHOLE_MASS_NUMBER})")
    text, why = _page_latex(found.group(1))
    if why:
        return None, None, (f"equation ({_WHOLE_MASS_NUMBER}) of "
                            f"{_WHOLE_MASS_PAGE} is written with {why}")
    num, den, why = _as_quotient(text, f"equation ({_WHOLE_MASS_NUMBER})")
    if why:
        return None, None, why
    return _balance_atoms(num), _balance_atoms(den), ""


def test_the_whole_mass_balance_is_published_then_reduced():
    """Corps of Engineers, Lowe & Karafiath and Morgenstern-Price print equation
    (12) in full before they print this model's, and the full form is the page's.

    None of the three solves that equation — each marches equations (6) and (7)
    slice by slice — but each prints the balance the march sums to, cites it as
    equation (12) of the force-equilibrium derivation, and evaluates it. Printed
    reduced-only under that number it named a published equation and showed a
    shorter one: the seismic force, the tension-crack water force and the support
    terms the page carries were gone with nothing to say they exist, on a model
    that happens to have none of them. Janbu's section prints the identical
    balance under its own page's number, in full and then reduced, and this is
    the same discipline on the same registry.

    Three claims are checked. The full form carries every contribution the
    registry declares for the equation, so it is the published form and not a
    reduction wearing its number. It says what the page says: term for term and
    sign for sign, numerator against numerator. And the sentence below it
    accounts for every term that then goes.

    The comparison with the page is on the terms rather than on the string. The
    same registry-assembled form is printed under Janbu's equation (7), and the
    two pages that publish this one balance write it differently — one sum over
    the bracket against a sum per term, and the surface load ahead of the support
    forces on one page and behind them on the other. Distributing the sums and
    comparing the signed terms as a set survives both and still catches a term
    dropped, a sign flipped or a symbol renamed, which is what a transcription
    claims.
    """
    fails = []
    from xslope.report import (FORCE_TERMS, NotApplicable, WHOLE_MASS_CONSUMERS,
                               WHOLE_MASS_BALANCE_METHODS)

    want_num, want_den, why = _whole_mass_page_form()
    if why:
        return [why]

    declared = []
    for term in FORCE_TERMS:
        for consumer in WHOLE_MASS_CONSUMERS:
            got = getattr(term, consumer)
            declared += [(term, c.symbol) for c in
                         (got.published if isinstance(got, NotApplicable)
                          else got)]

    printed_full = {}
    for method in WHOLE_MASS_BALANCE_METHODS:
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            fails.append(f"{method}: no calculation to read the balance of")
            continue
        published, sentence, reduced = _whole_mass_pair(section)
        if len(published) != 1 or len(reduced) != 1:
            maths = [b.notation for b in section.blocks if b.kind == "math"]
            fails.append(f"{method}: the section prints no equation (12) "
                         f"published-then-reduced; its equations are {maths}")
            continue
        printed_full[method] = published[0]
        fails += _whole_mass_is_the_page(published[0], want_num, want_den, method)
        # The published form, not a longer reduction: every term the registry
        # declares for the equation stands in it.
        for term, symbol in declared:
            if symbol not in published[0]:
                fails.append(f"{method}: equation (12) is printed without "
                             f"{symbol!r}, which the registry declares for it")
        # And the sentence below it accounts for everything that then went.
        gone = [symbol for _term, symbol in declared
                if symbol not in reduced[0]]
        for term, symbol in declared:
            if symbol in reduced[0] or symbol not in gone:
                continue
            if term.feature not in sentence and symbol not in sentence:
                fails.append(f"{method}: {symbol!r} is dropped from equation "
                             f"(12) and the sentence accounts for neither it "
                             f"nor the {term.feature}: {sentence!r}")
        for term in FORCE_TERMS:
            if term.feature and term.feature in sentence:
                still = [s for t, s in declared
                         if t is term and s in reduced[0]]
                if still:
                    fails.append(f"{method}: the sentence says the model "
                                 f"carries no {term.feature}, and the reduced "
                                 f"equation (12) prints {still}")

    if not printed_full:
        return fails + ["no section printed equation (12) to check"]

    # The mutations, run on the form the check reads. A term dropped from the
    # published equation has to stop matching the page, and the reduced form
    # printed in its place has to stop being the published form.
    method, full = sorted(printed_full.items())[0]
    dropped = full.replace(" + sum{kW}", "")
    if dropped == full:
        fails.append("the mutation dropped nothing, so it tests nothing")
    elif not _whole_mass_is_the_page(dropped, want_num, want_den, "the mutation"):
        fails.append("a term dropped from the printed equation (12) still "
                     "matched the page")

    report, _bundle = _calc_report(method)
    reduced = _whole_mass_pair(_calc_section(report))[2][0]
    if reduced == full:
        fails.append(f"{method}: its model drops nothing from equation (12), "
                     f"so printing the reduced form in place of the published "
                     f"one tests nothing")
    elif all(symbol in reduced for _term, symbol in declared):
        fails.append("the reduced equation printed as the published form "
                     "carried every term the registry declares, so printing it "
                     "there would not have been caught")
    return fails


def _shared_slice_counts(rows):
    """What a section's engine inputs claim the slice count of the analysis is."""
    return [value for label, value in rows if label == "Slices"]


def test_the_shared_slice_count():
    """The number of slices stands among the shared engine inputs only where the
    featured methods share one.

    Run a search per method and each method finds its own critical surface, cut
    into its own number of slices. The section's Analysis Inputs printed one of
    those counts — the first method's — as an input of the analysis, where it was
    false of every other method the section documented. It is a property of a
    surface, and where the surfaces differ there is no shared value to print.

    So: the row goes where the counts differ, and every method's own section
    states the count its own sums were taken over, which is where a reader needs
    it. Where the counts agree — one method, or several on one specified surface
    — there is a shared value and the row stands.
    """
    fails = []

    # Per-method surfaces: two methods, two frames, two counts.
    report, counts = _per_method_report([("oms", 16), ("bishop", 12)])
    sizes = {n for _method, n in counts}
    if len(sizes) < 2:
        return [f"the two frames came out the same size ({sizes}), so a report "
                f"of methods that do not share a slice count was not built"]

    shared = _shared_slice_counts(_analysis_inputs(report))
    if shared:
        fails.append(f"the featured methods carry {dict(counts)} slices and the "
                     f"engine inputs print a shared count of {shared}")
    for method, n in counts:
        node = _method_section(report, method)
        if node is None:
            fails.append(f"{method}: the report carries no section for it")
            continue
        prose = [b.text for _lvl, sub in node.walk() for b in sub.blocks
                 if b.kind == "prose"]
        if not any(f"the {n} slices" in t for t in prose):
            fails.append(f"{method}: its section never states the {n} slices "
                         f"its sums were taken over")
        wrong = [other for _m, other in counts
                 if other != n and any(f"the {other} slices" in t
                                       for t in prose)]
        if wrong:
            fails.append(f"{method}: its section is solved on {n} slices and "
                         f"states {wrong}")

    # One surface, and the count is an input of the analysis like any other.
    same, same_counts = _per_method_report([("oms", 15), ("bishop", 15)])
    sizes = {n for _method, n in same_counts}
    if len(sizes) != 1:
        fails.append(f"the two frames were cut to one size and came out "
                     f"{same_counts}, so the shared case was not built")
    else:
        want = [str(sizes.pop())]
        got = _shared_slice_counts(_analysis_inputs(same))
        if got != want:
            fails.append(f"the featured methods share {want} slices and the "
                         f"engine inputs print {got}")
    one = _calc_report("oms")[0]
    if one is not None:
        got = _shared_slice_counts(_analysis_inputs(one))
        if len(got) != 1:
            fails.append(f"a report of one method prints {got} for its slice "
                         f"count")

    # The mutation: the row put back on a report whose methods do not share a
    # count, read by the same reader the check above uses.
    mutated = _analysis_inputs(report) + [("Slices", str(counts[0][1]))]
    if not _shared_slice_counts(mutated):
        fails.append("a shared slice count restored to the engine inputs was "
                     "not read, so its absence above proves nothing")
    return fails


def test_the_maximum_depth_is_an_input_at_any_elevation():
    """The maximum surface depth stands among the engine inputs whenever the
    model declares one — elevation zero included.

    It is an ELEVATION, and zero is a place. Read as a truth value, a model whose
    depth is 0.0 stated no maximum depth here while the model figure beside it
    drew the line and the Project Definition named it: three statements about one
    input, one of them missing. A model that declares none — a section given by
    polygons — still gets no row.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from xslope.report import build_report

    def rows_for(model):
        with tempfile.TemporaryDirectory() as tmp:
            with contextlib.redirect_stdout(io.StringIO()):
                slope_data, solutions = _solved()
                report = build_report(model, solutions,
                                      {"input_path": REINF_XLSX,
                                       "method": "bishop", "pd_figure": False,
                                       "lem_inputs_figure": False,
                                       "lem_search_figure": False,
                                       "lem_solution_figure": False,
                                       "lem_slice_key": False,
                                       "lem_slice_table": False,
                                       "lem_calculations": False}, tmp)
        return dict(_analysis_inputs(report))

    label = "Maximum surface depth (elevation)"
    slope_data, _solutions = _solved()
    if slope_data.get("max_depth") != 0.0:
        fails.append(f"the sample's maximum depth is "
                     f"{slope_data.get('max_depth')!r}, so elevation zero — the "
                     f"case that was dropped — is not under test")
    got = rows_for(slope_data)
    if label not in got:
        fails.append(f"a model whose maximum depth is "
                     f"{slope_data.get('max_depth')!r} states no maximum depth: "
                     f"{sorted(got)}")
    elif got[label] != "0":
        fails.append(f"the maximum depth is printed {got[label]!r}, not the "
                     f"elevation the model declares")
    # Declared deeper, and it is the model's number that is printed.
    deep = rows_for(dict(slope_data, max_depth=-15.0))
    if deep.get(label) != "-15":
        fails.append(f"a maximum depth of -15 is printed {deep.get(label)!r}")
    # Not declared at all, and there is no row to print.
    none = rows_for(dict(slope_data, max_depth=None))
    if label in none:
        fails.append(f"a model that declares no maximum depth is given one: "
                     f"{none[label]!r}")
    return fails


def _whole_mass_is_the_page(printed, want_num, want_den, where):
    """One printed whole-mass balance held against equation (12) of the page."""
    num, den, why = _as_quotient(printed, f"{where}: the whole-mass balance")
    if why:
        return [why]
    fails = []
    for got, want, side in ((_balance_atoms(num), want_num, "numerator"),
                            (_balance_atoms(den), want_den, "denominator")):
        if got != want:
            fails.append(f"{where}: the {side} of equation (12) is printed as "
                         f"{got}; {_WHOLE_MASS_PAGE} publishes {want}")
    return fails


#: The base normal, by the method that solves it alongside its quotient: the page
#: and the number the page publishes it under. Janbu's page names the denominator
#: m_α in its equation (1) and writes equation (6) over it; the report prints both
#: lines, so the two are joined before the comparison.
#: ``method -> (page, number, denominator)``, where ``denominator`` is the
#: equation that page names the base-normal denominator in, for a page that names
#: it rather than writing it out. Janbu's equation (1) introduces m_α beside the
#: base normal it divides; the report prints that definition on a line of its own
#: and it is held against the equation that introduces it.
#:
#: Janbu's is the only one. Bishop's section prints no base normal at all: what
#: it solves is its equation (10), written in N_v — the GROUP of vertical forces
#: the base normal is formed from — so N' appears in none of its equations, and
#: an equation for a quantity nothing on the page uses is a page a reader has no
#: use for. Janbu's equation (7) is written in ΣN'·sin α, so its section gives
#: N' where the reader meets that sum. :func:`test_bishop_prints_no_base_normal`
#: is the other half of this: it holds Bishop's section to printing none.
_NORMAL_FORMS = {
    "janbu": ("lem/janbu.md", "6", "1"),
}

#: The strength mobilized on the slice base, as the two pages and the report each
#: write it. One product — c·Δl·sin α divided by F — with the fraction bar in
#: three places: Bishop's page divides c·Δl and multiplies by sin α, Janbu's
#: takes 1/F outside the product, and the report divides the whole product. They
#: are the same term, and where the bar is set is not what a transcription is
#: checked on; every other difference still is, because nothing but these three
#: spellings collapses.
_MOBILIZED_COHESION = re.compile(
    r"frac\{1\}\{F\}\s*c\s*Δl\s*sin\s*α"
    r"|frac\{c\s*Δl\}\{F\}\s*sin\s*α"
    r"|frac\{c·Δl·sin α\}\{F\}")


#: How a page spells a symbol the report prints in Unicode, for the one place a
#: symbol has to be FOUND in a page rather than compared with one.
_MATH_NAMES = {"m_α": r"m_\alpha"}


def _normal_canonical(text):
    """One base-normal equation reduced to what it says, with the mobilized
    cohesion written one way."""
    out, why = _page_latex(text)
    if why:
        return None, why
    out = _MOBILIZED_COHESION.sub("MOB", out)
    return "".join(c for c in out if c not in _GROUPING), ""


def test_the_normal_force_is_published_then_reduced():
    """Janbu's section prints the base normal its page publishes, then this
    model's reduction of it.

    It cites its derivation's own number for it — equation (6) — and used to
    print, under that number, an equation with the pore pressure, the
    reinforcement, the pile force and the line load already taken out. A reader
    following the reference met a shorter equation than the one it named. So the
    base normal is held to the same discipline as the equation the method solves:
    the page's form first, carrying every vertical force a slice can take, then
    the sentence saying what this model does not carry, then what is left.

    Three things are required. The published form is the page's, symbol for
    symbol. The reduction really is a reduction — every term it drops is one the
    published form had. And the sentence between them accounts for every one.
    """
    fails = []
    from xslope.report import FORCE_TERMS, NotApplicable

    for method, (path, number, denominator) in sorted(_NORMAL_FORMS.items()):
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            fails.append(f"{method}: no calculation to read the base normal of")
            continue
        sentence, full, reduced = _normal_force_pair(section)
        if not full:
            fails.append(f"{method}: the section prints no base normal above a "
                         f"sentence reducing it: "
                         f"{[b.notation for b in section.blocks if b.kind == 'math']}")
            continue
        if not reduced:
            fails.append(f"{method}: the base normal is transcribed and never "
                         f"reduced to this model")
            continue

        # --- the published form is the page's ---
        with open(os.path.join(_REPO, "docs", path), encoding="utf-8") as f:
            page = f.read()
        # Janbu's m_α is a line of its own here and a numbered equation there;
        # the comparison is against the equation the number names, so the
        # denominator is substituted back before the two are held together.
        def against(source, printed, named):
            """The printed equation reduced to what it says, against the page's."""
            want, why = _normal_canonical(source)
            got, why_got = _normal_canonical(printed)
            if why or why_got:
                return [f"{method}: equation ({named}) is written with "
                        f"{why or why_got}"]
            if got != want:
                return [f"{method}: the section prints equation ({named}) as "
                        f"{got!r}; {path} publishes {want!r}"]
            return []

        found = re.search(r"\$([^$]*?)\\q?quad ?\(%s\)\$" % re.escape(number),
                          page)
        if not found:
            fails.append(f"{method}: {path} publishes no equation ({number})")
        else:
            fails += against(found.group(1), full[0], number)
        # A page that NAMES the denominator rather than writing it out: the
        # report prints the definition on a line of its own, and the equation
        # that introduces it on the page has to carry that definition.
        if denominator:
            if len(full) < 2:
                fails.append(f"{method}: {path} writes the base normal over a "
                             f"named denominator and the section never gives it")
            else:
                # The definition itself, out of the equation that introduces it:
                # that equation carries the base normal as well, and what the
                # report prints on this line is the denominator alone.
                name = full[1].split(" = ", 1)[0]
                found = re.search(
                    r"(%s\s*=\s*.*?)\s*\\q?quad ?\(%s\)"
                    % (re.escape(_MATH_NAMES.get(name, name)),
                       re.escape(denominator)), page)
                if not found:
                    fails.append(f"{method}: equation ({denominator}) of {path} "
                                 f"gives no {name}")
                else:
                    fails += against(found.group(1), full[1], denominator)

        # --- the reduction drops only what the published form had, and the
        # sentence accounts for every one of them ---
        published_terms = _terms_printed(full[0].split("}{", 1)[0] + "}")
        reduced_terms = _terms_printed(reduced[0].split("}{", 1)[0] + "}")
        dropped = [s for s in published_terms if s not in reduced_terms]
        if not dropped:
            fails.append(f"{method}: the reduced base normal drops nothing and "
                         f"is printed under a sentence that says it does: "
                         f"{sentence!r}")
        owner = {}
        for term in FORCE_TERMS:
            got = term.normal
            published = got.published if isinstance(got, NotApplicable) else got
            for contribution in published:
                owner[contribution.symbol] = term
        for symbol in dropped:
            term = owner.get(symbol)
            if term is None:
                fails.append(f"{method}: the base normal drops {symbol!r}, which "
                             f"belongs to no force in the registry")
            elif term.feature not in sentence and symbol not in sentence:
                fails.append(f"{method}: the base normal drops {symbol!r} and "
                             f"the sentence accounts for neither it nor the "
                             f"{term.feature}: {sentence!r}")
        # And nothing the sentence calls absent survives into the reduced form.
        for term in FORCE_TERMS:
            if not term.feature or term.feature not in sentence:
                continue
            carried = [s for s, t in owner.items()
                       if t is term and s in reduced_terms]
            if carried:
                fails.append(f"{method}: the sentence says the model carries no "
                             f"{term.feature}, and the reduced base normal "
                             f"prints {carried}")

    # The mutation: the section that prints the reduced form under the page's
    # number and nothing else — the report the defect was found in. With the
    # published run taken away there is no pair, and the check has to say so.
    report, _bundle = _calc_report("janbu")
    section = _calc_section(report)
    _sentence, full, reduced = _normal_force_pair(section)
    if full == reduced:
        fails.append("Janbu's published and reduced base normals are the same "
                     "equation, so the pair tests nothing")
    if not _is_normal_force(reduced):
        fails.append(f"the reduced base normal is not read as one: {reduced}")
    return fails


def test_bishop_prints_no_base_normal():
    """Bishop's section prints no equation for N'.

    Its equation (10) is written in N_v, the group of vertical forces the base
    normal is formed from, and in nothing else: N' appears in no equation the
    section prints. It printed one anyway — the derivation's equation (8),
    published then reduced, a page and a half of algebra for a quantity the
    section's own equations never carry. What a reader needs about equation (8)
    is where N_v comes from, and the sentence that introduces equation (10) says
    it in words and names the equation.

    N' is still a COLUMN of the slice table, here as in every section, and still
    defined in the nomenclature: this is about the equation, not the symbol.
    """
    fails = []
    report, _bundle = _calc_report("bishop")
    section = _calc_section(report) if report is not None else None
    if section is None:
        fails.append("bishop: no calculation section to read")
        return fails
    printed = [b.notation for b in section.blocks if b.kind == "math"]
    normals = [n for n in printed if n.startswith("N' = ")]
    if normals:
        fails.append(f"bishop: the section still prints a base normal: {normals}")
    # The sentence that names equation (8) is what replaces it, and it says what
    # the equation is for rather than writing it out.
    said = " ".join(b.text for b in section.blocks if b.kind == "prose")
    if "equation (8)" not in said:
        fails.append("bishop: nothing in the section says where N_v comes from")
    if "N_v" not in said:
        fails.append(f"bishop: the section never names N_v: {said[:200]!r}")

    # N' has not gone with the equation: it is still a column of the slice table,
    # still defined in that table's own footnote, and still the quantity the
    # resisting moment is built on. The nomenclature defines the symbols of the
    # EQUATIONS, and N' is no longer one of them; the column's own legend is
    # where a reader of the table meets it, which is where the table is read.
    table = next((b for top in report.sections for _lvl, node in top.walk()
                  for b in node.blocks
                  if b.kind == "table" and b.caption.startswith("Slice data")),
                 None)
    if table is None:
        fails.append("bishop: the section prints no slice table")
        return fails
    if not any(str(h).startswith("N'") for h in table.headers):
        fails.append(f"bishop: the slice table lost its N' column: "
                     f"{table.headers}")
    legend = dict(table.legend or [])
    defined = [term for term in legend if str(term).startswith("N'")]
    if not defined:
        fails.append(f"bishop: the slice table's footnote no longer defines N': "
                     f"{sorted(legend)}")
    uses = [term for term, text in legend.items() if "N'" in str(text)]
    if not uses:
        fails.append("bishop: no column of the slice table is written in N'")

    # And Janbu's, which is written in ΣN'·sin α, still prints one: this removes
    # the equation from the section that does not use it, not from both.
    other, _b = _calc_report("janbu")
    theirs = _calc_section(other) if other is not None else None
    if theirs is None or not [b.notation for b in theirs.blocks
                              if b.kind == "math"
                              and b.notation.startswith("N' = ")]:
        fails.append("janbu: its section no longer prints its own base normal")
    return fails


#: The equations a method's section prints from its OWN documentation page, and
#: the number each is published under there.
#:
#: Morgenstern-Price is the method this exists for. Its march is the
#: force-equilibrium page's, pinned against that page by :data:`_FULL_FORMS`;
#: what defines the METHOD is published on its own page and printed here: the
#: interslice assumption, the force function the solution was solved with, and
#: the moment about the coordinate origin the second condition sums over.
_OWN_PAGE_FORMS = {
    "mprice": ("lem/mprice.md", (
        ("2", "tan θ_j = λ·f(x_j)"),
        ("4", "f(x_j) = sin(π·frac{x_j − x_L}{x_R − x_L})"),
        ("8", "M_O = x_F·F_y − y_F·F_x"),
    )),
}


def _prints_its_page(method, path, page, number, notation, printed):
    """``[]`` where the section prints that equation and it is the page's, or the
    failures saying which of the two it is not."""
    if notation not in printed:
        return [f"{method}: {path} publishes equation ({number}) and the "
                f"section prints no {notation!r}: {printed}"]
    found = re.search(r"\$([^$]*?)\\q?quad ?\(%s\)\$" % re.escape(number), page)
    if not found:
        return [f"{method}: {path} publishes no equation ({number})"]
    want, why = _canonical(found.group(1))
    if why:
        return [f"{method}: equation ({number}) of {path} is written with {why}"]
    got, why = _canonical(notation)
    if why:
        return [f"{method}: the section prints {notation!r}, which carries {why}"]
    if got != want:
        return [f"{method}: the section prints equation ({number}) as {got!r}; "
                f"{path} publishes {want!r}"]
    return []


def test_the_method_prints_its_own_pages_equations():
    """Every equation a section takes from its own method's page is that page's,
    symbol for symbol.

    A method whose derivation publishes the assumption that DEFINES it has to
    print that assumption, not describe it: Morgenstern-Price's interslice
    inclination tan θ_j = λ·f(x_j), the f it was solved with, and the moment
    about the coordinate origin its second condition sums, each held against the
    equation its page numbers it.
    """
    fails = []
    pages = {}
    for method, (path, wanted) in _OWN_PAGE_FORMS.items():
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            fails.append(f"{method}: no calculation to read the equations of")
            continue
        printed = [b.notation for b in section.blocks if b.kind == "math"]
        with open(os.path.join(_REPO, "docs", path), encoding="utf-8") as f:
            pages[method] = page = f.read()
        for number, notation in wanted:
            fails += _prints_its_page(method, path, page, number, notation,
                                      printed)

    # The mutations. A section that stopped printing the assumption, and a page
    # whose display has drifted from what the section prints, both have to be
    # caught — the first is what the check is for, the second is what makes it a
    # pin on the page and not on the string written here.
    path, wanted = _OWN_PAGE_FORMS["mprice"]
    page = pages.get("mprice")
    if page is None:
        return fails + ["Morgenstern-Price produced no section to mutate"]
    number, notation = wanted[0]
    report, _bundle = _calc_report("mprice")
    printed = [b.notation for b in _calc_section(report).blocks
               if b.kind == "math"]
    dropped = [n for n in printed if n != notation]
    if dropped == printed:
        fails.append("the mutation dropped nothing, so it tests nothing")
    elif not _prints_its_page("mprice", path, page, number, notation, dropped):
        fails.append("a section that prints no interslice assumption still "
                     "passed")
    drifted = page.replace(r"\tan \theta_j = \lambda f(x_j)",
                           r"\tan \theta_j = \lambda")
    if drifted == page:
        fails.append("the mutation changed no equation on the page, so it tests "
                     "nothing")
    elif not _prints_its_page("mprice", path, drifted, number, notation,
                              printed):
        fails.append("a printed equation that drifted from its page still "
                     "matched it")
    return fails


def _compiled_scripts(notation):
    """Every subscript and superscript the compiled equation really carries, as
    ``(mark, base, script)`` — read back out of the Word math, not out of the
    notation it was written in."""
    from docx.oxml.ns import qn
    from xslope.report_docx import omath

    def text(node):
        return "".join(t.text or "" for t in node.iter(qn("m:t")))

    out = []
    math = omath(notation)
    for tag, mark in ((qn("m:sSub"), "_"), (qn("m:sSup"), "^")):
        for node in math.iter(tag):
            script = node.find(qn("m:sub") if mark == "_" else qn("m:sup"))
            out.append((mark, text(node.find(qn("m:e"))),
                        text(script) if script is not None else ""))
    return out


#: A symbol written with a subscript or superscript of more than one character,
#: which is what the compiler used to cut short.
_SCRIPTED = re.compile(r"([A-Za-zΑ-Ωα-ω])([_^])([A-Za-z0-9]{1,})")


def _scripts_are_whole(notation, where):
    """Every scripted symbol in one equation, compiled and read back."""
    compiled = _compiled_scripts(notation)
    fails = []
    for base, mark, script in _SCRIPTED.findall(notation):
        if (mark, base, script) not in compiled:
            fails.append(f"{where}: {base}{mark}{script} in {notation!r} "
                         f"compiles to {compiled}, and not to {script!r} under "
                         f"{base!r}")
    return fails


def _math_room(doc=None):
    """``(room in twips, measuring face, size in points)`` — the line an equation
    of this report is set on, read off the shipped template."""
    from docx import Document
    from xslope.report_docx import (DEFAULT_TEMPLATE, MATH_FILL, MATH_FONT,
                                    _measuring_face, _style, _table_font,
                                    _usable_twips)
    doc = doc or Document(DEFAULT_TEMPLATE)
    size = _style(doc, "Normal")
    size_pt = float(size.font.size.pt) if size is not None and size.font.size \
        else 11.0
    face = (MATH_FONT if _measuring_face(MATH_FONT, size_pt)
            else (_table_font(doc) or MATH_FONT))
    return _usable_twips(doc.sections[0]) * MATH_FILL, face, size_pt


def test_prose_symbols_are_set_as_symbols():
    """A symbol a sentence names is SET as that symbol, not spelled with an
    underscore.

    The prose names the letters the equations above it carry — the numerator N_S,
    the base-normal denominator m_α, the interslice resultant Z_{i+1}, the columns
    F_R and F_D — and every one of them reached the page as the raw string a
    builder types: "N_S", which is not the N-sub-S the equation printed, in the
    one place the two have to be the same thing.

    Every prose block of every method's report is compiled here and read back out
    of the paragraph: the underscore is gone from the text runs, and the symbol is
    in the paragraph as Word math.
    """
    fails = []
    from docx import Document
    from docx.oxml.ns import qn
    from xslope.report import Prose
    from xslope.report_docx import DEFAULT_TEMPLATE, INLINE_MATH, _render_prose

    doc = Document(DEFAULT_TEMPLATE)
    seen = set()
    named = 0
    reports = [_build()]
    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        if report is not None:
            reports.append(report)
    for report in reports:
        for block in report.blocks("prose"):
            found = [m.group(0) for m in INLINE_MATH.finditer(block.text)]
            if not found:
                continue
            named += 1
            seen.update(found)
            paragraph = _render_prose(doc, Prose(block.text, links=block.links,
                                                 bold=block.bold))
            text = "".join(t.text or "" for t in paragraph._p.iter(qn("w:t")))
            for symbol in found:
                if symbol in text:
                    fails.append(f"{symbol!r} reaches the page as plain text: "
                                 f"{block.text!r}")
            if "_" in text:
                fails.append(f"an underscore reaches the page: {text!r}")
            # And no symbol is cut short. Read off the paragraph rather than off
            # the pattern that marked it: a pattern that took a_d and left the x
            # satisfies everything above — the underscore is gone and the piece
            # it kept is set as math — and puts half a moment arm on the page.
            # A symbol ends at a word boundary, so the run after one never opens
            # on a letter or a digit.
            after = False
            for child in paragraph._p:
                if child.tag == qn("m:oMath"):
                    after = True
                    continue
                if not after:
                    continue
                after = False
                run = "".join(t.text or "" for t in child.iter(qn("w:t")))
                if run[:1].isalnum():
                    fails.append(f"a symbol of {block.text!r} is cut short: the "
                                 f"text after it opens on {run[:12]!r}")
            if not list(paragraph._p.iter(qn("m:oMath"))):
                fails.append(f"a sentence naming {found} carries no math: "
                             f"{block.text!r}")
    if named < 10:
        fails.append(f"only {named} sentences name a symbol; the sweep is not "
                     f"reaching the sections")
    # The ones the owner named, and the shapes that are easy to get wrong: a
    # Greek subscript, a braced one, and a subscript of more than one letter.
    for symbol in ("N_S", "m_α", "Z_{i+1}", "F_h", "f_o"):
        if symbol not in seen:
            fails.append(f"no sentence in any report names {symbol}, so its "
                         f"shape is untested")
    return fails


def test_table_symbols_are_set_as_symbols():
    """A symbol in a table is set as that symbol too — in its header, in the
    Symbol column that defines it, and in the legend under it.

    The prose was put in the report's own notation and the tables were left in
    the notation a builder types: the nomenclature's own Symbol column printed
    "N_S", "D_W" and "a_dy" as three characters each, the slice table's headers
    printed "x_c" and "M_R", and the legend under it defined "M_R" as
    "(c·Δl + N'·tan φ)·R" with the R set and the M_R not. On one page the
    equations set D_D as a D with a subscript D while the line defining it,
    inches below, spelled it with an underscore.

    Every cell and every legend line of every table of a full report is written
    here and read back: no underscore reaches the page, and what stood where one
    was is a real subscript, at the size of the text around it. The columns are
    measured on the printed form (:func:`xslope.report_docx._plain`), so a table
    whose symbols set shorter than their notation is not left with a column of
    white space — which is checked where every other width is, against the widths
    the document declares.
    """
    fails = []
    from docx import Document
    from docx.oxml.ns import qn
    from xslope.report_docx import (DEFAULT_TEMPLATE, INLINE_MATH, LEGEND_PT,
                                    TABLE_PT, _cell_text, _math_runs, _plain)

    doc = Document(DEFAULT_TEMPLATE)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    def written(text, legend=False):
        """One string as it reaches the page: ``(all its text, the scripts)``."""
        if legend:
            p = doc.add_paragraph()
            _math_runs(p, text, LEGEND_PT)
            size = LEGEND_PT
        else:
            _cell_text(cell, text, TABLE_PT)
            p = cell.paragraphs[0]
            size = TABLE_PT
        scripts = []
        for run in p.runs:
            if run.font.subscript:
                scripts.append(run.text)
                if run.font.size is None or run.font.size.pt != size:
                    fails.append(f"a subscript of {text!r} is set at "
                                 f"{run.font.size} and its cell at {size} pt")
        return "".join(run.text for run in p.runs), scripts

    seen, swept = set(), 0
    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        if report is None:
            continue
        for block in report.blocks("table"):
            lines = [(h, False) for h in block.headers]
            lines += [(c, False) for row in block.rows for c in row]
            lines += [(f"{term}: {definition}", True)
                      for term, definition in (block.legend or [])]
            for text, legend in lines:
                found = [m.group(0) for m in INLINE_MATH.finditer(str(text))]
                if not found:
                    continue
                swept += 1
                seen.update(found)
                plain, scripts = written(text, legend)
                for symbol in found:
                    if symbol in plain:
                        fails.append(f"{method}: {symbol!r} reaches the page as "
                                     f"plain text in {text!r}")
                    script = symbol.partition("_")[2].strip("{}")
                    if script not in scripts:
                        fails.append(f"{method}: the subscript of {symbol!r} in "
                                     f"{text!r} is not set as one: {scripts}")
                if "_" in plain:
                    fails.append(f"{method}: an underscore reaches the page in "
                                 f"{text!r}: {plain!r}")
    if swept < 10:
        fails.append(f"only {swept} table lines name a symbol; the sweep is not "
                     f"reaching the tables")
    # The three the owner named, and one of each kind of line: a header, a
    # nomenclature Symbol, and a legend definition.
    for symbol in ("N_S", "D_W", "a_dy", "M_R"):
        if symbol not in seen:
            fails.append(f"no table of any report names {symbol}, so its shape "
                         f"is untested")

    # The width a symbol is measured at is the width it prints at, not the width
    # of the notation it is typed in.
    from xslope.report_docx import _table_font, _text_width
    family = _table_font(doc)
    for notation in ("N_S", "Z_{i+1}", "M_R (lb·ft per ft)"):
        if _text_width(notation, family, TABLE_PT) != _text_width(
                _plain(notation), family, TABLE_PT):
            fails.append(f"{notation!r} is measured as its notation and set as "
                         f"its symbol; its column carries the difference")
    if _plain("input_template.xlsx") != "input_template.xlsx":
        fails.append("a file name is being read as a symbol")

    # The mutation: a cell written as one plain run, which is what every one of
    # them was. The sweep above has to be able to see the underscore in it.
    _cell_text(cell, "", TABLE_PT)
    p = cell.paragraphs[0]
    p.add_run("N_S (lb/ft)")
    plain = "".join(t.text or "" for t in p._p.iter(qn("w:t")))
    if "N_S" not in plain or "_" not in plain:
        fails.append("the mutation left no underscore on the page, so the "
                     "sweep above tests nothing")
    if any(run.font.subscript for run in p.runs):
        fails.append("the mutation set a subscript of its own")
    return fails


def test_the_base_shear_arm_is_named_for_the_surface():
    """The slice table's resisting-moment column is defined at the arm this
    surface has.

    That column is ``(c·Δl + N'·tan φ)·a_S``, and ``a_S`` is the moment arm of
    the base shear. On a circular surface it is the radius, and the section
    around the table prints and defines R: the footnote's ``a_S`` was a symbol
    the report defined nowhere. On a composite surface the arm is not the radius,
    it stays ``a_S``, and the section prints the general moment arms — so the
    letter is defined where it is used.
    """
    fails = []
    from xslope.report import _legend_arm, equation_symbols

    for method in ("oms", "bishop"):
        report, _bundle = _calc_report(method)
        if report is None:
            fails.append(f"{method}: no report to read the slice table of")
            continue
        legends = [line for block in report.blocks("table")
                   for line in (block.legend or [])
                   if "Resisting moment" in line[1]]
        if not legends:
            fails.append(f"{method}: no slice table defines a resisting moment")
        for label, definition in legends:
            if "a_S" in definition:
                fails.append(f"{method}: {label} is defined at an arm the "
                             f"section never names: {definition!r}")
            if "·R" not in definition:
                fails.append(f"{method}: {label} is defined at no arm at all on "
                             f"a circular surface: {definition!r}")
        # And R is a symbol that report defines.
        section = _calc_section(report)
        rows = [row[0] for block in section.blocks if block.kind == "table"
                and "Nomenclature" in block.caption for row in block.rows]
        if "R" not in rows:
            fails.append(f"{method}: the footnote names R and the nomenclature "
                         f"defines {rows}")

    # The composite surface, which no model of the corpus cuts: the arm stays as
    # the column registry writes it, and a section printing it defines it.
    legend = [("M_R", "Resisting moment this slice contributes about the "
                      "center of rotation, (c·Δl + N'·tan φ)·a_S.")]
    for where, calc in (("a composite surface", {"radius": 0.0}),
                        ("no calculation at all", None)):
        if _legend_arm(legend, calc) != legend:
            fails.append(f"{where}: the arm is renamed to the radius the "
                         f"surface does not have")
    if _legend_arm(legend, {"radius": 42.0}) == legend:
        fails.append("a circular surface left the general arm in the footnote")
    defined = dict(equation_symbols("(c·Δl + N'·tan φ)·a_S"))
    if "a_S" not in defined:
        fails.append(f"a section printing the general arms defines {defined}")
    return fails


def test_calculation_leads_read_once():
    """The sentence that opens a calculation says each of its facts once, and
    attributes each equation to the page that publishes it.

    Four readings this pins. The two moment methods divide two sums that are
    BOTH moments over the radius, and a lead that ended "over the driving moment
    about the center of rotation, divided by the radius" left the division
    hanging on the driving side alone; "each written as its moment about the
    center of rotation divided by the radius" divided both and still let "each"
    distribute over the driving forces it stood next to, so the lead names the
    two sides instead of quantifying over them. Bishop's base normal is formed by
    its page's equation (8), which equation (10) then consumes — the lead named
    (10) for both. And a report that runs one analysis says so: "every analysis
    in this report" counts something the reader can see there is one of.
    """
    fails = []
    from xslope.report import TRANSCRIPTIONS

    for method in ("oms", "bishop"):
        lead = TRANSCRIPTIONS[method].lead
        if "driving moment about the center of rotation, divided by the " \
                "radius" in lead:
            fails.append(f"{method}: only the driving side is divided by the "
                         f"radius: {lead!r}")
        if "both sides written as moments about the center of rotation " \
                "divided by the radius" not in lead:
            fails.append(f"{method}: the lead does not say what the two sums "
                         f"are: {lead!r}")
        if "each written as its moment" in lead:
            fails.append(f"{method}: 'each' distributes over the driving forces "
                         f"it follows, and the numerator is a moment over the "
                         f"radius too: {lead!r}")
    bishop = TRANSCRIPTIONS["bishop"].lead
    if "vertical forces that equation (8) forms the base normal from" not in \
            bishop:
        fails.append(f"Bishop's lead does not give the base normal to the "
                     f"equation that forms it: {bishop!r}")

    # And the Project Definition is written in the register the owner set: the
    # subject is the problem, the figure reference is inside the sentence, and
    # the zones are named. A report of several analyses says once that they all
    # run on this one section; a report of one counts nothing. Built off one
    # model with one engine's solutions and then with two, so the sentences are
    # read as the report writes them. The figure is geometry: naming a member
    # here would name a line the figure no longer draws.
    from xslope.fileio import load_slope_data
    from xslope.report import (DEFAULT_OPTIONS, _Counter,
                               _project_definition_section)

    slope_data = load_slope_data(REINF_XLSX)
    _report, bundle = _calc_report("bishop")
    if bundle is None:
        return fails + ["no solution to build a project definition on"]
    said = {}
    with tempfile.TemporaryDirectory() as tmp:
        for how, solutions in (("one", {"lem": [bundle]}),
                               ("several", {"lem": [bundle],
                                            "seep": [{"solution": None}]})):
            opts = dict(DEFAULT_OPTIONS, method="bishop", pd_figure=True)
            section = _project_definition_section(
                slope_data, solutions, opts, _Counter(), tmp)
            said[how] = " ".join(b.text for b in section.blocks
                                 if b.kind == "prose")
    for how, text in said.items():
        if "The problem cross section is defined by 2 material zones described " \
                "with profile lines." not in text:
            fails.append(f"{how}: the section is not defined noun-first: {text!r}")
        if "The zones are named" not in text:
            fails.append(f"{how}: the material zones are not named: {text!r}")
        if "The problem definition is displayed in Figure 1, including the " \
                "geometry, the material zones and the maximum depth line" not in text:
            fails.append(f"{how}: the figure is not referred to from inside the "
                         f"sentence that says what it shows: {text!r}")
        for retired in ("is run on the model of", "The section is defined by"):
            if retired in text:
                fails.append(f"{how}: the retired wording {retired!r} is back: "
                             f"{text!r}")
        # The reinforced fixture's members are drawn by the engine, not here —
        # read off the figure sentence itself, which is the only one making a
        # claim about what Figure 1 carries.
        shown = text.partition("The problem definition is displayed in")[2]
        shown = shown.partition(".")[0]
        for member in ("reinforcement", "pile"):
            if member in shown:
                fails.append(f"{how}: the figure sentence names the {member}, "
                             f"which the Project Definition figure no longer "
                             f"draws: {shown!r}")
    if "Every analysis in this report is run on this cross section." in said["one"]:
        fails.append(f"a report of one analysis counts them: {said['one']!r}")
    if "Every analysis in this report is run on this cross section." not in \
            said["several"]:
        fails.append(f"a report of several analyses does not say they share the "
                     f"section: {said['several']!r}")
    return fails


def test_the_wide_quotient_is_narrowed():
    """No equation is left wider than the line it is set on, and ``F =`` never
    leaves the fraction it belongs to.

    The factor of safety of the four force methods is one quotient carrying every
    force a slice can take, and it is wider than any text column. There is nowhere
    inside a fraction that a line breaks, so the only break either program can
    find is the equation's own equals sign: Word takes it, and prints F alone on
    one line with = opening the next; LibreOffice does not, and runs the fraction
    out past the right margin. Neither is the equation.

    So a quotient whose DENOMINATOR will not fit the line is narrowed instead —
    that denominator stacked inside the bar — and what is checked here is the
    result: every such equation the report prints is narrowed, the narrowed form
    stands inside the line, it is still a quotient with its head on it, and the
    stack really reaches the document as the construction Word and LibreOffice
    both set (``m:eqArr``).

    The condition is the denominator's own width, because that is what has no
    break in it. An equation merely over the estimate — which is deliberately
    conservative — has its own operators to break at, or sets inside the column
    anyway, and is left alone.
    """
    fails = []
    from xslope.report_docx import (_math_lines, _math_segments, _math_width,
                                    _parse_math, _sole_fraction,
                                    _stacked_quotient, omath)

    room, face, size_pt = _math_room()
    stacked = 0
    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            continue
        for block in section.blocks:
            if block.kind != "math":
                continue
            for line in _math_lines(block.notation, room / 0.92, face, size_pt):
                split = _sole_fraction(line)
                if split is None:
                    continue
                head, _num, den = split
                space = room - _math_width(_parse_math(head)[0], face, size_pt)
                if _math_width(_parse_math(den)[0], face, size_pt) <= space \
                        or len(_math_segments(den)) < 2:
                    continue
                narrowed = _stacked_quotient(line, room, face, size_pt)
                if narrowed is None:
                    fails.append(f"{method}: the denominator of {line!r} is "
                                 f"wider than the line it is set on and nothing "
                                 f"narrows it")
                    continue
                stacked += 1
                if _math_width(narrowed, face, size_pt) > room:
                    fails.append(f"{method}: {line!r} is still wider than the "
                                 f"line after stacking")
                # The head — "F =" — is still attached to its fraction: the
                # stacking returns one sequence, and the fraction is in it.
                if not any(kind == "frac" for kind, _payload in narrowed):
                    fails.append(f"{method}: the narrowed {line!r} is no longer "
                                 f"a quotient: {narrowed}")
                if line.split("frac", 1)[0] != "".join(
                        p for k, p in narrowed if k == "text"):
                    fails.append(f"{method}: the head of {line!r} was dropped "
                                 f"or rewritten by the stacking")
    if not stacked:
        fails.append("no equation of any method needed narrowing, so nothing "
                     "here was exercised")

    # And the construction reaches the document. Read off the compiled OMML, so
    # a stack that is built and then not emitted is caught.
    wide = ("F = frac{sum{(c·Δl + N'·tan φ)·cos α}}"
            "{sum{N'·sin α} + sum{u·Δl·sin α} + sum{kW} + sum{T} "
            "− sum{D·sin β} − sum{P cos ψ} − sum{H cos θ_p} − sum{L cos δ}}")
    if _sole_fraction(wide) is None:
        fails.append("the wide quotient is not read as a quotient at all")
    narrowed = _stacked_quotient(wide, room, face, size_pt)
    if narrowed is None:
        fails.append(f"the published quotient is not narrowed: it measures "
                     f"{_math_width(_parse_math(wide)[0], face, size_pt):.0f} "
                     f"twips against a line of {room:.0f}")
    else:
        from lxml import etree
        from xslope.report_docx import omath_nodes
        out = etree.tostring(omath_nodes(narrowed)).decode()
        if "eqArr" not in out:
            fails.append(f"the narrowed quotient carries no equation array: "
                         f"{out[:400]}")
        if out.count("</m:e>") < 2:
            fails.append("the equation array holds fewer than two lines")

    # And the RENDERER reaches for it. Building the array and then writing the
    # unnarrowed equation anyway is exactly the state this is here to prevent, so
    # the document itself is read back.
    from xslope.report import Math, Report, Section
    from xslope.report_docx import render_docx
    tmp = tempfile.mkdtemp(prefix="xslope_wide_")
    path = render_docx(
        Report(sections=[Section("Equation", blocks=[Math(wide)])],
               meta={"title": "wide", "date": ""}),
        os.path.join(tmp, "wide.docx"))
    with zipfile.ZipFile(path) as z:
        written = z.read("word/document.xml").decode("utf-8")
    if "eqArr" not in written:
        fails.append("a document written with the wide quotient carries no "
                     "equation array; the renderer did not narrow it")
    return fails


def test_scripts_are_not_cut_short():
    """A subscript runs to the end of the word it opens.

    ``F_corr`` is the corrected factor of safety, one symbol. It reached the page
    as F subscript c followed by the roman letters orr — "F c orr" — and so did
    every moment arm of the general equations: a_dx, a_ry, a_ey. The compiler
    took one character after the mark and set the rest as text beside it.

    Every equation every method prints is compiled here and read back out of the
    Word math, so the rule is checked where it is used and not only on the case
    that was reported.
    """
    fails = []
    import xslope.report_docx as report_docx

    # The pins: the symbol that was reported, and a moment arm, whose subscripts
    # have to arrive whole and under the right letter.
    for notation, want in (
            ("F_corr = f_o·F = 1.07702·1.634 = 1.760",
             [("_", "F", "corr"), ("_", "f", "o")]),
            ("sum{D cos β·a_dx} − sum{(H cos θ_p·a_ey + H sin θ_p·a_ex)}",
             [("_", "a", "dx"), ("_", "θ", "p"), ("_", "a", "ey"),
              ("_", "a", "ex")])):
        compiled = _compiled_scripts(notation)
        for one in want:
            if one not in compiled:
                fails.append(f"{notation!r} compiles to {compiled}, without "
                             f"{one}")

    # And the sweep: every equation of every method's calculation.
    swept = 0
    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            continue
        for block in section.blocks:
            if block.kind == "math":
                swept += 1
                fails += _scripts_are_whole(block.notation, method)
    if swept < len(CALC_METHODS):
        fails.append(f"only {swept} equations were compiled; the sweep is not "
                     f"reaching the sections")

    # Mutation: put the one-character rule back and the sweep has to go red.
    saved = report_docx._script_span
    report_docx._script_span = lambda src, at: at + 1
    try:
        caught = _scripts_are_whole("F_corr = f_o·F", "the mutation")
    finally:
        report_docx._script_span = saved
    if not caught:
        fails.append("a subscript cut to one character was not caught, so "
                     "nothing here tests the rule")
    return fails


#: What a printed equation carries besides symbols: the operators, the fence
#: characters, and the three functions and two macros the notation is written
#: with. Everything else that is a letter has to be a symbol the section defines.
_MATH_WORDS = ("frac", "sum", "sin", "cos", "tan", "π")
_MATH_PUNCTUATION = "·−-+=/{}[]()., '"


def test_printed_symbols_resolve():
    """Every symbol a printed equation carries is defined where it is printed.

    An equation prints a letter and the nomenclature under it says what the
    letter means; a letter that reaches the page with no row is a term the reader
    cannot look up. Both models here print one that used to: T, the tension-crack
    water force, which only Spencer's letter V was defined for, and P on a model
    whose reinforcement is axial, where the tangent column is zero on every slice
    and is not printed.
    """
    import re
    fails = []

    for where, xlsx in (("the tension-crack model", TENSION_XLSX),
                        ("the axially reinforced model", AXIAL_XLSX)):
        for method in CALC_METHODS:
            report, _bundle = _calc_report(method, xlsx=xlsx)
            if report is None:
                continue
            section = _calc_section(report)
            if section is None:
                continue
            defined = []
            for block in section.blocks:
                if block.kind == "table" and "Nomenclature" in block.caption:
                    defined += [row[0].split(" (")[0] for row in block.rows]
            for block in section.blocks:
                if block.kind != "math":
                    continue
                # The symbols first, longest first, and the numbers after them:
                # a subscript is a digit, and stripping the numbers ahead of the
                # symbols leaves R_ where R_1 was printed.
                text = block.notation
                for symbol in sorted(set(defined) | set(_MATH_WORDS),
                                     key=len, reverse=True):
                    text = text.replace(symbol, " ")
                text = re.sub(r"\d[\d,]*\.?\d*(?:e[+-]?\d+)?", " ", text)
                left = "".join(ch for ch in text
                               if ch not in _MATH_PUNCTUATION
                               and not ch.isspace())
                if left:
                    fails.append(f"{where} under {method}: {block.notation!r} "
                                 f"prints {left!r}, which the nomenclature "
                                 f"under it does not define")
    return fails


#: The symbols the report's equations use, and the LaTeX each is written as on
#: the documentation pages. The report prints Unicode and the pages print LaTeX,
#: so a notation check has to translate before it compares; where a page spells
#: one of them more than one way, every spelling it uses is listed.
DOC_SYMBOLS = {
    "α": (r"\alpha",),
    "β": (r"\beta",),
    "φ": (r"\phi",),
    "ψ": (r"\psi",),
    "θ": (r"\theta",),
    "Δl": (r"\Delta \ell", r"\Delta\ell"),
    "N'": ("N'",),
    "a_S": ("a_S",),
    "a_N": ("a_N",),
    "x_r": ("x_r",),
    "a_dx": ("a_{dx}",),
    "a_dy": ("a_{dy}",),
    "a_k": ("a_k",),
    "a_t": ("a_t",),
    "a_ry": ("a_{ry}",),
    "a_rx": ("a_{rx}",),
    "a_ey": ("a_{ey}",),
    "a_ex": ("a_{ex}",),
    "a_fx": ("a_{fx}",),
    "a_fy": ("a_{fy}",),
    "m_α": (r"m_\alpha", r"m_{\alpha}"),
    "θ_p": (r"\theta_p",),
    "y_Q": ("y_Q",),
    "x_b": ("x_b",),
    "R_1": ("R_1",),
    "R_2": ("R_2",),
    "Z_n": ("Z_n", "Z_{i+1}"),
    "M_O": ("M_O",),
    "F_x": ("F_x",),
    "F_y": ("F_y",),
    "x_F": ("x_F",),
    "y_F": ("y_F",),
    "θ_j": (r"\theta_j",),
    "λ": (r"\lambda",),
    "f(x_j)": ("f(x_j)",),
    "x_j": ("x_j",),
    "x_L": ("x_L",),
    "x_R": ("x_R",),
    "f_o": ("f_o",),
    "F_corr": ("F_{corr}",),
    "c_m": ("c_m",),
    "φ_m": (r"\phi_m",),
    "Z_i": ("Z_{i}", "Z_i"),
    "Z_{i+1}": ("Z_{i+1}",),
    "θ_i": (r"\theta_i",),
    "θ_{i+1}": (r"\theta_{i+1}",),
    "F_v": ("F_v",),
    "F_h": ("F_h",),
    "P_p": (r"P\cos \psi", r"P \cos \psi"),
    "H_p": (r"H \cos \theta_p",),
}


#: Constructions that describe the DOCUMENT instead of the analysis. The report
#: is engineering documentation: it states facts about the slope, the model and
#: the solution, and a sentence about the section, the table, the process of
#: presenting or what the reader should do with it is a defect (Norm). Each
#: phrase here was really in the prose and was cut; the list is what keeps the
#: voice from drifting back.
#:
#: "section" alone is not bannable — it is also the cross-section of the slope —
#: so only the unambiguous forms are listed.
#: The second group is about the PAGE — how much of it an equation takes, and
#: what was therefore done about it. "It runs several times the width of the
#: page, so it is written here as a quotient of its named sums" states a fact
#: about the typesetting and reaches the reader as an apology for the equation;
#: the equation's own decomposition is a fact about the equation and needs no
#: such excuse. Every phrase here was really in the prose and was cut.
_DOCUMENT_VOICE = (
    "a reader", "the reader", "a reviewer", "the engineer who",
    "this section", "the section prints", "the section says",
    "the section arrives", "the table lists", "the table below",
    "the figure below shows", "is worked through below",
    "the evaluations are not numbered",
    "width of the page", "on the page", "fits on", "too wide", "one line",
    "written here", "printed here", "shown here", "set out here",
    "for brevity", "abbreviated here",
)


def test_prose_is_about_the_analysis():
    """No paragraph of the report describes the report.

    The prose states facts about the section, the model and the solution. Self
    narration — what the section is doing, what a reader should re-form, why the
    document is arranged as it is — reads as an explanation of the work rather
    than the work, and every instance of it was cut. This is the guard.
    """
    fails = []

    reports = [("the default report", _build()),
               ("the seepage report", _engine_report("seep")),
               # A march writes prose no steady report has — the basis it is, the
               # boundary that resolves per step, a sentence per state documented —
               # and none of it was under this sweep.
               ("the transient report", _tseep_report()),
               ("the strength reduction report", _engine_report("fem")),
               ("the reinforcement report",
                _engine_report("fem", xlsx=FEM_REINF_XLSX)),
               ("the piles report", _engine_report("fem", xlsx=FEM_PILES_XLSX))]
    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        if report is not None:
            reports.append((f"the {method} report", report))

    def narration(report):
        return [(b.text, phrase) for b in report.blocks("prose")
                for phrase in _DOCUMENT_VOICE if phrase in b.text.lower()]

    for where, report in reports:
        for text, phrase in narration(report):
            fails.append(f"{where} says {phrase!r}, which describes the "
                         f"document and not the analysis: {text!r}")

    # The mutation: the lead that explained how much of the page an equation
    # takes, and what was done about it, put back. It stood in Bishop's and the
    # Ordinary Method's sections and this check did not see it.
    import dataclasses
    from xslope.report import TRANSCRIPTIONS
    spec = TRANSCRIPTIONS["bishop"]
    key = ("bishop", (), REINF_XLSX)
    saved = _CALC.pop(key, None)
    TRANSCRIPTIONS["bishop"] = dataclasses.replace(spec, lead=(
        "Equation (10) of the derivation is the factor of safety this method "
        "solves. It runs several times the width of the page, so it is written "
        "here as a quotient of its named sums, one per force:"))
    try:
        narrated, _bundle = _calc_report("bishop")
        caught = narration(narrated) if narrated is not None else []
    finally:
        TRANSCRIPTIONS["bishop"] = spec
        _CALC.pop(key, None)
        if saved is not None:
            _CALC[key] = saved
    if not caught:
        fails.append("a lead narrating the width of the page and what was done "
                     "about it went through, so nothing here would stop that "
                     "phrasing coming back")
    return fails


def test_the_equation_is_cited_for_what_it_is():
    """No section cites a derivation for an equation that derivation does not
    publish.

    The force-equilibrium page derives a slice-by-slice march — its equations
    (6) and (10) — and Morgenstern-Price's derives no quotient at all. What
    Corps of Engineers, Lowe & Karafiath and Morgenstern-Price print is the
    horizontal force balance of the whole sliding mass at the converged
    solution, which is true because the march closed and is not what the page
    publishes. Those three say so, and still link the derivation for the march
    that reaches the solution.

    Janbu's page writes that balance directly as its equation (7), and Bishop's,
    the Ordinary Method of Slices' and Spencer's pages publish the equations
    their sections print, so all four cite the derivation for the equation.
    """
    fails = []
    from xslope.report import (WHOLE_MASS_BALANCE_METHODS, method_doc_url,
                               method_label)

    # Written out here rather than imported: a check that reads the same list
    # the prose is chosen from moves with it, and would pass on all seven
    # sections citing the derivation for an equation none of them prints.
    march = ("corps", "lowe", "mprice")
    if tuple(WHOLE_MASS_BALANCE_METHODS) != march:
        fails.append(f"{tuple(WHOLE_MASS_BALANCE_METHODS)} print the balance of "
                     f"the whole mass, and {march} is what publishes a march")

    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        section = _calc_section(report) if report is not None else None
        if section is None:
            fails.append(f"{method}: no calculation to read the citation of")
            continue
        intro = next((b for b in section.blocks if b.kind == "prose"), None)
        if intro is None:
            fails.append(f"{method}: the calculation opens with no statement of "
                         f"what its equation is")
            continue
        label = method_label(method)
        if (label, method_doc_url(method)) not in intro.links:
            fails.append(f"{method}: the opening statement does not link the "
                         f"published derivation")
        published = "the derivation published for" in intro.text
        if method in march:
            if "symbols of the derivation published" in intro.text:
                fails.append(f"{method}: cites the derivation for an equation "
                             f"that page does not publish: {intro.text!r}")
            if "march" not in intro.text:
                fails.append(f"{method}: the derivation is linked without "
                             f"saying what it derives: {intro.text!r}")
            fails += _quotient_is_introduced(section, method)
        elif not published:
            fails.append(f"{method}: prints the equation its page publishes and "
                         f"does not cite it: {intro.text!r}")
    return fails


#: What the sentence above the quotient has to establish, in the words it uses.
#: The reader's question is where the equation came from — the owner's, on the
#: report that printed it bare — and each of these is part of the answer: what
#: summing the march does, what is left, which numbered equation of which
#: derivation that is, and that it is not what the factor of safety was computed
#: from.
#:
#: The last one is pinned with its nouns in it. The sentence once closed "this is
#: the balance that holds at the value it reaches", two pronouns pointing
#: opposite ways in five words: "this" forward to the equation about to be
#: printed, "it" back past the balance to the march.
_QUOTIENT_LEAD = ("Summing the march's equation (6) over the slices",
                  "cancels the interslice forces",
                  "horizontal equilibrium of the whole sliding mass",
                  "equation (12) of the force-equilibrium derivation",
                  "not solved directly for F",
                  "The march is what solves for F",
                  "equation (12) is the balance that holds at the factor of "
                  "safety the march reaches")


def _quotient_is_introduced(section, method):
    """The quotient a marching method prints is introduced where it is printed.

    Not at the head of the section: the reader meets the equation pages after the
    sentence that was supposed to have accounted for it, with the march and its
    reduction in between. The sentence has to be the block directly above the
    equation, it has to say what the equation is and which numbered equation of
    which derivation, and it has to link that derivation — the number is a
    reference, and a reference the reader cannot follow is not one.
    """
    from xslope.report import WHOLE_MASS_BALANCE_PAGE, docs_url

    fails = []
    blocks = list(section.blocks)
    at = next((i for i, b in enumerate(blocks)
               if b.kind == "math" and b.notation.startswith("F = frac")), None)
    if at is None:
        return [f"{method}: prints no quotient for the factor of safety"]
    lead = blocks[at - 1] if at and blocks[at - 1].kind == "prose" else None
    if lead is None:
        return [f"{method}: the quotient is printed with no sentence above it "
                f"saying what it is"]
    for phrase in _QUOTIENT_LEAD:
        if phrase not in lead.text:
            fails.append(f"{method}: the sentence above the quotient does not "
                         f"say {phrase!r}: {lead.text!r}")

    # No sentence here is long enough to lose the reader in, and no word does two
    # jobs in one of them. The cancellation, what survives it and which numbered
    # equation that is were once a single 57-word sentence in which "leaves"
    # meant a force leaving a slice and then a sum leaving a balance behind, and
    # the second reading attached itself to the first.
    for sentence in re.split(r"(?<=[.:])\s+", lead.text):
        words = sentence.split()
        if len(words) > 40:
            fails.append(f"{method}: a {len(words)}-word sentence above the "
                         f"quotient: {sentence!r}")
        for word in ("leave", "leaves", "left", "leaving"):
            if sum(w.strip(",;:.").lower() == word for w in words) > 1:
                fails.append(f"{method}: {word!r} twice in one sentence, and it "
                             f"is a force leaving a slice or a sum leaving a "
                             f"balance: {sentence!r}")

    url = docs_url(WHOLE_MASS_BALANCE_PAGE)
    if not any(str(target).rstrip("/") == url.rstrip("/")
               for _text, target in (lead.links or ())):
        fails.append(f"{method}: the sentence names an equation of "
                     f"{WHOLE_MASS_BALANCE_PAGE} and does not link it: "
                     f"{lead.links}")
    return fails


def test_calculation_notation_matches_the_docs():
    """Every symbol the printed equations use is a symbol the method's own
    documentation page uses.

    The report and the derivation have to be readable side by side. This is a
    light guard — it compares symbols, not equations — but it is what catches a
    silent drift in notation, which is the way the two would come apart.

    A section is read against its own page AND against any page it transcribes an
    equation from: Morgenstern-Price's march is the force-equilibrium
    derivation's per-slice system, and its own page writes none of that system's
    symbols.
    """
    import re
    fails = []
    from xslope.report import METHOD_DOC_PAGES, MOMENT_PART_SYMBOLS

    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        if report is None:
            continue
        section = _calc_section(report)
        if section is None:
            fails.append(f"{method}: no calculation to check the notation of")
            continue
        paths = [METHOD_DOC_PAGES[method]]
        paths += [p for p, _n in _EQUATION_NUMBERS.get(method, ())
                  if p not in paths]
        source = ""
        for path in paths:
            page = os.path.join(_REPO, "docs", path)
            if not os.path.exists(page):
                fails.append(f"{method}: the documentation page {page} is "
                             f"missing")
                continue
            with open(page, encoding="utf-8") as f:
                source += f.read()
        if not source:
            continue
        for block in section.blocks:
            if block.kind != "math":
                continue
            # Longest first: Δl must not be tested as Δ and then l.
            text = block.notation
            # The letters the two moment methods' quotient is named in parts by
            # are the report's own and appear on no page. What pins them is the
            # recomposition — substituted back they have to give the published
            # equation — so they are taken out before the sweep for symbols the
            # documentation does not carry.
            for part in sorted(MOMENT_PART_SYMBOLS, key=len, reverse=True):
                text = text.replace(part, " ")
            for symbol in sorted(DOC_SYMBOLS, key=len, reverse=True):
                if symbol not in text:
                    continue
                text = text.replace(symbol, " ")
                if not any(spelling in source
                           for spelling in DOC_SYMBOLS[symbol]):
                    fails.append(
                        f"{method}: the equations use {symbol!r} but "
                        f"{', '.join(paths)} never writes "
                        f"{DOC_SYMBOLS[symbol][0]!r}")
            # Nothing left over that looks like a symbol we have not declared.
            for stray in re.findall(r"[A-Za-zΑ-Ωα-ω]_\{?[A-Za-zΑ-Ωα-ω0-9]+", text):
                fails.append(f"{method}: {stray!r} in {block.notation!r} is not "
                             f"in the notation map, so nothing checks it "
                             f"against the documentation")
    return fails


#: A phrase from each method's summary that only that method's summary carries —
#: the assumption or the equilibrium condition that distinguishes it. Written out
#: here rather than derived, so that one paragraph pasted over another fails: the
#: check requires each phrase in its own method's block AND out of every other's.
_SUMMARY_MARKS = {
    "oms": ("about the center of the circle, and no other condition",
            "no iteration"),
    "bishop": ("vertical force equilibrium of each slice and moment equilibrium",
               "assumed horizontal"),
    "janbu": ("horizontal force equilibrium of the sliding mass as a whole",
              "correction factor f_o"),
    "corps": ("the ground surface at each slice boundary",),
    "lowe": ("average of the ground-surface slope",),
    "spencer": ("one constant inclination θ",),
    "mprice": ("tan θ = λ·f(x)",),
}


def test_method_summary_opens_each_block():
    """Every method's block opens by saying what the method assumes and which
    equilibrium conditions it satisfies.

    The block has to stand on its own: a reader who arrives at a factor of
    safety or a slice table must be able to tell what produced it without paging
    back to the Calculations section, which may be switched off. So the paragraph
    is required as the FIRST block of the method's own section — above its
    search, its results and its slice table alike — and required again with the
    calculations off.

    It also has to be about the method it opens. Each method carries a phrase
    that appears in its summary and in no other, which is what makes a
    copy-pasted paragraph a failure rather than a pass.
    """
    fails = []
    from xslope.report import METHOD_SUMMARIES, method_label, method_summary

    def method_section(report, method):
        want = method_label(method)
        for section in report.sections:
            for _lvl, node in section.walk():
                if node.title == want:
                    return node
        return None

    for method in CALC_METHODS:
        for label, options in (("calculations on", None),
                               ("calculations off", {"lem_calculations": False})):
            report, _bundle = _calc_report(method, options)
            if report is None:
                continue
            node = method_section(report, method)
            if node is None:
                fails.append(f"{method} ({label}): no {method_label(method)} "
                             f"section to open")
                continue
            if not node.blocks or node.blocks[0].kind != "prose":
                kind = node.blocks[0].kind if node.blocks else "nothing"
                fails.append(f"{method} ({label}): the section opens with "
                             f"{kind}, not the method summary")
                continue
            opening = node.blocks[0].text
            if opening != method_summary(method):
                fails.append(f"{method} ({label}): the opening paragraph is not "
                             f"the method summary: {opening[:80]!r}")
                continue
            if "equilibrium" not in opening:
                fails.append(f"{method} ({label}): the opening paragraph names "
                             f"no equilibrium condition")

    # Every method the report can document has a summary, and each summary is
    # about its own method: its marks are in it and out of all the others.
    for method in CALC_METHODS:
        if not method_summary(method):
            fails.append(f"{method}: no summary paragraph is written for it")
            continue
        for mark in _SUMMARY_MARKS[method]:
            if mark not in METHOD_SUMMARIES[method]:
                fails.append(f"{method}: its summary no longer says {mark!r}, so "
                             f"nothing distinguishes it from the others")
            for other in CALC_METHODS:
                if other != method and mark in METHOD_SUMMARIES.get(other, ""):
                    fails.append(f"{other}: its summary carries {mark!r}, which "
                                 f"belongs to {method} — the two paragraphs say "
                                 f"the same thing")
    return fails


def test_docs_links():
    """The published documentation is linked, at the URL the site really uses."""
    fails = []
    from xslope.report import (DOCS_BASE_URL, METHOD_DOC_PAGES, docs_url,
                               method_doc_url)

    mkdocs = os.path.join(_REPO, "mkdocs.yml")
    with open(mkdocs, encoding="utf-8") as f:
        for line in f:
            if line.startswith("site_url:"):
                declared = line.split(":", 1)[1].strip()
                if declared.rstrip("/") != DOCS_BASE_URL.rstrip("/"):
                    fails.append(f"the report links to {DOCS_BASE_URL}, the site "
                                 f"is published at {declared}")
                break
        else:
            fails.append("mkdocs.yml declares no site_url to check against")

    if docs_url("lem/oms.md") != DOCS_BASE_URL + "lem/oms/":
        fails.append(f"docs_url('lem/oms.md') = {docs_url('lem/oms.md')!r}")
    if docs_url("usage/claude/index.md") != DOCS_BASE_URL + "usage/claude/":
        fails.append(f"an index page maps to {docs_url('usage/claude/index.md')!r}")

    for method, page in METHOD_DOC_PAGES.items():
        if not os.path.exists(os.path.join(_REPO, "docs", page)):
            fails.append(f"{method} cites docs/{page}, which does not exist")
        if not method_doc_url(method).endswith("/"):
            fails.append(f"{method}'s URL is {method_doc_url(method)!r}")

    # The finite element members cite the pages their formulations come from.
    from xslope.report import FEM_DETAIL_DOC_PAGES
    for kind, page in FEM_DETAIL_DOC_PAGES.items():
        if not os.path.exists(os.path.join(_REPO, "docs", page)):
            fails.append(f"the {kind} section cites docs/{page}, which does "
                         f"not exist")
        if not docs_url(page).endswith("/"):
            fails.append(f"the {kind} page's URL is {docs_url(page)!r}")

    from xslope.report import supported_methods
    for method in supported_methods():
        if method not in METHOD_DOC_PAGES:
            fails.append(f"the solver offers {method!r} and no documentation "
                         f"page is mapped for it")
    return fails


def test_calculation_in_the_document():
    """The section reaches the .docx as Word math, with its links live.

    Three things a reader depends on: the equations are text (Word's own math,
    not pictures), the reference to the slice table jumps to the slice table, and
    the citation of the derivation opens the published page.
    """
    import re
    fails = []
    from xslope.report import SLICE_TABLE_BOOKMARK, method_doc_url

    def write(method, options=None):
        import matplotlib
        matplotlib.use("Agg")
        from xslope.fileio import load_slope_data
        from xslope.report import generate_report
        from xslope.slice import generate_slices
        from xslope.solve import solve_selected

        slope_data = load_slope_data(REINF_XLSX)
        ok, out = generate_slices(slope_data, circle=slope_data["circles"][0],
                                  num_slices=15)
        df, surface = out[0].copy(), out[1]
        with contextlib.redirect_stdout(io.StringIO()):
            results = solve_selected(method, df)
        opts = {"method": method, "title": "Calculation", "pd_figure": False,
                "lem_search_figure": False, "lem_solution_figure": False}
        opts.update(options or {})
        tmp = tempfile.mkdtemp(prefix="xslope_calc_")
        path = os.path.join(tmp, f"{method}.docx")
        ok, info = generate_report(
            slope_data,
            {"lem": [{"slice_df": df, "failure_surface": surface,
                      "results": results, "search": None, "method": method}]},
            opts, path)
        if not ok:
            return None, None
        with zipfile.ZipFile(path) as z:
            return (z.read("word/document.xml").decode(),
                    z.read("word/_rels/document.xml.rels").decode())

    doc, rels = write("bishop")
    if doc is None:
        return ["the report could not be written"]

    if "<m:oMath>" not in doc:
        fails.append("the equations are not Word math")
    if doc.count("<m:oMathPara>") < 3:
        fails.append(f"only {doc.count('<m:oMathPara>')} displayed equations "
                     f"reached the document")
    for wanted in ("<m:f>", "<m:nary>", "<m:sSub>"):
        if wanted not in doc:
            fails.append(f"no {wanted} in the document: the notation is not "
                         f"compiling to real math")
    if "m:val=\"∑\"" not in doc:
        fails.append("no summation sign in the math")
    # The numbers are in the math, not in a picture of it.
    if not re.search(r"<m:t[^>]*>[^<]*1\.9", doc):
        fails.append("the factor of safety is not text inside the equation")

    # One bookmark per method's slice table: a report that documents several
    # carries several, and each calculation links to its own.
    mark = f"{SLICE_TABLE_BOOKMARK}_bishop"
    if f'w:name="{mark}"' not in doc:
        fails.append("the slice table carries no bookmark to link to")
    if f'w:anchor="{mark}"' not in doc:
        fails.append("nothing links to the slice table's bookmark")
    else:
        order = (doc.index(f'w:anchor="{mark}"'), doc.index(f'w:name="{mark}"'))
        if order[0] < order[1]:
            fails.append("the cross-reference is written before its bookmark")

    external = re.findall(r'Target="([^"]+)"[^>]*TargetMode="External"', rels)
    if method_doc_url("bishop") not in external:
        fails.append(f"the document links to {external}, not to Bishop's page "
                     f"{method_doc_url('bishop')}")

    # And the citation follows the method the report documents.
    doc2, rels2 = write("spencer")
    external2 = re.findall(r'Target="([^"]+)"[^>]*TargetMode="External"', rels2)
    if method_doc_url("spencer") not in external2:
        fails.append(f"a Spencer report links to {external2}")
    if method_doc_url("bishop") in external2:
        fails.append("a Spencer report still links to Bishop's page")

    # Switched off, the section and its equations are gone.
    doc3, _rels3 = write("bishop", {"lem_calculations": False})
    if "Calculations" in doc3:
        fails.append("lem_calculations=False left the section in the document")
    # The equations go with the section. What math is left is the slice table's
    # own — its headers and the legend under it name the columns' symbols, and a
    # symbol is set as a symbol wherever it is printed — so what has to be gone
    # is every construction only an equation has: a displayed line, a fraction,
    # a summation. A symbol is a base and a subscript and none of those.
    for tag, what in (("<m:oMathPara>", "a displayed equation"),
                      ("<m:f>", "a fraction"), ("<m:nary>", "a summation")):
        if tag in doc3:
            fails.append(f"lem_calculations=False left {what} in the document")
    if f'w:anchor="{mark}"' in doc3:
        fails.append("lem_calculations=False left a link to the slice table")
    return fails


#: The slice force diagram each method's Calculations section opens on, written
#: here from the documentation rather than imported from the report: it is the
#: figure the method's own derivation page displays as the complete set of forces
#: on a slice. Janbu's derivation takes the Ordinary Method's figure, the two
#: force-equilibrium methods share the one on the page they share, and
#: Morgenstern-Price takes Spencer's, which its page displays. A mapping that
#: quietly changed — Bishop's section headed by the force-equilibrium slice —
#: would print an equation over a picture of different forces, and only a
#: second copy of the mapping can catch that.
EXPECTED_DIAGRAMS = {
    "oms": "oms_complete.png",
    "janbu": "oms_complete.png",
    "bishop": "bishop_complete.png",
    "corps": "slice_fe_complete.png",
    "lowe": "slice_fe_complete.png",
    "spencer": "spencer3_forces.png",
    "mprice": "spencer3_forces.png",
}


def _calculations_section(block):
    """The Calculations section of a method's block, or None."""
    if block is None:
        return None
    return next((node for _lvl, node in block.walk()
                 if node.title == "Calculations"), None)


def _glyph_heights(path):
    """``(heights, image width)`` — the height of every dark blob on a drawing
    that could be a glyph, and the drawing it is on.

    The band is the same one the lettering was read in: four to forty pixels
    tall, three to forty wide, which excludes the slice outline and the long
    dimension lines. What it does NOT exclude is the drawings' own small marks —
    the tapered tails of the curved arrows, and slivers of the dashed
    construction lines, which land among the letters. Which of these blobs is
    lettering is not decided here; it is decided by a human and written down in
    :data:`xslope.report.FORCE_DIAGRAM_GLYPH_PX`. This function exists to anchor
    that pin to the artwork.
    """
    import numpy as np
    from PIL import Image
    from scipy import ndimage

    grey = np.array(Image.open(path).convert("L"))
    labels, _n = ndimage.label(grey < 100)
    heights = []
    for rows, cols in ndimage.find_objects(labels):
        h, w = rows.stop - rows.start, cols.stop - cols.start
        if 4 <= h <= 40 and 3 <= w <= 40:
            heights.append(int(h))
    return heights, grey.shape[1]


#: The tan the soil block is filled with on all four drawings, and how far a
#: pixel may sit from it and still be part of the block. The fill is flat, so the
#: tolerance is for the edges the drawing program anti-aliased.
_BLOCK_RGB = (253, 239, 227)
_BLOCK_TOLERANCE = 12


def _block_px(path):
    """``(block width, image width)`` in pixels — the tan soil block on a force
    diagram, and the drawing it is printed on.

    The block is the LARGEST connected run of the fill colour. Taking the colour
    wholesale takes stray anti-aliased pixels with it — the force-equilibrium
    drawing has one at each end, eighty pixels outside the block, which measured
    it seventy per cent too wide.
    """
    import numpy as np
    from PIL import Image
    from scipy import ndimage

    rgb = np.array(Image.open(path).convert("RGB")).astype(int)
    mask = np.abs(rgb - np.array(_BLOCK_RGB)).max(axis=2) <= _BLOCK_TOLERANCE
    labels, count = ndimage.label(mask)
    if not count:
        return None, rgb.shape[1]
    sizes = ndimage.sum(mask, labels, range(1, count + 1))
    biggest = int(np.argmax(sizes)) + 1
    columns = np.where((labels == biggest).any(axis=0))[0]
    return int(columns.max() - columns.min() + 1), rgb.shape[1]


def test_force_diagram_prints_one_slice_at_one_size():
    """Every force diagram prints its soil block at the same size, and the four
    widths that do it are pinned.

    The four drawings frame the same slice differently — the block is half the
    width of the Ordinary Method's image and a quarter of Spencer's — so printing
    them at one width, or at widths derived from their lettering, printed the
    slice itself at four different sizes. What a reader compares from section to
    section is the free body, so it is the block that is held equal and the image
    width that follows from it.

    The widths are stated in the module the report reads them from rather than
    computed there: these are fixed drawings, and their sizes were chosen against
    a rendered page. The block behind those numbers is re-measured here off the
    PNGs, so a redrawn diagram cannot keep a width that no longer prints its
    block at the common size.

    The smallest glyph cannot be measured the same way — the drawings' own arrow
    tails and dash slivers land among the lettering, and no rule separates them —
    so it is pinned by eye per drawing and anchored here: the pinned height must
    still be present on the artwork, within a pixel, or the pin has been left
    behind by a redraw and must be read again.
    """
    fails = []
    from xslope.report import (FORCE_DIAGRAM_BLOCK_IN, FORCE_DIAGRAM_BLOCK_PX,
                               FORCE_DIAGRAM_GLYPH_PX, FORCE_DIAGRAM_LABEL_IN,
                               FORCE_DIAGRAM_WIDTHS, FORCE_DIAGRAMS,
                               force_diagram, force_diagram_width)

    def lettering(name, path, width):
        """What the pinned smallest glyph says about a drawing printed ``width``
        wide: that the pin is still on the artwork, and that it prints legibly.

        The pin is not measured — the drawings' curved arrows taper into
        fragments that no rule separates from lettering, so which blob is a glyph
        is decided by eye and written into FORCE_DIAGRAM_GLYPH_PX. What is
        measured is that the pinned height is still THERE, within a pixel, so a
        redrawn PNG forces a re-pin rather than carrying a number read off a
        drawing that no longer exists.
        """
        out = []
        glyph = FORCE_DIAGRAM_GLYPH_PX.get(name)
        heights, label_pixels = _glyph_heights(path)
        if not heights:
            return [f"{name}: no lettering could be measured on it"]
        if glyph is None:
            return [f"{name}: no smallest glyph is pinned for it"]
        if not any(abs(h - glyph) <= 1 for h in heights):
            out.append(f"{name}: its pinned smallest glyph is {glyph} px, and no "
                       f"blob within a pixel of that is on the drawing (nearest "
                       f"{min(heights, key=lambda h: abs(h - glyph))} px). The "
                       f"drawing has changed; re-pin it by eye.")
        printed = glyph * width / label_pixels
        if printed < FORCE_DIAGRAM_LABEL_IN - 1e-9:
            out.append(f"{name}: at {width} in its smallest glyph prints "
                       f"{printed:.4f} in, under the {FORCE_DIAGRAM_LABEL_IN} in "
                       f"it is legible at")
        return out

    drawn = set(FORCE_DIAGRAMS.values())
    for what, stated in (("widths", FORCE_DIAGRAM_WIDTHS),
                         ("blocks", FORCE_DIAGRAM_BLOCK_PX),
                         ("glyphs", FORCE_DIAGRAM_GLYPH_PX)):
        if set(stated) != drawn:
            fails.append(f"the {what} are stated for {sorted(stated)} and the "
                         f"sections print {sorted(drawn)}")
    blocks = {}
    for method, name in sorted(FORCE_DIAGRAMS.items()):
        path = force_diagram(method)
        if not path:
            fails.append(f"{method}: the diagram {name} is not in the package")
            continue
        got = force_diagram_width(method)

        # --- the block is where the module says it is ---
        block_px, pixels = _block_px(path)
        if block_px is None:
            fails.append(f"{name}: no soil block could be found on it")
            continue
        stated = FORCE_DIAGRAM_BLOCK_PX.get(name)
        if stated != (block_px, pixels):
            fails.append(f"{name}: its block measures {(block_px, pixels)} px; "
                         f"the module states {stated}")
        # --- and it prints at the one size ---
        printed_block = got * block_px / pixels
        blocks[name] = printed_block
        if abs(printed_block - FORCE_DIAGRAM_BLOCK_IN) > 0.01:
            fails.append(f"{name}: at {got} in its block prints "
                         f"{printed_block:.4f} in, not the "
                         f"{FORCE_DIAGRAM_BLOCK_IN} in every section draws it at")

        # --- and the lettering survives the size that gives ---
        fails += lettering(name, path, got)

    # The four blocks agree with each other and not merely with the constant.
    if len(blocks) == len(drawn) and max(blocks.values()) - min(blocks.values()) > 0.01:
        fails.append(f"the blocks print at {blocks}, which is not one size")

    # The widths are not all the same: four drawings that frame the same block in
    # different amounts of white cannot print it equal at one width.
    if len(set(FORCE_DIAGRAM_WIDTHS.values())) < 2:
        fails.append(f"every diagram prints at one width again: "
                     f"{FORCE_DIAGRAM_WIDTHS}")

    # And every one of them fits the text column it is printed in.
    if any(w > 6.5 for w in FORCE_DIAGRAM_WIDTHS.values()):
        fails.append(f"a diagram is wider than the 6.5 in text column: "
                     f"{FORCE_DIAGRAM_WIDTHS}")

    # --- the two ways a pin goes wrong, and the two things that catch them ---
    #
    # They are independent, and each is caught by only one of the two rules. A
    # pin dropped onto one of the drawing's own arrow fragments is still ON the
    # artwork, so the anchor accepts it and only the printed-size floor objects.
    # A pin that has drifted off the artwork can be large enough to clear the
    # floor, so the floor accepts it and only the anchor objects. Neither rule
    # alone is the guard.
    probe = "oms_complete.png"
    probe_path = force_diagram("oms")
    probe_width = FORCE_DIAGRAM_WIDTHS[probe]
    heights, _px = _glyph_heights(probe_path)
    real = FORCE_DIAGRAM_GLYPH_PX[probe]

    # (i) lowered onto a fragment: the smallest blob on the drawing is one of the
    #     arrow tails, well under the pinned glyph.
    fragment = min(heights)
    if fragment >= real:
        fails.append(f"{probe}: its smallest blob is {fragment} px and its "
                     f"pinned glyph {real} px, so there is no fragment under the "
                     f"pin and this mutation proves nothing")
    else:
        FORCE_DIAGRAM_GLYPH_PX[probe] = fragment
        try:
            caught = lettering(probe, probe_path, probe_width)
        finally:
            FORCE_DIAGRAM_GLYPH_PX[probe] = real
        if not any("under the" in c for c in caught):
            fails.append(f"{probe}: pinned at {fragment} px — an arrow fragment, "
                         f"not a glyph — and the printed-size floor let it pass")
        if any("re-pin" in c for c in caught):
            fails.append(f"{probe}: the anchor objected to a fragment pin, which "
                         f"is on the artwork; only the floor should catch it")

    # (ii) drifted off the artwork: a height no blob on the drawing carries, and
    #      big enough that the floor is happy with it.
    absent = next((h for h in range(real + 2, max(heights) + 40)
                   if not any(abs(k - h) <= 1 for k in heights)), None)
    if absent is None:
        fails.append(f"{probe}: every height near the pin is on the drawing, so "
                     f"the anchor mutation proves nothing")
    else:
        FORCE_DIAGRAM_GLYPH_PX[probe] = absent
        try:
            caught = lettering(probe, probe_path, probe_width)
        finally:
            FORCE_DIAGRAM_GLYPH_PX[probe] = real
        if not any("re-pin" in c for c in caught):
            fails.append(f"{probe}: pinned at {absent} px, which no blob on the "
                         f"drawing is within a pixel of, and the anchor let it "
                         f"pass")
        if any("under the" in c for c in caught):
            fails.append(f"{probe}: the floor objected to the {absent} px pin, "
                         f"which clears it; only the anchor should catch it")
    return fails


def _diagram_figures(report):
    """Every force-diagram figure in a report, by caption."""
    return [f for f in report.figures() if f.caption.startswith("Forces on a slice")]


def test_force_diagram_heads_the_calculations():
    """Every Calculations section opens on the free body its equations are
    written about.

    The equations print in the symbols of the derivation, and the derivation
    draws those symbols on a slice. Printed without it, the letters meet the
    reader for the first time in a quotient. The figure is therefore the FIRST
    block of the section — before the sentence that introduces the equation —
    and it is the diagram that method's own documentation page displays.

    It prints exactly where the working prints: with the calculations switched
    off there is no section and no diagram, and a model whose equilibrium cannot
    be worked through gets the sentence that says so and no picture of a
    quotient that was never printed.
    """
    fails = []
    from xslope.report import METHOD_DOC_PAGES, method_label

    for method in CALC_METHODS:
        built, _bundle = _calc_report(method)
        if built is None:
            fails.append(f"the sample model did not solve with {method}")
            continue
        label = method_label(method)
        sec = _calculations_section(_method_block(built, method))
        if sec is None:
            fails.append(f"{method}: the block carries no Calculations section")
            continue
        if not sec.blocks or sec.blocks[0].kind != "figure":
            fails.append(f"{method}: the Calculations section opens on "
                         f"{[b.kind for b in sec.blocks[:3]]}, not its force "
                         f"diagram")
            continue
        fig = sec.blocks[0]
        if fig.caption != f"Forces on a slice — {label}":
            fails.append(f"{method}: the diagram is captioned {fig.caption!r}")
        if fig.number < 1:
            fails.append(f"{method}: the diagram carries figure number "
                         f"{fig.number}")
        if fig.landscape:
            fails.append(f"{method}: the diagram asks for a landscape page")
        # The width the section prints it at is the one pinned for that drawing —
        # the width that puts its soil block at the common size — and it fits the
        # text column.
        from xslope.report import force_diagram_width
        if abs(fig.width_in - force_diagram_width(method)) > 1e-9:
            fails.append(f"{method}: the diagram prints {fig.width_in} in wide, "
                         f"not the {force_diagram_width(method)} in pinned for it")
        if not 1.0 <= fig.width_in <= 6.5:
            fails.append(f"{method}: the diagram prints {fig.width_in} in wide")

        # And the section cites it, by the number the counter gave it — a
        # technical report refers to every figure it prints. The citation belongs
        # in the sentence that first sums the forces the diagram draws, which is
        # where a reader is sent to the picture. Spencer's is the section that
        # carries one today; the other six are their own round.
        if method == "spencer":
            import re
            says = next((b.text for b in sec.blocks
                         if b.kind == "prose" and "F_h and F_v" in b.text), "")
            cited = set(re.findall(r"Figure (\d+)", says))
            if str(fig.number) not in cited:
                fails.append(f"{method}: the diagram is Figure {fig.number} and "
                             f"the sentence that sums its forces cites "
                             f"{sorted(cited) or 'no figure'}: {says!r}")
            if cited - {str(fig.number)}:
                fails.append(f"{method}: the sentence that sums the forces of "
                             f"Figure {fig.number} cites "
                             f"{sorted(cited - {str(fig.number)})} as well")

        want = os.path.join(_REPO, "xslope", "resources",
                            EXPECTED_DIAGRAMS[method])
        if not os.path.exists(fig.path):
            fails.append(f"{method}: the diagram file {fig.path} is not there")
        elif open(fig.path, "rb").read() != open(want, "rb").read():
            fails.append(f"{method}: the section prints "
                         f"{os.path.basename(fig.path)}; the derivation draws "
                         f"{EXPECTED_DIAGRAMS[method]}")

        # The mapping's premise: the diagram is what the method's own published
        # derivation displays. A redraw under a new name would leave the report
        # printing a figure the documentation no longer shows.
        page = os.path.join(_REPO, "docs", METHOD_DOC_PAGES[method])
        if os.path.exists(page):
            with open(page, encoding="utf-8") as fh:
                text = fh.read()
            if EXPECTED_DIAGRAMS[method] not in text:
                fails.append(f"{method}: {METHOD_DOC_PAGES[method]} does not "
                             f"display {EXPECTED_DIAGRAMS[method]}")

        # One diagram per method block, at the head of the calculations and
        # nowhere else.
        found = _diagram_figures(built)
        if len(found) != 1:
            fails.append(f"{method}: the report carries {len(found)} force "
                         f"diagrams for one method")

    # --- switched off: no section, so no diagram ---
    off, _bundle = _calc_report("spencer", options={"lem_calculations": False})
    if off is not None and _diagram_figures(off):
        fails.append("with the calculations switched off the report still "
                     "carries a force diagram")

    # --- a model whose equilibrium cannot be worked through ---
    #
    # Passive support divides by the factor of safety on the driving side of the
    # force methods, so those blocks print the sentence instead of a quotient.
    # The moment methods can show it, and their diagrams stand.
    from xslope.fileio import load_slope_data
    from xslope.report import planned_figures, resolve_options

    passive = load_slope_data(PASSIVE_XLSX)
    for method, wanted in (("janbu", 0), ("corps", 0), ("lowe", 0),
                           ("mprice", 0), ("spencer", 0),
                           ("bishop", 1), ("oms", 1)):
        built, bundle = _calc_report(method, xlsx=PASSIVE_XLSX)
        if built is None:
            fails.append(f"the passive model did not solve with {method}")
            continue
        got = len(_diagram_figures(built))
        if got != wanted:
            fails.append(f"{method}: the passive model prints {got} force "
                         f"diagrams and {wanted} calculations")
        # The count a caller is shown follows the same refusal: a diagram that
        # is not printed is not planned either.
        opts = resolve_options({"input_path": PASSIVE_XLSX, "method": method,
                                "pd_figure": False, "lem_search_figure": False,
                                "lem_solution_figure": False})
        with contextlib.redirect_stdout(io.StringIO()):
            planned = planned_figures(passive, {"lem": [bundle]}, opts)
        drawn = len(built.figures())
        if planned != drawn:
            fails.append(f"{method}: {planned} figures were planned for the "
                         f"passive model and {drawn} were built")
    return fails


# --------------------------------------------------------------------------
# F. the shared plot
# --------------------------------------------------------------------------

def test_profile_lines_name_their_materials():
    """Every profile line's legend entry names the material it bounds.

    The model figure draws numbered profile lines and the materials table names
    materials; with nothing tying the two together a reader has no way to tell
    which line is the core and which the shell. The name goes in the legend
    entry — "Profile 2 (Core)" — and it is the same base geometry every mode of
    the plot draws, so every mode carries it.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    from matplotlib.figure import Figure as MplFigure
    from xslope.plot import plot_inputs

    dam = load_slope_data_cached(SEEP_XLSX)
    lines = dam.get("profile_lines") or []
    mats = dam.get("materials") or []
    if len(lines) < 2 or not mats:
        return ["the dam carries no layered profile lines; the legend is untested"]

    # What each line SHOULD say, read off the model rather than spelled here.
    want = []
    for i, line in enumerate(lines):
        idx = line.get("mat_id")
        idx = idx if isinstance(idx, int) and 0 <= idx < len(mats) else i
        name = str((mats[idx] or {}).get("name") or "").strip()
        want.append(f"Profile {i + 1}" + (f" ({name})" if name else ""))
    if len({w for w in want}) < 2:
        fails.append(f"every profile line would carry the same legend: {want}")

    for mode in ("shared", "lem", "seep", "fem"):
        fig = MplFigure(figsize=(9, 5.5))
        FigureCanvasAgg(fig)
        with contextlib.redirect_stdout(io.StringIO()):
            plot_inputs(dam, fig=fig, mode=mode, show_title=False,
                        frame="content")
        labels = fig.axes[0].get_legend_handles_labels()[1]
        for entry in want:
            if entry not in labels:
                fails.append(f"mode={mode!r}: the legend has no {entry!r} "
                             f"({[l for l in labels if l.startswith('Profile')]})")
    return fails


def test_shared_plot():
    """mode="shared" draws the SECTION and nothing else: its geometry and its
    material zones, with the trial surfaces, the members, the loads and the water
    lines all left to the engine that reads them.

    Every one of those is read by a particular analysis — a piezometric line sets
    pore pressure, a pool becomes a load, a bar carries tension — and each
    stability engine's own model figure draws the ones it reads. Drawn on the
    Project Definition figure too they were the same lines twice, one page apart
    (the owner's ruling, twice: fem_reinforce on the loads, fem_noncircular on
    the piezometric line). What the engine views draw is checked here as well, so
    the suppression above is a MOVE and not a loss.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    from matplotlib.figure import Figure
    from xslope.fileio import load_slope_data
    from xslope.plot import plot_derived_water_lines, plot_inputs
    from xslope.water import water_line_for_stage

    def draw(slope_data, mode):
        fig = Figure(figsize=(9, 5.5))
        FigureCanvasAgg(fig)
        plot_inputs(slope_data, fig=fig, mode=mode, show_title=False,
                    frame="content")
        return fig.axes[0]

    reinf = load_slope_data(REINF_XLSX)
    lem_ax = draw(reinf, "lem")
    shared_ax = draw(reinf, "shared")
    lem_labels = set(lem_ax.get_legend_handles_labels()[1])
    shared_labels = set(shared_ax.get_legend_handles_labels()[1])
    if not any("Circle" in l or "Surface" in l for l in lem_labels):
        fails.append(f"the LEM plot draws no trial circles: {sorted(lem_labels)}")
    for suppressed in ("Starting Circle", "Non-Circular Surface",
                       "Distributed Load", "Tension Crack Depth"):
        if suppressed in shared_labels:
            fails.append(f"the shared plot still draws {suppressed!r}")
    # …and the load it no longer draws is drawn by the engine that applies it.
    if not any("Distributed Load" in l for l in lem_labels):
        fails.append(f"the limit equilibrium view draws no distributed load, so "
                     f"the suppression above is a loss and not a move: "
                     f"{sorted(lem_labels)}")
    # The members leave. Counted on the drawn lines rather than on the legend
    # alone: only the first line of a set carries a label, so a legend test would
    # miss every line after it. Each member kind is counted in its own engine's
    # view first, so a zero in the shared view is the suppression and not an
    # empty fixture.
    def member_lines(ax, colors, label):
        """Every line of one member kind on ``ax``, by its own drawn color."""
        from matplotlib.colors import to_rgba
        want = {to_rgba(c) for c in colors}
        return [ln for ln in ax.lines
                if to_rgba(ln.get_color()) in want
                and (ln.get_label() == label or not ln.get_label().strip()
                     or ln.get_label().startswith("_"))]

    reinf_lem = len(member_lines(lem_ax, ("darkgray",), "Reinforcement Line"))
    reinf_shared = len(member_lines(shared_ax, ("darkgray",), "Reinforcement Line"))
    if not reinf_lem:
        fails.append(f"the reinforced fixture draws no reinforcement in its own "
                     f"engine view, so its absence below proves nothing: "
                     f"{sorted(lem_labels)}")
    if reinf_shared:
        fails.append(f"the shared plot still draws {reinf_shared} reinforcement "
                     f"lines")
    if "Reinforcement Line" in shared_labels:
        fails.append("the shared plot still names the reinforcement in its legend")

    piled = load_slope_data(PILES_XLSX)
    pile_lem = draw(piled, "lem")
    pile_shared = draw(piled, "shared")

    def piles_on(ax):
        lines = len(member_lines(ax, ("green",), "Pile"))
        heads = len([t for t in ax.texts if t.get_text().startswith("H=")])
        return lines, heads

    if piles_on(pile_lem) == (0, 0):
        fails.append("the pile fixture draws no pile in its own engine view, so "
                     "its absence below proves nothing")
    still = piles_on(pile_shared)
    if still != (0, 0):
        fails.append(f"the shared plot still draws the piles: "
                     f"{still[0]} lines, {still[1]} head labels")
    if "Pile" in set(pile_shared.get_legend_handles_labels()[1]):
        fails.append("the shared plot still names the piles in its legend")

    # The water lines leave too, both kinds: a piezometric line is a pore
    # pressure input and a pool is where a load comes from, and each is drawn
    # where the engine that reads it is documented.
    piezo_model = load_slope_data(NONCIRC_FEM_XLSX)
    if len(piezo_model.get("piezo_line") or []) < 2:
        fails.append("the piezometric fixture carries no line, so its absence "
                     "from the shared view proves nothing")
    else:
        for mode, wanted in (("shared", False), ("lem", True), ("fem", True)):
            drawn = {ln.get_gid() for ln in draw(piezo_model, mode).lines}
            has = any("PIEZO" in (g or "") for g in drawn)
            if has is not wanted:
                fails.append(
                    f"mode={mode!r} {'draws' if has else 'draws no'} "
                    f"piezometric line; it should "
                    f"{'draw one' if wanted else 'draw none'}")

    # The tension crack is a limit equilibrium input alone — it caps the driving
    # side of a slice and carries the crack water force — so it is drawn on the
    # limit equilibrium view and nowhere else. On the shared section it was the
    # same red dotted line one page from the figure that means something by it,
    # and it was named in neither sentence (the owner's ruling, tension_crack
    # review). The maximum-depth line is NOT suppressed with it: that one bounds
    # the geometry the section is cut on, so it stays where the section is drawn.
    cracked = load_slope_data(TENSION_XLSX)
    if cracked.get("tcrack_surface") is None:
        fails.append("the tension-crack fixture states no crack, so its absence "
                     "from the shared view proves nothing")
    else:
        for mode, wanted in (("shared", False), ("lem", True), ("fem", False),
                             ("seep", False)):
            ax = draw(cracked, mode)
            drawn = any(ln.get_gid() == "TENSION_CRACK" for ln in ax.lines)
            if drawn is not wanted:
                fails.append(
                    f"mode={mode!r} {'draws' if drawn else 'draws no'} tension "
                    f"crack; it should "
                    f"{'draw one' if wanted else 'draw none'}")
            named = "Tension Crack Depth" in set(
                ax.get_legend_handles_labels()[1])
            if named is not wanted:
                fails.append(f"mode={mode!r} "
                             f"{'names' if named else 'does not name'} the "
                             f"tension crack in its legend")
        # The section's own bound stays on the shared figure: the suppression
        # above is the crack alone and not everything drawn in that corner.
        if not any(ln.get_gid() == "MAX_DEPTH"
                   for ln in draw(cracked, "shared").lines):
            fails.append("the shared plot lost the maximum-depth line, which "
                         "bounds the section rather than any one engine's model")

    # A model with a reservoir/head boundary: the water line is on the ENGINE
    # figures, beside the load derived from it, and not on the shared section.
    dam = load_slope_data(DAM_XLSX)
    derived = water_line_for_stage(dam, stage=1)
    if not derived.get("points"):
        return fails + ["the dam sample no longer states a pool; the water-line "
                        "check has nothing to see"]
    if any(ln.get_gid() == "WATER_LINE" for ln in draw(dam, "shared").lines):
        fails.append("the shared plot draws the pool a head boundary states")
    ax = draw(dam, "lem")
    lines = [ln for ln in ax.lines if ln.get_gid() == "WATER_LINE"]
    if not lines:
        fails.append("the limit equilibrium view of a model with a reservoir "
                     "boundary draws no water line, so the load it does draw "
                     "has no water above it")
    else:
        # The drawn pool stands at the level the boundary states — read from the
        # boundary itself, not from the drawn line, so this is an oracle and not
        # a restatement. (The derived line lies ON the ground where it is dry, so
        # its own maximum is a hilltop, not a water level.)
        from xslope.water import bc_pool_levels
        levels = bc_pool_levels(dam["seepage_bc"], dam["ground_surface"], dam)
        want = max(level for level, _anchors, _label in levels)
        ys = [y for ln in lines for y in ln.get_ydata()]
        if abs(max(ys) - want) > 1e-6:
            fails.append(f"the drawn water line tops out at {max(ys)}, the "
                         f"boundary's pool at {want}")
        # The pool ends at its shoreline: at the drawn line's downstream end the
        # ground stands at the water level, so the load that follows it is not
        # run up a dry face.
        from xslope.water import _y_on
        x_end = max(x for ln in lines for x in ln.get_xdata())
        y_ground = _y_on(dam["ground_surface"], float(x_end))
        if y_ground is None or abs(y_ground - want) > 1e-6:
            fails.append(f"the water line ends at x = {x_end} where the ground is "
                         f"at {y_ground}, not at the {want} shoreline")
    if "Water Surface" not in set(ax.get_legend_handles_labels()[1]):
        fails.append("the water line has no legend key")

    # The layer itself, called directly: a dry model draws nothing.
    fig = Figure(figsize=(6, 4))
    FigureCanvasAgg(fig)
    if plot_derived_water_lines(fig.add_subplot(111), reinf):
        fails.append("a model with no head boundaries drew a water line")

    # The mesh belongs to the mesh figures, not to the shared model — and an
    # engine's model figure can ask for the same, which is what the report's
    # seepage and finite element inputs do: each gives the mesh a figure of its
    # own immediately below, and drawn twice the underlay is a grid over the
    # zones the model figure is there to show.
    if dam.get("mesh") is not None:
        from matplotlib.collections import LineCollection

        def meshed(mode, **kw):
            fig = Figure(figsize=(9, 5.5))
            FigureCanvasAgg(fig)
            plot_inputs(dam, fig=fig, mode=mode, show_title=False,
                        frame="content", **kw)
            return any(isinstance(c, LineCollection)
                       for c in fig.axes[0].collections)

        if meshed("shared"):
            fails.append("the shared plot drew the analysis mesh")
        if not meshed("fem"):
            fails.append("an engine view drew no mesh even when asked for one; "
                         "the suppression below proves nothing")
        for mode in ("fem", "seep"):
            if meshed(mode, show_mesh=False):
                fails.append(f"mode={mode!r} with show_mesh=False still drew the "
                             f"analysis mesh")
    return fails


def test_model_figure_coordinate_labels():
    """The model figure names its points, and the toggle really reaches the plot.

    Geometry is one of the things the model figure is there to show, so the
    coordinate labels are on by default and the dialog's checkbox is for the
    models they crowd. A checkbox wired to nothing looks exactly like one that
    works, so this follows the option from the report's own builder to the
    keyword ``plot_inputs`` is called with, and then to the artists that keyword
    produces: with it on, a vertex the model actually has is annotated with its
    own coordinates; with it off, not one coordinate label is drawn and the
    figure — which stays in the report either way — is a different PNG.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from xslope import plot as plot_mod
    from xslope.report import DEFAULT_OPTIONS, build_report

    slope_data, solutions = _solved()
    if DEFAULT_OPTIONS.get("pd_coords") is not True:
        fails.append("the point coordinates are not on by default")

    # A vertex the model really has, written the way the annotator writes it.
    lines = slope_data.get("profile_lines") or []
    if not lines:
        return fails + ["the sample model has no profile lines; the coordinate "
                        "labels have nothing to name"]
    vx, vy = lines[0]["coords"][0]
    wanted = f"({float(vx):g}, {float(vy):g})"

    passed, drawn, pngs = [], [], []
    real = plot_mod.plot_inputs

    def spy(sd, *args, **kw):
        passed.append(kw.get("label_coordinates", "<not passed>"))
        out = real(sd, *args, **kw)
        fig = kw.get("fig")
        drawn.append([t.get_text() for ax in (fig.axes if fig else [])
                      for t in ax.texts if t.get_gid() == "COORD_LABEL"])
        return out

    base = {"traceability": False, "lem": False, "model_checks": False}
    plot_mod.plot_inputs = spy
    try:
        for state in (True, False):
            with tempfile.TemporaryDirectory() as tmp:
                report = build_report(slope_data, solutions,
                                      dict(base, pd_coords=state), tmp)
                figs = report.figures()
                if len(figs) != 1:
                    fails.append(f"pd_coords={state} left the project definition "
                                 f"with {len(figs)} figures, not one — the labels "
                                 f"option moved the figure itself")
                    pngs.append(None)
                else:
                    with open(figs[0].path, "rb") as fh:
                        pngs.append(fh.read())
                # In the CAPTION: the labels are a property of the figure, and
                # in the sentence they stood in a list of what the model carries
                # — its geometry, its loads, its members — which a labeling is
                # not one of.
                said = " ".join(f.caption for f in report.figures())
                names_them = "labeled with its coordinates" in said
                if names_them is not state:
                    fails.append(f"pd_coords={state} and the caption "
                                 f"{'announces' if names_them else 'does not announce'} "
                                 f"the coordinate labels")
                loose = " ".join(b.text for b in report.blocks("prose"))
                if "labeled with its coordinates" in loose:
                    fails.append(f"the coordinate labels are announced in the "
                                 f"prose as well as the caption: {loose!r}")

        # And with no figure at all the plot is not called, so nothing is added
        # to `passed` and the count below still reads [True, False].
        with tempfile.TemporaryDirectory() as tmp:
            build_report(slope_data, solutions,
                         dict(base, pd_figure=False, pd_coords=True), tmp)
    finally:
        plot_mod.plot_inputs = real

    if passed != [True, False]:
        fails.append(f"the plot was asked for label_coordinates={passed}, not "
                     f"[True, False] — the toggle does not reach the figure")
    if len(drawn) >= 2:
        if wanted not in drawn[0]:
            fails.append(f"the labeled figure does not name the vertex "
                         f"{wanted}; it drew {len(drawn[0])} coordinate labels")
        if drawn[1]:
            fails.append(f"the unlabeled figure still drew {len(drawn[1])} "
                         f"coordinate labels")
    if None not in pngs and pngs[0] == pngs[1]:
        fails.append("the figure is byte-identical with the labels on and off")
    return fails


def _coord_labels(xlsx, figsize=(11.0, 5.0)):
    """``(axes, [(text, box, leader, point), ...])`` for one model's report
    figure.

    The boxes are the TEXT extents, measured off the drawn artists: an
    annotation's own window extent is the union of its text and its leader, and
    a leader is a hairline that is meant to run under things.

    ``leader`` is ``None`` where the label carries none, and otherwise the two
    ends of the line as DRAWN — taken off the arrow patch's own path, so what is
    measured is where the hairline stops rather than where it was aimed, with
    the vertex end second.
    """
    import contextlib as _c
    import io as _io
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    from matplotlib.figure import Figure as MplFigure
    from matplotlib.text import Text
    from xslope.fileio import load_slope_data
    from xslope.plot import plot_inputs

    with _c.redirect_stdout(_io.StringIO()):
        slope_data = load_slope_data(xlsx)
    fig = MplFigure(figsize=figsize)
    FigureCanvasAgg(fig)
    with _c.redirect_stdout(_io.StringIO()):
        plot_inputs(slope_data, fig=fig, mode="shared", show_title=False,
                    frame="content", label_coordinates=True)
    fig.canvas.draw()
    ax = fig.axes[0]
    renderer = fig.canvas.get_renderer()

    def drawn(t, point):
        patch = getattr(t, "arrow_patch", None)
        if patch is None:
            return None
        ends = patch.get_path().transformed(patch.get_transform()).vertices
        a, b = tuple(ends[0]), tuple(ends[-1])
        if math.dist(a, point) < math.dist(b, point):
            a, b = b, a
        return a, b

    out = []
    for t in ax.texts:
        if t.get_gid() != "COORD_LABEL":
            continue
        point = tuple(ax.transData.transform(t.xy))
        out.append((t.get_text(), Text.get_window_extent(t, renderer),
                    drawn(t, point), point))
    return ax.get_window_extent(renderer), out


#: How far a coordinate label may sit from the point it names, in multiples of
#: its own text height, when nothing is near enough to push it out.
#:
#: The placement offers a label a character's width off its vertex and walks
#: outward only while that is objectionable, so an uncrowded vertex is labelled
#: at the tightest offset there is. Over the ninety labels the five sample
#: sections below carry at two figure sizes, that offset measures 0.56 to 0.61
#: text heights — a character is a little over half the height of the line it
#: sits on. The bound is set at 0.85: forty per cent clear of the widest hug
#: measured, so a change of face or of figure size does not trip it, and well
#: under the 1.08 heights an offset ring keyed to the line height rather than to
#: the character put every uncrowded label at, so the bound tells the two apart.
HUG = 0.85

#: How near its vertex a leader may stop, and how far short of it it must, in
#: multiples of the label's own text height.
#:
#: A callout approaches the point it names and stops just short: a hairline run
#: into a corner is ink on the section — at a dam crest the leader arrives along
#: the crest line and the two read as one bent line — while one that halts well
#: out has to be read across the gap to see what it points at. The trim is half
#: a character, which measures 0.31 text heights over the leaders the sample
#: sections below draw at two figure sizes. The bounds are set a third under
#: that and half again over it, so a change of face or of figure size does not
#: trip them while a pull-back of a single point — 0.14 heights, which the width
#: of the line drawn through the corner swallows whole — does not pass. The
#: upper bound is well inside the 0.56 heights the tightest label offset
#: measures, so a leader always stops nearer its corner than any label stands.
LEADER_GAP = (0.35, 0.55)


def test_coordinate_labels_are_placed_clear():
    """Every coordinate label is readable where it landed.

    The labels name the vertices of the section, and on a dam several of those
    vertices are a few feet apart — a crest corner, the core beneath it — so the
    placement is solved rather than offset by a constant. What the solution owes
    a reader is checked on the artists it produced, not on the pixels: no label
    lies over another, none runs off the axes it belongs to, any label the
    solver had to move away from its point carries a leader back to it, since a
    coordinate that could belong to either of two corners belongs to neither,
    and NO TWO LEADERS CROSS — a reader who follows the wrong line out of a
    crossing reads the wrong coordinate, and nothing on the figure says so. A
    leader that is drawn stops short of its vertex rather than running into it,
    by a gap measured off the type and bounded both ways: a line that meets the
    corner is ink on the section, and one that halts well out is a line the
    reader has to carry across the white.

    And a label HUGS the point it names unless something made it travel. A
    coordinate floating in the white is read against the drawing rather than
    against its corner, so the placement offers every label a character's width
    off its vertex and moves it only where that offer is objectionable — and a
    label it did move carries a leader. The two facts are checked from both ends:
    a label with no leader has not travelled, so it stands within :data:`HUG` of
    its own text height of its point; and on a section where every vertex has
    somewhere to put its coordinate, NO label carries a leader at all, so a
    placement that walked past a clear position for a softer preference is caught
    even where the label it moved is honestly labelled as having moved.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from xslope.plot import _segments_cross

    # A zoned dam (the crest and the core top are four points inside one label
    # width), a reinforced slope with two vertices two feet apart, a layered
    # section, and two bare slopes: five placements the same rule has to solve.
    # All but the dam have somewhere to put every coordinate — the reinforced
    # slope's two close pairs are resolved by putting the two labels on opposite
    # sides of their points, not by moving either — so nothing on them should
    # travel. Each is drawn at the report's figure size and again at three
    # quarters of it — the same section with the type forty per cent larger
    # against it, which is what a smaller window gives Studio.
    for (label, xlsx, roomy), size in itertools.product(
            (("the dam", SEEP_XLSX, False),
             ("the reinforced slope", REINF_XLSX, True),
             ("the layered section", NONCIRC_XLSX, True),
             ("the simple slope", SIMPLE1_XLSX, True),
             ("the zoned dam", DAM_XLSX, False)),
            ((11.0, 5.0), (8.0, 3.6))):
        label = f"{label} at {size[0]:g}x{size[1]:g}"
        frame, labels = _coord_labels(xlsx, figsize=size)
        if len(labels) < 4:
            fails.append(f"{label}: {len(labels)} coordinate labels — too few to "
                         f"exercise the placement")
            continue
        for i in range(len(labels)):
            for j in range(i + 1, len(labels)):
                if labels[i][1].overlaps(labels[j][1]):
                    fails.append(f"{label}: {labels[i][0]} and {labels[j][0]} "
                                 f"are printed over each other")
        # The leader a reader follows: the vertex, to the middle of the box the
        # coordinate is printed in. Taken for every label, drawn or not — a
        # label sitting beside its point is read along that line just the same.
        runs = [(text, ((ax_px, ay_px),
                        ((box.x0 + box.x1) / 2, (box.y0 + box.y1) / 2)))
                for text, box, _leader, (ax_px, ay_px) in labels]
        for i, (t1, r1) in enumerate(runs):
            for t2, r2 in runs[i + 1:]:
                if _segments_cross(r1, r2):
                    fails.append(f"{label}: the leaders of {t1} and {t2} cross")
        # And a label nearer some other vertex than the one it names carries a
        # leader: along a stack of vertices one label-height apart every label
        # can land beside its neighbour's point without any two leaders ever
        # meeting, and only the line back says which coordinate is whose.
        for text, box, leader, own in labels:
            def gap(p):
                return math.hypot(max(box.x0 - p[0], 0.0, p[0] - box.x1),
                                  max(box.y0 - p[1], 0.0, p[1] - box.y1))
            nearer = [t for t, _b, _l, p in labels if p != own and gap(p) < gap(own)]
            if nearer and not leader:
                fails.append(f"{label}: {text} sits nearer the vertex "
                             f"{nearer[0]} names than its own, with no leader "
                             f"tying it back")
        for text, box, leader, (ax_px, ay_px) in labels:
            if not frame.contains(box.x0, box.y0) or not frame.contains(box.x1, box.y1):
                fails.append(f"{label}: {text} is drawn off the axes "
                             f"({box.x0:.0f}..{box.x1:.0f}, "
                             f"{box.y0:.0f}..{box.y1:.0f} in "
                             f"{frame.x0:.0f}..{frame.x1:.0f}, "
                             f"{frame.y0:.0f}..{frame.y1:.0f})")
            # A label that had to travel needs a leader: the reach it may sit
            # inside without one is its own height, so the rule scales with the
            # type rather than with the model.
            reach = 1.5 * box.height
            gap = min(abs(ax_px - box.x0), abs(ax_px - box.x1)) \
                if not (box.x0 <= ax_px <= box.x1) else 0.0
            gap = max(gap, 0.0 if box.y0 <= ay_px <= box.y1
                      else min(abs(ay_px - box.y0), abs(ay_px - box.y1)))
            if gap > reach and not leader:
                fails.append(f"{label}: {text} sits {gap:.0f} px from the point "
                             f"it names with no leader tying it back")

        # A label with no leader HUGS the point it names. The two are the same
        # statement read from either end: the figure carries a leader for every
        # coordinate that had to move, so a coordinate standing on its own has
        # not moved, and it is beside its corner rather than adrift in the white.
        # Which of the two a given label gets is the placement's business; that
        # it is one or the other is the reader's.
        for text, box, leader, own in labels:
            if leader:
                continue
            gap = math.hypot(
                0.0 if box.x0 <= own[0] <= box.x1
                else min(abs(own[0] - box.x0), abs(own[0] - box.x1)),
                0.0 if box.y0 <= own[1] <= box.y1
                else min(abs(own[1] - box.y0), abs(own[1] - box.y1)))
            if gap > HUG * box.height:
                fails.append(f"{label}: {text} carries no leader and yet sits "
                             f"{gap / box.height:.2f} text heights off the point "
                             f"it names (the bound is {HUG:g})")

        # And a leader that IS drawn approaches its vertex without meeting it,
        # by :data:`LEADER_GAP` of its own text height — measured on the path
        # the hairline was drawn along, so what is checked is where the line
        # stops rather than where it was aimed. A leader shorter than the gap it
        # stops short by is a smudge and not a line: that label is against its
        # point already and carries no leader at all, which the hug bound above
        # then holds to.
        for text, box, leader, own in labels:
            if not leader:
                continue
            stop = math.dist(leader[1], own) / box.height
            run = math.dist(leader[0], leader[1]) / box.height
            if not LEADER_GAP[0] <= stop <= LEADER_GAP[1]:
                fails.append(f"{label}: the leader of {text} stops {stop:.2f} "
                             f"text heights from the point it names, outside "
                             f"{LEADER_GAP[0]:g} to {LEADER_GAP[1]:g}")
            elif run < stop:
                fails.append(f"{label}: the leader of {text} is {run:.2f} text "
                             f"heights long and stops {stop:.2f} short of its "
                             f"point — the trim has eaten the line")

        # And read the other way: on a section with room for every coordinate,
        # a leader is the mark of a label that travelled, and nothing had cause
        # to. One here means the placement passed over a clear position beside
        # the point for something it liked better further out.
        if roomy:
            travelled = [t for t, _b, leader, _p in labels if leader]
            if travelled:
                fails.append(f"{label}: {travelled} carry leaders on a section "
                             f"where every vertex has room for its own label")

    # The dam's crest cluster is the case the placement exists for: four
    # vertices inside one label width, so something has to move.
    _frame, labels = _coord_labels(SEEP_XLSX)
    from xslope.plot import _label_geometry
    from xslope.fileio import load_slope_data
    with contextlib.redirect_stdout(io.StringIO()):
        slope_data = load_slope_data(SEEP_XLSX)
    points, _segments = _label_geometry(slope_data)
    named = {f"({x:g}, {y:g})" for x, y in points}
    if {t for t, _b, _l, _p in labels} != named:
        fails.append(f"the labels {sorted(t for t, _b, _l, _p in labels)} are not "
                     f"the model's vertices {sorted(named)}")
    if not any(leader for _t, _b, leader, _p in labels):
        fails.append("no label on the dam carries a leader; the crest cluster is "
                     "the case the placement exists for and nothing moved")
    return fails


# --------------------------------------------------------------------------
# K. the seepage and finite element sections
#
# Both engines are documented from SHIPPED solutions: the sample models carry
# the companion files their solvers write ({base}_seep.csv, {base}_fem_*.csv),
# and reading one back is the same bundle a run emits. Nothing here re-solves —
# a check of what the report SAYS about a solution has no business spending a
# seepage iteration or an SSRM bisection to get one.
#
# The reading is xslope.report.solutions_from_sidecars, the same call a caller
# makes to report an already-solved model, so every check below runs on what a
# caller gets rather than on a recipe kept only here.
# --------------------------------------------------------------------------

#: A dam with a solved unconfined seepage analysis beside it. Its four materials
#: are all on the linear-front conductivity model.
SEEP_XLSX = os.path.join(_REPO, "docs", "lem", "files", "xslope_gsat_seep.xlsx")

#: The same dam with its two zones on van Genuchten instead — the corpus's model
#: whose relative conductivity genuinely falls through decades, and the one that
#: settles what the conductivity figures' ordinate scale is chosen by.
VG_SEEP_XLSX = os.path.join(_REPO, "docs", "seep", "files",
                            "xslope_earth_dam1_vg.xlsx")

#: A loaded slope with a solved SSRM run beside it.
FEM_XLSX = os.path.join(_REPO, "docs", "fem", "files",
                        "xslope_griffiths1_load.xlsx")

#: The two finite element models that carry one-dimensional members: six
#: reinforcement lines, and two piles. Both ship a solved strength reduction run
#: with member forces, and both are ALSO solved here for one gravity trial — a
#: second or two each — because that trial is the converged, no-mechanism run
#: the checks need a fixture for, and because it puts a real bar force and a
#: real pile moment in front of the report on one recipe. The shipped runs are
#: read through :func:`_restored` by the checks that are about a mechanism.
#: (The run that records no member forces at all is built in the check that is
#: about it, :func:`test_a_solution_without_member_forces_says_so`.)
FEM_REINF_XLSX = os.path.join(_REPO, "docs", "fem", "files",
                              "xslope_reinforce_fem.xlsx")
FEM_PILES_XLSX = os.path.join(_REPO, "docs", "fem", "files",
                              "xslope_piles_fem.xlsx")
FEM_1D_MODELS = (FEM_REINF_XLSX, FEM_PILES_XLSX)

#: A benchmark whose saved run records a factor of safety (1.606) and NO trial
#: factor of its own — the model the invented "last converged F" was found on.
#: Its at-failure sidecar records the trial that WAS run, 1.847.
RS2_28A_XLSX = os.path.join(_REPO, "docs", "verification", "files",
                            "rocscience", "rs2_28a.xlsx")

#: A reservoir section solved by both engines off one mesh, its materials taking
#: pore pressure from the computed seepage field.
JOHNSON_XLSX = os.path.join(_REPO, "docs", "seep", "files",
                            "xslope_johnson_res.xlsx")

_ENGINE = {}


def _restored(xlsx, notes=None):
    """``(slope_data, solutions)`` for a model, with every engine solution that
    ships beside it read back through the public helper.

    One recipe, used here and by anyone assembling a report of an already-solved
    model: whatever the checks below assert about a restored bundle, they assert
    about the bundle a caller gets.
    """
    from xslope.fileio import load_slope_data
    from xslope.report import solutions_from_sidecars

    with contextlib.redirect_stdout(io.StringIO()):
        slope_data = load_slope_data(xlsx)
        solutions = solutions_from_sidecars(xlsx, slope_data, notes)
    return slope_data, solutions


def _seep_bundle(xlsx=SEEP_XLSX):
    """``(slope_data, bundle)`` for a seepage analysis, read back from the
    solution shipped beside the model. The first boundary condition set: a model
    that carries two is read by the checks that are about two."""
    key = ("seep", xlsx)
    if key not in _ENGINE:
        slope_data, solutions = _restored(xlsx)
        _ENGINE[key] = (slope_data, solutions["seep"][0])
    return _ENGINE[key]


def _fem_bundle(xlsx=FEM_XLSX):
    """``(slope_data, bundle)`` for a strength reduction run, read back from the
    solution shipped beside the model — the shape Studio's FEM runner emits."""
    key = ("fem", xlsx)
    if key not in _ENGINE:
        slope_data, solutions = _restored(xlsx)
        _ENGINE[key] = (slope_data, solutions["fem"])
    return _ENGINE[key]


def _fem_1d_bundle(xlsx):
    """``(slope_data, bundle)`` for a model carrying reinforcement or piles.

    One gravity trial is solved here and cached for every check that reads it,
    so both models arrive by one recipe whatever ships beside them. A trial that
    left every bar and every pile at zero force would exercise the report's
    member sections on a mechanism that never engaged, so the checks that read
    them assert the forces are real.
    """
    key = ("fem1d", xlsx)
    if key in _ENGINE:
        return _ENGINE[key]
    from xslope.fem import build_fem_data, solve_fem
    from xslope.fileio import load_slope_data

    with contextlib.redirect_stdout(io.StringIO()):
        slope_data = load_slope_data(xlsx)
        fem_data = build_fem_data(slope_data, slope_data["mesh"])
        solution = solve_fem(fem_data, F=1.0, debug_level=0, max_iterations=60)
    bundle = {"fem_data": fem_data, "solution": solution, "FS": None,
              "analysis": "single", "failure_solution": None}
    _ENGINE[key] = (slope_data, bundle)
    return _ENGINE[key]


def _fem_any_bundle(xlsx):
    """The finite element bundle for a model: read back where a solution ships
    beside it, solved where none does."""
    return _fem_1d_bundle(xlsx) if xlsx in FEM_1D_MODELS else _fem_bundle(xlsx)


def _built_report(slope_data, solutions, options):
    """A report of a model built here rather than read from a file — how a check
    asks what a model with one property changed would report."""
    from xslope.report import build_report

    opts = dict(FAST_FIGURES)
    opts.update(options or {})
    tmp = tempfile.mkdtemp(prefix="xslope_built_")
    with contextlib.redirect_stdout(io.StringIO()):
        return build_report(slope_data, solutions, opts, tmp)


def _engine_report(engine, options=None, bundle=None, xlsx=None):
    """A report of one engine's model with that engine's section built.

    The limit equilibrium section is switched off: these models are loaded, not
    solved for stability, and a report of a seepage run is a real thing to ask
    for.
    """
    from xslope.report import build_report

    xlsx = xlsx or (SEEP_XLSX if engine == "seep" else FEM_XLSX)
    slope_data, default = (_seep_bundle(xlsx) if engine == "seep"
                           else _fem_any_bundle(xlsx))
    opts = {"input_path": xlsx, "lem": False, "pd_figure": False}
    opts.update(FAST_FIGURES)
    opts.update(options or {})
    tmp = tempfile.mkdtemp(prefix=f"xslope_{engine}_")
    with contextlib.redirect_stdout(io.StringIO()):
        return build_report(slope_data, {engine: bundle or default}, opts, tmp)


def _titles(report):
    return [t for _lvl, t in report.section_titles()]


def _prose(report):
    return [b.text for b in report.blocks("prose")]


def _planned_matches(report, engine, options=None, bundle=None, xlsx=None,
                     slope_data=None):
    """``planned_figures`` against what the build produced, for one engine."""
    from xslope.report import planned_figures, resolve_options

    xlsx = xlsx or (SEEP_XLSX if engine == "seep" else FEM_XLSX)
    read, default = (_seep_bundle(xlsx) if engine == "seep"
                     else _fem_any_bundle(xlsx))
    slope_data = read if slope_data is None else slope_data
    opts = {"input_path": xlsx, "lem": False, "pd_figure": False}
    opts.update(FAST_FIGURES)
    opts.update(options or {})
    planned = planned_figures(slope_data, {engine: bundle or default},
                              resolve_options(opts))
    return planned, len(report.figures())


def _kr_axes(materials, abscissa="suction"):
    """The Axes a set of materials' conductivity curves are drawn on, off the
    real figure rather than off the code that draws it."""
    import matplotlib.figure as mplfig
    from xslope.plot import plot_material_kr_set
    fig = mplfig.Figure()
    plot_material_kr_set(materials, fig=fig, abscissa=abscissa,
                         show_legend=False)
    return fig.axes[0]


def _kr_ordinate_scale_checks():
    """The conductivity figures are read on the scale their models are shaped by.

    The linear front is a straight ramp: kr falls linearly from 1 at the ground
    water surface to kr0 one head interval below it, and the position of that
    front is the whole of what the model says. Drawn against a logarithmic
    ordinate the ramp becomes a decaying curve that falls off a cliff, and the
    front cannot be read off it at all. van Genuchten and Gardner do fall through
    decades, and a linear ordinate lays all of that flat against zero.

    So the scale follows the models assigned — linear where every curve on the
    axes is a linear front, logarithmic as soon as one of them spans decades —
    and the same rule governs the report's paired figures and the material
    editor's single one, so an engineer reads one curve on one axis wherever it
    is met.
    """
    fails = []
    from xslope.plot import (plot_material_kr, _kr_yscale, _kr_model)
    from xslope.report import _kr_materials
    import xslope.plot as plot_mod
    import matplotlib.figure as mplfig
    import numpy as np

    linear_mats = _kr_materials(load_slope_data_cached(SEEP_XLSX))
    vg_mats = _kr_materials(load_slope_data_cached(VG_SEEP_XLSX))
    # A set of both is not in the corpus and does not need to be: the rule is
    # about composition, so the composition is built here.
    mixed_mats = (linear_mats[:1] or []) + (vg_mats[:1] or [])

    if not linear_mats or not all(_kr_model(m) == "lf" for m in linear_mats):
        fails.append(f"the seepage sample is not all linear-front — "
                     f"{[_kr_model(m) for m in linear_mats]} — so the scale "
                     f"the defect was found on is untested")
    if not vg_mats or not any(_kr_model(m) == "vg" for m in vg_mats):
        fails.append(f"{os.path.basename(VG_SEEP_XLSX)} carries no van "
                     f"Genuchten material, so the log ordinate is untested")

    # Each composition, on both abscissae: the pair is one statement of the
    # model, and a pair read on two ordinates does not mirror.
    for what, mats, want in (("all linear front", linear_mats, "linear"),
                             ("all van Genuchten", vg_mats, "log"),
                             ("mixed", mixed_mats, "log")):
        if not mats:
            continue
        for abscissa in ("suction", "head"):
            ax = _kr_axes(mats, abscissa)
            got = ax.get_yscale()
            if got != want:
                fails.append(f"the {what} materials are drawn against "
                             f"{abscissa} on a {got} kr ordinate, not {want}")
            lo, hi = ax.get_ylim()
            if want == "linear" and not (abs(lo) < 1e-12 and 1.0 < hi <= 1.2):
                fails.append(f"the {what} kr ordinate spans {lo:g} to {hi:g}; "
                             f"kr runs 0 to 1 and the axis should show that")
            if want == "log" and not (0 < lo < 1 and hi > 1.0):
                fails.append(f"the {what} kr ordinate spans {lo:g} to {hi:g} on "
                             f"a log axis, which does not frame the curves")

        # The material editor draws the same curve on the same axis. Studio and
        # the report disagreeing about a scale is one model with two shapes.
        one = mats[0]
        fig = mplfig.Figure()
        ax = fig.add_subplot(111)
        plot_material_kr(ax, one)
        alone = _kr_yscale([one])
        if ax.get_yscale() != alone:
            fails.append(f"the material editor draws {_kr_model(one)!r} on a "
                         f"{ax.get_yscale()} ordinate; the report's rule gives "
                         f"{alone}")

    # A linear front on a mixed axes is still its own numbers — a log ordinate
    # is the right axis for that set, and the ramp is correctly plotted on it.
    if mixed_mats and linear_mats:
        alone = {ln.get_label(): ln.get_data()[1]
                 for ln in _kr_axes(linear_mats[:1]).get_lines()}
        together = {ln.get_label(): ln.get_data()[1]
                    for ln in _kr_axes(mixed_mats).get_lines()}
        for label, kr in alone.items():
            other = together.get(label)
            if other is None or not np.allclose(kr, other):
                fails.append(f"the linear-front curve for {label!r} changes when "
                             f"it is drawn beside a model that spans decades")

    # Mutation: the check has to be able to fail. Forcing the log ordinate back
    # onto the all-linear-front set is the figure the defect was reported on,
    # and it has to go red here.
    real = plot_mod._kr_yscale
    plot_mod._kr_yscale = lambda materials: "log"
    try:
        forced = [_kr_axes(linear_mats, a).get_yscale()
                  for a in ("suction", "head")] if linear_mats else []
    finally:
        plot_mod._kr_yscale = real
    if linear_mats and set(forced) != {"log"}:
        fails.append(f"the ordinate scale is not read from the models: forcing "
                     f"log still drew {forced}")
    if linear_mats and _kr_yscale(linear_mats) != "linear":
        fails.append("the rule itself does not call the all-linear-front set "
                     "linear, so the mutation above proves nothing")
    return fails


def test_seep_section():
    """A report of a seepage run says what was solved, on what, and what flowed."""
    fails = []
    _slope_data, bundle = _seep_bundle()
    report = _engine_report("seep")

    expected = [(1, "Traceability"), (1, "Project Definition"),
                (1, "Seepage Analysis"), (2, "Analysis Inputs"),
                (2, "Seepage Mesh"), (2, "Results")]
    got = report.section_titles()
    if got != expected:
        fails.append(f"the seepage report's sections are {got}, expected {expected}")

    # What was solved: the sample is unconfined, and both boundary counts are
    # read off the same array the solver used rather than stated as a constant
    # here.
    bc_type = list(bundle["seep_data"]["bc_type"])
    n_head = sum(1 for t in bc_type if int(t) == 1)
    n_exit = sum(1 for t in bc_type if int(t) == 2)
    if not n_exit:
        fails.append("the sample is not an unconfined problem, so the report's "
                     "unconfined wording is never exercised")
    text = " ".join(_prose(report))
    if "unconfined" not in text:
        fails.append("the seepage section does not say the problem was unconfined")
    for count in (f"{n_head:,}", f"{n_exit:,}"):
        if count not in text:
            fails.append(f"the boundary node count {count} is not stated: {text!r}")

    # One term for the thing the figure draws: the prose, the figure's own legend
    # and the seepage documentation all call them flow lines. That the term is
    # never spelled the other way, anywhere it is written for a reader, is
    # test_one_term_for_flow_lines.
    if "the flow lines" not in text:
        fails.append(f"the section never names the flow lines the figure draws: "
                     f"{text!r}")

    # The flow: the number the solution carries, in the prose, in bold.
    q = bundle["solution"]["flowrate"]
    stated = [b for b in report.blocks("prose") if f"{q:.4g}" in b.text]
    if not stated:
        fails.append(f"the computed flow of {q:.4g} is stated nowhere in the "
                     f"seepage section: {text!r}")
    elif not any(f"{q:.4g}" in " ".join(b.bold) for b in stated):
        fails.append("the flow is stated but not set in bold")

    # The inputs: the mesh, and the conductivities the flow was solved with.
    seep = next((s for s in report.sections if s.title == "Seepage Analysis"),
                None)
    inputs = next((c for c in (seep.children if seep else [])
                   if c.title == "Analysis Inputs"), None)
    if inputs is None:
        fails.append("the seepage section has no Analysis Inputs")
    else:
        kv = [b for b in inputs.blocks if b.kind == "keyvalues"]
        labels = [l for b in kv for l, _v in b.items]
        if "Mesh" not in labels:
            fails.append(f"the seepage inputs name no mesh: {labels}")
        table = next((b for b in inputs.blocks if b.kind == "table"), None)
        if table is None:
            fails.append("the seepage inputs carry no material properties table")
        elif not any("k" in h for h in table.headers):
            fails.append(f"the seepage material table has no conductivity "
                         f"column: {table.headers}")

    # The figures the build produced are the figures it planned, and each is a
    # different reading of the run: the model the flow was solved on, the
    # conductivity curves it reduces each material by, the mesh with the boundary
    # conditions on it, and the field that came out.
    from xslope.report import SEEP_PANELS
    planned, drawn = _planned_matches(report, "seep")
    if planned != drawn:
        fails.append(f"the seepage report planned {planned} figures and built {drawn}")
    sources = [f.source for f in report.figures()]
    for wanted in ["seep model", "seep kr", "seep kr_head", "seepage bc1 mesh"] \
            + [f"seepage bc1 {p['variable']}" for p in SEEP_PANELS]:
        if wanted not in sources:
            fails.append(f"the seepage report has no {wanted!r} figure: {sources}")
    # The mesh figure is captioned for what it carries. This sample carries both
    # boundary types, so its caption names them; a mesh carrying neither is
    # captioned "Seepage mesh" (see test_seep_boundaries_not_on_record).
    mesh_caption = next((f.caption for f in report.figures()
                         if f.source == "seepage bc1 mesh"), None)
    if mesh_caption != "Seepage mesh and boundary conditions":
        fails.append(f"the mesh of a model carrying both boundary types is "
                     f"captioned {mesh_caption!r}")
    if drawn != 4 + len(SEEP_PANELS):
        fails.append(f"the seepage report drew {drawn} figures, expected the "
                     f"model, the unsaturated conductivity curves against "
                     f"suction and against pressure head, the mesh with its "
                     f"boundary conditions, and the {len(SEEP_PANELS)} fields "
                     f"the solve produced")

    # Every result panel is introduced by a sentence that cites it: a field drawn
    # and never named is a page the reader is sent to by nothing.
    numbered = {f.source: f.number for f in report.figures()}
    said = " ".join(_seep_results_prose(report))
    for panel in SEEP_PANELS:
        number = numbered.get(f"seepage bc1 {panel['variable']}")
        if number is None:
            continue
        if f"Figure {number}" not in said:
            fails.append(f"the {panel['variable']!r} panel is Figure {number} and "
                         f"the results cite it nowhere: {said!r}")

    # The unsaturated curves stand in the inputs, after the properties table
    # whose parameters they draw and before the mesh they are solved on, and the
    # two conventions stand together: they are one statement of the conductivity
    # model, and a page between them makes them two. The mesh follows under a
    # heading of its own — what all of it was discretized onto, met after it.
    mesh_sub = next((c for c in seep.children if c.title == "Seepage Mesh"), None)
    if inputs is not None:
        order = [b.source if b.kind == "figure" else b.kind
                 for b in inputs.blocks if b.kind in ("figure", "table")]
        order += [b.source if b.kind == "figure" else b.kind
                  for b in (mesh_sub.blocks if mesh_sub else [])
                  if b.kind in ("figure", "table")]
        want = ["seep model", "table", "seep kr", "seep kr_head",
                "seepage bc1 mesh"]
        if order != want:
            fails.append(f"the seepage inputs are ordered {order}, not {want}")
        if mesh_sub is None:
            fails.append("the seepage mesh has no heading of its own")

    # Both conductivity figures are cited, from one sentence that presents them
    # as the same models in the two conventions. A figure drawn and not named is
    # a page the reader is never sent to.
    numbers = {f.source: f.number for f in report.figures()}
    citing = [b for b in report.blocks("prose")
              if all(f"Figure {numbers[s]}" in b.text
                     for s in ("seep kr", "seep kr_head") if s in numbers)
              and "conductivity" in b.text]
    if not citing:
        fails.append("no sentence cites both conductivity figures: "
                     + repr([b.text for b in report.blocks("prose")
                             if "conductivity" in b.text]))
    else:
        said = citing[0].text
        for named in ("matric suction", "pressure head"):
            if named not in said:
                fails.append(f"the sentence citing the conductivity figures "
                             f"does not name {named!r}: {said!r}")
        from xslope.report import cite_anchor
        targets = [t for _text, t in citing[0].links]
        for source in ("seep kr", "seep kr_head"):
            if f"#{cite_anchor('Figure', numbers[source])}" not in targets:
                fails.append(f"the sentence naming the {source!r} figure does "
                             f"not link to it: {targets}")

    # The two are the same curves under two sign conventions, so a material keeps
    # its color and its dash across the pair. Read off the artists, not asserted
    # of the code that draws them: the report ships PNGs, and a color that only
    # matches in the source is a pair the reader cannot follow.
    from xslope.plot import plot_material_kr_set
    from xslope.report import _kr_materials
    import matplotlib.figure as mplfig
    import numpy as np
    mats = _kr_materials(_slope_data)
    if not mats:
        fails.append("the seepage sample carries no unsaturated material, so "
                     "the conductivity figures prove nothing")
    else:
        keyed = []
        for abscissa in ("suction", "head"):
            f = mplfig.Figure()
            plot_material_kr_set(mats, fig=f, abscissa=abscissa,
                                 show_legend=False)
            ax = f.axes[0]
            keyed.append({ln.get_label(): (ln.get_color(), ln.get_linestyle())
                          for ln in ax.get_lines()})
        if keyed[0] != keyed[1]:
            fails.append(f"the conductivity pair keys its materials "
                         f"differently: {keyed[0]} against {keyed[1]}")
        if len(keyed[0]) < 1:
            fails.append("the conductivity figures drew no labelled curve")

    # And they are the same numbers: the head figure is the suction figure
    # mirrored, drawn over the negative abscissa and reaching zero, where the
    # material is saturated.
    if mats:
        f = mplfig.Figure()
        plot_material_kr_set(mats, fig=f, abscissa="head", show_legend=False,
                             unit_labels={"length": "ft"})
        ax = f.axes[0]
        if "pressure head" not in ax.get_xlabel():
            fails.append(f"the head figure's abscissa is labelled "
                         f"{ax.get_xlabel()!r}")
        if "(ft)" not in ax.get_xlabel():
            fails.append(f"the head abscissa carries no length unit: "
                         f"{ax.get_xlabel()!r}")
        lo, hi = ax.get_xlim()
        if not (lo < 0 and abs(hi) < 1e-9):
            fails.append(f"the head figure spans {lo:g} to {hi:g}, not the "
                         f"negative range up to saturation at zero")
        g = mplfig.Figure()
        plot_material_kr_set(mats, fig=g, abscissa="suction", show_legend=False)
        by_label = {ln.get_label(): ln for ln in g.axes[0].get_lines()}
        for line in ax.get_lines():
            twin = by_label.get(line.get_label())
            if twin is None:
                continue
            x, y = line.get_data()
            u, v = twin.get_data()
            if not (np.allclose(x, -np.asarray(u))
                    and np.allclose(y, np.asarray(v))):
                fails.append(f"the head curve for {line.get_label()!r} is not "
                             f"the suction curve mirrored")
            if abs(float(y[0]) - 1.0) > 1e-6:
                fails.append(f"the head curve for {line.get_label()!r} does not "
                             f"reach kr = 1 at saturation: kr = {float(y[0]):g}")

    fails += _kr_ordinate_scale_checks()

    # Every panel is a field plot of the same solve: no element edges over any of
    # them, the base material the flow lines are scaled to is chosen rather than
    # left at one, and each panel asks for its own variable and only that variable's
    # overlay.
    from xslope.plot_seep import flownet_base_material
    calls = _plot_seep_calls(SEEP_XLSX)
    by_variable = {kw.get("variable"): kw for kw, _sd, _sol in calls}
    if sorted(by_variable) != sorted(p["variable"] for p in SEEP_PANELS):
        fails.append(f"the seepage results drew {sorted(by_variable)}, not the "
                     f"{[p['variable'] for p in SEEP_PANELS]} SEEP_PANELS names")
    chosen = flownet_base_material(bundle["seep_data"], bundle["solution"])
    for panel in SEEP_PANELS:
        passed = by_variable.get(panel["variable"])
        if passed is None:
            continue
        if passed.get("mesh") is not False:
            fails.append(f"the {panel['variable']!r} panel is drawn with "
                         f"mesh={passed.get('mesh')!r}; element edges chop the "
                         f"contours and hide the field")
        if passed.get("base_mat") != chosen:
            fails.append(f"the {panel['variable']!r} panel is scaled to material "
                         f"{passed.get('base_mat')!r}, not the {chosen} its "
                         f"conductivities call for")
    # The velocity arrows are ON the velocity magnitude, and on nothing else: a
    # field's direction belongs over its own magnitude, and arrows across a
    # gradient plot are a second quantity nobody asked that figure for. Read off
    # the figure, and against the sentence printed above it — a panel described as
    # carrying vectors and drawn without them describes a figure that is not there.
    numbered = {f.source: f.number for f in report.figures()}
    sentences = {}
    for block in report.blocks("prose"):
        for panel in SEEP_PANELS:
            number = numbered.get(f"seepage bc1 {panel['variable']}")
            # The sentence is the one that refers to this figure by number, in
            # whatever register it is written in — the panel sentences are
            # noun-first ("The pore pressure field is shown in Figure 5.").
            if number is not None and f"Figure {number}" in block.text \
                    and "shown in" in block.text:
                sentences[panel["variable"]] = block.text
    for kw, sd, sol in calls:
        variable = kw.get("variable")
        panel = next((p for p in SEEP_PANELS if p["variable"] == variable), None)
        if panel is None:
            continue
        arrows = "VELOCITY" in _figure_gids(sd, sol, kw)
        if arrows is not (variable == "v_mag"):
            fails.append(f"the {variable!r} panel "
                         f"{'draws' if arrows else 'draws no'} velocity vectors")
        said = sentences.get(variable, "")
        if not said:
            fails.append(f"the {variable!r} panel's figure is introduced by no "
                         f"sentence")
            continue
        claimed = "vector" in said
        if claimed is not arrows:
            fails.append(f"the {variable!r} panel's sentence "
                         f"{'names' if claimed else 'does not name'} the velocity "
                         f"vectors, and the figure "
                         f"{'draws' if arrows else 'draws none'}: {said!r}")

    # Every sentence about a figure says what the figure contains, and nothing
    # else: the gradient panel was introduced as showing where the section is
    # checked against piping, which is a claim about WHERE the largest gradient
    # falls — and on this solve it falls inside the section, not at the exit face a
    # piping check is made on. A sentence that names a place, a use or a
    # consequence cannot be held against the figure, so none is written.
    for variable, said in sorted(sentences.items()):
        for word in ("piping", "checked against", "read together", "taken "
                     "against", "effective stress", "so the ", "which is",
                     "whose largest"):
            if word in said:
                fails.append(f"the {variable!r} panel's sentence says "
                             f"{word!r}, which is not something the figure "
                             f"contains: {said!r}")

    # Each option carries its own figures, and switching one off takes only
    # those. The conductivity option governs the pair — the same models in the
    # two conventions are one statement, and half of it is not a report of the
    # conductivity model — so it takes two; the flow net takes the head panel, and
    # the field-plot option takes the other three.
    variable_sources = tuple(f"seepage bc1 {p['variable']}" for p in SEEP_PANELS
                            if p["option"] == "seep_variable_figures")
    for option, gone in (("seep_inputs_figure", ("seep model",)),
                         ("seep_kr_figure", ("seep kr", "seep kr_head")),
                         ("seep_mesh_figure", ("seepage bc1 mesh",)),
                         ("seep_flownet", ("seepage bc1 head",)),
                         ("seep_variable_figures", variable_sources)):
        off = _engine_report("seep", options={option: False})
        got = [f.source for f in off.figures()]
        for source in gone:
            if source in got:
                fails.append(f"{option}=False still drew the {source!r} figure")
        if len(got) != drawn - len(gone):
            fails.append(f"{option}=False left {len(got)} figures, not the "
                         f"other {drawn - len(gone)}: {got}")
        planned, off_drawn = _planned_matches(off, "seep",
                                              options={option: False})
        if planned != off_drawn:
            fails.append(f"{option}=False planned {planned} figures and built "
                         f"{off_drawn}")

    # A model that carries no unsaturated parameters — a confined problem, where
    # the flow never leaves the saturated zone — has no curve to draw, and
    # NEITHER figure is drawn: it is absent rather than blank, with the sentence
    # that would cite it.
    saturated = copy.deepcopy(_slope_data)
    for m in saturated.get("materials") or []:
        m.update(kr0=0.0, h0=0.0, vg_a=0.0, vg_n=0.0)
    dry = _built_report(saturated, {"seep": bundle},
                        {"input_path": SEEP_XLSX, "lem": False,
                         "pd_figure": False})
    for source in ("seep kr", "seep kr_head"):
        if source in [f.source for f in dry.figures()]:
            fails.append(f"a model with no unsaturated material still drew the "
                         f"{source!r} conductivity figure")
    dry_text = " ".join(_prose(dry))
    for named in ("matric suction", "pressure head"):
        if named in dry_text:
            fails.append(f"a model with no unsaturated material still describes "
                         f"the {named} conductivity figure it did not draw")
    planned, drawn = _planned_matches(dry, "seep", slope_data=saturated)
    if planned != drawn:
        fails.append(f"a saturated model planned {planned} figures and built "
                     f"{drawn}")
    return fails


def test_report_figures_carry_their_axis_labels_whole():
    """A figure the report writes is cropped to include its axis labels.

    ``bbox_inches="tight"`` measures a y-axis label with its height collapsed to a
    point, and takes the crop from the same measurement — so a label taller than
    the axes it belongs to is cropped through. A velocity or gradient colorbar on a
    wide, shallow section is short and its label is not: the figure printed
    "Velocity Magnitude (ft/c" for "Velocity Magnitude (ft/day)".

    Held on the crop the report really asks for, against the label as really drawn,
    and against the default crop that does not hold it — which is what makes the
    measurement mean anything on this figure.
    """
    fails = []
    import matplotlib.figure as mplfig
    from xslope.report import _render, resolve_options

    _slope_data, bundle = _seep_bundle()
    opts = resolve_options({"input_path": SEEP_XLSX})

    def draw(fig):
        from xslope.plot_seep import plot_seep_solution
        plot_seep_solution(bundle["seep_data"],
                           dict(bundle["solution"], unconfined=True), fig=fig,
                           show_title=False, mesh=False, variable="v_mag",
                           vectors=True)

    seen = {}
    real = mplfig.Figure.savefig

    def spy(self, *args, **kwargs):
        seen["figure"] = self
        seen["extra"] = kwargs.get("bbox_extra_artists")
        seen["bbox_inches"] = kwargs.get("bbox_inches")
        return real(self, *args, **kwargs)

    mplfig.Figure.savefig = spy
    tmp = tempfile.mkdtemp(prefix="xslope_crop_")
    try:
        with contextlib.redirect_stdout(io.StringIO()):
            written = _render(draw, os.path.join(tmp, "panel.png"), opts)
    finally:
        mplfig.Figure.savefig = real
    if not written:
        fails.append("the figure was not written at all")
        return fails
    if seen.get("bbox_inches") != "tight":
        fails.append(f"the figure is written with bbox_inches="
                     f"{seen.get('bbox_inches')!r}, so this check measures the "
                     f"wrong crop")
        return fails

    fig = seen["figure"]
    renderer = fig.canvas.get_renderer()
    crop = fig.get_tightbbox(renderer, bbox_extra_artists=seen.get("extra"))
    bare = fig.get_tightbbox(renderer)
    labels = [ax.yaxis.label for ax in fig.axes if ax.yaxis.label.get_text()]
    if not labels:
        fails.append("the figure carries no y-axis label, so the crop this check "
                     "is about is never taken")
    held = False
    for label in labels:
        box = label.get_window_extent(renderer)
        lo, hi = box.y0 / fig.dpi, box.y1 / fig.dpi        # pixels → inches
        if crop.y0 > lo + 1e-6 or crop.y1 < hi - 1e-6:
            fails.append(f"{label.get_text()!r} spans {lo:.2f}–{hi:.2f} in and the "
                         f"figure is cropped to {crop.y0:.2f}–{crop.y1:.2f}, so it "
                         f"is printed with its end cut off")
        if bare.y0 > lo + 1e-6 or bare.y1 < hi - 1e-6:
            held = True
    if not held:
        fails.append("every label on this figure fits the crop matplotlib takes on "
                     "its own, so naming them as extra artists proves nothing")

    # The legend went with them: the extras REPLACE the crop's default artists, and
    # a figure cropped to its labels alone loses the legend under the plot.
    from matplotlib.legend import Legend
    legends = [a for a in fig.findobj(Legend) if a.get_visible()]
    if not legends:
        fails.append("the figure carries no legend, so losing one could not show "
                     "here")
    for legend in legends:
        box = legend.get_window_extent(renderer)
        lo, hi = box.y0 / fig.dpi, box.y1 / fig.dpi
        if crop.y0 > lo + 1e-6 or crop.y1 < hi - 1e-6:
            fails.append(f"the legend spans {lo:.2f}–{hi:.2f} in and the figure is "
                         f"cropped to {crop.y0:.2f}–{crop.y1:.2f}: naming the "
                         f"labels dropped the crop's own artists")
    return fails


def test_seep_panels_mirror_the_seep_view():
    """The report presents every field the seepage results view offers, in the
    same order, and each is a figure the reader can switch off.

    An engine's section carries that engine's full figure sequence: a report that
    printed the flow net alone documented a run the user had read four ways on the
    screen. The two lists are held against each other so neither can grow a field
    the other does not have. The finite element pair is the same arrangement and is
    held the same way, by :func:`test_fem_panels_mirror_the_fem_view` — which this
    docstring asserted before that check existed, and this one never looked at a
    finite element panel.
    """
    fails = []
    from studio.dialogs import SEEP_VARIABLES
    from studio.report_dialog import CONTENT_TREE
    from xslope.report import DEFAULT_OPTIONS, SEEP_PANELS

    def disagree(printed, offered):
        """What is wrong with a report list against a view list — the comparison
        itself, so it can be run on lists that ARE wrong."""
        if list(printed) == list(offered):
            return []
        missing = [v for v in offered if v not in printed]
        extra = [v for v in printed if v not in offered]
        if missing or extra:
            return [f"the report prints {printed} and the seepage results view "
                    f"offers {offered}: {missing} is on the view and not in the "
                    f"report, {extra} the other way about"]
        return [f"the report prints {printed} in a different order from the "
                f"view's {offered}"]

    offered = [key for key, _label in SEEP_VARIABLES]
    printed = [panel["variable"] for panel in SEEP_PANELS]
    fails += disagree(printed, offered)

    # And by NAME, not by key alone: the view called the fourth field "Gradient
    # magnitude" while the report captioned the same figure "Hydraulic gradient
    # magnitude", so a reader who picked a field on the screen and went looking for
    # it in the report was looking for a name nothing there carried. Compared
    # case-insensitively, because the two write for different places — a control
    # label and a sentence — and the WORDS are what a reader matches on.
    named = [label.lower() for _key, label in SEEP_VARIABLES]
    fields = [panel["field"].lower() for panel in SEEP_PANELS]
    fails += disagree(fields, named)
    for panel, label in zip(SEEP_PANELS, named):
        if panel["variable"] != "head" and panel["caption"].lower() != label:
            fails.append(f"the {panel['variable']!r} figure is captioned "
                         f"{panel['caption']!r} and the field is called {label!r}")

    # The comparison really discriminates: a panel dropped from either side, the
    # same four in a different order, and a field renamed on one side, are each
    # caught. A parity check that compares two lists it has already agreed on
    # proves nothing.
    for mutant, what in ((disagree(printed[:-1], offered), "a panel dropped from "
                          "the report"),
                         (disagree(printed, offered[:-1]), "a variable dropped "
                          "from the view"),
                         (disagree(list(reversed(printed)), offered), "the panels "
                          "reordered"),
                         (disagree(fields[:-1] + ["gradient magnitude"], named),
                          "a field renamed on one side")):
        if not mutant:
            fails.append(f"the parity comparison passes {what}")

    # Each panel is a switch the user has, on by default, under the seepage branch:
    # a figure nobody can turn off is not part of a composed report.
    rows = {}
    for key, _label, _tip, children in CONTENT_TREE:
        for child_key, _l, _t in children:
            rows[child_key] = key
    for panel in SEEP_PANELS:
        option = panel["option"]
        if option not in DEFAULT_OPTIONS:
            fails.append(f"the {panel['variable']!r} panel is printed under "
                         f"{option!r}, which is not an option the builder reads")
        elif DEFAULT_OPTIONS[option] is not True:
            fails.append(f"{option!r} is off by default, so the "
                         f"{panel['variable']!r} field is not in a report unless "
                         f"it is asked for")
        if option not in rows:
            fails.append(f"the dialog offers no {option!r} row")
        elif rows[option] != "seep":
            fails.append(f"the {option!r} row is a child of {rows[option]!r}, not "
                         f"of the seepage section that prints it")

    # And each is its own figure with its own name: two panels sharing a caption or
    # a source would be one figure printed twice as far as the report is concerned.
    for field in ("caption", "field"):
        names = [panel[field] for panel in SEEP_PANELS]
        if len(set(names)) != len(names):
            fails.append(f"two seepage panels share a {field}: {names}")
    for panel in SEEP_PANELS[1:]:
        if not panel["draws"]:
            fails.append(f"the {panel['variable']!r} panel has no sentence to "
                         f"introduce it, so its figure stands unexplained")
    if SEEP_PANELS[0]["draws"] is not None:
        fails.append("the head panel carries its own sentence as well as the one "
                     "the results paragraph writes for it")
    return fails


def test_seep_mesh_legend_names_its_boundaries():
    """The mesh figure names its boundary conditions the way the paragraph beside it
    does, and carries no code numbers.

    The legend read "Fixed Head (bc_type=1)" and "Exit Face (bc_type=2)" under a
    paragraph saying that so many nodes carry a SPECIFIED head: two names for one
    boundary in one page opening, and an internal array code printed in a figure an
    engineer signs.
    """
    fails = []
    import matplotlib.pyplot as plt
    from xslope.plot_seep import plot_seep_data

    _slope_data, bundle = _seep_bundle()
    seep_data = bundle["seep_data"]
    fig = plt.figure(figsize=(6, 4))
    try:
        with contextlib.redirect_stdout(io.StringIO()):
            plot_seep_data(seep_data, fig=fig, show_title=False, show_bc=True)
        from matplotlib.legend import Legend
        labels = [t.get_text() for legend in fig.findobj(Legend)
                  for t in legend.get_texts()]
    finally:
        plt.close(fig)

    if not labels:
        fails.append("the mesh figure carries no legend, so its names are untested")
    for wanted in ("Specified head", "Exit face"):
        if wanted not in labels:
            fails.append(f"the mesh legend does not name its boundary "
                         f"{wanted!r}: {labels}")
    for label in labels:
        if "bc_type" in label:
            fails.append(f"the mesh legend prints the internal code in "
                         f"{label!r}")
        if "Fixed Head" in label:
            fails.append(f"the mesh legend calls a specified-head boundary "
                         f"{label!r}, which is not what the paragraph beside it "
                         f"calls the same nodes")

    # And the report's own words for those boundaries are the legend's words: the
    # figure and the sentence pointing at it are read together.
    said = " ".join(_seep_results_prose(_engine_report("seep")))
    for label in ("Specified head", "Exit face"):
        if label.lower() not in said.lower():
            fails.append(f"the mesh legend says {label!r} and the results never "
                         f"use the term: {said!r}")
    return fails


def test_the_mesh_sentence_counts_what_the_mesh_holds():
    """The finite element mesh sentence counts the mesh out, including the
    one-dimensional elements, and says they stand on the same nodes.

    "It consists of 3,180 nodes, 1,521 six-node quadratic triangular elements
    (tri6), and N 1D elements for the piles. The 2D and 1D elements share the
    same nodes and ..." — the owner's own shape. The counts are read off the
    fem_data the run was solved on, so the sentence cannot state a mesh that is
    not the one in the figure above it, and the shared-node claim is checked
    against the arrays rather than taken on trust.
    """
    fails = []
    import numpy as np

    def sentence(report):
        for b in report.blocks("prose"):
            if "The finite element mesh constructed for the problem" in b.text:
                return b.text
        return ""

    for label, xlsx, member, element in (
            ("piles", FEM_PILES_XLSX, "piles", "beam"),
            ("reinforcement", FEM_REINF_XLSX, "reinforcement lines", "truss"),
            ("no members", FEM_XLSX, None, None)):
        _sd, bundle = _fem_1d_bundle(xlsx) if member else _fem_bundle()
        fem_data = bundle["fem_data"]
        said = sentence(_engine_report("fem", xlsx=xlsx))
        if not said:
            fails.append(f"{label}: the mesh figure is introduced by no sentence")
            continue
        for want in (f"{len(fem_data['nodes']):,} nodes",
                     f"{len(fem_data['elements']):,} "):
            if want not in said:
                fails.append(f"{label}: the mesh sentence does not count "
                             f"{want!r}: {said!r}")
        types = sorted({int(t) for t in fem_data["element_types"]})
        if types == [6] and "six-node quadratic triangular elements (tri6)" \
                not in said:
            fails.append(f"{label}: the element kind is not named in full: "
                         f"{said!r}")
        if member is None:
            if "1D" in said or "share the same nodes" in said:
                fails.append(f"{label}: a mesh with no one-dimensional element "
                             f"is credited with them: {said!r}")
            continue

        # The 1D elements, counted per kind and named for what they are.
        e1 = fem_data.get("elements_1d")
        n_1d = 0 if e1 is None else len(e1)
        if not n_1d:
            fails.append(f"{label}: the fixture carries no 1D element, so the "
                         f"count below proves nothing")
            continue
        mask = np.asarray(fem_data.get("pile_elem_mask", np.zeros(n_1d)), bool)
        n = int(mask.sum()) if member == "piles" else int((~mask).sum())
        want = f"{n:,} two-node {element} elements for the {member}"
        if want not in said:
            fails.append(f"{label}: the mesh sentence does not count the 1D "
                         f"elements as {want!r}: {said!r}")
        if "share the same nodes" not in said:
            fails.append(f"{label}: the mesh sentence does not say the 2D and 1D "
                         f"elements share nodes: {said!r}")

        # And the claim is TRUE of the arrays: every node a 1D element is built
        # on indexes the same node array the 2D elements do, and is a node of a
        # 2D element. A sentence stating a formulation fact must be falsifiable
        # by the formulation.
        used = np.unique(np.asarray(e1)[:, :2])
        if used.max() >= len(fem_data["nodes"]):
            fails.append(f"{label}: a 1D element indexes a node outside the 2D "
                         f"mesh, so 'share the same nodes' is false")
        elif not np.all(np.isin(used, np.unique(np.asarray(fem_data["elements"])))):
            fails.append(f"{label}: a 1D element stands on a node no 2D element "
                         f"uses, so 'share the same nodes' is false")
    return fails


def test_the_mesh_is_counted_out_once():
    """A report states the size of the mesh once.

    The node and element counts stood in the traceability stamp, again in the
    analysis inputs, and again in the sentence under the mesh figure — three times
    in a report of one solved set, and four in a report of two, because each solved
    set restated them under its own mesh figure.

    WHERE the one statement stands differs by engine, and each has a reason. The
    finite element mesh is drawn once per report, so its counts are read out in
    the sentence that introduces that figure — the register the owner set, which
    also carries the one-dimensional element counts and the shared-node fact,
    neither of which fits a key-value line. The seepage mesh may be drawn once per
    boundary condition set, so its counts stay in the analysis inputs above, where
    they are stated once however many sets are documented. A section that draws no
    mesh figure at all states them in its key-value lines, and the traceability
    stamp carries them only for a report with no analysis section to put them in.
    """
    fails = []
    from xslope.report import mesh_summary

    for label, report, slope_data, stated_in in (
            ("a seepage report", _engine_report("seep"),
             load_slope_data_cached(SEEP_XLSX), "keyvalues"),
            ("a report of two boundary condition sets",
             _seep_report(RAPID_SEEP_XLSX),
             load_slope_data_cached(RAPID_SEEP_XLSX), "keyvalues"),
            ("a strength reduction report", _engine_report("fem"),
             load_slope_data_cached(FEM_XLSX), "prose")):
        mesh = slope_data.get("mesh") or {}
        if not mesh_summary(mesh):
            fails.append(f"{label} is of a model with no mesh, so counting the "
                         f"statements of it proves nothing")
            continue
        # The node count, which both the key-value shorthand and the mesh
        # sentence's long form spell the same way — so one search counts the
        # statements whichever form carries them.
        counted = f"{len(mesh['nodes']):,} nodes"
        rows = [f"{l}: {v}" for b in report.blocks("keyvalues") for l, v in b.items
                if counted in str(v)]
        prose = [b.text for b in report.blocks("prose") if counted in b.text]
        if len(rows) + len(prose) != 1:
            fails.append(f"{label} states the mesh {len(rows) + len(prose)} "
                         f"times: {rows + prose}")
        elif stated_in == "keyvalues" and not rows:
            fails.append(f"{label} states the mesh in prose rather than among the "
                         f"analysis inputs: {prose}")
        elif stated_in == "prose" and not prose:
            fails.append(f"{label} states the mesh in a key-value line rather "
                         f"than in the sentence that introduces its figure: "
                         f"{rows}")

    # Two engines run on ONE mesh state it once. The seepage section counts it;
    # the finite element section, which is built after, cites that count instead
    # of setting the same numbers again four pages later.
    slope_data, solutions = _restored(JOHNSON_XLSX)
    if not (solutions.get("seep") and solutions.get("fem")):
        fails.append("johnson_res no longer carries both engines' solutions, so "
                     "the shared-mesh case is untested")
    else:
        from xslope.report import mesh_summary as _summary
        seep_mesh = _summary(solutions["seep"][0]["seep_data"])
        fem_mesh = _summary(solutions["fem"]["fem_data"])
        if not seep_mesh or seep_mesh != fem_mesh:
            fails.append(f"the two engines report different meshes "
                         f"({seep_mesh!r} vs {fem_mesh!r}); the shared case is "
                         f"not what this model is")
        quiet = {"input_path": JOHNSON_XLSX, "lem": False, "pd_figure": False,
                 "seep_inputs_figure": False, "seep_mesh_figure": False,
                 "seep_kr_figure": False, "seep_figures": False,
                 "fem_inputs_figure": False, "fem_mesh_figure": False,
                 "fem_figure": False}
        both = _built_report(slope_data, solutions, quiet)
        counted = f"{len(solutions['seep'][0]['seep_data']['nodes']):,} nodes"
        rows = [f"{l}: {v}" for b in both.blocks("keyvalues") for l, v in b.items
                if counted in str(v)]
        prose = [b.text for b in both.blocks("prose") if counted in b.text]
        if len(rows) + len(prose) != 1:
            fails.append(f"a two-engine report on one mesh states it "
                         f"{len(rows) + len(prose)} times: {rows + prose}")
        cited = [b.text for b in both.blocks("prose")
                 if "the same mesh" in b.text]
        if len(cited) != 1:
            fails.append(f"the second engine does not cite the mesh the first "
                         f"counted: {cited}")
        elif "Section" not in cited[0]:
            fails.append(f"the citation names no section: {cited[0]!r}")

        # Two DIFFERENT meshes are two facts, and both are counted.
        import numpy as np
        other = dict(solutions["fem"])
        other["fem_data"] = dict(other["fem_data"],
                                 element_types=np.full(
                                     len(other["fem_data"]["elements"]), 3))
        split = _built_report(slope_data, dict(solutions, fem=other), quiet)
        counts = [v for b in split.blocks("keyvalues") for l, v in b.items
                  if l == "Mesh"]
        if len(counts) != 2:
            fails.append(f"two engines on different meshes counted {len(counts)} "
                         f"of them: {counts}")
        if any("one mesh" in b.text for b in split.blocks("prose")):
            fails.append("two different meshes were reported as one")

    # A model that carries a mesh and is reported with no analysis section of the
    # engine that reads it has nowhere else to say so, and the stamp says it.
    _slope_data, bundle = _seep_bundle()
    summary = mesh_summary(bundle["seep_data"])
    bare = _built_report(_slope_data, {"seep": [bundle]},
                         {"input_path": SEEP_XLSX, "lem": False, "seep": False,
                          "pd_figure": False})
    stated = [f"{l}: {v}" for b in bare.blocks("keyvalues") for l, v in b.items
              if summary in str(v)]
    if len(stated) != 1:
        fails.append(f"a report with no analysis section states the mesh "
                     f"{len(stated)} times: {stated}")
    return fails


def test_seep_head_figure_draws_the_boundary_water_levels():
    """The head figure carries the water level each specified-head boundary holds,
    and a model with no such boundary neither draws one nor is said to.

    The pool a flow net is driven by is otherwise readable only off the contour
    values. The overlay draws nothing where the boundary geometry is not there, so
    what is asked for and what the paragraph says the figure carries are decided by
    the one predicate that knows.
    """
    fails = []
    import copy
    from xslope.plot_seep import seep_has_bc_levels

    _slope_data, bundle = _seep_bundle()
    if not seep_has_bc_levels(bundle["seep_data"], bundle["solution"]):
        fails.append("the seepage sample carries no boundary water level, so the "
                     "overlay this check is about is never drawn")

    said = " ".join(_seep_results_prose(_engine_report("seep")))
    claim = "the water level held by each specified-head boundary"
    if claim not in said:
        fails.append(f"the head figure carries the boundary water levels and the "
                     f"results do not say so: {said!r}")

    # Asked for on the head figure and nowhere else — a pore pressure plot with a
    # waterline across it is a second, unexplained line on a field it is not the
    # level of — and really drawn there.
    for kw, sd, sol in _plot_seep_calls(SEEP_XLSX):
        variable = kw.get("variable")
        want = variable == "head"
        if bool(kw.get("show_bc_levels")) is not want:
            fails.append(f"the {variable!r} panel is drawn with show_bc_levels="
                         f"{kw.get('show_bc_levels')!r}")
        gids = _figure_gids(sd, sol, kw)
        if ("BC_LEVEL" in gids) is not want:
            fails.append(f"the {variable!r} panel draws {sorted(gids)}, and the "
                         f"water levels {'are missing' if want else 'are on it'}")

    # A model whose boundary geometry is not there draws nothing and claims
    # nothing: the same sample with its specified heads removed, so what changed is
    # the boundary and not the model.
    stripped = copy.deepcopy(bundle)
    (stripped["seep_data"].get("seepage_bc") or {})["specified_heads"] = []
    if seep_has_bc_levels(stripped["seep_data"], stripped["solution"]):
        fails.append("a boundary condition set with no specified head still "
                     "reports a water level to draw")
    bare = _seep_report(SEEP_XLSX, bundles=[stripped])
    if claim in " ".join(_seep_results_prose(bare)):
        fails.append("a figure with no water level on it is still said to draw one")
    for kw, sd, sol in _plot_seep_calls(SEEP_XLSX, bundles=[stripped]):
        if kw.get("show_bc_levels"):
            fails.append(f"the {kw.get('variable')!r} panel of a set with no "
                         f"specified head is drawn with show_bc_levels=True")
        if "BC_LEVEL" in _figure_gids(sd, sol, kw):
            fails.append(f"the {kw.get('variable')!r} panel of a set with no "
                         f"specified head draws a water level anyway")

    # The one the corpus ships whose boundaries are not on record at all: no
    # geometry, no overlay, no claim.
    nobc = _seep_report(NOBC_SEEP_XLSX)
    if claim in " ".join(_seep_results_prose(nobc)):
        fails.append("a solve with no boundaries on record is said to draw their "
                     "water levels")
    return fails


#: A confined seepage analysis: an impervious blanket over a pervious foundation,
#: solved with no exit face anywhere on it. The corpus's confined models are the
#: only ones that exercise the confined branch of the results paragraph, and until
#: one was reported the branch had never been built.
CONFINED_SEEP_XLSX = os.path.join(_REPO, "docs", "seep", "files",
                                  "xslope_clay_blanket.xlsx")

#: A model whose saved seepage solution is a nodal field with no boundary
#: conditions behind it — the mesh carries neither a specified head nor an exit
#: face, and the file records no flow rate. Both counts zero is the degenerate
#: case: it is not a confined problem, it is a solve whose boundaries are not on
#: record.
NOBC_SEEP_XLSX = os.path.join(_REPO, "docs", "verification", "files",
                              "rocscience", "rs2_9.xlsx")

#: A solved unconfined analysis whose saved solution carries no flow rate footer —
#: twelve of the corpus's forty do. The flow is the results subsection's only
#: number, and the stream function it would space its flow lines by is flat.
NOFLOW_SEEP_XLSX = os.path.join(_REPO, "docs", "verification", "files",
                                "rocscience", "rs2_67b.xlsx")


def _seep_report(xlsx, options=None, bundles=None):
    """A report of EVERY boundary condition set a model was solved for.

    ``_engine_report`` documents one set, which is what a model solved once has.
    A rapid drawdown model was solved for two, and a section that documents two
    is a different thing to check.

    ``bundles`` reports sets other than the ones shipped beside the model — how a
    check builds the case the corpus does not happen to contain.
    """
    from xslope.report import build_report, seep_bundles

    key = ("seep_all", xlsx)
    if key not in _ENGINE:
        slope_data, solutions = _restored(xlsx)
        _ENGINE[key] = (slope_data, seep_bundles(solutions))
    slope_data, shipped = _ENGINE[key]
    bundles = shipped if bundles is None else bundles
    opts = {"input_path": xlsx, "lem": False, "pd_figure": False}
    opts.update(FAST_FIGURES)
    opts.update(options or {})
    tmp = tempfile.mkdtemp(prefix="xslope_seepall_")
    with contextlib.redirect_stdout(io.StringIO()):
        return build_report(slope_data, {"seep": bundles}, opts, tmp)


#: The subsections of a seepage section that describe its INPUTS. The mesh has
#: a heading of its own — it is what the domain and its properties were
#: discretized onto, and it is met after them (the owner's sequencing) — so a
#: check reading "the inputs" reads both.
_SEEP_INPUT_TITLES = ("Analysis Inputs", "Seepage Mesh")


def _seep_results_prose(report):
    """The prose of every results subsection of a seepage section, as one string
    per subsection, in the order the section prints them."""
    sec = next((s for s in report.sections if s.title == "Seepage Analysis"), None)
    out = []
    for _lvl, node in (sec.walk() if sec else []):
        if node.title in _SEEP_INPUT_TITLES or node is sec:
            continue
        out.append(" ".join(b.text for b in node.blocks if b.kind == "prose"))
    return out


def _seep_inputs_prose(report):
    """The prose of a seepage section's inputs — its Analysis Inputs and the
    mesh subsection under it — as one string."""
    sec = next((s for s in report.sections if s.title == "Seepage Analysis"), None)
    return " ".join(b.text for c in (sec.children if sec else [])
                    if c.title in _SEEP_INPUT_TITLES
                    for b in c.blocks if b.kind == "prose")


def _plot_seep_calls(xlsx, options=None, bundles=None):
    """Every call the flow nets of a model make to ``plot_seep_solution``, as
    ``(kwargs, seep_data, solution)`` — how a check asks what the figure was drawn
    from rather than re-deriving it."""
    import xslope.plot_seep as ps

    real = ps.plot_seep_solution
    seen = []

    def spy(seep_data, solution, **kw):
        seen.append((dict(kw), seep_data, solution))
        return real(seep_data, solution, **kw)

    ps.plot_seep_solution = spy
    try:
        _seep_report(xlsx, dict(options or {},
                                seep_inputs_figure=False, seep_mesh_figure=False,
                                seep_kr_figure=False), bundles=bundles)
    finally:
        ps.plot_seep_solution = real
    return seen


def test_seep_solution_file_records_the_solve():
    """A saved seepage solution carries what the solve was, and a solution saved
    before it did still reads back.

    ``run_seepage_analysis`` answers three things no column can hold — whether the
    problem was solved unconfined, whether it converged, and the closure error it
    converged to. The file dropped all three, so a solution read back could not say
    whether its negative pore pressures were a phreatic surface or the ordinary
    suction of a saturated potential field, and the plotter guessed.
    """
    fails = []
    from xslope.seep import (build_seep_data, export_seep_solution,
                             import_seep_solution)
    from xslope.fileio import load_slope_data

    with contextlib.redirect_stdout(io.StringIO()):
        slope_data = load_slope_data(CONFINED_SEEP_XLSX)
        seep_data = build_seep_data(slope_data["mesh"], slope_data, seep_bc=1)
        shipped = import_seep_solution(
            seep_data, os.path.splitext(CONFINED_SEEP_XLSX)[0] + "_seep.csv")

    # The companion shipped beside the model was written by a solve that recorded all
    # three, so all three read back. The sample is the confined one, and what it
    # records is what makes its negative pore pressures readable: unconfined=False
    # says they are the ordinary suction of a saturated potential field, not a
    # phreatic surface the plotter would otherwise have guessed at.
    for key in ("unconfined", "converged", "closure_error"):
        if key not in shipped:
            fails.append(f"the confined sample's saved solution does not record "
                         f"{key!r}")
    if shipped.get("unconfined") is not False:
        fails.append(f"the confined sample reads back as unconfined="
                     f"{shipped.get('unconfined')!r}, but it is a confined problem")
    if shipped.get("converged") is not True:
        fails.append(f"the confined sample reads back as converged="
                     f"{shipped.get('converged')!r}")
    if shipped.get("flowrate") is None:
        fails.append("the confined sample's saved flow rate was lost")

    solved = dict(shipped, unconfined=False, converged=True,
                  closure_error=1.25e-8)
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "written_seep.csv")
        with contextlib.redirect_stdout(io.StringIO()):
            export_seep_solution(seep_data, solved, path)
            back = import_seep_solution(seep_data, path)
        for key, want in (("unconfined", False), ("converged", True)):
            if back.get(key) != want:
                fails.append(f"a solution exported with {key}={want!r} reads "
                             f"back as {back.get(key)!r}")
        if abs(back.get("closure_error", 0) - 1.25e-8) > 1e-14:
            fails.append(f"a solution exported with a closure error of 1.25e-08 "
                         f"reads back as {back.get('closure_error')!r}")
        if back.get("flowrate") != shipped["flowrate"]:
            fails.append(f"the flow rate did not survive the round trip: "
                         f"{back.get('flowrate')} for {shipped['flowrate']}")

        # The footer lines are comments appended after the flowrate line, so a file
        # written today is a file written before it plus those lines. Stripping them
        # IS a file written before the footer existed, and it still reads: the three
        # keys come back ABSENT rather than guessed, so unknown stays distinguishable
        # from recorded-as-False. Built by stripping rather than by pointing at a
        # companion in the corpus, because every shipped companion is now written by
        # a solve that records the footer.
        older = os.path.join(tmp, "older_seep.csv")
        with open(path) as f:
            kept = [l for l in f if not l.startswith(
                ("# Unconfined:", "# Converged:", "# Closure Error:"))]
        with open(older, "w") as f:
            f.writelines(kept)
        with contextlib.redirect_stdout(io.StringIO()):
            old = import_seep_solution(seep_data, older)
        for key in ("unconfined", "converged", "closure_error"):
            if key in old:
                fails.append(f"a solution file with no {key} footer reads back "
                             f"carrying {key!r} anyway")
        if old.get("flowrate") != shipped["flowrate"]:
            fails.append("stripping the solve footer lost the flow rate too")
        if len(old["head"]) != len(seep_data["nodes"]):
            fails.append("stripping the solve footer lost the nodal table")

        # A solution that answers none of the three — a pore pressure field
        # imported from another program — is written without the footer rather
        # than with a guess in it.
        bare = os.path.join(tmp, "bare_seep.csv")
        unknown = {k: v for k, v in shipped.items()
                   if k not in ("unconfined", "converged", "closure_error")}
        with contextlib.redirect_stdout(io.StringIO()):
            export_seep_solution(seep_data, unknown, bare)
        with open(bare) as f:
            wrote = [l.strip() for l in f if l.startswith("#")]
        if any(l.startswith(("# Unconfined:", "# Converged:", "# Closure Error:"))
               for l in wrote):
            fails.append(f"a solution carrying none of the solve facts was "
                         f"exported with them anyway: {wrote}")
    return fails


def test_shipped_seep_companions_record_their_solve():
    """Every steady seepage companion the corpus ships from a solve records that
    solve, so a report of an already-solved model states its convergence instead of
    going quiet.

    ``tools/make_seep_sidecars.py`` owns those files and is the list of them: a model
    added to its registry whose companion is not regenerated fails here rather than
    reaching a reader as a solution that cannot say whether it converged. The tool
    also carries the corpus's exclusions, each held to the opposite contract: none of
    them may carry a solve fact. They are excluded for reasons that differ — a vendor
    field or a transient frame no steady solve produced, a solved field whose builder
    writes a format of its own, a solved field deliberately shipped with no mesh — but
    a footer written by THIS tool onto any of them would be a claim about a solve it
    did not run.
    """
    fails = []
    tools_dir = os.path.join(_REPO, "tools")
    if tools_dir not in sys.path:
        sys.path.insert(0, tools_dir)
    try:
        import make_seep_sidecars as tool
    except Exception as exc:
        return [f"tools/make_seep_sidecars.py could not be imported: {exc!r}"]

    facts = ("# Unconfined:", "# Converged:", "# Closure Error:")
    for stem_rel, bcs, _settings in tool.MODELS:
        for bc in bcs:
            suffix = "_seep.csv" if bc == 1 else "_seep2.csv"
            path = os.path.join(_REPO, stem_rel + suffix)
            if not os.path.exists(path):
                fails.append(f"{os.path.basename(path)} is in the regeneration "
                             f"registry but is not in the corpus")
                continue
            with open(path) as f:
                footer = [l.strip() for l in f if l.startswith("#")]
            missing = [k for k in facts
                       if not any(l.startswith(k) for l in footer)]
            if missing:
                fails.append(f"{os.path.basename(path)} records no "
                             f"{', '.join(k.strip('# :') for k in missing)} — "
                             f"re-run tools/make_seep_sidecars.py")
            if any(l == "# Converged: False" for l in footer):
                fails.append(f"{os.path.basename(path)} records a solve that did "
                             f"not converge, so its flow rate is not the flow "
                             f"through the section")

    for stem_rel in tool.EXCLUDED:
        path = os.path.join(_REPO, stem_rel + "_seep.csv")
        if not os.path.exists(path):
            continue
        with open(path) as f:
            footer = [l.strip() for l in f if l.startswith("#")]
        claimed = [k for k in facts if any(l.startswith(k) for l in footer)]
        if claimed:
            fails.append(f"{os.path.basename(path)} is excluded from the steady "
                         f"regeneration but carries {', '.join(claimed)} — a solve "
                         f"fact for a solve that did not run")
    return fails


#: The dam with the solve facts taken back off its saved solution: a model whose
#: seepage field is on record and whose solve is not, which is what every companion
#: written before ``export_seep_solution`` grew its footer looks like, and what a
#: field imported from another program looks like today.
#:
#: MADE HERE, NOT COMMITTED. Every steady companion the corpus ships is written by
#: tools/make_seep_sidecars.py and records its solve — a footerless one kept in the
#: corpus as a fixture would be the single file contradicting that, and the guard
#: above (:func:`test_shipped_seep_companions_record_their_solve`) is what it would
#: have to be excused from. So the silent sample is cut from a copy instead: the
#: model and its companions go to a temporary directory and the copy's footer loses
#: its solve-fact lines, leaving the flow rate and the nodal table untouched.
_SILENT = {}


def _silent_seep_xlsx():
    """Path to the copied model, made once per run and removed at exit."""
    if "xlsx" not in _SILENT:
        _SILENT["tmp"] = tempfile.TemporaryDirectory(prefix="xslope_silent_")
        stem = _sidecar_copy(os.path.splitext(SEEP_XLSX)[0], _SILENT["tmp"].name)
        path = f"{stem}_seep.csv"
        with open(path) as f:
            lines = f.readlines()
        kept = [l for l in lines if not l.startswith(
            ("# Unconfined:", "# Converged:", "# Closure "))]
        if len(kept) == len(lines):
            raise AssertionError(f"{os.path.basename(path)} carries no solve "
                                 f"facts to take off — the silent sample is a "
                                 f"copy of a file that was already silent")
        with open(path, "w") as f:
            f.writelines(kept)
        _SILENT["xlsx"] = f"{stem}.xlsx"
    return _SILENT["xlsx"]


def test_seep_convergence_is_stated():
    """A solution that records whether it converged says so; one that does not
    records nothing about it.

    An unconfined solve is iterative and can stop short of closing, and a flow rate
    read off a solve that never closed is not the flow through the section. The
    report states convergence only where the solution carries it — a report that
    said "converged" over a solution that never recorded it would be a guess in the
    voice of a result.
    """
    fails = []
    xlsx = _silent_seep_xlsx()
    _slope_data, bundle = _seep_bundle(xlsx)

    quiet = " ".join(_seep_results_prose(_engine_report("seep", xlsx=xlsx)))
    if "converge" in quiet:
        if bundle["solution"].get("converged") is not None:
            fails.append("the sample's saved solution records convergence, so "
                         "the silent case this check is about is never taken")
        else:
            fails.append(f"a solution that records nothing about convergence is "
                         f"reported as converging anyway: {quiet!r}")

    def reported(**facts):
        edited = dict(bundle, solution=dict(bundle["solution"], **facts))
        return " ".join(_seep_results_prose(
            _engine_report("seep", bundle=edited, xlsx=xlsx)))

    closed = reported(converged=True, closure_error=3.25e-6)
    if "The solution converged" not in closed:
        fails.append(f"a converged solve does not say so: {closed!r}")
    if "3.25e-06" not in closed:
        fails.append(f"a converged solve does not state the closure error it "
                     f"converged to: {closed!r}")

    exact = reported(converged=True, closure_error=0.0)
    if "The solution converged." not in exact:
        fails.append(f"a direct solve does not say it converged: {exact!r}")
    if "closing the boundary inflow" in exact:
        fails.append(f"a solve that closed exactly is given a closure error to "
                     f"have closed to: {exact!r}")

    short = reported(converged=False, closure_error=12.5)
    if "did not converge" not in short:
        fails.append(f"a solve that never closed is not said to have failed: "
                     f"{short!r}")
    if "not reliable" not in short:
        fails.append(f"a solve that never closed still stands behind its flow: "
                     f"{short!r}")
    if "The flow through the section is" not in short:
        fails.append(f"the flow of a solve that never closed is withheld rather "
                     f"than qualified: {short!r}")
    return fails


def test_seep_confined_section():
    """A confined analysis is reported as a confined analysis, and its figure
    draws no phreatic surface.

    A confined solve is a single saturated Laplace solve whose head is a potential,
    not a water level. The p = 0 contour of such a field is an artifact, and a
    figure that draws it contradicts its own flow net — flow lines correctly filling
    the whole saturated domain under a line implying most of it is dry.
    """
    fails = []
    from xslope.report import _bc_counts
    from xslope.plot_seep import flownet_has_phreatic

    report = _seep_report(CONFINED_SEEP_XLSX)
    titles = report.section_titles()
    if (1, "Seepage Analysis") not in titles or (2, "Results") not in titles:
        fails.append(f"the confined model's report has no seepage results: "
                     f"{titles}")

    _slope_data, bundles = _ENGINE[("seep_all", CONFINED_SEEP_XLSX)]
    if len(bundles) != 1:
        fails.append(f"the confined sample restored {len(bundles)} boundary "
                     f"condition sets, expected one")
    seep_data = bundles[0]["seep_data"]
    n_head, n_exit = _bc_counts(seep_data)
    if n_exit:
        fails.append(f"the confined sample carries {n_exit} exit-face nodes, so "
                     f"it is not a confined problem and the branch this check is "
                     f"about is never taken")
    if not n_head:
        fails.append("the confined sample carries no specified-head node")

    text = " ".join(_seep_results_prose(report))
    # Discriminating both ways: "solved as an unconfined problem" contains the
    # word confined and a substring test on it alone passes over the very defect.
    if "solved as a confined problem" not in text:
        fails.append(f"the confined analysis is not reported as confined: "
                     f"{text!r}")
    if "unconfined" in text:
        fails.append(f"the confined analysis is reported as unconfined: {text!r}")
    if f"{n_head:,}" not in text:
        fails.append(f"the confined analysis does not state its {n_head} "
                     f"specified-head nodes: {text!r}")
    if "exit face" in text:
        fails.append(f"the confined analysis describes an exit face it has none "
                     f"of: {text!r}")
    if "phreatic surface" in text:
        fails.append(f"the confined analysis says its figure draws a phreatic "
                     f"surface: {text!r}")

    # The inputs of a confined analysis carry no unsaturated apparatus: a confined
    # solve is a single saturated Laplace solve and never evaluates k_r. The
    # section printed the unsaturated columns of the materials table and two pages
    # of flat k_r = 1.0 curves, two pages before the results correctly said every
    # node of the mesh flows saturated.
    sec = next((s for s in report.sections if s.title == "Seepage Analysis"), None)
    inputs = next((c for c in (sec.children if sec else [])
                   if c.title == "Analysis Inputs"), None)
    if inputs is None:
        fails.append("the confined model's report has no Analysis Inputs")
    else:
        table = next((b for b in inputs.blocks if b.kind == "table"), None)
        if table is None:
            fails.append("the confined inputs carry no material properties table")
        else:
            unsat_cols = [h for h in table.headers
                          if h in ("Unsaturated model", "k_r0", "a", "n")
                          or h.startswith("h₀")]
            if unsat_cols:
                fails.append(f"the materials table of a confined analysis carries "
                             f"the unsaturated columns {unsat_cols}")
            if not any(h.startswith("k₁") for h in table.headers):
                fails.append(f"the conductivities went with them: {table.headers}")
        lead = _seep_inputs_prose(report)
        # The mesh figure marks the boundary types the mesh has: this model has no
        # exit face, and was described as having every exit-face node marked.
        if "with every specified-head node marked" not in lead:
            fails.append(f"the mesh of a model whose only boundary is a specified "
                         f"head does not say so: {lead!r}")
        if "exit-face" in lead:
            fails.append(f"the inputs mark exit-face nodes on a mesh that carries "
                         f"none: {lead!r}")
        if "Above the phreatic surface" in lead:
            fails.append(f"the confined inputs say the conductivity is reduced "
                         f"above a phreatic surface the solve has none of: "
                         f"{lead!r}")
        if "unsaturated" in lead.lower():
            fails.append(f"the confined inputs describe an unsaturated model the "
                         f"solve never reads: {lead!r}")
    sources = [f.source for f in report.figures()]
    for gone in ("seep kr", "seep kr_head"):
        if gone in sources:
            fails.append(f"a confined analysis prints its {gone!r} figure — a "
                         f"page of flat k_r = 1.0 lines: {sources}")
    planned, drawn = _planned_matches(report, "seep", bundle=bundles,
                                      xlsx=CONFINED_SEEP_XLSX)
    if planned != drawn:
        fails.append(f"the confined report planned {planned} figures and built "
                     f"{drawn}")
    # The materials themselves DO carry unsaturated curves — this model is a linear
    # front model — so what rules them off is the confined solve and not an empty
    # column. Read unconfined, the very same model prints both.
    from xslope.report import _kr_materials
    if not _kr_materials(_slope_data):
        fails.append("the confined sample's materials carry no unsaturated curve "
                     "to begin with, so dropping the figures proves nothing")
    if not _kr_materials(_slope_data, [dict(bundles[0], solution=dict(
            bundles[0]["solution"], unconfined=True))]):
        fails.append("the same model read as an unconfined solve still draws no "
                     "conductivity curves")

    # The figure: the plotter is handed the decided answer rather than left to
    # default, and it draws no p = 0 contour.
    calls = _plot_seep_calls(CONFINED_SEEP_XLSX)
    from xslope.report import SEEP_PANELS
    if len(calls) != len(SEEP_PANELS):
        fails.append(f"the confined model drew {len(calls)} result panels, "
                     f"expected the {len(SEEP_PANELS)} SEEP_PANELS names")
    # EVERY panel, not only the flow net: the phreatic rule is a fact about the
    # solve, so a pore pressure or a gradient plot of a confined solve carries no
    # p = 0 contour either. The decided answer travels on the solution each panel is
    # drawn from, so no panel can fall back to the plotter's unconfined default.
    for kw, sd, sol in calls:
        variable = kw.get("variable")
        if sol.get("unconfined") is not False:
            fails.append(f"the {variable!r} panel of a confined solve is drawn "
                         f"with unconfined={sol.get('unconfined')!r}")
        if flownet_has_phreatic(sd, sol):
            fails.append(f"the confined {variable!r} panel draws a phreatic "
                         f"surface")
    if calls:
        by_variable = {kw.get("variable"): (kw, sd, sol) for kw, sd, sol in calls}
        head_kw, head_sd, head_sol = by_variable["head"]
        fig = next((f for f in report.figures()
                    if f.source == "seepage bc1 head"), None)
        gids = _figure_gids(head_sd, head_sol, head_kw)
        if "PHREATIC" in gids:
            fails.append(f"the confined flow net contains a phreatic contour: "
                         f"{sorted(gids)}")
        if "FLOWLINES" not in gids:
            fails.append(f"the confined flow net draws no flow lines: "
                         f"{sorted(gids)}")
        if fig is None or fig.caption != "Flow net":
            fails.append(f"the confined flow net is captioned "
                         f"{getattr(fig, 'caption', None)!r}")

        # This sample's own field never goes into suction, so its figures carry no
        # p = 0 contour whichever branch is taken, and the absence above would be
        # satisfied by a plotter that had never heard of confined flow. What decides
        # it is pinned on the same field driven negative, ON EVERY PANEL: read as
        # confined the contour stays off, and read as unconfined the very same
        # field draws it. Held on the flow net alone, a pore pressure plot could
        # still contradict the paragraph above it.
        import numpy as np
        for variable, (kw, sd, sol) in sorted(by_variable.items()):
            suction = dict(sol, u=np.asarray(sol["u"], dtype=float) - 1.0)
            held = _figure_gids(sd, dict(suction, unconfined=False), kw)
            loose = _figure_gids(sd, dict(suction, unconfined=True), kw)
            if "PHREATIC" in held:
                fails.append(f"the {variable!r} panel of a confined solve whose "
                             f"field goes into suction is still given a phreatic "
                             f"surface")
            if "PHREATIC" not in loose:
                fails.append(f"the {variable!r} panel of an unconfined field in "
                             f"suction is denied its phreatic surface, so "
                             f"suppressing it when confined proves nothing")
        if flownet_has_phreatic(seep_data, dict(
                by_variable["head"][2],
                u=np.asarray(by_variable["head"][2]["u"], dtype=float) - 1.0,
                unconfined=False)):
            fails.append("flownet_has_phreatic reports a phreatic surface for a "
                         "confined solve")
    return fails


def _longest_arrow(seep_data, solution, kwargs):
    """The length of the longest velocity arrow a flow figure draws, in the section's
    own units — how a check asks what an arrow MEANS rather than what it was asked
    for. ``None`` where the figure draws no arrows.

    The quiver is drawn with ``angles="xy", scale_units="xy", scale=1``, so its U
    and V are the arrow as it lands on the section.
    """
    import matplotlib.pyplot as plt
    import numpy as np
    from matplotlib.quiver import Quiver
    from xslope.plot_seep import plot_seep_solution

    fig = plt.figure(figsize=(6, 4))
    try:
        with contextlib.redirect_stdout(io.StringIO()):
            plot_seep_solution(seep_data, solution, fig=fig,
                               **{k: v for k, v in kwargs.items() if k != "fig"})
        for artist in fig.findobj(Quiver):
            u = np.asarray(artist.U, dtype=float)
            v = np.asarray(artist.V, dtype=float)
            if u.size:
                return float(np.max(np.hypot(u, v)))
        return None
    finally:
        plt.close(fig)


def _figure_gids(seep_data, solution, kwargs):
    """The gids of everything ``plot_seep_solution`` draws for one call — how a
    check reads what is IN a figure rather than what was asked for."""
    import matplotlib.pyplot as plt
    from xslope.plot_seep import plot_seep_solution

    fig = plt.figure(figsize=(6, 4))
    try:
        with contextlib.redirect_stdout(io.StringIO()):
            plot_seep_solution(seep_data, solution, fig=fig,
                               **{k: v for k, v in kwargs.items()
                                  if k not in ("fig",)})
        gids = set()
        for ax in fig.axes:
            for artist in ax.get_children():
                gid = artist.get_gid()
                if gid:
                    gids.add(gid)
        return gids
    finally:
        plt.close(fig)


#: The pages that write the term for a reader. Identifiers do not count — the
#: plotter's ``flowlines=`` keyword is an argument name and stays as it is — so
#: what is swept is prose, captions, dialog labels and documentation text.
_FLOW_LINE_PAGES = [
    os.path.join(_REPO, "docs", "verification", "seep.md"),
    os.path.join(_REPO, "docs", "studio", "interface.md"),
    os.path.join(_REPO, "docs", "usage", "claude", "index.md"),
    os.path.join(_REPO, "docs", "usage", "claude", "xslope.md"),
]

#: A line of documentation that only mentions the term inside code — a call
#: writing ``flowlines=True``, or a keyword named in backticks — is naming the
#: argument, not the lines.
_FLOW_LINE_CODE = ("flowlines=", "flowlines:", '"flowlines"', "'flowlines'",
                   "`flowlines`", "self.flowlines", ".flowlines",
                   "flownet_has_flowlines")


def _flowline_hits(texts):
    """``(where, text)`` for every piece of reader-facing text that spells the
    term as one word, ignoring the places it is an identifier."""
    out = []
    for where, text in texts:
        for line in text.splitlines():
            if "flowline" not in line.lower():
                continue
            stripped = line
            for code in _FLOW_LINE_CODE:
                stripped = stripped.replace(code, "")
            if "flowline" in stripped.lower():
                out.append((where, line.strip()))
    return out


def test_one_term_for_flow_lines():
    """The lines a flow net draws are called flow lines wherever they are named
    for a reader — in the report, in the dialog that builds it, and in the
    documentation of both.

    The plotter draws them under a legend entry reading "flow line", and the
    prose, the figure captions, the dialog row that switches the flow net on and
    the seepage documentation had all written "flowlines" against it. The keyword
    that turns them on is an identifier and keeps its spelling; everything written
    in words does not.
    """
    fails = []
    from studio.report_dialog import CONTENT_TREE

    texts = []

    # Every report's prose and every figure caption it prints.
    reports = [("the default report", _build()),
               ("the seepage report", _engine_report("seep")),
               ("the seepage report of a solve with no boundaries on record",
                _seep_report(NOBC_SEEP_XLSX)),
               ("the strength reduction report", _engine_report("fem")),
               ("the reinforcement report",
                _engine_report("fem", xlsx=FEM_REINF_XLSX)),
               ("the piles report", _engine_report("fem", xlsx=FEM_PILES_XLSX))]
    for method in CALC_METHODS:
        report, _bundle = _calc_report(method)
        if report is not None:
            reports.append((f"the {method} report", report))
    for where, report in reports:
        for block in report.blocks("prose"):
            texts.append((f"{where}, in its prose", block.text))
        for figure in report.figures():
            texts.append((f"{where}, in a figure caption", figure.caption))

    # Every row the dialog offers, by the label it wears and the description
    # under it.
    rows = 0
    for key, label, tip, children in CONTENT_TREE:
        for row_key, row_label, row_tip in [(key, label, tip)] + list(children):
            rows += 1
            texts.append((f"the dialog's {row_key!r} row", f"{row_label}\n{row_tip}"))
    if rows < 2:
        fails.append(f"the dialog sweep read {rows} rows, so it proves nothing")

    # And the documentation that describes both.
    for path in _FLOW_LINE_PAGES:
        if not os.path.exists(path):
            fails.append(f"{path} is swept for the term but is not there, so the "
                         f"sweep passes over it")
            continue
        with open(path, encoding="utf-8") as fh:
            texts.append((os.path.relpath(path, _REPO), fh.read()))

    for where, line in _flowline_hits(texts):
        fails.append(f"{where} writes the term as one word, against a legend "
                     f"reading 'flow line': {line!r}")

    # The mutation: each of the six spellings this replaced, put back one at a
    # time. A sweep that reads the right sources and matches nothing in them is
    # indistinguishable from a sweep that reads nothing.
    for spelling in (
            "3. `plot_seep_solution()` — head contours, flowlines, and phreatic "
            "surface",
            "The flow net (head contours and flowlines) for the s/T = 0.5 case:",
            "flowlines/vectors on the Seep · Solution view; plot type,",
            "and plots head contours with flowlines and the phreatic surface",
            "Head contours with the flowlines, and the phreatic surface where "
            "the problem is unconfined.",
            "'k1_by_mat' (optional, for flowline calculation).",
    ):
        if not _flowline_hits([("a reverted spelling", spelling)]):
            fails.append(f"the sweep passes {spelling!r}, which is the wording "
                         f"it exists to catch")

    # The identifiers really are left alone: the keyword the plotter takes reads
    # the same way it always did, and the sweep does not object to it.
    for identifier in ("plot_seep_solution(seep_data, solution, flowlines=True)",
                       'opts.get("flowlines", True)',
                       "self.flowlines = QCheckBox(\"Flow lines\")"):
        if _flowline_hits([("an identifier", identifier)]):
            fails.append(f"the sweep objects to {identifier!r}, which names the "
                         f"argument rather than the lines")
    return fails


def test_a_figure_sentence_agrees_with_its_subject():
    """The verb in a figure sentence follows the subject, not the number of
    things listed.

    "The pore pressure field is shown in Figure 5" and "The head contours and
    the flow lines are shown in Figure 4" are settled by the count. "The head
    contours", alone on a state carrying no phreatic surface and no boundary
    water level, is not: one item in the list, a plural noun, and a verb chosen
    by counting wrote "The head contours is shown in Figure 3". The caller that
    knows its subject's number says so.
    """
    fails = []
    from xslope.report import _shown_in

    for subjects, kwargs, want in (
            ("the pore pressure field", {}, "The pore pressure field is shown "
             "in Figure 5."),
            (["the head contours", "the flow lines"], {},
             "The head contours and the flow lines are shown in Figure 5."),
            (["the head contours"], {"plural": True},
             "The head contours are shown in Figure 5."),
            (["the phreatic surface"], {"plural": False},
             "The phreatic surface is shown in Figure 5."),
            ("the magnitude of the seepage velocity",
             {"when": "t = 0 days",
              "tail": ", with the velocity vectors drawn over it"},
             "The magnitude of the seepage velocity at t = 0 days is shown in "
             "Figure 5, with the velocity vectors drawn over it."),
            ([], {}, ""),
            (["", None], {}, "")):
        got = _shown_in(subjects, "Figure 5", **kwargs)
        if got != want:
            fails.append(f"{subjects!r} {kwargs}: {got!r}, not {want!r}")

    # And the mesh counts agree too: a one-element mesh is degenerate, but the
    # sentence would print "1 elements" rather than refuse to print.
    import numpy as np
    from xslope.report import mesh_counts, one_d_counts
    single = mesh_counts({"nodes": np.zeros((1, 2)),
                          "elements": np.zeros((1, 6)),
                          "element_types": np.array([6])})
    if single != ["1 node", "1 six-node quadratic triangular element (tri6)"]:
        fails.append(f"a one-element mesh is counted {single!r}")
    lone = one_d_counts({"elements_1d": np.zeros((1, 3), int),
                         "element_types_1d": np.array([2]),
                         "pile_elem_mask": np.array([True])})
    if lone != ["1 two-node beam element for the piles"]:
        fails.append(f"a single pile element is counted {lone!r}")
    return fails


def test_seep_boundaries_not_on_record():
    """A solution restored without the boundary conditions that produced it says
    so, rather than describing a boundary problem that does not exist.

    Both counts zero read as confined, and the paragraph stated that every node of
    the mesh flows saturated and 0 of them carry a specified head — a sentence
    about a solve with no boundaries at all.
    """
    fails = []
    from xslope.report import _bc_counts, _seep_unconfined

    report = _seep_report(NOBC_SEEP_XLSX)
    _slope_data, bundles = _ENGINE[("seep_all", NOBC_SEEP_XLSX)]
    n_head, n_exit = _bc_counts(bundles[0]["seep_data"])
    if (n_head, n_exit) != (0, 0):
        fails.append(f"the sample carries {n_head} head and {n_exit} exit-face "
                     f"nodes, so the degenerate case this check is about is "
                     f"never taken")
    if _seep_unconfined(bundles[0]["seep_data"], bundles[0]["solution"]) is not None:
        fails.append("a solve with neither boundary type on record is read as "
                     "one or the other rather than as unknown")

    text = " ".join(_seep_results_prose(report))
    if "does not record the boundary conditions" not in text:
        fails.append(f"a solution with no boundary conditions on record does not "
                     f"say so: {text!r}")
    for wrong in ("confined problem", "flows saturated", "exit face"):
        if wrong in text:
            fails.append(f"a solve with no boundaries on record is described as "
                         f"{wrong!r}: {text!r}")
    if " 0 " in text or "0 nodes" in text or "0 of them" in text:
        fails.append(f"a solve with no boundaries on record counts them anyway: "
                     f"{text!r}")

    # The inputs describe no boundaries either. The model figure was credited with
    # "the water surface each specified-head boundary states" and the mesh figure
    # with "every specified-head and exit-face node marked", over a mesh that
    # carries neither.
    lead = _seep_inputs_prose(report)
    for wrong in ("specified-head", "exit-face", "node marked"):
        if wrong in lead:
            fails.append(f"the inputs of a solve with no boundaries on record "
                         f"describe {wrong!r}: {lead!r}")
    if "including the flow domain and its material zones" not in lead:
        fails.append(f"the model figure does not say what it does show: {lead!r}")
    if "The seepage mesh constructed for the problem is shown in" not in lead:
        fails.append(f"the inputs do not say what the mesh figure is: {lead!r}")

    # And the head figure's sentence agrees with its own subject. This solve
    # carries neither a phreatic surface nor a boundary water level, so the
    # subject is the contours ALONE — one item in the list and a plural noun —
    # and a verb chosen by counting the list wrote "The head contours is shown".
    if "The head contours are shown in" not in text:
        fails.append(f"the head figure's sentence does not agree with its "
                     f"subject: {text!r}")

    # And they say it ONCE. The sentence stood in the inputs and in the results,
    # on facing pages, for the same model. It belongs to the results — the fact is
    # about the saved solution, and the results are where the boundary counts are
    # stated — so the inputs simply do not raise the subject.
    if "does not record the boundary conditions" in lead:
        fails.append(f"the boundary conditions not being on record is stated in "
                     f"the inputs as well as the results, twice for one model: "
                     f"{lead!r}")
    said = sum(b.text.count("does not record the boundary conditions")
               for b in report.blocks("prose"))
    if said != 1:
        fails.append(f"the boundary conditions not being on record is stated "
                     f"{said} times in one report; it is one fact, said once")

    # The mesh figure is captioned as what it is. "Seepage mesh and boundary
    # conditions" sat over prose saying the boundary conditions are not on record.
    captions = [f.caption for f in report.figures()
                if f.caption.startswith("Seepage mesh")]
    if captions != ["Seepage mesh"]:
        fails.append(f"the mesh of a model carrying no boundary condition is "
                     f"captioned {captions!r}, not ['Seepage mesh']")
    return fails


def test_seep_stale_sidecar_says_so():
    """A solution whose record contradicts the boundaries now on the model tells
    ONE story, and says which of the two it is telling.

    The solution file records the branch the solver took; the boundary arrays are
    rebuilt from the spreadsheet every time the model is opened. Edit the
    boundaries after the solution was saved and the two describe different
    problems, and the paragraph stated both as fact: "Flow was solved as a confined
    problem: every node of the mesh flows saturated. 50 nodes carry a specified
    head and 57 lie on an exit face" — a confined problem with an exit face on it.
    """
    fails = []
    from xslope.report import _bc_counts, _seep_bc_stale

    xlsx = _silent_seep_xlsx()
    _slope_data, bundle = _seep_bundle(xlsx)
    n_head, n_exit = _bc_counts(bundle["seep_data"])
    if not (n_head and n_exit):
        fails.append(f"the sample carries {n_head} head and {n_exit} exit-face "
                     f"nodes, so a flipped flag would not contradict them")
        return fails
    if bundle["solution"].get("unconfined") is not None:
        fails.append("the sample's saved solution already records what it was, so "
                     "the flip below is not the disagreement this check is about")

    def reported(**facts):
        edited = dict(bundle, solution=dict(bundle["solution"], **facts))
        return " ".join(_seep_results_prose(
            _engine_report("seep", bundle=edited, xlsx=xlsx)))

    # The record and the mesh agree: the counts are the solve's own boundaries and
    # are stated as such, with no stale-sidecar language anywhere.
    agreed = reported(unconfined=True)
    if "not the ones the saved solution was solved under" in agreed:
        fails.append(f"a solution whose record matches its boundaries is called "
                     f"stale: {agreed!r}")
    if f"{n_head:,} nodes carry a specified head and {n_exit:,} lie on an exit " \
            f"face" not in agreed:
        fails.append(f"a solve whose boundaries are its own does not state them: "
                     f"{agreed!r}")

    # Flipped: the record says confined over a mesh carrying exit faces. What was
    # solved is what the record says; the counts are what the model carries now.
    if not _seep_bc_stale(bundle["seep_data"],
                          dict(bundle["solution"], unconfined=False)):
        fails.append("a confined record over a mesh with exit faces is not read "
                     "as a disagreement at all")
    stale = reported(unconfined=False)
    if "solved as a confined problem" not in stale:
        fails.append(f"a solution recorded as confined is not reported as the "
                     f"confined solve it records: {stale!r}")
    if "unconfined problem" in stale:
        fails.append(f"a solution recorded as confined is reported as unconfined "
                     f"as well: {stale!r}")
    told = (f"The boundary conditions now on the model are not the ones the saved "
            f"solution was solved under: {n_head:,} nodes carry a specified head "
            f"and {n_exit:,} lie on an exit face")
    if told not in stale:
        fails.append(f"the boundaries that contradict the record are stated as "
                     f"the solve's own: {stale!r}")
    # The two halves as fact is exactly the defect: the count sentence standing on
    # its own, with nothing between it and "every node of the mesh flows saturated".
    if "flows saturated. " + f"{n_head:,} nodes" in stale:
        fails.append(f"the contradicting counts are stated straight after the "
                     f"confined sentence, as though both were the same solve: "
                     f"{stale!r}")

    # And the other way: a confined model whose record says unconfined.
    _cd, cbundle = _seep_bundle(CONFINED_SEEP_XLSX)
    cn_head, cn_exit = _bc_counts(cbundle["seep_data"])
    edited = dict(cbundle, solution=dict(cbundle["solution"], unconfined=True))
    loose = " ".join(_seep_results_prose(
        _engine_report("seep", bundle=edited, xlsx=CONFINED_SEEP_XLSX)))
    if "solved as an unconfined problem" not in loose:
        fails.append(f"a solution recorded as unconfined over a mesh with no exit "
                     f"face is not reported as what it records: {loose!r}")
    if "not the ones the saved solution was solved under" not in loose:
        fails.append(f"a mesh with no exit face under an unconfined record is not "
                     f"said to differ from it: {loose!r}")
    if "exit face" in loose:
        fails.append(f"a mesh with no exit face on it is given one: {loose!r}")
    if f"{cn_head:,} nodes carry a specified head" not in loose:
        fails.append(f"the boundaries the model does carry are not stated: "
                     f"{loose!r}")
    if cn_exit:
        fails.append(f"the confined sample carries {cn_exit} exit-face nodes, so "
                     f"the case this half is about is never taken")
    return fails


def test_seep_without_a_flowrate():
    """A solution that records no flow rate says so, and its figure is not called
    a flow net.

    The flow is the results subsection's only number and the paragraph simply
    dropped the sentence when the file carried none. Flow lines are contours of the
    stream function SPACED BY the flow they carry, so a solution with no flow rate
    has none of them: the figure is head contours, and calling it a flow net names
    a thing it does not contain.
    """
    fails = []
    from xslope.plot_seep import flownet_has_flowlines

    report = _seep_report(NOFLOW_SEEP_XLSX)
    _slope_data, bundles = _ENGINE[("seep_all", NOFLOW_SEEP_XLSX)]
    if bundles[0]["solution"].get("flowrate") is not None:
        fails.append("the sample's saved solution does record a flow rate, so "
                     "the branch this check is about is never taken")

    text = " ".join(_seep_results_prose(report))
    if "records no flow rate" not in text:
        fails.append(f"a solution with no flow rate passes over it in silence: "
                     f"{text!r}")
    if "The flow through the section is" in text:
        fails.append(f"a solution with no flow rate states one anyway: {text!r}")
    if "flow line" in text:
        fails.append(f"a figure with no flow lines in it is described as drawing "
                     f"them: {text!r}")

    calls = [c for c in _plot_seep_calls(NOFLOW_SEEP_XLSX)
             if c[0].get("variable") == "head"]
    if len(calls) != 1:
        fails.append(f"the model drew {len(calls)} head panels, expected one")
    for kw, sd, sol in calls:
        if flownet_has_flowlines(sd, sol):
            fails.append("a solution with no flow rate is said to carry flow lines")
        gids = _figure_gids(sd, sol, kw)
        if "FLOWLINES" in gids:
            fails.append(f"the figure of a solution with no flow rate draws flow "
                         f"lines: {sorted(gids)}")
        if "HEAD_CONTOURS" not in gids:
            fails.append(f"the figure draws no head contours either: {sorted(gids)}")
    fig = next((f for f in report.figures()
                if f.source == "seepage bc1 head"), None)
    if fig is None or fig.caption != "Head contours":
        fails.append(f"a figure with no flow lines is captioned "
                     f"{getattr(fig, 'caption', None)!r}, not for what it draws")

    # This sample's stream function is flat as well as its flow rate absent, so its
    # figure carries no flow lines on either count and the absence above would be
    # satisfied by a plotter that only ever looked at the stream function. The flow
    # rate on its own is pinned on a solution whose stream function has real range:
    # take the rate away and the lines go, because there is nothing left to space
    # them by.
    _sd, dam = _seep_bundle()
    if not flownet_has_flowlines(dam["seep_data"], dam["solution"]):
        fails.append("the dam's solution draws no flow lines to begin with, so "
                     "taking its flow rate away proves nothing")
    stripped = dict(dam["solution"], flowrate=None)
    if flownet_has_flowlines(dam["seep_data"], stripped):
        fails.append("a solution whose flow rate is gone still claims flow lines")
    gids = _figure_gids(dam["seep_data"], stripped,
                        {"mesh": False, "show_title": False})
    if "FLOWLINES" in gids:
        fails.append(f"a figure of a solution with no flow rate draws flow lines "
                     f"spaced by nothing: {sorted(gids)}")
    if "HEAD_CONTOURS" not in gids:
        fails.append(f"it draws no head contours either: {sorted(gids)}")
    return fails


def _pressure_only_bundle():
    """A seepage bundle whose saved solution is a bare nodal pore pressure field —
    what arrives from another program, and what no model in the corpus ships.

    Written and read back through the real pair (:func:`~xslope.seep.export_seep_u`
    and :func:`~xslope.seep.import_seep_solution`), so the solution carries exactly
    the NaNs an imported field carries rather than a set of NaNs invented here.
    """
    from xslope.seep import export_seep_u, import_seep_solution

    slope_data, bundle = _seep_bundle()
    seep_data = bundle["seep_data"]
    tmp = tempfile.mkdtemp(prefix="xslope_useep_")
    path = os.path.join(tmp, "field_seep.csv")
    with contextlib.redirect_stdout(io.StringIO()):
        export_seep_u(seep_data["nodes"], bundle["solution"]["u"], path,
                      float(seep_data["unit_weight"]))
        solution = import_seep_solution(seep_data, path)
    return slope_data, dict(bundle, solution=solution)


def test_seep_panels_the_solution_cannot_draw():
    """A field the saved solution has nothing in is not drawn, and the subsection
    says which fields those are.

    Two ways a panel has no figure in it. A pore pressure field imported from
    another program carries head and pore pressure and nothing else, and the flow
    quantities read back as NaN; several saved solutions in the corpus record a
    velocity and a gradient of zero at every node, which has no contour in it. Drawn
    anyway, the first is a blank page and the second was an exception the figure
    never came back from — so the panels are dropped and the reader is told, rather
    than left to wonder which figures went missing.
    """
    fails = []
    import numpy as np
    from xslope.report import SEEP_PANELS, seep_panels, resolve_options

    flow_fields = [p["variable"] for p in SEEP_PANELS
                   if p["variable"] not in ("head", "u")]
    kept = ["head", "u"]

    # --- one value everywhere: the corpus's own zero-velocity solutions ---
    report = _seep_report(NOFLOW_SEEP_XLSX)
    _slope_data, bundles = _ENGINE[("seep_all", NOFLOW_SEEP_XLSX)]
    for variable in flow_fields:
        field = np.asarray(bundles[0]["solution"][variable], dtype=float)
        if float(np.ptp(field[np.isfinite(field)])) > 0.0:
            fails.append(f"the sample's {variable!r} field has range in it, so the "
                         f"flat-field branch this check is about is never taken")
    sources = [f.source for f in report.figures()]
    for variable in flow_fields:
        if f"seepage bc1 {variable}" in sources:
            fails.append(f"a {variable!r} field with one value everywhere was "
                         f"drawn anyway: {sources}")
    for variable in kept:
        if f"seepage bc1 {variable}" not in sources:
            fails.append(f"the {variable!r} field has range in it and was not "
                         f"drawn: {sources}")
    said = " ".join(_seep_results_prose(report))
    if "one value everywhere" not in said:
        fails.append(f"a flat field is passed over in silence: {said!r}")
    for variable in flow_fields:
        name = next(p["field"] for p in SEEP_PANELS if p["variable"] == variable)
        if name not in said:
            fails.append(f"the {name} field is not drawn and is not named as one "
                         f"of the flat ones: {said!r}")
    planned, drawn = _planned_matches(report, "seep", bundle=bundles,
                                      xlsx=NOFLOW_SEEP_XLSX)
    if planned != drawn:
        fails.append(f"the flat-field report planned {planned} figures and built "
                     f"{drawn}")
    for kw, _sd, _sol in _plot_seep_calls(NOFLOW_SEEP_XLSX):
        if kw.get("variable") in flow_fields:
            fails.append(f"a {kw.get('variable')!r} figure was drawn from a field "
                         f"with one value everywhere")

    # And the flatness is what decides it: the same sets given a field with range
    # in them print the panels. Otherwise a builder that never printed them at all
    # would pass everything above.
    with_range = []
    for bundle in bundles:
        solution = dict(bundle["solution"])
        head = np.asarray(solution["head"], dtype=float)
        for variable in flow_fields:
            solution[variable] = np.abs(head - float(np.mean(head)))
        solution["velocity"] = np.column_stack([solution[flow_fields[0]],
                                                solution[flow_fields[0]]])
        with_range.append(dict(bundle, solution=solution))
    back = _seep_report(NOFLOW_SEEP_XLSX, bundles=with_range)
    got = [f.source for f in back.figures()]
    for variable in flow_fields:
        if f"seepage bc1 {variable}" not in got:
            fails.append(f"a {variable!r} field WITH range in it is still not "
                         f"drawn, so dropping the flat one proves nothing: {got}")
    if "one value everywhere" in " ".join(_seep_results_prose(back)):
        fails.append("a solution whose fields all have range in them is still said "
                     "to carry a flat one")

    # --- carried at all: a bare imported pore pressure field ---
    slope_data, bundle = _pressure_only_bundle()
    for variable in flow_fields:
        field = np.asarray(bundle["solution"][variable], dtype=float)
        if np.isfinite(field).any():
            fails.append(f"the imported field carries a {variable!r} after all, so "
                         f"the missing-field branch is never taken")
    imported = _seep_report(SEEP_XLSX, bundles=[bundle])
    sources = [f.source for f in imported.figures()]
    for variable in flow_fields:
        if f"seepage bc1 {variable}" in sources:
            fails.append(f"a {variable!r} field the solution does not carry was "
                         f"drawn: {sources}")
    for variable in kept:
        if f"seepage bc1 {variable}" not in sources:
            fails.append(f"the imported field carries {variable!r} and it was not "
                         f"drawn: {sources}")
    said = " ".join(_seep_results_prose(imported))
    if "carries no" not in said:
        fails.append(f"an imported field passes over what it does not carry: "
                     f"{said!r}")
    for variable in flow_fields:
        name = next(p["field"] for p in SEEP_PANELS if p["variable"] == variable)
        if name not in said:
            fails.append(f"the {name} the solution does not carry is not named: "
                         f"{said!r}")
    if "one value everywhere" in said:
        fails.append(f"a field that is not there is described as a flat one: "
                     f"{said!r}")
    planned, drawn = _planned_matches(imported, "seep", bundle=[bundle],
                                      xlsx=SEEP_XLSX, slope_data=slope_data)
    if planned != drawn:
        fails.append(f"the imported-field report planned {planned} figures and "
                     f"built {drawn}")
    for kw, _sd, _sol in _plot_seep_calls(SEEP_XLSX, bundles=[bundle]):
        if kw.get("variable") in flow_fields:
            fails.append(f"a {kw.get('variable')!r} figure was drawn from a field "
                         f"of NaN")

    # The count and the section are decided by one call, so they cannot come apart:
    # the panels the section draws are the panels planned_figures counts.
    opts = resolve_options({"input_path": SEEP_XLSX, "lem": False})
    for solution, want in ((bundle["solution"], kept),
                           (_seep_bundle()[1]["solution"],
                            [p["variable"] for p in SEEP_PANELS])):
        got = [p["variable"] for p in seep_panels(solution, opts)]
        if got != want:
            fails.append(f"seep_panels gives {got}, expected {want}")

    # The report leaves a flat field out because a blank figure is not worth a page,
    # not because the plotter cannot draw one. The seepage results view offers the
    # same four variables on the same solution, and asking it for a velocity that is
    # zero everywhere raised "Contour levels must be increasing" instead of drawing
    # the flat field it is.
    flat_gids = _figure_gids(bundles[0]["seep_data"], bundles[0]["solution"],
                             {"variable": flow_fields[0], "mesh": False,
                              "show_title": False})
    if "CONTOUR_FILL" not in flat_gids:
        fails.append(f"a field with one value everywhere is not drawn even when it "
                     f"is asked for directly: {sorted(flat_gids)}")

    # A model whose fields are all there is unaffected: no sentence about anything
    # missing on a report that draws every panel.
    whole = " ".join(_seep_results_prose(_engine_report("seep")))
    for wording in ("carries no", "one value everywhere"):
        if wording in whole:
            fails.append(f"a complete solution is said to be missing something: "
                         f"{whole!r}")
    return fails


def test_seep_dual_section():
    """A model solved for two boundary condition sets documents both, and draws
    them on one scale.

    Each set is its own flow problem on the same mesh, so each gets its own
    subsection with its own boundary counts and its own flow. Auto-scaled
    independently the two figures carried the same colors for different heads and
    spaced their flow lines by different amounts of flow, so the pair read as two
    unrelated problems rather than as one section before and after drawdown.
    """
    fails = []
    import numpy as np
    from xslope.plot_seep import flownet_base_material

    report = _seep_report(RAPID_SEEP_XLSX)
    _slope_data, bundles = _ENGINE[("seep_all", RAPID_SEEP_XLSX)]
    if len(bundles) != 2:
        fails.append(f"the rapid drawdown sample restored {len(bundles)} "
                     f"boundary condition sets, expected two")
        return fails

    expected = [(1, "Traceability"), (1, "Project Definition"),
                (1, "Seepage Analysis"), (2, "Analysis Inputs"),
                (2, "Seepage Mesh"),
                (2, "Boundary Condition Set 1"), (2, "Boundary Condition Set 2")]
    got = report.section_titles()
    if got != expected:
        fails.append(f"the dual-solution report's sections are {got}, expected "
                     f"{expected}")

    # The model, the two conductivity conventions, a mesh per set, and every result
    # panel per set. A set whose mesh figure went missing would leave its paragraph
    # pointing at the other set's boundaries.
    from xslope.report import SEEP_PANELS
    sources = [f.source for f in report.figures()]
    want = ["seep model", "seep kr", "seep kr_head", "seepage bc1 mesh",
            "seepage bc2 mesh"] + [f"seepage bc{n} {p['variable']}"
                                   for n in (1, 2) for p in SEEP_PANELS]
    if sorted(sources) != sorted(want):
        fails.append(f"the dual-solution report drew {sources}, expected {want}")
    planned, drawn = _planned_matches(report, "seep", bundle=bundles,
                                      xlsx=RAPID_SEEP_XLSX)
    if planned != drawn or drawn != len(want):
        fails.append(f"the dual-solution report planned {planned} figures and "
                     f"built {drawn}, of {len(want)}")

    # Each set states its own boundaries and its own flow, and they differ: one
    # set's numbers printed twice would be a section that documented one solve.
    proses = _seep_results_prose(report)
    if len(proses) != 2:
        fails.append(f"the dual-solution section has {len(proses)} results "
                     f"subsections, expected two")
    else:
        from xslope.report import _bc_counts
        for text, bundle in zip(proses, bundles):
            n_head, n_exit = _bc_counts(bundle["seep_data"])
            q = bundle["solution"]["flowrate"]
            for want_str in (f"{n_head:,}", f"{n_exit:,}", f"{q:.4g}"):
                if want_str not in text:
                    fails.append(f"boundary condition set "
                                 f"{bundle['options']['bc']} does not state "
                                 f"{want_str}: {text!r}")
        heads = [_bc_counts(b["seep_data"])[0] for b in bundles]
        flows = [b["solution"]["flowrate"] for b in bundles]
        if heads[0] == heads[1] or flows[0] == flows[1]:
            fails.append(f"the two sets carry the same head count {heads} or the "
                         f"same flow {flows}, so a section printing one twice "
                         f"would pass this check")
        if proses[0] == proses[1]:
            fails.append("both boundary condition sets are described in the same "
                         "words")

    # One scale PER VARIABLE: each field's pair of panels is drawn on one contour
    # range, spanning both sets so neither is clipped, and every panel on one base
    # material. Held on the head alone, a pore pressure colour or a gradient colour
    # would mean a different amount in the two halves of the same drawdown.
    calls = _plot_seep_calls(RAPID_SEEP_XLSX)
    if len(calls) != 2 * len(SEEP_PANELS):
        fails.append(f"the dual-solution report drew {len(calls)} result panels, "
                     f"expected {2 * len(SEEP_PANELS)}")
    mats = {kw.get("base_mat") for kw, _sd, _sol in calls}
    if len(mats) != 1:
        fails.append(f"the panels are scaled to different base materials {mats}, "
                     f"so a flow channel means a different flow in each")
    for panel in SEEP_PANELS:
        variable = panel["variable"]
        pair = [kw for kw, _sd, _sol in calls if kw.get("variable") == variable]
        if len(pair) != 2:
            fails.append(f"the pair drew {len(pair)} {variable!r} panels")
            continue
        ranges = {(kw.get("vmin"), kw.get("vmax")) for kw in pair}
        if len(ranges) != 1 or None in list(ranges)[0]:
            fails.append(f"the two {variable!r} panels are drawn on {ranges}, not "
                         f"on one shared contour range")
            continue
        vmin, vmax = sorted(ranges)[0]
        own = [(float(np.min(b["solution"][variable])),
                float(np.max(b["solution"][variable]))) for b in bundles]
        lo = min(o[0] for o in own)
        hi = max(o[1] for o in own)
        if vmin > lo or vmax < hi:
            fails.append(f"the shared {variable!r} range {vmin}–{vmax} clips the "
                         f"fields, which span {lo}–{hi}")
        # And the pair really would scale apart: the two sets span different
        # amounts of this field, so a per-panel auto-scale would not land here.
        if own[0] == own[1]:
            fails.append(f"both sets span the same {variable!r} values {own}, so "
                         f"an independently scaled pair would pass this check")

    # The velocity ARROWS are on one scale too. An arrow is a length standing for a
    # speed, and each panel scaled to its own maximum drew its longest arrow the
    # same length in both sets — the same mark meaning 109 ft/day before drawdown
    # and 269 after. Measured off the arrows as drawn: their lengths are in the
    # ratio of the speeds they stand for.
    arrows = [(kw, sd, sol) for kw, sd, sol in calls if kw.get("vectors")]
    if len(arrows) != 2:
        fails.append(f"the pair drew {len(arrows)} panels with velocity arrows, "
                     f"expected the two velocity magnitude panels")
    else:
        pinned = {kw.get("vector_max") for kw, _sd, _sol in arrows}
        peaks = [float(np.max(b["solution"]["v_mag"])) for b in bundles]
        if len(pinned) != 1 or None in pinned:
            fails.append(f"the two velocity panels are drawn with vector_max "
                         f"{pinned}, not on one arrow scale")
        elif abs(list(pinned)[0] - max(peaks)) > 1e-9 * max(peaks):
            fails.append(f"the arrows are pinned to {list(pinned)[0]:g}, not the "
                         f"pair's own maximum speed of {max(peaks):g}")
        lengths = [_longest_arrow(sd, sol, kw) for kw, sd, sol in arrows]
        if None in lengths:
            fails.append(f"a velocity panel drew no arrows to measure: {lengths}")
        elif abs(peaks[0] - peaks[1]) < 1e-9:
            fails.append(f"both sets reach the same peak speed {peaks}, so an "
                         f"independently scaled pair would pass this check")
        else:
            drawn_ratio = lengths[0] / lengths[1]
            speed_ratio = peaks[0] / peaks[1]
            if abs(drawn_ratio - speed_ratio) > 1e-6 * speed_ratio:
                fails.append(f"the longest arrows are in the ratio "
                             f"{drawn_ratio:.4f} and the speeds they stand for in "
                             f"the ratio {speed_ratio:.4f}: an arrow of a given "
                             f"length means a different speed in each set")

    # The base material, held against a pair that would NOT choose it on its own.
    # Both sets of this model happen to call for the same zone, so any assertion
    # made on the sets as they stand is satisfied by a per-set recompute — delete
    # the shared choice and the check still passes. The base material follows the
    # flow rate, so set 2's is multiplied until the zone it calls for alone is a
    # different one; the pair must still be drawn on the zone set 1 calls for.
    want_mat = flownet_base_material(bundles[0]["seep_data"],
                                     bundles[0]["solution"])
    apart = [bundles[0],
             dict(bundles[1], solution=dict(bundles[1]["solution"],
                  flowrate=float(bundles[1]["solution"]["flowrate"]) * 5.0))]
    alone = flownet_base_material(apart[1]["seep_data"], apart[1]["solution"])
    if alone == want_mat:
        fails.append(f"the second set still calls for material {alone} on its "
                     f"own, so a per-set choice would pass this check")
    got = [kw.get("base_mat") for kw, _sd, _sol in
           _plot_seep_calls(RAPID_SEEP_XLSX, bundles=apart)]
    if set(got) != {want_mat}:
        fails.append(f"the panels were drawn on base materials {sorted(set(got))}, "
                     f"not all on the {want_mat} the first set's conductivities "
                     f"call for — the second set left to itself calls for {alone}")
    return fails


#: The transient samples: a homogeneous dam drawn down over 45 days, and the
#: zoned Johnson Reservoir. Each ships the march it was figured from —
#: ``{stem}_tseep.csv`` and its ledger — beside the model.
TSEEP_XLSX = os.path.join(_REPO, "docs", "seep", "files",
                          "xslope_earth_dam_tseep.xlsx")
TSEEP_ZONED_XLSX = os.path.join(_REPO, "docs", "seep", "files",
                                "xslope_johnson_res_tseep.xlsx")


def _tseep_solutions(xlsx=TSEEP_XLSX):
    """``(slope_data, solutions)`` for a transient model, read back from the march
    saved beside it."""
    key = ("tseep", xlsx)
    if key not in _ENGINE:
        _ENGINE[key] = _restored(xlsx)
    return _ENGINE[key]


def _tseep_report(xlsx=TSEEP_XLSX, options=None, solutions=None):
    """A report of a transient march."""
    from xslope.report import build_report

    slope_data, shipped = _tseep_solutions(xlsx)
    opts = {"input_path": xlsx, "lem": False, "pd_figure": False}
    opts.update(FAST_FIGURES)
    opts.update(options or {})
    tmp = tempfile.mkdtemp(prefix="xslope_tseep_")
    with contextlib.redirect_stdout(io.StringIO()):
        return build_report(slope_data, solutions or shipped, opts, tmp)


def _tseep_section(report):
    """The subsection of a seepage section that documents the march, or None."""
    sec = next((s for s in report.sections if s.title == "Seepage Analysis"), None)
    for _lvl, node in (sec.walk() if sec else []):
        if node is sec or node.title == "Analysis Inputs":
            continue
        text = " ".join(b.text for b in node.blocks if b.kind == "prose")
        if "transient" in text:
            return node
    return None


def _plot_tseep_calls(xlsx=TSEEP_XLSX, options=None, solutions=None):
    """Every call a transient report's frame figures make to
    ``plot_seep_solution``, as ``(kwargs, seep_data, solution)``."""
    import xslope.plot_seep as ps

    real = ps.plot_seep_solution
    seen = []

    def spy(seep_data, solution, **kw):
        seen.append((dict(kw), seep_data, solution))
        return real(seep_data, solution, **kw)

    ps.plot_seep_solution = spy
    try:
        _tseep_report(xlsx, dict(options or {}, seep_inputs_figure=False,
                                 seep_mesh_figure=False, seep_kr_figure=False),
                      solutions=solutions)
    finally:
        ps.plot_seep_solution = real
    return seen


def test_tseep_discovered_beside_the_model():
    """A transient march saved beside a model is read back, and one that cannot be
    used is noted rather than raised.

    The two transient samples ship ``{stem}_tseep.csv`` and its ledger. Read back
    they give the bundle a fresh march emits — every frame a full solution the
    plotter takes — so documenting a model that has been marched costs no marching.
    A companion read against a mesh it was not solved on is the failure that
    actually happens, and it costs its own section, not the report.
    """
    fails = []
    import numpy as np
    from xslope.report import (seep_bundles, seepage_documented,
                               solutions_from_sidecars, tseep_bundles)

    for xlsx in (TSEEP_XLSX, TSEEP_ZONED_XLSX):
        _slope_data, solutions = _tseep_solutions(xlsx)
        bundles = tseep_bundles(solutions)
        if len(bundles) != 1:
            fails.append(f"{os.path.basename(xlsx)} restored "
                         f"{len(bundles)} transient bundles, expected one")
            continue
        bundle = bundles[0]
        march = bundle["transient"]
        if len(march["frames"]) < 4:
            fails.append(f"{os.path.basename(xlsx)} restored "
                         f"{len(march['frames'])} frames; a march documented at "
                         f"four states needs at least four")
        if len(march["times"]) != len(march["frames"]):
            fails.append(f"{os.path.basename(xlsx)} restored "
                         f"{len(march['times'])} times for "
                         f"{len(march['frames'])} frames")
        if bundle.get("frames") is not march["frames"]:
            fails.append("the bundle's frames are not the march's own")
        # Every frame is a plottable solution: head and pore pressure read from
        # the file, velocity and gradient derived on load, and NO stream function.
        for i, frame in enumerate(march["frames"]):
            for key in ("head", "u", "velocity", "v_mag", "gradient", "i_mag"):
                field = frame.get(key)
                if field is None or not np.asarray(field).size:
                    fails.append(f"{os.path.basename(xlsx)} frame {i} carries no "
                                 f"{key}")
            if frame.get("phi") is not None:
                fails.append(f"{os.path.basename(xlsx)} frame {i} carries a "
                             f"stream function; a storage-release state has none")
        if not seepage_documented(solutions):
            fails.append(f"{os.path.basename(xlsx)} has a march and no seepage "
                         f"section would be written for it")
        # The transient samples carry no steady set, which is what makes them the
        # models that prove the section is not gated on one.
        if seep_bundles(solutions):
            fails.append(f"{os.path.basename(xlsx)} also ships a steady set, so "
                         f"it does not prove a march alone earns the section")

    # A march saved against a different mesh: noted, named, and not raised. The
    # mesh is coarsened so the node counts disagree, which is exactly what a
    # rebuilt mesh does to a companion saved before it.
    from xslope.fileio import load_slope_data
    from xslope.mesh import build_mesh_from_polygons, get_material_polygons
    with contextlib.redirect_stdout(io.StringIO()):
        slope_data = load_slope_data(TSEEP_XLSX)
        coarse = build_mesh_from_polygons(get_material_polygons(slope_data),
                                          12.0, "tri3")
    stale = dict(slope_data, mesh=coarse)
    if len(coarse["nodes"]) == len(slope_data["mesh"]["nodes"]):
        fails.append("the coarsened mesh has the same node count, so a stale "
                     "march would still load and this check proves nothing")
    notes = []
    with contextlib.redirect_stdout(io.StringIO()):
        solutions = solutions_from_sidecars(TSEEP_XLSX, stale, notes)
    if tseep_bundles(solutions):
        fails.append("a march saved against a different mesh was read as if it "
                     "were this model's")
    if not any("tseep.csv" in note for note in notes):
        fails.append(f"a march that could not be read is not in the notes: {notes}")

    # A model that was never marched is not a fault, and says nothing.
    notes = []
    with contextlib.redirect_stdout(io.StringIO()):
        solutions = solutions_from_sidecars(SEEP_XLSX, None, notes)
    if tseep_bundles(solutions):
        fails.append(f"{os.path.basename(SEEP_XLSX)} was never marched and a "
                     f"transient bundle was invented for it")
    if any("tseep" in note for note in notes):
        fails.append(f"a model that was never marched is noted for it: {notes}")
    return fails


def test_tseep_section():
    """A report of a transient march says it is one, states the march's own
    numbers, and draws the section at states through it.

    A head figure does not say whether the field under it was solved at one
    instant or reached through time, and a pore pressure read off the wrong basis
    is a different number. So the basis is stated, and everything stated with it —
    how long the march ran, how many states it kept, how well it conserved water —
    is the march's own record rather than a number this section computed.
    """
    fails = []
    from xslope.report import SEEP_PANELS, transient_ledger, tseep_bundles

    _slope_data, solutions = _tseep_solutions()
    bundle = tseep_bundles(solutions)[0]
    ledger = transient_ledger(bundle)
    report = _tseep_report()

    expected = [(1, "Traceability"), (1, "Project Definition"),
                (1, "Seepage Analysis"), (2, "Analysis Inputs"),
                (2, "Seepage Mesh"), (2, "Results")]
    got = report.section_titles()
    if got != expected:
        fails.append(f"the transient report's sections are {got}, expected "
                     f"{expected}")

    # The rest of the report knows a march is a seepage analysis: the model checks
    # are filtered to one, and the section is one a citation elsewhere can land on
    # and one the Project Definition can send a reader to.
    from xslope.report import _engine_sections, report_analyses, resolve_options
    resolved = resolve_options({"lem": False, "pd_figure": False})
    if "seep" not in report_analyses(solutions, resolved):
        fails.append(f"a march is not a seepage analysis to the model checks: "
                     f"{report_analyses(solutions, resolved)}")
    listed = [key for key, _anchor, _name in _engine_sections(solutions, resolved)]
    if "seep" not in listed:
        fails.append(f"the section a march is documented in is not among the "
                     f"engine sections a citation can reach: {listed}")

    sub = _tseep_section(report)
    if sub is None:
        fails.append("the transient report has no subsection describing a march")
        return fails
    text = " ".join(b.text for b in sub.blocks if b.kind == "prose")

    if "transient" not in text:
        fails.append(f"the section never says the analysis was transient: {text!r}")
    # The march's own numbers, each read off the ledger rather than written here.
    for label, value in (("duration", f"{ledger['duration']:g}"),
                         ("saved states", f"{len(ledger['times']):,}")):
        if value not in text:
            fails.append(f"the {label} {value} is stated nowhere: {text!r}")
    if "day" not in text:
        fails.append(f"the time unit the march ran in is never named: {text!r}")

    # In the past tense, and in days. The march is a run that happened, and the
    # sentence said "The march runs from t = 0 to t = 360 day and saved 12
    # states" — half of it present, and a duration carrying the unit label as an
    # axis carries it rather than as an amount of time.
    want = (f"The march ran from t = 0 to t = {ledger['duration']:g} days and "
            f"saved {len(ledger['times']):,} states.")
    if want not in text:
        fails.append(f"the march is not stated as {want!r}: {text!r}")
    if "The march runs" in text:
        fails.append(f"the march is described in the present tense: {text!r}")
    if f"{ledger['duration']:g} day " in text:
        fails.append(f"a duration of {ledger['duration']:g} is given in the "
                     f"singular: {text!r}")

    # The march closed, and how well: the closure is a transient solve's own
    # statement of how much water it failed to account for.
    if ledger["converged"] is not True:
        fails.append("the sample march did not converge, so the converged wording "
                     "is never exercised")
    if ledger["closure"] is None:
        fails.append("the sample march records no closure, so the closure wording "
                     "is never exercised")
    elif f"{100.0 * ledger['closure']:.3g}" not in text:
        fails.append(f"the mass-balance closure "
                     f"{100.0 * ledger['closure']:.3g} percent is not stated: "
                     f"{text!r}")

    # The boundary the march is driven by, and the law its nodes follow.
    from xslope.plot_seep import reservoir_face_mask
    n_face = int(reservoir_face_mask(bundle["seep_data"]).sum())
    if not n_face:
        fails.append("the sample carries no series-driven reservoir boundary, so "
                     "the sentence that names one is never exercised")
    if "drains freely once the level falls below it" not in text:
        fails.append(f"the march never says what its reservoir face does as the "
                     f"level passes a node: {text!r}")

    # Four states, each drawn for every field the frame carries, plus the history.
    # The pool falls from t = 2 to t = 47, so the states are the full pool, one
    # part way down, the drawn-down state, and the section long after.
    sources = [f.source for f in report.figures()]
    times = [0.0, 15.0, 47.0, 360.0]
    want = (["seep model", "seep kr", "seep kr_head", "seepage bc1 mesh"]
            + [f"seepage tseep {t:g} {p['variable']}"
               for t in times for p in SEEP_PANELS]
            + ["seepage tseep history"])
    if sources != want:
        fails.append(f"the transient report drew {sources}, expected {want}")
    from xslope.report import planned_figures, resolve_options
    planned = planned_figures(_slope_data, solutions, resolve_options(
        dict(FAST_FIGURES, input_path=TSEEP_XLSX, lem=False, pd_figure=False)))
    if planned != len(report.figures()):
        fails.append(f"the transient report planned {planned} figures and built "
                     f"{len(report.figures())}")

    # Each state is captioned for the instant it is, in time order.
    captioned = [f.caption for f in report.figures()
                 if f.source.startswith("seepage tseep ")
                 and f.source != "seepage tseep history"]
    stamps = [c.split("—")[-1].strip() for c in captioned]
    if stamps != [f"t = {t:g} day" for t in times for _p in SEEP_PANELS]:
        fails.append(f"the frame captions are {stamps}")
    return fails


def test_tseep_mesh_figure_marks_what_it_says():
    """The mesh figure of a marched model promises only the boundary it marks, and
    says what the rest of the boundary is.

    A reservoir face carries no boundary type to mark: which type each of its nodes
    has is decided at every step by where the water line stands, so the figure draws
    all 28 of them — 35 on the zoned sample — as plain mesh. It was captioned for
    the boundary conditions and credited with "every specified-head and exit-face
    node marked", and a reader counting the marks against the counts the section
    states would come up short by exactly that face.
    """
    fails = []
    import numpy as np
    from xslope.plot_seep import reservoir_face_mask
    from xslope.report import _bc_counts, tseep_bundles

    for xlsx in (TSEEP_XLSX, TSEEP_ZONED_XLSX):
        name = os.path.basename(xlsx)
        _slope_data, solutions = _tseep_solutions(xlsx)
        data = tseep_bundles(solutions)[0]["seep_data"]
        face = reservoir_face_mask(data)
        n_face = int(face.sum()) if face is not None else 0
        if not n_face:
            fails.append(f"{name} carries no reservoir face, so the figure this "
                         f"checks has nothing it cannot mark")
            continue
        # The premise: not one node of the face carries a type the figure marks.
        typed = np.asarray(data["bc_type"])[face]
        if int(np.count_nonzero(typed)) != 0:
            fails.append(f"{name}: {int(np.count_nonzero(typed))} of the "
                         f"{n_face} reservoir-face nodes carry a fixed boundary "
                         f"type, so the face is partly marked after all")
        n_head, n_exit = _bc_counts(data)
        if not (n_head and n_exit):
            fails.append(f"{name} marks {n_head} specified-head and {n_exit} "
                         f"exit-face nodes; the wording this checks names both")

        report = _tseep_report(xlsx, options={"seep_transient_figures": False,
                                              "seep_transient_history": False})
        caption = next((f.caption for f in report.figures()
                        if f.source == "seepage bc1 mesh"), None)
        if caption != "Seepage mesh and the boundary conditions fixed on it":
            fails.append(f"{name}: the mesh figure is captioned {caption!r}, "
                         f"which claims a boundary it does not carry marks for")
        lead = _seep_inputs_prose(report)
        if "with every specified-head and exit-face node marked" in lead:
            fails.append(f"{name}: the mesh figure is credited with marking every "
                         f"specified-head node while the {n_face} of the "
                         f"reservoir face are drawn as plain mesh: {lead!r}")
        want = (f"The {n_face:,} nodes of the reservoir face are not marked: each "
                f"takes its boundary type at every step of the march")
        if want not in lead:
            fails.append(f"{name}: the inputs never say the reservoir face is "
                         f"unmarked, or why: {lead!r}")
        # And the face is counted once. It was counted with the mesh figure and
        # again in the march, the same number in two sentences.
        said = sum(b.text.count(f"{n_face:,} nodes")
                   for b in report.blocks("prose"))
        if said != 1:
            fails.append(f"{name}: the reservoir face is counted out {said} times "
                         f"in one report; it is one fact, said once")
    return fails


def test_tseep_frame_selection():
    """The states a march is documented at are the first, the last, and the rest
    weighted onto the drawdown — including one strictly inside the fall.

    Chosen among the SAVED states rather than at equal times, because the save
    schedule is where the modeller said the answer moves. That alone is not enough:
    spaced evenly through the twelve states Johnson Reservoir saves, the states
    documented were 0, 80, 300 and 1000 while the pool fell between t = 5 and
    t = 50, so every saved state of the drawdown was stepped over and the drawdown
    report never showed the drawdown. A state strictly inside the fall is the only
    one that shows the pool part way down — at the fall's own end it has already
    arrived — so both samples are held to having one.
    """
    fails = []
    from xslope.plot_seep import level_fall_interval
    from xslope.report import (resolve_options, transient_frame_times,
                               tseep_bundles)

    for xlsx in (TSEEP_XLSX, TSEEP_ZONED_XLSX):
        name = os.path.basename(xlsx)
        _slope_data, solutions = _tseep_solutions(xlsx)
        bundle = tseep_bundles(solutions)[0]
        times = [float(t) for t in bundle["transient"]["times"]]

        def picked(_bundle=bundle, **options):
            return transient_frame_times(_bundle, resolve_options(options))

        fall = level_fall_interval(bundle["seep_data"])
        if fall is None:
            fails.append(f"{name} is a drawdown sample whose series never falls, "
                         f"so the selection this checks is never exercised")
            continue
        start, end = fall
        inside = [t for t in times if start < t < end]
        if not inside:
            fails.append(f"{name} saved no state strictly inside its fall "
                         f"{fall}, so there is none to require")
            continue

        got = picked()
        if got[0] != times[0] or got[-1] != times[-1]:
            fails.append(f"{name}: the states documented, {got}, do not open on "
                         f"the first saved state {times[0]:g} and close on the "
                         f"last {times[-1]:g}")
        if len(got) != 4:
            fails.append(f"{name}: the default draws {len(got)} states, expected "
                         f"four")
        if got != sorted(got) or len(set(got)) != len(got):
            fails.append(f"{name}: the states are not in time order, or repeat: "
                         f"{got}")
        if any(t not in times for t in got):
            fails.append(f"{name}: a state was documented that the march never "
                         f"saved: {got}")

        # Not the states equal time steps would pick, and not the states even
        # spacing through the saved ones would pick either: both step over the
        # fall on one sample or the other.
        even_in_time = [min(times, key=lambda t, want=want: abs(t - want))
                        for want in [times[0] + i * (times[-1] - times[0]) / 3.0
                                     for i in range(4)]]
        if got == even_in_time:
            fails.append(f"{name}: the states {got} are the ones equal time steps "
                         f"would pick, so the schedule is not being followed")

        # The fall itself: a state STRICTLY inside it, at every count that has
        # room for one. The end of the fall does not count — the pool has arrived
        # by then — and admitting it is what let the earth dam pass while Johnson
        # documented no state of its drawdown at all.
        for wanted in range(3, len(times) + 1):
            states = picked(seep_transient_frames=wanted)
            if not [t for t in states if start < t < end]:
                fails.append(f"{name}: {wanted} states documented, {states}, and "
                             f"not one of them is inside the fall from "
                             f"{start:g} to {end:g}; the saved states there are "
                             f"{inside}")

        # The count is the caller's, and the toggle takes the frames away entirely.
        for wanted in (2, 3, 6, len(times), len(times) + 5):
            states = picked(seep_transient_frames=wanted)
            if len(states) != min(wanted, len(times)):
                fails.append(f"{name}: asking for {wanted} states drew "
                             f"{len(states)}: {states}")
            if states and (states[0] != times[0] or states[-1] != times[-1]):
                fails.append(f"{name}: asking for {wanted} states dropped an end: "
                             f"{states}")
        if picked(seep_transient_frames=0):
            fails.append(f"{name}: asking for no states still drew some")
        if picked(seep_transient_figures=False):
            fails.append(f"{name}: the frames toggle off still drew states")

    # And the toggle really removes them from the report, leaving the march it
    # still describes and the history it still draws.
    report = _tseep_report(options={"seep_transient_figures": False})
    sources = [f.source for f in report.figures()]
    if any(s.startswith("seepage tseep ") and s != "seepage tseep history"
           for s in sources):
        fails.append(f"the frames toggle off left frame figures: {sources}")
    if "seepage tseep history" not in sources:
        fails.append("the frames toggle took the history figure with it")
    if _tseep_section(report) is None:
        fails.append("the frames toggle took the whole transient subsection")
    return fails


def test_tseep_frames_share_one_scale():
    """Every state a march is documented at is drawn on ONE range per variable.

    A drawdown is one story. Scaled to their own fields the late states — whose
    field is nearly flat — re-normalize into a bullseye, and the same colour then
    means a different head at every time, so the fall of the water table is
    invisible. Held to one range a colour means the same amount at every instant
    and the change between states is the change in the field.
    """
    fails = []
    import numpy as np
    from xslope.report import SEEP_PANELS, transient_frame_times, tseep_bundles

    _slope_data, solutions = _tseep_solutions()
    bundle = tseep_bundles(solutions)[0]
    calls = _plot_tseep_calls()
    times = transient_frame_times(bundle, __import__(
        "xslope.report", fromlist=["x"]).resolve_options({}))
    if len(calls) != len(times) * len(SEEP_PANELS):
        fails.append(f"the march drew {len(calls)} frame panels, expected "
                     f"{len(times) * len(SEEP_PANELS)}")

    frames = [fr for fr in bundle["frames"]
              if any(abs(float(fr["time"]) - t) < 1e-9 for t in times)]
    mats = {kw.get("base_mat") for kw, _sd, _sol in calls}
    if len(mats) != 1:
        fails.append(f"the states are scaled to different base materials {mats}")

    for panel in SEEP_PANELS:
        variable = panel["variable"]
        drawn = [kw for kw, _sd, _sol in calls if kw.get("variable") == variable]
        if len(drawn) != len(times):
            fails.append(f"{variable!r} was drawn {len(drawn)} times, expected "
                         f"{len(times)}")
            continue
        ranges = {(kw.get("vmin"), kw.get("vmax")) for kw in drawn}
        if len(ranges) != 1 or None in list(ranges)[0]:
            fails.append(f"the {variable!r} states are drawn on {ranges}, not on "
                         f"one shared range")
            continue
        vmin, vmax = list(ranges)[0]
        own = [(float(np.nanmin(fr[variable])), float(np.nanmax(fr[variable])))
               for fr in frames]
        if vmin > min(o[0] for o in own) or vmax < max(o[1] for o in own):
            fails.append(f"the shared {variable!r} range {vmin}–{vmax} clips the "
                         f"fields, which span {min(o[0] for o in own)}–"
                         f"{max(o[1] for o in own)}")
        # And the states really would scale apart, so an independently scaled
        # series would not land on one range by accident.
        if len(set(own)) < len(own):
            fails.append(f"two states span the same {variable!r} values {own}, so "
                         f"an independently scaled series could pass this check")

    # The velocity ARROWS too: an arrow is a length standing for a speed, and each
    # state scaled to its own maximum drew its longest arrow the same length
    # however fast the section was draining.
    arrows = {kw.get("vector_max") for kw, _sd, _sol in calls if kw.get("vectors")}
    if len(arrows) != 1 or None in arrows:
        fails.append(f"the velocity arrows are scaled to {arrows}, not to one "
                     f"maximum across the march")
    return fails


def test_tseep_carries_no_flow_net():
    """A transient state is drawn as head contours, and said to be.

    A flow net is a stream function contoured into channels of equal flow, and
    that requires divergence-free through-flow. A state releasing water from
    storage is not that, so no stream function is stored and none is invented: the
    figures are contours, the captions call them contours, and the section says
    why before the first of them.
    """
    fails = []
    from xslope.plot_seep import flownet_has_flowlines
    from xslope.report import tseep_bundles

    _slope_data, solutions = _tseep_solutions()
    bundle = tseep_bundles(solutions)[0]
    report = _tseep_report()

    for i, frame in enumerate(bundle["frames"]):
        if flownet_has_flowlines(bundle["seep_data"], frame):
            fails.append(f"frame {i} would be drawn with flow lines")

    captions = [f.caption for f in report.figures()
                if f.source.endswith(" head")]
    if not captions:
        fails.append("the transient report drew no head figure")
    for caption in captions:
        if not caption.startswith("Head contours"):
            fails.append(f"a transient head figure is captioned {caption!r}")

    sub = _tseep_section(report)
    text = " ".join(b.text for b in (sub.blocks if sub else [])
                    if b.kind == "prose")
    if "no stream function" not in text:
        fails.append(f"the section never says why there is no flow net: {text!r}")
    if "flow net" in text:
        fails.append(f"the section calls a transient figure a flow net: {text!r}")
    if "flow lines on them" not in text:
        fails.append(f"the section does not say the head figures carry no flow "
                     f"lines: {text!r}")
    # And nothing anywhere in the report claims the figures draw flow lines.
    for block in report.blocks("prose"):
        if "draws the flow lines" in block.text or "the flow lines" in block.text:
            fails.append(f"a transient report names flow lines: {block.text!r}")

    # The figures really were asked for without them: a flow-line request on a
    # state with no stream function is a request the plotter has to refuse, and
    # refusing is not the same as never asking.
    for kw, _sd, solution in _plot_tseep_calls():
        if kw.get("variable") == "head" and solution.get("phi") is not None:
            fails.append("a head figure was drawn from a state carrying a stream "
                         "function")
    return fails


def test_tseep_history_figure():
    """The march over time is drawn, and the sentence names the traces it carries.

    The frame figures are the field at four instants. The history is every instant
    the march saved, of the quantities that have one number: the level the section
    is being driven by, the water table and the seepage face it reaches, and the
    two boundary rates whose difference IS the water coming out of storage.
    """
    fails = []
    import numpy as np
    from xslope.plot_seep import transient_has_history, transient_history
    from xslope.report import tseep_bundles

    _slope_data, solutions = _tseep_solutions()
    bundle = tseep_bundles(solutions)[0]
    seep_data, march = bundle["seep_data"], bundle["transient"]

    if not transient_has_history(seep_data, march):
        fails.append("the sample march carries no history, so the figure is "
                     "never drawn")
        return fails
    history = transient_history(seep_data, march)
    for key in ("level", "phreatic", "exit_point", "inflow", "outflow"):
        trace = history[key]
        if trace is None:
            fails.append(f"the sample march reads no {key} trace")
            continue
        if len(trace) != len(march["times"]):
            fails.append(f"the {key} trace has {len(trace)} values for "
                         f"{len(march['times'])} saved states")
        if not np.any(np.isfinite(trace)):
            fails.append(f"the {key} trace is entirely undefined")
    # The traces are the march, not constants: a level that never fell and a
    # water table that never moved would draw a history of flat lines.
    for key in ("level", "phreatic", "exit_point", "outflow"):
        trace = np.asarray(history[key], dtype=float)
        if float(np.nanmax(trace) - np.nanmin(trace)) <= 0:
            fails.append(f"the {key} trace never changes over the march")
    # The drawdown is read off the series' own breakpoints, not entered here.
    if history["drawdown"] != (2.0, 47.0):
        fails.append(f"the drawdown interval read off the schedule is "
                     f"{history['drawdown']}, expected (2.0, 47.0)")

    report = _tseep_report()
    figure = next((f for f in report.figures()
                   if f.source == "seepage tseep history"), None)
    if figure is None:
        fails.append("the transient report drew no history figure")
        return fails
    sub = _tseep_section(report)
    text = " ".join(b.text for b in (sub.blocks if sub else []) if b.kind == "prose")
    for named in ("the level the reservoir boundary is held at",
                  "the phreatic elevation at x = ",
                  "the top of the seepage face",
                  "the boundary inflow", "the boundary outflow"):
        if named not in text:
            fails.append(f"the history sentence does not name {named!r}: {text!r}")
    station = f"{history['station']:.4g}"
    if station not in text:
        fails.append(f"the station the water table is followed at, {station}, is "
                     f"not stated: {text!r}")

    # The toggle takes the figure and its sentence, and leaves the states.
    off = _tseep_report(options={"seep_transient_history": False})
    sources = [f.source for f in off.figures()]
    if "seepage tseep history" in sources:
        fails.append("the history toggle off still drew the history")
    if not any(s.startswith("seepage tseep ") for s in sources):
        fails.append("the history toggle took the frame figures with it")
    off_text = " ".join(b.text for b in off.blocks("prose"))
    if "the top of the seepage face" in off_text:
        fails.append("the history toggle left the sentence describing the figure")

    # A march with no boundary rates and no reservoir has no history to draw, and
    # is not asked for one.
    bare = {"frames": [{"time": t, "head": fr["head"], "u": fr["u"]}
                       for t, fr in zip(march["times"], march["frames"])],
            "times": march["times"]}
    empty = dict(seep_data)
    empty["seepage_bc"] = {}
    empty["head_series_bindings"] = []
    if transient_has_history(empty, bare):
        fails.append("a march with neither a water level nor a boundary rate was "
                     "reported to carry a history")
    return fails


def test_tseep_dual_basis():
    """A model documented on both bases presents both, the march after the states
    it runs between, and each says which it is.

    A rapid drawdown can be documented as the two states it runs between and as
    the march between them, and the two are different answers for the same
    section: the steady drawn-down state is the field the pool would reach given
    time, and the march is the field it actually has on the day the pool stopped
    falling. No model in the corpus ships both, so the pair is built here — the
    march's own first state, restored as a steady solution, beside the march.
    """
    fails = []
    from xslope.report import seep_bundles, tseep_bundles

    slope_data, solutions = _tseep_solutions()
    transient = tseep_bundles(solutions)[0]
    steady = {"seep_data": transient["seep_data"],
              "solution": dict(transient["frames"][0], flowrate=None,
                               unconfined=True),
              "options": {"bc": 1}}
    both = {"seep": [steady], "tseep": [transient]}
    report = _tseep_report(solutions=both)

    expected = [(1, "Traceability"), (1, "Project Definition"),
                (1, "Seepage Analysis"), (2, "Analysis Inputs"),
                (2, "Seepage Mesh"), (2, "Results"), (2, "Transient Analysis")]
    got = report.section_titles()
    if got != expected:
        fails.append(f"the two-basis report's sections are {got}, expected "
                     f"{expected}")

    # Each basis says which it is, and the march comes second.
    proses = _seep_results_prose(report)
    if len(proses) != 2:
        fails.append(f"the two-basis section has {len(proses)} results "
                     f"subsections, expected two")
        return fails
    if "steady state" not in proses[0]:
        fails.append(f"the steady basis does not say it is a steady state: "
                     f"{proses[0]!r}")
    if "transient" in proses[0]:
        fails.append(f"the steady basis calls itself transient: {proses[0]!r}")
    if "transient analysis" not in proses[1]:
        fails.append(f"the march does not say it is transient: {proses[1]!r}")
    if "steady state" in proses[1]:
        fails.append(f"the march calls itself a steady state: {proses[1]!r}")

    # The mesh is drawn once: the march runs on the first set's own boundaries, so
    # a second figure of it would be the same picture under a second number.
    sources = [f.source for f in report.figures()]
    if sources.count("seepage bc1 mesh") != 1:
        fails.append(f"the two-basis report drew the mesh "
                     f"{sources.count('seepage bc1 mesh')} times: {sources}")
    if not any(s.startswith("seepage bc1 ") and s.endswith(" head")
               for s in sources):
        fails.append(f"the steady basis drew no head figure: {sources}")
    if not any(s.startswith("seepage tseep ") for s in sources):
        fails.append(f"the march drew no figures: {sources}")
    if sources.index("seepage bc1 head") > min(
            i for i, s in enumerate(sources) if s.startswith("seepage tseep ")):
        fails.append(f"the march is documented before the states it runs "
                     f"between: {sources}")

    # And a model documented on ONE basis does not carry the sentence that tells
    # two apart: there is nothing for it to distinguish it from.
    alone = " ".join(_seep_results_prose(_tseep_report()))
    if "steady state" in alone:
        fails.append(f"a march documented alone is told apart from a basis the "
                     f"report does not carry: {alone!r}")
    steady_alone = " ".join(_seep_results_prose(_seep_report(SEEP_XLSX)))
    if "steady state" in steady_alone:
        fails.append(f"a steady solve documented alone is told apart from a "
                     f"basis the report does not carry: {steady_alone!r}")
    return fails


def test_tseep_reaches_the_report_from_studio():
    """A march solved in Studio reaches the report as the basis it is.

    The seepage runner emits one bundle per boundary condition set for a steady
    run and one march for a transient one, and the window keeps them under
    different keys. Handing the report the steady ones alone documented a session
    whose only run was a march as a session that had solved nothing.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    _app()
    from studio.main_window import MainWindow
    from xslope.report import seep_bundles, tseep_bundles

    slope_data, solutions = _tseep_solutions()
    bundle = tseep_bundles(solutions)[0]
    # The shape the runner emits, built as the runner builds it (studio/runners.py
    # ``_run_transient``): the march, its plottable frames, and the run options.
    from_runner = {"mode": "transient", "seep_data": bundle["seep_data"],
                   "transient": bundle["transient"],
                   "frames": bundle["frames"], "options": {"mode": "transient"}}

    mw = MainWindow()
    try:
        mw.doc.slope_data = slope_data
        mw.doc.results["transient_seep"] = from_runner
        got = mw.report_solutions()
        forwarded = tseep_bundles(got)
        if len(forwarded) != 1:
            fails.append(f"a session that marched a model hands the report "
                         f"{len(forwarded)} marches: {sorted(got)}")
            return fails
        if forwarded[0] is not from_runner:
            fails.append("the runner's own bundle is not what reaches the report")
        if seep_bundles(got):
            fails.append("a transient march was forwarded as a steady set")
        # The action a report is started from is live on a march alone.
        mw._update_run_actions()
        if not mw.act_report.isEnabled():
            fails.append("a session whose only run is a march leaves Generate "
                         "Report dimmed")
        # And the bundle really builds the section, so forwarding it is not a key
        # the builder ignores.
        report = _tseep_report(solutions={"tseep": [from_runner]})
        if _tseep_section(report) is None:
            fails.append("the runner's bundle produced no transient subsection")
    finally:
        mw.close()
        mw.deleteLater()
    return fails


def test_tseep_dialog_rows():
    """The dialog offers the transient figures under the seepage branch, and the
    boxes reach the builder."""
    fails = []
    _app()
    from studio.report_dialog import CONTENT_TREE, ReportDialog
    from xslope.report import DEFAULT_OPTIONS

    parents = {}
    for key, _label, _tip, children in CONTENT_TREE:
        for child_key, _l, _t in children:
            parents[child_key] = key
    for key in ("seep_transient_figures", "seep_transient_history"):
        if key not in DEFAULT_OPTIONS:
            fails.append(f"{key} is not an option the builder reads")
        elif DEFAULT_OPTIONS[key] is not True:
            fails.append(f"{key} is off by default; documenting a march means "
                         f"documenting what it did")
        if parents.get(key) != "seep":
            fails.append(f"the {key} row is a child of {parents.get(key)!r}, not "
                         f"of the seepage section that prints it")
    # The count is a builder option, not a box: a number has no checkbox.
    if "seep_transient_frames" in parents:
        fails.append("the dialog offers seep_transient_frames as a tick box; it "
                     "is a count")
    if not isinstance(DEFAULT_OPTIONS.get("seep_transient_frames"), int):
        fails.append("seep_transient_frames is not a count")

    slope_data, solutions = _tseep_solutions()
    dlg = ReportDialog(slope_data=slope_data, solutions=solutions,
                       model_path=TSEEP_XLSX, default_method="bishop")
    try:
        opts = dlg.options()
        for key in ("seep_transient_figures", "seep_transient_history"):
            if opts.get(key) is not True:
                fails.append(f"the dialog opens with {key}={opts.get(key)!r}")
            item = dlg._items.get(key) if hasattr(dlg, "_items") else None
            if item is None:
                fails.append(f"the dialog has no widget for {key}")
        from PySide6.QtCore import Qt
        for key in ("seep_transient_figures", "seep_transient_history"):
            dlg._items[key].setCheckState(0, Qt.Unchecked)
        opts = dlg.options()
        for key in ("seep_transient_figures", "seep_transient_history"):
            if opts.get(key) is not False:
                fails.append(f"unticking {key} still reports {opts.get(key)!r}")
    finally:
        dlg.close()
        dlg.deleteLater()
    return fails


def test_fem_section():
    """A report of a strength reduction run states its factor of safety, in bold,
    and a report of a single trial states no factor of safety at all."""
    fails = []
    from xslope.report import FEM_PANELS
    _slope_data, bundle = _fem_bundle()
    report = _engine_report("fem")

    # The mesh comes LAST of the inputs, under a heading of its own: it is what
    # the model and everything on it were discretized onto, and a reader meets it
    # having already met the properties and the loads it carries (the owner's
    # sequencing, fem_johnson_res review).
    expected = [(1, "Traceability"), (1, "Project Definition"),
                (1, "Deformation and Strength Reduction"),
                (2, "Analysis Inputs"), (2, "Loads"),
                (2, "Finite Element Mesh"), (2, "Results")]
    got = report.section_titles()
    if got != expected:
        fails.append(f"the SSRM report's sections are {got}, expected {expected}")

    fs = bundle["FS"]
    stated = [b for b in report.blocks("prose") if f"{fs:.3f}" in b.text]
    if not stated:
        fails.append(f"the strength reduction factor of safety {fs:.3f} is stated "
                     f"nowhere: {_prose(report)!r}")
    elif not any(f"{fs:.3f}" in " ".join(b.bold) for b in stated):
        fails.append(f"the factor of safety {fs:.3f} is stated but not set in "
                     f"bold, which is how the report states one")
    if "factor of safety" not in " ".join(_prose(report)):
        fails.append("the SSRM section never uses the words factor of safety")

    # The inputs: the stiffness columns are what make this table the finite
    # element one rather than a second copy of the strength table.
    sec = next((s for s in report.sections
                if s.title == "Deformation and Strength Reduction"), None)
    inputs = next((c for c in (sec.children if sec else [])
                   if c.title == "Analysis Inputs"), None)
    if inputs is None:
        fails.append("the finite element section has no Analysis Inputs")
    else:
        table = next((b for b in inputs.blocks if b.kind == "table"), None)
        if table is None:
            fails.append("the finite element inputs carry no properties table")
        else:
            for header in ("E", "ν"):
                if not any(h == header or h.startswith(header + " ")
                           for h in table.headers):
                    fails.append(f"the finite element material table has no "
                                 f"{header} column: {table.headers}")

    planned, drawn = _planned_matches(report, "fem")
    if planned != drawn:
        fails.append(f"the SSRM report planned {planned} figures and built {drawn}")
    sources = [f.source for f in report.figures()]
    # The model, the mesh, one panel per field, and the search that reached the
    # factor of safety — the corpus run records its trials, so the search is
    # drawn and is part of what a full strength reduction report is.
    for wanted in ("fem model", "fem mesh", "fem run1 deformation",
                   "fem run1 shear_strain", "fem run1 displace_vector",
                   "fem run1 search"):
        if wanted not in sources:
            fails.append(f"the SSRM report has no {wanted!r} figure: {sources}")
    if drawn != 3 + len(FEM_PANELS):
        fails.append(f"the SSRM report drew {drawn} figures, expected the model, "
                     f"the mesh, the {len(FEM_PANELS)} result panels and the "
                     f"search: {sources}")

    # The mesh figure is Studio's FEM data view, from the same function, with the
    # boundary conditions on it: a mesh drawn without them does not say what the
    # section was held by, and a report that draws its own version of a plot
    # Studio already draws is a second answer to one question.
    import xslope.plot_fem as pf
    real_fem_data = pf.plot_fem_data
    seen = {}

    def spy_fem_data(fem_data, **kw):
        seen.update(kw)
        return real_fem_data(fem_data, **kw)

    pf.plot_fem_data = spy_fem_data
    try:
        _engine_report("fem", options={"fem_figure": False,
                                       "fem_inputs_figure": False})
    finally:
        pf.plot_fem_data = real_fem_data
    if not seen:
        fails.append("the finite element mesh figure is not drawn by "
                     "plot_fem_data, the function Studio's FEM data view draws")
    elif seen.get("show_bc") is not True:
        fails.append(f"the mesh figure is drawn with show_bc="
                     f"{seen.get('show_bc')!r}; the boundary conditions are "
                     f"half of what the figure says")
    mesh_caption = next((f.caption for f in report.figures()
                         if f.source == "fem mesh"), "")
    if "boundary condition" not in mesh_caption.lower():
        fails.append(f"the mesh figure is captioned {mesh_caption!r}, which "
                     f"does not name the boundary conditions it carries")

    # Each figure carries its own option, and switching one off takes only it.
    for option, gone in (("fem_inputs_figure", 1), ("fem_mesh_figure", 1),
                         ("fem_figure", len(FEM_PANELS))):
        off = _engine_report("fem", options={option: False})
        planned, off_drawn = _planned_matches(off, "fem", options={option: False})
        if planned != off_drawn:
            fails.append(f"{option}=False planned {planned} figures and built "
                         f"{off_drawn}")
        if off_drawn != drawn - gone:
            fails.append(f"{option}=False left {off_drawn} figures, not the "
                         f"{drawn - gone} the other options draw")

    # A single trial is not a strength reduction run: it reports displacements
    # and claims no factor of safety, and the heading says so.
    single = dict(bundle, analysis="single", FS=None)
    plain = _engine_report("fem", bundle=single)
    titles = _titles(plain)
    if "Deformation and Strength Reduction" in titles:
        fails.append("a single-trial run is headed as a strength reduction run")
    if "Deformation Analysis" not in titles:
        fails.append(f"a single-trial run has no deformation section: {titles}")
    text = " ".join(_prose(plain))
    if f"{fs:.3f}" in text:
        fails.append("a single-trial run reports the strength reduction factor "
                     "of safety anyway")
    if "no factor of safety" not in text:
        fails.append(f"a single-trial run does not say it reports no factor of "
                     f"safety: {text!r}")
    return fails


def test_fem_mesh_legend_names_what_it_holds():
    """The mesh figure's fixity legend says what each symbol holds, in words, and
    the roller entry names the axis of the nodes it is drawn on.

    It read "Fixed (bc_type=1)" and "Y-Roller (bc_type=3)" — an internal array
    code in a figure an engineer signs, and the second one drawn on the nodes
    build_fem_data marks 2 (u = 0, v free), which is the OTHER axis. No node in
    any model carries a 3 at all, so the name could never have been right.
    """
    fails = []
    import numpy as np
    import matplotlib.pyplot as plt
    from matplotlib.legend import Legend
    from xslope.plot_fem import plot_fem_data

    _slope_data, bundle = _fem_bundle()
    fem_data = bundle["fem_data"]

    # What the mesh actually carries, so the legend is checked against the nodes
    # it is drawn on rather than against another sentence about them.
    bc = np.asarray(fem_data["bc_type"])
    present = sorted({int(t) for t in np.unique(bc)})
    if 2 not in present:
        fails.append(f"the fixture's mesh carries no roller node, so the roller "
                     f"legend entry is untested: bc types {present}")
    if 3 in present:
        fails.append(f"a node carries bc_type 3, which build_fem_data never "
                     f"assigns: bc types {present}")
    # The roller nodes are held in x and free in y — the restraint the legend
    # names. Read off the constraint the assembler applies, not off the label.
    rollers = np.where(bc == 2)[0]

    fig = plt.figure(figsize=(6, 4))
    try:
        with contextlib.redirect_stdout(io.StringIO()):
            plot_fem_data(fem_data, fig=fig, show_title=False, show_bc=True,
                          show_nodes=False)
        labels = [t.get_text() for legend in fig.findobj(Legend)
                  for t in legend.get_texts()]
    finally:
        plt.close(fig)

    if not labels:
        fails.append("the mesh figure carries no legend, so its names are untested")
    for label in labels:
        if "bc_type" in label:
            fails.append(f"the mesh legend prints the internal code in {label!r}")
    if "Fixed" not in labels:
        fails.append(f"the mesh legend does not name the fixed nodes: {labels}")
    roller = [l for l in labels if "oller" in l]
    if len(rollers) and not roller:
        fails.append(f"{len(rollers)} nodes are held on one axis and the legend "
                     f"names no roller: {labels}")
    for label in roller:
        if "Y-Roller" in label or "y-roller" in label:
            fails.append(f"the roller entry is {label!r}; it is drawn on the "
                         f"nodes held in x, which are free VERTICALLY")
        if "vertical" not in label.lower():
            fails.append(f"the roller entry is {label!r} and does not say which "
                         f"axis those nodes are free on")
    return fails


def test_the_fem_section_states_its_pore_pressure_basis():
    """The finite element section says where the pore pressure its elements were
    solved with came from, and a dry model says nothing.

    johnson_res's materials are all u='seep' and build_fem_data feeds the saved
    field straight into the assembly, with the seepage section that computed it
    four pages up. The section named neither — a reader could not tell an
    effective-stress run from a dry one.
    """
    fails = []
    from xslope.report import (FEM_PORE_SOURCES, PORE_SOURCES, _fem_pore_basis)

    def basis(fem_data, solutions=None, opts=None):
        from xslope.report import resolve_options
        block = _fem_pore_basis(fem_data, solutions or {},
                                resolve_options(opts or {}))
        return block.text if block is not None else ""

    # Every source the assembly has a branch for gets its own words, and they are
    # the words the limit equilibrium section uses for the same source.
    for option, said in FEM_PORE_SOURCES.items():
        text = basis({"pp_option": option})
        if said not in text:
            fails.append(f"a model on u={option!r} is described as {text!r}, "
                         f"which does not name {said!r}")
        if PORE_SOURCES.get(option) != said:
            fails.append(f"the finite element section calls u={option!r} "
                         f"{said!r} and the limit equilibrium section calls it "
                         f"{PORE_SOURCES.get(option)!r}")
    # A dry model, and one whose option the assembly has no branch for, are not
    # credited with a source.
    for option in ("none", "", None, "cons"):
        if basis({"pp_option": option}):
            fails.append(f"a model on u={option!r} is credited with a pore "
                         f"pressure source: {basis({'pp_option': option})!r}")

    # On the real two-engine model: the sentence is there, and it sends the
    # reader to the analysis that computed the field.
    slope_data, solutions = _restored(JOHNSON_XLSX)
    fem_data = (solutions.get("fem") or {}).get("fem_data") or {}
    if str(fem_data.get("pp_option")) != "seep":
        fails.append(f"johnson_res reads pp_option={fem_data.get('pp_option')!r}; "
                     f"the seepage-fed case is untested")
    report = _built_report(slope_data, solutions,
                           {"input_path": JOHNSON_XLSX, "lem": False,
                            "pd_figure": False, "seep_inputs_figure": False,
                            "seep_mesh_figure": False, "seep_kr_figure": False,
                            "seep_figures": False, "fem_inputs_figure": False,
                            "fem_mesh_figure": False, "fem_figure": False})
    stated = [b for b in report.blocks("prose")
              if "Pore pressure at every node" in b.text]
    if len(stated) != 1:
        fails.append(f"the finite element section states its pore pressure "
                     f"basis {len(stated)} times")
    else:
        text = stated[0].text
        if PORE_SOURCES["seep"] not in text:
            fails.append(f"the basis is stated as {text!r}")
        if "Section" not in text:
            fails.append(f"the field is an outcome of an analysis this report "
                         f"documents and the reader is not sent to it: {text!r}")
        if not any(url.startswith("#") for _phrase, url in stated[0].links):
            fails.append(f"the section citation is not a link: {stated[0].links}")

    # A report that does NOT document the flow analysis states the source and
    # points nowhere.
    text = basis(fem_data, solutions, {"seep": False})
    if PORE_SOURCES["seep"] not in text:
        fails.append(f"the source is dropped with the seepage section: {text!r}")
    if "Section" in text:
        fails.append(f"a report with no seepage section still cites one: {text!r}")

    # A dry finite element model says nothing about water.
    dry = _engine_report("fem")
    if any("Pore pressure at every node" in b.text
           for b in dry.blocks("prose")):
        fails.append("a model whose materials take no pore pressure is given a "
                     "pore pressure basis")
    return fails


def test_the_ssrm_paragraph_describes_what_runs():
    """The strength reduction paragraph describes bracketing and bisection, and
    names the tolerance and the failure criterion where the run recorded them.

    It said the solution was "repeated at increasing factors until equilibrium
    can no longer be reached" — a staged march. solve_ssrm brackets the critical
    factor and halves the bracket to a tolerance, and what counts as a failed
    trial is decided by a criterion (hybrid by default, which demands
    displacement evidence). None of that was on the page.
    """
    fails = []
    import json
    from xslope.report import SSRM_CRITERIA, _ssrm_criterion, _ssrm_procedure

    said = " ".join(_prose(_engine_report("fem")))
    if "increasing factors" in said:
        fails.append(f"the paragraph still describes a staged march: {said!r}")
    for wanted in ("bracket", "halved", "midpoint"):
        if wanted not in said:
            fails.append(f"the paragraph never uses the word {wanted!r}: {said!r}")

    # A run that recorded neither, so neither is claimed. Built here by stripping
    # the record out of a copy of the corpus model's own metadata, rather than by
    # pointing at a corpus file that happens not to carry it: every shipped run
    # now records its bracket and its criterion, and a check whose subject is the
    # SILENT case must not be able to lose its subject to a regeneration.
    silent = _silent_run_bundle()
    for key in ("tolerance", "failure_criterion", "method"):
        if (silent.get("meta") or {}).get(key) is not None:
            fails.append(f"the silent fixture records {key}; the case is not "
                         f"tested")
    quiet = _ssrm_procedure(silent)
    if "narrower than the solution tolerance" not in quiet:
        fails.append(f"a run recording no tolerance names one anyway: {quiet!r}")
    for sentence in SSRM_CRITERIA.values():
        if sentence in quiet:
            fails.append("a run recording no failure criterion is given one")

    # A run that DID record them says so, in its own numbers. The corpus records
    # the criterion as solve_ssrm's method string, so that is what is fed in.
    recorded = json.load(open(os.path.join(
        _REPO, "docs", "fem", "files", "xslope_griffiths1_fem_meta.json")))
    if recorded.get("tolerance") is None or recorded.get("method") is None:
        fails.append("the corpus meta this check reads records no tolerance or "
                     "method, so the derived case is untested")
    else:
        text = _ssrm_procedure({"meta": recorded})
        if f"narrower than {recorded['tolerance']:g}" not in text:
            fails.append(f"a run recording tolerance={recorded['tolerance']} "
                         f"does not state it: {text!r}")
        kind = _ssrm_criterion(recorded["method"])
        if kind != "hybrid":
            fails.append(f"the recorded method {recorded['method']!r} reads as "
                         f"{kind!r}")
        elif SSRM_CRITERIA["hybrid"] not in text:
            fails.append(f"a hybrid run does not say what a failed trial is: "
                         f"{text!r}")

    # Hybrid's own method string contains the word non-convergence; matching it
    # to the wrong criterion is the mistake this ordering exists to prevent.
    for method, kind in (
            ("SSRM — Hybrid (non-convergence corroborated by displacement "
             "evidence)", "hybrid"),
            ("SSRM — Non-Convergence (Griffiths & Lane 1999; equilibrium test "
             "after Dawson, Roth & Drescher 1999)", "non_convergence"),
            ("SSRM — Displacement Limit", "displacement_limit"),
            ("SSRM — Displacement Catastrophe (Sun et al. 2021)",
             "displacement_increase"),
            ("something else entirely", None)):
        got = _ssrm_criterion(method)
        if got != kind:
            fails.append(f"{method!r} reads as {got!r}, not {kind!r}")

    # The final bracket, where a run records one.
    text = _ssrm_procedure({"solution": {"final_interval": (1.34, 1.35)}})
    if "1.340 to 1.350" not in text:
        fails.append(f"a recorded final bracket is not stated: {text!r}")
    if "which is" in _ssrm_procedure({}):
        fails.append("a run recording no bracket is given one")
    return fails


def _fem_run_forgetting(keys):
    """``(slope_data, bundle)`` for the corpus strength reduction run with the
    named facts struck out of its record — the shape of a run saved before that
    fact was ever persisted, which is what every older file on a reader's disk
    still is.

    Built by subtraction from a real bundle rather than found among the shipped
    models. A check about what the report does when a fact is MISSING used to be
    aimed at whichever corpus file happened not to carry it; regenerating that
    file retired the case silently, leaving the check passing on a model that no
    longer poses the question. Subtracting names the absence out loud.
    """
    slope_data, bundle = _fem_bundle()
    meta = {k: v for k, v in (bundle.get("meta") or {}).items() if k not in keys}
    solution = {k: v for k, v in bundle["solution"].items() if k not in keys}
    forgot = dict(bundle, meta=meta, solution=solution)
    for key in keys:
        forgot.pop(key, None)
    return slope_data, forgot


def _silent_run_bundle():
    """The corpus run with every record of HOW it was solved struck out, keeping
    only the answer."""
    from xslope.report import SSRM_RECORD_KEYS

    _slope_data, bundle = _fem_run_forgetting(set(SSRM_RECORD_KEYS) - {"FS"})
    return bundle


def _restored_without_member_forces():
    """``(stem, slope_data, solutions)`` for the reinforcement sample read back
    from companions that carry no bar forces.

    The model and its companions are copied and the two force files are left
    behind, which is what a run saved before those files existed looks like —
    and what the sample itself was until its companions were regenerated. Cached:
    two checks ask the same question of it.
    """
    key = ("noforces", FEM_REINF_XLSX)
    if key not in _ENGINE:
        tmp = tempfile.mkdtemp(prefix="xslope_noforce_")
        stem = _sidecar_copy(os.path.splitext(FEM_REINF_XLSX)[0], tmp,
                             drop=("_fem_reinf.csv", "_fem_failure_reinf.csv"))
        slope_data, solutions = _restored(f"{stem}.xlsx")
        _ENGINE[key] = (stem, slope_data, solutions)
    return _ENGINE[key]


def _declared_fem_model(tmp, k0=None, t_cut=None):
    """``(slope_data, bundle)`` for a finite element model that DECLARES an input
    no shipped model does — an at-rest coefficient, a tensile cap, or both.

    The declaration goes through the file: the corpus model is saved back out
    with the value set, and read again, so what the report is asked about is what
    the loader makes of a workbook carrying it, not a dict this check assembled.
    The mesh companion is copied alongside so the reloaded model discretizes onto
    the same section, and the finite element model is rebuilt from the reloaded
    inputs — nothing is solved.
    """
    from xslope.fem import build_fem_data
    from xslope.fileio import (default_template_path, load_slope_data,
                               save_slope_data_to_xlsx)

    source, bundle = _fem_bundle()
    edited = dict(source)
    if k0 is not None:
        edited["k0"] = k0
    if t_cut is not None:
        edited["materials"] = [dict(m) for m in source["materials"]]
        edited["materials"][0]["t_cut"] = t_cut
    out = os.path.join(tmp, "declared.xlsx")
    shutil.copy(os.path.splitext(FEM_XLSX)[0] + "_mesh.json",
                os.path.join(tmp, "declared_mesh.json"))
    with contextlib.redirect_stdout(io.StringIO()):
        save_slope_data_to_xlsx(edited, out, template=str(default_template_path()))
        loaded = load_slope_data(out)
        fem_data = build_fem_data(loaded, loaded["mesh"])
    return loaded, dict(bundle, fem_data=fem_data)


def test_the_in_situ_stress_assumption_is_stated():
    """Every finite element section says what stress state the section started
    from — the at-rest coefficient where the model declares one, and the gravity
    turn-on where it does not.

    K0 is a run option no corpus model declares, so the ``Initial stress state``
    row had never been printed by anything. The bigger half is the other branch:
    a run without K0 does not start from no assumption, it starts from
    sigma_h = nu/(1-nu) sigma_v, which is the STIFFNESS choosing the lateral
    stress. Two runs of one section from those two states reach different factors
    of safety, and the report named neither.
    """
    fails = []
    from xslope.report import FEM_IN_SITU_GRAVITY, FEM_IN_SITU_K0

    # --- the common case: nothing declared, and the default is on the page ----
    said = " ".join(_prose(_engine_report("fem")))
    if FEM_IN_SITU_GRAVITY not in said:
        fails.append(f"a run with no at-rest coefficient does not say what "
                     f"in-situ state it started from: {said!r}")
    if FEM_IN_SITU_K0 in said:
        fails.append("a run with no at-rest coefficient is credited with "
                     "building its initial stress from the overburden")

    # --- the declared case: the row, and the sentence that goes with it -------
    tmp = tempfile.mkdtemp(prefix="xslope_k0_")
    slope_data, bundle = _declared_fem_model(tmp, k0=1.0)
    if float(bundle["fem_data"].get("k0") or 0) != 1.0:
        fails.append(f"a workbook declaring K0 = 1 loads and builds a model "
                     f"carrying k0={bundle['fem_data'].get('k0')!r}")
    report = _built_report(slope_data, {"fem": bundle},
                           {"input_path": os.path.join(tmp, "declared.xlsx"),
                            "lem": False, "pd_figure": False})
    rows = {}
    for block in report.blocks("keyvalues"):
        rows.update(dict(block.items))
    if rows.get("Initial stress state") != "K₀ = 1":
        fails.append(f"a model declaring K0 = 1 does not print it: "
                     f"{rows.get('Initial stress state')!r}")
    told = " ".join(_prose(report))
    if FEM_IN_SITU_K0 not in told:
        fails.append(f"a declared at-rest coefficient gets no sentence saying "
                     f"what it built: {told!r}")
    if FEM_IN_SITU_GRAVITY in told:
        fails.append("a run that built its initial stress from the overburden "
                     "is also described as a gravity turn-on")
    return fails


def test_a_tensile_cap_reaches_the_materials_table():
    """A material with a tension cutoff prints its σ_t in the finite element
    materials table.

    The column exists and no corpus model fills it, so a cap that halves a
    factor of safety — the RS2-62 failure mode, where dropped tensile caps
    doubled the strength reduction answer — was reported as if it were not there.
    """
    fails = []
    tmp = tempfile.mkdtemp(prefix="xslope_tcut_")
    slope_data, bundle = _declared_fem_model(tmp, t_cut=1.5)
    caps = [m.get("t_cut") for m in slope_data["materials"]]
    if 1.5 not in caps:
        fails.append(f"a workbook declaring a tensile cap loads without one: "
                     f"{caps}")
        return fails
    report = _built_report(slope_data, {"fem": bundle},
                           {"input_path": os.path.join(tmp, "declared.xlsx"),
                            "lem": False, "pd_figure": False})

    def fem_materials(rep):
        for t in rep.blocks("table"):
            if "Finite element material" in (t.caption or ""):
                return t
        return None

    table = fem_materials(report)
    if table is None:
        fails.append(f"the finite element materials table is absent: "
                     f"{[t.caption for t in report.blocks('table')]}")
        return fails
    headers = [str(h) for h in table.headers]
    cap_col = [i for i, h in enumerate(headers) if h.startswith("σ_t")]
    if not cap_col:
        fails.append(f"a model carrying a tensile cap prints no σ_t column: "
                     f"{headers}")
        return fails
    printed = [row[cap_col[0]] for row in table.rows]
    if "1.5" not in printed:
        fails.append(f"the declared cap of 1.5 is not in the σ_t column: "
                     f"{printed}")

    # A model with no cap does not carry the column at all — the report prints
    # the properties the analysis reads, not every column the schema has.
    plain = fem_materials(_engine_report("fem"))
    if plain is not None and any(str(h).startswith("σ_t") for h in plain.headers):
        fails.append(f"a model with no tensile cap prints a σ_t column: "
                     f"{plain.headers}")
    return fails


def test_a_low_order_mesh_carries_its_caution():
    """A finite element analysis run on tri3 or quad4 says on the page what those
    elements do to its answer.

    solve_ssrm prints the volumetric-locking warning to stdout, where a reader of
    the report never sees it: the section is drawn, the factor of safety is
    printed in bold, and nothing on the page says it is overestimated. tri6 is
    the default and every shipped model uses a quadratic element, so no report
    had ever carried the sentence.
    """
    fails = []
    import numpy as np
    from xslope.report import _low_order_caution

    _slope_data, bundle = _fem_bundle()
    fem_data = bundle["fem_data"]

    # The quadratic mesh the model was actually solved on says nothing.
    if _low_order_caution(fem_data):
        fails.append("a quadratic mesh is cautioned about volumetric locking")
    if _low_order_caution({}):
        fails.append("a container carrying no mesh is cautioned about its "
                     "element order")

    # Each low-order kind is named by the name the mesh row uses for it.
    for code, name in ((3, "tri3"), (4, "quad4")):
        types = np.full(len(fem_data["element_types"]), code)
        said = _low_order_caution(dict(fem_data, element_types=types))
        if name not in said:
            fails.append(f"a {name} mesh is not told it is {name}: {said!r}")
        for wanted in ("overestimated", "tri6, quad8 and quad9"):
            if wanted not in said:
                fails.append(f"a {name} mesh's caution does not say {wanted!r}: "
                             f"{said!r}")

    # And it reaches the page, under the mesh it is about.
    mixed = np.array(fem_data["element_types"])
    mixed = np.where(np.arange(len(mixed)) < 3, 3, mixed)
    low = dict(bundle, fem_data=dict(fem_data, element_types=mixed))
    told = " ".join(_prose(_engine_report("fem", bundle=low)))
    if "low-order elements" not in told:
        fails.append(f"a report of a mesh carrying tri3 elements does not "
                     f"carry the caution: {told!r}")
    return fails


def test_a_catastrophe_run_is_described_as_one():
    """A displacement catastrophe run is described by the search it ran, not by
    the bracketing search the other three criteria run.

    ``failure_criterion='displacement_increase'`` sweeps a series of factors,
    reads the viscoplastic displacement of a characteristic point off each, takes
    the interval where that displacement jumps hardest and halves it. BOTH ends of
    that interval can reach equilibrium, so "bracketed between a trial the section
    stands under and one it does not" describes something that did not happen —
    and it was printed on every run, because the criterion sentence that would
    have qualified it was unreachable: solve_ssrm records this run as "SSRM —
    Displacement Catastrophe (Sun et al. 2021)" and nothing matched that string.
    """
    fails = []
    import inspect
    from xslope.fem import _ssrm_displacement_increase
    from xslope.report import (SSRM_BISECTION, SSRM_CATASTROPHE, SSRM_CRITERIA,
                               _ssrm_criterion, _ssrm_procedure)

    # The two strings solve_ssrm can leave behind for this criterion: the option
    # name a live run records, and the method string a saved one does.
    for record in ("displacement_increase",
                   "SSRM — Displacement Catastrophe (Sun et al. 2021)"):
        if _ssrm_criterion(record) != "displacement_increase":
            fails.append(f"{record!r} reads as "
                         f"{_ssrm_criterion(record)!r}, so a catastrophe run is "
                         f"described as some other search")
        text = _ssrm_procedure({"meta": {"method": record, "tolerance": 0.05,
                                         "final_interval": [1.20, 1.25]}})
        if "bracketed between a trial" in text:
            fails.append(f"a catastrophe run recorded as {record!r} is described "
                         f"as a bracketing run: {text!r}")
        for wanted in ("series of factors", "jumps hardest",
                       "both ends of the interval may have reached one"):
            if wanted not in text:
                fails.append(f"a catastrophe run does not say {wanted!r}: "
                             f"{text!r}")
        if "narrower than 0.05" not in text or "1.200 to 1.250" not in text:
            fails.append(f"a catastrophe run drops its own tolerance or interval:"
                         f" {text!r}")
        # It is described ONCE: the criterion sentences belong to the searches
        # that have a failed trial, and this search's own paragraph has already
        # said what moves its interval.
        for said in SSRM_CRITERIA.values():
            if said in text:
                fails.append(f"a catastrophe run is given a failed-trial rule as "
                             f"well as its own description: {text!r}")

    # The other three still get the bracketing description — the branch takes only
    # what it names.
    for record in ("hybrid", "non_convergence", "displacement_limit"):
        text = _ssrm_procedure({"meta": {"failure_criterion": record}})
        if "bracketed between a trial" not in text:
            fails.append(f"a {record!r} run is no longer described as bracketing:"
                         f" {text!r}")
        if "jumps hardest" in text:
            fails.append(f"a {record!r} run is described as a catastrophe sweep: "
                         f"{text!r}")
        if SSRM_CRITERIA[record] not in text:
            fails.append(f"a {record!r} run does not say what a failed trial is")

    # And the description matches the code that runs. The two searches differ in
    # the two ways the paragraphs claim: one sweeps a series of factors before it
    # refines, the other does not; one moves its interval on a displacement RATIO,
    # the other on a trial's verdict.
    source = inspect.getsource(_ssrm_displacement_increase)
    for token, why in (("np.linspace(F_min, F_max, n_sweep)",
                        "the sweep the paragraph describes"),
                       ("ratio_left_half", "the displacement ratio the paragraph "
                        "says moves the interval")):
        if token not in source:
            fails.append(f"the catastrophe solver no longer carries {why} "
                         f"({token!r}), so the paragraph describes something else")
    if "converged" in SSRM_CATASTROPHE and "may have reached one" not in SSRM_CATASTROPHE:
        fails.append("the catastrophe paragraph claims something about "
                     "convergence without saying both ends can converge")
    if SSRM_BISECTION == SSRM_CATASTROPHE:
        fails.append("the two searches are described by the same words")
    return fails


def test_an_unrecorded_analysis_is_not_called_a_single_trial():
    """A solution whose companions record no analysis type says so, rather than
    asserting that no strength reduction was run.

    "No strength reduction was run, so this analysis reports no factor of safety"
    was printed on every finite element solution that had no factor of safety —
    including the reinforcement sample, which ships no meta sidecar at all and
    arrives as "loaded". Nothing on disk said what had been run; the sentence said
    it anyway.
    """
    fails = []
    from xslope.report import FEM_SOLVE_KINDS

    _slope_data, bundle = _fem_bundle()
    said = {}
    for analysis in ("single", "loaded", None):
        report = _engine_report("fem", bundle=dict(bundle, analysis=analysis,
                                                   FS=None))
        said[analysis] = " ".join(_prose(report))

    if "No strength reduction was run" not in said["single"]:
        fails.append(f"a recorded single trial no longer says a strength "
                     f"reduction was not run: {said['single']!r}")
    for unrecorded in ("loaded", None):
        text = said[unrecorded]
        if "No strength reduction was run" in text:
            fails.append(f"a run recorded as {unrecorded!r} — which is no record "
                         f"at all — is stated to have had no strength reduction: "
                         f"{text!r}")
        if "does not record which analysis produced it" not in text:
            fails.append(f"a run recorded as {unrecorded!r} does not say its "
                         f"analysis is unrecorded: {text!r}")
        if "factor of safety" not in text:
            fails.append(f"a run recorded as {unrecorded!r} says nothing about a "
                         f"factor of safety at all: {text!r}")

    # And it really does arrive that way off disk: a solved model whose node and
    # element companions have no meta beside them — which is what the sample this
    # was found on used to be, and what any run saved before the metadata existed
    # still is.
    tmp = tempfile.mkdtemp(prefix="xslope_unrecorded_")
    stem = _sidecar_copy(os.path.splitext(FEM_XLSX)[0], tmp,
                         drop=("_fem_meta.json",))
    _sd, restored = _restored(f"{stem}.xlsx")
    got = str((restored.get("fem") or {}).get("analysis"))
    if got != "loaded":
        fails.append(f"a solution with no meta beside it restores as {got!r}, "
                     f"so the unrecorded case is not reached from a file")
    if "loaded" in FEM_SOLVE_KINDS:
        fails.append("'loaded' — the word for no record — is listed as a kind of "
                     "solve that was recorded")
    return fails


def test_the_run_record_survives_the_file():
    """What a strength reduction run chose, and what its trials found, is written
    beside the solution and read back with it.

    solve_ssrm returns the criterion, the final interval, the trial record and the
    step count on its RESULT; the bundle carried only ``result['last_solution']``,
    the field, so every one of them was dropped before anything was saved. A
    reloaded run could then say what its factor of safety was and nothing about
    how it was reached — a zone-confined answer read exactly like a whole-section
    one.
    """
    fails = []
    import json
    from xslope.fem import (import_fem_meta, export_fem_solution,
                            ssrm_run_record)
    from xslope.report import _ssrm_procedure

    _slope_data, bundle = _fem_bundle()

    # The record, off a result shaped as solve_ssrm returns one.
    result = {"FS": 1.36, "converged": True,
              "last_solution": bundle["solution"],
              "failure_criterion": "hybrid",
              "method": "SSRM — Hybrid (non-convergence corroborated by "
                        "displacement evidence)",
              "final_interval": (1.34, 1.38), "interval_width": 0.04,
              "iterations_ssrm": 7,
              "trials": [{"F": 1.0, "role": "lower", "stable": True},
                         {"F": 1.8, "role": "upper", "stable": False}]}
    options = {"tolerance": 0.01, "F_min": 1.0, "F_max": 1.8,
               "ssr_exclude": ["Bedrock"]}
    fem_data = dict(bundle["fem_data"])
    fem_data["ssr_zones"] = [{"kind": "reduce", "polygon": [(0, 0), (1, 0), (1, 1)]},
                             {"kind": "hold", "polygon": [(0, 0), (1, 0), (1, 1)]}]
    record = ssrm_run_record(result, fem_data, options)
    for key, want in (("failure_criterion", "hybrid"),
                      ("final_interval", [1.34, 1.38]),
                      ("iterations_ssrm", 7), ("tolerance", 0.01),
                      ("ssr_exclude", ["Bedrock"]),
                      ("ssr_zones", {"reduce": 1, "hold": 1})):
        if record.get(key) != want:
            fails.append(f"the run record's {key!r} is {record.get(key)!r}, "
                         f"not {want!r}")
    if len(record.get("trials") or []) != 2:
        fails.append(f"the run record drops the trials: {record.get('trials')!r}")
    # A run that chose nothing is credited with nothing.
    if ssrm_run_record({}, {}, {}):
        fails.append(f"a run recording nothing yields a record anyway: "
                     f"{ssrm_run_record({}, {}, {})!r}")
    # An explicit search-area polygon is a reduce zone however it was expressed.
    if ssrm_run_record({}, {}, {"ssr_zone": [(0, 0), (1, 0), (1, 1)]}) \
            .get("ssr_zones") != {"reduce": 1}:
        fails.append("an ssr_zone polygon is not counted as a search area")

    # It is JSON, it survives the sidecar, and it comes back whole.
    tmp = tempfile.mkdtemp(prefix="xslope_runrec_")
    stem = os.path.join(tmp, "run")
    meta = dict(record)
    meta.update({"FS": 1.36, "analysis": "ssrm"})
    with contextlib.redirect_stdout(io.StringIO()):
        export_fem_solution(bundle["fem_data"], bundle["solution"], stem,
                            meta=meta)
    read = import_fem_meta(stem) or {}
    for key in record:
        if read.get(key) != json.loads(json.dumps(record[key])):
            fails.append(f"{key!r} does not survive the sidecar: wrote "
                         f"{record[key]!r}, read {read.get(key)!r}")
    if read.get("FS") != 1.36 or read.get("analysis") != "ssrm":
        fails.append("the run record displaced the facts the writer states "
                     "itself")

    # And the paragraph reads it: the tolerance, the interval and the confinement
    # are all stated from a bundle carrying nothing but this meta.
    text = _ssrm_procedure({"meta": read})
    for wanted in ("narrower than 0.01", "1.340 to 1.380",
                   "Strength was reduced only", "Bedrock"):
        if wanted not in text:
            fails.append(f"the paragraph does not state {wanted!r} from the "
                         f"restored record: {text!r}")
    # Each of them really comes from the record: drop one and the claim goes.
    for key, gone in (("tolerance", "narrower than 0.01"),
                      ("final_interval", "1.340 to 1.380"),
                      ("ssr_zones", "Strength was reduced only"),
                      ("ssr_exclude", "Bedrock")):
        without = _ssrm_procedure({"meta": {k: v for k, v in read.items()
                                            if k != key}})
        if gone in without:
            fails.append(f"a run recording no {key!r} states {gone!r} anyway")

    # The runner really emits it. solve_ssrm is stood in for — what is under test
    # is the line that builds the bundle from its result, which used to keep the
    # field and throw the rest away.
    try:
        from studio import runners
        from studio.main_window import MainWindow
    except Exception:
        return fails

    import xslope.fem as fem_module
    real_solve = fem_module.solve_ssrm
    fem_module.solve_ssrm = lambda fem_data, **kw: dict(
        result, last_solution=bundle["solution"], failure_solution=None)
    emitted = []
    try:
        runner = runners.FemRunner(dict(_slope_data), dict(options, analysis="ssrm"))
        runner.succeeded.connect(emitted.append)
        with contextlib.redirect_stdout(io.StringIO()):
            runner.run()
    finally:
        fem_module.solve_ssrm = real_solve
    if not emitted:
        fails.append("the finite element runner emitted no bundle for a run that "
                     "reached a factor of safety")
    else:
        got = (emitted[0].get("meta") or {})
        for key in ("failure_criterion", "final_interval", "trials",
                    "iterations_ssrm", "tolerance"):
            if key not in got:
                fails.append(f"the runner's bundle drops {key!r}: {sorted(got)}")
        if emitted[0].get("FS") != result["FS"]:
            fails.append("the runner's bundle no longer carries the factor of "
                         "safety")

    # The Studio writer lays these under the three facts it states itself, so
    # neither can silently overwrite the other.
    written = MainWindow._fem_meta({"FS": 1.36, "analysis": "ssrm",
                                    "solution": {"F": 1.34}, "meta": record})
    if written.get("failure_criterion") != "hybrid":
        fails.append("the Studio writer drops the run record from the sidecar")
    for key, want in (("FS", 1.36), ("analysis", "ssrm"), ("F", 1.34)):
        if written.get(key) != want:
            fails.append(f"the Studio writer's {key!r} is {written.get(key)!r}, "
                         f"not {want!r}")
    return fails


def test_fem_result_figures_carry_no_title():
    """The finite element result panels are captioned like every other figure in
    the report, and carry no title of their own.

    They were the only three that did: every other producer is passed
    show_title=False, and these were not. The consequence was the factor of
    safety printed three times on one page opening at two roundings — 1.345 in
    the paragraph, "FS = 1.35" on each of three panels. What the title said and
    the caption does not is the deformation exaggeration, which the paragraph
    now states, derived from the function the panel scales itself by.
    """
    fails = []
    from xslope.plot_fem import deformation_scale

    _slope_data, bundle = _fem_bundle()
    report = _engine_report("fem")

    # Read off the rendered axes, not off the call: whatever reaches the page.
    import matplotlib.figure as mplfig
    from xslope.plot_fem import plot_fem_results
    from xslope.report import FEM_PANELS
    titled = []
    for panel, _caption, _shows in FEM_PANELS:
        fig = mplfig.Figure(figsize=(4.0, 2.5))
        with contextlib.redirect_stdout(io.StringIO()):
            plot_fem_results(bundle["fem_data"], bundle["solution"],
                             plot_type=[panel], fig=fig,
                             fs=bundle.get("FS"),
                             failure_solution=bundle.get("failure_solution"),
                             show_title=False)
        titled += [ax.get_title() for ax in fig.axes if ax.get_title()]
    if titled:
        fails.append(f"a result panel drawn for the report carries a title: "
                     f"{titled}")

    # The report asks for them that way — the option, not just the default.
    seen = {}
    import xslope.plot_fem as pf
    real = pf.plot_fem_results

    def spy(fem_data, solution, **kw):
        seen.setdefault("show_title", kw.get("show_title"))
        return real(fem_data, solution, **kw)

    pf.plot_fem_results = spy
    try:
        _engine_report("fem", options={"fem_inputs_figure": False,
                                       "fem_mesh_figure": False})
    finally:
        pf.plot_fem_results = real
    if seen.get("show_title") is not False:
        fails.append(f"the report draws its result panels with show_title="
                     f"{seen.get('show_title')!r}")

    # The factor of safety is stated once, in the paragraph, and nowhere else.
    fs = bundle["FS"]
    for rounding in (f"{fs:.2f}", f"{fs:.3f}"):
        hits = [b.text for b in report.blocks("prose") if rounding in b.text]
        hits += [f.caption for f in report.figures() if rounding in f.caption]
        if len(hits) > 1:
            fails.append(f"the factor of safety appears as {rounding} in "
                         f"{len(hits)} places: {hits}")

    # And the one thing the title uniquely carried survives, at the value the
    # panel is drawn at. Read at the LAST CONVERGED state: at failure the section
    # has already moved far enough to be drawn life-size, and a panel drawn at
    # 1.0x exercises no exaggeration sentence.
    converged = _engine_report("fem", options={"fem_state_failure": False,
                                               "fem_state_converged": True})
    scale = deformation_scale(bundle["fem_data"], bundle["solution"])
    if scale <= 1.0:
        fails.append(f"the fixture is drawn at {scale}x, so the exaggeration "
                     f"sentence proves nothing")
    else:
        shown = f"{scale:.0f}" if scale >= 10 else f"{scale:.1f}"
        said = " ".join(_prose(converged))
        if f"{shown} times the computed displacement" not in said:
            fails.append(f"the deformed grid is drawn at {shown}x and the "
                         f"report does not say so: {said!r}")
    return fails


def _fem_with_failure_state(bundle):
    """The bundle with an at-failure snapshot on it, for a model that ships none.

    The snapshot is the converged field itself, which is what makes it useful
    here: whatever the report then says about WHICH state it read cannot have come
    from the numbers, only from the selection.
    """
    return dict(bundle, failure_solution=dict(bundle["solution"]))


def test_the_field_state_toggles():
    """The result panels are drawn at the field state (or states) asked for, and
    the report says which.

    A strength reduction run leaves two fields — the mechanism it developed past
    its critical factor, and the last trial that reached equilibrium — and the
    report drew one of them, whichever was there, with no way to ask for the
    other. Each is a box now: at failure on by default, last converged off, both
    together giving both sets of panels.
    """
    fails = []
    from xslope.report import DEFAULT_OPTIONS, FEM_PANELS

    xlsx = RS2_28A_XLSX
    slope_data, solutions = _restored(xlsx)
    bundle = solutions["fem"]
    if bundle.get("failure_solution") is None:
        fails.append("the model this is checked on carries no at-failure "
                     "snapshot, so the selection cannot be exercised")
        return fails

    if DEFAULT_OPTIONS["fem_state_failure"] is not True:
        fails.append("the at-failure state is off by default")
    if DEFAULT_OPTIONS["fem_state_converged"] is not False:
        fails.append("the last-converged state is on by default, so every report "
                     "carries both sets of panels")

    def built(extra):
        opts = {"input_path": xlsx, "lem": False, "pd_figure": False}
        opts.update(FAST_FIGURES)
        opts.update(extra)
        tmp = tempfile.mkdtemp(prefix="xslope_state_")
        with contextlib.redirect_stdout(io.StringIO()):
            report = build_report_(slope_data, {"fem": bundle}, opts, tmp)
            planned = planned_figures_(slope_data, {"fem": bundle},
                                       resolve_options_(opts))
        return report, planned

    from xslope.report import (build_report as build_report_,
                               planned_figures as planned_figures_,
                               resolve_options as resolve_options_)

    n = len(FEM_PANELS)
    cases = [("the default", {}, ["fem run1 shear_strain"], n),
             ("last converged alone", {"fem_state_failure": False,
                                       "fem_state_converged": True},
              ["fem run1 shear_strain"], n),
             ("both", {"fem_state_converged": True},
              ["fem run1 shear_strain failure",
               "fem run1 shear_strain converged"], 2 * n),
             ("neither", {"fem_state_failure": False}, [], 0)]
    for label, extra, wanted, panels in cases:
        report, planned = built(extra)
        sources = [f.source for f in report.figures()]
        panel_figures = [s for s in sources if s not in ("fem model", "fem mesh")]
        if len(panel_figures) != panels:
            fails.append(f"{label}: {len(panel_figures)} result panels, expected "
                         f"{panels}: {panel_figures}")
        if planned != len(report.figures()):
            fails.append(f"{label}: planned {planned} figures and built "
                         f"{len(report.figures())}")
        for source in wanted:
            if source not in sources:
                fails.append(f"{label}: no {source!r} figure: {sources}")

        said = " ".join(_prose(report))
        captions = [f.caption for f in report.figures()]
        if label == "both":
            # Each state is NAMED, on the figures and in the sentence: two sets of
            # identical captions would be six figures of one thing.
            for state in ("at failure", "last converged"):
                if not any(state in c for c in captions):
                    fails.append(f"both: no figure caption names the {state!r} "
                                 f"state: {captions}")
                if state not in said:
                    fails.append(f"both: the paragraph does not name the "
                                 f"{state!r} state: {said!r}")
            if "drawn twice" not in said:
                fails.append(f"both: the paragraph does not say the fields are "
                             f"drawn at two states: {said!r}")
            if len(set(captions)) != len(captions):
                fails.append(f"both: two figures share a caption: {captions}")
            # Each state's figures are named in a sentence of their own, led by
            # the state. Named in one sentence with the state hung off the end of
            # what each draws, six figures read "where the section is shearing at
            # failure … how the section is moving last converged".
            for lead in ("At failure,", "At the last converged trial,"):
                if lead not in said:
                    fails.append(f"both: no sentence opens {lead!r}: {said!r}")
            for run_on in ("shearing at failure", "moving last converged",
                           "section last converged"):
                if run_on in said:
                    fails.append(f"both: the state is hung off what a figure "
                                 f"draws — {run_on!r}: {said!r}")
            # Each fact once: what a panel draws is described in the FIRST
            # state's sentence, and the second state's sentence cites its
            # figures without re-describing them — the two sentences repeated
            # the three descriptions verbatim in one paragraph.
            for _p, _c, shows in FEM_PANELS:
                if said.count(shows) != 1:
                    fails.append(f"both: {shows!r} is said "
                                 f"{said.count(shows)} times in one paragraph")
        elif panels:
            for state in ("at failure", "last converged"):
                if any(state in c for c in captions):
                    fails.append(f"{label}: a caption qualifies the state on the "
                                 f"only set of panels there is: {captions}")
        if label == "the default" and "mechanism at failure" not in said:
            fails.append(f"the default does not say it draws the mechanism at "
                         f"failure: {said!r}")
        if label == "last converged alone":
            if "last trial that reached equilibrium" not in said:
                fails.append(f"{label}: the paragraph does not say which field "
                             f"was drawn: {said!r}")
            if "mechanism at failure" in said:
                fails.append(f"{label}: the paragraph claims the mechanism at "
                             f"failure was drawn: {said!r}")

    # A state asked for that the run never captured: the converged field is drawn
    # and the sentence says the snapshot is missing, rather than the converged
    # field being passed off as the mechanism.
    # Built by copying a solved model WITHOUT its at-failure companions — which
    # is what every run saved before the snapshot was captured looks like on
    # disk, and what the fallback is for. It used to be a corpus model that
    # simply had none; they carry one now.
    no_snap = tempfile.mkdtemp(prefix="xslope_nosnap_")
    no_snap_stem = _sidecar_copy(
        os.path.splitext(FEM_XLSX)[0], no_snap,
        drop=("_fem_failure_nodes.csv", "_fem_failure_elements.csv",
              "_fem_failure_meta.json"))
    load_sd, load_solutions = _restored(f"{no_snap_stem}.xlsx")
    load_bundle = load_solutions["fem"]
    if load_bundle.get("failure_solution") is not None:
        fails.append("the fallback is checked on a model that DOES carry a "
                     "snapshot, so it proves nothing")
    for extra in ({}, {"fem_state_converged": True}):
        opts = {"input_path": f"{no_snap_stem}.xlsx", "lem": False,
                "pd_figure": False}
        opts.update(FAST_FIGURES)
        opts.update(extra)
        tmp = tempfile.mkdtemp(prefix="xslope_fallback_")
        with contextlib.redirect_stdout(io.StringIO()):
            report = build_report_(load_sd, {"fem": load_bundle}, opts, tmp)
            planned = planned_figures_(load_sd, {"fem": load_bundle},
                                       resolve_options_(opts))
        said = " ".join(_prose(report))
        sources = [f.source for f in report.figures()]
        # The field panels alone: the model, the mesh and the search figure are
        # drawn once whatever state is asked for, and it is the panels that count
        # the states.
        panel_figures = [s for s in sources
                         if s not in ("fem model", "fem mesh", "fem run1 search")]
        if len(panel_figures) != n:
            fails.append(f"a run with no snapshot drew {len(panel_figures)} "
                         f"panels for {extra}, not the {n} of one state: "
                         f"{panel_figures}")
        if planned != len(report.figures()):
            fails.append(f"the fallback planned {planned} figures and built "
                         f"{len(report.figures())}")
        if "No at-failure snapshot was captured" not in said:
            fails.append(f"a run whose snapshot is missing does not say so: "
                         f"{said!r}")
        if "mechanism at failure — the trial" in said:
            fails.append(f"a run with no snapshot is said to draw the mechanism "
                         f"at failure: {said!r}")
    return fails


def test_each_state_is_drawn_at_its_own_scale():
    """Drawn at two states, each state's panels are scaled to its OWN field.

    A strength reduction run drawn at both states carries a mechanism whose
    strains and displacements are orders above the standing section's. Pinned to
    the pair, the last converged trial's strain panel is one flat colour under a
    colorbar it never reaches, its deformed grid is drawn at the collapse's
    exaggeration and shows no deformation at all, and its arrows are dust — three
    figures a reader gets nothing from (the owner's reading of
    fem_griffiths1_load, Figures 7-9). So the colour range, the exaggeration and
    the arrow scale are each state's own, and the paragraph states the
    exaggeration each grid was drawn at.

    The measurement is on the CALL the plotter was made with and on the artists
    that call produced: the converged state's strain panel must resolve its own
    field, which on this model is a small fraction of the failure field's.
    """
    fails = []
    import numpy as np
    import xslope.plot_fem as pf
    from xslope.fem_details import field_solution
    from xslope.plot_fem import deformation_scale, shear_strain_field
    from xslope.report import FEM_PANELS, build_report

    xlsx = FEM_XLSX          # griffiths1_load: the model the owner read
    slope_data, solutions = _restored(xlsx)
    bundle = solutions["fem"]

    real = pf.plot_fem_results
    calls = []

    def spy(fem_data, solution, **kw):
        calls.append(kw)
        return real(fem_data, solution, **kw)

    def draw(extra):
        calls.clear()
        opts = {"input_path": xlsx, "lem": False, "pd_figure": False,
                "fem_inputs_figure": False, "fem_mesh_figure": False}
        opts.update(FAST_FIGURES)
        opts.update(extra)
        tmp = tempfile.mkdtemp(prefix="xslope_scale_")
        pf.plot_fem_results = spy
        try:
            with contextlib.redirect_stdout(io.StringIO()):
                report = build_report(slope_data, {"fem": bundle}, opts, tmp)
        finally:
            pf.plot_fem_results = real
        return list(calls), report

    both, report = draw({"fem_state_converged": True, "fem_figure": True})
    if len(both) != 2 * len(FEM_PANELS):
        fails.append(f"both states drew {len(both)} panels, not "
                     f"{2 * len(FEM_PANELS)}")
    states = [kw.get("field_state") for kw in both]
    if set(states) != {"failure", "converged"}:
        fails.append(f"the panels were not drawn at the two states: {states}")

    # Nothing is pinned ACROSS the pair: each panel takes its own field's range
    # and its own arrow scale.
    for key in ("vmin", "vmax", "vector_max"):
        pinned = [kw.get(key) for kw in both if kw.get(key) is not None]
        if pinned:
            fails.append(f"{key} was pinned to {pinned}; each state scales "
                         f"itself")

    # …and the exaggerations really are two, one per state, each the one that
    # state's own field asks for.
    fields = {s: field_solution(bundle["solution"], s,
                                failure_solution=bundle.get("failure_solution"))
              for s in ("failure", "converged")}
    own = {s: deformation_scale(bundle["fem_data"], f)
           for s, f in fields.items()}
    if abs(own["failure"] - own["converged"]) < 1e-9:
        fails.append("the two states ask for the same exaggeration, so drawing "
                     "them apart proves nothing")
    for kw in both:
        want = own[kw["field_state"]]
        if kw.get("deform_scale") != want:
            fails.append(f"the {kw['field_state']} panels were drawn at "
                         f"exaggeration {kw.get('deform_scale')}, not their own "
                         f"field's {want}")

    # The paragraph states BOTH multipliers, so a reader of either grid knows
    # what it is drawn at.
    said = " ".join(_prose(report))
    stated = 0
    for state, scale in own.items():
        # A grid drawn at true scale is not exaggerated and says nothing; one
        # that is has to name its multiplier.
        if scale <= 1.0:
            continue
        stated += 1
        shown = f"{scale:.0f}" if scale >= 10 else f"{scale:.1f}"
        if f"drawn at {shown} times" not in said:
            fails.append(f"the {state} exaggeration ({shown}) is not stated")
    if not stated:
        fails.append("neither state is exaggerated, so the per-state multiplier "
                     "is untested on this fixture")

    # The consequence, measured on the drawn artists: the converged strain panel
    # resolves its OWN field. Pinned to the pair it spanned the failure field,
    # which on this model is orders larger — one flat colour.
    import matplotlib.figure as mplfig

    def strain_top(state, **kw):
        fig = mplfig.Figure(figsize=(4.0, 3.0))
        with contextlib.redirect_stdout(io.StringIO()):
            pf.plot_fem_results(bundle["fem_data"], bundle["solution"],
                                plot_type=["shear_strain"], fig=fig,
                                show_title=False, field_state=state,
                                failure_solution=bundle.get("failure_solution"),
                                **kw)
        tops = [float(c.get_clim()[1]) for ax in fig.axes
                for c in ax.collections if c.get_array() is not None]
        return max(tops) if tops else None

    peaks = {}
    for state, field in fields.items():
        values = shear_strain_field(bundle["fem_data"], field)
        peaks[state] = (float(np.nanmax(values)) if values is not None else None)
    if None in peaks.values():
        fails.append("the strain field could not be read; the colour range "
                     "check has no oracle")
    elif not peaks["failure"] > 5.0 * peaks["converged"]:
        fails.append(f"the two states' strains are of a size ({peaks}); a "
                     f"pinned colorbar would have been readable anyway")
    else:
        own_top = strain_top("converged")
        pinned_top = strain_top("converged", vmin=0.0, vmax=peaks["failure"])
        if own_top is None or pinned_top is None:
            fails.append("the strain panel drew nothing to measure")
        elif not own_top < 0.5 * pinned_top:
            fails.append(f"the converged strain panel is drawn to {own_top}, no "
                         f"smaller than the pair-pinned {pinned_top} — its own "
                         f"structure is still not resolved")

    # And a run drawn at ONE state is unchanged: its panels scale themselves and
    # its grid is drawn at its own field's exaggeration.
    alone, _report = draw({"fem_figure": True})
    for kw in alone:
        for key in ("vmin", "vmax", "vector_max"):
            if kw.get(key) is not None:
                fails.append(f"a run drawn at one state was pinned to a "
                             f"{key} of {kw.get(key)!r}")
    if any(kw.get("deform_scale") is None for kw in alone):
        fails.append("the exaggeration the paragraph states is not the one the "
                     "panel was drawn at")
    return fails


def test_the_member_forces_follow_the_state_the_panels_are_drawn_at():
    """The member forces are read at the state the section is drawn at, and the
    subsection says which — once.

    The members were always read at the at-failure state, whatever the panels
    above them showed. A run drawn at both states does not get two readings of
    the same bars: the subsection follows the PRIMARY state — at failure when
    that was asked for and captured, the converged field otherwise.
    """
    fails = []
    from xslope.report import _fem_primary_state, resolve_options

    slope_data, plain = _fem_1d_bundle(FEM_REINF_XLSX)
    bundle = _fem_with_failure_state(plain)

    for label, extra, state, words in (
            ("the default", {}, "failure", "developed mechanism at failure"),
            ("both", {"fem_state_converged": True}, "failure",
             "developed mechanism at failure"),
            ("last converged alone", {"fem_state_failure": False,
                                      "fem_state_converged": True},
             "converged", "last converged field")):
        opts = resolve_options(extra)
        got = _fem_primary_state(bundle, opts)
        if got != state:
            fails.append(f"{label}: the members are read at {got!r}, not "
                         f"{state!r}")
        report = _engine_report("fem", options=extra, bundle=bundle,
                               xlsx=FEM_REINF_XLSX)
        sec = _member_section(report, "Reinforcement Forces")
        if sec is None:
            fails.append(f"{label}: the run carries no reinforcement subsection")
            continue
        # The reinforcement carries no summary table (the owner's ruling): its
        # capacities are in the properties table of the inputs and its peaks are
        # annotated on the figure of the line itself.
        tables = [b for b in sec.blocks if b.kind == "table"]
        if tables:
            fails.append(f"{label}: the reinforcement subsection printed "
                         f"{len(tables)} table(s) of member forces")
        said = " ".join(b.text for b in sec.blocks if b.kind == "prose")
        if said.count(words) != 1:
            fails.append(f"{label}: the state the forces were read at is said "
                         f"{said.count(words)} times, not once: {said!r}")
        other = ("last converged field" if state == "failure"
                 else "developed mechanism at failure")
        if other in said:
            fails.append(f"{label}: the sentence names both states: {said!r}")

    # A run that captured no snapshot reads the converged field however the boxes
    # are set — there is only one field to read.
    if _fem_primary_state(plain, resolve_options({})) != "converged":
        fails.append("a run with no snapshot claims to read its members at a "
                     "state it never captured")
    return fails


def test_the_search_figure_draws_the_trials():
    """The search that reached the factor of safety is drawn from the trials the
    run recorded.

    plot_ssrm_convergence read ``F_history`` and ``convergence_history``, which
    solve_ssrm has never emitted under any criterion: it printed "No SSRM
    convergence history found" and returned nothing, on every run there has ever
    been. It reads the ``trials`` the solver does record.
    """
    fails = []
    import inspect
    import matplotlib.figure as mplfig
    from xslope.fem import _ssrm_displacement_limit
    from xslope.plot_fem import (plot_ssrm_convergence, ssrm_has_convergence_history,
                                 ssrm_interval_history, ssrm_trials)

    # The keys it reads are the keys the solver writes.
    recorder = inspect.getsource(_ssrm_displacement_limit)
    for key in ('"F"', '"role"', '"stable"'):
        if key not in recorder:
            fails.append(f"the solver no longer records {key} on a trial, so the "
                         f"figure is reading something else")
    # The keys the function itself uses, not the docstring — which says what it
    # used to read, and why.
    used = [c for c in plot_ssrm_convergence.__code__.co_consts
            if isinstance(c, str) and c is not plot_ssrm_convergence.__doc__]
    for gone in ("F_history", "convergence_history"):
        if gone in used:
            fails.append(f"the figure still reads {gone!r}, which the solver "
                         f"does not emit")
    if "trials" not in [c for c in ssrm_trials.__code__.co_consts
                        if isinstance(c, str)]:
        fails.append("the figure's trial reader no longer reads 'trials'")

    record = {"FS": 1.36, "tolerance": 0.01,
              "trials": [{"F": 1.0, "role": "lower", "stable": True},
                         {"F": 1.8, "role": "upper", "stable": False},
                         {"F": 1.4, "role": "bisect", "stable": False},
                         {"F": 1.2, "role": "bisect", "stable": True},
                         {"F": 1.3, "role": "bisect", "stable": True},
                         {"F": 1.35, "role": "bisect", "stable": True}]}
    if not ssrm_has_convergence_history(record):
        fails.append("a run with six recorded trials has no history to draw")
    for bare in ({}, {"trials": []}, {"trials": [{"F": 1.0}]},
                 {"trials": [{"role": "lower"}, {"role": "upper"}]}):
        if ssrm_has_convergence_history(bare):
            fails.append(f"a run recording {bare!r} is said to have a search to "
                         f"draw")

    # The interval really narrows, and it is the solver's own rule that moves it.
    history = ssrm_interval_history(record)
    widths = [hi - lo for lo, hi in history if lo is not None and hi is not None]
    if widths != sorted(widths, reverse=True):
        fails.append(f"the interval drawn does not narrow: {history}")
    if history[-1] != (1.35, 1.4):
        fails.append(f"the interval after the last trial is {history[-1]}, not "
                     f"the (1.35, 1.4) the recorded verdicts leave")

    # It draws: one axes, every trial on it, and the factor of safety ruled.
    fig = mplfig.Figure(figsize=(6.0, 4.0))
    plot_ssrm_convergence(record, fig=fig, show_title=False)
    if len(fig.axes) != 1:
        fails.append(f"the search figure draws {len(fig.axes)} axes, not the one "
                     f"bounded figure it is meant to be")
    else:
        ax = fig.axes[0]
        plotted = sum(len(line.get_xdata()) for line in ax.lines
                      if line.get_linestyle() == "None")
        if plotted != len(record["trials"]):
            fails.append(f"{plotted} trials are marked, of "
                         f"{len(record['trials'])} recorded")
        labels = [t.get_text() for t in (ax.get_legend().get_texts()
                                         if ax.get_legend() else [])]
        for wanted in ("stood", "did not stand", "1.360"):
            if not any(wanted in l for l in labels):
                fails.append(f"the figure's legend does not name {wanted!r}: "
                             f"{labels}")
    # A run with nothing to draw refuses rather than drawing an empty axes.
    try:
        plot_ssrm_convergence({"trials": []}, fig=mplfig.Figure())
        fails.append("a run recording no trials still produced a figure")
    except ValueError:
        pass

    # And it reaches the report, under its own sentence, with its own switch.
    slope_data, bundle = _fem_bundle()
    carried = dict(bundle, meta=dict(bundle.get("meta") or {}, **record))
    report = _engine_report("fem", bundle=carried)
    sources = [f.source for f in report.figures()]
    if "fem run1 search" not in sources:
        fails.append(f"the search figure does not reach the report: {sources}")
    said = " ".join(_prose(report))
    if f"each of the {len(record['trials'])} trials at its own" not in said:
        fails.append(f"the report does not say how many trials the search took: "
                     f"{said!r}")
    planned, drawn = _planned_matches(report, "fem", bundle=carried)
    if planned != drawn:
        fails.append(f"a report carrying the search planned {planned} figures "
                     f"and built {drawn}")
    off = _engine_report("fem", options={"fem_convergence_figure": False},
                         bundle=carried)
    if "fem run1 search" in [f.source for f in off.figures()]:
        fails.append("the search figure cannot be switched off")

    # A run that kept no trials gets neither the figure nor a sentence about it.
    _sd, forgot = _fem_run_forgetting({"trials"})
    plain = _engine_report("fem", bundle=forgot)
    if "fem run1 search" in [f.source for f in plain.figures()]:
        fails.append("a run recording no trials drew a search figure anyway")
    if "the search that reached it" in " ".join(_prose(plain)):
        fails.append("a run recording no trials is given a sentence about a "
                     "search figure it has not got")
    return fails


def test_which_result_panels_draw_a_legend():
    """The deformation panel names its two grids in a legend; the strain and
    vector panels draw none.

    The plots stay as they are (Norm's ruling). What is checked here is that the
    Studio panel beside them tells the truth about them: the note where the legend
    controls would go said "the single-panel FEM result plots draw no legend",
    and the deformation panel has always drawn one — original grid, deformed grid,
    and the reinforcement in both configurations where the model carries any.
    """
    fails = []
    import inspect
    import matplotlib.figure as mplfig
    from xslope.plot_fem import plot_fem_results
    from studio import display_panels

    _slope_data, bundle = _fem_bundle()
    drawn = {}
    for panel in ("deformation", "shear_strain", "displace_vector"):
        fig = mplfig.Figure(figsize=(4.0, 3.0))
        with contextlib.redirect_stdout(io.StringIO()):
            plot_fem_results(bundle["fem_data"], bundle["solution"],
                             plot_type=[panel], fig=fig, show_title=False)
        legends = [text.get_text()
                   for ax in fig.axes if ax.get_legend() is not None
                   for text in ax.get_legend().get_texts()]
        drawn[panel] = legends
    if not drawn["deformation"]:
        fails.append("the deformation panel draws no legend, so the note beside "
                     "the controls is describing a plot that no longer exists")
    for panel in ("shear_strain", "displace_vector"):
        if drawn[panel]:
            fails.append(f"the {panel!r} panel draws a legend {drawn[panel]}, "
                         f"which the note beside the controls says it does not")

    # The note itself: it must not assert the absence the panels contradict.
    source = inspect.getsource(display_panels.FemResultsDisplayPanel)
    notes = [line.strip() for line in source.splitlines()
             if line.strip().startswith("#") and "legend" in line.lower()]
    if not notes:
        fails.append("the results panel carries no note about the legends at all")
    for note in notes:
        if "draw no legend" in note or "draws no legend" in note:
            fails.append(f"the note still says the result plots draw no legend: "
                         f"{note!r}")
    return fails


def test_one_name_for_the_shear_strain_field():
    """One field, one name — in the caption, in the sentence citing it, and on
    the colorbar.

    It had four: the caption said "Maximum shear strain", which is the name of
    the OTHER strain column (the total one, ``max_shear_strain``); the plotted
    array is ``vp_shear_strain``; the colorbar said "VP Max Shear Strain". The
    documentation calls it the viscoplastic shear strain, so that is the name.
    """
    fails = []
    from xslope.plot_fem import SHEAR_STRAIN_LABEL
    from xslope.report import FEM_PANELS

    name = SHEAR_STRAIN_LABEL
    caption = next((c for p, c, _s in FEM_PANELS if p == "shear_strain"), None)
    shows = next((s for p, _c, s in FEM_PANELS if p == "shear_strain"), "")
    if caption != name:
        fails.append(f"the panel caption is {caption!r} and the plotting label "
                     f"is {name!r}")
    if name.lower() not in shows.lower():
        fails.append(f"the sentence citing the figure calls it {shows!r}, which "
                     f"does not use {name!r}")

    # The colorbar, off the rendered figure — BOTH ways the panel draws one.
    # plot_fem_results places the bar itself only where every requested panel is
    # one whose internal colorbar it can defer; mix in a panel it cannot (any
    # diagnostic field) and plot_shear_strain_contours draws its own inline,
    # labelled from its own call to _plot_nodal_contours. That second label was
    # left out of this check, and reverting it alone to "VP Max Shear Strain"
    # passed clean: the report and the results view both take the deferred path.
    import matplotlib.figure as mplfig
    from xslope.plot_fem import plot_fem_results
    _slope_data, bundle = _fem_bundle()
    for panels, drawn_by in ((["shear_strain"], "the deferred colorbar"),
                             (["shear_strain", "stress"], "the panel's own "
                              "inline colorbar")):
        fig = mplfig.Figure(figsize=(4.0, 3.0 * len(panels)))
        with contextlib.redirect_stdout(io.StringIO()):
            plot_fem_results(bundle["fem_data"], bundle["solution"],
                             plot_type=panels, fig=fig, show_title=False)
        labels = [ax.get_ylabel() or ax.get_xlabel() for ax in fig.axes]
        labels = [l for l in labels if l]
        if not labels:
            fails.append(f"the shear strain panel draws no labelled colorbar "
                         f"as {drawn_by}")
        elif not any(name.lower() in l.lower() for l in labels):
            fails.append(f"{drawn_by} is labelled {labels}, and the field is "
                         f"called {name!r} everywhere else")
        for l in labels:
            if "VP Max" in l:
                fails.append(f"{drawn_by} still reads {l!r}")

    # The report's own figure caption, on the page.
    report = _engine_report("fem")
    captions = [f.caption for f in report.figures() if "shear" in f.caption.lower()]
    if not captions:
        fails.append("no shear strain figure reached the report")
    for c in captions:
        if not c.startswith(name):
            fails.append(f"the figure is captioned {c!r}, not {name!r}")

    # Studio's own two names for it: the results view's plot-type list, and the
    # report dialog row describing the figure it switches on. The view called it
    # "Shear strain" and the dialog row called it "the maximum shear strain",
    # which is the name of the OTHER strain field. Nothing pinned either. The
    # whole list — keys, order and names — is pinned by
    # test_fem_panels_mirror_the_fem_view; this is the one name, at the colorbar.
    from studio.display_panels import FEM_PLOT_TYPES
    from studio.report_dialog import CONTENT_TREE

    offered = dict(FEM_PLOT_TYPES).get("shear_strain")
    if offered != name:
        fails.append(f"the results view offers the field as {offered!r} and the "
                     f"colorbar beside it reads {name!r}")

    rows = [(key, label, tip)
            for _k, _l, _t, children in CONTENT_TREE for key, label, tip in children]
    described = next((tip for key, _l, tip in rows if key == "fem_figure"), None)
    if described is None:
        fails.append("the dialog offers no row for the finite element figure")
    else:
        if name.lower() not in described.lower():
            fails.append(f"the dialog row describes the figure as {described!r}, "
                         f"which does not call the field {name!r}")
        if "maximum shear strain" in described.lower():
            fails.append(f"the dialog row calls it the maximum shear strain, "
                         f"which is the other strain field: {described!r}")
    return fails


def test_fem_panels_mirror_the_fem_view():
    """The report presents every plot the finite element results view offers, in
    the same order and under the same names, and they are figures the reader can
    switch off.

    The counterpart of :func:`test_seep_panels_mirror_the_seep_view` for the other
    engine. The two lists carried the same three plots in DIFFERENT ORDERS — the
    report led with the deformed mesh, the view leads with the strain field — and
    called the deformation plot two things ("Deformation" on the screen,
    "Deformed mesh" in the report), so a reader moving between the screen and the
    page was matching neither position nor name.
    """
    fails = []
    from studio.display_panels import FEM_PLOT_TYPES
    from studio.report_dialog import CONTENT_TREE
    from xslope.report import DEFAULT_OPTIONS, FEM_PANELS

    def disagree(printed, offered, what):
        """What is wrong with a report list against a view list — the comparison
        itself, so it can be run on lists that ARE wrong."""
        if list(printed) == list(offered):
            return []
        missing = [v for v in offered if v not in printed]
        extra = [v for v in printed if v not in offered]
        if missing or extra:
            return [f"the report prints {what} {printed} and the finite element "
                    f"results view offers {offered}: {missing} is on the view and "
                    f"not in the report, {extra} the other way about"]
        return [f"the report prints {what} {printed} in a different order from "
                f"the view's {offered}"]

    offered = [key for key, _label in FEM_PLOT_TYPES]
    printed = [panel for panel, _c, _s in FEM_PANELS]
    fails += disagree(printed, offered, "the plots")

    # And by NAME, not by key alone — compared case-insensitively, because the two
    # write for different places (a control label and a figure caption) and the
    # WORDS are what a reader matches on.
    named = [label.lower() for _key, label in FEM_PLOT_TYPES]
    captions = [caption.lower() for _p, caption, _s in FEM_PANELS]
    fails += disagree(captions, named, "the captions")

    # The comparison really discriminates: a panel dropped from either side, the
    # same three in a different order, and a plot renamed on one side, are each
    # caught. A parity check that compares two lists it has already agreed on
    # proves nothing.
    for mutant, what in (
            (disagree(printed[:-1], offered, "x"), "a panel dropped from the "
             "report"),
            (disagree(printed, offered[:-1], "x"), "a plot dropped from the view"),
            (disagree(list(reversed(printed)), offered, "x"), "the panels "
             "reordered"),
            (disagree(captions[:-1] + ["deformation"], named, "x"),
             "a plot renamed on one side")):
        if not mutant:
            fails.append(f"the parity comparison passes {what}")

    # Every panel says what it draws, in its own words, and no two of them are the
    # same figure: a caption or a sentence repeated is one figure printed twice as
    # far as the report is concerned.
    for field, index in (("caption", 1), ("the sentence", 2)):
        names = [entry[index] for entry in FEM_PANELS]
        if len(set(names)) != len(names):
            fails.append(f"two finite element panels share {field}: {names}")
    for panel, caption, shows in FEM_PANELS:
        if not shows:
            fails.append(f"the {panel!r} panel has no sentence to introduce it, "
                         f"so its figure stands unexplained")

    # They are drawn under one switch, on by default, under the finite element
    # branch: the three are one reading of the solve, and the results view offers
    # them together.
    rows = {}
    for key, _label, _tip, children in CONTENT_TREE:
        for child_key, _l, _t in children:
            rows[child_key] = key
    if DEFAULT_OPTIONS.get("fem_figure") is not True:
        fails.append("'fem_figure' is off by default, so the finite element "
                     "fields are not in a report unless they are asked for")
    if rows.get("fem_figure") != "fem":
        fails.append(f"the 'fem_figure' row is a child of "
                     f"{rows.get('fem_figure')!r}, not of the finite element "
                     f"section that prints it")

    # The row's TITLE names the three plots too, and in the order they are
    # printed. It read "Deformation, shear strain and displacement plots" — the
    # old name and the old order — after the panels themselves were pinned, so
    # the title is held to the same list. The word each panel goes by in a row
    # title covers exactly the panels printed: a panel added or renamed fails
    # here until the title carries it.
    labels = {child_key: label for _k, _l, _t, children in CONTENT_TREE
              for child_key, label, _t2 in children}
    row_title = labels.get("fem_figure") or ""
    words = {"shear_strain": "shear strain", "deformation": "deformed mesh",
             "displace_vector": "displacement"}
    if set(words) != set(printed):
        fails.append(f"the row-title words cover {sorted(words)}, not the "
                     f"panels printed: {printed}")
    positions = [row_title.lower().find(words[p]) for p in printed
                 if p in words]
    if -1 in positions:
        fails.append(f"the 'fem_figure' row title {row_title!r} does not name "
                     f"every plot it switches on: "
                     f"{[words[p] for p in printed if p in words]}")
    elif positions != sorted(positions):
        fails.append(f"the 'fem_figure' row title {row_title!r} names the "
                     f"plots out of the printed order {printed}")

    # And the plot types are ones plot_fem_results actually draws: a name on
    # either list that the plotter rejects is a control that cannot be used.
    import matplotlib.figure as mplfig
    from xslope.plot_fem import plot_fem_results
    _slope_data, bundle = _fem_bundle()
    for panel in offered:
        fig = mplfig.Figure(figsize=(3.0, 2.0))
        try:
            with contextlib.redirect_stdout(io.StringIO()):
                plot_fem_results(bundle["fem_data"], bundle["solution"],
                                 plot_type=[panel], fig=fig, show_title=False)
        except Exception as exc:
            fails.append(f"the {panel!r} plot both lists offer cannot be drawn: "
                         f"{exc!r}")
    return fails


def test_the_model_figure_names_only_the_members_it_draws():
    """The finite element model figure's sentence names the reinforcement lines
    and the piles only where the model carries them.

    johnson_res has neither a reinforcement line nor a pile, and its sentence
    credited the figure with members anyway. The members are named for what they
    are — "the reinforcement lines", "the piles" — as the limit equilibrium model
    sentence names them, rather than lumped under "the members the solution
    carries".
    """
    fails = []
    from xslope.fileio import load_slope_data

    def sentence(report):
        for section in report.sections:
            for _lvl, sec in section.walk():
                if sec.title != "Analysis Inputs":
                    continue
                for b in sec.blocks:
                    if b.kind == "prose" and "specific to the finite element " \
                            "model are shown in" in b.text:
                        return b.text
        return ""

    bare = load_slope_data_cached(FEM_XLSX)
    if bare.get("reinforcement_lines") or bare.get("pile_lines"):
        fails.append("the member-less fixture carries members after all")
    said = sentence(_engine_report("fem"))
    if not said:
        fails.append("the model figure's sentence was not found")
    elif "reinforcement" in said or "pile" in said:
        fails.append(f"a model with no member is credited with members: {said!r}")
    elif "the material zones" not in said:
        fails.append(f"the model figure's sentence does not name the material "
                     f"zones it does draw: {said!r}")

    reinforced = load_slope_data_cached(FEM_REINF_XLSX)
    if not reinforced.get("reinforcement_lines"):
        fails.append("the reinforced fixture carries no line, so the positive "
                     "case is untested")
    said = sentence(_engine_report("fem", xlsx=FEM_REINF_XLSX))
    if "the reinforcement lines" not in said:
        fails.append(f"a model carrying six reinforcement lines does not say the "
                     f"figure draws them: {said!r}")
    if "pile" in said:
        fails.append(f"a model with no pile is credited with piles: {said!r}")
    return fails


#: Wording the finite element section is no longer allowed to print, with what
#: was wrong with each. Every one of them was on a shipped page and was read by
#: the owner: they are the failures the prose audit was called for, pinned so a
#: revert or a re-derivation cannot bring them back unnoticed.
FEM_RETIRED_WORDING = (
    ("shows the finite element model",
     "the bare Figure-N-shows-X register the owner superseded"),
    ("shows the finite element mesh",
     "the bare Figure-N-shows-X register the owner superseded"),
    ("is the section the analysis was run on",
     "the model figure introduced as a thing rather than shown"),
    ("the material zones the properties below belong to",
     "an appositive naming the zones by what belongs to them"),
    ("the members the solution carries",
     "members lumped instead of named"),
    ("is the mesh the section was discretized onto",
     "the mesh sentence written backwards"),
    ("fixities", "jargon for the boundary conditions the legend names plainly"),
    ("the boundary conditions are marked on it",
     "the marks named instead of the restraint they stand for"),
    ("counted in Section",
     "a cross-reference that counts rather than describes"),
    ("its own weight is switched on in one step",
     "colloquial mechanism-speak for applying gravity"),
    ("the one plane-strain elasticity produces from the vertical",
     "an appositive with the referent buried"),
    ("which is about 0.43 of it", "a dangling referent for sigma_v"),
    ("is the search that reached it",
     "the search figure introduced as a thing rather than shown"),
    ("past the trial's own elastic scale and still growing",
     "the failure criterion compressed past the point of reading"),
    ("A trial counts as failed only when the solution cannot reach equilibrium "
     "and the computed displacements are both large",
     "the failure criterion demanding BOTH displacement signals to fail a "
     "trial, which is stricter than the test that runs"),
    ("In the strength reduction method the cohesion",
     "the method's mechanism given before the method is named or defined"),
)

#: What every finite element section must say, in the register the owner set on
#: the fem_piles review: noun-first sentences with the figure reference inside
#: them, and the method named in full and defined before any of its machinery.
FEM_REQUIRED_WORDING = (
    ("The problem inputs specific to the finite element model are shown in",
     "the model figure has no plain sentence"),
    ("The finite element mesh constructed for the problem is shown in",
     "the mesh figure has no plain sentence"),
    ("The mesh is colored by material.",
     "the mesh sentence does not say what is drawn on the figure"),
    ("The nodes along the base of the mesh are fixed in both directions",
     "the mesh sentence does not say what was held at the boundary"),
    ("No at-rest coefficient K₀ is specified",
     "the in-situ assumption is not stated plainly"),
    ("σ_h = ν/(1−ν)·σ_v", "the in-situ horizontal stress is not given"),
)

#: What a STRENGTH REDUCTION section must say, over and above the above: the
#: method named in full and defined before any of its machinery. A gravity run
#: reduces no strength and says none of it.
SSRM_REQUIRED_WORDING = (
    ("shear strength reduction method",
     "the method is never named in full"),
    ("The shear strength of every material is divided by a trial factor",
     "the method is never defined"),
    ("the largest reduction the section can withstand",
     "the factor of safety the method reports is never defined"),
    ("A trial that cannot reach equilibrium counts as failed unless",
     "the failure criterion is not stated in plain terms"),
)


#: A shipped strength reduction run whose bracket was closed by trials the
#: hybrid criterion counted as failed WITHOUT their displacement growing — the
#: case a criterion sentence demanding both signals describes wrongly.
NONCIRC_FEM_XLSX = os.path.join(_REPO, "docs", "fem", "files",
                                "xslope_noncircular_fem.xlsx")


def _hybrid_thresholds():
    """The two displacement thresholds the hybrid criterion decides on, off
    :mod:`xslope.fem` rather than copied here."""
    from xslope import fem
    return fem._HYBRID_U_STUCK_MAX, fem._HYBRID_GROWTH_MIN


def test_the_inputs_an_engine_reads_are_stated_where_it_is_documented():
    """Every input a stability model figure draws is named in that figure's own
    sentence, and every input the analysis reads is set down in full.

    The Project Definition figure is the section: geometry and material zones.
    The water lines, the loads and the members are each read by a particular
    analysis, and each is drawn and named where that analysis is documented (the
    owner's rulings on fem_reinforce, fem_johnson_res and fem_noncircular). What
    the model carries decides what is said: a dry section is credited with no
    water surface, and a section with no load is credited with no load.

    Three inputs the report used to draw and never state are stated here: the
    ponded-water load the engine derives (a row of the loads table like any
    other, marked as derived, with the surface it was computed from named), the
    geometry of a piezometric line (its coordinates, or the one elevation a level
    line stands at), and what a row of piles IS in a plane-strain section.
    """
    fails = []
    from xslope.report import _model_figure_shows, water_features

    def inputs_sentence(report, engine):
        want = f"specific to the {engine} model are shown in"
        for section in report.sections:
            for _lvl, sec in section.walk():
                for b in sec.blocks:
                    if b.kind == "prose" and want in b.text:
                        return b.text
        return ""

    # --- the feature list is the model's own -------------------------------
    #
    # johnson_res: a pool from head boundaries, a derived load, no member.
    # noncircular_fem: a piezometric line, no load, no member.
    for xlsx, present, absent in (
            (SEEP_XLSX, ("the water surface", "the distributed loads"),
             ("the piezometric line", "the reinforcement lines", "the piles")),
            (NONCIRC_FEM_XLSX, ("the piezometric line",),
             ("the water surface", "the distributed loads", "the piles"))):
        model = load_slope_data_cached(xlsx)
        shows = _model_figure_shows(model, water_features(model))
        for named in present:
            if named not in shows:
                fails.append(f"{os.path.basename(xlsx)}: the figure draws "
                             f"{named} and the list does not name it: {shows}")
        for named in absent:
            if named in shows:
                fails.append(f"{os.path.basename(xlsx)}: the list names {named}, "
                             f"which this model does not carry: {shows}")

    # …and it reaches the finite element section's own sentence.
    said = inputs_sentence(_engine_report("fem", xlsx=NONCIRC_FEM_XLSX),
                           "finite element")
    if "the piezometric line" not in said:
        fails.append(f"a model whose figure draws a piezometric line does not "
                     f"say so: {said!r}")

    # …and the Project Definition figure's sentence names NONE of them: it is
    # the section, and nothing an analysis puts on it or reads off it.
    pd_said = ""
    report = _build({"pd_figure": True})
    for section in report.sections:
        if section.title != "Project Definition":
            continue
        for b in section.blocks:
            if b.kind == "prose" and "displayed in Figure" in b.text:
                pd_said = b.text
    if "the geometry" not in pd_said or "the material zones" not in pd_said:
        fails.append(f"the project definition figure's sentence does not name "
                     f"the section it draws: {pd_said!r}")
    for banned in ("water", "load", "piezometric", "reinforcement", "pile"):
        if banned in pd_said.lower():
            fails.append(f"the project definition figure is credited with "
                         f"{banned!r}, which it does not draw: {pd_said!r}")

    # The maximum depth line IS geometry — the profile lines are closed at it to
    # make the material zones — so the sentence names it on a model that defines
    # one and on no other (the owner's ruling). The figure draws it under exactly
    # that condition.
    from xslope.report import DEFAULT_OPTIONS, _Counter, _project_definition_section

    depth_model = load_slope_data_cached(REINF_XLSX)
    if depth_model.get("max_depth") is None:
        fails.append("the sample defines no maximum depth; the clause is "
                     "untested")
    else:
        with tempfile.TemporaryDirectory() as tmp:
            def _pd_said(model):
                opts = dict(DEFAULT_OPTIONS, method="bishop", pd_figure=True)
                section = _project_definition_section(
                    model, {}, opts, _Counter(), tmp)
                return " ".join(b.text for b in section.blocks
                                if b.kind == "prose")

            named = "the maximum depth line"
            if named not in _pd_said(depth_model):
                fails.append(f"the figure draws the maximum depth line and the "
                             f"sentence does not name it: "
                             f"{_pd_said(depth_model)!r}")
            without = _pd_said(dict(depth_model, max_depth=None))
            if named in without:
                fails.append(f"a model with no maximum depth is credited with "
                             f"the line: {without!r}")

    # --- the derived water load is a row of the loads table ------------------
    from xslope.report import _loads_section, _Counter
    from xslope.water import derived_blocks, with_water_loads

    dam = load_slope_data_cached(SEEP_XLSX)
    if not derived_blocks(with_water_loads(dam), 1):
        fails.append("the seepage sample derives no water load; the row is "
                     "untested")
    else:
        loads = _loads_section(dam, water_features(dam), _Counter())
        table = next((b for b in loads.blocks if b.kind == "table"), None)
        prose = " ".join(b.text for b in loads.blocks if b.kind == "prose")
        if table is None:
            fails.append("a model whose only load is derived printed no loads "
                         "table; the load reaches the figure and not the page")
        else:
            if "Source" not in table.headers:
                fails.append(f"the derived load is not told from a typed one: "
                             f"{table.headers}")
            elif not any("Derived" in str(r[-1]) for r in table.rows):
                fails.append(f"no row is marked as derived: {table.rows}")
            # The pressures printed are the derivation's own.
            printed = {v for r in table.rows for v in re.findall(r"@ ([\d.]+)",
                                                                 str(r[1]))}
            peak = max(float(pt["Normal"]) for blk in
                       derived_blocks(with_water_loads(dam), 1) for pt in blk)
            if printed and f"{peak:g}" not in printed:
                fails.append(f"the table's pressures {sorted(printed)} do not "
                             f"carry the derivation's peak {peak:g}")
        # And the paragraph says where the water it was computed from stands.
        if "computed by the engine rather than entered" not in prose:
            fails.append(f"the derived load is not said to be computed: "
                         f"{prose!r}")
        from xslope.report import _load_sources
        source = (_load_sources(dam) or {}).get(1, "")
        if source and source not in prose:
            fails.append(f"the water surface the load was computed from is not "
                         f"named ({source!r}): {prose!r}")

    # --- a piezometric line's geometry --------------------------------------
    from xslope.report import _piezo_sections, _level_elevation, _piezo_lines

    level = load_slope_data_cached(NONCIRC_FEM_XLSX)
    lines = _piezo_lines(level)
    if not lines:
        fails.append("the piezometric fixture carries no line")
    elif _level_elevation(lines[0][1]) is None:
        fails.append("the piezometric fixture's line is not level; the "
                     "one-elevation branch is untested")
    else:
        sections, numbered = _piezo_sections(level, _Counter())
        said = " ".join(b.text for b in sections[0].blocks
                        if b.kind == "prose")
        if numbered or [b for b in sections[0].blocks if b.kind == "table"]:
            fails.append(f"a level line was given a table of one elevation "
                         f"repeated: {numbered}")
        if "is level at elevation" not in said:
            fails.append(f"a level line's elevation is not stated: {said!r}")
        if f"{_level_elevation(lines[0][1]):g}" not in said:
            fails.append(f"the elevation stated is not the line's: {said!r}")

    # A line with a shape gets its coordinates, every point of them.
    shaped = copy.deepcopy(level)
    shaped["piezo_line"] = [(0.0, 5.0), (10.0, 4.0), (20.0, 2.0)]
    sections, numbered = _piezo_sections(shaped, _Counter())
    table = next((b for b in sections[0].blocks if b.kind == "table"), None)
    said = " ".join(b.text for b in sections[0].blocks if b.kind == "prose")
    if table is None:
        fails.append("a piezometric line with a shape printed no coordinates")
    else:
        if len(table.rows) != 3:
            fails.append(f"3 points printed as {len(table.rows)} rows")
        if table.rows and table.rows[1][1:] != ["10", "4"]:
            fails.append(f"the second point is printed {table.rows[1][1:]}")
        if f"listed in Table {table.number}" not in said:
            fails.append(f"the coordinates table is not introduced: {said!r}")

    # --- the piles' formulation ---------------------------------------------
    piled = load_slope_data_cached(FEM_PILES_XLSX)
    said = " ".join(_prose(_engine_report("fem", xlsx=FEM_PILES_XLSX)))
    spacing = {p.get("S") for p in piled.get("pile_lines") or []}
    for want in ("continuous out of plane", "divided by that spacing",
                 "approximation for a row of separate piles", "arches"):
        if want not in said:
            fails.append(f"the pile section does not state the formulation "
                         f"assumption ({want!r})")
    for s in spacing:
        if s and f"{float(s):g}" not in said:
            fails.append(f"the spacing the stiffnesses were divided by ({s}) is "
                         f"not stated")
    # A limit equilibrium report takes no stiffness off a pile and says none of
    # it: the sentence belongs to the engine that assembles the beam.
    lem_said = " ".join(b.text for b in _build().blocks("prose"))
    if "continuous out of plane" in lem_said:
        fails.append("a limit equilibrium report states the beam idealization "
                     "it never makes")

    # --- a peak reported with its depth, both peaks --------------------------
    sec = _member_section(_engine_report("fem", xlsx=FEM_PILES_XLSX),
                          "Pile Forces")
    table = next((b for b in (sec.blocks if sec else []) if b.kind == "table"),
                 None)
    if table is None:
        fails.append("the pile subsection carries no forces table")
    else:
        heads = [str(h) for h in table.headers]
        shear = next((i for i, h in enumerate(heads)
                      if h.startswith("Peak shear")), None)
        moment = next((i for i, h in enumerate(heads)
                       if h.startswith("Peak moment")), None)
        if shear is None or moment is None:
            fails.append(f"the pile table lost a peak column: {heads}")
        else:
            # The depth columns are located BEFORE any value is read out of one:
            # a table missing them has nothing at those indices to compare, and a
            # check that read them anyway would die of whatever happened to be
            # there instead of failing with its own sentence.
            paired = True
            for what, i in (("shear", shear), ("moment", moment)):
                if i + 1 >= len(heads) or not heads[i + 1].startswith("At depth"):
                    fails.append(f"the peak {what} is printed without its "
                                 f"depth: {heads}")
                    paired = False
            if paired:
                pile_data, pile_bundle = _fem_1d_bundle(FEM_PILES_XLSX)
                for row, profile in zip(table.rows,
                                        _profiles(pile_data, pile_bundle,
                                                  "pile")):
                    for key, i in (("max_shear_depth", shear + 1),
                                   ("max_moment_depth", moment + 1)):
                        want = profile.get(key)
                        if want is None:
                            continue
                        if abs(float(row[i]) - float(want)) > 0.01:
                            fails.append(f"{profile['label']}: {key} printed "
                                         f"{row[i]}, measured {want}")
    return fails


def test_the_hybrid_criterion_sentence_is_true_of_the_runs_that_shipped():
    """The sentence describing the hybrid criterion states the test that runs,
    not a stricter one.

    ``solve_ssrm`` counts a trial as standing only where it converged or where
    :func:`xslope.fem.classify_nonconvergence` returned STABLE_STUCK, and
    STABLE_STUCK needs BOTH signals absent: displacement at elastic scale AND no
    longer growing. One signal without the other is AMBIGUOUS, and AMBIGUOUS
    counts as FAILED. So "a trial counts as failed only when ... the
    displacements are both large and still growing" describes a stricter test
    than the one that decided the answer, and noncircular_fem is the
    counterexample on disk: two of its trials were counted as failed with
    displacement past elastic scale and growth of 0.006 and 0.008, and the
    second is the upper end of the bracket its factor of safety is the midpoint
    of.

    Checked three ways: the fixture really poses the question, the sentence the
    report prints does not claim the strict conjunction, and the rule the
    sentence states is the rule the classifier applies.
    """
    fails = []
    import json
    from xslope.fem import classify_nonconvergence
    from xslope.report import SSRM_CRITERIA

    said = SSRM_CRITERIA["hybrid"]
    stuck_max, growth_min = _hybrid_thresholds()

    # --- the fixture really poses the question ------------------------------
    meta = json.load(open(os.path.splitext(NONCIRC_FEM_XLSX)[0]
                          + "_fem_meta.json"))
    if meta.get("failure_criterion") != "hybrid":
        fails.append(f"the counterexample run was solved under "
                     f"{meta.get('failure_criterion')!r}, not the hybrid "
                     f"criterion this check is about")
    counted = [t for t in (meta.get("trials") or [])
               if not t.get("converged") and not t.get("stable")
               and t.get("growth") is not None
               and float(t["growth"]) <= growth_min]
    if not counted:
        fails.append("no shipped trial is counted as failed without its "
                     "displacement growing, so the overclaim is untested")
    # And one of them set the answer: the bracket the factor of safety is the
    # midpoint of has this trial at its upper end.
    interval = [float(v) for v in (meta.get("final_interval") or [0, 0])]
    if not any(abs(float(t["F"]) - interval[-1]) < 1e-9 for t in counted):
        fails.append(f"none of the not-growing failures is the bracket end "
                     f"{interval[-1]}; the case is weaker than it reads")

    # --- the sentence does not claim the strict conjunction -----------------
    for wrong in ("both large", "and still growing. An iteration",
                  "counts as failed only when the solution cannot reach "
                  "equilibrium and"):
        if wrong in said:
            fails.append(f"the criterion sentence claims a stricter test than "
                         f"the one that runs: {wrong!r}")
    # --- and it states the rule the classifier applies ----------------------
    for wanted in ("unless", "stayed at the elastic scale", "stopped growing"):
        if wanted not in said:
            fails.append(f"the criterion sentence does not state the standing "
                         f"case as the exception it is ({wanted!r} missing): "
                         f"{said!r}")

    # The classifier itself, on the two shapes the sentence distinguishes: a
    # trial at elastic scale and steady stands; the same trial past that scale
    # stands no longer, growth or no growth. Built from the thresholds fem.py
    # declares, so the prose cannot drift from the code.
    steady = [1.0] * 12
    verdict, _u, _g = classify_nonconvergence(steady, 1.0, "iteration_cap")
    if verdict != "STABLE_STUCK":
        fails.append(f"a trial at elastic scale and steady reads {verdict!r}, "
                     f"so the sentence's standing case is not the code's")
    past = [stuck_max + 0.3] * 12
    verdict, _u, _g = classify_nonconvergence(past, 1.0, "iteration_cap")
    if verdict == "STABLE_STUCK":
        fails.append("a trial past elastic scale and steady is counted as "
                     "standing, so the sentence overstates what fails")

    # The report prints it, on the run that poses the question.
    printed = " ".join(_prose(_engine_report("fem", xlsx=NONCIRC_FEM_XLSX)))
    if said not in printed:
        fails.append(f"the counterexample run's own section does not carry the "
                     f"criterion sentence: {printed!r}")
    return fails


def test_the_fem_prose_reads_as_documentation():
    """Every sentence the finite element section prints is one a practising
    engineer reads once.

    The section was written in a register of its own — "Figure 2 is the section
    the analysis was run on: the material zones the properties below belong to",
    "with the fixities the solution was found under marked on the nodes that
    carry them", a strength reduction paragraph that opened on viscoplasticity
    and never said what the shear strength reduction method is. The limit
    equilibrium section's reviewed sentences are the register: "Figure 2 shows
    the limit equilibrium model: the section and its materials, ...". These pins
    hold the swept wording in place from both directions — the retired phrasings
    cannot come back, and the facts they carried cannot go missing with them.

    Read off two shipped models, so both branches of the sentences that have two
    are covered: one with members and one without, one with a shared seepage
    mesh and one with a mesh of its own.
    """
    fails = []
    ssrm = " ".join(_prose(_engine_report("fem")))
    gravity = " ".join(_prose(_engine_report("fem", xlsx=FEM_REINF_XLSX)))
    for label, said in (("griffiths1 (strength reduction)", ssrm),
                        ("reinforce_fem (gravity)", gravity)):
        for phrase, why in FEM_RETIRED_WORDING:
            if phrase in said:
                fails.append(f"{label}: {why} — {phrase!r} is back on the page")
        for phrase, why in FEM_REQUIRED_WORDING:
            if phrase not in said:
                fails.append(f"{label}: {why} ({phrase!r} is not on the page)")
    for phrase, why in SSRM_REQUIRED_WORDING:
        if phrase not in ssrm:
            fails.append(f"the strength reduction section: {why} ({phrase!r} is "
                         f"not on the page)")
        if phrase in gravity:
            fails.append(f"a gravity run is described as a strength reduction "
                         f"run: {phrase!r}")
    # The method is defined BEFORE its machinery: a reader meets the name and the
    # definition, then the finite element model, then the search.
    order = [ssrm.find("shear strength reduction method"),
             ssrm.find("Each trial is solved by the finite element method"),
             ssrm.find("The critical factor is bracketed")]
    if -1 in order or order != sorted(order):
        fails.append(f"the paragraph does not run definition, then model, then "
                     f"search: {order}")
    return fails


#: How every table in the report is introduced, one entry per table class: the
#: sentence shape the owner set, and the sentence shape it replaced.
#:
#: The register is noun-first — the subject is what the table holds, the table
#: reference sits inside the sentence, and the columns are named after it. What
#: it replaced was an inventory: the table as subject, a verb of possession, and
#: a colon followed by a list ("Table 4 gives each pile: its head and tip, …").
#: Every one of these was rewritten on the owner's review, and each is pinned
#: from BOTH directions — the shape that must be there, and the shape that must
#: not come back. One direction is not enough: a required phrase that is merely
#: ABSENT from a report that never prints that table passes vacuously, and a
#: retired phrase can return under a required one that is still true elsewhere.
TABLE_INTRO_REGISTER = (
    ("limit equilibrium materials",
     "The material properties associated with the limit equilibrium analysis "
     "are shown in",
     "gives the properties of every material in the"),
    ("seepage materials",
     "The material properties associated with the seepage analysis are shown in",
     "gives the properties of every material in the"),
    ("finite element materials",
     "finite element analysis are shown in",
     "gives the properties every element is solved with"),
    ("distributed loads",
     "The distributed loads defined for this problem are listed in",
     "gives each distributed load:"),
    ("piezometric line",
     "defined for this problem is listed in",
     "gives the line's own points:"),
    ("reinforcement properties",
     "The geometry and properties for each of the reinforcement lines defined "
     "for this problem are listed in",
     "gives each reinforcement line:"),
    ("pile properties",
     "The geometry and properties for each of the piles defined for this "
     "problem are listed in",
     "gives each pile:"),
    ("pile forces",
     "The forces developed in each pile are listed in",
     "gives every pile the analysis solved:"),
    ("factors of safety",
     "factor of safety reported by each method is listed in",
     "gives the factor of safety each method reported"),
    ("slice table",
     "The geometry, forces and strengths of every slice on the",
     "holds the geometry, forces and strengths of every"),
    ("nomenclature",
     "The symbols used above are defined in",
     "the symbols the equations above use"),
)

#: The retired register, as a shape rather than as a list of sentences: a table
#: as the subject of a verb of possession. This is the pin that does not have to
#: be updated when a table is added — whatever the new table is called, it
#: cannot be introduced in the register the owner retired.
TABLE_AS_SUBJECT = re.compile(r"\bTables? \d+(?:[–-]\d+)? "
                              r"(gives|holds|lists|shows|summari[sz]es)\b")


def test_a_wide_table_is_fitted_by_the_lines_it_makes():
    """A table too wide for the page is cut so that no cell wraps further than
    it has to, and no value wraps at all.

    "Keep every column, but be smarter about the widths" was the owner's ruling
    on the pile forces table, which had come out with one column of prose broken
    into four lines beside a column of five-character displacements printed at
    an inch and a quarter. The cause was a measurement, not a preference: a
    column of numbers was floored at its widest line INCLUDING its header, so
    the longest thing written at the top of a table set the width of the table.

    What is measured here is the outcome and not the recipe. The shipped
    document is read back, every cell is wrapped at the width it was actually
    given, and three things have to hold: no value is broken; the tallest cell
    in the table is as short as the page allows, proved by showing that one line
    fewer does not fit; and two columns under the same header are the same
    width.
    """
    fails = []
    import re

    from docx import Document

    from xslope.report_docx import (_column_widths, _table_font, _text_width,
                                    _usable_twips, _width_within,
                                    _wrapped_lines)

    slope_data, bundle = _fem_1d_bundle(FEM_PILES_XLSX)
    with tempfile.TemporaryDirectory() as tmp:
        opts = {"input_path": FEM_PILES_XLSX, "lem": False, "pd_figure": False,
                "fem_inputs_figure": False, "fem_mesh_figure": False,
                "fem_figure": False, "fem_pile_figures": False}
        opts.update(FAST_FIGURES)
        _report, doc_xml, path = _written("the piles report", slope_data,
                                          {"fem": bundle}, opts, tmp)
        document = Document(path)
        family = _table_font(document)
        usable = _usable_twips(document.sections[0])
        wide = None
        for tbl in re.findall(r"<w:tbl>.*?</w:tbl>", doc_xml, re.S):
            columns = _table_cell_texts(tbl)
            if columns and columns[0] and columns[0][0] == "Pile":
                wide = (tbl, columns)
        if wide is None:
            return ["the piles report carries no pile forces table"]
        tbl, columns = wide
        grid = [int(w) for w in re.findall(r'<w:gridCol w:w="(\d+)"/>', tbl)]
        size = _table_point_size(tbl)
        margin = _cell_margin(tbl, "", "TableGrid")
        pad = 2 * (DEFAULT_CELL_MARGIN if margin is None else margin)
        if len(grid) != len(columns) or size is None:
            return [f"the pile table reads as {len(columns)} columns of text "
                    f"for {len(grid)} grid columns at {size} pt"]
        if sum(grid) <= usable - 1:
            return [f"the pile table fits the page with {usable - sum(grid)} "
                    f"twips to spare, so nothing below is being measured"]

        # No NUMBER is broken. A number over two lines is the defect the whole
        # measurement exists to prevent; a sentence in a cell is meant to wrap,
        # and the column it wraps in is what the rest of this measures.
        from xslope.columns import is_number
        for j, texts in enumerate(columns):
            for text in texts[1:]:
                if is_number(text) and \
                        _wrapped_lines(text, family, size, grid[j] - pad) > 1:
                    fails.append(f"column {j} prints the number {text!r} on "
                                 f"more than one line at {grid[j]} twips")

        # The tallest cell is as short as the page allows.
        tallest = max(_wrapped_lines(t, family, size, grid[j] - pad)
                      for j, texts in enumerate(columns) for t in texts)
        if tallest > 1:
            need = 0.0
            for j, texts in enumerate(columns):
                lo = max(_text_width(w, family, size)
                         for t in texts for w in t.split())
                hi = max(_text_width(t, family, size) for t in texts)
                need += _width_within(texts, family, size, lo, hi,
                                      tallest - 1) + pad
            if need <= usable:
                fails.append(f"the table wraps to {tallest} lines where "
                             f"{tallest - 1} would fit in {need:.0f} of "
                             f"{usable} twips")

        # One header, one width.
        by_header = {}
        for j, texts in enumerate(columns):
            by_header.setdefault(texts[0], []).append(j)
        for header, members in by_header.items():
            if len({grid[j] for j in members}) > 1:
                fails.append(f"the two {header!r} columns print at "
                             f"{[grid[j] for j in members]} twips")

    # …and the measurement that decides it does not read the header as a value:
    # a header far longer than anything under it may wrap.
    long_header = [["Head movement (ft)", "-6.308", "-8.135"],
                   ["State", "within capacity", "within capacity"]]
    widths = _column_widths(long_header, "Calibri", 9.0, 2000, 86,
                            nowrap=[True, False], header_rows=1)
    if widths[0] >= _text_width("Head movement (ft)", "Calibri", 9.0):
        fails.append(f"a column of five-character numbers is held open by its "
                     f"own header: {widths}")
    return fails


def test_the_tension_crack_is_documented_where_it_is_read():
    """The tension crack is drawn on the limit equilibrium model figure, named
    in that figure's own sentence, and named there once.

    It was drawn on the Project Definition figure as well — the same red dotted
    line one page from the figure that means something by it — and named in
    neither sentence. No engine but the method of slices reads it: it caps the
    driving side of a slice and carries the crack water force, and a seepage or
    finite element run reads neither. Where it is drawn is checked on the plots
    themselves ("the shared-model plot"); this is the sentence that names it.
    """
    fails = []
    report = _calc_report("spencer", xlsx=TENSION_XLSX)[0]
    if report is None:
        return ["the tension-crack model no longer solves, so nothing here is "
                "checked"]
    if not load_slope_data_cached(TENSION_XLSX).get("tcrack_depth"):
        fails.append("the fixture states no tension crack depth")
    said = [b.text for b in report.blocks("prose")
            if "specific to the limit equilibrium model are shown in" in b.text]
    if len(said) != 1:
        fails.append(f"the limit equilibrium model figure has {len(said)} "
                     f"sentences introducing it")
    elif "tension crack" not in said[0]:
        fails.append(f"the figure draws the tension crack and its sentence does "
                     f"not name it: {said[0]!r}")

    # Named where it is drawn and nowhere else: the Project Definition figure's
    # own sentence describes a section without one.
    shared = [b.text for b in report.blocks("prose")
              if "The problem definition is displayed in" in b.text]
    for text in shared:
        if "tension crack" in text:
            fails.append(f"the shared section figure is credited with the "
                         f"tension crack it no longer draws: {text!r}")

    # A model with no crack is credited with none.
    plain = _build()
    if load_slope_data_cached(REINF_XLSX).get("tcrack_depth"):
        fails.append("the crack-less fixture states a crack after all")
    for b in plain.blocks("prose"):
        if "are shown in" in b.text and "tension crack" in b.text:
            fails.append(f"a model with no tension crack is credited with one: "
                         f"{b.text!r}")
    return fails


def test_every_table_is_introduced_in_the_owners_register():
    """Every table the report prints is introduced by a noun-first sentence that
    says what the table holds, with the table's own reference inside it.

    The reports shipped for review introduced their tables as inventories —
    "Table 4 gives each pile: its head and tip, its diameter and out-of-plane
    spacing, …" — a register the owner superseded across every table class in
    one review. The rewrite is pinned here from both directions and across every
    class, because a rewrite that only one check watches is a rewrite one revert
    undoes: the finite element materials intro was put back to its retired form
    and the whole suite still passed.

    Read off real reports carrying between them every table the builder can
    print: the limit equilibrium tree (materials, the loads, the reinforcement
    properties, factors of safety, the slice table and its nomenclature), a
    model whose piezometric line has a shape and so gets a table of its points,
    the seepage tree, and two finite element trees (materials, the piles, and
    the forces developed in each). The reinforcement carries no forces table by
    the owner's own ruling, so there is none to introduce.
    """
    fails = []
    trees = {
        "limit equilibrium": _build(options={"method": ["spencer", "bishop"]}),
        "piezometric line": _calc_report(
            "spencer", options={"lem_inputs_figure": False},
            xlsx=PIEZO_XLSX)[0],
        "seepage": _engine_report("seep"),
        "finite element (strength reduction)": _engine_report("fem"),
        "finite element (piles)": _engine_report("fem", xlsx=FEM_PILES_XLSX),
        "finite element (reinforcement)": _engine_report("fem",
                                                         xlsx=FEM_REINF_XLSX),
    }
    said = {name: " ".join(_prose(tree)) for name, tree in trees.items()}
    everything = " ".join(said.values())

    # Every table class the builder can print is printed by one of these trees,
    # so a required sentence missing from all of them is a table introduced some
    # other way — not a table these fixtures happen not to carry.
    for label, required, retired in TABLE_INTRO_REGISTER:
        if required not in everything:
            fails.append(f"the {label} table is not introduced in the owner's "
                         f"register ({required!r} is on no page)")
        for name, text in said.items():
            if retired in text:
                fails.append(f"{name}: the {label} table is introduced as an "
                             f"inventory again — {retired!r}")

    # …and the shape itself, whatever the table.
    for name, text in said.items():
        for m in TABLE_AS_SUBJECT.finditer(text):
            fails.append(f"{name}: a table is made the subject of a verb of "
                         f"possession — {text[max(0, m.start() - 40):m.end() + 60]!r}")

    # The tables really are there: a register check reading prose alone would
    # pass a report that stopped printing the tables the sentences introduce.
    for name, tree in trees.items():
        if not tree.blocks("table"):
            fails.append(f"{name}: the tree carries no table at all, so the "
                         f"sentences checked above introduce nothing")
    return fails


def test_the_boundary_sentence_is_derived_from_what_was_held():
    """The mesh sentence states the restraint the solve actually ran under,
    read off its constraint array and not off a fixed list.

    ``main!D22`` chooses what the sides of the domain are held by, and the two
    choices are two different problems: rollers let the truncated ground settle
    under its own weight, and clamped sides add a shear restraint the ground
    beyond the mesh does not have. A sentence that names rollers whatever the
    model asked for is a false statement about what was solved, and every
    finite element report carries it. The derivation was pinned by nothing: the
    roller codes were replaced with a hardcoded list and the whole suite passed.

    Both directions, and on a real model rather than on an array written here:
    the shipped sample is rebuilt with its sides clamped, through the same
    ``build_fem_data`` the solve uses, and the sentence has to follow. The code
    that stands for the other roller is exercised too, so neither roller can be
    the one the sentence always names.
    """
    fails = []
    import numpy as np
    from xslope.fem import build_fem_data
    from xslope.fileio import load_slope_data
    from xslope.report import FEM_RESTRAINTS, _fem_boundary_conditions

    def bc_of(side_bc):
        """``(bc_type, sentence)`` for the sample model held that way."""
        with contextlib.redirect_stdout(io.StringIO()):
            sd = load_slope_data(FEM_XLSX)
            if side_bc is not None:
                sd["side_bc"] = side_bc
            fem_data = build_fem_data(sd, sd["mesh"])
        arr = np.asarray(fem_data.get("bc_type")).astype(int)
        return arr, _fem_boundary_conditions(fem_data)

    rollered, said_rollers = bc_of("rollers")
    clamped, said_fixed = bc_of("fixed")

    # The fixture is a fixture: the two models really were held differently.
    if not (rollered == 2).any():
        fails.append("the sample model on rollers carries no roller node, so "
                     "the sentence below is checked against nothing")
    if (clamped == 2).any() or (clamped == 3).any():
        fails.append(f"the clamped model still carries roller nodes, so it is "
                     f"not the fixture this check needs: "
                     f"{sorted(set(clamped.tolist()))}")
    if not (clamped == 1).any():
        fails.append("the clamped model holds nothing at all")

    if "on rollers" not in said_rollers:
        fails.append(f"a model whose sides are on rollers does not say so: "
                     f"{said_rollers!r}")
    if FEM_RESTRAINTS[2] not in said_rollers:
        fails.append(f"the roller sentence does not say what the rollers hold: "
                     f"{said_rollers!r}")
    if "roller" in said_fixed:
        fails.append(f"a model whose sides are clamped is reported on rollers — "
                     f"the sentence is not derived from what was held: "
                     f"{said_fixed!r}")
    if said_fixed != ("Every node on the boundary of the mesh is fixed in both "
                      "directions."):
        fails.append(f"a clamped model does not say what it was held by: "
                     f"{said_fixed!r}")

    # Either roller reaches the sentence, so neither can be the one it names
    # whatever the array says.
    for code in (2, 3):
        arr = np.array([1, 1, code, code], dtype=int)
        said = _fem_boundary_conditions({"bc_type": arr})
        if FEM_RESTRAINTS[code] not in said:
            fails.append(f"a mesh held by restraint {code} is described as "
                         f"something else: {said!r}")
        other = 5 - code                                # 2 <-> 3
        if FEM_RESTRAINTS[other] in said:
            fails.append(f"a mesh held by restraint {code} is also credited "
                         f"with restraint {other}: {said!r}")
    return fails


def test_fem_solve_facts_are_recorded_not_assumed():
    """A reloaded finite element solution carries the solve facts its file
    records, and none that it does not — and the report follows the record.

    ``import_fem_solution`` set ``converged=True`` on every file it read: a saved
    field was taken as proof the solve that made it had closed. And the largest
    displacement, the single-trial paragraph's only number, was never written
    down at all, so that paragraph silently dropped it. Both now travel in the
    meta sidecar, written off the solution being exported and read back where
    recorded.
    """
    fails = []
    from xslope.fem import (export_fem_solution, import_fem_meta,
                            import_fem_solution)

    slope_data, bundle = _fem_bundle()
    fem_data = bundle["fem_data"]
    saved = bundle["solution"]

    # A file that records none of them reloads claiming none. This is the whole
    # point: the file exists, and that is not evidence of anything. The four
    # facts are struck out of a copy of the corpus model's metadata, so the case
    # is the absence itself and not a corpus file that happens to be poor.
    facts_tmp = tempfile.mkdtemp(prefix="xslope_nofacts_")
    silent_stem = _sidecar_copy(
        os.path.splitext(FEM_XLSX)[0], facts_tmp,
        lambda meta: {k: v for k, v in meta.items()
                      if k not in ("converged", "iterations", "residual",
                                   "max_displacement")})
    with contextlib.redirect_stdout(io.StringIO()):
        unrecorded = import_fem_solution(fem_data, silent_stem)
    for key in ("converged", "iterations", "residual", "max_displacement"):
        if key in unrecorded:
            fails.append(f"the reloaded solution carries {key}="
                         f"{unrecorded[key]!r}, which its meta sidecar never "
                         f"recorded")

    # Written down, they come back — through the public pair, at the values the
    # solve had.
    facts = {"converged": False, "iterations": 137, "residual": 4.2e-5,
             "max_displacement": 0.012345}
    tmp = tempfile.mkdtemp(prefix="xslope_femmeta_")
    stem = os.path.join(tmp, "roundtrip")
    with contextlib.redirect_stdout(io.StringIO()):
        export_fem_solution(fem_data, dict(saved, **facts), stem,
                            meta={"FS": None, "analysis": "single"})
        back = import_fem_solution(fem_data, stem)
    meta = import_fem_meta(stem) or {}
    for key, want in facts.items():
        if meta.get(key) != want:
            fails.append(f"the meta sidecar records {key}={meta.get(key)!r}, "
                         f"and the solve had {want!r}")
        if back.get(key) != want:
            fails.append(f"the reload restored {key}={back.get(key)!r}, "
                         f"and the file records {want!r}")

    # A solution that knows none of them writes none, and the reload leaves the
    # keys absent rather than filling them in.
    bare = {k: v for k, v in saved.items() if k not in facts}
    stem2 = os.path.join(tmp, "silent")
    with contextlib.redirect_stdout(io.StringIO()):
        export_fem_solution(fem_data, bare, stem2,
                            meta={"FS": None, "analysis": "single"})
        back2 = import_fem_solution(fem_data, stem2)
    for key in facts:
        if key in (import_fem_meta(stem2) or {}):
            fails.append(f"a solve that recorded no {key} still wrote one")
        if key in back2:
            fails.append(f"a file recording no {key} reloaded with "
                         f"{key}={back2[key]!r}")

    # --- the report follows the record ---------------------------------------
    def prose_of(solution):
        run = dict(bundle, analysis="single", FS=None, solution=solution)
        return " ".join(_prose(_engine_report("fem", bundle=run)))

    closed = prose_of(dict(saved, converged=True))
    if "The solution converged." not in closed:
        fails.append(f"a run recorded as converged is not said to be: {closed!r}")
    stopped = prose_of(dict(saved, converged=False))
    if "The solution did not converge" not in stopped:
        fails.append(f"a run recorded as not converged is not said to be: "
                     f"{stopped!r}")
    quiet = prose_of(unrecorded)
    for claim in ("The solution converged.", "The solution did not converge"):
        if claim in quiet:
            fails.append(f"a run whose file records no convergence is reported "
                         f"as {claim!r}")

    # The largest displacement, the single-trial paragraph's only number.
    moved = prose_of(dict(unrecorded, max_displacement=0.012345))
    if f"{0.012345:.4g}" not in moved:
        fails.append(f"the recorded largest displacement is not stated: {moved!r}")
    if "largest computed displacement" in quiet:
        fails.append(f"a run that records no displacement states one anyway: "
                     f"{quiet!r}")
    shutil.rmtree(tmp, ignore_errors=True)
    return fails


def test_no_trial_factor_is_invented():
    """A saved run that recorded no strength reduction factor is reported without
    one; the factor of safety never stands in for it.

    rs2_28a records FS = 1.606 and no trial factor. The reader fell back to the
    factor of safety and set it on the solution as ``F``, from where the result
    panels printed "rendered at last converged F = 1.61" — a trial that was
    never run. The only trial factor that run recorded is the 1.847 in its
    at-failure sidecar.
    """
    fails = []

    def num(v):
        try:
            return float(v)
        except (TypeError, ValueError):
            return None

    _slope_data, solutions = _restored(RS2_28A_XLSX)
    bundle = solutions.get("fem")
    if not bundle:
        fails.append("rs2_28a ships no finite element solution to read")
        return fails

    from xslope.fem import import_fem_meta
    meta = import_fem_meta(os.path.splitext(RS2_28A_XLSX)[0]) or {}
    if meta.get("F") is not None:
        fails.append(f"rs2_28a's meta now records F={meta['F']!r}; the fixture "
                     f"no longer exercises the fallback it was chosen for")
    if num(meta.get("FS")) is None:
        fails.append("rs2_28a's meta records no FS, so there is nothing that "
                     "could have stood in for the trial factor")
    if "F" in bundle["solution"]:
        fails.append(f"the run recorded no trial factor and the solution came "
                     f"back carrying F={bundle['solution']['F']!r}")

    # The trial the run DID record survives — the fix removes an invention, not
    # the fact.
    captured = num((bundle.get("failure_solution") or {}).get("F"))
    if captured is None:
        fails.append("the captured at-failure trial's own F was lost")
    elif abs(captured - num(meta.get("FS"))) < 0.01:
        fails.append(f"the captured trial is {captured}, indistinguishable from "
                     f"the factor of safety; the fixture cannot show the two apart")

    # And a run that DOES record its trial factor keeps it.
    _sd, johnson = _restored(JOHNSON_XLSX)
    j_meta = import_fem_meta(os.path.splitext(JOHNSON_XLSX)[0]) or {}
    recorded = num(j_meta.get("F"))
    if recorded is None:
        fails.append("johnson_res records no trial factor, so keeping one is "
                     "untested")
    elif num((johnson.get("fem") or {}).get("solution", {}).get("F")) != recorded:
        fails.append(f"a run recording F={recorded} came back without it")
    return fails


def test_member_detail_figures_are_readable():
    """Every member's detail figure states its peak where a reader can read it.

    The peak is a point ON the profile, and on a member well inside its capacity
    that profile runs along the envelope's own descent — so a label offset by
    quadrant lands on the dashes it is meant to be read against. Checked on the
    artists: the label's box stays inside its panel, covers no other label and
    not the legend, and crosses none of the curves the panel draws. Six
    reinforcement lines and two piles, because where the peak falls is what
    decides where its label can go.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    from matplotlib.figure import Figure as MplFigure
    from matplotlib.text import Text
    from xslope.plot import _box_crosses_segment
    from xslope.plot_fem_details import plot_detail
    from xslope.report import FIGURE_SIZE

    def utilized(profile, share=0.89):
        """The same line worked up to ``share`` of its capacity.

        Under gravity alone the sample's bars sit at a few percent, and a peak
        near the axis is the easy case: the label has the whole panel over it. A
        bar near capacity puts its peak ON the envelope, which is where a rule
        that picks a quadrant lands the label on the dashes it is meant to be
        read against — the case this exists for, made from the real profile by
        scaling the one series it is a peak of.
        """
        import numpy as np
        peak = profile.get("peak_utilization")
        if not peak or not np.isfinite(peak) or peak <= 0:
            return None
        k = share / peak
        out = dict(profile)
        out["T"] = np.asarray(profile["T"], dtype=float) * k
        out["peak_T"] = float(profile["peak_T"]) * k
        out["peak_utilization"] = share
        return out

    def worked_pile(profile, share=0.89):
        """The same pile bent up to ``share`` of its moment capacity.

        The pile's counterpart of ``utilized``, and it exists for the same
        reason: under gravity the samples reach 14% and 20% of Mcap, so the peak
        label has the width of the moment panel to itself and the placement rule
        is never asked a hard question. Bent to near capacity the peak sits ON
        the Mcap rule — a dashed line with its own rotated label on it — in the
        narrowest of four panels sharing a depth axis, which is the case the
        solved offset exists for.
        """
        import numpy as np
        cap = profile.get("M_cap")
        peak = profile.get("max_moment")
        if not cap or not peak or not np.isfinite(cap) or not np.isfinite(peak):
            return None
        k = share * float(cap) / abs(float(peak))
        out = dict(profile)
        out["moment"] = np.asarray(profile["moment"], dtype=float) * k
        out["max_moment"] = float(peak) * k
        out["peak_utilization"] = share
        return out

    for xlsx, kind in ((FEM_REINF_XLSX, "reinforcement"),
                       (FEM_PILES_XLSX, "pile")):
        slope_data, bundle = _fem_1d_bundle(xlsx)
        profiles = _profiles(slope_data, bundle, kind)
        if not profiles:
            fails.append(f"{os.path.basename(xlsx)} carries no {kind} profile")
            continue
        make = utilized if kind == "reinforcement" else worked_pile
        worked = [p for p in (make(q) for q in profiles) if p]
        if not worked:
            fails.append(f"no {kind} could be worked up to near its capacity; "
                         f"the crowded case goes untested")
        profiles = profiles + worked
        for profile in profiles:
            fig = MplFigure(figsize=FIGURE_SIZE)
            FigureCanvasAgg(fig)
            with contextlib.redirect_stdout(io.StringIO()):
                plot_detail(profile, fig=fig)
            fig.canvas.draw()
            renderer = fig.canvas.get_renderer()
            where = f"{os.path.basename(xlsx)} {profile['label']}"

            # The peak labels are the bold ones; they are the only annotations
            # these panels place against the data rather than against the axes.
            # A reinforcement line has one — where it is most utilized. A pile
            # has two, one per panel: the largest shear at its depth and the
            # largest moment at its own (the owner's fem_piles ruling — a peak
            # reported without its depth cannot be found on the pile).
            peaks = [(ax, t) for ax in fig.axes for t in ax.texts
                     if t.get_fontweight() == "bold"]
            want = 2 if profile["kind"] == "pile" else 1
            if len(peaks) != want:
                fails.append(f"{where}: {len(peaks)} peak labels on the figure, "
                             f"not {want}")
                continue
            if want == 2:
                heads = sorted(t.get_text().split()[0] for _a, t in peaks)
                if heads != ["Mmax", "Vmax"]:
                    fails.append(f"{where}: the peak labels are {heads}")
                if len({id(ax) for ax, _t in peaks}) != 2:
                    fails.append(f"{where}: both peak labels are on one panel")
            for ax, label in peaks:
                box = Text.get_window_extent(label, renderer)
                frame = ax.get_window_extent(renderer)
                if not (frame.x0 <= box.x0 and box.x1 <= frame.x1
                        and frame.y0 <= box.y0 and box.y1 <= frame.y1):
                    fails.append(f"{where}: the peak label {label.get_text()!r} "
                                 f"runs outside the panel it belongs to")
                for other in ax.texts:
                    if other is label:
                        continue
                    if box.overlaps(Text.get_window_extent(other, renderer)):
                        fails.append(f"{where}: the peak label is printed over "
                                     f"{other.get_text()!r}")
                legend = ax.get_legend()
                if legend is not None and box.overlaps(
                        legend.get_window_extent(renderer)):
                    fails.append(f"{where}: the peak label is printed over the "
                                 f"legend")
            # The two series the peak is read AGAINST: the profile it is a point
            # of, and the capacity it is a fraction of. A label over either
            # hides the comparison it was printed to make. (A hairline the label
            # cannot avoid — a step's vertical riser in a crowded panel — sits
            # under an opaque backing and hides nothing, so it is not asked
            # about here.)
                crossed = []
                for line in ax.lines:
                    if line.get_gid() not in ("DETAIL_PROFILE",
                                              "DETAIL_CAPACITY"):
                        continue
                    name = str(line.get_label())
                    pts = line.get_xydata()
                    if pts is None or len(pts) < 2:
                        continue
                    px = ax.transData.transform(pts)
                    if any(_box_crosses_segment(box, a, b)
                           for a, b in zip(px[:-1], px[1:])):
                        crossed.append(name)
                if crossed:
                    fails.append(f"{where}: the peak label lies across {crossed}")
                if label.get_bbox_patch() is None:
                    fails.append(f"{where}: the peak label carries no backing; "
                                 f"where a panel leaves it nowhere clear it "
                                 f"dissolves into what is behind it")
    return fails


def _profiles(slope_data, bundle, kind):
    """The member profiles the report builds its member table from — read
    through the builder's own reader, so a check that compares a printed number
    to a profile is comparing it to the series it was printed from."""
    from xslope.report import _detail_profiles
    return _detail_profiles(slope_data, bundle, kind)


def _member_section(report, title):
    """The finite element run's Reinforcement or Piles subsection, or None."""
    for section in report.sections:
        for _lvl, sec in section.walk():
            if sec.title == title:
                return sec
    return None


def _column(table, prefix):
    """The index of the column whose header starts with ``prefix``, or None."""
    for i, head in enumerate(table.headers):
        if str(head) == prefix or str(head).startswith(prefix + " "):
            return i
    return None


def _cell_numbers(text):
    """The numbers in a member-table cell, as floats.

    A cell holds one number, or the two ends of a stretch written ``"a to b"`` —
    what the force and position columns print for a line that reaches its
    greatest utilization at more than one point. Raises ValueError on anything
    that is neither.
    """
    return [float(part.replace(",", "").strip())
            for part in str(text).split(" to ")]


def _member_faults(table, profiles, force_prefix, force_key):
    """What a member table has to say about the profiles it was built from.

    One row per member, in the order the solver assigns them, carrying that
    member's own peak force and its own utilization. The last two rules are the
    honest ones: a row cannot claim a share of capacity that the force printed
    beside it could not have produced, and a table in which EVERY force is zero
    is not a measurement of anything — that is what a solution carrying no
    member forces at all produces, and it printed as a result.
    """
    faults = []
    printed_forces = []
    if len(table.rows) != len(profiles):
        return [f"{table.caption!r} has {len(table.rows)} rows for "
                f"{len(profiles)} members"]
    force_col = _column(table, force_prefix)
    util_col = _column(table, "Utilization")
    if force_col is None or util_col is None:
        return [f"{table.caption!r} has no {force_prefix} or Utilization "
                f"column: {table.headers}"]
    for row, profile in zip(table.rows, profiles):
        where = f"{table.caption!r}, {profile['label']}"
        if row[0] != profile["label"]:
            faults.append(f"{where}: the row is headed {row[0]!r}")
        util = profile.get("peak_utilization")
        stated = row[util_col]
        want = "" if util is None else f"{util:.0%}"
        if stated != want:
            faults.append(f"{where}: the utilization is printed {stated!r}, "
                          f"and the profile peaks at {want!r}")
        force = profile.get(force_key)
        # A member at its greatest utilization along a stretch prints the range
        # of force over that stretch, and the profile carries the same two ends.
        span = profile.get(f"{force_key}_span")
        want = [float(v) for v in span] if span else (
            [float(force)] if force is not None else [])
        try:
            got = _cell_numbers(row[force_col])
        except ValueError:
            faults.append(f"{where}: the force column reads {row[force_col]!r}")
            continue
        if want and (len(got) != len(want) or any(
                abs(g - w) > 0.05 + 0.005 * abs(w)
                for g, w in zip(got, want))):
            faults.append(f"{where}: the force is printed {got} and the profile "
                          f"carries {want}")
        printed = max(got) if got else 0.0
        printed_forces.append(printed)
        if stated not in ("", "0%") and printed == 0.0:
            faults.append(f"{where}: the row claims {stated} of capacity with "
                          f"no force in the member")
    if printed_forces and not any(printed_forces):
        faults.append(f"{table.caption!r} prints a force of zero in every one of "
                      f"its {len(printed_forces)} rows; a table of zeros is what "
                      f"a solution carrying no member forces produces, and it is "
                      f"not a measurement")
    return faults


def _drawn_member_forces(slope_data, bundle, kind="reinforcement"):
    """Every axial force the detail figures of one run actually DRAW.

    Read off the mobilized-force curve each figure plots (``DETAIL_PROFILE``),
    at the state the report reads its members at, so what is measured is what a
    reader sees rather than what a table said.
    """
    import matplotlib
    matplotlib.use("Agg")
    from matplotlib.figure import Figure as MplFigure
    from xslope.plot_fem_details import plot_detail
    from xslope.report import (_detail_profiles, _fem_primary_state,
                               resolve_options)

    state = _fem_primary_state(bundle, resolve_options({}))
    out = []
    for profile in _detail_profiles(slope_data, bundle, kind, state):
        fig = MplFigure(figsize=(6.0, 4.0))
        with contextlib.redirect_stdout(io.StringIO()):
            plot_detail(profile, fig=fig)
        for ax in fig.axes:
            for line in ax.lines:
                if line.get_gid() == "DETAIL_PROFILE":
                    out += [float(v) for v in line.get_ydata()]
    return out


def test_fem_members_are_reported():
    """A finite element run that carries reinforcement lines or piles reports
    what the solution put in them, and one that carries neither reports nothing.

    The forces have to be REAL: a bar the mesh never loaded would let every
    number in the table be zero and every rule about it hold, so the mechanism
    is asserted to have engaged before anything is asked about how it is
    printed.
    """
    fails = []
    from xslope.fem_details import list_lines

    # --- reinforcement: six lines, each drawn --------------------------------
    slope_data, bundle = _fem_1d_bundle(FEM_REINF_XLSX)
    lines = list_lines(bundle["fem_data"], bundle["solution"], slope_data,
                       field_state="failure")
    n_lines = len(lines)
    if n_lines != len(slope_data.get("reinforcement_lines") or []):
        fails.append(f"the model owns elements for {n_lines} of its "
                     f"{len(slope_data.get('reinforcement_lines') or [])} "
                     f"reinforcement lines")
    if n_lines < 2:
        fails.append(f"the reinforcement model carries {n_lines} line; a report "
                     f"of one member proves nothing about drawing them all")

    report = _engine_report("fem", xlsx=FEM_REINF_XLSX)
    sec = _member_section(report, "Reinforcement Forces")
    if sec is None:
        fails.append(f"a run carrying {n_lines} reinforcement lines has no "
                     f"Reinforcement subsection: {_titles(report)}")
        return fails

    profiles = _profiles(slope_data, bundle, "reinforcement")
    peak = max((abs(p["peak_T"]) for p in profiles if p.get("peak_T")),
               default=0.0)
    if peak <= 0.0:
        fails.append("every reinforcement line came out of the solve at zero "
                     "force; the fixture never engages the bars and proves "
                     "nothing about reporting them")
    # NO summary table: T_max and T_res are already in the properties table of
    # this engine's inputs, and where the force peaks, how far it is utilized and
    # what state the line is in are read off the annotated figure of the line
    # itself (the owner's ruling, fem_reinforce review). What the table used to
    # guarantee — that a printed force is the profile's own — is guaranteed on
    # the drawn curve instead, below.
    if [b for b in sec.blocks if b.kind == "table"]:
        fails.append("the Reinforcement subsection still prints a summary table")
    named_in = [t for t in report.tables()
                if t.caption == "Finite element reinforcement lines"]
    if len(named_in) != 1:
        fails.append(f"{len(named_in)} reinforcement properties tables; the "
                     f"member names have nowhere to be anchored")
    else:
        anchored = " ".join(b.text for b in sec.blocks if b.kind == "prose")
        if f"Table {named_in[0].number}" not in anchored:
            fails.append(f"the subsection does not name the table its members "
                         f"are labeled in: {anchored!r}")
        for header in ("T_max",):
            if _column(named_in[0], header) is None:
                fails.append(f"the reinforcement properties table declares no "
                             f"{header} capacity: {named_in[0].headers}")

    # The locator opens the subsection and is not one of the details; every line
    # the analysis solved is drawn after it
    # (:func:`test_the_member_subsections_locate_their_members`).
    figures = [b for b in sec.blocks
               if b.kind == "figure" and not b.source.endswith(" map")]
    if len(figures) != len(profiles):
        fails.append(f"{n_lines} lines drew {len(figures)} detail figures; every "
                     f"line the analysis solved is drawn")
    else:
        captioned = " | ".join(f.caption for f in figures)
        for profile in profiles:
            if profile["label"] not in captioned:
                fails.append(f"{profile['label']} has no detail figure: "
                             f"{captioned!r}")
    planned, drawn = _planned_matches(report, "fem", xlsx=FEM_REINF_XLSX)
    if planned != drawn:
        fails.append(f"the reinforcement report planned {planned} figures and "
                     f"built {drawn}")

    # --- piles: two, both drawn ---------------------------------------------
    pile_data, pile_bundle = _fem_1d_bundle(FEM_PILES_XLSX)
    pile_profiles = _profiles(pile_data, pile_bundle, "pile")
    if not pile_profiles:
        fails.append("the pile model carries no pile the solver owns elements "
                     "for")
        return fails
    if max((abs(p.get("max_moment") or 0.0) for p in pile_profiles),
           default=0.0) <= 0.0:
        fails.append("every pile came out of the solve at zero moment; the "
                     "fixture never engages the beams")
    pile_report = _engine_report("fem", xlsx=FEM_PILES_XLSX)
    pile_sec = _member_section(pile_report, "Pile Forces")
    if pile_sec is None:
        fails.append(f"a run carrying {len(pile_profiles)} piles has no Piles "
                     f"subsection: {_titles(pile_report)}")
    else:
        pile_table = next((b for b in pile_sec.blocks if b.kind == "table"), None)
        if pile_table is None:
            fails.append("the Piles subsection carries no table")
        else:
            fails += _member_faults(pile_table, pile_profiles, "Peak moment",
                                    "max_moment")
        pile_figures = [b for b in pile_sec.blocks
                        if b.kind == "figure" and not b.source.endswith(" map")]
        if len(pile_figures) != len(pile_profiles):
            fails.append(f"{len(pile_profiles)} piles drew "
                         f"{len(pile_figures)} detail figures; every pile the "
                         f"analysis solved is drawn")
    planned, drawn = _planned_matches(pile_report, "fem", xlsx=FEM_PILES_XLSX)
    if planned != drawn:
        fails.append(f"the piles report planned {planned} figures and built "
                     f"{drawn}")

    # --- neither member, and neither subsection ------------------------------
    plain = _engine_report("fem")
    for title in ("Reinforcement Forces", "Pile Forces"):
        if _member_section(plain, title) is not None:
            fails.append(f"a run carrying no member was given a {title} "
                         f"subsection")
    if _member_section(_build(), "Reinforcement Forces") is not None:
        fails.append("a limit equilibrium report carries a finite element "
                     "Reinforcement subsection")

    # --- the toggles ---------------------------------------------------------
    off = _engine_report("fem", {"fem_reinforcement": False},
                         xlsx=FEM_REINF_XLSX)
    if _member_section(off, "Reinforcement Forces") is not None:
        fails.append("the reinforcement toggle left the subsection standing")
    no_figures = _engine_report("fem", {"fem_reinforcement_figure": False},
                                xlsx=FEM_REINF_XLSX)
    quiet = _member_section(no_figures, "Reinforcement Forces")
    if quiet is None:
        fails.append("switching the detail figures off took the subsection with "
                     "them")
    elif [b for b in quiet.blocks if b.kind == "figure"]:
        fails.append("the detail figures were drawn with their toggle off")
    else:
        # The terms the figures were read in are still owed to a reader who is
        # now shown none of them.
        left = " ".join(b.text for b in quiet.blocks if b.kind == "prose")
        if "utilization at a point" not in left:
            fails.append(f"with no figure and no table the subsection defines "
                         f"nothing: {left!r}")

    # --- the field the forces are read at ------------------------------------
    #
    # An SSRM run is asked about the mechanism it developed. Hand the same run
    # an at-failure snapshot whose bar forces are twice the converged ones, and
    # the table has to move to it and say which field it read.
    import numpy as np
    snapshot = dict(bundle["solution"])
    snapshot["forces_1d"] = 2.0 * np.asarray(bundle["solution"]["forces_1d"],
                                             dtype=float)
    at_failure = dict(bundle, analysis="ssrm", FS=1.207,
                      failure_solution=snapshot)
    failed = _engine_report("fem", bundle=at_failure, xlsx=FEM_REINF_XLSX)
    sec = _member_section(failed, "Reinforcement Forces")
    if sec is None:
        fails.append("the at-failure run lost its Reinforcement subsection")
    else:
        said = " ".join(b.text for b in sec.blocks if b.kind == "prose")
        if "at failure" not in said:
            fails.append(f"the profiles were read at failure and the section "
                         f"does not say so: {said!r}")
        # Measured on the DRAWN curve, which is what the reader has now that the
        # summary table is gone: the force profile has to be the snapshot's.
        drew = max((abs(v) for v in _drawn_member_forces(
            slope_data, at_failure)), default=0.0)
        if abs(drew - 2.0 * peak) > 0.05 + 0.005 * 2.0 * peak:
            fails.append(f"the at-failure figures peak at {drew:,.1f} where the "
                         f"snapshot carries {2.0 * peak:,.1f}; the converged "
                         f"field was drawn instead")

    # Mutation: the subsection is absent because the model owns no member, not
    # because an option happened to be off. Hand the builder a member for the
    # model that has none, and the absence rule above has to notice.
    import xslope.report as report_mod
    invented = {"kind": "reinforcement", "index": 1, "label": "invented",
                "field_state": "converged", "peak_T": 0.0, "peak_s": 0.0,
                "peak_utilization": 0.0, "status": "within capacity",
                "units": {}}
    saved = report_mod._detail_profiles
    report_mod._detail_profiles = (
        lambda sd, b, kind, *a: [dict(invented)] if kind == "reinforcement"
        else saved(sd, b, kind, *a))
    try:
        mutated = _engine_report("fem", {"fem_reinforcement_figure": False})
        if _member_section(mutated, "Reinforcement Forces") is None:
            fails.append("a member invented for a model with none printed no "
                         "subsection; the absence rule cannot fail")
    finally:
        report_mod._detail_profiles = saved

    # Mutation: the table that remains — the piles' — has teeth. A force column
    # zeroed while the utilization beside it still claims a share of capacity is
    # the defect the table exists to make impossible to hide.
    saved = report_mod._detail_profiles
    report_mod._detail_profiles = (
        lambda sd, b, kind, *a: [dict(p, max_moment=0.0, max_shear=0.0)
                                 for p in saved(sd, b, kind, *a)]
        if kind == "pile" else saved(sd, b, kind, *a))
    try:
        zeroed = _engine_report("fem", {"fem_piles_figure": False},
                                xlsx=FEM_PILES_XLSX)
        sec = _member_section(zeroed, "Pile Forces")
        table = next((b for b in (sec.blocks if sec else [])
                      if b.kind == "table"), None)
        if table is None or not _member_faults(table, pile_profiles,
                                               "Peak moment", "max_moment"):
            fails.append("every peak moment was zeroed and the table still "
                         "agreed with its profiles; the honesty rule has no "
                         "teeth")
    finally:
        report_mod._detail_profiles = saved
    return fails


def test_a_broken_tie_stretch_is_excepted():
    """A line at capacity everywhere along a stretch BUT one point is drawn that
    way: the highlight breaks at the hole.

    The greatest utilization is held over a stretch on most reinforcement lines.
    On the reinforcement sample's own mechanism, line 4 stands at capacity at
    every sample from 1.00 to 19.00 except the one at 5.00, and line 6 from 7.00
    to 19.00 except 15.00 — and the two ends alone read as an unbroken run, which
    is a different bar. The figure is now the only place that stretch is
    reported (the summary table went with the owner's fem_reinforce ruling), so
    what is measured here is the drawn highlight: one thickened run per unbroken
    stretch, and no run that spans a sample the line drops below capacity at.

    Read off the shipped at-failure snapshot, which is the field the deliverable
    reports and the one the holes are in.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    import numpy as np
    from matplotlib.figure import Figure as MplFigure
    from xslope.plot_fem_details import plot_detail
    from xslope.report import _detail_profiles

    slope_data, solutions = _restored(FEM_REINF_XLSX)
    bundle = solutions.get("fem")
    if not bundle:
        return ["the reinforcement sample ships no solved run"]
    profiles = _detail_profiles(slope_data, bundle, "reinforcement")
    broken = {p["label"]: [float(v) for v in p.get("peak_gap_s", [])]
              for p in profiles if len(p.get("peak_gap_s", []))}
    if not broken:
        return fails + ["no line on the shipped mechanism holds its greatest "
                        "utilization over a BROKEN stretch, so the exception "
                        "this check is about could not arise"]

    def highlight_runs(profile):
        """The x-ranges the figure thickens on one line."""
        fig = MplFigure(figsize=(6.5, 4.5))
        with contextlib.redirect_stdout(io.StringIO()):
            plot_detail(profile, fig=fig)
        runs = []
        for ax in fig.axes:
            for line in ax.lines:
                if line.get_gid() == "DETAIL_PEAK_SPAN":
                    xs = [float(v) for v in line.get_xdata()]
                    if xs:
                        runs.append((min(xs), max(xs)))
        return sorted(runs)

    for profile in profiles:
        label = profile["label"]
        runs = highlight_runs(profile)
        gaps = broken.get(label)
        if gaps is None:
            # An unbroken stretch is one run, or none where the peak is a point.
            if len(runs) > 1:
                fails.append(f"{label}: the stretch is unbroken and the figure "
                             f"draws {len(runs)} separate runs: {runs}")
            continue
        if not runs:
            fails.append(f"{label}: the line holds its greatest utilization "
                         f"over a stretch and the figure highlights none of it")
            continue
        # No run may span a position the line drops below capacity at.
        for gap in gaps:
            spanning = [r for r in runs if r[0] < gap < r[1]]
            if spanning:
                fails.append(f"{label}: a highlighted run {spanning} spans "
                             f"{gap:.2f}, where the line stands below capacity")
        if len(runs) < len(gaps) + 1:
            fails.append(f"{label}: {len(runs)} highlighted run(s) for a stretch "
                         f"broken at {len(gaps)} position(s)")

    # Mutation: the contiguous claim restored — the highlight drawn from the two
    # ends alone, exactly as a chord across the dip. Every broken line has to be
    # caught.
    holed = next(p for p in profiles if p["label"] in broken)
    tied = np.asarray(holed.get("peak_indices", []), dtype=int)
    chord = dict(holed, peak_indices=np.arange(int(tied.min()),
                                               int(tied.max()) + 1))
    runs = highlight_runs(chord)
    gaps = broken[holed["label"]]
    if not any(r[0] < gaps[0] < r[1] for r in runs):
        fails.append(f"a highlight drawn from the two ends alone still broke at "
                     f"{gaps[0]:.2f}: {runs} — the measurement cannot fail")
    return fails


def test_a_solution_without_member_forces_says_so():
    """A run reloaded from companions that carry no member forces reports none.

    A field saved without its reinf sidecar carries no forces_1d at all.
    fem_details substituted zeros for the array it could not find, and the
    subsection printed six rows of 0.0 lb at 0% of capacity, six flat-zero detail
    figures, and the sentence "The forces are read from the last converged field"
    — over a field in which nothing of the kind was recorded. The saved solution
    now says so in one sentence, and the live-solve path is untouched.

    The reinforcement sample used to BE this case, because its shipped companions
    were written without the bar forces. They are written with them now, so the
    case is built here instead: the model and its companions are copied, and the
    two force files are left behind.
    """
    fails = []
    import numpy as np
    from xslope.report import _member_forces_recorded

    # --- the restored run: no forces recorded, and none reported -------------
    stem, slope_data, solutions = _restored_without_member_forces()
    restored = solutions.get("fem")
    if not restored:
        fails.append("the copied reinforcement model carries no finite element "
                     "companions, so the case this check is about cannot arise")
        return fails
    if "forces_1d" in (restored.get("solution") or {}):
        fails.append("the restored solution carries forces_1d; the fixture no "
                     "longer exercises the absence it was chosen for")
    if _member_forces_recorded(restored, "reinforcement"):
        fails.append("a solution carrying no forces_1d is read as recording "
                     "member forces")
    if not _detail_profiles_exist(slope_data, restored, "reinforcement"):
        fails.append("the restored run owns no reinforcement profile, so the "
                     "subsection would be absent for the wrong reason")

    report = _built_report(slope_data, solutions,
                           {"input_path": f"{stem}.xlsx", "lem": False,
                            "pd_figure": False})
    sec = _member_section(report, "Reinforcement Forces")
    if sec is None:
        fails.append(f"the restored run lost its Reinforcement subsection "
                     f"entirely: {_titles(report)}")
        return fails
    if [b for b in sec.blocks if b.kind == "table"]:
        fails.append("a solution recording no member forces still printed a "
                     "forces table")
    if [b for b in sec.blocks if b.kind == "figure"]:
        fails.append("a solution recording no member forces still drew detail "
                     "figures")
    said = " ".join(b.text for b in sec.blocks if b.kind == "prose")
    if "records no forces" not in said:
        fails.append(f"the subsection does not say the solution records no "
                     f"forces: {said!r}")
    if "read from" in said:
        fails.append(f"the subsection still claims the forces were read from a "
                     f"field: {said!r}")
    planned, drawn = _planned_matches(report, "fem", xlsx=f"{stem}.xlsx",
                                      bundle=restored, slope_data=slope_data)
    if planned != drawn:
        fails.append(f"the restored reinforcement report planned {planned} "
                     f"figures and built {drawn}")

    # --- the live solve is untouched -----------------------------------------
    solved_data, solved = _fem_1d_bundle(FEM_REINF_XLSX)
    if not _member_forces_recorded(solved, "reinforcement"):
        fails.append("a freshly solved run is read as recording no member "
                     "forces; the predicate refuses real results")
    live = _member_section(_engine_report("fem", xlsx=FEM_REINF_XLSX),
                           "Reinforcement Forces")
    if live is None or not [b for b in live.blocks if b.kind == "figure"]:
        fails.append("the live-solve path lost its detail figures")
    elif "records no forces" in " ".join(b.text for b in live.blocks
                                         if b.kind == "prose"):
        fails.append("a run that DID record member forces is refused")

    # --- the same rule for piles, on the piles' own arrays -------------------
    pile_data, pile_bundle = _fem_1d_bundle(FEM_PILES_XLSX)
    if not _member_forces_recorded(pile_bundle, "pile"):
        fails.append("the solved pile run is read as recording no forces")
    stripped = dict(pile_bundle, solution={
        k: v for k, v in pile_bundle["solution"].items()
        if not k.startswith("forces_pile")})
    if _member_forces_recorded(stripped, "pile"):
        fails.append("a solution with every pile force array removed is still "
                     "read as recording them")
    quiet = _member_section(_engine_report("fem", bundle=stripped,
                                           xlsx=FEM_PILES_XLSX), "Pile Forces")
    if quiet is None:
        fails.append("stripping the pile forces removed the Piles subsection")
    else:
        if [b for b in quiet.blocks if b.kind == "table"]:
            fails.append("a solution recording no pile forces printed a table")
        if "records no forces" not in " ".join(b.text for b in quiet.blocks
                                               if b.kind == "prose"):
            fails.append("the Piles subsection does not say the solution "
                         "records no forces")
    return fails


def _member_overlay_words(bundle):
    """Every word the finite element results figure prints about the members it
    draws: the legend entries, and the labels of the colorbars beside them.

    Rendered through :func:`plot_fem_results` on the panel list the report
    prints, so what is asserted is what reaches the page rather than what one
    helper returns in isolation.
    """
    import matplotlib.figure as mplfig
    from xslope.plot_fem import plot_fem_results
    from xslope.report import FEM_PANELS

    fig = mplfig.Figure(figsize=(7.0, 8.0))
    with contextlib.redirect_stdout(io.StringIO()):
        plot_fem_results(bundle["fem_data"], bundle["solution"],
                         plot_type=[p for p, _c, _s in FEM_PANELS], fig=fig,
                         failure_solution=bundle.get("failure_solution"),
                         show_title=False)
    legend = []
    for ax in fig.axes:
        legend += ax.get_legend_handles_labels()[1]
    bars = [l for l in (ax.get_ylabel() or ax.get_xlabel() for ax in fig.axes) if l]
    return legend, bars


def test_a_member_overlay_claims_no_force_it_was_not_given():
    """The members drawn over the results figure are classified by force only
    where a force was solved for them.

    ``plot_reinforcement_forces`` substituted a zero array for the ``forces_1d``
    it could not find, and zero classifies: every one of the six bars fell into
    the no-tension branch and the figure legend read "Inactive (no tension)" —
    six green bars making a measurement, beside the sentence in the same
    section saying the solution records no forces in the lines. Where the array
    is absent the members are now drawn as geometry, in one neutral color,
    named by kind. Piles are judged on their own arrays, and a live solve is
    classified exactly as before.
    """
    fails = []
    from xslope.plot_fem import (MEMBER_FORCE_LEGEND_LABELS,
                                 REINFORCEMENT_GEOMETRY_LABEL)
    import xslope.plot_fem as plot_fem_mod

    def claims(legend, bars):
        """The force claims on a figure: classification legend entries, and the
        force colorbars, which are equally a measurement."""
        return ([l for l in legend if l in MEMBER_FORCE_LEGEND_LABELS]
                + [b for b in bars if "Force" in b])

    # --- the restored run: no force array, so no force claim -----------------
    _stem, _sd, solutions = _restored_without_member_forces()
    restored = solutions.get("fem")
    if not restored:
        fails.append("the copied reinforcement model carries no finite element "
                     "companions, so the case this check is about cannot arise")
        return fails
    if "forces_1d" in (restored.get("solution") or {}):
        fails.append("the restored solution carries forces_1d; the fixture no "
                     "longer exercises the absence it was chosen for")
    legend, bars = _member_overlay_words(restored)
    said = claims(legend, bars)
    if said:
        fails.append(f"a solution carrying no member forces draws a figure "
                     f"claiming {said}")
    if REINFORCEMENT_GEOMETRY_LABEL not in legend:
        fails.append(f"six reinforcement lines are drawn and the legend does "
                     f"not name them: {legend}")

    # --- the live solve: classified, as before -------------------------------
    _sd, solved = _fem_1d_bundle(FEM_REINF_XLSX)
    live_legend, live_bars = _member_overlay_words(solved)
    if not claims(live_legend, live_bars):
        fails.append(f"a run that solved its bar forces classifies none of "
                     f"them: legend {live_legend}, colorbars {live_bars}")
    if REINFORCEMENT_GEOMETRY_LABEL in live_legend:
        fails.append("a run that solved its bar forces draws them as bare "
                     "geometry")

    # Mutation, absent -> present: the zero substitution restored. The gate is
    # the predicate, so answering it yes over a solution with no array is the
    # defect exactly.
    saved = plot_fem_mod._member_forces_in_solution
    plot_fem_mod._member_forces_in_solution = lambda solution, key: True
    try:
        legend, bars = _member_overlay_words(restored)
        if not claims(legend, bars):
            fails.append("a force-less solution read as carrying forces still "
                         "classified nothing; the absence rule cannot fail")
    finally:
        plot_fem_mod._member_forces_in_solution = saved

    # Mutation, present -> absent: a solved run drawn as geometry. The live
    # figure has to notice its classifications are gone.
    plot_fem_mod._member_forces_in_solution = lambda solution, key: False
    try:
        legend, bars = _member_overlay_words(solved)
        if claims(legend, bars):
            fails.append("a solved run read as carrying no forces still "
                         "classified its bars; the live path is not pinned")
        if REINFORCEMENT_GEOMETRY_LABEL not in legend:
            fails.append("the geometry-only draw names no member kind")
    finally:
        plot_fem_mod._member_forces_in_solution = saved

    # --- piles are judged on the pile arrays, not the bar array --------------
    _sd, piles = _fem_1d_bundle(FEM_PILES_XLSX)
    stripped = dict(piles, solution={
        k: v for k, v in piles["solution"].items()
        if not k.startswith("forces_pile")})
    legend, bars = _member_overlay_words(stripped)
    if [b for b in bars if "Pile" in b]:
        fails.append(f"a solution recording no pile forces draws a pile force "
                     f"colorbar: {bars}")
    if plot_fem_mod.PILE_GEOMETRY_LABEL not in legend:
        fails.append(f"the piles vanished from the figure instead of being "
                     f"drawn plain: {legend}")
    return fails


#: The overlay's magenta, read off the drawn artists rather than restated: the
#: colour ``plot_reinforcement_forces`` draws its "At residual" class in.
_RESIDUAL_RGBA = (1.0, 0.0, 1.0, 0.9)


def _overlay_render(bundle, field_state, solution=None):
    """The finite element results figure for one bundle at one field state.

    ``solution`` replaces the bundle's converged field, which is how a check
    puts a state the sample never reached in front of the drawing code.
    """
    import matplotlib.figure as mplfig
    from xslope.plot_fem import plot_fem_results
    from xslope.report import FEM_PANELS

    fig = mplfig.Figure(figsize=(7.0, 8.0))
    with contextlib.redirect_stdout(io.StringIO()):
        plot_fem_results(bundle["fem_data"],
                         bundle["solution"] if solution is None else solution,
                         plot_type=[p for p, _c, _s in FEM_PANELS], fig=fig,
                         failure_solution=bundle.get("failure_solution"),
                         field_state=field_state, show_title=False)
    return fig


def _overlay_legend(fig):
    labels = []
    for ax in fig.axes:
        labels += ax.get_legend_handles_labels()[1]
    return labels


def _residual_segments(fig):
    """How many member elements the figure drew as "at residual" — counted off
    the magenta collections themselves, so the count is what reaches the page."""
    import numpy as np
    from matplotlib.collections import LineCollection

    n = 0
    for ax in fig.axes:
        for coll in ax.collections:
            if not isinstance(coll, LineCollection):
                continue
            colors = np.asarray(coll.get_colors())
            if len(colors) == 1 and np.allclose(colors[0], _RESIDUAL_RGBA):
                n += len(coll.get_segments())
    return n


def test_the_deformed_members_are_a_different_color_from_the_original():
    """On the deformation panel, a member's original and deformed positions are
    drawn in colors a reader can tell apart.

    Both pile overlays were green — the same green, one over the other — so at
    any exaggeration the panel showed one pile and the deformation of the piles
    could not be read at all. Norm: "a terrible idea."

    The convention the panel already uses for the reinforcement is the one the
    piles now follow: the ORIGINAL configuration is the muted reference (the
    reinforcement's gray, the pile's own green, the dashed gray mesh outline),
    and the DEFORMED configuration is red. Red on this panel therefore means
    "this is the deformed position" and not "this is a reinforcement line" —
    the reinforcement's geometry color elsewhere is gray, never red — so a model
    carrying both kinds reads consistently, with the two kinds held apart by
    line weight and by their legend entries.
    """
    fails = []
    from matplotlib.colors import to_rgba
    from xslope.plot_fem import (PILE_COLOR, PILE_DEFORMED_COLOR,
                                 plot_fem_results)

    if to_rgba(PILE_COLOR)[:3] == to_rgba(PILE_DEFORMED_COLOR)[:3]:
        fails.append(f"the original and deformed pile colors are the same "
                     f"({PILE_COLOR!r}, {PILE_DEFORMED_COLOR!r})")

    for label, xlsx in (("piles", FEM_PILES_XLSX),
                        ("reinforcement", FEM_REINF_XLSX)):
        _sd, bundle = _fem_1d_bundle(xlsx)
        import matplotlib.figure as _mplfig
        fig = _mplfig.Figure(figsize=(8.0, 4.5))
        with contextlib.redirect_stdout(io.StringIO()):
            plot_fem_results(bundle["fem_data"], bundle["solution"],
                             plot_type=["deformation"], fig=fig,
                             show_title=False)
        ax = fig.axes[0]
        gid = "PILES" if label == "piles" else "REINFORCEMENT"
        drawn = [c for c in ax.collections if c.get_gid() == gid]
        if len(drawn) != 2:
            fails.append(f"{label}: the deformation panel draws {len(drawn)} "
                         f"{gid} overlays, not the original and the deformed")
            continue
        colors = [tuple(c.get_colors()[0][:3]) for c in drawn]
        if colors[0] == colors[1]:
            fails.append(f"{label}: the original and deformed overlays are both "
                         f"drawn in {colors[0]}, so neither can be read")
        # Far enough apart to survive printing: not two shades of one hue.
        if max(abs(a - b) for a, b in zip(*colors)) < 0.25:
            fails.append(f"{label}: the two overlays differ by less than a "
                         f"quarter of a channel: {colors}")
        named = set(ax.get_legend_handles_labels()[1])
        kind = "Pile" if label == "piles" else "Reinforcement"
        for want in (f"Original {kind}", f"Deformed {kind}"):
            if want not in named:
                fails.append(f"{label}: the legend does not name {want!r}: "
                             f"{sorted(named)}")
    return fails


def test_the_member_overlay_marks_the_state_the_solver_recorded():
    """The overlay's "At residual" is the softening latch, not the yield latch.

    The solver keeps two: ``failed_1d_elements``, set for reporting the moment
    an element reaches the capacity it was given, and ``softened_1d_elements``,
    set when an element DROPS off that capacity onto its residual. The overlay
    classified from the first, so the reinforcement sample's at-failure figure —
    where forty-five elements are holding their full 800 and not one has
    softened — printed "At residual (Tres)" in magenta over most of the bars.
    They were at Tmax. Softening is the mark; an element at its peak rides the
    force ramp at the colour its force earns.

    Pinned in both fields of the sample, because they differ: the converged
    field has five back-end elements that softened with nothing left to hold and
    are marked pulled out, and the at-failure field has none.
    """
    fails = []
    import numpy as np
    import xslope.plot_fem as plot_fem_mod

    _sd, solutions = _restored(FEM_REINF_XLSX)
    bundle = solutions.get("fem")
    if not bundle:
        return ["the reinforcement sample ships no finite element solution"]
    fem_data = bundle["fem_data"]
    converged, failure = bundle["solution"], bundle.get("failure_solution")
    if failure is None:
        return ["the reinforcement sample carries no at-failure snapshot, so "
                "the two fields this check compares are one field"]

    n_1d = len(fem_data["elements_1d"])
    t_res = np.asarray(fem_data["t_res_by_1d_elem"], dtype=float)

    def counts(solution):
        """(softened onto a residual, softened with nothing left) — what the
        figure should mark, read off the solver's own arrays."""
        soft = np.asarray(solution.get("softened_1d_elements",
                                       np.zeros(n_1d, bool)), dtype=bool)
        force = np.asarray(solution.get("forces_1d", np.zeros(n_1d)), dtype=float)
        pulled = soft & (t_res < 1e-6) & (force < 1e-6)
        return int((soft & ~pulled).sum()), int(pulled.sum())

    # The sample's own two fields.
    for state, solution in (("failure", failure), ("converged", converged)):
        residual, pulled = counts(solution)
        fig = _overlay_render(bundle, state)
        legend = _overlay_legend(fig)
        drawn = _residual_segments(fig)
        if drawn != residual:
            fails.append(f"at the {state} field {residual} element(s) softened "
                         f"onto a residual and the figure marks {drawn}")
        if ("At residual (Tres)" in legend) != (residual > 0):
            fails.append(f"at the {state} field the legend {'names' if residual else 'does not name'} "
                         f"a residual state that {residual} element(s) are in: {legend}")
        if ("Pulled out" in legend) != (pulled > 0):
            fails.append(f"at the {state} field the legend disagrees with the "
                         f"{pulled} pulled-out element(s): {legend}")

    # The class itself, put in front of the drawing code: one element softened
    # onto a residual it can still hold is marked, and one that has merely
    # reached its peak is not.
    soft = np.zeros(n_1d, dtype=bool)
    holds = np.where((t_res > 1e-6)
                     & (np.asarray(failure["forces_1d"]) > 1e-6))[0]
    if len(holds) < 2:
        fails.append("the sample has no element carrying force against a "
                     "residual capacity, so the marked class cannot be pinned")
        return fails
    soft[holds[0]] = True
    synthetic = dict(failure, softened_1d_elements=soft,
                     failed_1d_elements=np.ones(n_1d, dtype=bool))
    fig = _overlay_render(dict(bundle, failure_solution=None), "converged",
                          solution=synthetic)
    if _residual_segments(fig) != 1:
        fails.append(f"one softened element with a residual left to hold is "
                     f"drawn as {_residual_segments(fig)} residual segments, "
                     f"with every other element latched as yielded")

    # Mutation: the yield latch put back where the softening latch belongs. The
    # at-failure field is the case it was wrong on — nothing there has softened,
    # and everything has yielded.
    real = plot_fem_mod._elem_flags
    plot_fem_mod._elem_flags = (
        lambda solution, key, n: real(solution, "failed_1d_elements", n))
    try:
        if not _residual_segments(_overlay_render(bundle, "failure")):
            fails.append("classifying from the yield latch marks nothing at "
                         "failure either; the reclassification is not pinned")
    finally:
        plot_fem_mod._elem_flags = real
    return fails


def test_the_member_subsections_locate_their_members():
    """Each member subsection opens on a figure of where its members are.

    A row of the forces table and a detail figure both name a member, and
    nothing in the report put that name anywhere on the slope — Norm's ask. One
    locator per kind of member, drawn before the details it locates, cited, and
    counted in the figures the caller is promised.
    """
    fails = []
    from xslope.report import DETAIL_KINDS

    for label, xlsx, kind in (("reinforcement", FEM_REINF_XLSX, "reinforcement"),
                              ("piles", FEM_PILES_XLSX, "pile")):
        report = _engine_report("fem", xlsx=xlsx)
        title = DETAIL_KINDS[kind]["title"]
        sub = next((c for s in report.sections for r in s.children
                    for c in r.children if c.title == title), None)
        if sub is None:
            fails.append(f"{label}: no {title} subsection")
            continue
        figures = [b for b in sub.blocks if b.kind == "figure"]
        if not figures:
            fails.append(f"{label}: the subsection carries no figure at all")
            continue
        first = figures[0]
        if not first.source.endswith(f"{kind} map"):
            fails.append(f"{label}: the subsection opens on {first.source!r}, "
                         f"not on the locator")
        if len(figures) < 2:
            fails.append(f"{label}: the locator is the only figure, so it "
                         f"locates nothing")
        cites = [b.text for b in sub.blocks
                 if b.kind == "prose" and f"Figure {first.number}" in b.text]
        if not cites:
            fails.append(f"{label}: nothing cites the locator (Figure "
                         f"{first.number})")
        elif "in their positions on the section" not in cites[0]:
            fails.append(f"{label}: the locator's sentence does not say what it "
                         f"shows: {cites[0]!r}")

        # The drawing names every member the table has a row for, so a row and a
        # place on the slope are the same member.
        from xslope.fem_details import member_lines
        slope_data, bundle = _fem_any_bundle(xlsx)
        named = _map_labels(slope_data, bundle, kind)
        want = {m["label"] for m in
                member_lines(bundle["fem_data"], slope_data, kind)}
        if named != want:
            fails.append(f"{label}: the locator names {sorted(named)} for the "
                         f"{len(want)} member(s) {sorted(want)}")

        # And the caller's count includes it.
        planned, drawn = _planned_matches(report, "fem", xlsx=xlsx)
        if planned != drawn:
            fails.append(f"{label}: the caller is promised {planned} figures "
                         f"and the build produced {drawn}")

        # The highlight the Studio inset selects with picks out one member and
        # leaves the rest as they were — the same drawing, one member marked.
        picked = member_lines(bundle["fem_data"], slope_data, kind)[-1]
        if _map_labels(slope_data, bundle, kind,
                       highlight=picked["index"]) != want:
            fails.append(f"{label}: highlighting a member changes which members "
                         f"the drawing names")
    return fails


def _map_labels(slope_data, bundle, kind, highlight=None):
    """Every name the member locator draws, read off the rendered figure."""
    import matplotlib.figure as mplfig
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    from xslope.plot_fem_details import plot_member_map

    fig = mplfig.Figure(figsize=(6.5, 3.5))
    FigureCanvasAgg(fig)
    with contextlib.redirect_stdout(io.StringIO()):
        plot_member_map(bundle["fem_data"], slope_data, kind,
                        highlight=highlight, fig=fig)
    return {t.get_text() for ax in fig.axes for t in ax.texts}


def test_a_definition_follows_the_thing_it_defines():
    """No paragraph opens on a definition of a term nothing has used yet.

    The member subsections led with "The utilization at a point is the force
    mobilized there against the capacity available there" — a definition set in
    front of a reader who had not yet been told that anything was being measured,
    let alone where to find it. Norm, on the fem_piles review: refer to the figure
    or the table first, and define the term later in the paragraph. Ordering is a
    review criterion, so it is a check.
    """
    fails = []
    from xslope.report import DETAIL_KINDS, UTILIZATION_DEFINED

    for label, xlsx, kind in (("reinforcement", FEM_REINF_XLSX, "reinforcement"),
                              ("piles", FEM_PILES_XLSX, "pile")):
        report = _engine_report("fem", xlsx=xlsx)
        title = DETAIL_KINDS[kind]["title"]
        sub = next((c for s in report.sections for r in s.children
                    for c in r.children if c.title == title), None)
        if sub is None:
            fails.append(f"{label}: no {title} subsection")
            continue
        defining = [b.text for b in sub.blocks
                    if b.kind == "prose" and UTILIZATION_DEFINED[kind] in b.text]
        if len(defining) != 1:
            fails.append(f"{label}: utilization is defined in {len(defining)} "
                         f"paragraphs of its subsection")
            continue
        para = defining[0]
        at = para.index(UTILIZATION_DEFINED[kind])
        if at == 0:
            fails.append(f"{label}: the paragraph opens on the definition of "
                         f"utilization, with nothing before it to define: "
                         f"{para[:120]!r}")
            continue
        # And what stands before it is the reference the definition is for: the
        # table where there is one, and the figures the term is read off where
        # there is not (the reinforcement, whose summary table the owner
        # dropped).
        opening = para[:at]
        if "Table" not in opening and "Figure" not in opening:
            fails.append(f"{label}: the definition follows neither the table nor "
                         f"the figures it is read in: {opening!r}")
        if "utilization" not in opening.lower():
            fails.append(f"{label}: the word is defined before it is used: "
                         f"{opening!r}")
    return fails


def test_the_detail_figures_are_explained():
    """The member detail figures are explained panel by panel.

    They are the densest figures the report prints — four profiles on a shared
    depth axis for a pile, a force curve against a shaped capacity envelope with
    a bond-transfer trace beneath it for a reinforcement line — and the sentence
    citing them named the quantities and stopped. Norm, on the fem_piles review:
    educate the reader. Explaining what an engineering figure PLOTS is not the
    self-narration the prose rule bans; writing about the document is.
    """
    fails = []
    from xslope.report import DETAIL_FIGURE_READING, DETAIL_KINDS

    # What each explanation has to account for: the panels the plotter really
    # draws, and the marks a reader cannot name from the axis labels.
    for kind, wanted in (
            ("pile", ("four panels", "share one depth axis", "head at the top",
                      "Vcap", "Mcap", "largest moment", "Ito and Matsui")),
            ("reinforcement", ("upper panel", "lower panel", "envelope",
                               "pullout length", "T_max", "dT/ds",
                               "how fast the force in the panel above is "
                               "building"))):
        for word in wanted:
            if word not in DETAIL_FIGURE_READING[kind]:
                fails.append(f"the {kind} figure explanation does not account "
                             f"for {word!r}")
        # And it explains the FIGURE, not the report. A sentence whose subject is
        # the document is the defect the prose rule exists for.
        for banned in ("this section", "the report", "the reader", "below the",
                       "as noted"):
            if banned in DETAIL_FIGURE_READING[kind].lower():
                fails.append(f"the {kind} figure explanation writes about the "
                             f"document: {banned!r}")

    # It reaches the page, once, under the figures it explains.
    for label, xlsx, kind in (("reinforcement", FEM_REINF_XLSX, "reinforcement"),
                              ("piles", FEM_PILES_XLSX, "pile")):
        report = _engine_report("fem", xlsx=xlsx)
        title = DETAIL_KINDS[kind]["title"]
        sub = next((c for s in report.sections for r in s.children
                    for c in r.children if c.title == title), None)
        if sub is None:
            fails.append(f"{label}: no {title} subsection")
            continue
        blocks = sub.blocks
        carrying = [i for i, b in enumerate(blocks)
                    if b.kind == "prose" and DETAIL_FIGURE_READING[kind] in b.text]
        if len(carrying) != 1:
            fails.append(f"{label}: the figures are explained in {len(carrying)} "
                         f"paragraphs")
            continue
        after = [b for b in blocks[carrying[0] + 1:] if b.kind == "figure"]
        if not after:
            fails.append(f"{label}: the explanation stands under no figure")
        # The sentence it opens with agrees with its own subject: a single
        # quantity drawn over six figures is drawn, not are drawn.
        opening = blocks[carrying[0]].text.split(". ")[0]
        want = "is drawn" if kind == "reinforcement" else "are drawn"
        if want not in opening:
            fails.append(f"{label}: the citing sentence does not agree with its "
                         f"subject ({want!r} expected): {opening!r}")
    return fails


def test_the_member_terms_are_defined_where_they_are_used():
    """Utilization and the band are defined on the page that uses them.

    A member subsection printed a utilization column, a State column reading
    "at capacity", and figures carrying a shaded band, and defined none of the
    three. Norm, reading it: "A good clear explanation would help tremendously."

    Each definition is written once per report — the term is one term, however
    many runs and however many kinds of member the report describes — and the
    state language matches the classes the solver actually records: at capacity
    is a member holding its full capacity, softened is one that has dropped onto
    its residual, pulled out is one with no residual left carrying nothing.

    The band is the one term that is really two. On a strength reduction run it
    marks the mechanism; on a run that converged under gravity there is no
    mechanism, and the shared sentence had a stable section's members crossed by
    a failure the analysis never found.
    """
    fails = []
    from xslope.report import BAND_DEFINED, UTILIZATION_DEFINED

    # What each definition has to say, as against merely being present: the
    # page is compared to the constant below, so the constant is compared to
    # the fact it exists to state.
    for kind, wanted in (
            ("reinforcement", ("mobilized", "capacity available",
                               "ramps up from each end",
                               "not the point of greatest force")),
            ("pile", ("mobilized against the capacity", "Vcap", "Mcap"))):
        for word in wanted:
            if word not in UTILIZATION_DEFINED[kind]:
                fails.append(f"the {kind} definition of utilization does not "
                             f"say {word!r}: {UTILIZATION_DEFINED[kind]!r}")

    # The band means one thing on a run that developed a mechanism and another
    # on one that converged under gravity, and the two definitions are held
    # apart: the converged one cannot claim a failure the analysis never found.
    # And each is written for the ARTIST the figures carry — a shaded stretch, or
    # the dashed line a crossing confined to one element is drawn as. The owner,
    # reading fem_piles: "I don't see a band on the figures. Just a dashed line.
    # Confusing." A sentence describing a mark the figure does not draw is the
    # defect, whichever way it is cured.
    for word in ("failure mechanism passes through", "shear strain field"):
        if word not in BAND_DEFINED[("failure", "band")]:
            fails.append(f"the at-failure band definition does not say {word!r}: "
                         f"{BAND_DEFINED[('failure', 'band')]!r}")
    if "computed shear strain concentrates" not in \
            BAND_DEFINED[("converged", "band")]:
        fails.append(f"the converged band definition does not say where the "
                     f"strain concentrates: "
                     f"{BAND_DEFINED[('converged', 'band')]!r}")
    if "failure" in BAND_DEFINED[("converged", "band")]:
        fails.append(f"a run that reached no failure is told about one: "
                     f"{BAND_DEFINED[('converged', 'band')]!r}")
    for state in ("failure", "converged"):
        if "shaded" not in BAND_DEFINED[(state, "band")]:
            fails.append(f"the {state} band definition does not say the mark is "
                         f"shaded: {BAND_DEFINED[(state, 'band')]!r}")
        if "dashed line" not in BAND_DEFINED[(state, "line")]:
            fails.append(f"the {state} single-element definition does not say "
                         f"the mark is a dashed line: "
                         f"{BAND_DEFINED[(state, 'line')]!r}")
        if "band" in BAND_DEFINED[(state, "line")]:
            fails.append(f"a mark drawn as a dashed line is still called a band: "
                         f"{BAND_DEFINED[(state, 'line')]!r}")

    # And the sentence a page carries is the one for the mark that page's figures
    # DRAW, measured on the artists: a shaded span is an axvspan or an axhspan, a
    # single-element crossing a dashed rule.
    import matplotlib
    matplotlib.use("Agg")
    from matplotlib.figure import Figure as MplFigure
    from matplotlib.colors import to_rgb
    from xslope.plot_fem_details import C_BAND, plot_detail
    from xslope.report import _band_artist, _detail_profiles

    for label, xlsx, kind in (("reinforcement", FEM_REINF_XLSX, "reinforcement"),
                              ("piles", FEM_PILES_XLSX, "pile")):
        sd, bundle = _fem_1d_bundle(xlsx)
        profiles = _detail_profiles(sd, bundle, kind)
        artist = _band_artist(profiles)
        if not artist:
            continue
        shaded = 0
        for profile in profiles:
            fig = MplFigure(figsize=(6.5, 4.5))
            with contextlib.redirect_stdout(io.StringIO()):
                plot_detail(profile, fig=fig)
            shaded += sum(
                1 for ax in fig.axes for patch in ax.patches
                if patch.get_alpha() is not None
                and 0.0 < patch.get_alpha() < 0.5
                and to_rgb(patch.get_facecolor()[:3]) == to_rgb(C_BAND))
        if artist == "band" and not shaded:
            fails.append(f"{label}: the prose calls the mark a shaded band and "
                         f"the figures shade nothing")
        if artist == "line" and shaded:
            fails.append(f"{label}: the prose calls the mark a dashed line and "
                         f"the figures shade {shaded} span(s)")
        said = " ".join(_prose(_engine_report("fem", xlsx=xlsx)))
        state = "converged"
        if BAND_DEFINED[(state, artist)] not in said:
            fails.append(f"{label}: the page does not carry the sentence for the "
                         f"{artist!r} its figures draw")
        wrong = "line" if artist == "band" else "band"
        if BAND_DEFINED[(state, wrong)] in said:
            fails.append(f"{label}: the page describes a {wrong!r} its figures "
                         f"do not draw")

    # The two sample models are documented from a gravity trial each — no
    # snapshot, no mechanism — so their pages take the converged definition and
    # not the other one.
    for label, xlsx, kind in (("reinforcement", FEM_REINF_XLSX, "reinforcement"),
                              ("piles", FEM_PILES_XLSX, "pile")):
        said = " ".join(_prose(_engine_report("fem", xlsx=xlsx)))
        artist = _band_artist(_detail_profiles(
            *_fem_1d_bundle(xlsx), kind)) or "band"
        for what, phrase in (("utilization", UTILIZATION_DEFINED[kind]),
                             ("the band", BAND_DEFINED[("converged", artist)])):
            n = said.count(phrase)
            if n != 1:
                fails.append(f"{label}: {what} is defined {n} times on the page "
                             f"({phrase[:48]!r}…)")
        if BAND_DEFINED[("failure", artist)] in said:
            fails.append(f"{label}: a converged gravity run's band is called the "
                         f"failure mechanism's")

    # And the run that DID develop one: the shipped strength reduction solution,
    # which is what the reinforcement deliverable is built from.
    _sd, ssrm = _restored(FEM_REINF_XLSX)
    if not ssrm.get("fem"):
        fails.append("the reinforcement sample ships no strength reduction run, "
                     "so the at-failure definition is untested")
    else:
        at_failure = " ".join(_prose(_engine_report(
            "fem", bundle=ssrm["fem"], xlsx=FEM_REINF_XLSX)))
        ssrm_artist = _band_artist(_detail_profiles(
            _sd, ssrm["fem"], "reinforcement", "failure")) or "band"
        if at_failure.count(BAND_DEFINED[("failure", ssrm_artist)]) != 1:
            fails.append(
                f"a run carrying a mechanism defines the failure band "
                f"{at_failure.count(BAND_DEFINED[('failure', ssrm_artist)])} "
                f"times")
        if BAND_DEFINED[("converged", ssrm_artist)] in at_failure:
            fails.append("a run carrying a mechanism calls its band the "
                         "converged field's")
        # A report of both kinds of run defines each band once: the term is two
        # terms, and the first must not silence the second.
        _gd, gravity = _fem_1d_bundle(FEM_REINF_XLSX)
        mixed = " ".join(_prose(_built_report(
            _sd, {"fem": [ssrm["fem"], gravity]},
            {"input_path": FEM_REINF_XLSX, "lem": False, "pd_figure": False})))
        for state, art in (("failure", ssrm_artist),
                           ("converged", _band_artist(_detail_profiles(
                               _sd, gravity, "reinforcement")) or "band")):
            n = mixed.count(BAND_DEFINED[(state, art)])
            if n != 1:
                fails.append(f"a report of a mechanism run and a gravity run "
                             f"defines the {state} band {n} times")

    # The states, where the model declares a residual capacity for them to be
    # reachable — and the definition that they are the softening latch's, not
    # the yield latch's.
    said = " ".join(_prose(_engine_report("fem", xlsx=FEM_REINF_XLSX)))
    for phrase in ("A line reported at capacity is holding the full capacity "
                   "declared for it",
                   "dropped onto its residual capacity is reported as softened",
                   "whose residual is nothing and which now carries nothing as "
                   "pulled out"):
        if phrase not in said:
            fails.append(f"the reinforcement subsection does not say what a "
                         f"state means: {phrase!r} is not on the page")

    # Two runs of the same model: two member subsections, and each definition
    # still written once. The term is one term whatever the report describes.
    slope_data, bundle = _fem_1d_bundle(FEM_REINF_XLSX)
    two_runs = _built_report(
        slope_data, {"fem": [bundle, bundle]},
        {"input_path": FEM_REINF_XLSX, "lem": False, "pd_figure": False})
    subsections = [t for _lvl, t in two_runs.section_titles()
                   if t == "Reinforcement Forces"]
    if len(subsections) < 2:
        fails.append(f"a report of two runs carries {len(subsections)} member "
                     f"subsection(s), so a repeated definition could not arise")
    twice = " ".join(_prose(two_runs))
    drawn = _band_artist(_detail_profiles(slope_data, bundle,
                                          "reinforcement")) or "band"
    for what, phrase in (("utilization", UTILIZATION_DEFINED["reinforcement"]),
                         ("the band", BAND_DEFINED[("converged", drawn)])):
        n = twice.count(phrase)
        if n != 1:
            fails.append(f"a report of two runs defines {what} {n} times")

    # A model whose lines declare no residual reaches none of those states, and
    # says nothing about them.
    plain = copy.deepcopy(slope_data)
    for line in plain.get("reinforcement_lines") or []:
        line["t_res"] = 0.0
    quiet = " ".join(_prose(_built_report(
        plain, {"fem": bundle},
        {"input_path": FEM_REINF_XLSX, "lem": False, "pd_figure": False})))
    if "reported as softened" in quiet:
        fails.append("a model whose lines declare no residual capacity is told "
                     "what softening would mean")
    if UTILIZATION_DEFINED["reinforcement"] not in quiet:
        fails.append("the utilization definition went with the residual "
                     "capacities; it belongs to every member subsection")
    return fails


def _detail_profiles_exist(slope_data, bundle, kind):
    """Whether the run owns member profiles at all — so a subsection's silence
    can be told apart from a model with no member in it."""
    from xslope.report import _detail_profiles
    return bool(_detail_profiles(slope_data, bundle, kind))


def test_engine_sections_follow_their_solutions():
    """Neither section is built without its engine's solution, and each toggle
    removes what it names."""
    fails = []
    from xslope.report import FEM_PANELS, SEEP_PANELS, build_report

    # The LEM sample carries neither a seepage nor a finite element solution, and
    # the default report of it has neither section.
    titles = [t for _lvl, t in _build().section_titles()]
    for gone in ("Seepage Analysis", "Deformation and Strength Reduction",
                 "Deformation Analysis"):
        if gone in titles:
            fails.append(f"a report with no seepage or FEM solution carries "
                         f"{gone!r}")

    # And the same model, reported with the engine's own solution absent from
    # the mapping: the option is on, and there is still no section.
    for engine, heading in (("seep", "Seepage Analysis"),
                            ("fem", "Deformation and Strength Reduction")):
        slope_data, _bundle = (_seep_bundle() if engine == "seep"
                               else _fem_bundle())
        tmp = tempfile.mkdtemp(prefix="xslope_absent_")
        with contextlib.redirect_stdout(io.StringIO()):
            report = build_report(slope_data, {}, {"lem": False,
                                                   "pd_figure": False}, tmp)
        if heading in [t for _lvl, t in report.section_titles()]:
            fails.append(f"{heading!r} was built from a solutions mapping that "
                         f"carries no {engine!r}")

    # Each case names what its option takes out: a whole section, every block of
    # one kind, or the figures with these sources — and nothing else goes with it.
    cases = [
        ("seep", {"seep": False}, "Seepage Analysis", None, ()),
        ("seep", {"seep_materials": False}, None, "table", ()),
        ("seep", {"seep_inputs_figure": False}, None, None, ("seep model",)),
        # One option, both conventions: the pair is one statement of the
        # conductivity model and goes or stays together.
        ("seep", {"seep_kr_figure": False}, None, None,
         ("seep kr", "seep kr_head")),
        ("seep", {"seep_mesh_figure": False}, None, None, ("seepage bc1 mesh",)),
        ("seep", {"seep_flownet": False}, None, None, ("seepage bc1 head",)),
        # One option, the other three fields: they are one reading of the solve
        # beyond its flow net, and the seepage results view offers them together.
        ("seep", {"seep_variable_figures": False}, None, None,
         tuple(f"seepage bc1 {p['variable']}" for p in SEEP_PANELS
               if p["option"] == "seep_variable_figures")),
        ("fem", {"fem": False}, "Deformation and Strength Reduction", None, ()),
        # The loads table has to go with the properties table for the section to
        # carry none: this model is loaded, and its loads are an input of the
        # analysis rather than a property of a material.
        ("fem", {"fem_materials": False, "fem_loads": False}, None, "table", ()),
        ("fem", {"fem_loads": False}, "Loads", None, ()),
        ("fem", {"fem_inputs_figure": False}, None, None, ("fem model",)),
        ("fem", {"fem_mesh_figure": False}, None, None, ("fem mesh",)),
        ("fem", {"fem_figure": False}, None, None,
         tuple(f"fem run1 {panel}" for panel, _c, _s in FEM_PANELS)),
    ]
    heads = {"seep": "Seepage Analysis",
             "fem": "Deformation and Strength Reduction"}
    for engine, options, gone_section, gone_kind, gone_figures in cases:
        full = _engine_report(engine)
        off = _engine_report(engine, options)
        if gone_section is not None:
            if gone_section in _titles(off):
                fails.append(f"{options} left {gone_section!r} in the report")
            continue
        if heads[engine] not in _titles(off):
            fails.append(f"{options} removed the {heads[engine]!r} section, not "
                         f"only what it names")
        if gone_kind is not None:
            under = _blocks_under(off, heads[engine], gone_kind)
            was = _blocks_under(full, heads[engine], gone_kind)
            if not was:
                fails.append(f"the {engine} report has no {gone_kind} to remove, "
                             f"so {options} proves nothing")
            if under:
                fails.append(f"{options} left {len(under)} {gone_kind}(s) in the "
                             f"{heads[engine]} section")
        if gone_figures:
            was = {f.source for f in full.figures()}
            now = {f.source for f in off.figures()}
            for source in gone_figures:
                if source not in was:
                    fails.append(f"the {engine} report draws no {source!r}, so "
                                 f"{options} proves nothing")
                if source in now:
                    fails.append(f"{options} left the {source!r} figure standing")
            if now != was - set(gone_figures):
                fails.append(f"{options} changed the figures to {sorted(now)}, "
                             f"not only by dropping {sorted(gone_figures)}")
        planned, drawn = _planned_matches(off, engine, options)
        if planned != drawn:
            fails.append(f"{options} planned {planned} figures and built {drawn}")
    return fails


def _blocks_under(report, heading, kind):
    """Every block of one kind in the section headed ``heading`` and below it."""
    sec = next((s for s in report.sections if s.title == heading), None)
    if sec is None:
        return []
    return [b for _lvl, node in sec.walk() for b in node.blocks if b.kind == kind]


#: A rapid drawdown analysis: two boundary condition sets, two saved solutions.
RAPID_SEEP_XLSX = os.path.join(_REPO, "docs", "lem", "files",
                               "xslope_earth_dam_rapid.xlsx")

#: A model whose saved seepage solution has no mesh beside it to be placed on.
#:
#: THE MISSING MESH IS THE FIXTURE. xslope_levee1 ships a ``_seep.csv`` with no
#: ``_seep_mesh.json`` next to it, and that is what makes it the only model in the
#: corpus that exercises the "a solution with nothing to place it on" path. A
#: corpus tidy-up that generates the missing mesh — an obvious kindness, since
#: every other seepage model has one — silently retires that path, leaving the
#: check passing on a model that no longer poses the question. Leave the mesh
#: missing, or move this fixture to a model that is deliberately broken.
NOMESH_SEEP_XLSX = os.path.join(_REPO, "docs", "seep", "files",
                                "xslope_levee1.xlsx")


def _shipped_flowrate(path):
    """The total flowrate a saved seepage solution records in its own footer."""
    with open(path) as f:
        for line in f:
            if line.startswith("# Total Flowrate:"):
                return float(line.split(":", 1)[1])
    return None


def _sidecar_copy(stem, tmp, meta_edit=None, drop=()):
    """A copy of a model and its solution sidecars in ``tmp``, for the checks
    that damage one and ask what is made of it. ``meta_edit`` rewrites the FEM
    run metadata where it is given, and ``drop`` names suffixes to leave behind —
    how a check builds a model whose companions are missing a file. Returns the
    copied stem."""
    import glob
    for path in glob.glob(stem + "*"):
        if os.path.isfile(path) and not any(path.endswith(x) for x in drop):
            shutil.copy(path, tmp)
    out = os.path.join(tmp, os.path.basename(stem))
    meta_path = f"{out}_fem_meta.json"
    if meta_edit is not None:
        with open(meta_path) as f:
            meta = json.load(f)
        with open(meta_path, "w") as f:
            json.dump(meta_edit(meta), f)
    return out


def _truncated(path, rows=5):
    """``path`` with its last few NODE rows cut off — a solution that no longer
    has a row per node, which is what a saved solution looks like once its mesh
    has been rebuilt under it.

    The rows are counted past whatever footer the file ends with, and the footer
    is left on. Counting from the end of the file instead cuts the footer and
    nothing else once it is as long as the count: the file keeps its row per node,
    the reader finds no mismatch, and the check passes on a solution it never
    damaged.
    """
    with open(path) as f:
        lines = f.readlines()
    footer = 0
    while footer < len(lines) and lines[-1 - footer].startswith("#"):
        footer += 1
    body = lines[:len(lines) - footer]
    with open(path, "w") as f:
        f.writelines(body[:-rows] + lines[len(lines) - footer:])


def test_sidecars_assemble_the_solutions():
    """A model solved once is reported without solving it again: the helper reads
    every engine solution saved beside it back into the mapping the report
    consumes, and nothing it returns was computed here.
    """
    fails = []
    from xslope.report import seep_bundles

    # The dam: one boundary condition set, and the flow the saved solution
    # states is the flow the restored bundle carries.
    _sd, solutions = _restored(SEEP_XLSX)
    if set(solutions) != {"seep"}:
        fails.append(f"the dam restored {sorted(solutions)}, expected seepage "
                     f"alone")
    bundles = seep_bundles(solutions)
    shipped = _shipped_flowrate(f"{os.path.splitext(SEEP_XLSX)[0]}_seep.csv")
    got = bundles[0]["solution"]["flowrate"] if bundles else None
    if shipped is None:
        fails.append("the dam's saved seepage solution records no flowrate, so "
                     "there is nothing to restore it against")
    elif got != shipped:
        fails.append(f"the restored seepage solution flows {got}, and the file "
                     f"beside the model says {shipped}")
    if bundles and set(bundles[0]) != {"seep_data", "solution", "options"}:
        fails.append(f"the restored seepage bundle is {sorted(bundles[0])}, not "
                     f"the shape the report reads")

    # Rapid drawdown: two sets solved, two solutions saved, and they come back in
    # the order the section documents them.
    _sd, solutions = _restored(RAPID_SEEP_XLSX)
    bcs = [b["options"]["bc"] for b in seep_bundles(solutions)]
    if bcs != [1, 2]:
        fails.append(f"the rapid drawdown model restored boundary condition "
                     f"sets {bcs}, expected [1, 2]")

    # The strength reduction run: its factor of safety and what it was.
    _sd, solutions = _restored(FEM_XLSX)
    fem = solutions.get("fem")
    if fem is None:
        fails.append("the strength reduction run saved beside the FEM model was "
                     "not restored")
    else:
        if fem["analysis"] != "ssrm":
            fails.append(f"the restored run calls itself {fem['analysis']!r}, "
                         f"not a strength reduction run")
        if not (1.0 < (fem["FS"] or 0) < 2.0):
            fails.append(f"the restored factor of safety is {fem['FS']}")

    # A model that was never solved has nothing beside it, and asking is not an
    # error — the caller gets an empty mapping and reports what it has.
    notes = []
    _sd, solutions = _restored(REINF_XLSX, notes)
    if solutions != {}:
        fails.append(f"a model with no saved solutions restored "
                     f"{sorted(solutions)}")
    if notes:
        fails.append(f"a model with no saved solutions was faulted for it: "
                     f"{notes}")
    return fails


def test_unusable_sidecars_are_reported_not_raised():
    """A saved solution that no longer fits its model costs its own section and
    nothing else: the engine is left out, the caller is told which file and why,
    and the call returns.
    """
    fails = []

    # Stale: the solution was saved against a mesh that has since been rebuilt,
    # so its node count no longer matches. The message names both counts.
    with tempfile.TemporaryDirectory() as tmp:
        stem = _sidecar_copy(os.path.splitext(FEM_XLSX)[0], tmp)
        _truncated(f"{stem}_fem_nodes.csv")
        notes = []
        _sd, solutions = _restored(f"{stem}.xlsx", notes)
        if "fem" in solutions:
            fails.append("a finite element solution saved against a different "
                         "mesh was restored anyway")
        if not any(os.path.basename(stem) in n and "does not match this mesh" in n
                   for n in notes):
            fails.append(f"a stale finite element solution was passed over "
                         f"without saying why: {notes}")

    # The same for a seepage solution, and the reason is the node count again.
    with tempfile.TemporaryDirectory() as tmp:
        stem = _sidecar_copy(os.path.splitext(SEEP_XLSX)[0], tmp)
        _truncated(f"{stem}_seep.csv")
        notes = []
        _sd, solutions = _restored(f"{stem}.xlsx", notes)
        if "seep" in solutions:
            fails.append("a seepage solution saved against a different mesh was "
                         "restored anyway")
        if not any(f"{os.path.basename(stem)}_seep.csv" in n
                   and "does not match this mesh" in n for n in notes):
            fails.append(f"a stale seepage solution was passed over without "
                         f"saying why: {notes}")

    # A saved solution with no mesh beside it has nothing to be placed on, and
    # that is a different reason, said differently.
    notes = []
    _sd, solutions = _restored(NOMESH_SEEP_XLSX, notes)
    if solutions:
        fails.append(f"a model with no mesh restored {sorted(solutions)}")
    if not any("no mesh" in n for n in notes):
        fails.append(f"a saved solution with no mesh to place it on was passed "
                     f"over without saying why: {notes}")

    # A run saved under the older key still names itself. The finite element
    # companion Studio writes says "analysis"; the ones the benchmark figures
    # were built with say "analysis_type", and a strength reduction run read
    # back as a loaded one is one whose factor of safety the report will not
    # state.
    with tempfile.TemporaryDirectory() as tmp:
        stem = _sidecar_copy(
            os.path.splitext(FEM_XLSX)[0], tmp,
            lambda meta: {("analysis_type" if k == "analysis" else k): v
                          for k, v in meta.items()})
        _sd, solutions = _restored(f"{stem}.xlsx")
        got = (solutions.get("fem") or {}).get("analysis")
        if got != "ssrm":
            fails.append(f"a strength reduction run saved under the older key "
                         f"reads back as {got!r}")
    return fails


def _rewrite_member_csv(path, edit):
    """``path`` with ``edit`` applied to its data rows — the comment header and the
    column line are kept, since that is what a real member sidecar looks like and
    the reader skips the one and needs the other."""
    with open(path) as f:
        lines = f.readlines()
    head = [ln for ln in lines if ln.startswith("#")]
    body = [ln for ln in lines if not ln.startswith("#")]
    with open(path, "w") as f:
        f.writelines(head + [body[0]] + edit(body[1:]))


def test_member_companions_that_are_not_this_models_are_refused():
    """Member forces are restored from a file only when the file is this model's.

    ``{stem}_fem_reinf.csv`` and ``{stem}_fem_piles.csv`` are found by name beside
    the field, and a name carries no model identity: a file holding another
    model's members was grafted on, and rows addressing elements this model does
    not have were dropped one at a time, so a partial set of forces arrived
    looking like a solved result. Each file is measured against the model's own
    element count, and one that disagrees is left out whole and named — the same
    treatment, through the same notes, a field saved against a rebuilt mesh gets.

    The measure is the exact SET of element ids the model's own export would
    write for that kind, not a count and a range. A model carrying both
    reinforcement and piles splits one 1D element list between the two, so a
    reinforcement file whose rows land on the pile slots has the right number of
    rows and every id inside the list — and would have had every force grafted
    onto the wrong element.
    """
    fails = []

    # The model's own files: the members come back, and nothing is said.
    with tempfile.TemporaryDirectory() as tmp:
        stem = _sidecar_copy(os.path.splitext(FEM_PILES_XLSX)[0], tmp)
        notes = []
        _sd, solutions = _restored(f"{stem}.xlsx", notes)
        sol = (solutions.get("fem") or {}).get("solution") or {}
        if "forces_pile_lateral" not in sol:
            fails.append("a model read beside its own pile companion got no "
                         "pile forces")
        if notes:
            fails.append(f"a model read beside its own companions was faulted "
                         f"for it: {notes}")

    # A file holding more members than the model has. Count alone settles it:
    # every row could address an element that exists and the file would still be
    # a different model's.
    with tempfile.TemporaryDirectory() as tmp:
        stem = _sidecar_copy(os.path.splitext(FEM_PILES_XLSX)[0], tmp)
        _rewrite_member_csv(f"{stem}_fem_piles.csv", lambda rows: rows[:-1])
        notes = []
        _sd, solutions = _restored(f"{stem}.xlsx", notes)
        sol = (solutions.get("fem") or {}).get("solution") or {}
        if "forces_pile_lateral" in sol:
            fails.append("a pile companion with the wrong number of members was "
                         "grafted onto the model anyway")
        if not any("_fem_piles.csv" in n and "not this model's" in n
                   for n in notes):
            fails.append(f"a pile companion that was refused is not in the "
                         f"notes: {notes}")

    # A file whose rows address elements the model does not carry. The row count
    # is right; the ids are not, and every row that missed used to vanish in
    # silence.
    with tempfile.TemporaryDirectory() as tmp:
        stem = _sidecar_copy(os.path.splitext(FEM_REINF_XLSX)[0], tmp)

        def _shift(rows):
            out = []
            for ln in rows:
                cells = ln.split(",")
                cells[0] = str(int(cells[0]) + 100000)
                out.append(",".join(cells))
            return out

        _rewrite_member_csv(f"{stem}_fem_reinf.csv", _shift)
        notes = []
        _sd, solutions = _restored(f"{stem}.xlsx", notes)
        sol = (solutions.get("fem") or {}).get("solution") or {}
        if "forces_1d" in sol:
            fails.append("a reinforcement companion addressing elements the "
                         "model does not have was grafted onto it anyway")
        if not any("_fem_reinf.csv" in n and "not this model's" in n
                   for n in notes):
            fails.append(f"a reinforcement companion that was refused is not in "
                         f"the notes: {notes}")

    # A file with the right count whose ids are all IN RANGE and are not the
    # model's set. Range and count both pass; the file addresses one element
    # sixty times and the other fifty-nine not at all.
    with tempfile.TemporaryDirectory() as tmp:
        stem = _sidecar_copy(os.path.splitext(FEM_REINF_XLSX)[0], tmp)

        def _collapse(rows):
            out = []
            for ln in rows:
                cells = ln.split(",")
                cells[0] = "0"
                out.append(",".join(cells))
            return out

        _rewrite_member_csv(f"{stem}_fem_reinf.csv", _collapse)
        notes = []
        _sd, solutions = _restored(f"{stem}.xlsx", notes)
        sol = (solutions.get("fem") or {}).get("solution") or {}
        if "forces_1d" in sol:
            fails.append("a reinforcement companion whose ids are in range but "
                         "are not the model's set was grafted onto it anyway")
        if not any("_fem_reinf.csv" in n and "not this model's" in n
                   for n in notes):
            fails.append(f"a reinforcement companion with the wrong id set was "
                         f"not refused: {notes}")

    # And the case only set membership can see: a model carrying BOTH kinds
    # splits one 1D element list between them, so a reinforcement file whose
    # rows land on the PILE slots has the right count and every id inside the
    # list. Built by marking half of a reinforced model's 1D elements as piles,
    # which is what such a model's fem_data looks like, and asked of the guard
    # directly — the two kinds' files are told apart before either is grafted.
    import numpy as np
    import pandas as pd
    from xslope.fem import _1d_sidecar_mismatch

    _sd, bundle = _fem_1d_bundle(FEM_REINF_XLSX)
    n_1d = len(bundle["fem_data"]["elements_1d"])
    if n_1d < 4:
        fails.append("the reinforced fixture carries too few 1D elements to "
                     "split between two kinds")
    else:
        mask = np.zeros(n_1d, dtype=bool)
        mask[n_1d // 2:] = True           # the back half is piles
        reinf_ids = np.flatnonzero(~mask)
        pile_ids = np.flatnonzero(mask)
        n_reinf = len(reinf_ids)

        def refusal(ids):
            return _1d_sidecar_mismatch(
                "x_fem_reinf.csv", pd.DataFrame({"element_id": ids}),
                "element_id", n_reinf, n_1d, "reinforcement",
                expected=reinf_ids)

        if refusal(reinf_ids) is not None:
            fails.append(f"the model's own reinforcement ids were refused: "
                         f"{refusal(reinf_ids)!r}")
        wrong = refusal(pile_ids[:n_reinf])
        if wrong is None:
            fails.append("a reinforcement file whose rows address the PILE "
                         "slots was accepted: same count, every id in range")
        elif "not this model's" not in wrong:
            fails.append(f"the refusal does not say the file is not this "
                         f"model's: {wrong!r}")

        # The pile file is identified by element_id, not by pile_index: the
        # latter is 0..n-1 on every model and identifies nothing.
        def pile_refusal(element_ids):
            return _1d_sidecar_mismatch(
                "x_fem_piles.csv",
                pd.DataFrame({"pile_index": np.arange(len(element_ids)),
                              "element_id": element_ids}),
                "pile_index", len(pile_ids), len(pile_ids), "pile",
                expected=np.arange(len(pile_ids)),
                also=(("element_id", pile_ids),))

        if pile_refusal(pile_ids) is not None:
            fails.append(f"the model's own pile ids were refused: "
                         f"{pile_refusal(pile_ids)!r}")
        if pile_refusal(reinf_ids[:len(pile_ids)]) is None:
            fails.append("a pile file whose element_id column addresses the "
                         "reinforcement slots was accepted")

        # An older sidecar carries no element_id at all, and is still accepted
        # on what it does carry: the guard hardens, it does not break reloads.
        older = _1d_sidecar_mismatch(
            "old_fem_piles.csv",
            pd.DataFrame({"pile_index": np.arange(len(pile_ids))}),
            "pile_index", len(pile_ids), len(pile_ids), "pile",
            expected=np.arange(len(pile_ids)),
            also=(("element_id", pile_ids),))
        if older is not None:
            fails.append(f"a sidecar written before element_id existed was "
                         f"refused: {older!r}")

        # And the READER hands the guard that set, rather than the whole 1D
        # element list. Driven through the reloader on the same split model, so
        # what is measured is the wiring and not the predicate alone.
        from pathlib import Path

        from xslope.fem import _import_1d_result_sidecars
        mixed = dict(bundle["fem_data"])
        mixed["pile_elem_mask"] = mask
        mixed["n_pile_elements"] = len(pile_ids)
        mixed["pile_elem_indices"] = pile_ids
        columns = ["element_id", "line_id", "x_start", "y_start", "x_end",
                   "y_end", "axial_force", "t_allow", "t_cap", "t_res",
                   "mobilization", "failed", "softened"]

        def written(tmp, ids):
            frame = pd.DataFrame({c: np.zeros(len(ids)) for c in columns})
            frame["element_id"] = ids
            frame["failed"] = False
            frame["softened"] = False
            frame.to_csv(os.path.join(tmp, "m_fem_reinf.csv"), index=False)
            solution, notes = {}, []
            _import_1d_result_sidecars(mixed, solution,
                                       Path(os.path.join(tmp, "m")), "fem")
            return solution

        with tempfile.TemporaryDirectory() as tmp:
            got = written(tmp, pile_ids[:n_reinf])
            if not (got.get("sidecar_notes") or []):
                fails.append("the reloader accepted a reinforcement file "
                             "addressing the pile slots of a model carrying "
                             "both kinds")
            if "forces_1d" in got:
                fails.append("forces from the wrong kind's slots were grafted "
                             "onto the solution")
        with tempfile.TemporaryDirectory() as tmp:
            got = written(tmp, reinf_ids)
            if got.get("sidecar_notes"):
                fails.append(f"the reloader refused the model's own "
                             f"reinforcement file: {got['sidecar_notes']}")
    return fails


# --------------------------------------------------------------------------
# H. nothing is said that is not so
# --------------------------------------------------------------------------

def test_water_prose_is_conditional():
    """A dry model gets one plain statement about water; a model with a pool gets
    the rows that describe it.

    The sample defines no piezometric line, no specified-head boundary and no
    material pore pressure, and the report still explained how its water loads
    were derived. Prose that describes a feature the model does not have is not a
    harmless extra sentence — it is a statement about the analysis that is false.
    """
    fails = []
    from xslope.fileio import load_slope_data
    from xslope.report import build_report, water_features

    slope_data, solutions = _solved()
    feats = water_features(slope_data)
    if feats["any"]:
        return ["the sample model now carries water; the dry-model check has "
                "nothing to see"]

    report = _build()
    # Water is a stability input, so it is stated where the stability analysis
    # is documented, beside the materials whose pore pressures it sets.
    lem = next((s for s in report.sections
                if s.title == "Limit Equilibrium Analysis"), None)
    water = next((s for _l, s in (lem.walk() if lem else [])
                  if s.title == "Materials"), None)
    if water is None:
        return ["the limit equilibrium section has no Materials section"]
    dry = [b.text for b in water.blocks
           if b.kind == "prose" and "dry" in b.text]
    if not dry:
        fails.append(f"a dry model gets no statement that it is analyzed dry: "
                     f"{[b.text for b in water.blocks if b.kind == 'prose']}")
    else:
        text = dry[0]
        for want in ("no groundwater", "dry"):
            if want not in text:
                fails.append(f"the dry-model statement does not say {want!r}: "
                             f"{text!r}")
        if "zero" not in text:
            fails.append("the dry-model statement does not say pore pressures are "
                         "zero")
        # One sentence carries it (Norm: the long inventory of absent features
        # was itself unnecessary).
        if text.count(".") > 1:
            fails.append(f"the dry-model statement is more than one sentence: "
                         f"{text!r}")
    if any(b.kind == "keyvalues" for b in water.blocks):
        fails.append("a dry model was given water rows")

    # No sentence anywhere in a dry report describes water the model lacks.
    banned = ("water surface", "water loads", "piezometric line", "seepage",
              "ponded", "standing water", "water level")
    for block in report.blocks("prose"):
        low = block.text.lower()
        for phrase in banned:
            if phrase in low and "no " + phrase not in low and \
                    "no piezometric lines" not in low and \
                    "no seepage head boundaries" not in low:
                fails.append(f"a dry report says {phrase!r}: {block.text!r}")
    for block in report.blocks("keyvalues"):
        for label, _value in block.items:
            if "water" in label.lower():
                fails.append(f"a dry report carries the key-value row {label!r}")

    # The positive path: a model that really has a pool states it.
    dam = load_slope_data(DAM_XLSX)
    wet = water_features(dam)
    if not wet["any"]:
        fails.append("the dam sample no longer carries water; the positive path "
                     "is untested")
    else:
        with tempfile.TemporaryDirectory() as tmp:
            dam_report = build_report(dam, solutions,
                                      {"pd_figure": False,
                                       "lem_search_figure": False,
                                       "lem_solution_figure": False}, tmp)
        sub = next((s for _l, s in _sections(dam_report)
                    if s.title == "Materials"), None)
        if sub is None:
            fails.append("the dam report has no Materials section")
        else:
            kv = [b for b in sub.blocks if b.kind == "keyvalues"]
            if not kv:
                fails.append("a model with a reservoir got the dry statement")
            else:
                labels = [l for l, _v in kv[0].items]
                if not any("Water surface" in l for l in labels):
                    fails.append(f"the pool is not stated: {labels}")
                if not any(l == "Water loads" for l in labels):
                    fails.append(f"how the water loads are applied is not stated: "
                                 f"{labels}")
    fails += _water_stage_language()
    fails += _water_surface_cites_the_seepage_section()
    return fails


def _water_stage_language():
    """Stage numbering appears only where the analysis has stages.

    "Water surface (stage 1)" on a model solved at one water state is rapid
    drawdown vocabulary applied to something that is not staged: it tells the
    reader there is a stage 2 to compare it against, and there is not.
    """
    fails = []
    from xslope.fileio import load_slope_data
    from xslope.report import water_features, _water_items

    # A single-state model: one pool, held. Its one water surface is not
    # numbered, and no row it carries says "stage" at all.
    for name, xlsx in (("the zoned dam", DAM_XLSX), ("the dam", SEEP_XLSX)):
        sd = load_slope_data(xlsx)
        feats = water_features(sd)
        if len(feats["surfaces"]) != 1:
            fails.append(f"{name} no longer carries exactly one water surface, "
                         f"so the unstaged wording is untested")
            continue
        for seepage_section in (False, True):
            rows = _water_items(sd, feats, seepage_section=seepage_section)
            staged = [(l, v) for l, v in rows if "stage" in f"{l} {v}".lower()]
            if staged:
                fails.append(f"{name} is solved at one water state and its rows "
                             f"still use stage language: {staged}")
            if not any(l == "Water surface" for l, _v in rows):
                fails.append(f"{name} states no plain 'Water surface' row: "
                             f"{[l for l, _v in rows]}")

    # And the staged analyses keep their numbering: the two states of a rapid
    # drawdown are told apart by it, and a report that dropped it would be
    # naming two different water surfaces the same thing.
    for name, xlsx in (("the earth dam rapid drawdown", RAPID_SEEP_XLSX),
                       ("the transient reservoir drawdown",
                        os.path.join(_REPO, "docs", "lem", "files",
                                     "xslope_johnson_res_rapid.xlsx"))):
        sd = load_slope_data(xlsx)
        feats = water_features(sd)
        rows = _water_items(sd, feats, seepage_section=True)
        surfaces = [l for l, _v in rows if l.startswith("Water surface")]
        if not surfaces:
            fails.append(f"{name} states no water surface at all: "
                         f"{[l for l, _v in rows]}")
        for label in surfaces:
            if "stage" not in label.lower():
                fails.append(f"{name} is a staged analysis and its water "
                             f"surface row is unnumbered: {label!r}")
        if len(set(surfaces)) != len(surfaces):
            fails.append(f"{name} labels two water surfaces the same: {surfaces}")
    return fails


def _water_surface_cites_the_seepage_section():
    """Where the report solves the flow, the stability section cites it for the
    water surface instead of restating the head boundaries.

    The elevations of the head boundaries are stated once, in the section that
    imposes them. A second copy under the stability analysis is a number that can
    only ever come to disagree with the first.
    """
    fails = []
    from xslope.report import water_features, SEEPAGE_ANCHOR, section_anchor

    slope_data, _bundle = _seep_bundle()
    feats = water_features(slope_data)
    if not (feats["heads"] and feats["surfaces"]):
        return ["the seepage sample's head boundaries no longer put water above "
                "the ground surface; the citation path is untested"]

    report = _cite_report(SEEP_XLSX, ("spencer",), engines=("seep",))
    titles = [t for _l, t in report.section_titles()]
    if "Seepage Analysis" not in titles or \
            "Limit Equilibrium Analysis" not in titles:
        return [f"the seepage sample no longer reports both engines: {titles}"]

    seep = next(s for s in report.sections if s.title == "Seepage Analysis")
    if seep.anchor != section_anchor(SEEPAGE_ANCHOR):
        fails.append(f"the Seepage Analysis section carries the anchor "
                     f"{seep.anchor!r}, which nothing can cite it by")
    number = next(n for n, _l, s in report.section_numbers() if s is seep)

    lem = next(s for s in report.sections
               if s.title == "Limit Equilibrium Analysis")
    mats = next((s for _l, s in lem.walk() if s.title == "Materials"), None)
    if mats is None:
        return fails + ["the limit equilibrium section has no Materials section"]

    cites = [b for b in mats.blocks
             if b.kind == "prose" and f"Section {number}" in b.text]
    if not cites:
        fails.append(f"the stability analysis does not send the reader to "
                     f"Section {number} for its water surface: "
                     f"{[b.text for b in mats.blocks if b.kind == 'prose']}")
    else:
        targets = [t for _text, t in cites[0].links]
        if f"#{section_anchor(SEEPAGE_ANCHOR)}" not in targets:
            fails.append(f"the reference to the seepage section is not a live "
                         f"cross-reference: {targets}")

    # And the boundary elevations are set down once. The row names the analysis
    # the surface comes from; the numbers stay in the section that imposes them.
    from xslope.water import water_line_for_stage
    stated = water_line_for_stage(slope_data, stage=1)["source"]
    detail = stated[stated.find("(") + 1:stated.rfind(")")] if "(" in stated else ""
    if not detail:
        fails.append(f"the seepage sample's water source names no boundary "
                     f"detail ({stated!r}), so the duplication is untested")
    for block in mats.blocks:
        if block.kind != "keyvalues":
            continue
        for label, value in block.items:
            if detail and detail in value:
                fails.append(f"the stability analysis restates the seepage "
                             f"boundaries the seepage section already gives: "
                             f"{label!r} = {value!r}")
    return fails


def _sections(report):
    """Every section in the report, as ``(level, section)``."""
    out = []
    for s in report.sections:
        out.extend(s.walk())
    return out


def _under(report, heading):
    """The titles of every section under one top-level heading."""
    sec = next((s for s in report.sections if s.title == heading), None)
    return [] if sec is None else [s.title for _l, s in sec.walk()]


def _table_in(report, heading, caption):
    """The table with ``caption`` printed under ``heading``, or None."""
    sec = next((s for s in report.sections if s.title == heading), None)
    if sec is None:
        return None
    for _l, sub in sec.walk():
        for block in sub.blocks:
            if block.kind == "table" and block.caption == caption:
                return block
    return None


#: What each engine's member tables must and must not carry, in the words their
#: headers are built with. Written out here rather than read off the registry the
#: tables are built from: a registry that starts handing the finite element
#: analysis a direction, or the method of slices a Young's modulus, has to fail
#: rather than pass on its own new claim.
#:
#: Each entry is the properties that engine reads and the ones the OTHER engine
#: reads and it does not. The evidence is the code: the limit equilibrium force
#: on a slice is built in ``slice.py`` from the capacity, the pullout envelope,
#: the direction and the application; the finite element element is assembled in
#: ``fem.py`` from the modulus and area, softening to the residual capacity, and
#: reads neither the direction nor whether the force is applied or mobilized.
_MEMBER_COLUMNS = {
    ("reinforcement", "lem"): (("T_max", "L_p1", "L_p2", "Direction",
                                "Applied"), ("T_res", "E", "Area")),
    ("reinforcement", "fem"): (("T_max", "T_res", "L_p1", "L_p2",
                                "E", "Area"), ("Direction", "Applied")),
    ("piles", "lem"): (("H", "θ (deg)", "V_cap", "M_cap", "Applied"),
                       ("E", "I", "Head fixity")),
    ("piles", "fem"): (("V_cap", "M_cap", "E", "Head fixity"),
                       ("H", "θ (deg)", "Applied")),
}


def _member_columns_are_the_engines(table, member, engine, where):
    """One member table carries the properties its own engine reads, and none of
    the ones only the other engine reads."""
    fails = []
    wanted, refused = _MEMBER_COLUMNS[(member, engine)]
    heads = [h.split(" (")[0].strip() for h in table.headers]
    for name in wanted:
        if name.split(" (")[0] not in heads:
            fails.append(f"{where}: the {engine} {member} table does not give "
                         f"{name}, which that engine reads: {table.headers}")
    for name in refused:
        if name.split(" (")[0] in heads:
            fails.append(f"{where}: the {engine} {member} table gives {name}, "
                         f"which only the other engine reads: {table.headers}")
    return fails


def test_members_stand_with_the_engine_that_reads_them():
    """Reinforcement and piles are stated where the analysis that reads them is
    documented, with the properties THAT analysis reads.

    They were in the general description of the model, which said they were a
    property of the section. They are not: they are structure an analysis acts
    on, and the two engines read different things off the same line. The method
    of slices resolves a capacity onto the slice base and needs the pullout
    envelope, the direction and whether the force is applied or mobilized; the
    finite element analysis carries the line as an element and needs the modulus
    and area it is assembled from and the residual it softens to. So each engine
    states its own, and a report of both prints two tables that share a geometry
    and nothing else.

    Each is its own section, present only where the model carries that member:
    one heading for the reinforcement, one for the piles, and neither hiding
    under the other.
    """
    fails = []
    from xslope.report import MEMBER_CAPTIONS, build_report

    caption = MEMBER_CAPTIONS
    slope_data, solutions = _solved()
    report = _build()
    if not slope_data.get("reinforcement_lines"):
        return ["the sample carries no reinforcement; the split is untested"]
    if slope_data.get("pile_lines"):
        return ["the sample now carries piles; the 'absent' half is untested"]
    # Two tables off the same lines need two names in the list of tables.
    if len(set(caption.values())) != len(caption):
        fails.append(f"two member tables share a caption: {caption}")

    # --- not in the general description of the model ---
    for gone in ("Reinforcement", "Piles", "Reinforcement and Piles"):
        if gone in _under(report, "Project Definition"):
            fails.append(f"Project Definition still carries {gone!r}")
    lem_titles = _under(report, "Limit Equilibrium Analysis")
    if "Reinforcement" not in lem_titles:
        fails.append(f"the limit equilibrium section does not state the "
                     f"reinforcement it was run on: {lem_titles}")
    if "Piles" in lem_titles:
        fails.append("a model with no piles was given a Piles section")

    # --- the columns are the engine's own ---
    table = _table_in(report, "Limit Equilibrium Analysis",
                      caption[("reinforcement", "lem")])
    if table is None:
        fails.append("the limit equilibrium section prints no reinforcement table")
    else:
        fails += _member_columns_are_the_engines(
            table, "reinforcement", "lem", "the sample")

    # --- a model with piles and no reinforcement gets the other half ---
    piled = dict(slope_data)
    piled["reinforcement_lines"] = []
    piled["pile_lines"] = load_slope_data_cached(PILES_XLSX)["pile_lines"]
    with tempfile.TemporaryDirectory() as tmp:
        report = build_report(piled, solutions, {"pd_figure": False}, tmp)
    titles = _under(report, "Limit Equilibrium Analysis")
    if "Piles" not in titles:
        fails.append(f"a model with piles got no Piles section: {titles}")
    if "Reinforcement" in titles:
        fails.append("a model with no reinforcement got a Reinforcement section")
    table = _table_in(report, "Limit Equilibrium Analysis",
                      caption[("piles", "lem")])
    if table is None:
        fails.append("the limit equilibrium section prints no piles table")
    else:
        fails += _member_columns_are_the_engines(table, "piles", "lem",
                                                 "a piled model")

    # --- the finite element analysis states its own, off the same lines ---
    for xlsx, member in ((FEM_REINF_XLSX, "reinforcement"),
                         (FEM_PILES_XLSX, "piles")):
        fem_data, bundle = _fem_1d_bundle(xlsx)
        with tempfile.TemporaryDirectory() as tmp:
            report = build_report(fem_data, {"fem": [bundle]},
                                  dict(FAST_FIGURES, pd_figure=False), tmp)
        # One gravity trial is a deformation analysis, not a strength reduction,
        # and the section is headed for what it is.
        head = next((s.title for s in report.sections
                     if s.title.startswith("Deformation")), "")
        titles = _under(report, head)
        wanted = "Reinforcement" if member == "reinforcement" else "Piles"
        if wanted not in titles:
            fails.append(f"{os.path.basename(xlsx)}: the finite element section "
                         f"does not state the {member} it carries: {titles}")
        named = caption[(member, "fem")]
        table = _table_in(report, head, named)
        if table is None:
            fails.append(f"{os.path.basename(xlsx)}: the finite element section "
                         f"prints no {named!r} table")
            continue
        fails += _member_columns_are_the_engines(
            table, member, "fem", os.path.basename(xlsx))
        # A finite-element-only report puts them in the only engine section
        # there is, and nowhere else.
        if _table_in(report, "Project Definition", named) is not None:
            fails.append(f"{os.path.basename(xlsx)}: the {named!r} table is in "
                         f"the Project Definition of a report that has an "
                         f"engine to state it under")

    # --- the two engines' tables really are different ---
    #
    # Both sides of the split assembled from one registry can agree by accident;
    # what makes the split real is that the same line prints different columns
    # under each engine.
    from xslope.report import _reinforcement_table, _Counter
    lem_t = _reinforcement_table(slope_data, _Counter(), "lem")
    fem_t = _reinforcement_table(slope_data, _Counter(), "fem")
    if lem_t is not None and fem_t is not None and lem_t.headers == fem_t.headers:
        fails.append(f"the two engines print the same reinforcement columns, so "
                     f"the split states nothing: {lem_t.headers}")
    return fails


def test_the_model_checks_stay_out_of_the_report():
    """A report carries no model-check findings, whatever it is handed.

    The checks tell whoever is BUILDING a model what is wrong with it, and that
    is a conversation with the interface: a submittal that prints its own
    warnings is a submittal arguing with itself (the owner's ruling,
    fem_reinforce review). The machinery stays — Studio runs it, and it still
    gates a run — and the report simply has no such section: no findings table,
    no severity words, and no sentence saying the checks found nothing either.

    Handed a preflight report and asked for the retired option by name, the
    builder must still print none of it.
    """
    fails = []
    from xslope import report as report_mod
    from xslope.preflight import Finding, PreflightReport, preflight, rules

    if "model_checks" in report_mod.DEFAULT_OPTIONS:
        fails.append("the report still declares a model_checks option")
    if "preflight" in report_mod.DEFAULT_OPTIONS:
        fails.append("the report still takes a preflight report as an option")
    if hasattr(report_mod, "_model_checks_section"):
        fails.append("the model checks section builder is still in the module")

    slope_data, solutions = _solved()
    every = next((r for r in rules() if "*" in r.analyses), None)
    if every is None:
        return fails + ["the rule registry offers no every-analysis rule to "
                        "hand the builder"]
    findings = [Finding(every.id, "error", "a finding about the model itself")]
    report = _build({"model_checks": True,
                     "preflight": PreflightReport(analysis="lem",
                                                  findings=findings)})
    titles = [t for _l, t in report.section_titles()]
    if "Model Checks" in titles:
        fails.append(f"the report built a Model Checks section: {titles}")
    said = " ".join(b.text for b in report.blocks("prose"))
    for banned in ("a finding about the model itself",
                   "raised no findings", "xslope checks a model"):
        if banned in said:
            fails.append(f"the report still says {banned!r}")
    for table in report.tables():
        if table.caption == "Model check findings":
            fails.append("the findings table is still printed")

    # The machinery the interface uses is untouched: the same rules still run
    # against the same model and still return findings.
    live = preflight(slope_data, "lem")
    if not hasattr(live, "findings"):
        fails.append("the preflight machinery no longer returns findings")
    if report_mod.report_analyses(solutions, {"lem": True}) != ["lem"]:
        fails.append(f"a LEM report says it documents "
                     f"{report_mod.report_analyses(solutions, {'lem': True})}")
    return fails


def _fem_check_rows(slope_data, bundle, prefix):
    """The model-check findings a strength reduction run of this model raises
    whose rule id starts with ``prefix``, in the shape the section that used to
    print them used: ``[severity word, message, rule id]``.

    Read from the checker itself. The report no longer carries the findings —
    they belong to the interface (:func:`test_the_model_checks_stay_out_of_the_report`)
    — and what these checks are about is the RULES: a modulus a thousandth of
    what its material's own strength implies leaves the factor of safety intact
    and corrupts every displacement beside it, and the checker has to say so
    whoever is listening.
    """
    from xslope.preflight import preflight

    words = {"error": "Error", "warning": "Warning", "info": "Note"}
    with contextlib.redirect_stdout(io.StringIO()):
        found = preflight(slope_data, "ssrm").findings
    rows = [[words.get(f.severity, f.severity.title()), f.message, f.rule_id]
            for f in found if str(f.rule_id).startswith(prefix)]
    return rows, ""


def test_implausible_elastic_properties_are_flagged():
    """The checker says when a material's elastic constants are not a soil's, and
    never fills one in.

    The factor of safety a strength reduction reaches does not depend on the
    elastic constants — a perfectly plastic collapse load is independent of them —
    so a modulus in the wrong stress unit leaves the answer intact and corrupts
    every displacement reported beside it. The checker measures each modulus
    against its own material's soil type in the declared unit system, and each
    Poisson's ratio against the range a geomaterial has, and it does so under the
    strength reduction rules — the rules a limit equilibrium check never
    evaluates, which is how a model whose modulus was a thousandth of its soil
    type once passed in silence.
    """
    import copy

    fails = []
    slope_data, bundle = _fem_bundle()

    # As shipped, the model says nothing about its elastic constants.
    rows, prose = _fem_check_rows(slope_data, bundle, "mat.")
    if rows is None:
        return ["the strength reduction rules produced no findings at all"]
    for row in rows:
        if row[2] in ("mat.E_off_soil_type_band", "mat.nu_implausible",
                      "mat.nu_unusable", "mat.E_unusable"):
            fails.append(f"a model with ordinary elastic constants was faulted "
                         f"for them: {row[1]!r}")

    # A modulus a thousandth of what the material's own strength implies.
    soft = copy.deepcopy(slope_data)
    soft["materials"][0]["E"] = float(soft["materials"][0]["E"]) / 1000.0
    rows, _prose = _fem_check_rows(soft, bundle, "mat.E")
    if not any(r[2] == "mat.E_off_soil_type_band" for r in rows):
        fails.append(f"a modulus a thousandth of its soil type's is not reported: "
                     f"{rows}")
    else:
        said = next(r[1] for r in rows if r[2] == "mat.E_off_soil_type_band")
        # The finding names the material, its value and the declared system, so
        # the reader can act on it without opening the checker.
        for want in ("Material 1", "700", "imperial"):
            if want not in said:
                fails.append(f"the modulus finding does not name {want!r}: "
                             f"{said!r}")
    # Flagged, never filled: the model still holds what was entered.
    if float(soft["materials"][0]["E"]) != float(slope_data["materials"][0]["E"]) / 1000.0:
        fails.append("the checker rewrote the modulus it was asked to report on")

    # A Poisson's ratio no geomaterial has.
    flat = copy.deepcopy(slope_data)
    flat["materials"][0]["nu"] = 0.02
    rows, _prose = _fem_check_rows(flat, bundle, "mat.nu")
    if not any(r[2] == "mat.nu_implausible" for r in rows):
        fails.append(f"a Poisson's ratio of 0.02 is not reported: {rows}")
    if float(flat["materials"][0]["nu"]) != 0.02:
        fails.append("the checker rewrote the Poisson's ratio it was reporting on")

    # And one outside the admissible range is the harder finding.
    bad = copy.deepcopy(slope_data)
    bad["materials"][0]["nu"] = 0.7
    rows, _prose = _fem_check_rows(bad, bundle, "mat.nu")
    hit = [r for r in rows if r[2] == "mat.nu_unusable"]
    if not hit:
        fails.append(f"a Poisson's ratio of 0.7 is not reported: {rows}")
    elif hit[0][0] != "Error":
        fails.append(f"a Poisson's ratio the solver refuses is reported as "
                     f"{hit[0][0]!r}")
    return fails


def test_title_page_omits_empty_rows():
    """A title-page row with nothing in it is not printed.

    Not every project has a number, an organization or a named author, and a label
    with a blank beside it is a question the title page asks and does not answer.
    """
    fails = []
    from xslope.report import generate_report

    slope_data, solutions = _solved()
    opts = {"input_path": REINF_XLSX, "title": "Bare Title Page",
            "pd_figure": False, "lem_search_figure": False,
            "lem_solution_figure": False, "lem_slice_table": False}
    with tempfile.TemporaryDirectory() as tmp:
        out_path = os.path.join(tmp, "bare.docx")
        ok, out = generate_report(slope_data, solutions, opts, out_path)
        if not ok:
            return [f"generate_report failed: {out}"]
        _names, xml = _docx_parts(out_path)
        doc = xml.get("word/document.xml", "")
        for absent in ("Project:", "Author:"):
            if absent in doc:
                fails.append(f"a report with no project number and no author "
                             f"still prints {absent!r}")
        if "Date:" not in doc:
            fails.append("the date row is missing; it always has a value")
        if "Bare Title Page" not in doc:
            fails.append("the title is not on the title page")

    # With the fields filled, the rows are there and read as asked.
    opts = dict(opts, project_number="26-001", author="A. Engineer",
                organization="Example Engineering")
    with tempfile.TemporaryDirectory() as tmp:
        out_path = os.path.join(tmp, "full.docx")
        ok, out = generate_report(slope_data, solutions, opts, out_path)
        if not ok:
            return fails + [f"generate_report failed: {out}"]
        _names, xml = _docx_parts(out_path)
        doc = xml.get("word/document.xml", "")
        for want in ("Project:", "Author:", "Date:", "26-001", "A. Engineer"):
            if want not in doc:
                fails.append(f"the filled title page is missing {want!r}")
        if "Project number" in doc:
            fails.append("the title page still labels the row 'Project number'")
    return fails


# --------------------------------------------------------------------------
# J. every figure and every table is cited
#
# For a technical report it is customary to cite every table and every figure
# by number, and a block nothing points at is a block the reader meets without
# being sent to it. Three things are checked, over the option combinations and
# the models that switch the numbered blocks on and off:
#
#   1. EVERY block that prints is cited. The citation has to sit in the block's
#      own section or in a section above it — Bishop's block naming "Table 5"
#      says nothing about Spencer's slice table, and a sentence a page away
#      under another method is not a citation of this one.
#   2. NOTHING is cited that does not print. An option that suppresses a table
#      has to take the sentence that names it with it, or the report sends the
#      reader to a number that is not in the document.
#   3. A citation is a live cross-reference: the phrase carries a link to a
#      bookmark, which report_docx places on the caption line of every numbered
#      block.
#
# The numbers are read off the built tree, never off a literal here: a citation
# is only right if it is the number the counter assigned.
# --------------------------------------------------------------------------

#: A citation, as the prose writes it.
CITATION = re.compile(r"\b(Figure|Table) (\d+)\b")

#: A run of consecutive numbered blocks cited as one range — "Figures 8\u201313".
#: Both ends carry a cross-reference and the numbers between them are cited by
#: inclusion, so the range counts as a citation of every block it spans.
CITATION_RANGE = re.compile(r"\b(Figure|Table)s (\d+)[\u2013-](\d+)\b")

#: What a section citation writes in place of a number the builder cannot know
#: yet. Nothing that survives the build should still carry it.
_MARK_CHAR = "\ue000"

#: Low-resolution figures: these checks read the tree a build produced, not the
#: pixels, and every combination below draws the full set.
FAST_FIGURES = {"dpi": 60, "figsize": (4.0, 2.5)}

PILES_XLSX = os.path.join(_REPO, "docs", "lem", "files", "xslope_piles.xlsx")

_CITE_REPORTS = {}


def _cite_report(xlsx, methods, options=None, engines=()):
    """A report of ``xlsx`` solved by ``methods``, with the figures drawn.

    ``engines`` names the other engines whose solutions the report documents
    alongside — ``("seep",)`` for a model with a solved flow field. Each is read
    back from the solution shipped beside the model. A case naming no method
    switches the limit equilibrium section off: a report of a seepage run or a
    strength reduction run is a report in its own right.

    Cached on its arguments: several combinations share a model, and each is a
    solve plus a set of plots.
    """
    key = (xlsx, tuple(methods), tuple(sorted((options or {}).items())),
           tuple(engines))
    if key in _CITE_REPORTS:
        return _CITE_REPORTS[key]
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.report import build_report
    from xslope.slice import generate_slices
    from xslope.solve import solve_selected

    with contextlib.redirect_stdout(io.StringIO()):
        slope_data = load_slope_data(xlsx)
    solutions = {}
    if methods:
        circles = slope_data.get("circles") or []
        surface_kw = ({"circle": circles[0]} if circles
                      else {"non_circ": slope_data.get("non_circ")})
        ok, out = generate_slices(slope_data, num_slices=15, **surface_kw)
        if not ok:
            raise RuntimeError(f"{os.path.basename(xlsx)} produced no slices: {out}")
        df, surface = out[0], out[1]
        bundles = []
        for name in methods:
            work = df.copy()
            with contextlib.redirect_stdout(io.StringIO()):
                results = solve_selected(name, work)
            if not isinstance(results, dict):
                continue
            bundles.append({"slice_df": work, "failure_surface": surface,
                            "results": results, "search": None, "method": name})
        if not bundles:
            raise RuntimeError(f"no method converged on {os.path.basename(xlsx)}")
        # A synthetic search on the first bundle, so the search figure and the
        # sentence that cites it are built by every combination that asks for them.
        bundles[0]["search"] = {
            "kind": "circular" if circles else "noncircular",
            "fs_cache": [{"Xo": 10.0, "Yo": 50.0, "Depth": 0.0,
                          "FS": b["results"]["FS"], "slices": b["slice_df"],
                          "failure_surface": surface,
                          "solver_result": b["results"]} for b in bundles],
            "search_path": [{"x": 10.0, "y": 50.0,
                             "FS": bundles[0]["results"]["FS"]}],
            "circle_cache": None,
        }
        solutions["lem"] = bundles
    for engine in engines:
        _sd, bundle = (_seep_bundle(xlsx) if engine == "seep"
                       else _fem_any_bundle(xlsx))
        solutions[engine] = bundle
    opts = {"input_path": xlsx, "method": list(methods)}
    if not methods:
        opts["lem"] = False
    opts.update(FAST_FIGURES)
    opts.update(options or {})
    tmp = tempfile.mkdtemp(prefix="xslope_cite_")
    with contextlib.redirect_stdout(io.StringIO()):
        report = build_report(slope_data, solutions, opts, tmp)
    _CITE_REPORTS[key] = report
    return report


#: The reports the citation rule is checked on. Each names an option
#: combination or a model feature that puts a numbered block into the tree, or
#: takes one out and has to take its citation with it.
CITATION_CASES = [
    ("the shipped defaults", REINF_XLSX, ("spencer", "bishop"), {}, ()),
    ("the calculations switched off", REINF_XLSX, ("spencer",),
     {"lem_calculations": False}, ()),
    ("the slice table switched off", REINF_XLSX, ("spencer",),
     {"lem_slice_table": False}, ()),
    ("the slice key switched off", REINF_XLSX, ("spencer",),
     {"lem_slice_key": False}, ()),
    ("every figure switched off", REINF_XLSX, ("spencer",),
     {"pd_figure": False, "lem_search_figure": False,
      "lem_solution_figure": False, "lem_slice_key": False}, ()),
    ("the search switched off", REINF_XLSX, ("spencer",),
     {"lem_search": False}, ()),
    ("three methods in detail", REINF_XLSX, ("spencer", "bishop", "oms"), {}, ()),
    ("every method in detail", REINF_XLSX,
     ("oms", "bishop", "janbu", "spencer", "corps", "lowe", "mprice"), {}, ()),
    ("a model carrying piles", PILES_XLSX, ("spencer",), {}, ()),
    ("a model with neither reinforcement nor loads", DAM_XLSX, ("spencer",),
     {}, ()),
    ("a model whose working is refused", PASSIVE_XLSX, ("spencer",), {}, ()),
    # The other two engines: a seepage run reported beside the stability it
    # feeds, a seepage run reported on its own, and a strength reduction run.
    ("a seepage run beside the stability analysis", SEEP_XLSX, ("spencer",),
     {}, ("seep",)),
    ("a seepage run on its own", SEEP_XLSX, (), {}, ("seep",)),
    ("a seepage run with no flow net", SEEP_XLSX, (),
     {"seep_flownet": False}, ("seep",)),
    ("a seepage run with no model figure", SEEP_XLSX, (),
     {"seep_inputs_figure": False}, ("seep",)),
    ("a seepage run with no mesh figure", SEEP_XLSX, (),
     {"seep_mesh_figure": False}, ("seep",)),
    ("a seepage run with no conductivity curves", SEEP_XLSX, (),
     {"seep_kr_figure": False}, ("seep",)),
    ("a strength reduction run", FEM_XLSX, (), {}, ("fem",)),
    ("a strength reduction run with no figures", FEM_XLSX, (),
     {"fem_figure": False}, ("fem",)),
    ("a strength reduction run with no model figure", FEM_XLSX, (),
     {"fem_inputs_figure": False}, ("fem",)),
    ("a strength reduction run with no mesh figure", FEM_XLSX, (),
     {"fem_mesh_figure": False}, ("fem",)),
    # The members a run can carry: each puts a table and its detail figures into
    # the tree, and the reinforcement model has more lines than the figure
    # budget, so its table is cited a second time by the sentence that says so.
    ("a run carrying reinforcement", FEM_REINF_XLSX, (), {}, ("fem",)),
    ("a run carrying piles", FEM_PILES_XLSX, (), {}, ("fem",)),
    ("a run whose member figures are switched off", FEM_REINF_XLSX, (),
     {"fem_reinforcement_figure": False}, ("fem",)),
    ("a strength reduction run with no properties table", FEM_XLSX, (),
     {"fem_materials": False}, ("fem",)),
    # Both engines on one model: the finite element section carries the loads the
    # stability section already presented, and cites that section for them.
    ("both engines on one model", FEM_REINF_XLSX, ("spencer",), {}, ("fem",)),
]


def _numbered_blocks(report):
    """Every figure and table that prints, as ``(kind, number, path, caption)``
    where ``path`` is the tuple of section titles it sits under."""
    out = []

    def walk(node, path):
        here = path + (node.title,)
        for block in node.blocks:
            if block.kind in ("figure", "table"):
                out.append(("Figure" if block.kind == "figure" else "Table",
                            block.number, here, block.caption))
        for child in node.children:
            walk(child, here)

    for node in report.sections:
        walk(node, ())
    return out


#: A citation of a section, by the number Word prints in front of its heading.
SECTION_CITATION = re.compile(r"\bSection (\d+(?:\.\d+)*)\b")


def _section_citations(report):
    """Every citation of a section the prose makes, as ``(number, path, block)``."""
    out = []

    def walk(node, path):
        here = path + (node.title,)
        for block in node.blocks:
            if block.kind == "prose":
                for number in SECTION_CITATION.findall(block.text):
                    out.append((number, here, block))
        for child in node.children:
            walk(child, here)

    for node in report.sections:
        walk(node, ())
    return out


def _citations(report):
    """Every citation the prose makes, as ``(kind, number, path, block)``.

    A range cites every block it spans: "Figures 8\u201313" is a citation of
    Figure 8, of Figure 13, and of the four between them. Written out that way it
    is checked exactly as an enumeration is — a range over a number the report
    does not print still fails rule 2, and a figure inside no range and no
    enumeration still fails rule 1.
    """
    out = []

    def walk(node, path):
        here = path + (node.title,)
        for block in node.blocks:
            if block.kind == "prose":
                spanned = set()
                for kind, low, high in CITATION_RANGE.findall(block.text):
                    for number in range(int(low), int(high) + 1):
                        out.append((kind, number, here, block))
                        spanned.add((kind, number))
                for kind, number in CITATION.findall(block.text):
                    if (kind, int(number)) not in spanned:
                        out.append((kind, int(number), here, block))
        for child in node.children:
            walk(child, here)

    for node in report.sections:
        walk(node, ())
    return out


def _reaches(citing, block_path):
    """Does a citation made under ``citing`` reach a block under
    ``block_path``? Only from the block's own section or one above it."""
    return tuple(citing) == tuple(block_path[:len(citing)])


def test_every_block_is_cited():
    """Every numbered block is cited, from its own section or one above it, and
    nothing is cited that does not print."""
    fails = []

    # What the rule refuses, stated on the shape it is applied to: a sentence in
    # one method's block is not a citation of another method's table, and a
    # sentence in a sibling section is not a citation of what stands here.
    lem = ("Limit Equilibrium Analysis",)
    spencer_table = lem + ("Spencer's Method", "Slice Table")
    for citing, allowed in ((spencer_table, True),
                            (lem + ("Spencer's Method",), True),
                            (lem, True),
                            (lem + ("Spencer's Method", "Calculations"), False),
                            (lem + ("Bishop's Simplified Method", "Slice Table"),
                             False)):
        if _reaches(citing, spencer_table) is not allowed:
            fails.append(f"a citation under {' > '.join(citing)} "
                         f"{'does not reach' if allowed else 'reaches'} a block "
                         f"under {' > '.join(spencer_table)}")

    for label, xlsx, methods, options, engines in CITATION_CASES:
        try:
            report = _cite_report(xlsx, methods, options, engines)
        except Exception as exc:
            fails.append(f"{label}: the report could not be built: {exc!r}")
            continue
        blocks = _numbered_blocks(report)
        cites = _citations(report)
        if not blocks:
            fails.append(f"{label}: the report carries no numbered block")
            continue

        for kind, number, path, caption in blocks:
            if not number:
                fails.append(f"{label}: {kind} {caption!r} printed with no number")
                continue
            reached = [c for c in cites
                       if c[0] == kind and c[1] == number and _reaches(c[2], path)]
            if not reached:
                elsewhere = [" > ".join(c[2]) for c in cites
                             if c[0] == kind and c[1] == number]
                fails.append(
                    f"{label}: {kind} {number} ({caption!r}), under "
                    f"{' > '.join(path)}, is cited by no sentence of its own "
                    f"section or of a section above it"
                    + (f"; it is named only under {elsewhere}" if elsewhere else ""))

        printed = {(k, n) for k, n, _p, _c in blocks}
        for kind, number, path, block in cites:
            if (kind, number) not in printed:
                fails.append(
                    f"{label}: a sentence under {' > '.join(path)} cites {kind} "
                    f"{number}, which this report does not print: {block.text!r}")

        # The numbers run 1..n with no gap, in each series: a citation is only
        # useful if the reader can find the number by counting captions.
        for kind in ("Figure", "Table"):
            got = [n for k, n, _p, _c in blocks if k == kind]
            if got != list(range(1, len(got) + 1)):
                fails.append(f"{label}: the {kind.lower()}s are numbered {got}")

        # A section citation names a section this report carries, and every
        # number a citation had to wait for was filled in.
        carried = {n for n, _lvl, _s in report.section_numbers()}
        for number, path, block in _section_citations(report):
            if number not in carried:
                fails.append(
                    f"{label}: a sentence under {' > '.join(path)} cites Section "
                    f"{number}, which this report does not carry: {block.text!r}")
        for block in report.blocks("prose"):
            if _MARK_CHAR in block.text:
                fails.append(f"{label}: a section citation was never numbered: "
                             f"{block.text!r}")
    return fails


def test_consecutive_figures_are_cited_as_a_range():
    """A run of three or more consecutive figures is cited as a range, and
    anything else is read out in full.

    Six reinforcement lines printed "Figure 8, Figure 9, Figure 10, Figure 11,
    Figure 12 and Figure 13" — the word Figure six times for one fact. Both ends
    of a range carry their own cross-reference; the numbers between them are
    cited by inclusion, which is what :func:`_citations` reads.
    """
    fails = []
    from xslope.report import CITE_RANGE_DASH, cite_anchor, cite_range

    # A gap is not a range: the reader would be sent to a figure the sentence
    # does not mean.
    for numbers, wanted in (([8, 9, 10, 11, 12, 13], f"Figures 8{CITE_RANGE_DASH}13"),
                            ([8, 9, 10], f"Figures 8{CITE_RANGE_DASH}10"),
                            ([8, 9], "Figure 8 and Figure 9"),
                            ([8], "Figure 8"),
                            ([8, 9, 11], "Figure 8, Figure 9 and Figure 11"),
                            ([], "")):
        phrase, links = cite_range("Figure", numbers)
        if phrase != wanted:
            fails.append(f"{numbers} is cited as {phrase!r}, not {wanted!r}")
        # Every end the phrase names is a live cross-reference to that block.
        for number in ({numbers[0], numbers[-1]} if numbers else set()):
            target = f"#{cite_anchor('Figure', number)}"
            if not any(t == target for _d, t in links):
                fails.append(f"{numbers}: Figure {number} is named without a "
                             f"link to it: {links}")
        for display, _target in links:
            if display not in phrase:
                fails.append(f"{numbers}: the link phrase {display!r} is not in "
                             f"the sentence {phrase!r}")

    # And the report really writes one, over the six bars of the reinforced
    # model, with the interior figures counted as cited.
    report = _engine_report("fem", xlsx=FEM_REINF_XLSX)
    said = " ".join(_prose(report))
    ranges = CITATION_RANGE.findall(said)
    if not ranges:
        fails.append(f"six reinforcement figures are still enumerated one by "
                     f"one: {said!r}")
    else:
        for kind, low, high in ranges:
            spanned = set(range(int(low), int(high) + 1))
            printed = {n for k, n, _p, _c in _numbered_blocks(report) if k == kind}
            missing = sorted(spanned - printed)
            if missing:
                fails.append(f"the range {kind}s {low}{CITE_RANGE_DASH}{high} "
                             f"spans {missing}, which the report does not print")
        cited = {n for k, n, _p, _c in _citations(report) if k == "Figure"}
        for _kind, low, high in ranges:
            for number in range(int(low), int(high) + 1):
                if number not in cited:
                    fails.append(f"Figure {number} lies inside a cited range and "
                                 f"is still read as uncited")
    return fails


def test_citations_are_cross_references():
    """Every cited number is a live cross-reference in the prose, and the .docx
    carries the bookmark it lands on."""
    fails = []
    from xslope.report import cite, cite_anchor

    phrase, links = cite("Figure", 4)
    if phrase != "Figure 4" or links != [("Figure 4", "#xslope_figure_4")]:
        fails.append(f"cite('Figure', 4) = {(phrase, links)!r}")
    if cite("Table", 0) != ("", []):
        fails.append("a block with no number cites as something")

    report = _cite_report(REINF_XLSX, ("spencer", "bishop"), {})
    for kind, number, path, block in _citations(report):
        targets = [t for text, t in (block.links or [])
                   if text == f"{kind} {number}"]
        if not targets:
            fails.append(f"{kind} {number} is named under {' > '.join(path)} "
                         f"without a link to it: {block.text!r}")
        elif not any(t.startswith("#") for t in targets):
            fails.append(f"{kind} {number} links to {targets!r}, which is not a "
                         f"bookmark in this document")

    # And the bookmarks those links name are really placed, on the caption of
    # the block whose number they carry.
    from xslope.report_docx import render_docx

    tmp = tempfile.mkdtemp(prefix="xslope_cite_docx_")
    path = render_docx(report, os.path.join(tmp, "cited.docx"))
    with zipfile.ZipFile(path) as z:
        doc = z.read("word/document.xml").decode("utf-8")
    for kind, number, _path, _caption in _numbered_blocks(report):
        name = cite_anchor(kind, number)
        if f'w:name="{name}"' not in doc:
            fails.append(f"{kind} {number}'s caption carries no bookmark for a "
                         f"citation to land on")
        elif f'w:anchor="{name}"' not in doc:
            fails.append(f"nothing in the document links to {kind} {number}")

    # A link covers its own phrase and nothing else. A renderer that moves runs
    # into a hyperlink element by hand can take the words beside them too, or
    # nest one link inside another, and both read as a sentence in which half a
    # paragraph is a link to one figure.
    fails += _link_spans_fails(report, doc)
    fails += _section_reference_fails()
    return fails


def _link_spans_fails(report, doc_xml):
    """Every hyperlink in the document is exactly one of the report's own link
    phrases, and no link is inside another."""
    import re
    fails = []
    declared = {text for block in report.blocks("prose")
                for text, _target in (block.links or [])}
    for link in re.findall(r"<w:hyperlink[ >].*?</w:hyperlink>", doc_xml, re.S):
        if link.count("<w:hyperlink") > 1:
            inner = re.findall(r"<w:t[^>]*>([^<]*)</w:t>", link)
            fails.append(f"a hyperlink is nested inside another: {inner!r}")
            continue
        text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", link))
        if text and text not in declared:
            fails.append(f"a hyperlink covers {text!r}, which is no phrase the "
                         f"report asked to be linked")
    return fails


def _engine_inputs_sentence(report):
    """The Project Definition sentence that says where the inputs each analysis
    reads off the shared model are stated, or None.

    Found by what it says rather than by one of its wordings: with one analysis
    it names that analysis's section directly, and with two it enumerates them.
    """
    for section in report.sections:
        if section.title != "Project Definition":
            continue
        for block in section.blocks:
            if block.kind == "prose" and "reads off this model" in block.text:
                return block
    return None


def test_the_project_definition_sends_the_reader_on():
    """The Project Definition says where the inputs one analysis reads off the
    shared model are stated, by cross-reference to the sections that carry them.

    The section prints the whole model in one figure, so a reader who meets the
    geometry there looks for the strengths, the loads and the members beside it.
    They are under the engines' own headings, because the engines do not read the
    same things off the section: a material's shear strength is a stability input
    and its conductivity is a seepage input, and the members are read for a
    capacity by one engine and a stiffness by the other.

    The sentence names the sections this report actually carries — one engine one
    section, two engines two — and each name is a live cross-reference resolving
    to that section's own number. A report of the flow solution alone prints
    nothing: its conductivities and its boundary conditions are in the only
    analysis section it has.
    """
    fails = []
    from xslope.report import (FEM_ANCHOR, LEM_ANCHOR, SEEPAGE_ANCHOR,
                               section_anchor)

    cases = (
        ("the stability analysis alone",
         (REINF_XLSX, ("spencer", "bishop"), {}, ()), (LEM_ANCHOR,)),
        ("both stability engines on one model",
         (FEM_REINF_XLSX, ("spencer",), {}, ("fem",)),
         (LEM_ANCHOR, FEM_ANCHOR)),
        ("a seepage run beside the stability analysis",
         (SEEP_XLSX, ("spencer",), {}, ("seep",)), (SEEPAGE_ANCHOR, LEM_ANCHOR)),
        ("a strength reduction run", (FEM_XLSX, (), {}, ("fem",)),
         (FEM_ANCHOR,)),
        ("a seepage run on its own", (SEEP_XLSX, (), {}, ("seep",)), ()),
    )
    for name, args, want in cases:
        report = _cite_report(*args)
        block = _engine_inputs_sentence(report)
        if not want:
            if block is not None:
                fails.append(f"{name}: has no analysis to send a reader to and "
                             f"prints {block.text!r}")
            continue
        if block is None:
            fails.append(f"{name}: never says where the inputs each analysis "
                         f"reads off the model are stated")
            continue
        numbers = {sec.anchor: number
                   for number, _lvl, sec in report.section_numbers()}
        targets = [t.lstrip("#") for _text, t in (block.links or [])]
        if targets != [section_anchor(a) for a in want]:
            fails.append(f"{name}: the sentence cites {targets} and the report "
                         f"carries {[section_anchor(a) for a in want]}")
            continue
        for anchor in targets:
            if anchor not in numbers:
                fails.append(f"{name}: cites {anchor}, which is no section of "
                             f"this report: {block.text!r}")
            elif f"Section {numbers[anchor]}" not in block.text:
                fails.append(f"{name}: cites {anchor}, which is Section "
                             f"{numbers[anchor]}, and reads {block.text!r}")

    # The mutation: a citation of a section the report does not carry. It has to
    # stay visible as an unresolved mark rather than reading as a sentence
    # somebody wrote without a number in it.
    from xslope.report import (Prose, Report, Section, cite_section,
                               _resolve_section_citations)
    phrase, links = cite_section("no_such_engine")
    probe = Report(sections=[Section(
        "Project Definition",
        blocks=[Prose(f"The materials are given in {phrase}.", links=links)])])
    _resolve_section_citations(probe)
    said = probe.sections[0].blocks[0].text
    if "\ue000" not in said:
        fails.append(f"a citation of a section the report does not carry "
                     f"resolved to {said!r}")
    return fails


def test_the_material_zones_are_named_in_bold():
    """The Project Definition names the material zones, and each name is set in
    bold where it is named.

    They are the words the legend of every figure and the first column of every
    materials table use, and a reader who meets them here should be able to find
    them again by their shape. Set in the same face as the sentence around them,
    they are three words in a paragraph of forty.

    The names are read off the model, so the bold phrases are checked to be the
    materials the analysis actually references — a sentence that bolded a fixed
    list would bold nothing on somebody else's section — and the document is read
    back to confirm the runs came out bold.
    """
    import re as _re
    fails = []
    from xslope.report import referenced_materials

    slope_data, solutions = _solved()
    names = [str(m.get("name") or "").strip()
             for _i, m in referenced_materials(slope_data)]
    names = [n for n in names if n]
    if len(names) < 2:
        return ["the sample model names fewer than two materials; the sentence "
                "this checks cannot be exercised"]

    with tempfile.TemporaryDirectory() as tmp:
        report, doc, _path = _written(
            "the material zones", slope_data, solutions,
            {"input_path": REINF_XLSX, "title": "Material Zones",
             "method": "spencer", "pd_figure": False,
             "lem_search_figure": False, "lem_slice_key": False,
             "lem_solution_figure": False}, tmp)

    block = next((b for s in report.sections if s.title == "Project Definition"
                  for b in s.blocks
                  if b.kind == "prose" and "zones are" in b.text
                  or b.kind == "prose" and "zone is" in b.text), None)
    if block is None:
        fails.append("the Project Definition never names the material zones")
        return fails
    said = f"The zones are named {_join_names(names)}." if len(names) > 1 else \
        f"The zone is named {names[0]}."
    if said not in block.text:
        fails.append(f"the zones are named as {block.text!r}, not {said!r}")
    if list(getattr(block, "bold", None) or []) != names:
        fails.append(f"the sentence sets {list(getattr(block, 'bold', []) or [])} "
                     f"in bold; the zones are {names}")

    # And the document really came out with them bold: the run that holds each
    # name carries <w:b/>, and the words around it do not.
    for name in names:
        runs = _re.findall(
            r"<w:r>(?:(?!</w:r>).)*?<w:t[^>]*>%s</w:t>" % _re.escape(name),
            doc, _re.S)
        if not runs:
            fails.append(f"{name!r} is not a run of its own in the document; it "
                         f"cannot be the only bold word in the sentence")
        elif not any("<w:b/>" in r for r in runs):
            fails.append(f"{name!r} reaches the document in no bold run")
    if _re.search(r"<w:r>(?:(?!</w:r>).)*?<w:b/>(?:(?!</w:r>).)*?"
                  r"<w:t[^>]*>The problem cross section is defined by</w:t>",
                  doc, _re.S):
        fails.append("the sentence that introduces the zones is itself bold")
    return fails


def _join_names(names):
    """The report's own way of listing names in a sentence."""
    from xslope.report import _join
    return _join(names)


def test_the_factor_of_safety_is_defined_as_a_ratio():
    """The limit equilibrium section defines the factor of safety as the ratio
    the method computes.

    F is the available shear strength along the failure surface over the shear
    required along it to hold the mass in static equilibrium — F = s/τ, which is
    what ``docs/lem/overview.md`` publishes and what every method in the section
    solves for. The section defined it instead as the factor the strength is
    DIVIDED by to bring the mass to limiting equilibrium, which is the strength
    reduction method's definition and not this one's: nothing in this section
    reduces a strength, and a reader who took that definition to the equations
    would not find it in any of them.
    """
    fails = []
    slope_data, solutions = _solved()
    report = _cite_report(REINF_XLSX, ("spencer",), {}, ())
    lem = next((s for s in report.sections
                if s.title == "Limit Equilibrium Analysis"), None)
    if lem is None:
        return ["the report carries no limit equilibrium section"]
    said = next((b.text for b in lem.blocks if b.kind == "prose"
                 and "factor of safety is" in b.text), "")
    if not said:
        return ["the section never says what the factor of safety is"]
    want = ("The factor of safety is the ratio of the shear strength available "
            "along the failure surface to the shear required along it to hold "
            "the sliding mass in static equilibrium.")
    if want not in said:
        fails.append(f"the section defines the factor of safety as {said!r}")
    # And it does not define it the way a strength reduction analysis does.
    for phrase in ("must be divided", "reduced until", "limiting equilibrium"):
        if phrase in said:
            fails.append(f"the definition reads {phrase!r}, which is the "
                         f"strength reduction method's: {said!r}")
    # The page it agrees with says the same thing, in its own symbols.
    with open(os.path.join(_REPO, "docs", "lem", "overview.md"),
              encoding="utf-8") as f:
        page = f.read()
    if "$F = \\dfrac{s}{\\tau}$" not in page:
        fails.append("docs/lem/overview.md no longer publishes F = s/τ; the "
                     "sentence and the page have to be changed together")
    return fails


def test_the_reinforcement_direction_is_named_as_the_column_prints_it():
    """The reinforcement subsection names the two force directions the Direction
    column can hold, and names them the way that column and the documentation do.

    The column prints "tangent" or "axial", and the reinforcement page defines
    them as tangent to the slip surface and along the line's own axis. The
    sentence said the force acts "along the slice base or along the bar", which
    names a shape of reinforcement the setting does not mean: a geosynthetic is
    not a bar, and the tangent case is the one a geosynthetic is analysed under.
    """
    fails = []
    from xslope.report import _REINFORCEMENT_PROSE

    said = _REINFORCEMENT_PROSE["lem"]
    if "bar" in said.split():
        fails.append(f"the sentence still calls the reinforcement a bar: {said!r}")
    want = ("whether the force acts tangent to the slice base or along the axis "
            "of the reinforcement line")
    if want not in said:
        fails.append(f"the sentence names the directions as {said!r}")

    # It reaches the page, over a table whose Direction column prints those very
    # words.
    slope_data, solutions = _solved()
    report = _cite_report(REINF_XLSX, ("spencer",), {}, ())
    table = next((b for top in report.sections for _lvl, node in top.walk()
                  for b in node.blocks if b.kind == "table"
                  and b.caption.startswith("Reinforcement lines")), None)
    if table is None:
        return fails + ["the report prints no reinforcement table"]
    prose = next((b.text for top in report.sections for _lvl, node in top.walk()
                  for b in node.blocks
                  if b.kind == "prose" and "out-of-plane spacing" in b.text), "")
    if want not in prose:
        fails.append(f"the subsection reads {prose!r}")
    column = next((j for j, h in enumerate(table.headers) if h == "Direction"),
                  None)
    if column is None:
        fails.append(f"the reinforcement table has no Direction column: "
                     f"{table.headers}")
    else:
        printed = {str(r[column]) for r in table.rows}
        if not printed <= {"tangent", "axial"}:
            fails.append(f"the Direction column prints {sorted(printed)}, and "
                         f"the sentence describes only tangent and axial")
    # Both values the column can hold are described, in the terms the
    # reinforcement page defines them in — tangent to the slip surface, and along
    # the line's own axis. A model whose lines are all tangent still prints the
    # sentence, and a reader of that model has to be able to tell what the other
    # setting would have meant.
    with open(os.path.join(_REPO, "docs", "lem", "reinforcement.md"),
              encoding="utf-8") as f:
        page = f.read()
    for what, in_prose, on_page in (
            ("tangent", "tangent to the slice base", "Tangent to slip surface"),
            ("axial", "along the axis of the reinforcement line",
             "the inclination of the reinforcement line itself")):
        if in_prose not in prose:
            fails.append(f"the sentence never describes the {what} case: "
                         f"{prose!r}")
        if on_page not in page:
            fails.append(f"docs/lem/reinforcement.md no longer defines the "
                         f"{what} case as {on_page!r}; the sentence and the page "
                         f"have to be changed together")
    return fails


def _section_reference_fails():
    """A citation of a section is a link AND a field: Word computes the number.

    Checked on a report carrying both engines, which is where one section cites
    another — the finite element analysis carries the loads the limit equilibrium
    section already presented, and points at the section rather than printing the
    table twice.
    """
    import re
    fails = []
    from xslope.report import section_anchor

    report = _cite_report(FEM_REINF_XLSX, ("spencer",), {}, ("fem",))
    cited = _section_citations(report)
    if not cited:
        return ["a report of both engines cites no section, so the section "
                "cross-reference is never exercised"]

    numbers = {sec.anchor: number for number, _lvl, sec in report.section_numbers()}
    for number, path, block in cited:
        targets = [t for text, t in (block.links or [])
                   if text == f"Section {number}"]
        if not targets:
            fails.append(f"Section {number} is named under {' > '.join(path)} "
                         f"without a link to it: {block.text!r}")
        elif not any(t.startswith("#" + section_anchor("")) for t in targets):
            fails.append(f"Section {number} links to {targets!r}, which is no "
                         f"heading of this document")
        elif numbers.get(targets[0][1:]) != number:
            fails.append(f"the citation reads Section {number} but points at "
                         f"{targets[0]}, which is Section "
                         f"{numbers.get(targets[0][1:])!r}")

    from xslope.report_docx import render_docx
    tmp = tempfile.mkdtemp(prefix="xslope_secref_")
    path = render_docx(report, os.path.join(tmp, "sections.docx"))
    with zipfile.ZipFile(path) as z:
        doc = z.read("word/document.xml").decode("utf-8")

    # This report puts three figure links in one sentence, which is where a
    # renderer that moves runs by hand takes the words between them.
    fails += _link_spans_fails(report, doc)

    # Every heading is bookmarked, so a reference written later has a target.
    marked = set(re.findall(r'w:name="(xslope_section_[^"]*)"', doc))
    for _number, _lvl, sec in report.section_numbers():
        if sec.anchor not in marked:
            fails.append(f"the heading {sec.title!r} carries no bookmark")

    # And each citation is a REF field on that bookmark, printing its number,
    # with the number Word will compute already cached in it.
    for number, _path, block in cited:
        anchor = next((t[1:] for text, t in (block.links or [])
                       if text == f"Section {number}"), "")
        if not anchor:
            continue
        instr = f"REF {anchor} \\r"
        if instr not in doc:
            fails.append(f"Section {number} is written as text, not as a "
                         f"cross-reference field: no {instr!r} in the document")
            continue
        if f'w:anchor="{anchor}"' not in doc:
            fails.append(f"the field for Section {number} is in no hyperlink, so "
                         f"clicking the reference goes nowhere")
        start = doc.index(instr)
        after = doc.index('w:fldCharType="separate"', start)
        end = doc.index('w:fldCharType="end"', after)
        cached = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", doc[after:end]))
        if cached != number:
            fails.append(f"the field for Section {number} carries {cached!r} as "
                         f"its result; a reader sees that before Word updates it")
    return fails


# --------------------------------------------------------------------------
# G. the dialog
# --------------------------------------------------------------------------

def _app():
    from PySide6.QtWidgets import QApplication
    return QApplication.instance() or QApplication([])


def test_dialog():
    """The dialog constructs offscreen, and its controls reach the options."""
    fails = []
    _app()
    from PySide6.QtCore import Qt
    from studio.report_dialog import ReportDialog, default_output_path

    _slope_data, solutions = _solved()
    dlg = ReportDialog(slope_data=_slope_data, solutions=solutions,
                       model_path=REINF_XLSX, default_method="bishop")
    try:
        if dlg.output_format() != "docx":
            fails.append(f"the dialog opens on {dlg.output_format()!r}")
        if not dlg.output_path().endswith("_report.docx"):
            fails.append(f"the default output path is {dlg.output_path()!r}")
        # The Methods list is a MULTI-select over every method the solver offers,
        # opening ticked on the one the results view is showing.
        from xslope.report import supported_methods
        offered = [i.data(Qt.UserRole) for i in dlg._method_items()]
        if offered != supported_methods():
            fails.append(f"the methods list offers {offered}, not the solver's "
                         f"own {supported_methods()}")
        if dlg.selected_methods() != ["bishop"]:
            fails.append(f"the methods list opens on {dlg.selected_methods()}, "
                         f"not the results view's method")
        if dlg.options().get("method") != ["bishop"]:
            fails.append(f"the options carry {dlg.options().get('method')!r}")

        # Ticking a second one reports both, in list order.
        items = {i.data(Qt.UserRole): i for i in dlg._method_items()}
        items["spencer"].setCheckState(Qt.Checked)
        got = dlg.selected_methods()
        if got != ["bishop", "spencer"]:
            fails.append(f"two ticked methods came back as {got}")

        # And at least one is always ticked: unticking the last re-ticks it.
        for name in got:
            items[name].setCheckState(Qt.Unchecked)
        if not dlg.selected_methods():
            fails.append("every method could be unticked; a report with no "
                         "method has no results in it")
        items["bishop"].setCheckState(Qt.Checked)
        for name in dlg.selected_methods():
            if name != "bishop":
                items[name].setCheckState(Qt.Unchecked)

        # Only DOCX can be picked in S1; the rest are listed and dimmed.
        for i in range(dlg.format.count()):
            key = dlg.format.itemData(i)
            enabled = dlg.format.model().item(i).isEnabled()
            if (key == "docx") != enabled:
                fails.append(f"format {key!r} is {'enabled' if enabled else 'dimmed'}")

        opts = dlg.options()
        for key in ("traceability", "project_definition", "lem",
                    "lem_slice_table", "lem_materials"):
            if opts.get(key) is not True:
                fails.append(f"{key} is not on by default")
        if opts.get("signature_lines") is not False:
            fails.append("signature lines are on by default")
        # The boxes open on the builder's own defaults — one declaration, not two.
        from xslope.report import DEFAULT_OPTIONS
        if "model_checks" in dlg._items:
            fails.append("the dialog still offers a Model checks row; the "
                         "findings belong to the interface, not the report")
        for key in dlg._items:
            if opts.get(key) is not bool(DEFAULT_OPTIONS.get(key, True)):
                fails.append(f"the {key!r} box opens on {opts.get(key)!r}, not the "
                             f"builder's default "
                             f"{bool(DEFAULT_OPTIONS.get(key, True))!r}")

        # The point-coordinate labels are a row of their own under the model
        # figure, on by default, and the box moves the option it names.
        coords = dlg._items.get("pd_coords")
        if coords is None:
            fails.append("the dialog has no Point coordinates row")
        else:
            if coords.parent() is not dlg._items["project_definition"]:
                fails.append("the Point coordinates row is not under Project "
                             "definition")
            if "Point coordinates" not in coords.text(0):
                fails.append(f"the coordinates row reads {coords.text(0)!r}")
            if opts.get("pd_coords") is not True:
                fails.append("the point coordinates open off; geometry is one of "
                             "the things the model figure is for")
            for state in (Qt.Unchecked, Qt.Checked):
                coords.setCheckState(0, state)
                want = state == Qt.Checked
                if dlg.options().get("pd_coords") is not want:
                    fails.append(f"the coordinates box set to {want} came back "
                                 f"as {dlg.options().get('pd_coords')!r}")
            # And it is the figure's option: it never carries the figure with it.
            coords.setCheckState(0, Qt.Unchecked)
            if dlg.options().get("pd_figure") is not True:
                fails.append("turning the coordinate labels off took the model "
                             "figure with them")
            coords.setCheckState(0, Qt.Checked)

        # A toggle reaches the options, and a parent takes its children.
        dlg._items["lem_slice_table"].setCheckState(0, Qt.Unchecked)
        if dlg.options().get("lem_slice_table") is not False:
            fails.append("unchecking the slice table did not reach the options")
        dlg._items["project_definition"].setCheckState(0, Qt.Unchecked)
        after = dlg.options()
        if after.get("pd_coords") is not False:
            fails.append("a section turned off left its children on")
        if not dlg._items["pd_coords"].isDisabled():
            fails.append("a section turned off left its children live")
        # And took nothing that is not its own: the reinforcement is read by the
        # stability analysis and its box is that analysis's.
        if after.get("lem_members") is not True:
            fails.append("turning Project definition off took the limit "
                         "equilibrium section's member properties with it")

        # And the options it produces really build a report.
        from xslope.report import build_report
        dlg._items["project_definition"].setCheckState(0, Qt.Checked)
        opts = dlg.options()
        opts["lem_search_figure"] = False
        with tempfile.TemporaryDirectory() as tmp:
            report = build_report(_slope_data, solutions, opts, tmp)
        titles = [t for _l, t in report.section_titles()]
        if "Slice Table" in titles:
            fails.append("the dialog's options did not carry the slice table off")
        if "Reinforcement" not in titles:
            fails.append("the dialog's options lost the reinforcement section")
    finally:
        dlg.close()

    # Without a saved project the path still points somewhere writable.
    if not default_output_path(None).endswith(".docx"):
        fails.append("an unsaved project gets no default report path")
    return fails


def test_dialog_settings():
    """What the dialog remembers survives a second construction."""
    fails = []
    _app()
    from PySide6.QtCore import QSettings, Qt
    from studio.report_dialog import SETTINGS_PREFIX, ReportDialog

    _slope_data, solutions = _solved()
    with tempfile.TemporaryDirectory() as tmp:
        ini = os.path.join(tmp, "settings.ini")
        settings = QSettings(ini, QSettings.IniFormat)

        first = ReportDialog(slope_data=_slope_data, solutions=solutions,
                             model_path=REINF_XLSX, settings=settings)
        first.organization.setText("Example Engineering")
        first.author.setText("A. Engineer")
        first.signature_lines.setChecked(True)
        first._items["lem_slice_table"].setCheckState(0, Qt.Unchecked)
        for item in first._method_items():
            if item.data(Qt.UserRole) in ("bishop", "corps"):
                item.setCheckState(Qt.Checked)
            elif item.flags() & Qt.ItemIsUserCheckable:
                item.setCheckState(Qt.Unchecked)
        first.title.setText("A One-Off Title")
        first.remember()
        first.close()

        again = ReportDialog(slope_data=_slope_data, solutions=solutions,
                             model_path=REINF_XLSX, settings=settings)
        try:
            if again.organization.text() != "Example Engineering":
                fails.append(f"the organization was not remembered "
                             f"({again.organization.text()!r})")
            if again.author.text() != "A. Engineer":
                fails.append("the author was not remembered")
            if not again.signature_lines.isChecked():
                fails.append("the signature-lines choice was not remembered")
            if again._items["lem_slice_table"].checkState(0) != Qt.Unchecked:
                fails.append("the content selections were not remembered")
            if again._items["traceability"].checkState(0) != Qt.Checked:
                fails.append("a selection that was left on came back off")
            if again.selected_methods() != ["bishop", "corps"]:
                fails.append(f"the methods were not remembered "
                             f"({again.selected_methods()})")
            # The project's own fields are NOT carried between projects.
            if again.title.text() == "A One-Off Title":
                fails.append("the project title was remembered; it belongs to the "
                             "project, not to the user")
        finally:
            again.close()

        # And it really went to disk, under this dialog's own prefix.
        stored = QSettings(ini, QSettings.IniFormat)
        if stored.value(SETTINGS_PREFIX + "author") != "A. Engineer":
            fails.append("the remembered fields are not under the report prefix")

        # A dialog with no settings store remembers nothing and does not fail.
        plain = ReportDialog(slope_data=_slope_data, solutions=solutions,
                             model_path=REINF_XLSX)
        try:
            if plain.author.text():
                fails.append("a dialog with no settings pre-filled the author")
            plain.remember()
        finally:
            plain.close()
    return fails


def test_open_output():
    """The finished document is opened, and the formats offered are the ones
    that exist.

    A LaTeX row that could never be picked was a promise on the screen where the
    choice is made; the report is a Word document, with PDF still to come.
    """
    fails = []
    _app()
    from studio import report_dialog

    formats = [key for key, *_rest in report_dialog.FORMATS]
    if "latex" in formats:
        fails.append(f"LaTeX is still offered: {formats}")
    if formats != ["docx", "pdf"]:
        fails.append(f"the formats offered are {formats}, not ['docx', 'pdf']")

    opened = []
    real = report_dialog.QDesktopServices.openUrl
    report_dialog.QDesktopServices.openUrl = lambda url: opened.append(url.toLocalFile())
    try:
        with tempfile.TemporaryDirectory() as tmp:
            docx = os.path.join(tmp, "r.docx")
            open(docx, "wb").close()
            if report_dialog.open_output(docx, "docx") != "document":
                fails.append("a .docx did not open as a document")
            if opened[-1] != docx:
                fails.append(f"the document opened was {opened[-1]!r}")
    finally:
        report_dialog.QDesktopServices.openUrl = real
    return fails


#: What a runner check builds: two methods, so the figures the progress bar
#: counts are several and two of them are the same figure of different methods,
#: at a resolution that costs nothing to render.
_RUNNER_OPTIONS = {"input_path": REINF_XLSX, "title": "Sample Levee",
                   "method": ["spencer", "bishop"], "lem_search_figure": False,
                   "dpi": 60, "figsize": (4.0, 2.5)}


def _run_report_runner(path, options=None, cancel_after=None, fmt=None):
    """Drive a :class:`studio.runners.ReportRunner` to completion and collect
    what it emitted.

    ``run()`` is called on this thread rather than started as a thread: the
    signals then arrive as direct calls, so the check reads them in order
    without an event loop. What is being checked is what the worker emits, which
    is the same either way.

    ``cancel_after`` cancels the run from inside the progress handler once that
    many steps have arrived — which is what the Cancel button does, at the point
    in the build where it does it. The Word finish is never asked for.
    """
    from studio.runners import ReportRunner

    slope_data, solutions = _solved()
    opts = dict(_RUNNER_OPTIONS)
    opts.update(options or {})
    runner = ReportRunner(slope_data, solutions, opts, path, fmt=fmt,
                          finalize=False)
    seen = {"progress": [], "ok": [], "failed": [], "cancelled": []}

    def on_progress(done, total, label):
        seen["progress"].append((done, total, label))
        if cancel_after is not None and len(seen["progress"]) >= cancel_after:
            runner.cancel()

    runner.progress.connect(on_progress)
    runner.succeeded.connect(lambda out: seen["ok"].append(out))
    runner.failed.connect(lambda msg: seen["failed"].append(msg))
    runner.cancelled.connect(lambda: seen["cancelled"].append(True))
    runner.run()
    return seen


def test_report_runner_progress():
    """The report builds on a worker thread, counting the figures by name.

    The bar is determinate over the figures plus the document itself, every
    figure the build renders is announced with its own label, and the count the
    bar is ranged on is the count the builder plans — the same agreement
    ``planned_figures`` promises, arriving where the user can see it.
    """
    fails = []
    _app()
    import matplotlib
    matplotlib.use("Agg")
    from studio.runners import (REPORT_FINALIZE_LABEL, REPORT_WRITE_LABEL,
                                REPORT_WRITE_STEPS)
    from xslope.report import planned_figures, resolve_options

    slope_data, solutions = _solved()
    planned = planned_figures(slope_data, solutions,
                              resolve_options(dict(_RUNNER_OPTIONS)))
    if planned < 2:
        fails.append(f"the runner check plans {planned} figures; it is meant to "
                     f"exercise several")

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "runner.docx")
        seen = _run_report_runner(path)

        if seen["failed"] or seen["cancelled"]:
            return fails + [f"the run did not succeed: {seen['failed']} "
                            f"{seen['cancelled']}"]
        if len(seen["ok"]) != 1:
            fails.append(f"succeeded fired {len(seen['ok'])} times")
        out = seen["ok"][0] if seen["ok"] else {}
        if not os.path.isfile(path):
            fails.append("the runner emitted success with no document written")
        if out.get("path") != path:
            fails.append(f"the result names {out.get('path')!r}")
        if out.get("finalized") is not False:
            fails.append("a run that was not asked to finish in Word says it did")

        steps = seen["progress"]
        total = planned + REPORT_WRITE_STEPS
        if {t for _d, t, _l in steps} != {total}:
            fails.append(f"the bar was ranged on {sorted({t for _d, t, _l in steps})}, "
                         f"not {total} (= {planned} figures + the document)")
        # The figure steps: one per figure, numbered 1..N in order.
        figure_steps = [(d, l) for d, _t, l in steps
                        if l not in (REPORT_WRITE_LABEL, "rendering the figures")]
        if [d for d, _l in figure_steps] != list(range(1, planned + 1)):
            fails.append(f"the figure steps are {[d for d, _l in figure_steps]}, "
                         f"not 1..{planned}")
        # And they are the figures that were built: the same count, each named.
        drawn = out.get("figures") or []
        if len(figure_steps) != len(drawn):
            fails.append(f"{len(figure_steps)} steps were announced and "
                         f"{len(drawn)} figures reached the document")
        if any(not l.strip() for _d, l in figure_steps):
            fails.append("a figure was announced with no label")
        labels = [l for _d, l in figure_steps]
        # Two methods mean the per-method figures are announced per method, so
        # the labels name which one is rendering rather than repeating.
        for name in ("Spencer", "Bishop"):
            if not any(name in l for l in labels):
                fails.append(f"no step names {name}: {labels}")
        if len(set(labels)) != len(labels):
            fails.append(f"two steps carry the same label: {labels}")

        # The document is a step of its own, named, and the bar reaches the end
        # on it. Word was not asked for and is not announced.
        write = [(d, t) for d, t, l in steps if l == REPORT_WRITE_LABEL]
        if not write:
            fails.append("the document write is not a phase of its own")
        elif write[-1] != (total, total):
            fails.append(f"the bar ends at {write[-1]}, not ({total}, {total})")
        if any(l == REPORT_FINALIZE_LABEL for _d, _t, l in steps):
            fails.append("a run that was not asked to finish in Word announced it")
    return fails


def test_report_runner_cancel():
    """Cancel stops the build at the next figure, and leaves nothing behind.

    The builder takes no cancel argument; what it has is a progress callback it
    calls before each figure, guarded by ``except Exception`` so that a broken
    progress line can never cost a report. ``ReportCancelled`` is a
    ``BaseException`` for exactly that reason — it is the one thing that guard
    lets through — and the builder's own ``finally`` removes the figures on the
    way out. Widening either guard would turn Cancel into a button that does
    nothing, so the seam itself is checked here, not just the button.
    """
    fails = []
    _app()
    import matplotlib
    matplotlib.use("Agg")
    import glob
    from studio.runners import ReportCancelled

    if isinstance(ReportCancelled(), Exception):
        fails.append("ReportCancelled is an Exception; the builder's progress "
                     "guard would swallow it and the run would not stop")

    # The builder renders into a temporary directory of its own and removes it
    # on the way out; a cancelled build must do the same. Counted as a
    # difference, because other checks here keep their figure directories.
    figures = os.path.join(tempfile.gettempdir(), "xslope_report_*")
    before = set(glob.glob(figures))

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "cancelled.docx")
        # Two steps in: the first figure has been announced, so the next one the
        # builder announces is where the run unwinds.
        seen = _run_report_runner(path, cancel_after=2)
        if not seen["cancelled"]:
            fails.append("cancelling mid-build did not report a cancelled run")
        if seen["ok"]:
            fails.append("a cancelled run reported success")
        if seen["failed"]:
            fails.append(f"a cancelled run reported a failure: {seen['failed']}")
        if os.path.exists(path):
            fails.append("a cancelled run left a half-written document")
        # It stopped where it was cancelled rather than running to the end.
        from xslope.report import planned_figures, resolve_options
        slope_data, solutions = _solved()
        planned = planned_figures(slope_data, solutions,
                                  resolve_options(dict(_RUNNER_OPTIONS)))
        if len(seen["progress"]) > planned:
            fails.append(f"the cancelled run announced {len(seen['progress'])} "
                         f"steps of {planned} figures; it did not stop")

    leaked = set(glob.glob(figures)) - before
    if leaked:
        fails.append(f"a cancelled build left its figures behind: {sorted(leaked)}")
    return fails


def test_report_runner_failure():
    """A build that cannot produce a document says so, in the builder's own
    words, instead of succeeding quietly or taking the window down with it."""
    fails = []
    _app()
    import matplotlib
    matplotlib.use("Agg")

    # A format that does not exist yet: refused before anything is rendered.
    seen = _run_report_runner("nowhere.pdf", fmt="pdf")
    if seen["ok"]:
        fails.append("a report in an unavailable format reported success")
    if not seen["failed"]:
        fails.append("an unavailable format produced no failure message")
    elif "not available" not in seen["failed"][0]:
        fails.append(f"the failure reads {seen['failed'][0]!r}")

    # A document that cannot be written: the figures render, the write fails,
    # and the message is the one the builder produced.
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "no_such_folder", "r.docx")
        with contextlib.redirect_stdout(io.StringIO()):
            seen = _run_report_runner(
                path, options={"lem": False, "project_definition": False,
                               "seep": False, "fem": False})
        if seen["ok"]:
            fails.append("a report that could not be written reported success")
        if not seen["failed"]:
            fails.append("an unwritable path produced no failure message")
        elif "could not be written" not in seen["failed"][0]:
            fails.append(f"the write failure reads {seen['failed'][0]!r}")
        if not seen["progress"]:
            fails.append("a failing run showed no progress at all")
    return fails


def test_report_runs_off_the_gui_thread():
    """Generating a report never blocks the window.

    The window hands the build to the runner and waits for its signals: no
    ``generate_report`` on the GUI thread, and no wait cursor standing in for a
    progress bar. While it runs, the report action is dimmed with the rest of
    the run actions, so a second report cannot be started over the first, and
    Cancel reaches the report runner like any other.
    """
    fails = []
    import inspect
    import matplotlib
    matplotlib.use("Agg")
    _app()
    from studio import report_dialog
    from studio.main_window import MainWindow
    from studio.runners import ReportRunner

    src = inspect.getsource(MainWindow.generate_report)
    if "ReportRunner(" not in src:
        fails.append("the window does not hand the report to the runner")
    if "generate_report as _generate" in src or "_generate(" in src:
        fails.append("the window still builds the report on the GUI thread")
    if "setOverrideCursor" in src:
        fails.append("the report still puts a wait cursor up; the window stays "
                     "live now")
    # The dialog collects options; it does not build reports either.
    dialog_src = inspect.getsource(report_dialog)
    if "generate_report(" in dialog_src:
        fails.append("the dialog builds the report itself")
    # The Word finish the runner calls has no Qt in it.
    finish_src = inspect.getsource(report_dialog.word_finish)
    for qt in ("QApplication", "setOverrideCursor", "_status_line"):
        if qt in finish_src:
            fails.append(f"word_finish touches {qt} off the GUI thread")

    slope_data, solutions = _solved()
    mw = MainWindow()
    try:
        mw.doc.slope_data = slope_data
        mw.doc.results["lem_solution"] = solutions["lem"][0]
        mw._last_lem_opts = {"method": "spencer"}
        mw._update_run_actions()
        if not mw.act_report.isEnabled():
            fails.append("the report action is dimmed with a solved model")

        # A run in flight dims it, and the Run action with it.
        mw._report_runner = ReportRunner(slope_data, solutions, {}, "x.docx",
                                         finalize=False)
        mw._update_run_actions()
        if mw.act_report.isEnabled():
            fails.append("a second report can be started over the first")
        if mw.act_run.isEnabled():
            fails.append("an analysis can be started while the report builds")
        # Cancel goes to the report runner — the one in flight, which here is
        # this one and not started, so it says so.
        cancelled = []
        mw._report_runner.cancel = lambda: cancelled.append(True)
        mw._report_runner.isRunning = lambda: True
        mw.cancel_btn.setVisible(True)
        mw._cancel_run()
        if not cancelled:
            fails.append("Cancel does not reach the report runner")
        # Word's stretch is indeterminate and takes Cancel out with it.
        mw.cancel_btn.setEnabled(True)
        mw._on_report_progress(0, -1, "finalizing…")
        if mw.cancel_btn.isEnabled():
            fails.append("Cancel is still live while Word holds the document")

        mw._report_runner = None
        mw._update_run_actions()
        if not mw.act_report.isEnabled():
            fails.append("the report action stayed dimmed after the run")
    finally:
        mw._report_runner = None
        mw.close()
    return fails


def _noncircular_solutions():
    """The cached solve, presented as a non-circular surface.

    A slice table states its own surface family through the circle radius it
    carries, and blanking that is what makes the same slices non-circular — so
    this costs no second solve and exercises the path the dialog and the report
    both read (``surface_family``).
    """
    slope_data, solutions = _solved()
    bundle = dict(solutions["lem"][0])
    df = bundle["slice_df"].copy()
    df["r"] = float("nan")
    bundle["slice_df"] = df
    bundle["search"] = None
    return dict(slope_data, circular=False), {"lem": [bundle]}


def test_noncircular_dims_the_moment_methods():
    """On a non-circular surface the circular-only methods are dimmed, and the
    list still opens on something the surface can take.

    A moment method needs a center of rotation. Offering one on a non-circular
    surface invites a report section that can only say no, and the ways it goes
    wrong are quiet: the results view may be SHOWING a circular-only method when
    the dialog opens, and a remembered selection may name nothing this surface
    can run. Both have to land on a method that works, because the list always
    keeps one ticked.
    """
    fails = []
    _app()
    from PySide6.QtCore import QSettings, Qt
    from studio.report_dialog import SETTINGS_PREFIX, ReportDialog
    from xslope.preflight import method_surface_reason
    from xslope.report import supported_methods, surface_family

    slope_data, solutions = _noncircular_solutions()
    if surface_family(slope_data, solutions) != "noncircular":
        return ["the fixture is not read as a non-circular surface"]

    circular_only = [n for n in supported_methods()
                     if method_surface_reason(n, "noncircular")]
    if not circular_only:
        return ["no method is circular-only; the check proves nothing"]

    dlg = ReportDialog(slope_data=slope_data, solutions=solutions,
                       model_path=REINF_XLSX, default_method="spencer")
    try:
        for item in dlg._method_items():
            name = item.data(Qt.UserRole)
            checkable = bool(item.flags() & Qt.ItemIsUserCheckable)
            if (name in circular_only) == checkable:
                fails.append(f"{name} is {'live' if checkable else 'dimmed'} on a "
                             f"non-circular surface")
            if name in circular_only:
                if item.checkState() == Qt.Checked:
                    fails.append(f"{name} opens ticked but cannot run")
                if not item.toolTip().startswith(
                        method_surface_reason(name, "noncircular")[:30]):
                    fails.append(f"{name} is dimmed without saying why: "
                                 f"{item.toolTip()!r}")
        # The same model, circular: the very same methods are live again, so the
        # dimming is the surface's doing and not a permanently disabled row.
        circular = ReportDialog(slope_data=dict(slope_data, circular=True),
                                solutions=_solved()[1], model_path=REINF_XLSX)
        try:
            live = [i.data(Qt.UserRole) for i in circular._method_items()
                    if i.flags() & Qt.ItemIsUserCheckable]
            if live != supported_methods():
                fails.append(f"a circular surface dims {set(supported_methods()) - set(live)}")
        finally:
            circular.close()
    finally:
        dlg.close()

    # The results view showing a circular-only method must not leave the dialog
    # opening on one — nor on an arbitrary method, when one that was RUN works.
    dlg = ReportDialog(slope_data=slope_data, solutions=solutions,
                       model_path=REINF_XLSX, default_method="bishop")
    try:
        if dlg.selected_methods() != ["spencer"]:
            fails.append(f"with the view on Bishop and Spencer run, the list "
                         f"opened on {dlg.selected_methods()}")
    finally:
        dlg.close()

    # And a remembered selection this surface cannot take falls back rather than
    # leaving the dialog with nothing ticked.
    with tempfile.TemporaryDirectory() as tmp:
        settings = QSettings(os.path.join(tmp, "s.ini"), QSettings.IniFormat)
        settings.setValue(SETTINGS_PREFIX + "methods", list(circular_only))
        dlg = ReportDialog(slope_data=slope_data, solutions=solutions,
                           model_path=REINF_XLSX, default_method="spencer",
                           settings=settings)
        try:
            got = dlg.selected_methods()
            if not got:
                fails.append("a remembered circular-only selection left the "
                             "dialog with no method ticked")
            if set(got) & set(circular_only):
                fails.append(f"the dialog restored {got}, which this surface "
                             f"cannot run")
        finally:
            dlg.close()

    # Asked for one anyway, the report says so under its own heading rather than
    # dropping the block — a missing section reads as an answer withheld.
    from xslope.report import build_report, method_label
    with tempfile.TemporaryDirectory() as tmp:
        report = build_report(slope_data, solutions,
                              {"method": [circular_only[0], "spencer"],
                               "pd_figure": False, "lem_search_figure": False,
                               "lem_solution_figure": False,
                               "lem_slice_key": False}, tmp)
    titles = [t for _l, t in report.section_titles()]
    head = method_label(circular_only[0])
    # Asked for one anyway, the report documents the method that WAS run and
    # says nothing about the one that was not: a report describes the analysis,
    # and the dialog above has already told the user why the method is dimmed —
    # which is where a question about what can be run belongs.
    if head in titles:
        fails.append(f"{circular_only[0]} got a section on a surface it cannot "
                     f"run on: {titles}")
    if method_label("spencer") not in titles:
        fails.append(f"the method that was run is not documented: {titles}")
    if [t for t in report.tables()
            if t.landscape and method_label(circular_only[0]) in t.caption]:
        fails.append(f"{circular_only[0]} got a slice table on a surface it "
                     f"cannot run on")
    return fails


def test_slice_numbers_display_option():
    """Studio's LEM solution view can label the slices, and the toggle really
    reaches the plot.

    The same labels the report's slice key carries. A display checkbox wired to
    nothing looks exactly like one that works, so this follows the value from the
    panel through the display-options dict to the keyword the canvas hands the
    plotting function.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    _app()
    from studio import canvas as canvas_mod
    from studio.display_panels import SolutionDisplayPanel

    panel = SolutionDisplayPanel()
    try:
        if "slice_numbers" not in panel.options():
            return ["the LEM solution display panel has no Slice numbers option"]
        if panel.options()["slice_numbers"] is not False:
            fails.append("slice numbers are on by default")
        box = panel._boxes["slice_numbers"]
        if box.text() != "Slice numbers":
            fails.append(f"the checkbox reads {box.text()!r}")

        # The canvas defers its raster until it is visible, so the stored draw
        # function is called here with a figure of its own — the same call the
        # canvas makes, without needing a screen.
        from matplotlib.figure import Figure as MplFigure

        seen = []
        real = canvas_mod.plot_solution
        canvas_mod.plot_solution = lambda *a, **kw: seen.append(kw)
        try:
            slope_data, solutions = _solved()
            bundle = solutions["lem"][0]
            view = canvas_mod.MplCanvas()
            for state in (True, False):
                box.setChecked(state)
                view.render_solution(slope_data, bundle["slice_df"],
                                     bundle["failure_surface"], bundle["results"],
                                     opts=panel.options())
                view._draw_fn(MplFigure())
            view.deleteLater()
        finally:
            canvas_mod.plot_solution = real

        got = [kw.get("slice_numbers") for kw in seen]
        if got != [True, False]:
            fails.append(f"the canvas was asked for slice_numbers={got}, not "
                         f"[True, False] — the toggle does not reach the plot")
    finally:
        panel.deleteLater()
    return fails


def test_main_window_action():
    """The menu item and toolbar button exist, and are gated on a solution."""
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    _app()
    from studio.main_window import MainWindow

    from PySide6.QtWidgets import QMenu, QToolBar

    _slope_data, solutions = _solved()
    mw = MainWindow()
    try:
        act = getattr(mw, "act_report", None)
        if act is None:
            return ["there is no Generate Report action"]
        if "Report" not in act.text():
            fails.append(f"the action reads {act.text()!r}")
        menus = [m for m in mw.findChildren(QMenu) if act in m.actions()]
        if not any("File" in m.title() for m in menus):
            fails.append(f"the action is not on the File menu "
                         f"(found on {[m.title() for m in menus]})")
        if not any(act in tb.actions() for tb in mw.findChildren(QToolBar)):
            fails.append("the action is not on a toolbar")
        if act.isEnabled():
            fails.append("the action is live with nothing solved")
        if "Run an analysis" not in act.toolTip():
            fails.append(f"the dimmed action gives no reason: {act.toolTip()!r}")

        mw.doc.slope_data = _slope_data
        mw.doc.results["lem_solution"] = solutions["lem"][0]
        mw._last_lem_opts = {"method": "bishop"}
        mw._update_run_actions()
        if not act.isEnabled():
            fails.append("the action stays dimmed with a solved model")
        got = mw.report_solutions()
        if not got.get("lem"):
            fails.append("the window offers no LEM solution to the report")
        elif got["lem"][0].get("method") != "bishop":
            fails.append(f"the solution carries method "
                         f"{got['lem'][0].get('method')!r}, not the run's")
    finally:
        mw.close()
    return fails


def _handed_over(slope_data, results):
    """What a main window with ``results`` in its document hands the report."""
    from studio.main_window import MainWindow
    mw = MainWindow()
    try:
        mw.doc.slope_data = slope_data
        mw.doc.results.update(results)
        return mw.report_solutions()
    finally:
        mw.close()


def test_report_solutions_carry_every_engine():
    """Every engine the session has solved reaches the report.

    Studio runs three: limit equilibrium, seepage, and the finite element
    analysis. The report has a section for each, and each reads its engine's own
    bundle — so a window that hands over the LEM solution alone documents a
    seepage run and a strength reduction run nowhere, whatever was solved in
    them.

    What the sidecar says was run is part of it. The metadata Studio writes
    beside a solution names it ``analysis``, the metadata the benchmark figures
    were built with names it ``analysis_type``, and a restored strength reduction
    run that arrives as neither is documented as a deformation analysis with no
    factor of safety — which is a run that reached 1.345 reported as a run that
    reached nothing.
    """
    fails = []
    import matplotlib
    matplotlib.use("Agg")
    _app()
    from studio.main_window import MainWindow
    from xslope.report import (build_report, fem_bundles, lem_bundles,
                               seep_bundles)

    slope_data, solutions = _solved()
    seep_data, seep = _seep_bundle()
    fem_data, fem = _fem_bundle()

    mw = MainWindow()
    try:
        mw.doc.slope_data = slope_data
        mw.doc.results["lem_solution"] = solutions["lem"][0]
        mw._last_lem_opts = {"method": "bishop"}
        # Two boundary condition sets, put in out of order: the report documents
        # them in the order it is handed them.
        mw.doc.results["seep_solutions"] = {
            2: dict(seep, options={"bc": 2}), 1: seep}
        mw.doc.results["fem_solution"] = fem

        got = mw.report_solutions()
        for engine, bundles in (("lem", lem_bundles(got)),
                                ("seep", seep_bundles(got)),
                                ("fem", fem_bundles(got))):
            if not bundles:
                fails.append(f"a session that solved {engine} hands the report "
                             f"nothing for it: {sorted(got)}")
        if len(seep_bundles(got)) != 2:
            fails.append(f"two boundary condition sets were solved and "
                         f"{len(seep_bundles(got))} reached the report")
        bcs = [b.get("options", {}).get("bc") for b in seep_bundles(got)]
        if bcs != sorted(b for b in bcs if b is not None):
            fails.append(f"the seepage sets reach the report in the order {bcs}")

        # A solution from any engine is a report: the action no longer waits on
        # a limit equilibrium run.
        for key in ("lem_solution", "seep_solutions"):
            mw.doc.results.pop(key, None)
        mw._update_run_actions()
        if not mw.act_report.isEnabled():
            fails.append("a strength reduction run alone leaves Generate Report "
                         "dimmed, and its section is the report")
    finally:
        mw.close()

    # And what is handed over is what the sections are built from.
    for engine, model, xlsx, results, title in (
            ("seep", seep_data, SEEP_XLSX, {"seep_solutions": {1: seep}},
             "Seepage Analysis"),
            ("fem", fem_data, FEM_XLSX, {"fem_solution": fem},
             "Deformation and Strength Reduction")):
        opts = {"input_path": xlsx, "lem": False, "pd_figure": False}
        opts.update(FAST_FIGURES)
        tmp = tempfile.mkdtemp(prefix=f"xslope_studio_{engine}_")
        with contextlib.redirect_stdout(io.StringIO()):
            built = build_report(model, _handed_over(model, results), opts, tmp)
        if title not in _titles(built):
            fails.append(f"the {engine} bundle Studio hands over builds "
                         f"{_titles(built)}, with no {title!r} section")

    # The restored run: the sidecar of the sample SSRM model, rewritten to name
    # what was run the way the benchmark builder names it.
    from xslope.fileio import load_slope_data
    with tempfile.TemporaryDirectory() as tmp:
        stem = _sidecar_copy(
            os.path.splitext(FEM_XLSX)[0], tmp,
            lambda meta: {("analysis_type" if k == "analysis" else k): v
                          for k, v in meta.items()})
        mw = MainWindow()
        try:
            with contextlib.redirect_stdout(io.StringIO()):
                mw.doc.slope_data = load_slope_data(f"{stem}.xlsx")
                mw._restore_fem_sidecar(mw.doc.slope_data["mesh"], stem)
            restored = mw.doc.results.get("fem_solution")
            if not restored:
                fails.append("the FEM sidecar beside the sample model did not "
                             "restore, so what it says was run is untested")
            elif str(restored.get("analysis")) != "ssrm":
                fails.append(f"a restored strength reduction run says it was a "
                             f"{restored.get('analysis')!r} analysis, and its "
                             f"factor of safety {restored.get('FS')} goes "
                             f"unstated")
            # The trial factor the field was solved at, where the file records
            # one — and this one does.
            if restored:
                got_F = (restored.get("solution") or {}).get("F")
                if got_F is None:
                    fails.append("a sidecar recording its trial factor restored "
                                 "without it, and the panel titles lost the F "
                                 "the field was solved at")
        finally:
            mw.close()

    # A sidecar recording a factor of safety and NO trial factor of its own —
    # every benchmark meta the corpus builder writes. The restore invented one by
    # falling back to FS, and the panels titled it "F = 1.35" over a trial that
    # was never run: F is the last converged strength-reduction factor, FS the
    # bracket midpoint, and they are not the same number. Absent F is absent.
    with tempfile.TemporaryDirectory() as tmp:
        stem = _sidecar_copy(
            os.path.splitext(FEM_XLSX)[0], tmp,
            lambda meta: {k: v for k, v in meta.items() if k != "F"})
        mw = MainWindow()
        try:
            with contextlib.redirect_stdout(io.StringIO()):
                mw.doc.slope_data = load_slope_data(f"{stem}.xlsx")
                mw._restore_fem_sidecar(mw.doc.slope_data["mesh"], stem)
            restored = mw.doc.results.get("fem_solution") or {}
            invented = (restored.get("solution") or {}).get("F")
            if invented is not None:
                fails.append(f"a sidecar recording no trial factor restored one "
                             f"anyway: F = {invented!r}, which is its factor of "
                             f"safety {restored.get('FS')!r}")
            # And the panel title says nothing about a trial rather than naming
            # a made-up one.
            from xslope.plot_fem import _fs_title
            titled = _fs_title("Viscoplastic shear strain", invented)
            if "F" in titled.replace("Viscoplastic shear strain", ""):
                fails.append(f"the panel title of a run with no recorded trial "
                             f"factor still names one: {titled!r}")
        finally:
            mw.close()
    return fails


CHECKS = [
    ("the content tree and its section order", test_tree),
    ("the tables carry the model's numbers", test_tables_carry_the_model),
    ("every figure is rendered to a file", test_figures_rendered),
    ("the traceability stamp", test_traceability),
    ("the content toggles remove what they name", test_toggles),
    ("the method picker drives the detail only", test_method_picker),
    ("the summary compares only what is documented",
     test_fs_table_compares_only_what_was_documented),
    ("the summary says where its numbers came from",
     test_the_fs_summary_says_where_its_numbers_came_from),
    ("one full detail block per method", test_multi_method_detail),
    ("the summary reaches the document", test_fs_summary_reaches_the_document),
    ("each engine presents the loads it applies",
     test_each_engine_presents_the_loads_it_applies),
    ("critical is a word a search earns",
     test_critical_is_a_word_a_search_earns),
    ("the stability section opens on its own model", test_lem_inputs_figure),
    ("the search plot is read for the engineer",
     test_search_figure_is_read_for_the_engineer),
    ("the slice key stands before its table", test_slice_key_figure),
    ("the figures are counted for the caller", test_figure_progress_counts),
    ("the seepage section", test_seep_section),
    ("every figure is cropped to keep its labels",
     test_report_figures_carry_their_axis_labels_whole),
    ("the seepage panels are the seepage view's",
     test_seep_panels_mirror_the_seep_view),
    ("the mesh legend names its boundaries",
     test_seep_mesh_legend_names_its_boundaries),
    ("the mesh is counted out once", test_the_mesh_is_counted_out_once),
    ("the mesh sentence counts what the mesh holds",
     test_the_mesh_sentence_counts_what_the_mesh_holds),
    ("the head figure carries the boundary water levels",
     test_seep_head_figure_draws_the_boundary_water_levels),
    ("a field the solution cannot draw is not drawn",
     test_seep_panels_the_solution_cannot_draw),
    ("a march saved beside a model is read back",
     test_tseep_discovered_beside_the_model),
    ("the transient section states the basis it is", test_tseep_section),
    ("the mesh figure marks what it says it marks",
     test_tseep_mesh_figure_marks_what_it_says),
    ("the states a march is documented at", test_tseep_frame_selection),
    ("a march is drawn on one scale throughout",
     test_tseep_frames_share_one_scale),
    ("a transient state is contours, not a flow net",
     test_tseep_carries_no_flow_net),
    ("the march over time is drawn and named", test_tseep_history_figure),
    ("both bases are documented, and each says which",
     test_tseep_dual_basis),
    ("a march solved in Studio reaches the report",
     test_tseep_reaches_the_report_from_studio),
    ("the transient figures have their dialog rows", test_tseep_dialog_rows),
    ("a saved solution records what the solve was",
     test_seep_solution_file_records_the_solve),
    ("every shipped solution records its solve",
     test_shipped_seep_companions_record_their_solve),
    ("convergence is stated where it is recorded",
     test_seep_convergence_is_stated),
    ("a confined analysis is reported as one", test_seep_confined_section),
    ("one term for the lines a flow net draws", test_one_term_for_flow_lines),
    ("a solve whose boundaries are not on record",
     test_seep_boundaries_not_on_record),
    ("a figure sentence agrees with its subject",
     test_a_figure_sentence_agrees_with_its_subject),
    ("a solution whose record and boundaries disagree",
     test_seep_stale_sidecar_says_so),
    ("a solution that records no flow rate", test_seep_without_a_flowrate),
    ("two boundary condition sets, one scale", test_seep_dual_section),
    ("the strength reduction section", test_fem_section),
    ("the mesh legend says what it holds",
     test_fem_mesh_legend_names_what_it_holds),
    ("the finite element section states its pore pressure basis",
     test_the_fem_section_states_its_pore_pressure_basis),
    ("the strength reduction paragraph describes what runs",
     test_the_ssrm_paragraph_describes_what_runs),
    ("the in-situ stress assumption is stated",
     test_the_in_situ_stress_assumption_is_stated),
    ("a tensile cap reaches the materials table",
     test_a_tensile_cap_reaches_the_materials_table),
    ("a low-order mesh carries its caution",
     test_a_low_order_mesh_carries_its_caution),
    ("a catastrophe run is described as one",
     test_a_catastrophe_run_is_described_as_one),
    ("an unrecorded analysis is not called a single trial",
     test_an_unrecorded_analysis_is_not_called_a_single_trial),
    ("the run record survives the file", test_the_run_record_survives_the_file),
    ("the result figures carry no title",
     test_fem_result_figures_carry_no_title),
    ("which result panels draw a legend",
     test_which_result_panels_draw_a_legend),
    ("the field state toggles", test_the_field_state_toggles),
    ("each state is drawn at its own scale",
     test_each_state_is_drawn_at_its_own_scale),
    ("the member forces follow the state drawn",
     test_the_member_forces_follow_the_state_the_panels_are_drawn_at),
    ("the search figure draws the trials",
     test_the_search_figure_draws_the_trials),
    ("one name for the shear strain field",
     test_one_name_for_the_shear_strain_field),
    ("the panels mirror the finite element view",
     test_fem_panels_mirror_the_fem_view),
    ("the model figure names only the members it draws",
     test_the_model_figure_names_only_the_members_it_draws),
    ("the finite element prose reads as documentation",
     test_the_fem_prose_reads_as_documentation),
    ("a wide table is fitted by the lines it makes",
     test_a_wide_table_is_fitted_by_the_lines_it_makes),
    ("the tension crack is documented where it is read",
     test_the_tension_crack_is_documented_where_it_is_read),
    ("every table is introduced in the owner's register",
     test_every_table_is_introduced_in_the_owners_register),
    ("the boundary sentence is derived from what was held",
     test_the_boundary_sentence_is_derived_from_what_was_held),
    ("consecutive figures are cited as a range",
     test_consecutive_figures_are_cited_as_a_range),
    ("the hybrid criterion sentence is true of the runs that shipped",
     test_the_hybrid_criterion_sentence_is_true_of_the_runs_that_shipped),
    ("the solve facts are recorded, not assumed",
     test_fem_solve_facts_are_recorded_not_assumed),
    ("no trial factor is invented", test_no_trial_factor_is_invented),
    ("a solution carrying no member forces says so",
     test_a_solution_without_member_forces_says_so),
    ("the member overlay claims no force it was not given",
     test_a_member_overlay_claims_no_force_it_was_not_given),
    ("reinforcement and piles in the finite element section",
     test_fem_members_are_reported),
    ("a broken stretch of capacity excepts what it leaves out",
     test_a_broken_tie_stretch_is_excepted),
    ("every member detail figure is readable",
     test_member_detail_figures_are_readable),
    ("each engine's section follows its solution",
     test_engine_sections_follow_their_solutions),
    ("a solved model's companions are the solutions",
     test_sidecars_assemble_the_solutions),
    ("an unusable companion is reported, not raised",
     test_unusable_sidecars_are_reported_not_raised),
    ("a member companion that is not this model's is refused",
     test_member_companions_that_are_not_this_models_are_refused),
    ("each searched method draws its own search",
     test_each_searched_method_draws_its_own_search),
    ("a method the run did not solve is run for the report",
     test_a_method_the_run_did_not_solve_is_run_for_the_report),
    ("a report-run method is cut like the analysis",
     test_a_report_run_method_is_run_at_the_analysis_discretization),
    ("a method that is not run says why",
     test_a_method_that_is_not_run_says_why),
    ("the rapid drawdown block names the governing stage",
     test_rapid_drawdown_names_the_governing_stage),
    ("the water prose follows the model", test_water_prose_is_conditional),
    ("the inputs an engine reads are stated where it is documented",
     test_the_inputs_an_engine_reads_are_stated_where_it_is_documented),
    ("members stand with the engine that reads them",
     test_members_stand_with_the_engine_that_reads_them),
    ("the model checks stay out of the report",
     test_the_model_checks_stay_out_of_the_report),
    ("implausible elastic properties are flagged, never filled",
     test_implausible_elastic_properties_are_flagged),
    ("an empty title-page field prints no row", test_title_page_omits_empty_rows),
    ("the .docx and its structure", test_docx),
    ("the running head names the section",
     test_running_head_names_the_section),
    ("the tables are fitted to their content", test_table_geometry),
    ("a table's rows are one height", test_table_rows_are_one_height),
    ("section breaks take no room", test_section_breaks_take_no_room),
    ("every paragraph is on a page of the report",
     test_every_block_reaches_the_page),
    ("the contents page lists the report", test_contents_page),
    ("Word numbers the headings", test_heading_numbers),
    ("the report writes one file", test_report_writes_one_file),
    ("the shipped template is reproducible", test_docx_template),
    ("the slice-column registry", test_column_registry),
    ("every force is in every equation or says why not",
     test_force_term_registry),
    ("a converged solution gets its working",
     test_calculation_tolerance_follows_the_solver),
    ("a method block never goes quiet", test_a_method_block_never_goes_quiet),
    ("a refusal prints no number it cannot stand behind",
     test_a_refusal_prints_no_number_it_cannot_stand_behind),
    ("the factor of safety from the printed operands",
     test_calculation_reproduces_fs),
    ("the sums carry the digits to divide",
     test_calculation_sums_are_printed_precisely_enough),
    ("the equation follows the model", test_calculation_terms_follow_the_model),
    ("the per-slice terms are table columns", test_calculation_columns),
    ("the equilibrium residuals close", test_calculation_residuals),
    ("Spencer's force sums are printed and check out", test_spencer_force_sums),
    ("the mirror rule is the arithmetic behind Q", test_spencer_mirror_rule),
    ("no force the equations print is called absent",
     test_absent_features_are_really_absent),
    ("the equation numbers are in the prose",
     test_equation_numbers_are_in_the_prose),
    ("the published equation comes before this model's",
     test_the_published_equation_comes_first),
    ("the base normal is published then reduced",
     test_the_normal_force_is_published_then_reduced),
    ("Bishop prints no base normal", test_bishop_prints_no_base_normal),
    ("the moment quotient recomposes to its page's equation",
     test_the_moment_quotient_recomposes),
    ("the evaluated equation introduces its terms",
     test_the_evaluated_equation_introduces_its_terms),
    ("the general moment arms name no equation",
     test_the_general_arms_name_no_equation),
    ("a line load's moment is its page's", test_a_line_loads_moment_is_the_pages),
    ("every printed full form matches its page",
     test_the_full_forms_match_their_pages),
    ("the whole-mass balance is published then reduced",
     test_the_whole_mass_balance_is_published_then_reduced),
    ("the shared slice count", test_the_shared_slice_count),
    ("the maximum depth is an input at any elevation",
     test_the_maximum_depth_is_an_input_at_any_elevation),
    ("each method prints its own page's equations",
     test_the_method_prints_its_own_pages_equations),
    ("a symbol in a sentence is set as a symbol",
     test_prose_symbols_are_set_as_symbols),
    ("a symbol in a table is set as a symbol",
     test_table_symbols_are_set_as_symbols),
    ("the base shear's arm is named for the surface",
     test_the_base_shear_arm_is_named_for_the_surface),
    ("each calculation's lead reads once", test_calculation_leads_read_once),
    ("a wide quotient keeps its own equals sign",
     test_the_wide_quotient_is_narrowed),
    ("a subscript is not cut short", test_scripts_are_not_cut_short),
    ("every printed symbol is defined where it is printed",
     test_printed_symbols_resolve),
    ("the prose is about the analysis", test_prose_is_about_the_analysis),
    ("the equation is cited for what it is",
     test_the_equation_is_cited_for_what_it_is),
    ("the notation matches the documentation",
     test_calculation_notation_matches_the_docs),
    ("each method's block opens by saying what it satisfies",
     test_method_summary_opens_each_block),
    ("the documentation links resolve", test_docs_links),
    ("the calculations reach the document", test_calculation_in_the_document),
    ("each calculation opens on its force diagram",
     test_force_diagram_heads_the_calculations),
    ("every force diagram draws its slice at one size",
     test_force_diagram_prints_one_slice_at_one_size),
    ("every figure and table is cited", test_every_block_is_cited),
    ("a citation is a live cross-reference",
     test_citations_are_cross_references),
    ("the project definition sends the reader on",
     test_the_project_definition_sends_the_reader_on),
    ("the material zones are named in bold",
     test_the_material_zones_are_named_in_bold),
    ("the factor of safety is defined as a ratio",
     test_the_factor_of_safety_is_defined_as_a_ratio),
    ("the reinforcement direction is named as its column prints it",
     test_the_reinforcement_direction_is_named_as_the_column_prints_it),
    ("the member overlay marks the state the solver recorded",
     test_the_member_overlay_marks_the_state_the_solver_recorded),
    ("the deformed members are a different color from the original",
     test_the_deformed_members_are_a_different_color_from_the_original),
    ("the member terms are defined where they are used",
     test_the_member_terms_are_defined_where_they_are_used),
    ("a definition follows the thing it defines",
     test_a_definition_follows_the_thing_it_defines),
    ("the detail figures are explained",
     test_the_detail_figures_are_explained),
    ("the member subsections locate their members",
     test_the_member_subsections_locate_their_members),
    ("the shared-model plot", test_shared_plot),
    ("every profile line names its material",
     test_profile_lines_name_their_materials),
    ("the model figure's point-coordinate toggle",
     test_model_figure_coordinate_labels),
    ("the coordinate labels are placed clear",
     test_coordinate_labels_are_placed_clear),
    ("the dialog and its toggles", test_dialog),
    ("the dialog remembers the right things", test_dialog_settings),
    ("the finished report is opened", test_open_output),
    ("the report builds off the GUI thread", test_report_runs_off_the_gui_thread),
    ("the progress bar counts the figures", test_report_runner_progress),
    ("cancelling stops the build", test_report_runner_cancel),
    ("a build that cannot finish says so", test_report_runner_failure),
    ("a non-circular surface dims the moment methods",
     test_noncircular_dims_the_moment_methods),
    ("the slice-numbers display toggle", test_slice_numbers_display_option),
    ("the menu item and its gate", test_main_window_action),
    ("every engine's solution reaches the report",
     test_report_solutions_carry_every_engine),
]

#: Checks that need the Studio layer; skipped when PySide6 is absent.
_STUDIO_ONLY = {test_seep_panels_mirror_the_seep_view,
                test_fem_panels_mirror_the_fem_view,
                test_which_result_panels_draw_a_legend,
                test_the_run_record_survives_the_file,
                test_dialog, test_dialog_settings, test_open_output,
                test_report_runs_off_the_gui_thread, test_report_runner_progress,
                test_report_runner_cancel, test_report_runner_failure,
                test_noncircular_dims_the_moment_methods,
                test_slice_numbers_display_option, test_main_window_action,
                test_report_solutions_carry_every_engine}


def run():
    """Every check; returns a list of failure strings (empty = pass)."""
    checks = CHECKS
    try:
        import PySide6                                    # noqa: F401
    except Exception:
        print("Report: PySide6 not installed — Studio checks skipped.")
        checks = [c for c in CHECKS if c[1] not in _STUDIO_ONLY]

    failures = []
    for name, fn in checks:
        try:
            fs = fn()
        except Exception as exc:
            import traceback
            traceback.print_exc()
            fs = [f"{name} raised: {exc!r}"]
        print(f"  {name:48s} {'ok' if not fs else f'FAIL ({len(fs)})'}")
        failures += fs
    return failures


def main():
    print("Analysis Report:")
    failures = run()
    if failures:
        print("\nFAILURES:")
        for f in failures:
            print(f"  - {f}")
        raise SystemExit(1)
    print("\nAll report checks passed.")


if __name__ == "__main__":
    import matplotlib
    matplotlib.use("Agg")
    try:
        from PySide6.QtWidgets import QApplication
        _app_ = QApplication.instance() or QApplication([])
    except Exception:
        pass
    main()
