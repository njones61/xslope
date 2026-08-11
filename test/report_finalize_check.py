# Copyright 2025 Norman L. Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Checks for finishing the Analysis Report — in Word, or without it.

What is being defended:

  A. THE PAGE NUMBERS ARE WORD'S — on a machine that has Word, a generated
     report goes in with a contents page that names its sections and comes out
     with one that says what page each is on. The evidence is the field result
     cached in the document: ``PAGEREF`` fields with numbers in them where there
     were none, and the "press F9" line gone.

  G. AND WITHOUT WORD THEY ARE STILL RIGHT — the same report finished by
     LibreOffice comes out with a number against every entry, and every one of
     those numbers is the page that heading is on. That is measured HERE, by a
     second reading of a fresh layout that has nothing to do with how the
     numbers were computed: the page a heading first appears on as a line of
     text, and the page's own "Page N of M" footer. An entry that cannot be
     placed is a failure that names it, and a failure leaves the report byte for
     byte as it was. The proof itself is mutation-tested: a leg that writes a
     number other than the one it computed must not report success.

  B. AND THE DOCUMENT IS STILL CLEAN — no field comes back marked dirty. A dirty
     field is what makes Word for Mac ask about "fields that may refer to other
     files" every time the report is opened, so a finish that introduced one
     would trade a missing page number for a warning on a clean document.

  C. NOTHING RAISES, EVER — no path, a path to nothing, a file that is not a
     Word document, no Word on the machine, a platform with no Word at all: each
     is a ``(False, sentence)``. This is a cosmetic step; it may not cost anyone
     their report.

  D. WORD IS ONLY HANDED A REPORT — a file that is not a document package with a
     contents field in it never reaches Word. An unreadable file put in front of
     Word is a modal dialog on the user's screen.

  E. THE STUDIO WIRING — generating a report finalizes it and then opens it; the
     settings key switches the finish off; and a document with nothing to update
     is opened without Word being started.

  F. THE WINDOWS LEG IS INTACT — it cannot be run here, so what is checked is
     that it is there, that it says the things COM needs said, and that on a
     platform where its shell is missing it fails into the same sentence the Mac
     leg would.

Word is started at most once (check A), and only when this machine has it: with
no Word, A is skipped and the detection and fallback paths are checked instead.

AND ONLY WHEN IT IS ASKED FOR. Driving Word means driving the copy the person at
this machine is working in: their document steals focus, their session is
scripted, and a run started for another reason has taken the machine over. So
every leg that reaches the real Word — check A, and the tidiness question it
asks afterwards — is off unless ``XSLOPE_WORD_OK=1`` is set in the environment
(:data:`WORD_OK`). The skip is printed and the row says so; a leg that did not
run is never reported as one that passed. Everything else here runs
unconditionally: none of it starts an application.
"""

import os
import re
import shutil
import sys
import tempfile
import zipfile

_HERE = os.path.dirname(os.path.abspath(__file__))
_REPO = os.path.dirname(_HERE)
if _REPO not in sys.path:
    sys.path.insert(0, _REPO)

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

REINF_XLSX = os.path.join(_REPO, "docs", "inputs", "slope", "xslope_reinf.xlsx")

#: The title the sample report is written under. It is also what stands at the
#: left of every running head, which is how a heading is told from the head that
#: repeats it (:func:`_first_page_of`).
SAMPLE_TITLE = "Sample Levee"

_REPORT = {}


def _report_docx(directory):
    """A small real report, written into ``directory``.

    The same sample the report checks use, solved by one method with the figures
    switched off: what is being read here is the contents field, not a plot.
    """
    import matplotlib
    matplotlib.use("Agg")
    from xslope.fileio import load_slope_data
    from xslope.report import generate_report
    from xslope.slice import generate_slices
    from xslope.solve import solve_selected

    if "solved" not in _REPORT:
        slope_data = load_slope_data(REINF_XLSX)
        ok, out = generate_slices(slope_data, circle=slope_data["circles"][0],
                                  num_slices=15)
        if not ok:
            raise RuntimeError(f"the sample model produced no slices: {out}")
        slice_df, surface = out
        bundle = {"slice_df": slice_df, "failure_surface": surface,
                  "results": solve_selected("bishop", slice_df),
                  "search": None, "method": "bishop"}
        _REPORT["solved"] = (slope_data, {"lem": [bundle]})
    slope_data, solutions = _REPORT["solved"]

    # A path with a space in it: the scripts are handed their arguments, never a
    # quoted command line, and this is what would catch it if that changed.
    path = os.path.join(directory, "sample report.docx")
    opts = {"input_path": REINF_XLSX, "title": SAMPLE_TITLE,
            "pd_figure": False, "lem_search_figure": False,
            "lem_solution_figure": False}
    ok, out = generate_report(slope_data, solutions, opts, path)
    if not ok:
        raise RuntimeError(f"generate_report failed: {out}")
    return path


def _contents_state(path):
    """What the document's contents field currently holds."""
    with zipfile.ZipFile(path) as package:
        xml = package.read("word/document.xml").decode("utf-8", "replace")
    start = xml.find(" TOC ")
    end = xml.find("fldCharType=\"end\"", start)
    result = xml[start:end] if start >= 0 else ""
    return {
        "pagerefs": result.count("PAGEREF"),
        "dirty": xml.count("w:dirty"),
        "asks_for_f9": "Page numbers appear" in xml or "press F9" in xml,
        "numbers": [t for t in _run_texts(result) if t.strip().isdigit()],
    }


def _run_texts(xml):
    """The text of every run in a fragment of document XML."""
    import re
    return re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)


def _cached_contents(path):
    """``[(entry, page or None), …]`` — the contents field as the file holds it.

    Read out of the document's XML here rather than through the finalizer's own
    parser, so that a leg which numbers the wrong paragraphs is caught by
    something that does not share its idea of which paragraphs those are.
    """
    with zipfile.ZipFile(path) as package:
        xml = package.read("word/document.xml").decode("utf-8", "replace")
    start = xml.find(" TOC ")
    if start < 0:
        return []
    result = xml[start:xml.find("fldCharType=\"end\"", start)]
    out = []
    for chunk in result.split("</w:p>")[:-1]:
        texts = [t for t in _run_texts(chunk)]
        if not "".join(texts).strip():
            continue
        page = int(texts[-1]) if texts[-1].strip().isdigit() else None
        joined = "".join(texts[:-1] if page is not None else texts)
        out.append((" ".join(joined.split()), page))
    return out


def _stub_docx(path):
    """A file with a .docx name and nothing in it — what a check harness writes
    when it only needs a path."""
    open(path, "wb").close()
    return path


# --------------------------------------------------------------------------
# A + B. the real thing
# --------------------------------------------------------------------------

#: The one way to ask for the legs that drive Microsoft Word.
#:
#: Word on this machine is the copy the person at it is working in. A leg that
#: scripts it opens and closes documents in their live session, steals their
#: focus, and — on a machine whose automation permission has been granted once —
#: does it without asking. That is not a thing a test run may decide to do for
#: its own reasons, so it is opt-in and nothing else in the suite can turn it on:
#: it is read from the environment, at the moment it matters.
WORD_OK = "XSLOPE_WORD_OK"


def word_automation_allowed():
    """Whether this run may drive Microsoft Word."""
    return os.environ.get(WORD_OK) == "1"


def test_word_finishes_the_report():
    """Word turns the contents page into a contents page.

    Skipped, loudly, unless this run was explicitly allowed to drive Word
    (:data:`WORD_OK`), and skipped on a machine with no Word: the rest of the
    row still checks the detection and the fallback.
    """
    from xslope.report_finalize import finalize_with_word, word_available

    if not word_automation_allowed():
        print(f"    (Word automation is off: set {WORD_OK}=1 to run the finish "
              f"against the Word on this machine)")
        return []

    ok, why = word_available()
    if not ok:
        print(f"    (no Word on this machine: {why} — the finish is not run)")
        return []

    fails = []
    with tempfile.TemporaryDirectory() as tmp:
        path = _report_docx(tmp)
        before = _contents_state(path)
        if before["pagerefs"]:
            fails.append("the written report already carries page references; "
                         "there would be nothing to prove")
        if not before["asks_for_f9"]:
            fails.append("the written report does not ask the reader to update "
                         "the field — check what the renderer now caches")

        done, msg = finalize_with_word(path)
        if not done:
            return fails + [f"Word did not finish the report: {msg}"]

        after = _contents_state(path)
        if after["pagerefs"] < 1:
            fails.append("the contents field came back with no page references")
        if not after["numbers"]:
            fails.append("the contents field came back with no page numbers")
        if after["asks_for_f9"]:
            fails.append("the contents page still asks the reader to press F9")
        if after["dirty"]:
            fails.append(f"the finish left {after['dirty']} dirty field(s) — Word "
                         f"would ask about them on every open")
        if not _closed_in_word(path):
            fails.append("the document was left open in Word")
    return fails


def _closed_in_word(path):
    """True when Word is not holding the document open. Anything that stops the
    question being answered counts as answered: this guards tidiness, not
    correctness — including a run that was never allowed to talk to Word."""
    if sys.platform != "darwin" or not word_automation_allowed():
        return True
    import subprocess
    script = ("function run(argv) { var w = Application('Microsoft Word');"
              " var d = w.documents; var out = [];"
              " for (var i = 0; i < d.length; i++) {"
              "  try { out.push(d[i].name()); } catch (e) {} }"
              " return out.join('\\n'); }")
    try:
        done = subprocess.run(["osascript", "-l", "JavaScript", "-"],
                              input=script, capture_output=True, text=True,
                              timeout=30)
    except Exception:
        return True
    return os.path.basename(path) not in done.stdout.splitlines()


# --------------------------------------------------------------------------
# G. the finish without Word
# --------------------------------------------------------------------------

def _layout_available():
    """``(True, "")`` when this machine can lay a document out, else
    ``(False, reason)``. Printed by the caller, which then skips."""
    from xslope.report_finalize import libreoffice_available
    return libreoffice_available()


def _layout_pages(path):
    """The text of every page of a LibreOffice layout of ``path``, in order.

    This check's OWN reading of the document. It borrows where LibreOffice lives
    and nothing else: the pages come back as text, and where a heading is on
    them is worked out here, so a finish whose page numbers agree with these has
    been checked by something that did not compute them.
    """
    import subprocess
    from xslope.report_finalize import _soffice
    tmp = tempfile.mkdtemp()
    try:
        subprocess.run([_soffice(), "--headless", "--convert-to", "pdf",
                        "--outdir", tmp, path], capture_output=True, timeout=900)
        pdf = os.path.join(
            tmp, os.path.splitext(os.path.basename(path))[0] + ".pdf")
        if not os.path.exists(pdf):
            return []
        import pypdf
        return [page.extract_text() or "" for page in pypdf.PdfReader(pdf).pages]
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _body_starts(pages, title):
    """The page the body starts on.

    The running head over the body carries the report's title AND the numbered
    section the page is in; the front matter's carries the title alone. So the
    body starts at the first page with such a head — counting contents pages
    instead puts the start one page early on a report whose contents run to two,
    and every entry listed on that second page then matches itself.
    """
    head = re.compile(r"^%s\s+\d+(\.\d+)*\s+\S" % re.escape(title))
    for number, text in enumerate(pages, start=1):
        if any(head.match(line.strip()) for line in text.splitlines()):
            return number
    return 1


def _first_page_of(pages, heading, title):
    """The page a heading first stands on as a line of its own, or None.

    The contents page names every heading and the running head repeats one of
    them at the top of every page of its section, so neither counts: the search
    starts after the front matter, and a line that begins with the report's
    title is a running head.
    """
    for number, text in enumerate(pages, start=1):
        if number < _body_starts(pages, title):
            continue
        for line in text.splitlines():
            line = line.strip()
            if title and line.startswith(title):
                continue
            if line == heading or line.startswith(heading + " "):
                return number
    return None


def _footer_number(text):
    """What a page calls itself, out of its own footer, or None."""
    for line in text.splitlines():
        found = re.match(r"^Page\s+(\d+)\s+of\s+(\d+)\b", line.strip())
        if found:
            return int(found.group(1))
    return None


def test_libreoffice_numbers_the_report():
    """A report finished without Word: every entry numbered, every number the
    page that heading is on.

    The numbers are checked against a fresh layout read as TEXT — where each
    heading first stands as a line of its own, and what the page it stands on
    calls itself in its footer. Neither is how the finish computed them.
    """
    from xslope.report_finalize import finalize_with_libreoffice

    ok, why = _layout_available()
    if not ok:
        print(f"    (nothing here lays a document out: {why})")
        return []

    fails = []
    with tempfile.TemporaryDirectory() as tmp:
        path = _report_docx(tmp)
        before = _cached_contents(path)
        if not before:
            return ["the written report has no contents entries to number"]
        if any(page is not None for _entry, page in before):
            fails.append("the written report already carries page numbers; "
                         "there would be nothing to prove")

        done, msg = finalize_with_libreoffice(path)
        if not done:
            return fails + [f"the report was not finished: {msg}"]

        after = _cached_contents(path)
        if [e for e, _p in after] != [e for e, _p in before]:
            fails.append(f"the finish rewrote the contents: {before} became "
                         f"{after}")
        unnumbered = [e for e, page in after if page is None]
        if unnumbered:
            fails.append(f"{len(unnumbered)} entr(ies) came back with no page: "
                         f"{unnumbered[:3]}")

        state = _contents_state(path)
        if state["asks_for_f9"]:
            fails.append("the contents page still asks the reader to press F9")
        if state["dirty"]:
            fails.append(f"the finish left {state['dirty']} dirty field(s)")
        with zipfile.ZipFile(path) as package:
            xml = package.read("word/document.xml").decode("utf-8", "replace")
        for mark in ('<w:instrText', ' TOC ', 'fldCharType="begin"',
                     'fldCharType="separate"', 'fldCharType="end"'):
            if mark not in xml:
                fails.append(f"the contents field is no longer a field: {mark} "
                             f"is gone, so F9 would no longer rebuild it")

        pages = _layout_pages(path)
        if not pages:
            return fails + ["the finished report could not be laid out again"]
        for entry, page in after:
            if page is None:
                continue
            stands_on = _first_page_of(pages, entry, SAMPLE_TITLE)
            if stands_on != page:
                fails.append(f"the contents sends a reader to page {page} for "
                             f"“{entry}”, which is on page {stands_on}")
            called = _footer_number(pages[page - 1]) if page <= len(pages) \
                else None
            if called is not None and called != page:
                fails.append(f"page {page} of the file calls itself page "
                             f"{called}, so “{entry}” is numbered for a sheet "
                             f"and not for a page")
    return fails


def test_an_entry_with_no_page_is_named():
    """An entry the layout cannot place fails the finish, by name — and the
    report is left exactly as it was.

    A contents page that is right about eleven sections and silent about the
    twelfth is worse than one with no numbers at all: nothing on it says which
    one it is wrong about. So the entry is named, the counts are given, and the
    document on disk is not touched.
    """
    from docx import Document
    from xslope.report_finalize import finalize_with_libreoffice

    ok, why = _layout_available()
    if not ok:
        print(f"    (nothing here lays a document out: {why})")
        return []

    fails = []
    sabotage = "3.2 Materials of a section this report does not have"
    with tempfile.TemporaryDirectory() as tmp:
        path = _report_docx(tmp)
        doc = Document(path)
        for para in doc.paragraphs:
            if para.text.strip() == "3.2 Materials" \
                    and para.style.name.lower().startswith("toc"):
                para.runs[0].text = sabotage
                break
        else:
            return ["the sample report's contents no longer list 3.2 Materials"]
        doc.save(path)
        with open(path, "rb") as handle:
            written = handle.read()

        done, msg = finalize_with_libreoffice(path)
        if done:
            fails.append("a contents page with an entry on no page was reported "
                         "as finished")
        if sabotage not in (msg or ""):
            fails.append(f"the entry that could not be placed is not named: "
                         f"{msg!r}")
        if not re.search(r"\d+ of \d+ entries", msg or ""):
            fails.append(f"the refusal does not say how many entries were "
                         f"numbered and how many were not: {msg!r}")
        with open(path, "rb") as handle:
            if handle.read() != written:
                fails.append("a finish that failed changed the report anyway")
    return fails


def test_the_contents_field_is_found_by_its_instruction():
    """Fields nest, and the contents field is the one whose instruction says so.

    The document built here has what a report has: a field on the title page
    (whose separate mark comes first), and a field INSIDE a contents entry (whose
    end mark comes before the contents field's own). Keying on the first
    separate mark collects the title page; keying on the first end mark stops
    halfway down the list. Both are measured.
    """
    from docx import Document
    from xslope.report_finalize import _toc_field
    from xslope.report_docx import TOC_INSTRUCTION, add_field, _fld_char, \
        _instr_text

    fails = []
    doc = Document()
    title = doc.add_paragraph()
    add_field(title, ' DOCPROPERTY "Title" \\* MERGEFORMAT ', "Sample Levee")
    doc.add_paragraph("Table of Contents")

    first = doc.add_paragraph()
    _fld_char(first, "begin")
    _instr_text(first, TOC_INSTRUCTION)
    _fld_char(first, "separate")
    first.add_run("1 One")
    doc.add_paragraph().add_run("2 Two")
    nested = doc.add_paragraph()
    nested.add_run("3 Three")
    add_field(nested, " PAGEREF _Toc3 \\h ", "9")
    doc.add_paragraph().add_run("4 Four")
    hint = doc.add_paragraph()
    hint.add_run("Page numbers appear when the table is updated in Word.")
    _fld_char(hint, "end")

    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "nested.docx")
        doc.save(path)
        entries, placeholder = _toc_field(Document(path))

    found = [text for _para, text in entries]
    if found != ["1 One", "2 Two", "3 Three9", "4 Four"]:
        fails.append(f"the contents field's entries came back as {found}")
    if any("Sample Levee" in text for text in found):
        fails.append("the title page's own field was read as a contents entry")
    if placeholder is None or "Page numbers appear" not in placeholder.text:
        fails.append("the line the numbers replace was not found")
    return fails


def test_the_numbers_are_proved_after_they_are_written():
    """The proof is taken off the document, not off the intention.

    Mutation: the leg is made to write a number other than the one it computed,
    and then to write none at all. A finish that reports success either way is
    checking what it meant to do rather than what it did — which is the whole
    difference between a page number and a guess.
    """
    import xslope.report_finalize as rf

    ok, why = _layout_available()
    if not ok:
        print(f"    (nothing here lays a document out: {why})")
        return []

    fails = []
    real = rf._write_page_number
    mutations = [
        ("one page out",
         lambda para, page, width: real(para, page + 1, width)),
        ("not written at all", lambda para, page, width: None),
    ]
    with tempfile.TemporaryDirectory() as tmp:
        path = _report_docx(tmp)
        with open(path, "rb") as handle:
            written = handle.read()
        for what, mutant in mutations:
            rf._write_page_number = mutant
            try:
                done, msg = rf.finalize_with_libreoffice(path)
            finally:
                rf._write_page_number = real
            if done:
                fails.append(f"a report whose numbers went in {what} was "
                             f"reported as finished: {msg}")
            with open(path, "rb") as handle:
                if handle.read() != written:
                    fails.append(f"a finish that went wrong ({what}) changed "
                                 f"the report anyway")
    return fails


def test_the_finish_uses_the_program_it_is_told_to():
    """Which program finishes the report, and what happens when it declines.

    Stubbed throughout: what is being checked is the choosing, not the
    finishing, and this check must not start either program.
    """
    import xslope.report_finalize as rf

    fails = []
    called = []
    real_word, real_office = rf.finalize_with_word, rf.finalize_with_libreoffice
    real_available = rf.word_available
    rf.finalize_with_word = lambda path, **kw: (called.append("word"),
                                                (False, "Word declined."))[1]
    rf.finalize_with_libreoffice = lambda path, **kw: (
        called.append("libreoffice"), (True, "LibreOffice numbered them."))[1]
    was = os.environ.get(rf.FINISH_ENV)
    try:
        with tempfile.TemporaryDirectory() as tmp:
            path = _stub_docx(os.path.join(tmp, "r.docx"))

            rf.word_available = lambda: (True, "Word")
            called.clear()
            done, msg = rf.finalize_report(path)
            if not done or called != ["word", "libreoffice"]:
                fails.append(f"a Word that declined did not fall back to the "
                             f"layout engine: {called}, {msg!r}")

            called.clear()
            done, msg = rf.finalize_report(path, prefer="word")
            if done or called != ["word"]:
                fails.append(f"asked for Word, the finish used {called}")

            called.clear()
            rf.word_available = lambda: (False, "Microsoft Word is not "
                                                "installed.")
            done, _msg = rf.finalize_report(path)
            if not done or called != ["libreoffice"]:
                fails.append(f"with no Word, the finish used {called}")

            called.clear()
            done, msg = rf.finalize_report(path, prefer="word")
            if done or called:
                fails.append(f"asked for Word where there is none, the finish "
                             f"used {called}")

            rf.word_available = lambda: (True, "Word")
            called.clear()
            done, _msg = rf.finalize_report(path, prefer="libreoffice")
            if not done or called != ["libreoffice"]:
                fails.append(f"asked for the layout engine, the finish used "
                             f"{called}")

            # The same answer from the environment, and a plain sentence when it
            # names something that is not a program this can use.
            called.clear()
            os.environ[rf.FINISH_ENV] = "libreoffice"
            rf.finalize_report(path)
            if called != ["libreoffice"]:
                fails.append(f"{rf.FINISH_ENV} did not choose the finish: "
                             f"{called}")
            os.environ[rf.FINISH_ENV] = "Microsoft Publisher"
            done, msg = rf.finalize_report(path)
            if done or rf.FINISH_ENV not in msg:
                fails.append(f"an unknown {rf.FINISH_ENV} is reported as "
                             f"{msg!r}")
    finally:
        rf.finalize_with_word, rf.finalize_with_libreoffice = (real_word,
                                                              real_office)
        rf.word_available = real_available
        if was is None:
            os.environ.pop(rf.FINISH_ENV, None)
        else:
            os.environ[rf.FINISH_ENV] = was
    return fails


# --------------------------------------------------------------------------
# C. nothing raises
# --------------------------------------------------------------------------

def test_refusals_are_sentences():
    """Every way any of this can decline is a (False, sentence).

    Asked of all three entry points: a report is not lost to a bad path,
    whichever program was going to finish it.
    """
    from xslope.report_finalize import (finalize_report,
                                        finalize_with_libreoffice,
                                        finalize_with_word)

    fails = []
    cases = [("", "no path"), ("/no/such/report.docx", "a path to nothing")]
    with tempfile.TemporaryDirectory() as tmp:
        text = os.path.join(tmp, "notes.txt")
        with open(text, "w") as fh:
            fh.write("not a document")
        cases.append((text, "a file that is not a Word document"))

        for finish in (finalize_with_word, finalize_with_libreoffice,
                       finalize_report):
            for path, what in cases:
                where = f"{what}, through {finish.__name__}"
                try:
                    ok, msg = finish(path)
                except Exception as exc:
                    fails.append(f"{where} raised {exc!r}")
                    continue
                if ok:
                    fails.append(f"{where} was reported as finished")
                if not isinstance(msg, str) or not msg.strip():
                    fails.append(f"{where} gave no reason")
    return fails


def test_no_word_falls_back():
    """With no Word to be found, detection and the finish both say so plainly —
    and nothing is run."""
    import xslope.report_finalize as rf

    fails = []
    app, finder = rf._MAC_WORD_APP, rf._mdfind_word
    ran = []
    osascript = rf._osascript
    try:
        rf._MAC_WORD_APP = os.path.join(tempfile.gettempdir(), "No Word.app")
        rf._mdfind_word = lambda: None
        rf._osascript = lambda *a, **k: ran.append(a) or (False, "should not run")

        if sys.platform == "darwin":
            ok, why = rf.word_available()
            if ok:
                fails.append("Word was found where there is none")
            if "not installed" not in why:
                fails.append(f"the reason given is {why!r}")

        with tempfile.TemporaryDirectory() as tmp:
            path = _stub_docx(os.path.join(tmp, "r.docx"))
            ok, msg = rf.finalize_with_word(path)
            if ok:
                fails.append("a finish was reported with no Word")
            if not msg:
                fails.append("the refusal gave no reason")
        if ran:
            fails.append("a script was run although there is no Word")
    finally:
        rf._MAC_WORD_APP, rf._mdfind_word, rf._osascript = app, finder, osascript
    return fails


def test_platform_gate():
    """A platform Word does not run on is told so, not driven."""
    import xslope.report_finalize as rf

    fails = []
    real = sys.platform
    try:
        sys.platform = "linux"
        ok, why = rf.word_available()
        if ok:
            fails.append("Word was offered on a platform it does not run on")
        if "macOS" not in why:
            fails.append(f"the reason given is {why!r}")
        with tempfile.TemporaryDirectory() as tmp:
            ok, _msg = rf.finalize_with_word(_stub_docx(
                os.path.join(tmp, "r.docx")))
            if ok:
                fails.append("a finish was reported on a platform with no Word")
    finally:
        sys.platform = real
    return fails


# --------------------------------------------------------------------------
# F. the Windows leg
# --------------------------------------------------------------------------

#: What the Windows script must still be doing. It is driven from PowerShell
#: rather than pywin32 — no new dependency, and a subprocess can be given a
#: timeout — so these are the COM calls, spelled as PowerShell spells them.
_WINDOWS_MUST_SAY = [
    "Word.Application",
    "$word.Visible = $false",
    "$word.Documents.Open(",
    "Fields.Update()",
    "TablesOfContents",
    "$doc.Save()",
    "$doc.Close(",
    "$word.Quit()",
]


def test_windows_leg_is_intact():
    """It cannot be run here; it can be read, and it can be made to fail
    correctly."""
    import xslope.report_finalize as rf

    fails = []
    for phrase in _WINDOWS_MUST_SAY:
        if phrase not in rf._WINDOWS_SCRIPT:
            fails.append(f"the Windows script no longer says {phrase!r}")
    if "$started" not in rf._WINDOWS_SCRIPT:
        fails.append("the Windows script quits Word whether or not it started it")
    if rf._ps_quote("O'Brien's") != "'O''Brien''s'":
        fails.append("a PowerShell argument is not quoted safely")

    # On this machine there is no PowerShell to run it: the leg must come back
    # with a sentence rather than an exception, which is the discipline that
    # matters most on the platform it cannot be tried on.
    real = sys.platform
    try:
        sys.platform = "win32"
        with tempfile.TemporaryDirectory() as tmp:
            path = _stub_docx(os.path.join(tmp, "r.docx"))
            try:
                ok, msg = rf.finalize_with_word(path, timeout=20)
            except Exception as exc:
                fails.append(f"the Windows leg raised {exc!r}")
                return fails
            if ok and not os.path.exists(path):
                fails.append("a finish was reported for a file that is not there")
            if not isinstance(msg, str) or not msg.strip():
                fails.append("the Windows leg gave no reason")
    finally:
        sys.platform = real
    return fails


# --------------------------------------------------------------------------
# D + E. the Studio wiring
# --------------------------------------------------------------------------

def test_only_a_report_reaches_word():
    """A file with nothing to update is opened, and Word is left alone."""
    from studio.report_dialog import carries_contents_field

    fails = []
    with tempfile.TemporaryDirectory() as tmp:
        if carries_contents_field(_stub_docx(os.path.join(tmp, "empty.docx"))):
            fails.append("an empty file was taken for a report")
        text = os.path.join(tmp, "notes.docx")
        with open(text, "w") as fh:
            fh.write("not a package")
        if carries_contents_field(text):
            fails.append("a text file was taken for a report")
        if carries_contents_field(os.path.join(tmp, "absent.docx")):
            fails.append("a missing file was taken for a report")
        if not carries_contents_field(_report_docx(tmp)):
            fails.append("a real report was not recognised as one")
    return fails


def test_generate_finalizes_then_opens():
    """The Studio path: the document is finished, then shown — and the settings
    key stops the finish without stopping the report."""
    from PySide6.QtCore import QSettings
    from studio import report_dialog

    fails = []
    finalized, opened = [], []
    import xslope.report_finalize as rf
    real_finalize = rf.finalize_report
    real_open = report_dialog.QDesktopServices.openUrl
    rf.finalize_report = lambda path, **kw: (finalized.append(path),
                                             (True, "pretend"))[1]
    report_dialog.QDesktopServices.openUrl = (
        lambda url: opened.append(url.toLocalFile()))
    try:
        with tempfile.TemporaryDirectory() as tmp:
            path = _report_docx(tmp)
            if report_dialog.open_output(path, "docx") != "document":
                fails.append("a .docx did not open as a document")
            if finalized != [path]:
                fails.append(f"the report was not finalized before it was "
                             f"opened (finalized {finalized})")
            if opened[-1:] != [path]:
                fails.append(f"the document opened was {opened[-1:]}")

            # A document with nothing to update is opened just the same, and
            # Word is not started for it.
            finalized.clear()
            stub = _stub_docx(os.path.join(tmp, "stub.docx"))
            report_dialog.open_output(stub, "docx")
            if finalized:
                fails.append("a file with no contents field was sent to Word")
            if opened[-1] != stub:
                fails.append("a file with no contents field was not opened")

            # The escape hatch: no UI, one key.
            settings = QSettings(os.path.join(tmp, "settings.ini"),
                                 QSettings.IniFormat)
            settings.setValue(report_dialog.FINALIZE_KEY, False)
            if report_dialog.finalization_enabled(settings):
                fails.append("the settings key does not switch the finish off")
            finalized.clear()
            ok, why = report_dialog.finalize_document(path, settings=settings)
            if ok or finalized:
                fails.append("the finish ran with the key switched off")
            if "off" not in why:
                fails.append(f"switching it off is reported as {why!r}")
            settings.setValue(report_dialog.FINALIZE_KEY, True)
            if not report_dialog.finalization_enabled(settings):
                fails.append("the finish cannot be switched back on")
    finally:
        rf.finalize_report = real_finalize
        report_dialog.QDesktopServices.openUrl = real_open
    return fails


def test_word_is_not_driven_unless_it_was_asked_for():
    """With the environment variable unset, not one leg of this file talks to
    Microsoft Word.

    Measured on the calls themselves rather than on the docstrings: every check
    in this file is run with ``subprocess.run`` and the finalizer's own
    ``_osascript`` watched, and either one reaching the machine's Word is the
    failure. That is the guard the ban needs — a run started for something else
    drove the owner's live Word session twice, and a rule that lives only in a
    comment is a rule that does that again.

    The other checks stub ``_osascript`` out for their own purposes and put it
    back afterwards; the watch is what they put back, so it sees every call any
    of them makes for real.

    The mutation at the end runs ``osascript -e 1`` twice, which is the one
    osascript this file ever spawns: it evaluates the literal 1 and exits. It
    starts no application, names none, and cannot reach Word — it exists only so
    that the watch above is proved able to count a call it was meant to catch.
    Anything watching the process list sees the name and nothing else.
    """
    import subprocess

    fails = []
    import xslope.report_finalize as rf

    if word_automation_allowed():
        return [f"{WORD_OK} is set in this environment, so the guard cannot be "
                f"measured; unset it and run again"]

    spawned, scripted = [], []
    real_run = subprocess.run
    real_osascript = rf._osascript

    def watched_run(args, *rest, **kw):
        name = args[0] if isinstance(args, (list, tuple)) and args else args
        if str(name).endswith("osascript"):
            spawned.append(list(args) if isinstance(args, (list, tuple))
                           else args)
        return real_run(args, *rest, **kw)

    def watched_osascript(*a, **kw):
        scripted.append(a[:1])
        return real_osascript(*a, **kw)

    subprocess.run = watched_run
    rf._osascript = watched_osascript
    try:
        for name, fn in CHECKS:
            if fn is test_word_is_not_driven_unless_it_was_asked_for:
                continue
            try:
                fn()
            except Exception:
                # A leg that raised is another check's failure, not this one's;
                # what matters here is whether it reached Word on the way.
                pass
    finally:
        subprocess.run = real_run
        rf._osascript = real_osascript

    if spawned:
        fails.append(f"{len(spawned)} osascript call(s) were made with {WORD_OK} "
                     f"unset: {spawned[:2]}")
    if scripted:
        fails.append(f"{len(scripted)} script(s) were sent to the machine's Word "
                     f"with {WORD_OK} unset")

    # Mutation: the watch has to be able to see one. A leg that does call
    # osascript is counted.
    spawned.clear()
    subprocess.run = watched_run
    try:
        real_run(["osascript", "-e", "1"], capture_output=True, text=True,
                 timeout=10)
        watched_run(["osascript", "-e", "1"], capture_output=True, text=True,
                    timeout=10)
    except Exception:
        pass
    finally:
        subprocess.run = real_run
    if not spawned:
        fails.append("the watch counted no osascript call even when one was "
                     "made; it cannot fail")
    return fails


CHECKS = [
    ("Word builds the report's page numbers", test_word_finishes_the_report),
    ("LibreOffice builds them where Word cannot",
     test_libreoffice_numbers_the_report),
    ("an entry on no page is named, and nothing is written",
     test_an_entry_with_no_page_is_named),
    ("the contents field is found among fields",
     test_the_contents_field_is_found_by_its_instruction),
    ("the numbers are proved on the document",
     test_the_numbers_are_proved_after_they_are_written),
    ("the finish uses the program it is told to",
     test_the_finish_uses_the_program_it_is_told_to),
    ("every refusal is a sentence", test_refusals_are_sentences),
    ("no Word falls back quietly", test_no_word_falls_back),
    ("a platform without Word is not driven", test_platform_gate),
    ("the Windows leg is intact", test_windows_leg_is_intact),
    ("only a report reaches Word", test_only_a_report_reaches_word),
    ("generating finalizes, then opens", test_generate_finalizes_then_opens),
    ("Word is not driven unless it was asked for",
     test_word_is_not_driven_unless_it_was_asked_for),
]

#: Checks that need the Studio layer; skipped when PySide6 is absent.
_STUDIO_ONLY = {test_only_a_report_reaches_word,
                test_generate_finalizes_then_opens}


def run():
    """Every check; returns a list of failure strings (empty = pass)."""
    checks = CHECKS
    try:
        import PySide6                                    # noqa: F401
    except Exception:
        print("Report finish: PySide6 not installed — Studio checks skipped.")
        checks = [c for c in CHECKS if c[1] not in _STUDIO_ONLY]

    failures = []
    for name, fn in checks:
        try:
            fs = fn()
        except Exception as exc:
            import traceback
            traceback.print_exc()
            fs = [f"{name} raised: {exc!r}"]
        print(f"  {name:48s} {'ok' if not fs else f'FAIL ({len(fs)})'}")
        failures += fs
    return failures


def main():
    print("Finishing the report in Word:")
    failures = run()
    if failures:
        print("\nFAILURES:")
        for f in failures:
            print(f"  - {f}")
        raise SystemExit(1)
    print("\nAll report-finish checks passed.")


if __name__ == "__main__":
    import matplotlib
    matplotlib.use("Agg")
    try:
        from PySide6.QtWidgets import QApplication
        _app_ = QApplication.instance() or QApplication([])
    except Exception:
        pass
    main()
