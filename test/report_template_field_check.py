"""The report dialog's Template field, and what refuses a bad template.

A company template is how a report leaves the office on the firm's letterhead:
the page size and margins, the header and footer, the logo, and the Title,
Heading, Body Text and Caption styles the report is written in. Studio's report
dialog carries a Template row for it, and this pins both halves of that choice.

**The field.** It opens on the shipped template and says so in words. Browse
sets a path, the Shipped template button clears it, and either way ``options()``
carries ``template`` as the path or None — which is the key
:func:`xslope.report.generate_report` reads. The choice belongs to the person,
not the project, so it is remembered in QSettings; a remembered file that has
since been moved or deleted falls back to the shipped template and says which
file went, in the field's tooltip, rather than reporting on a different template
than last week without a word.

**The refusal.** A template is only a template of this report if it declares the
styles the report is written in. The renderer falls back to Normal for any style
it does not find, which for a Title or a Caption means a document with no
headings and no captions — a failure that is invisible in the result. So
:func:`xslope.report_docx.template_problem` states what is wrong with a chosen
template in one sentence naming the missing style, and three callers ask it: the
dialog on Generate, :func:`~xslope.report.generate_report` before it builds
anything, and :func:`~xslope.report_docx.render_docx` before it writes. None of
them raises a traceback at the user.

The refusal is mutation-tested from the template end: a copy of the shipped
template with one style deleted has to be named for that style, one style at a
time, for every style in ``REQUIRED_STYLES``. A check that only ever deleted
Caption would pass an implementation that hunted for Caption alone.

Qt-light: the dialogs are built offscreen, nothing is solved, and the only
documents written are style-stripped copies of the shipped template.

Run directly:  PYTHONPATH=. python3 test/report_template_field_check.py
"""

import os
import sys
import tempfile

_HERE = os.path.dirname(os.path.abspath(__file__))
_REPO = os.path.dirname(_HERE)
if _REPO not in sys.path:
    sys.path.insert(0, _REPO)

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

#: A model to open the dialog on — read only for its default title and output
#: path, so any shipped workbook does.
MODEL = os.path.join(_REPO, "docs", "inputs", "slope", "xslope_simple1.xlsx")

#: The copy of the shipped template the documentation hands out, so a reader can
#: edit it in Word and pick it here.
DOCS_TEMPLATE = os.path.join(_REPO, "docs", "studio", "files",
                             "report_template.docx")


def _app():
    from PySide6.QtWidgets import QApplication
    return QApplication.instance() or QApplication([])


def _dialog(settings=None):
    from studio.report_dialog import ReportDialog
    return ReportDialog(slope_data={}, solutions={}, model_path=MODEL,
                        settings=settings)


#: A one-pixel PNG, so a letterhead can carry a real image relationship without
#: a drawing library or a committed fixture.
ONE_PIXEL_PNG = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8"
    "z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg==")


def _letterhead_template(directory):
    """The shipped template with a letterhead in the header: a logo and a firm
    name linking to the firm's site, in a one-row table above the paragraph the
    report writes its running head into.

    Two relationship kinds in one table on purpose — an internal one to an image
    part and an external one to a URL — because a relationship id means nothing
    outside the part it was written in, and both have to be remade when the
    letterhead is copied into a section's own header.
    """
    import base64

    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    from xslope.report_docx import DEFAULT_TEMPLATE

    logo = os.path.join(directory, "logo.png")
    with open(logo, "wb") as fh:
        fh.write(base64.b64decode(ONE_PIXEL_PNG))

    doc = Document(DEFAULT_TEMPLATE)
    section = doc.sections[0]
    header = section.header
    width = section.page_width - section.left_margin - section.right_margin
    table = header.add_table(1, 2, width)
    left, right = table.rows[0].cells
    left.paragraphs[0].add_run().add_picture(logo, width=Inches(1.0))

    paragraph = right.paragraphs[0]
    run = paragraph.add_run("ACME Geotechnical")
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), header.part.relate_to("https://example.com",
                                               RT.HYPERLINK, is_external=True))
    paragraph._p.remove(run._r)
    link.append(run._r)
    paragraph._p.append(link)

    header.paragraphs[0]._p.addprevious(table._tbl)
    path = os.path.join(directory, "letterhead.docx")
    doc.save(path)
    return path


def _template_without(style_name, directory):
    """A copy of the shipped template with one style deleted."""
    from docx import Document

    from xslope.report_docx import DEFAULT_TEMPLATE

    doc = Document(DEFAULT_TEMPLATE)
    doc.styles[style_name].delete()
    path = os.path.join(directory,
                        f"no_{style_name.replace(' ', '_').lower()}.docx")
    doc.save(path)
    return path


# --------------------------------------------------------------------------
# The field
# --------------------------------------------------------------------------

def check_opens_on_the_shipped_template(failures):
    """No template chosen: the field says so, and the option is None."""
    _app()
    from studio.report_dialog import ReportDialog

    dlg = _dialog()
    try:
        if dlg.template.text() != ReportDialog.SHIPPED_TEMPLATE_TEXT:
            failures.append(f"the Template field opens reading "
                            f"{dlg.template.text()!r}, not "
                            f"{ReportDialog.SHIPPED_TEMPLATE_TEXT!r}")
        if not dlg.template.isReadOnly():
            failures.append("the Template field can be typed in; a template is "
                            "chosen by finding it")
        if dlg.template_path() is not None:
            failures.append(f"the dialog opens on {dlg.template_path()!r}")
        if dlg.options().get("template", "missing") is not None:
            failures.append(f"the options carry "
                            f"{dlg.options().get('template')!r} as the template")
    finally:
        dlg.close()


def check_browse_and_reset(failures):
    """Browse sets the path; Shipped template puts it back."""
    _app()
    from studio import report_dialog
    from studio.report_dialog import ReportDialog

    dlg = _dialog()
    real = report_dialog.QFileDialog.getOpenFileName
    picked = {}
    try:
        with tempfile.TemporaryDirectory() as tmp:
            chosen = os.path.join(tmp, "company.docx")
            from xslope.report_docx import DEFAULT_TEMPLATE
            import shutil
            shutil.copyfile(DEFAULT_TEMPLATE, chosen)

            def fake_open(parent, caption, directory, filt, *args, **kwargs):
                picked["filter"] = filt
                picked["caption"] = caption
                return chosen, ""

            report_dialog.QFileDialog.getOpenFileName = staticmethod(fake_open)
            dlg._browse_template()

            if picked.get("filter") != "Word templates (*.docx);;All files (*)":
                failures.append(f"the file dialog filters on "
                                f"{picked.get('filter')!r}")
            if dlg.template.text() != chosen:
                failures.append(f"after Browse the field reads "
                                f"{dlg.template.text()!r}")
            if dlg.template_path() != chosen:
                failures.append("Browse did not reach the template path")
            if dlg.options().get("template") != chosen:
                failures.append("Browse did not reach the options")

            # And a cancelled Browse leaves the choice alone.
            report_dialog.QFileDialog.getOpenFileName = staticmethod(
                lambda *a, **k: ("", ""))
            dlg._browse_template()
            if dlg.template_path() != chosen:
                failures.append("a cancelled Browse cleared the template")

            dlg.use_shipped_template()
            if dlg.template_path() is not None:
                failures.append(f"the Shipped template button left "
                                f"{dlg.template_path()!r}")
            if dlg.template.text() != ReportDialog.SHIPPED_TEMPLATE_TEXT:
                failures.append(f"after the reset the field reads "
                                f"{dlg.template.text()!r}")
            if dlg.options().get("template") is not None:
                failures.append("the reset did not reach the options")
    finally:
        report_dialog.QFileDialog.getOpenFileName = real
        dlg.close()


def check_remembered(failures):
    """The choice survives a second dialog, and a file that has gone does not."""
    _app()
    import shutil

    from PySide6.QtCore import QSettings

    from studio.report_dialog import SETTINGS_PREFIX, ReportDialog
    from xslope.report_docx import DEFAULT_TEMPLATE

    with tempfile.TemporaryDirectory() as tmp:
        ini = os.path.join(tmp, "settings.ini")
        settings = QSettings(ini, QSettings.IniFormat)
        chosen = os.path.join(tmp, "company.docx")
        shutil.copyfile(DEFAULT_TEMPLATE, chosen)

        first = _dialog(settings)
        first.set_template(chosen)
        first.remember()
        first.close()

        again = _dialog(settings)
        try:
            if again.template_path() != chosen:
                failures.append(f"the template was not remembered "
                                f"({again.template_path()!r})")
            if again.template.text() != chosen:
                failures.append("the remembered template is not shown")
        finally:
            again.close()

        if QSettings(ini, QSettings.IniFormat).value(
                SETTINGS_PREFIX + "template") != chosen:
            failures.append("the template is not stored under the report prefix")

        # The file goes away between sessions.
        os.remove(chosen)
        after = _dialog(settings)
        try:
            if after.template_path() is not None:
                failures.append(f"a remembered template that is no longer there "
                                f"came back as {after.template_path()!r}")
            if after.template.text() != ReportDialog.SHIPPED_TEMPLATE_TEXT:
                failures.append(f"the fallback field reads "
                                f"{after.template.text()!r}")
            tip = after.template.toolTip()
            if os.path.basename(chosen) not in tip:
                failures.append(f"the tooltip does not name the template that "
                                f"went missing: {tip!r}")
            if after.options().get("template") is not None:
                failures.append("the options carry a template that is not there")
        finally:
            after.close()

        # Clearing the choice is remembered too — a company template picked once
        # must not come back after the user goes back to the shipped one.
        back = _dialog(settings)
        back.use_shipped_template()
        back.remember()
        back.close()
        last = _dialog(settings)
        try:
            if last.template_path() is not None:
                failures.append("the shipped template was not remembered")
        finally:
            last.close()


def check_generate_refuses_a_bad_template(failures):
    """Generate names the missing style rather than starting a build."""
    _app()
    from PySide6.QtWidgets import QMessageBox

    dlg = _dialog()
    shown = []
    real = QMessageBox.warning
    QMessageBox.warning = staticmethod(
        lambda parent, title, text, *a, **k: shown.append((title, text)))
    try:
        with tempfile.TemporaryDirectory() as tmp:
            dlg.set_template(_template_without("Caption", tmp))
            accepted = []
            dlg.accepted.connect(lambda: accepted.append(True))
            dlg.accept()
            if accepted:
                failures.append("the dialog accepted a template the report "
                                "cannot be built on")
            if not shown:
                failures.append("a bad template produced no message")
            elif "Caption" not in shown[-1][1]:
                failures.append(f"the message does not name the missing style: "
                                f"{shown[-1][1]!r}")

            # A missing file, and a file that is not a .docx.
            shown.clear()
            dlg.set_template(os.path.join(tmp, "gone.docx"))
            dlg.accept()
            if not shown or "gone.docx" not in shown[-1][1]:
                failures.append(f"a template that is not there was not named: "
                                f"{shown}")
            shown.clear()
            not_word = os.path.join(tmp, "template.txt")
            open(not_word, "w").close()
            dlg.set_template(not_word)
            dlg.accept()
            if not shown or ".docx" not in shown[-1][1]:
                failures.append(f"a template that is not a .docx was not "
                                f"refused: {shown}")

            # And the shipped template goes through.
            shown.clear()
            dlg.use_shipped_template()
            dlg.path.setText(os.path.join(tmp, "report.docx"))
            dlg.accept()
            if shown:
                failures.append(f"the shipped template was refused: {shown}")
    finally:
        QMessageBox.warning = real
        dlg.close()


# --------------------------------------------------------------------------
# The refusal itself
# --------------------------------------------------------------------------

def check_every_required_style_is_named(failures):
    """Delete any one required style and the refusal names that style."""
    from xslope.report_docx import REQUIRED_STYLES, template_problem

    with tempfile.TemporaryDirectory() as tmp:
        for style_name in REQUIRED_STYLES:
            problem = template_problem(_template_without(style_name, tmp))
            if not problem:
                failures.append(f"a template with no {style_name!r} style was "
                                f"accepted")
            elif style_name not in problem:
                failures.append(f"a template with no {style_name!r} style was "
                                f"refused as {problem!r}")


def check_required_styles_are_the_renderer_s(failures):
    """The required list is the styles the renderer actually writes in, to the
    depth the report numbers headings to."""
    from xslope.report import HEADING_LEVELS
    from xslope.report_docx import (DEFAULT_TEMPLATE, REQUIRED_STYLES, STYLE,
                                    template_problem)

    headings = [s for s in REQUIRED_STYLES if s.startswith("Heading ")]
    if len(headings) != HEADING_LEVELS:
        failures.append(f"the report writes {HEADING_LEVELS} heading levels but "
                        f"requires {len(headings)} heading styles")
    for key in ("title", "body", "caption"):
        if STYLE[key] not in REQUIRED_STYLES:
            failures.append(f"the renderer writes in {STYLE[key]!r}, which no "
                            f"template is required to define")
    if template_problem(DEFAULT_TEMPLATE):
        failures.append(f"the shipped template is refused by its own rule: "
                        f"{template_problem(DEFAULT_TEMPLATE)}")
    if template_problem(None) or template_problem(""):
        failures.append("no template chosen was read as a bad template")
    if os.path.exists(DOCS_TEMPLATE) and template_problem(DOCS_TEMPLATE):
        failures.append(f"the template the documentation hands out is refused: "
                        f"{template_problem(DOCS_TEMPLATE)}")


def check_the_python_paths_refuse_too(failures):
    """``generate_report`` returns the sentence; ``render_docx`` raises it."""
    from xslope.report import generate_report
    from xslope.report_docx import TemplateError, render_docx

    with tempfile.TemporaryDirectory() as tmp:
        bad = _template_without("Title", tmp)
        out_path = os.path.join(tmp, "report.docx")
        ok, message = generate_report({}, {}, {"template": bad}, out_path)
        if ok:
            failures.append("generate_report built a report on a template "
                            "missing the Title style")
        elif "Title" not in str(message):
            failures.append(f"generate_report refused with {message!r}, which "
                            f"does not name the style")
        if os.path.exists(out_path):
            failures.append("a refused template still wrote a document")

        try:
            render_docx(None, out_path, template=bad)
        except TemplateError as exc:
            if "Title" not in str(exc):
                failures.append(f"render_docx raised {exc!r}, which does not "
                                f"name the style")
        except Exception as exc:                               # noqa: BLE001
            failures.append(f"render_docx raised {type(exc).__name__} rather "
                            f"than TemplateError: {exc}")
        else:
            failures.append("render_docx wrote a document on a template missing "
                            "the Title style")


def check_a_letterhead_carries_into_every_header(failures):
    """A letterhead in the template's header reaches the header of every section
    the report opens, with every relationship it carries remade there.

    Each section gets a header part of its own — the running head's tab stop is a
    property of the page it prints on — and a new part starts empty, so a
    letterhead left behind would print on the front matter and nowhere else. The
    ids are the point: an ``r:embed`` or an ``r:id`` copied verbatim names a
    relationship the destination part does not define, and Word calls a file
    carrying one unreadable.
    """
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn

    from xslope.report_docx import _carry_header_furniture

    with tempfile.TemporaryDirectory() as tmp:
        doc = Document(_letterhead_template(tmp))
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        _carry_header_furniture(doc, section)

        header = section.header
        tables = [el for el in header._element if el.tag == qn("w:tbl")]
        if not tables:
            failures.append("the template's letterhead table did not reach the "
                            "header of the section the report opened")
            return

        attributes = (qn("r:embed"), qn("r:link"), qn("r:id"))
        ids = [node.get(a) for table in tables for node in table.iter()
               for a in attributes if node.get(a)]
        if not ids:
            failures.append("the copied letterhead carries no relationship id, "
                            "so this check proves nothing about remapping")
            return

        rels = header.part.rels
        for rel_id in ids:
            if rel_id not in rels:
                failures.append(f"the copied letterhead still names {rel_id}, "
                                f"which the header part it was copied into does "
                                f"not define")
        kinds = {rels[i].reltype for i in ids if i in rels}
        if len(kinds) < 2:
            failures.append(f"only {len(kinds)} kind(s) of relationship were "
                            f"exercised; the image and the hyperlink are both "
                            f"meant to be carried")


CHECKS = [check_opens_on_the_shipped_template, check_browse_and_reset,
          check_remembered, check_generate_refuses_a_bad_template,
          check_every_required_style_is_named,
          check_required_styles_are_the_renderer_s,
          check_the_python_paths_refuse_too,
          check_a_letterhead_carries_into_every_header]


def run():
    """Returns a list of failure descriptions (empty on success)."""
    failures = []
    for chk in CHECKS:
        try:
            chk(failures)
        except Exception as e:                                 # noqa: BLE001
            failures.append(f"{chk.__name__} raised {type(e).__name__}: {e}")
    return failures


if __name__ == "__main__":
    bad = run()
    for f in bad:
        print(f"FAIL  {f}")
    print("report_template_field_check: "
          + (f"{len(bad)} failure(s)" if bad else f"{len(CHECKS)} checks passed"))
    sys.exit(1 if bad else 0)
