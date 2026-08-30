# Copyright 2025 Norman L. Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Build the three files tutorial W-3 ships: the example company template, and
the two reports the page links.

W-3 runs on ``docs/tutorials/files/xslope_johnson_res_solved.xlsx`` — COMBO-1's
model with its mesh, its steady seepage field and its strength reduction solution
beside it (``tools/build_combo01_solved.py``). Those attached solutions are
REPORTED, not re-solved, so the only analysis run here is the limit equilibrium
one, which no sidecar carries:

    report_template_example.docx        the shipped template with a letterhead on it
    w03_report.docx                     Bishop + Spencer, shipped template
    w03_report_example_template.docx    Spencer, built on the example template

The Spencer search is checked against COMBO-1's own published number before
anything is written — 1.248, from the ``<!-- test: -->`` tag on
``combo01_seepage_stability.md`` — because a report that ships as this model's
must carry this model's answers.

The example template is built rather than committed as a hand-edit so the
letterhead can be rebuilt when the shipped template changes: the logo is drawn
here, the header and footer are added with python-docx, and the result has to
pass :func:`xslope.report_docx.template_problem` or nothing is written. The
letterhead goes in a header TABLE and a footer TABLE, which is what survives the
build — the report writes its own running head and its own page numbers into the
first paragraph of each.

Both documents are finished with LibreOffice rather than Word, so the run never
takes over a Word the owner is working in.

    PYTHONPATH=. python3 tools/build_w03_report_files.py
"""

from __future__ import annotations

import contextlib
import io
import os
import subprocess
import sys
import tempfile
import time

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if REPO_ROOT not in sys.path:
    sys.path.insert(0, REPO_ROOT)

import matplotlib

matplotlib.use("Agg")

FILES = os.path.join(REPO_ROOT, "docs", "tutorials", "files")
MODEL = os.path.join(FILES, "xslope_johnson_res_solved.xlsx")
PAGE = os.path.join(REPO_ROOT, "docs", "tutorials",
                    "combo01_seepage_stability.md")
SHIPPED_TEMPLATE = os.path.join(REPO_ROOT, "docs", "studio", "files",
                                "report_template.docx")
EXAMPLE_TEMPLATE = os.path.join(FILES, "report_template_example.docx")
REPORT = os.path.join(FILES, "w03_report.docx")
REPORT_TEMPLATED = os.path.join(FILES, "w03_report_example_template.docx")

#: The firm the example letterhead belongs to. Invented, and named so on the page.
FIRM = "ACME Geotechnical"

#: The title page as W-3 has the reader fill it in, and the same date on both
#: documents: a report whose date moved every time the file was rebuilt would
#: show a different footer in the tutorial's figures than in the file it links.
META = {"title": "Johnson Reservoir Dam — Steady Seepage and Stability",
        "project_number": "2026-231",
        "organization": "Example Engineering",
        "author": "A. Engineer",
        "date": "August 30, 2026"}

#: The number of slices COMBO-1 runs its Spencer search at.
NUM_SLICES = 40

SOFFICE = ("soffice", "/Applications/LibreOffice.app/Contents/MacOS/soffice")


def _quiet(fn, *a, **k):
    with contextlib.redirect_stdout(io.StringIO()):
        return fn(*a, **k)


def draw_logo(path):
    """A placeholder mark for the letterhead: a rounded rectangle with the firm's
    name in it. Drawn rather than committed, because what it stands for is *a
    logo*, and a real one belongs to whoever owns it."""
    from PIL import Image, ImageDraw, ImageFont

    width, height, inset = 1200, 300, 4
    img = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    pen = ImageDraw.Draw(img)
    pen.rounded_rectangle([inset, inset, width - inset, height - inset],
                          radius=48, fill=(232, 238, 244, 255),
                          outline=(31, 73, 125, 255), width=6)
    font = None
    for candidate in ("/System/Library/Fonts/Supplemental/Arial Bold.ttf",
                      "/Library/Fonts/Arial Bold.ttf",
                      "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"):
        if os.path.exists(candidate):
            font = ImageFont.truetype(candidate, 84)
            break
    pen.text((width / 2, height / 2), FIRM, font=font,
             fill=(31, 73, 125, 255), anchor="mm")
    img.save(path)
    return path


def build_template(logo, path=EXAMPLE_TEMPLATE):
    """The shipped template with a letterhead on it, and nothing else changed.

    The logo and the firm's name go in a one-row table in the header, and the
    firm's line in a one-row table in the footer. The report clears the FIRST
    paragraph of each and writes the running head and *page N of M* into it, so
    anything a template wants to keep has to sit outside that paragraph. Each
    table is moved above that paragraph: the letterhead then prints over the
    running head, and the story still ends on a paragraph, which is how Word
    wants a header and a footer written.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    from xslope.report_docx import template_problem

    doc = Document(SHIPPED_TEMPLATE)
    section = doc.sections[0]
    text_width = section.page_width - section.left_margin - section.right_margin

    table = section.header.add_table(1, 2, text_width)
    table.autofit = False
    left, right = table.rows[0].cells
    left.width, right.width = Inches(2.6), Inches(3.9)
    left.paragraphs[0].add_run().add_picture(logo, width=Inches(1.9))
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{FIRM}, Inc.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    section.header.paragraphs[0]._p.addprevious(table._tbl)

    table = section.footer.add_table(1, 1, text_width)
    p = table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{FIRM} · Slope stability report")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    section.footer.paragraphs[0]._p.addprevious(table._tbl)

    doc.save(path)
    problem = template_problem(path)
    if problem:
        os.unlink(path)
        raise SystemExit(f"the example template is not one: {problem}")
    return path


def spencer_bundle(slope_data):
    """W-3's one solve: the Spencer auto search COMBO-1 publishes, checked against
    COMBO-1's own tag before it is allowed into a report."""
    import run_tests as RT
    from xslope.search import run_lem_analysis

    tag = next(t for t in RT.parse_test_tags(PAGE)
               if t.get("type") == "circular_search"
               and t.get("seep") == "steady")
    bundle = _quiet(run_lem_analysis, slope_data, "spencer",
                    analysis="auto_search", num_slices=NUM_SLICES)
    got = bundle["results"]["FS"]
    want, tol = float(tag["expected_fs"]), float(tag["tolerance"])
    print(f"  Spencer, auto search      {got:.4f}   tag {want:g} +/- {tol:g}   "
          f"{'ok' if abs(got - want) <= tol else 'MISMATCH'}")
    if abs(got - want) > tol:
        raise SystemExit(f"Spencer came out {got:.4f} against COMBO-1's {want:g} "
                         f"— nothing written")
    return bundle


def _soffice():
    import shutil

    for name in (os.environ.get("XSLOPE_SOFFICE"),) + SOFFICE:
        if not name:
            continue
        found = name if os.path.isabs(name) and os.path.exists(name) \
            else shutil.which(name)
        if found:
            return found
    raise SystemExit("LibreOffice is not on this machine — nothing finished")


def pages_in(path):
    """How many pages ``path`` lays out to, read off a LibreOffice conversion."""
    with tempfile.TemporaryDirectory(prefix="xslope_w03_") as tmp:
        subprocess.run([_soffice(), "--headless", "--convert-to", "pdf",
                        "--outdir", tmp, path],
                       check=True, capture_output=True, timeout=600)
        pdf = os.path.join(tmp, os.path.splitext(os.path.basename(path))[0] + ".pdf")
        out = subprocess.run(["pdfinfo", pdf], capture_output=True, text=True)
        return int(out.stdout.split("Pages:")[1].split()[0])


def build_report(slope_data, solutions, path, methods, template=None, meta=None):
    """One report, generated then finished with LibreOffice."""
    from xslope.report import generate_report
    from xslope.report_finalize import finalize_report

    options = dict(META, **(meta or {}))
    options.update(input_path=MODEL, method=methods)
    if template:
        options["template"] = template
    ok, out = _quiet(generate_report, slope_data, solutions, options, path)
    if not ok:
        raise SystemExit(f"the report would not build: {out}")
    ok, msg = _quiet(finalize_report, path, prefer="libreoffice")
    if not ok:
        raise SystemExit(f"the report would not finish: {msg}")
    n = pages_in(path)
    print(f"  {os.path.basename(path):<40} {len(out['figures']):>3} figures   "
          f"{n:>3} pages   {os.path.getsize(path) / 1024:>7.1f} KB")
    return out


def main():
    from xslope.fileio import load_slope_data
    from xslope.report import solutions_from_sidecars

    t0 = time.time()
    with tempfile.TemporaryDirectory(prefix="xslope_w03_logo_") as tmp:
        build_template(draw_logo(os.path.join(tmp, "logo.png")))
    print(f"Wrote {os.path.relpath(EXAMPLE_TEMPLATE, REPO_ROOT)}")

    slope_data = _quiet(load_slope_data, MODEL)
    solutions = _quiet(solutions_from_sidecars, MODEL, slope_data)
    print(f"\nAttached to the model: {', '.join(sorted(solutions))} "
          f"— reported, not re-solved")
    solutions["lem"] = [spencer_bundle(slope_data)]

    print("\nReports:")
    build_report(slope_data, solutions, REPORT, ["bishop", "spencer"])
    build_report(slope_data, solutions, REPORT_TEMPLATED, "spencer",
                 template=EXAMPLE_TEMPLATE,
                 meta={"organization": FIRM})
    print(f"\n{time.time() - t0:.0f}s")
    return 0


if __name__ == "__main__":
    sys.exit(main())
