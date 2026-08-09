# Copyright 2025 Norman L. Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""The Word renderer: a :class:`xslope.report.Report` as a ``.docx``.

The document is built ON a template — ``xslope/resources/report_template.docx``
by default — and inherits everything the template declares: the page size and
margins, the Title / Heading / Body / Caption styles, and the header and footer
frames. Nothing here sets a font or a margin; a company template dropped in its
place restyles the whole report (S4 makes that a dialog setting).

What this module does own is the document's *structure*: the title page, the
table-of-contents field, the header and footer content, the landscape section the
slice table needs, and the mapping from content blocks to Word objects.

**Fields, not frozen text.** The title page, the header and the footer carry real
Word fields — ``DOCPROPERTY`` for the metadata, ``PAGE``/``NUMPAGES`` for the page
count, ``TOC`` for the contents, and ``STYLEREF`` for the section the running head
names. Each is written with its result already cached, so the document reads
correctly the moment it is opened and in viewers that never update fields; Word
refreshes them on open or on print, which is what makes "page 3 of 17" and the
contents list true after an edit.

The one field with nothing cached in it is the ``STYLEREF`` pair, whose result
depends on which page it is being laid out on and so does not exist until there
are pages. Its head names the report alone until the fields are computed, which
is true, where a cached guess would name the wrong section.

The metadata maps onto the Word core properties Word's own field names reach:

    ==================  =================
    Report metadata     Core property
    ==================  =================
    Project title       Title
    Project number      Subject
    Organization        Company / Category
    Author              Author
    ==================  =================
"""

from __future__ import annotations

import os
import re
from functools import lru_cache

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips

#: The styles template every report is built on unless the caller names another.
DEFAULT_TEMPLATE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                "resources", "report_template.docx")

#: Style names the renderer asks the template for. A template that does not
#: define one falls back to the document's Normal style rather than failing.
STYLE = {
    "title": "Title",
    "subtitle": "Subtitle",
    "heading": "Heading %d",
    "body": "Body Text",
    "caption": "Caption",
    "bullet": "List Bullet",
    # A contents entry, by Word's internal name for it — the UI calls it "TOC 1".
    "toc": "toc %d",
    "table": "Table Grid",
    # The default table style, which an unstyled (borderless) table inherits its
    # cell margins from.
    "plain_table": "Normal Table",
}

#: The style whose paragraphs the running head names — the top-level sections.
#: A STYLEREF field is given a style's UI name, which is what Word's own field
#: dialog writes and what LibreOffice matches on import; the "heading 1" that
#: styles.xml stores for the same style is the internal name, and resolves in
#: neither.
RUNNING_HEAD_STYLE = STYLE["heading"] % 1

#: Font size for table body text, and the narrower size a wide table falls back
#: to so the slice table's twenty columns still fit a landscape page.
TABLE_PT = 8.5
WIDE_TABLE_PT = 7.0
WIDE_TABLE_COLUMNS = 12

#: Font size for the legend under a table — the line that defines each of its
#: columns, set below the body size so a table and its key read as one block.
LEGEND_PT = 7.5

#: Font size for a key-value block and for the title page's own tables.
KEYVALUE_PT = 10
TITLE_PT = 10.5

#: The table of contents: the field, how deep it goes, and how far each level is
#: indented. The depth is the field switch and the depth of the list written into
#: the field's cached result — the two are the same number so the contents a
#: reader sees cannot differ from the contents Word builds on an update. The
#: indent is Word's own for its latent ``toc N`` styles.
TOC_INSTRUCTION = ' TOC \\o "1-3" \\h \\z \\u '
TOC_LEVELS = 3
TOC_INDENT_TWIPS = 220

#: Word states table geometry in twentieths of a point (twips), and python-docx
#: in English Metric Units; these are the two conversions that need naming.
TWIPS_PER_PT = 20
EMU_PER_TWIP = 635                                  # 914400 EMU/in ÷ 1440 tw/in

#: The cell margin Word assumes when the table style declares none — 0.075 in,
#: the value in Word's own Table Normal.
DEFAULT_CELL_MARGIN = 108

#: The cell margins the report writes onto every table it builds, in twips:
#: ``(side, top_and_bottom)``.
#:
#: Word's own 0.075 in of padding either side of a cell is a body table's margin,
#: and on a twenty-two column slice table it is a third of the landscape page —
#: 4,752 twips of white space, which is more than the six widest columns of
#: numbers put together. The columns are then cut below the width of their own
#: values and Word breaks "-1416.5" across two lines, mid-number.
#:
#: So the report states its own, once, for every table: as little side padding as
#: still separates two columns of digits, and none at all above or below, which
#: with a single-spaced cell paragraph is what makes a dense table. Everything
#: that has to agree with it — the measurement pad, the indent that puts a
#: bordered table's left edge on the text margin — reads this same number.
CELL_MARGIN = (43, 0)

#: The border Word assumes when the table style declares none: a hairline, in
#: eighths of a point, which is the weight its own Table Grid carries.
DEFAULT_BORDER_SZ = 4

#: How much heavier than the table's own rules a totals rule is drawn. The style
#: already rules every row off from the next, so a totals rule at that weight
#: would read as one more row boundary. Doubling is the least that reads as a
#: rule while still taking its weight from the template rather than naming one.
TOTALS_RULE_FACTOR = 2

#: The size a section break's carrier paragraph is set at, in points — the
#: smallest Word's own dialog offers, and small enough that a hidden empty line
#: cannot push anything onto a page of its own (:func:`_collapse_sect_break`).
SECT_BREAK_PT = 1

#: The letters a table's ``align`` is written in, and what Word calls them. A
#: column of numbers only reads as a column when its digits line up, and Word has
#: no notion of a numeric column — the alignment has to be written on the cells.
CELL_ALIGN = {"l": WD_ALIGN_PARAGRAPH.LEFT,
              "c": WD_ALIGN_PARAGRAPH.CENTER,
              "r": WD_ALIGN_PARAGRAPH.RIGHT}

#: Metrically compatible stand-ins, tried in order when the font the template
#: names is not installed where the report is generated. A clone has the same
#: advance widths as the font it replaces, so measuring in one is measuring in
#: the other; anything further down is a ruler of a different length, which
#: still ranks the columns correctly because they are scaled to the page.
FONT_SUBSTITUTES = {
    "calibri": ("Carlito",),
    "cambria": ("Caladea",),
    "times new roman": ("Tinos", "Liberation Serif"),
    "arial": ("Arimo", "Liberation Sans", "Helvetica"),
    "helvetica": ("Arial", "Arimo", "Liberation Sans"),
}


# ---------------------------------------------------------------------------
# Word field plumbing
# ---------------------------------------------------------------------------

def _fld_char(paragraph, kind, dirty=False):
    """One of a field's three markers — begin, separate, end."""
    run = paragraph.add_run()._r
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), kind)
    if dirty:
        fld.set(qn("w:dirty"), "true")
    run.append(fld)


def _instr_text(paragraph, instruction):
    """The field's instruction — what Word recomputes the result from."""
    run = paragraph.add_run()._r
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run.append(instr)


def add_field(paragraph, instruction, cached="", dirty=False):
    """Append a Word field to ``paragraph``, with its result already cached.

    ``dirty`` marks the field for refresh when the document is opened. Nothing
    here sets it: on Word for Mac a dirty field is what raises the "this document
    contains fields that may refer to other files" prompt every time the report
    is opened, which reads as a warning about a clean document.
    """
    _fld_char(paragraph, "begin", dirty)
    _instr_text(paragraph, instruction)
    _fld_char(paragraph, "separate")
    result = paragraph.add_run(cached)
    _fld_char(paragraph, "end")
    return result


def _style(doc, name):
    """A style by name, or None when the template does not define it."""
    try:
        return doc.styles[name]
    except KeyError:
        return None


# ---------------------------------------------------------------------------
# Heading numbers
#
# The headings are numbered 1, 1.1, 1.1.1 by Word's own multilevel list, bound
# to the Heading styles — not by writing the numbers into the heading strings.
# That is what makes the numbers hold: a report generated with a section
# switched off renumbers everything below it without the builder knowing, the
# table of contents picks the numbers up when it is updated, and a cross-
# reference to a section resolves through Word's REF field to whatever the
# heading's number has become.
#
# Word writes exactly this pair when a user applies a multilevel list to the
# heading styles from the ribbon: a numbering definition whose levels name the
# styles, and the same numbering set on those styles. python-docx has no API for
# either, so both are built here as raw OXML.
# ---------------------------------------------------------------------------

#: The numbering definition the headings use. Word numbers list definitions from
#: 1 as a user creates them; a template carrying nine of them is not going to
#: reach this one.
HEADING_NUM_ID = 100


def _el(tag, **attrs):
    """One raw element with its attributes, each in the ``w:`` namespace."""
    el = OxmlElement(tag)
    for name, value in attrs.items():
        el.set(qn("w:" + name), str(value))
    return el


def _heading_level(ilvl, style_id):
    """One level of the heading list: ``%1``, then ``%1.%2``, then ``%1.%2.%3``,
    bound to the heading style of that depth.

    A space separates the number from the title rather than a tab: a tab lands on
    whatever tab stop the style has, and the heading styles have none.
    """
    lvl = _el("w:lvl", ilvl=ilvl)
    lvl.append(_el("w:start", val=1))
    lvl.append(_el("w:numFmt", val="decimal"))
    lvl.append(_el("w:pStyle", val=style_id))
    lvl.append(_el("w:suff", val="space"))
    lvl.append(_el("w:lvlText",
                   val=".".join(f"%{i + 1}" for i in range(ilvl + 1))))
    lvl.append(_el("w:lvlJc", val="left"))
    p_pr = OxmlElement("w:pPr")
    p_pr.append(_el("w:ind", left=0, firstLine=0))
    lvl.append(p_pr)
    return lvl


def _number_headings(doc):
    """Bind the multilevel list to the document's heading styles.

    Returns the number of levels bound — zero for a template that defines no
    heading styles or carries no numbering part, in which case the headings print
    unnumbered rather than the report failing to write.
    """
    from .report import HEADING_LEVELS

    ids = []
    for level in range(1, HEADING_LEVELS + 1):
        style = _style(doc, STYLE["heading"] % level)
        ids.append(None if style is None
                   else style.element.get(qn("w:styleId")))
    if not any(ids):
        return 0
    try:
        numbering = doc.part.numbering_part.element
    except Exception:
        return 0

    abstract = _el("w:abstractNum", abstractNumId=HEADING_NUM_ID)
    abstract.append(_el("w:multiLevelType", val="multilevel"))
    for ilvl, style_id in enumerate(ids):
        if style_id:
            abstract.append(_heading_level(ilvl, style_id))
    num = _el("w:num", numId=HEADING_NUM_ID)
    num.append(_el("w:abstractNumId", val=HEADING_NUM_ID))
    # Every w:abstractNum precedes every w:num, so the definition goes in ahead
    # of the first one the template already carries.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        first_num.addprevious(abstract)
    numbering.append(num)

    bound = 0
    for ilvl, style_id in enumerate(ids):
        if not style_id:
            continue
        p_pr = doc.styles[STYLE["heading"] % (ilvl + 1)].element.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = ilvl
        num_pr.get_or_add_numId().val = HEADING_NUM_ID
        bound += 1
    return bound


def _para(doc, text="", style=None, align=None, size=None, bold=None,
          space_after=None):
    p = doc.add_paragraph()
    if style and _style(doc, style) is not None:
        p.style = doc.styles[style]
    if text:
        run = p.add_run(text)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def _repeat_header_row(row):
    """Mark a table row as a header that repeats on every page it spans."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


#: Where ``w:noWrap`` goes in a cell's properties. Word rejects a table whose
#: ``w:tcPr`` children are out of schema order, and a cell here is written three
#: times over — its text, its rule, its width — so the position is named rather
#: than appended to whatever is already there.
NOWRAP_SUCCESSORS = ("w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign",
                     "w:hideMark", "w:cellIns", "w:cellDel", "w:cellMerge",
                     "w:tcPrChange")


def _no_wrap(cell):
    """Forbid Word from breaking this cell's content across lines.

    For a cell holding one number, which is the only place this is used. Word
    takes an ASCII hyphen-minus as a place a line may break, so a cell one twip
    too narrow prints "-" on one line and "1416.5" on the next — a number
    rendered as two. The columns are measured to hold their widest value
    (:func:`_column_widths`), and this is what makes that measurement binding
    rather than a preference.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:noWrap")) is not None:
        return
    tc_pr.insert_element_before(OxmlElement("w:noWrap"), *NOWRAP_SUCCESSORS)


#: Where ``w:vAlign`` goes in a cell's properties, for the same reason
#: :data:`NOWRAP_SUCCESSORS` exists: the schema fixes the order of ``w:tcPr``.
VALIGN_SUCCESSORS = ("w:hideMark", "w:cellIns", "w:cellDel", "w:cellMerge",
                     "w:tcPrChange")


def _center_cell(cell):
    """Center a cell's content between the top and bottom of its row.

    Word's default is to hang it from the top, which shows the moment two cells
    in a row are not the same height: the taller one sets the row and the shorter
    one's single line sits at the top of it, floating over white. A row of a
    generated table is read across, so its text sits on one line across.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:vAlign")):
        tc_pr.remove(existing)
    el = OxmlElement("w:vAlign")
    el.set(qn("w:val"), "center")
    tc_pr.insert_element_before(el, *VALIGN_SUCCESSORS)


#: Where ``w:sz`` goes in a run's properties, and what follows ``w:szCs`` — the
#: schema fixes their order too.
MARK_SIZE_SUCCESSORS = ("w:sz", "w:szCs", "w:highlight", "w:u", "w:effect",
                        "w:bdr", "w:shd", "w:fitText", "w:vertAlign", "w:rtl",
                        "w:cs", "w:em", "w:lang", "w:eastAsianLayout",
                        "w:specVanish", "w:oMath")


def _mark_size(p, size):
    """Set the paragraph MARK of ``p`` at ``size`` points.

    The invisible character at the end of a paragraph carries formatting of its
    own, and Word lays a line out to fit the tallest thing on it — the mark
    included. A cell written at eight and a half points whose mark is left at the
    document's eleven is laid out for eleven, so a row with one empty cell in it
    came out a fifth taller than the rows either side and its text hung at the
    top of the extra space. An empty cell is the case that shows it, because
    there is nothing else on that line, but the mark is set on every cell: one
    rule, and the row heights are then a property of the text alone.
    """
    if size is None:
        return
    p_pr = p._p.get_or_add_pPr()
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.insert_element_before(r_pr, "w:sectPr", "w:pPrChange")
    for tag, successors in (("w:sz", MARK_SIZE_SUCCESSORS),
                            ("w:szCs", MARK_SIZE_SUCCESSORS[1:])):
        for existing in r_pr.findall(qn(tag)):
            r_pr.remove(existing)
        el = OxmlElement(tag)
        # Word measures a font size in half-points.
        el.set(qn("w:val"), str(int(round(float(size) * 2))))
        r_pr.insert_element_before(el, *successors)


def _cell_text(cell, text, size, bold=False, align=None, nowrap=False):
    """Write one cell's text. The single place a cell is filled, so it is also
    the single place a cell's justification is set — a column centered here is
    centered in its header, its body and its totals alike.

    The paragraph is set tight: no space above or below and single line spacing,
    so a row is as tall as its text and a table of numbers reads as a block
    rather than as a list. ``nowrap`` keeps the cell's content on one line.

    Two things keep the rows of one table the same height and their text on one
    line across: the paragraph mark is set at the cell's own size
    (:func:`_mark_size`), so an empty cell does not lay its row out for the
    document's body size, and the cell is vertically centered
    (:func:`_center_cell`), so a cell that does wrap does not leave its
    neighbours hanging at the top of the row it makes.
    """
    cell.text = ""
    p = cell.paragraphs[0]
    # Emptying a cell leaves the run that held its text behind. It is dropped, so
    # that the cell's first run is the one written here: a caller reaching for
    # runs[0] to restyle a label would otherwise style an empty run and see
    # nothing change.
    for stray in list(p.runs):
        stray._r.getparent().remove(stray._r)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    _mark_size(p, size)
    _center_cell(cell)
    if align is not None:
        p.alignment = align
    if nowrap:
        _no_wrap(cell)
    # A cell that names a symbol sets it as one: the header of the M_R column,
    # the Symbol column of a nomenclature. A cell that names none is written as
    # the single run every caller reaching for ``runs[0]`` expects.
    if INLINE_MATH.search(str(text)):
        _math_runs(p, text, size, bold)
        return
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold


def _column_alignments(align, n_cols):
    """One Word alignment per column, from a table's ``align``.

    A single letter justifies the whole table; a letter per column — as a list,
    or as one string of letters — justifies them one at a time. Anything the map
    does not know falls back to the left, so a typo in a caller costs a
    justification rather than the report.
    """
    letters = ([str(a)[:1] for a in align] if not isinstance(align, str)
               else list(align))
    if len(letters) == 1:
        letters = letters * n_cols
    letters = (letters + ["l"] * n_cols)[:n_cols]
    return [CELL_ALIGN.get(x.lower(), WD_ALIGN_PARAGRAPH.LEFT) for x in letters]


# ---------------------------------------------------------------------------
# Table geometry
#
# Two things Word will not do for a generated table, and both of them show:
#
# 1. A table with no indent hangs its left border OUT into the margin. Word
#    anchors a table on the text in its leading cell, not on the cell's edge, so
#    the border sits one cell margin to the left of every paragraph on the page.
#    An explicit indent of exactly that cell margin puts it back on the text
#    edge — which is why the indent and the padding are one number
#    (:data:`CELL_MARGIN`, written onto every table by :func:`_set_cell_margins`).
#
# 2. An autofitting table gives "#" the same width as "Material". The columns
#    here are measured instead: every cell in a column is set in the report's own
#    font at the table's own size, and the column asks for its widest line
#    (:func:`_column_widths`). A table takes the width those measurements add up
#    to and stops there; only one that will not fit the page is cut down to it,
#    and only a table that asks to be stretched is stretched.
# ---------------------------------------------------------------------------

def _usable_twips(section):
    """The width of ``section``'s text column, in twips.

    Read from the section — the template owns the paper size and the margins, and
    a landscape section has already swapped them.
    """
    return int(round((section.page_width - section.left_margin
                      - section.right_margin) / EMU_PER_TWIP))


def _table_font(doc):
    """The font the document's tables are set in, as its Normal style names it."""
    style = _style(doc, "Normal")
    return (style.font.name if style is not None else None) or ""


@lru_cache(maxsize=None)
def _measuring_face(family, size_pt):
    """A FreeType face for ``family`` at ``size_pt``, or None when there is no
    font on this machine to measure with.

    A point is a point: the face is sized at 72 dpi, so FreeType's advances come
    back in the units Word lays the table out in.
    """
    try:
        from matplotlib import font_manager
        from matplotlib.ft2font import FT2Font
    except Exception:
        return None
    names = [n for n in (family,) + FONT_SUBSTITUTES.get(family.lower(), ()) if n]
    path = None
    for name in names:
        try:
            path = font_manager.findfont(font_manager.FontProperties(family=name),
                                         fallback_to_default=False)
            break
        except Exception:
            continue
    try:
        if path is None:
            path = font_manager.findfont(font_manager.FontProperties())
        face = FT2Font(path)
        face.set_size(size_pt, 72)
        return face
    except Exception:
        return None


@lru_cache(maxsize=None)
def _char_width(family, size_pt, ch):
    """The advance width of one character, in twips."""
    face = _measuring_face(family, size_pt)
    if face is None:
        # Nothing to measure with. Half an em per character is about the mean
        # advance of a proportional Latin face, and only the RELATIVE widths
        # decide the columns — the set is scaled to the page whatever the ruler.
        return size_pt / 2 * TWIPS_PER_PT
    face.set_text(ch)
    return face.get_width_height()[0] / 64.0 * TWIPS_PER_PT


def _text_width(text, family, size_pt):
    """How wide ``text`` sets in ``family`` at ``size_pt``, in twips.

    Advance widths summed character by character: kerning is left out, which
    makes the measurement very slightly generous and never short.

    A symbol is measured as it PRINTS (:func:`_plain`) and not as it is typed:
    "N_S" reaches the page as an N with a subscript S, two characters, and a
    column measured on the three of its notation carries a column of white space
    it never fills. The subscript sets smaller than the letter it rides, so the
    proxy is generous in the direction a column can afford.
    """
    return sum(_char_width(family, size_pt, ch) for ch in _plain(text))


def _apportion(widths, total):
    """``widths`` as whole twips summing to exactly ``total``.

    The floors are handed out first and the leftover twips go to the columns with
    the largest fractions, so no column loses a twip its neighbour did not gain.
    """
    out = [int(w) for w in widths]
    order = sorted(range(len(widths)), key=lambda j: widths[j] - out[j],
                   reverse=True)
    short = total - sum(out)
    while short > 0 and order:
        for j in order[:short]:
            out[j] += 1
        short = total - sum(out)
    while short < 0:
        j = max(range(len(out)), key=lambda k: out[k])
        out[j] -= 1
        short += 1
    return out


def _column_widths(columns, family, size_pt, usable, pad, fit="content",
                   nowrap=None):
    """Column widths in twips, summing to the table's width.

    ``columns`` is every string that will print in each column — its header, its
    cells and its total — one list per column.

    Each column asks for its widest line plus ``pad``, the cell margins that sit
    either side of the text, and floors at its longest single word: Word breaks a
    word that will not fit rather than widening the column, so a column narrower
    than that prints "Applicatio n". A word wider than an equal share of the page
    is the exception — a 64-character file digest is one — and its column's floor
    is capped there: past an equal share, refusing to break the word would starve
    every other column, and Word breaks it anyway.

    ``nowrap`` is one flag per column, true where nothing in the column may be
    broken — a column of numbers, which the cells themselves also declare
    (:func:`_no_wrap`). Such a column floors at its widest VALUE and that floor
    is not capped: an equal share of the page is a sensible place to stop
    widening a column of prose, and no place at all to cut a number in half. The
    set can still be scaled bodily if even the floors do not fit, which is the
    one case left where a table cannot hold what it is given.

    ``fit`` decides what happens to a table that does not need the whole page.
    Under ``"content"`` it keeps the width its content asked for and the table
    ends there: three columns of factors of safety ruled across a seven-inch page
    are a table pretending to be a page. Under ``"page"`` the set is grown in
    proportion to what the columns asked for until it fills the text width, so a
    "#" column stays narrow while the table spans the margins.

    A table wider than the page is cut to a single water level whatever the fit —
    it still has to fit. The widest column loses first and keeps losing until it
    is no wider than the next, and a column below the level is not touched at
    all. That is what lets a long Finding wrap while the Severity beside it still
    prints "Warning" on one line.
    """
    n = max(1, len(columns))
    fair_share = usable / n
    flags = list(nowrap or ())
    flags = (flags + [False] * n)[:n]
    want, floor = [], []
    for j, texts in enumerate(columns):
        widest = max((_text_width(t, family, size_pt) for t in texts), default=0.0)
        longest_word = max((_text_width(w, family, size_pt)
                            for t in texts for w in t.split()), default=0.0)
        want.append(widest + pad)
        # An empty column is still a column: one em is the narrowest that reads
        # as one rather than as a line.
        unbreakable = max(widest if flags[j] else longest_word,
                          size_pt * TWIPS_PER_PT) + pad
        floor.append(unbreakable if flags[j] else min(unbreakable, fair_share))

    natural = [max(want[j], floor[j]) for j in range(n)]
    if fit != "page" and sum(natural) <= usable:
        return _apportion(natural, int(round(sum(natural))))
    if sum(want) <= usable:
        w = [x * usable / max(sum(want), 1.0) for x in want]
    elif sum(floor) >= usable:
        # Every column is already at its floor: the table cannot hold its content
        # at this size, so it is scaled bodily and Word does the breaking.
        w = [x * usable / sum(floor) for x in floor]
    else:
        def at(level):
            return sum(max(min(want[j], level), floor[j]) for j in range(n))

        lo, hi = 0.0, max(want)
        for _ in range(64):
            mid = (lo + hi) / 2
            if at(mid) > usable:
                hi = mid
            else:
                lo = mid
        w = [max(min(want[j], lo), floor[j]) for j in range(n)]
    return _apportion(w, usable)


def _table_indent(table, twips):
    """Indent the whole table, so its left border lands on the text margin."""
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblInd")):
        tbl_pr.remove(existing)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), str(int(twips)))
    ind.set(qn("w:type"), "dxa")
    tbl_pr.insert_element_before(ind, "w:tblBorders", "w:shd", "w:tblLayout",
                                 "w:tblCellMar", "w:tblLook")


#: Where ``w:tblCellMar`` goes in a table's properties, for the same reason
#: :data:`NOWRAP_SUCCESSORS` exists.
CELL_MAR_SUCCESSORS = ("w:tblLook", "w:tblCaption", "w:tblDescription",
                       "w:tblPrChange")


def _set_cell_margins(table, side, vertical):
    """Declare a table's own cell padding, overriding whatever its style says.

    One number, three uses: it is the padding a column carries either side of its
    widest line, it is how far a table's left border overhangs the text when the
    table has no indent, and so it is exactly the indent that cures the overhang.
    """
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblCellMar")):
        tbl_pr.remove(existing)
    mar = OxmlElement("w:tblCellMar")
    for edge, value in (("top", vertical), ("left", side),
                        ("bottom", vertical), ("right", side)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(int(value)))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.insert_element_before(mar, *CELL_MAR_SUCCESSORS)


def _style_rule(doc, style_name):
    """The line a table style rules its rows off with, as ``(val, sz, color)``.

    Read rather than chosen, for the same reason the cell margin is: a totals
    rule drawn at a weight of its own would be the one line in the table that
    does not belong to the template. The inside horizontal border is the line
    between two rows, which is the line a totals rule replaces; a style that
    declares only an outline is asked for that instead, and a style that declares
    no borders at all — a borderless key-value table — falls back to Word's own
    hairline.
    """
    style = _style(doc, style_name)
    if style is not None:
        for edge in ("insideH", "top"):
            found = style.element.xpath(f"./w:tblPr/w:tblBorders/w:{edge}")
            if found:
                try:
                    sz = int(found[0].get(qn("w:sz")))
                except (TypeError, ValueError):
                    sz = DEFAULT_BORDER_SZ
                return (found[0].get(qn("w:val")) or "single", sz,
                        found[0].get(qn("w:color")) or "auto")
    return "single", DEFAULT_BORDER_SZ, "auto"


def _rule_above(cell, val, sz, color):
    """Draw a line along the top of ``cell``, over whatever the style draws there.

    Direct cell formatting outranks the table style, which is what lets a totals
    rule show through a table whose every row is already ruled.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), val)
    top.set(qn("w:sz"), str(int(sz)))
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), color)
    borders.append(top)
    tc_pr.insert_element_before(borders, "w:shd", "w:noWrap", "w:tcMar",
                                "w:textDirection", "w:tcFitText", "w:vAlign",
                                "w:hideMark")


def _fit_table(doc, table, section, columns, size, style_name=STYLE["table"],
               fit="content", nowrap=None):
    """Give ``table`` fixed, measured columns, and the indent that puts its left
    border on the text margin.

    ``columns`` is the text of each column, as :func:`_column_widths` reads it,
    and ``fit`` is that function's — ``"content"`` ends the table where its
    content ends, ``"page"`` stretches it to the text width.

    Word wants the width in three places — the table, its grid and every cell —
    and honours none of them on its own unless the layout is fixed; all three are
    set here, from one set of numbers. The table's own width is the sum of the
    columns and not the page, or Word draws the last column out to the margin
    whatever the grid says.

    Only a BORDERED table is indented. There the border is the edge the eye reads
    and it belongs on the text margin. A borderless table has no such edge: what
    the reader sees is its text, and that text belongs where a line of prose
    begins — which is where Word puts it when the table carries no indent at all.
    """
    margin = CELL_MARGIN[0]
    usable = _usable_twips(section)
    widths = _column_widths(columns, _table_font(doc), size, usable, 2 * margin,
                            fit, nowrap)

    _set_cell_margins(table, *CELL_MARGIN)
    table.autofit = False                       # w:tblLayout w:type="fixed"
    if style_name != STYLE["plain_table"]:
        _table_indent(table, margin)
    tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
    if tbl_w is not None:
        tbl_w.set(qn("w:w"), str(sum(widths)))
        tbl_w.set(qn("w:type"), "dxa")
    for j, width in enumerate(widths):
        table.columns[j].width = Twips(width)
        for cell in table.columns[j].cells:
            cell.width = Twips(width)
    return widths


# ---------------------------------------------------------------------------
# Equations
#
# Word has a math format of its own — OMML — and an equation in it is text: it
# sets in the document's own fonts, it copies, it is searchable, and it can be
# edited by whoever reviews the calculation. A picture of an equation is none of
# those things, so nothing here renders one.
#
# The content tree writes equations in a small notation (:func:`omath`), which
# this compiles to OMML. The notation exists so that the equations can live in
# report.py as readable strings rather than as trees of XML:
#
#     frac{a}{b}      a fraction
#     sum{x}          a summation over x
#     a_b   a^b       subscript and superscript of the preceding character
#     {...}           a group, wherever one character is not enough
#
# Everything else is literal text, Unicode included — α, φ and Δ are the symbols
# the documentation uses and they are what goes in the document.
# ---------------------------------------------------------------------------

#: The n-ary operators the notation offers, by the word that writes them.
NARY = {"sum": "∑"}


def _m(tag, *children):
    """One OMML element, in the math namespace."""
    el = OxmlElement(f"m:{tag}")
    for child in children:
        el.append(child)
    return el


def _m_run(text):
    """A run of literal math text."""
    run = OxmlElement("m:r")
    node = OxmlElement("m:t")
    node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    return run


def _m_val(tag, value):
    el = OxmlElement(f"m:{tag}")
    el.set(qn("m:val"), value)
    return el


def _m_element(nodes):
    """An ``m:e`` argument holding a parsed sequence."""
    return _m("e", *_emit(nodes))


def _stack(rows):
    """An ``m:eqArr`` — several expressions stacked, centered on each other.

    What a fraction's denominator is set in when the sum in it is wider than the
    text column (:func:`_stacked_quotient`).
    """
    return _m("eqArr", _m("eqArrPr", _m_val("baseJc", "center"),
                          _m_val("maxDist", "0")),
              *[_m("e", *_emit(row)) for row in rows])


def _emit(nodes):
    """OMML elements for a parsed sequence."""
    out = []
    for kind, payload in nodes:
        if kind == "text":
            out.append(_m_run(payload))
        elif kind == "stack":
            out.append(_stack(payload))
        elif kind == "frac":
            num, den = payload
            out.append(_m("f", _m("fPr", _m_val("type", "bar")),
                          _m("num", *_emit(num)), _m("den", *_emit(den))))
        elif kind == "nary":
            char, body = payload
            props = _m("naryPr", _m_val("chr", char), _m_val("limLoc", "undOvr"),
                       _m_val("subHide", "1"), _m_val("supHide", "1"))
            out.append(_m("nary", props, _m("sub"), _m("sup"),
                          _m_element(body)))
        elif kind in ("sub", "sup"):
            base, script = payload
            tag = "sSub" if kind == "sub" else "sSup"
            out.append(_m(tag, _m_element(base),
                          _m(kind, *_emit(script))))
    return out


def _script_span(src, at):
    """How far a subscript or superscript written without braces runs: to the end
    of the word it opens.

    ``F_corr`` is F subscript corr — one symbol, the corrected factor of safety —
    and not F subscript c followed by the roman letters orr, which is what the
    equations printed for every symbol whose subscript is more than one
    character: F_corr, m_α's neighbours a_dx and a_ry, every moment arm the
    general equations carry. A brace group is still the way to write a subscript
    that is not one word (``Z_{i+1}``).
    """
    end = at
    while end < len(src) and src[end].isalnum():
        end += 1
    return max(end, at + 1)


def _parse_math(src, pos=0, depth=0):
    """``(nodes, pos)`` — the notation, parsed to the sequence :func:`_emit`
    walks. Stops at an unmatched ``}``."""
    nodes = []
    text = ""

    def flush():
        nonlocal text
        if text:
            nodes.append(("text", text))
            text = ""

    while pos < len(src):
        ch = src[pos]
        if ch == "}":
            break
        word = None
        for name in ("frac",) + tuple(NARY):
            if src.startswith(name + "{", pos):
                word = name
                break
        if word is not None:
            flush()
            first, pos = _parse_math(src, pos + len(word) + 1, depth + 1)
            pos += 1                                    # the closing brace
            if word == "frac":
                if pos < len(src) and src[pos] == "{":
                    second, pos = _parse_math(src, pos + 1, depth + 1)
                    pos += 1
                else:
                    second = []
                nodes.append(("frac", (first, second)))
            else:
                nodes.append(("nary", (NARY[word], first)))
            continue
        if ch == "{":
            flush()
            group, pos = _parse_math(src, pos + 1, depth + 1)
            pos += 1
            nodes.extend(group)
            continue
        if ch in "_^" and pos + 1 < len(src):
            # The base is the character just written, or the group just closed:
            # every subscript in the report's equations is on a single symbol.
            base = []
            if text:
                base = [("text", text[-1])]
                text = text[:-1]
                flush()
            elif nodes:
                base = [nodes.pop()]
            else:
                text += ch
                pos += 1
                continue
            flush()
            if src[pos + 1] == "{":
                script, pos = _parse_math(src, pos + 2, depth + 1)
                pos += 1
            else:
                end = _script_span(src, pos + 1)
                script, pos = [("text", src[pos + 1:end])], end
            nodes.append(("sub" if ch == "_" else "sup", (base, script)))
            continue
        text += ch
        pos += 1
    flush()
    return nodes, pos


def omath_nodes(nodes):
    """An ``m:oMath`` element for an already-parsed sequence."""
    math = OxmlElement("m:oMath")
    for el in _emit(nodes):
        math.append(el)
    return math


def omath(notation):
    """An ``m:oMath`` element for one equation written in the report's notation."""
    return omath_nodes(_parse_math(str(notation))[0])


#: The font a Word equation sets in, and what to fall back on where the machine
#: has no copy of it. Only the WIDTH of the equation is wanted, so a face of
#: about the same set is a usable ruler.
MATH_FONT = "Cambria Math"

#: How wide the letters of an equation run against the plain text of the same
#: size. Word sets math in an italic face with its own spacing around the
#: operators, which is wider than the measured advances of the characters alone.
#: Measured against the equations this report prints, at the width they come out
#: on the page.
MATH_SET = 1.15

#: How much of the text column an equation may fill before it is broken over
#: another line. The measurement is an estimate — it cannot know Word's spacing
#: to the twip — and the margin is what keeps an estimate a little short from
#: running an equation into the margin.
MATH_FILL = 0.92


def _math_width(nodes, family, size_pt):
    """How wide a parsed equation sets, in twips.

    A fraction is as wide as the wider of its two arguments, an n-ary as wide as
    its operator and its body, and a script rides its base at about two thirds
    the size. Kerning and Word's own operator spacing are not modelled; what is
    wanted is a number good enough to decide where a line has to break.
    """
    total = 0.0
    for kind, payload in nodes:
        if kind == "text":
            total += _text_width(payload, family, size_pt) * MATH_SET
        elif kind == "stack":
            total += max((_math_width(row, family, size_pt) for row in payload),
                         default=0.0)
        elif kind == "frac":
            total += max(_math_width(payload[0], family, size_pt),
                         _math_width(payload[1], family, size_pt))
        elif kind == "nary":
            total += (_text_width(payload[0], family, size_pt) * MATH_SET
                      + _math_width(payload[1], family, size_pt))
        else:
            total += (_math_width(payload[0], family, size_pt)
                      + 0.65 * _math_width(payload[1], family, size_pt))
    return total


def _math_segments(notation):
    """One equation's top-level pieces, each opening on the operator that joins
    it to the one before — the places a long equation can be broken.

    Only the top level: an operator inside a fraction, a summation or any other
    group is inside a construction Word sets as one thing, and a break there is
    not a break Word or anything else will honor.
    """
    out, start, depth = [], 0, 0
    i = 0
    while i < len(notation):
        ch = notation[i]
        if ch in "{([":
            depth += 1
        elif ch in "})]":
            depth -= 1
        elif depth == 0 and notation.startswith((" + ", " − ", " = "), i) \
                and i > start:
            out.append(notation[start:i])
            start = i                           # the operator opens the next
            i += 1
            continue
        i += 1
    out.append(notation[start:])
    out = [s for s in out if s.strip()]
    # The two sides of the equation are joined back to the piece before them:
    # the equals sign may not open a line (see :func:`_math_lines`), and the
    # only way it cannot is for it never to be a place a line is broken.
    joined = []
    for piece in out:
        if joined and piece.lstrip().startswith("= "):
            joined[-1] += piece
        else:
            joined.append(piece)
    return joined


def _math_lines(notation, usable_twips, family, size_pt):
    """``notation`` broken into the lines it takes to fit ``usable_twips``.

    Greedy: a line takes segments until the next one would not fit, and the
    operator that joins it to the line above opens the one below. A segment that
    will not fit on a line of its own — a fraction wider than the page — stands
    anyway, because there is nowhere inside a fraction that Word will break, and
    a report that silently dropped it would be worse than one that runs it wide.

    The equals sign is not a break point. A line that ends or begins with one is
    an equation missing a side, which Word sets without complaint and every
    other reader of a .docx marks as an error where the operator should be.
    """
    room = usable_twips * MATH_FILL
    segments = _math_segments(notation)
    widths = [_math_width(_parse_math(s)[0], family, size_pt) for s in segments]
    return _fill(segments, widths, room)


def _brace_group(text, at):
    """``(contents, position after the closing brace)`` for the group opening at
    ``at``, or ``(None, at)`` where ``at`` is not an opening brace."""
    if at >= len(text) or text[at] != "{":
        return None, at
    depth, j = 0, at
    while j < len(text):
        if text[j] == "{":
            depth += 1
        elif text[j] == "}":
            depth -= 1
            if not depth:
                return text[at + 1:j], j + 1
        j += 1
    return None, at


def _sole_fraction(notation):
    """``(head, numerator, denominator)`` for an equation that is one fraction
    with something in front of it — ``F = frac{A}{B}`` — or None.

    The shape every method's factor of safety is printed in, and the only shape
    the stacking below knows how to narrow. Anything with a tail after the
    fraction, or a second fraction beside it, is left alone.
    """
    at, depth = 0, 0
    while at < len(notation):
        ch = notation[at]
        if ch in "([":
            depth += 1
        elif ch in ")]":
            depth -= 1
        elif ch == "{":
            _group, after = _brace_group(notation, at)
            at = after if _group is not None else at + 1
            continue
        elif depth == 0 and notation.startswith("frac{", at):
            num, after = _brace_group(notation, at + 4)
            den, after = _brace_group(notation, after)
            if num is None or den is None or notation[after:].strip():
                return None
            return notation[:at], num, den
        at += 1
    return None


def _fill(segments, widths, room):
    """``segments`` packed greedily into lines no wider than ``room``."""
    lines, line, width = [], "", 0.0
    for segment, w in zip(segments, widths):
        if line and width + w > room:
            lines.append(line)
            line, width = segment.lstrip(), w
        else:
            line += segment
            width += w
    return lines + ([line] if line else [])


def _stacked_quotient(notation, room, family, size_pt):
    """``notation`` parsed with its denominator stacked over as many lines as it
    takes to stand inside ``room`` — or None where it neither needs nor admits it.

    A quotient carrying every force a slice can take is wider than any text
    column, and there is nowhere in a fraction that a line breaks: Word breaks the
    equation at the only operator outside it, which strands ``F`` on a line of its
    own with ``=`` opening the next, and LibreOffice runs the fraction out into
    the margin. Neither is the equation. Stacking the denominator's sum over two
    lines inside the fraction narrows the whole equation instead, so ``F =`` stays
    against the bar it belongs to and nothing leaves the column.
    """
    split = _sole_fraction(notation)
    if split is None:
        return None
    head, num, den = split
    segments = _math_segments(den)
    if len(segments) < 2:
        return None
    widths = [_math_width(_parse_math(s)[0], family, size_pt) for s in segments]
    # Every row stands under the bar, and the bar stands after the head: the room
    # a row has is what the head leaves of the line.
    space = room - _math_width(_parse_math(head)[0], family, size_pt)
    rows = _fill(segments, widths, space)
    if len(rows) < 2:
        return None
    nodes, _pos = _parse_math(head)
    return nodes + [("frac", (_parse_math(num)[0],
                              [("stack", [_parse_math(r)[0] for r in rows])]))]


def _render_math(doc, block, section=None):
    """One displayed equation, centered, on as many lines as it takes.

    Nothing is set beside it: the equations come from several derivations, each
    numbering its own, and the number they carry there is given in the sentence
    that introduces them (:class:`xslope.report.Math`).

    An equation wider than the text column is broken at its top-level operators,
    each line a display of its own. Word wraps a long equation itself, but only
    at those same operators and only in Word: the break has to be in the document
    for the equation to be readable in anything else that opens it, and a
    paragraph apiece is the one break every reader of a .docx honors.

    A quotient has no top-level operator but its own equals sign, and breaking
    there leaves ``F`` standing alone above the fraction. One that will not fit is
    narrowed instead, by stacking its denominator inside the bar
    (:func:`_stacked_quotient`).
    """
    family = _table_font(doc) or MATH_FONT
    size = _style(doc, "Normal")
    size_pt = float(size.font.size.pt) if size is not None and size.font.size \
        else 11.0
    face = MATH_FONT if _measuring_face(MATH_FONT, size_pt) else family
    lines, room = [block.notation], 0.0
    if section is not None:
        room = _usable_twips(section) * MATH_FILL
        lines = _math_lines(block.notation, _usable_twips(section), face, size_pt)
    out = None
    for index, line in enumerate(lines):
        last = index == len(lines) - 1
        nodes, _pos = _parse_math(line)
        if room and _math_width(nodes, face, size_pt) > room:
            nodes = _stacked_quotient(line, room, face, size_pt) or nodes
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4 if not index else 0)
        p.paragraph_format.space_after = Pt(4 if last else 0)
        # One equation is one thing to read: its lines do not fall over a page
        # break away from each other.
        p.paragraph_format.keep_with_next = not last
        para = OxmlElement("m:oMathPara")
        props = OxmlElement("m:oMathParaPr")
        props.append(_m_val("jc", "center"))
        para.append(props)
        para.append(omath_nodes(nodes))
        p._p.append(para)
        out = out or p
    return out


# ---------------------------------------------------------------------------
# Links and bookmarks
# ---------------------------------------------------------------------------

def _bookmark(paragraph, name, ident):
    """Mark ``paragraph``'s position so a link elsewhere can reach it.

    The mark opens after the paragraph's properties and closes at its end, so it
    spans the paragraph — which is what a REF field to a numbered heading reads
    the number off.
    """
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(ident))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(ident))
    p_pr = paragraph._p.find(qn("w:pPr"))
    paragraph._p.insert(0 if p_pr is None else 1, start)
    paragraph._p.append(end)


def _link_style(run, doc):
    """Make a run look like the link it is part of."""
    if _style(doc, "Hyperlink") is not None:
        run.style = doc.styles["Hyperlink"]
    else:
        # No character style to inherit: a link still has to look like one.
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        run.font.underline = True
    return run


def _section_ref(paragraph, text, anchor, doc):
    """Write ``text`` — "Section 2.1" — as a live cross-reference to a heading.

    The number is a REF field on the heading's bookmark with the number switch,
    so it is Word's number and not a copy of one: a section inserted above the
    one being cited changes both the heading and the reference. The result is
    cached, as every field here is, so the sentence reads before anything is
    updated.

    Returns False when the phrase carries no number to make a field of, and the
    caller should write it as ordinary link text.
    """
    import re
    match = re.match(r"^(.*?)(\d+(?:\.\d+)*)$", text)
    if match is None:
        return False
    prefix, number = match.groups()
    if prefix:
        _link_style(paragraph.add_run(prefix), doc)
    _fld_char(paragraph, "begin")
    _instr_text(paragraph, f" REF {anchor} \\r \\h ")
    _fld_char(paragraph, "separate")
    _link_style(paragraph.add_run(number), doc)
    _fld_char(paragraph, "end")
    return True


def _link_run(paragraph, text, target, doc):
    """Append ``text`` to ``paragraph`` as a link.

    ``target`` beginning with ``#`` is a bookmark in this document — Word's own
    cross-reference, which jumps the reader to the table the numbers came from —
    and anything else is an external URL, related to the document part so that
    the link survives being sent to somebody.

    A link to a section is that and a field as well: the words are the link, and
    the number in them is computed by Word from the heading it points at.
    """
    from .report import SECTION_ANCHOR_PREFIX

    link = OxmlElement("w:hyperlink")
    if target.startswith("#"):
        link.set(qn("w:anchor"), target[1:])
    else:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        rel_id = doc.part.relate_to(target, RT.HYPERLINK, is_external=True)
        link.set(qn("r:id"), rel_id)

    # Everything written from here belongs inside the link. Runs are appended at
    # the end of the paragraph, so where the end was is where they start —
    # counted, not identified: an lxml proxy is not the node, and two proxies for
    # one node are not the same object.
    start = len(paragraph._p)
    section = target.startswith("#" + SECTION_ANCHOR_PREFIX)
    if not (section and _section_ref(paragraph, text, target[1:], doc)):
        _link_style(paragraph.add_run(text), doc)
    for child in list(paragraph._p)[start:]:
        paragraph._p.remove(child)
        link.append(child)
    paragraph._p.append(link)


def _mark_caption(caption, kind, block, state):
    """Bookmark a figure's or a table's caption line.

    Two names can land on it. The one every numbered block gets is built from
    its number, and is what the prose's "Figure 4" and "Table 5" cross-reference
    — the citation convention a technical report is read under. A block may also
    carry a bookmark of its own, under a name a particular paragraph knows: the
    slice table's, which the calculations reach for by method rather than by
    number. Both are placed, so both kinds of reference land on the line that
    names the block.
    """
    if state is None:
        return
    from .report import cite_anchor

    names = []
    if getattr(block, "number", 0):
        names.append(cite_anchor(kind, block.number))
    own = getattr(block, "bookmark", "")
    if own and own not in names:
        names.append(own)
    for name in names:
        state["bookmark"] = state.get("bookmark", 0) + 1
        _bookmark(caption, name, state["bookmark"])


def _phrase_spans(text, phrases):
    """Where each phrase falls in ``text``, as non-overlapping ``(start, end,
    kind, payload)`` spans in reading order.

    ``phrases`` is ``[(phrase, kind, payload), ...]``. Each takes the first
    occurrence no earlier phrase has claimed, so two links reading the same words
    mark the first and the second of them, and a phrase that does not occur is
    skipped in silence — a sentence rewritten in report.py should lose a link or
    a bold word, not raise.

    Overlap is what keeps a link and a bold phrase from clobbering each other: a
    span inside one already claimed cannot be marked, because a run is either a
    link or it is not, and the paragraph is written as one pass over the text.
    """
    spans = []
    for phrase, kind, payload in phrases:
        if not phrase:
            continue
        start = 0
        while True:
            at = text.find(str(phrase), start)
            if at < 0:
                break
            end = at + len(str(phrase))
            if not any(at < e and s < end for s, e, _k, _p in spans):
                spans.append((at, end, kind, payload))
                break
            start = at + 1
    return sorted(spans)


#: A symbol of the report's own notation standing in a sentence — N_S, a_S, m_α,
#: θ_p, Z_{i+1}, f_o.
#:
#: The prose names the letters the equations above it carry, and a letter set as
#: plain text with an underscore in it is not the letter the equation printed:
#: it leaves the reader to translate "N_S" into the N-sub-S they were shown, in
#: the one place the two have to be the same thing. So each is set as real math,
#: from the same notation and through the same compiler the displayed equations
#: use.
#:
#: A letter, an underscore, then either a brace group or a run of letters and
#: digits — which is what :func:`_script_span` reads as a subscript, so a symbol
#: is marked here exactly where it would be set as one.
INLINE_MATH = re.compile(
    r"(?<![0-9A-Za-z_])"
    r"[A-Za-zΑ-Ωα-ω]_(?:\{[^{}\s]+\}|[A-Za-z0-9Α-Ωα-ω]+)")


def _math_spans(text, claimed):
    """``(start, end, "math", None)`` for every symbol in ``text`` that no span
    in ``claimed`` already covers."""
    return [(m.start(), m.end(), "math", None)
            for m in INLINE_MATH.finditer(text)
            if not any(m.start() < e and s < m.end()
                       for s, e, _k, _p in claimed)]


def _plain(text):
    """``text`` with its symbols spelled as they PRINT — the notation's
    underscores and braces gone.

    A width is measured on this and not on the notation: "Z_{i+1}" is typed with
    four characters that never reach the page, and a column measured on them is
    a column of white space. Everything the pattern does not claim is left
    exactly as it stands, so a file name with an underscore in it is untouched.
    """
    return INLINE_MATH.sub(
        lambda m: m.group(0).translate({ord("_"): None, ord("{"): None,
                                        ord("}"): None}),
        str(text))


def _math_runs(p, text, size=None, bold=False):
    """Write ``text`` into paragraph ``p``, every symbol in it set as a symbol.

    The same marking the prose is written with (:data:`INLINE_MATH`): a table
    header, a nomenclature's Symbol column and the legend under a table all name
    the letters the equations carry, and a letter spelled with an underscore is
    not the letter the equation printed.

    Set as a real subscript rather than as Word math, which is what a paragraph
    of prose uses. A table is set at seven or eight and a half points and Word
    math takes the size of the document, not of the run it stands in: LibreOffice
    ignores a size on the math run altogether — asked three ways, it renders the
    symbol at its own base size — and a nomenclature came out with every symbol
    half again as tall as the meaning beside it and its rows uneven. A subscript
    run carries the cell's own size and weight, in both programs, and is still
    text a reader can copy and search.
    """
    text = str(text)
    at = 0

    def run(piece, subscript=False):
        r = p.add_run(piece)
        if size is not None:
            r.font.size = Pt(size)
        r.font.bold = bold
        r.font.subscript = subscript
        return r

    for m in INLINE_MATH.finditer(text):
        if m.start() > at:
            run(text[at:m.start()])
        base, _sep, script = m.group(0).partition("_")
        run(base)
        run(script.strip("{}"), subscript=True)
        at = m.end()
    if at < len(text) or not at:
        run(text[at:])
    return p


def _render_prose(doc, block):
    """A paragraph, with its linked phrases turned into links, its bold phrases
    set in bold, and the symbols it names set as math.

    The first two are found in the text by the phrase itself, so the sentence is
    written once in report.py, as a sentence, rather than assembled from
    fragments. The symbols are found by their own shape (:data:`INLINE_MATH`), so
    a sentence that names one is typeset with it without the builder having to
    say so twice.
    """
    p = _para(doc, "", style=STYLE["body"])
    marks = [(display, "link", target)
             for display, target in getattr(block, "links", None) or [] if target]
    marks += [(phrase, "bold", None)
              for phrase in getattr(block, "bold", None) or []]
    spans = _phrase_spans(block.text, marks)
    spans = sorted(spans + _math_spans(block.text, spans))
    at = 0
    for start, end, kind, payload in spans:
        if start > at:
            p.add_run(block.text[at:start])
        if kind == "link":
            _link_run(p, block.text[start:end], payload, doc)
        elif kind == "math":
            p._p.append(omath(block.text[start:end]))
        else:
            p.add_run(block.text[start:end]).font.bold = True
        at = end
    if at < len(block.text):
        p.add_run(block.text[at:])
    return p


# ---------------------------------------------------------------------------
# Page furniture
# ---------------------------------------------------------------------------

def _header_paragraph(section):
    """The one paragraph of ``section``'s running head, emptied and unlinked."""
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs[1:]):
        p._element.getparent().remove(p._element)
    p = header.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    return p


def _write_header(section, title):
    """The front-matter head: the project title, as a live document-property
    field.

    The front matter is the pages before the first section — the title page and
    the contents — so the head over them names the report and nothing else.
    """
    p = _header_paragraph(section)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, ' DOCPROPERTY "Title" \\* MERGEFORMAT ', title)
    for run in p.runs:
        run.font.size = Pt(9)


def _tab_to_right_margin(paragraph, section):
    """Give ``paragraph`` one tab stop, at ``section``'s right margin.

    The Header style carries stops of its own, measured for a portrait page.
    They are cleared rather than left to compete: a landscape section is three
    inches wider, and what the section name has to reach is the margin of the
    section it is printed in, not the margin the style was written for.
    """
    stops = paragraph.paragraph_format.tab_stops
    pos = section.page_width - section.left_margin - section.right_margin
    style = paragraph.style
    inherited = list(style.paragraph_format.tab_stops) if style is not None else []
    for stop in inherited:
        if stop.position != pos:
            stops.add_tab_stop(stop.position, WD_TAB_ALIGNMENT.CLEAR)
    stops.add_tab_stop(pos, WD_TAB_ALIGNMENT.RIGHT)


def _write_running_header(section, title):
    """The running head over the body: the report title at the left margin, the
    section the page is in at the right one.

    Neither string is typed into the header. The title is the document's own
    Title property, and the section is a pair of STYLEREF fields reading the
    :data:`RUNNING_HEAD_STYLE` heading that governs the page — its number, then
    its text, so the head reads "3 Limit Equilibrium Analysis", the way the
    heading itself prints. Word recomputes both as it lays the pages out, so a
    section that is added, renamed or renumbered carries its head with it and
    there is nothing to keep in step by hand.

    Nothing is cached in the STYLEREF fields. A cached result is what a reader
    sees until the fields are next computed, and there is no page at write time
    to read a heading from: an empty result leaves the head naming the report
    alone, which is true, where a guess would name the wrong section.
    """
    p = _header_paragraph(section)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _tab_to_right_margin(p, section)
    add_field(p, ' DOCPROPERTY "Title" \\* MERGEFORMAT ', title)
    p.add_run("\t")
    add_field(p, f' STYLEREF "{RUNNING_HEAD_STYLE}" \\n ')
    p.add_run(" ")
    add_field(p, f' STYLEREF "{RUNNING_HEAD_STYLE}" ')
    for run in p.runs:
        run.font.size = Pt(9)


def _write_footer(section, date):
    """The running foot: page N of M, and the report date."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs[1:]):
        p._element.getparent().remove(p._element)
    p = footer.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    add_field(p, " PAGE ", "1")
    p.add_run(" of ")
    add_field(p, " NUMPAGES ", "1")
    if date:
        p.add_run(f"     |     {date}")
    for run in p.runs:
        run.font.size = Pt(9)


def _title_page(doc, meta, section):
    """Title page: a left-aligned title block — organization, title, the document
    type — then the metadata, and only when asked for the prepared-by / checked-by
    signature lines.

    Left-aligned and ranged from one margin: a calculation package is a working
    document, and a centered title block reads as a cover slide. The hierarchy does
    the work instead — the organization small above the title, the title large and
    closed by the rule its style carries, the document type beneath, and white
    space between that block and the metadata.
    """
    for _ in range(4):
        _para(doc, "")

    if meta.get("organization"):
        p = _para(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
        add_field(p, ' DOCPROPERTY "Company" \\* MERGEFORMAT ', meta["organization"])
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.all_caps = True

    p = doc.add_paragraph()
    if _style(doc, STYLE["title"]) is not None:
        p.style = doc.styles[STYLE["title"]]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field(p, ' DOCPROPERTY "Title" \\* MERGEFORMAT ', meta.get("title", ""))

    _para(doc, meta.get("document_type", ""), style=STYLE["subtitle"],
          align=WD_ALIGN_PARAGRAPH.LEFT)

    for _ in range(2):
        _para(doc, "")

    # A row with nothing in it is left out: not every project has a number, an
    # organization or a named author, and a blank beside a label is a question the
    # title page asks and does not answer. "Author" rather than "Prepared by": the
    # signature block below owns that phrase.
    rows = [("Project:", meta.get("project_number", ""), "Subject"),
            ("Author:", meta.get("author", ""), "Author"),
            ("Date:", meta.get("date", ""), None)]
    rows = [r for r in rows if str(r[1]).strip()]
    if rows:
        table = doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (label, value, prop) in enumerate(rows):
            _cell_text(table.rows[i].cells[0], label, TITLE_PT, bold=True)
            cell = table.rows[i].cells[1]
            cell.text = ""
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(1)
            cp.paragraph_format.space_after = Pt(1)
            # This cell is filled by hand rather than through :func:`_cell_text`
            # — its value is a live document property, not a string — so it takes
            # the two rules that keep a row one height on its own.
            _mark_size(cp, TITLE_PT)
            _center_cell(cell)
            if prop:
                add_field(cp, f' DOCPROPERTY "{prop}" \\* MERGEFORMAT ', value)
            else:
                cp.add_run(value)
            for run in cp.runs:
                run.font.size = Pt(TITLE_PT)
        # Measured, not autofitted: the labels want a narrow column so the values
        # sit beside them rather than half a page away.
        _fit_table(doc, table, section,
                   [[label for label, _v, _p in rows],
                    [str(value) for _l, value, _p in rows]],
                   TITLE_PT, style_name=STYLE["plain_table"])

    if meta.get("signature_lines"):
        for _ in range(3):
            _para(doc, "")
        sig = doc.add_table(rows=2, cols=2)
        sig.alignment = WD_TABLE_ALIGNMENT.LEFT
        sig_cells = [[], []]
        for col, label in ((0, "Prepared by"), (1, "Checked by")):
            rule, caption = "_" * 34, f"{label}                    Date"
            _cell_text(sig.rows[0].cells[col], rule, TITLE_PT)
            _cell_text(sig.rows[1].cells[col], caption, 9)
            sig_cells[col] = [rule, caption]
        # The one table in the report that wants the whole page: two signature
        # blocks belong one at each margin, not side by side in the middle.
        _fit_table(doc, sig, section, sig_cells, TITLE_PT,
                   style_name=STYLE["plain_table"], fit="page")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _toc_paragraph(doc, level):
    """A contents-entry paragraph at ``level``, in the template's style for it.

    A template that declares no ``TOC N`` style still gets the indent, so the
    list reads as a list of levels rather than a flat column.
    """
    p = doc.add_paragraph()
    name = STYLE["toc"] % min(level, TOC_LEVELS)
    if _style(doc, name) is not None:
        p.style = doc.styles[name]
    else:
        p.paragraph_format.left_indent = Twips(TOC_INDENT_TWIPS * (level - 1))
        p.paragraph_format.space_after = Pt(2)
    return p


def _contents_page(doc, report):
    """The table of contents: a real TOC field whose cached result is the
    report's own heading list.

    The field is deliberately NOT marked dirty — see :func:`add_field` — so the
    reader gets no prompt on open, and therefore never gets Word's "right-click
    to update" placeholder either. What is cached instead is the contents
    themselves, taken from the same content tree the document is written from, so
    the list cannot disagree with the document: a section switched off is not in
    the tree and is not in the contents.

    Each entry carries its heading number, from the same walk that tells a
    cross-reference what a section is numbered — so what a reader sees before
    Word first updates the field is what Word will compute for it.

    No page numbers. Where the sections fall is Word's to compute, and a guessed
    number in a calculation package is worse than none; the line under the last
    entry says where they come from, and sits INSIDE the field result so Word's
    first update replaces it along with the rest.

    The page the contents end on is closed by the section break that opens the
    body — see :func:`_begin_body`, which is what starts the next page — not by
    a page break of its own.
    """
    _para(doc, "Table of Contents", size=14, bold=True, space_after=10)

    entries = [(lvl, f"{number} {sec.title}")
               for number, lvl, sec in report.section_numbers()
               if lvl <= TOC_LEVELS]
    first = _toc_paragraph(doc, entries[0][0] if entries else 1)
    _fld_char(first, "begin")
    _instr_text(first, TOC_INSTRUCTION)
    _fld_char(first, "separate")
    if entries:
        first.add_run(entries[0][1])
    for level, title in entries[1:]:
        _toc_paragraph(doc, level).add_run(title)

    hint = doc.add_paragraph()
    hint.paragraph_format.space_before = Pt(10)
    run = hint.add_run("Page numbers appear when the table is updated in Word "
                       "(right-click → Update Field, or select all and "
                       "press F9).")
    run.font.size = Pt(9)
    run.font.italic = True
    _fld_char(hint, "end")


def _set_orientation(section, landscape, margin_in):
    """Rotate a section, swapping the page dimensions with it."""
    width, height = section.page_width, section.page_height
    is_landscape = width > height
    if landscape != is_landscape:
        section.page_width, section.page_height = height, width
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    for side in ("left_margin", "right_margin"):
        setattr(section, side, Inches(margin_in))


def _paragraph_is_empty(p):
    """True when a paragraph puts nothing on the page — no text, no picture.

    Only such a paragraph may be collapsed to a hidden mark; see
    :func:`_collapse_sect_break`.
    """
    return not p.text.strip() and not p._p.findall(".//" + qn("w:drawing"))


def _sect_break_carrier(doc):
    """The empty paragraph that carries the closing section break just written.

    The section break is the last ``w:sectPr`` in the body's paragraphs. It is
    written on a paragraph of its own, and it stays on one: if it is found on a
    paragraph that carries content, it is moved to an empty paragraph appended
    after it, which is the paragraph returned and collapsed.
    """
    carrier = next((p for p in reversed(doc.paragraphs)
                    if p._p.find(qn("w:pPr")) is not None
                    and p._p.pPr.find(qn("w:sectPr")) is not None), None)
    if carrier is None:
        return doc.paragraphs[-1] if doc.paragraphs else None
    if _paragraph_is_empty(carrier):
        return carrier
    sect_pr = carrier._p.pPr.find(qn("w:sectPr"))
    own = doc.add_paragraph()
    own._p.get_or_add_pPr().append(sect_pr)
    return own


def _collapse_sect_break(doc):
    """Make the paragraph that carries a closing section break take no room.

    Word writes a section break as a paragraph, and that paragraph is a real,
    empty line at the end of the section it closes. Nothing prints on it, so it
    is invisible — until the section's content ends near the foot of a page,
    which a table that spans pages does sooner or later. Then the empty line has
    nowhere to go but a page of its own, and the report grows a blank landscape
    sheet between the slice table and the section after it.

    The break has to stay; only its paragraph is made to vanish, by the idiom
    Word's own users are told to use: the smallest size the format expresses, an
    exact line height to match, no space above or below, and the paragraph mark
    marked hidden. Applied to every section break the renderer writes rather than
    to the one that happened to show, because which one lands at the foot of a
    page is a property of the model being reported, not of the renderer.

    What is made to vanish is a paragraph mark and nothing else. A paragraph
    holding text or a picture is never given this treatment: an exact line height
    of a point crops whatever is set in it, so a sentence in a collapsed
    paragraph is present in the file and absent from the page. When the break
    lands on a paragraph that carries content — a refusal sentence at the foot of
    a landscape section is where it happens — the break is moved onto an empty
    paragraph of its own after it, and that one is collapsed.
    """
    p = _sect_break_carrier(doc)
    if p is None:
        return None
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(SECT_BREAK_PT)
    p_pr = p._p.get_or_add_pPr()
    # The paragraph MARK's own formatting: a w:rPr inside w:pPr, which sits after
    # everything else the properties carry and before the section break itself.
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.insert_element_before(r_pr, "w:sectPr", "w:pPrChange")
    for tag, value in (("w:vanish", None),
                       ("w:sz", str(int(SECT_BREAK_PT * 2))),
                       ("w:szCs", str(int(SECT_BREAK_PT * 2)))):
        el = OxmlElement(tag)
        if value is not None:
            el.set(qn("w:val"), value)
        r_pr.append(el)
    return p


def _open_section(doc):
    """Start a Word section, carrying the page furniture into it.

    A new section inherits the previous one's properties, including its
    "different first page" flag — which would blank the header and footer on the
    first page of every section the report opens. It is cleared here, and the
    footer is relinked, so the page count carries across rather than stopping at
    a section boundary.

    The head is written afresh instead of linked. It is the one piece of
    furniture whose layout is a property of the section it prints on: the
    section name sits at the right margin, and where that is depends on how wide
    the page is.
    """
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    # add_section leaves the closing properties on a new paragraph at the end of
    # the section just closed. That paragraph is never wanted; it is only where
    # Word keeps a section break.
    _collapse_sect_break(doc)
    section.different_first_page_header_footer = False
    section.footer.is_linked_to_previous = True
    return section


def _begin_body(doc, meta):
    """Close the front matter and open the body of the report.

    The break is a section break, not a page break, because the two parts of the
    document want different heads: the front matter has no section to name, and
    a STYLEREF over the contents page would reach forward and label it with the
    first heading of the report. Splitting them at the sectPr boundary is what
    keeps the contents page naming the report and the body pages naming their
    own sections.
    """
    body = _open_section(doc)
    _write_running_header(body, meta.get("title", ""))
    return body


def ensure_orientation(doc, state, landscape):
    """Put the document into the requested orientation, starting a new Word
    section when it is not already there.
    """
    if landscape == state["landscape"]:
        return state["section"]
    section = _open_section(doc)
    _set_orientation(section, landscape, 0.75 if landscape else 1.0)
    _write_running_header(section, state.get("title", ""))
    state["section"], state["landscape"] = section, landscape
    return section


# ---------------------------------------------------------------------------
# Blocks
# ---------------------------------------------------------------------------

def _content_width_in(section):
    return (section.page_width - section.left_margin - section.right_margin) / 914400


def _content_height_in(section):
    return (section.page_height - section.top_margin - section.bottom_margin) / 914400


def _para_spacing_in(doc, style_name):
    """The vertical space a paragraph in ``style_name`` puts around itself."""
    style = _style(doc, style_name)
    pf = getattr(style, "paragraph_format", None)
    extra = 0
    for gap in ("space_before", "space_after"):
        extra += getattr(pf, gap, None) or 0
    return extra / 914400


def _line_multiple(doc, style_name):
    """The line-spacing multiple ``style_name`` sets, or 1.

    It applies to a line holding a picture as much as to one holding text: a
    paragraph set at 1.08 lines draws a 6-inch image in a 6.5-inch line, which is
    the difference between a figure that fits its page and one that does not.
    """
    style = _style(doc, style_name)
    spacing = getattr(getattr(style, "paragraph_format", None), "line_spacing", None)
    return spacing if isinstance(spacing, float) and spacing > 0 else 1.0


#: How many lines of caption a figure leaves room for. Two, not one: a caption
#: naming a method wraps on a portrait page, and a figure that fitted only
#: because its own caption happened to be short would overflow the next report.
CAPTION_LINES = 2


def _caption_height_in(doc, lines=CAPTION_LINES):
    """How much vertical room a caption takes, read off the template's own
    Caption style rather than assumed. It is what a figure has to leave for the
    lines that name it."""
    style = _style(doc, STYLE["caption"])
    size = getattr(getattr(style, "font", None), "size", None) or Pt(10)
    lead = size * _line_multiple(doc, STYLE["caption"]) * lines
    return lead / 914400 + _para_spacing_in(doc, STYLE["caption"])


def _lead_height_in(doc, lines=1):
    """How much vertical room the sentence a figure is kept with takes.

    One line, because keep-with-next binds the LAST line of a paragraph to what
    follows: a lead that wraps leaves the rest of itself on the page above. Read
    off the template's body style, as the caption's room is read off the caption
    style.
    """
    style = _style(doc, STYLE["body"])
    size = getattr(getattr(style, "font", None), "size", None) or Pt(11)
    lead = size * _line_multiple(doc, STYLE["body"]) * lines
    return lead / 914400 + _para_spacing_in(doc, STYLE["body"])


def _render_table(doc, block, section, state=None):
    """One table, its caption above it and its legend beneath."""
    n_cols = max(1, len(block.headers))
    size = WIDE_TABLE_PT if n_cols > WIDE_TABLE_COLUMNS else TABLE_PT

    if block.caption:
        caption = _para(doc, f"Table {block.number}. {block.caption}",
                        style=STYLE["caption"])
        # A table's caption travels with its first row: left behind at the foot of
        # the previous page it names whatever happens to be above it.
        caption.paragraph_format.keep_with_next = True
        # The bookmark goes on the caption: a cross-reference to a table should
        # land on the line that names it, not inside its first cell.
        _mark_caption(caption, "Table", block, state)

    align = _column_alignments(getattr(block, "align", "l"), n_cols)
    totals = list(getattr(block, "totals", None) or [])

    # A cell holding a number is kept on one line, and its column is measured to
    # hold it. The header is not: "W (lb/ft)" over a column of four-digit forces
    # is meant to wrap, and a header that could not would set the width of every
    # column of numbers in the table.
    from .columns import is_number

    table = doc.add_table(rows=1, cols=n_cols)
    if _style(doc, STYLE["table"]) is not None:
        table.style = doc.styles[STYLE["table"]]
    for j, head in enumerate(block.headers):
        _cell_text(table.rows[0].cells[j], head, size, bold=True, align=align[j])
    _repeat_header_row(table.rows[0])
    # The rows the report singles out are set in bold — the whole row, so it reads
    # as one line rather than as a bold cell beside plain ones.
    bold_rows = set(getattr(block, "bold_rows", None) or [])
    for i, row in enumerate(block.rows):
        cells = table.add_row().cells
        for j in range(n_cols):
            text = row[j] if j < len(row) else ""
            _cell_text(cells[j], text, size, bold=i in bold_rows, align=align[j],
                       nowrap=is_number(text))
    if totals:
        # A totals row is not another datum: it is set in bold and ruled off from
        # the body, so a reader scanning the table sees where the data stop and
        # the sums begin. Columns with no meaningful sum are left empty, and the
        # rule still crosses them — a rule that stopped short would read as a
        # border around the cells that happen to hold numbers.
        cells = table.add_row().cells
        val, sz, color = _style_rule(doc, STYLE["table"])
        for j in range(n_cols):
            text = totals[j] if j < len(totals) else ""
            _cell_text(cells[j], text, size, bold=True, align=align[j],
                       nowrap=is_number(text))
            _rule_above(cells[j], val, sz * TOTALS_RULE_FACTOR, color)
    # Every string that prints, the totals row included — a sum is the widest
    # number in its column as often as not, and a column measured without it
    # prints "78357." over "0".
    printed = [[block.headers[j] if j < len(block.headers) else ""]
               + [r[j] if j < len(r) else "" for r in block.rows]
               + ([totals[j]] if j < len(totals) else [])
               for j in range(n_cols)]
    _fit_table(doc, table, section, printed, size,
               fit=getattr(block, "fit", "content"),
               nowrap=[all(is_number(t) for t in texts[1:] if t.strip())
                       and any(is_number(t) for t in texts[1:])
                       for texts in printed])

    if block.legend:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        for i, (term, definition) in enumerate(block.legend):
            if i:
                p.add_run("   ")
            # The term and the sentence that defines it both name symbols — the
            # column's own letter, and the equation the column holds a term of —
            # and each is set as the symbol the equations print.
            _math_runs(p, f"{term}: ", LEGEND_PT, bold=True)
            _math_runs(p, definition, LEGEND_PT)
    _para(doc, "")


def _render_keyvalues(doc, block, section):
    from .columns import is_number

    if block.title:
        _para(doc, block.title, bold=True, space_after=2)
    table = doc.add_table(rows=len(block.items), cols=2)
    for i, (label, value) in enumerate(block.items):
        _cell_text(table.rows[i].cells[0], label, KEYVALUE_PT)
        # A value that is a number is one here too — "16" slices belongs on one
        # line as much as a force in the slice table does.
        _cell_text(table.rows[i].cells[1], value, KEYVALUE_PT,
                   nowrap=is_number(value))
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    # Each column takes the width its own text needs: a definition list, not two
    # half-page columns with the values stranded away from their labels.
    _fit_table(doc, table, section,
               [[str(label) for label, _v in block.items],
                [str(value) for _l, value in block.items]],
               KEYVALUE_PT, style_name=STYLE["plain_table"])
    _para(doc, "")


def _height_limited_width(path, room_in):
    """The widest a picture can be drawn and still stand ``room_in`` inches tall.

    Its proportions come from the file, through python-docx's own image reader —
    the same one that will place it — so an image it cannot read simply imposes
    no limit rather than taking the report down.
    """
    if room_in <= 0:
        return float("inf")
    try:
        from docx.image.image import Image as DocxImage
        image = DocxImage.from_file(path)
        px_w, px_h = image.px_width, image.px_height
    except Exception:
        return float("inf")
    if not px_w or not px_h:
        return float("inf")
    return room_in * px_w / px_h


def _render_figure(doc, block, section, state=None, kept=False):
    # A figure asks for a width in inches, or for zero — "as wide as this page
    # allows", which is how the slice key fills the landscape page it shares with
    # the table it keys without naming a number that holds for one paper size.
    full = _content_width_in(section) - 0.05
    width = full if block.width_in <= 0 else min(block.width_in, full)
    if os.path.exists(block.path):
        # And never taller than the page it is on, less the line that names it: a
        # picture that overflows is moved to a page of its own by Word and leaves
        # its caption behind on the previous one, which reads as a caption for
        # whatever came before. Scaled down by its own proportions instead.
        # A figure kept with the sentence above it shares its page with that
        # sentence's last line, and has to leave room for it: sized to the page
        # without it, the three together overflow and Word moves the figure on,
        # taking the sentence with it and leaving most of a page blank. The
        # reinforcement report grew six pages that way.
        room = ((_content_height_in(section) - _caption_height_in(doc)
                 - _para_spacing_in(doc, "Normal")
                 - (_lead_height_in(doc) if kept else 0))
                / _line_multiple(doc, "Normal"))
        width = min(width, _height_limited_width(block.path, room))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # The picture and the line that names it are one thing: a page break
        # between them leaves a caption at the top of a page under nothing.
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(block.path, width=Inches(width))
    caption = _para(doc, f"Figure {block.number}. {block.caption}",
                    style=STYLE["caption"], align=WD_ALIGN_PARAGRAPH.CENTER)
    # A figure's caption ends the figure, and is bound to nothing after it. The
    # template's Caption style carries keep-with-next — right for a table, whose
    # caption stands above its first row, and wrong here: it chained each figure
    # to the heading, sentence and figure that followed, and once a lead sentence
    # was bound to its own figure as well the chain ran longer than a page. Word
    # and LibreOffice both then break it wherever they can, which put a sentence
    # at the foot of one page, its figure alone on the next, and half of each
    # page blank.
    caption.paragraph_format.keep_with_next = False
    _mark_caption(caption, "Figure", block, state)


def _render_blocks(doc, blocks, state):
    """Render a section's blocks, opening and closing a landscape page around any
    block that asks for one.

    A sentence that introduces a figure is kept on the page its figure lands on.
    Each figure of a transient march stands under a sentence naming the state it
    is, and four of those sentences per report were left at the foot of one page
    with the figure they name at the top of the next — one page of the Johnson
    report ran a third blank below the orphaned line. The property is Word's own
    keep-with-next, so the sentence moves to the figure rather than the figure
    being dragged back to the sentence.
    """
    kept = False
    for i, block in enumerate(blocks):
        ensure_orientation(doc, state, bool(getattr(block, "landscape", False)))
        under_a_sentence, kept = kept, False
        if block.kind == "prose":
            p = _render_prose(doc, block)
            following = blocks[i + 1] if i + 1 < len(blocks) else None
            # Not across a change of orientation: a figure that asks for a
            # landscape page opens one, and a section break is a page break the
            # sentence cannot follow it over.
            kept = (following is not None and following.kind == "figure"
                    and bool(getattr(following, "landscape", False))
                    == bool(state.get("landscape")))
            p.paragraph_format.keep_with_next = kept
        elif block.kind == "math":
            _render_math(doc, block, state["section"])
        elif block.kind == "bullets":
            for item in block.items:
                _para(doc, item, style=STYLE["bullet"])
        elif block.kind == "keyvalues":
            _render_keyvalues(doc, block, state["section"])
        elif block.kind == "figure":
            _render_figure(doc, block, state["section"], state,
                           kept=under_a_sentence)
        elif block.kind == "table":
            _render_table(doc, block, state["section"], state)


def _render_section(doc, section_node, level, state):
    from .report import HEADING_LEVELS

    # A heading always opens on a portrait page: a landscape table is a place the
    # report visits for one table, never a place the next section starts in.
    ensure_orientation(doc, state, False)
    style = STYLE["heading"] % min(level, HEADING_LEVELS)
    heading = _para(doc, section_node.title, style=style)
    # Every heading is bookmarked, cited or not: it costs the document nothing,
    # and a sentence written later has something to cross-reference.
    anchor = getattr(section_node, "anchor", "")
    if anchor and state is not None:
        state["bookmark"] = state.get("bookmark", 0) + 1
        _bookmark(heading, anchor, state["bookmark"])
    _render_blocks(doc, section_node.blocks, state)
    for child in section_node.children:
        _render_section(doc, child, level + 1, state)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def render_docx(report, path, template=None):
    """Write ``report`` (a :class:`xslope.report.Report`) to ``path`` as a
    ``.docx``, built on ``template`` (the shipped default when None).

    Returns the path written.
    """
    template = template or DEFAULT_TEMPLATE
    doc = Document(template) if os.path.exists(template) else Document()

    meta = report.meta or {}
    props = doc.core_properties
    props.title = meta.get("title", "")
    props.subject = meta.get("project_number", "")
    props.author = meta.get("author", "") or props.author
    props.category = meta.get("organization", "")
    props.comments = "Generated by xslope."

    # A template opened as a starting document may carry sample body content; the
    # report replaces it entirely, and only the styles and page setup are kept.
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    _number_headings(doc)

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    _write_header(section, meta.get("title", ""))
    _write_footer(section, meta.get("date", ""))

    _title_page(doc, meta, section)
    _contents_page(doc, report)

    section = _begin_body(doc, meta)
    state = {"section": section, "landscape": False, "bookmark": 0,
             "title": meta.get("title", "")}
    for node in report.sections:
        _render_section(doc, node, 1, state)

    # Never end on a landscape page: the report closes in the orientation it
    # opened in, so a document appended to it starts straight.
    ensure_orientation(doc, state, False)

    doc.save(path)
    return path
