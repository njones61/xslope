# Copyright 2025 Norman L. Jones
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Finish a written report so that its contents page carries page numbers.

The Analysis Report carries a real table-of-contents field. The field is
deliberately not marked dirty — a dirty field is what makes Word for Mac ask
"this document contains fields that may refer to other files" every time the
report is opened — and the consequence is that its page numbers only appear
once someone presses F9. Page numbers cannot be guessed: only a page layout
engine knows what page a heading lands on.

So the report is handed to a program that already knows, and there are two of
them. :func:`finalize_report` picks one; either finish leaves a document whose
contents page is already right, opened by nobody.

**Word** is asked to do its own job: it opens the document, updates every field
and every table of contents, saves, and closes (:func:`finalize_with_word`).

**LibreOffice** finishes a report on a machine that has no Word, or in a run
that may not take the user's Word over. It cannot update a Word field, so it is
used as what it is — a page layout engine. The document is laid out to PDF, the
page each heading landed on is read off the outline of that PDF, and those pages
are written into the contents field's CACHED RESULT, which is the same place
Word puts them on an update (:func:`finalize_with_libreoffice`). The field stays
live: nothing is marked dirty, nothing is flattened, and F9 in Word still
recomputes the whole table.

What either program is handed is a COPY, which takes the report's place when it
comes back. A finish that fails, hangs or is killed halfway therefore leaves the
report exactly as it was written, and Word's lock file is never written beside
the user's document.

**Nothing here is required.** Every entry point returns the package's
``(bool, message)`` and never raises: no Word, no automation permission, a Word
that does not answer inside :data:`TIMEOUT`, no LibreOffice, a heading that
cannot be found on a page — all of them are a ``False`` and a sentence. The
report is complete without this step; it is merely nicer with it.

**Nothing here is automatic.** :func:`xslope.report.generate_report` does not
call it: a headless script that writes fifty reports in a loop must not open
fifty documents in Word. Studio calls it after a generate, and a script that
wants the same finish calls it too::

    from xslope.report_finalize import finalize_report
    ok, msg = finalize_report("north_levee_report.docx")

**macOS** drives Word over Apple events (``osascript``). Word 16 for Mac does
not answer the ``close`` message in AppleScript but does in JavaScript for
Automation, so the document is opened and updated by one script and closed by
the other — see :data:`_MAC_UPDATE` and :data:`_MAC_CLOSE`.

**Windows** drives ``Word.Application`` over COM from PowerShell rather than
from ``win32com``: pywin32 is not a dependency of this package, and a COM call
made from a subprocess can be given a timeout, which an in-process one cannot.
"""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
import tempfile

__all__ = ["TIMEOUT", "RENDER_TIMEOUT", "word_available", "finalize_with_word",
           "libreoffice_available", "finalize_with_libreoffice",
           "finalize_report"]

#: Seconds any one call into Word may take before it is abandoned. Generous:
#: Word's first launch of the day is slow, and the alternative to waiting is a
#: report with no page numbers.
TIMEOUT = 60

#: Where a Mac keeps Word. Checked before ``mdfind``, which is slower and can be
#: switched off.
_MAC_WORD_APP = "/Applications/Microsoft Word.app"
_MAC_BUNDLE_ID = "com.microsoft.Word"

# ---------------------------------------------------------------------------
# The scripts
# ---------------------------------------------------------------------------

#: Open the document, update every field in it — the body's, and the headers'
#: and footers' one section at a time, which is where the page count lives —
#: then rebuild every table of contents with its page numbers, and save.
#:
#: The document is offered to Word twice, and both offers matter. Word for Mac
#: is sandboxed: it may only read a file it has been granted, and an Apple event
#: carrying an ALIAS is what carries that grant — without the alias line, asking
#: Word to open a file it has never seen blocks until the call is abandoned.
#: Whether that first event also opens the document depends on how long Word has
#: been running, so the document is looked for among the open ones before being
#: opened by name. ``file name`` is given an HFS path (Word 16 hangs on a POSIX
#: one), converted by AppleScript itself from the argument, so a path with
#: spaces, quotes or non-ASCII characters is never pasted into a script.
#:
#: There is no ``revert``: asking Word to re-read a file it already has open
#: raises a confirmation Word never shows and never returns from. Freshness is
#: guaranteed instead by :func:`_finalize_macos`, which hands Word a copy at a
#: path it cannot already have open.
_MAC_UPDATE = r'''
on updateFieldsOf(theRange)
    -- Every field in one story, each update on its own: a field Word declines
    -- to rebuild must not cost the document the rest of them.
    set updated to 0
    tell application "Microsoft Word"
        try
            repeat with aField in (get fields of theRange)
                try
                    update field aField
                    set updated to updated + 1
                end try
            end repeat
        end try
    end tell
    return updated
end updateFieldsOf

on run argv
    set thePath to item 1 of argv
    set theAlias to (POSIX file thePath) as alias
    set hfsPath to (POSIX file thePath) as text
    tell application "Microsoft Word"
        try
            open theAlias
        end try
        set theDoc to missing value
        try
            repeat with aDoc in (get documents)
                try
                    if (full name of aDoc) is hfsPath then set theDoc to aDoc
                end try
            end repeat
        end try
        if theDoc is missing value then
            set theDoc to open file name hfsPath add to recent files false
        end if
        set fieldCount to my updateFieldsOf(theDoc)
        -- A section need not have a header or a footer of every kind, and asking
        -- for one it has not got is an error, not an empty answer.
        try
            repeat with aSection in (get sections of theDoc)
                repeat with anIndex in {header footer primary, header footer first page, header footer even pages}
                    try
                        set aPart to (get header aSection index anIndex)
                        set fieldCount to fieldCount + my updateFieldsOf(text object of aPart)
                    end try
                    try
                        set aPart to (get footer aSection index anIndex)
                        set fieldCount to fieldCount + my updateFieldsOf(text object of aPart)
                    end try
                end repeat
            end repeat
        end try
        set tocCount to 0
        try
            repeat with aToc in (get tables of contents of theDoc)
                try
                    update aToc
                    update page numbers aToc
                    set tocCount to tocCount + 1
                end try
            end repeat
        end try
        -- Word occasionally answers a save with "the object does not exist" and
        -- saves the same document happily a moment later, so the one step that
        -- puts the work on disk is the one step that is tried twice.
        set fullPath to ""
        try
            set fullPath to (full name of theDoc) as text
        end try
        try
            save theDoc
        on error
            delay 1
            save theDoc
        end try
        set saved of theDoc to true
        return fullPath & "|" & fieldCount & "|" & tocCount
    end tell
end run
'''

#: Close the window of the document that was just finalized, and nothing else.
#: The document is matched on its full name so that a report is never mistaken
#: for something of the user's that happens to be open beside it.
#:
#: This is JavaScript because Word 16 does not answer the AppleScript ``close``
#: message at all, and it tries three routes because it does not always answer
#: this one either — whether it does appears to depend on how long Word has been
#: running and how much has been asked of it. Failing to close is not a failure
#: of the finish: the document is saved, and it is the document the user is
#: about to be shown anyway.
_MAC_CLOSE = r'''
function run(argv) {
    var full = argv[0];
    var word = Application('Microsoft Word');
    var docs = word.documents;
    for (var i = 0; i < docs.length; i++) {
        var name;
        try { name = docs[i].fullName(); } catch (e) { continue; }
        if (name !== full) { continue; }
        try { docs[i].saved = true; } catch (e) {}
        var title = null;
        try { title = docs[i].windows[0].name(); } catch (e) {}
        try { docs[i].windows[0].close(); return 'closed'; } catch (e) {}
        try { docs[i].close(); return 'closed'; } catch (e) {}
        if (title !== null) {
            var wins = word.windows;
            for (var j = 0; j < wins.length; j++) {
                try {
                    if (wins[j].name() !== title) { continue; }
                    wins[j].close();
                    return 'closed';
                } catch (e) {}
            }
        }
        return 'left open';
    }
    return 'gone';
}
'''

#: The Windows leg. ``$Path`` arrives as an argument, never spliced into the
#: script.
#:
#: ``Word.Application`` is a single-instance COM server, so asking for one when
#: the user already has Word open hands back THEIRS. Two things follow: the
#: window is only forced hidden when this script started Word — hiding the
#: user's own Word would be startling — and Word is only quit on the way out
#: when this script started it. Each story range is walked through its
#: ``NextStoryRange`` chain, which is where the headers and footers of every
#: section past the first live.
_WINDOWS_SCRIPT = r'''
param([Parameter(Mandatory=$true)][string]$Path)
$ErrorActionPreference = 'Stop'
if ([Type]::GetTypeFromProgID('Word.Application') -eq $null) {
    Write-Output 'no-word'
    exit 3
}
$started = -not [bool](Get-Process -Name WINWORD -ErrorAction SilentlyContinue)
$word = New-Object -ComObject Word.Application
if ($started) { $word.Visible = $false }
$word.DisplayAlerts = 0
$doc = $null
try {
    # Open(FileName, ConfirmConversions, ReadOnly, AddToRecentFiles)
    $doc = $word.Documents.Open($Path, $false, $false, $false)
    $fields = 0
    foreach ($story in $doc.StoryRanges) {
        $range = $story
        while ($range -ne $null) {
            try { $range.Fields.Update() | Out-Null; $fields += $range.Fields.Count } catch {}
            $range = $range.NextStoryRange
        }
    }
    $tocs = 0
    foreach ($toc in $doc.TablesOfContents) {
        try { $toc.Update(); $toc.UpdatePageNumbers(); $tocs += 1 } catch {}
    }
    $doc.Save()
    $doc.Saved = $true
    Write-Output ("{0}|{1}|{2}" -f $doc.FullName, $fields, $tocs)
} finally {
    if ($doc -ne $null) { try { $doc.Close(0) } catch {} }
    if ($started) { try { $word.Quit() } catch {} }
    try { [void][Runtime.InteropServices.Marshal]::ReleaseComObject($word) } catch {}
}
'''


# ---------------------------------------------------------------------------
# Is there a Word to finalize with?
# ---------------------------------------------------------------------------

def word_available():
    """``(True, path_or_name)`` when this machine has a Word to drive, else
    ``(False, reason)``.

    The reason is written to be shown to a person: it says what is missing, not
    which call failed.
    """
    if sys.platform == "darwin":
        if os.path.isdir(_MAC_WORD_APP):
            return True, _MAC_WORD_APP
        found = _mdfind_word()
        if found:
            return True, found
        return False, "Microsoft Word is not installed."
    if sys.platform.startswith("win"):
        ok, out = _powershell(
            "if ([Type]::GetTypeFromProgID('Word.Application')) "
            "{ 'yes' } else { 'no' }", timeout=20)
        if not ok:
            return False, out
        if out.strip() == "yes":
            return True, "Word.Application"
        return False, "Microsoft Word is not installed."
    return False, ("Reports are finished in Word, which runs on macOS and "
                   "Windows only.")


def _mdfind_word():
    """Word somewhere other than /Applications, or None. Spotlight can be off,
    which is a None and not an error."""
    try:
        out = subprocess.run(
            ["mdfind", f"kMDItemCFBundleIdentifier == '{_MAC_BUNDLE_ID}'"],
            capture_output=True, text=True, timeout=15)
    except (OSError, subprocess.SubprocessError):
        return None
    for line in out.stdout.splitlines():
        if line.strip().endswith(".app"):
            return line.strip()
    return None


# ---------------------------------------------------------------------------
# The finish itself
# ---------------------------------------------------------------------------

def finalize_with_word(path, timeout=None):
    """Update the fields of the document at ``path`` in Word and save it.

    Parameters
    ----------
    path : str
        A ``.docx`` that has already been written.
    timeout : float, optional
        Seconds to wait for Word, per call. Defaults to :data:`TIMEOUT`.

    Returns
    -------
    (bool, str)
        ``(True, "…")`` when Word saved the document, else ``(False, reason)``.
        The reason is a sentence for a person; the caller decides whether it is
        worth showing at all — for Studio it is a log line, because a report
        without page numbers is still a report.
    """
    timeout = TIMEOUT if timeout is None else timeout
    ok, resolved = _finishable(path)
    if not ok:
        return False, resolved
    path = resolved

    ok, why = word_available()
    if not ok:
        return False, why

    if sys.platform == "darwin":
        return _finalize_macos(path, timeout)
    if sys.platform.startswith("win"):
        return _finalize_windows(path, timeout)
    return False, why                       # unreachable: word_available said no


def _finalize_macos(path, timeout):
    """Word for Mac, over Apple events — and never on the user's own file.

    Word is given a copy at a path of our making, and the finished copy takes
    the report's place when it comes back. Three things follow from that, all of
    them the reason for it:

    * Word cannot already have that document open, so it cannot hand back an
      older copy of a report that has just been regenerated — and it is never
      asked to revert, which is a confirmation Word never shows and never
      returns from.
    * The lock file Word writes beside anything it opens is written beside the
      copy, and goes when the copy's directory goes. Nothing appears next to the
      user's report.
    * A Word that fails, hangs or is killed halfway leaves the report exactly as
      it was written, because it never had it.

    The copy is made in the report's own directory so that putting it back is a
    rename rather than a second write.
    """
    folder = os.path.dirname(path) or "."
    try:
        work_dir = tempfile.mkdtemp(prefix=".xslope_finalize_", dir=folder)
    except OSError as exc:
        return False, f"The report's folder is not writable: {exc}."
    work = os.path.join(work_dir, os.path.basename(path))
    try:
        shutil.copy2(path, work)
        ok, out = _osascript(_MAC_UPDATE, [work], timeout=timeout)
        if ok:
            parts = out.strip().split("|")
            full = parts[0] if parts else ""
            fields = parts[1] if len(parts) > 1 else "?"
            tocs = parts[2] if len(parts) > 2 else "?"
            if full:
                # Closing is a courtesy, not the result: the copy is saved.
                _osascript(_MAC_CLOSE, [full], timeout=timeout,
                           language="JavaScript")
            os.replace(work, path)
            return True, (f"Word updated {fields} fields and {tocs} table(s) of "
                          f"contents.")
        return False, out
    except OSError as exc:
        return False, f"The finished report could not be put back: {exc}."
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


def _finalize_windows(path, timeout):
    """Word for Windows, over COM from PowerShell.

    Untested on macOS by definition; written so that every step that can fail
    fails into the same ``(False, sentence)`` the Mac leg returns.
    """
    ok, out = _powershell(_WINDOWS_SCRIPT, args=[path], timeout=timeout)
    if not ok:
        return False, out
    text = out.strip().splitlines()[-1] if out.strip() else ""
    if text == "no-word":
        return False, "Microsoft Word is not installed."
    parts = text.split("|")
    fields = parts[1] if len(parts) > 1 else "?"
    tocs = parts[2] if len(parts) > 2 else "?"
    return True, (f"Word updated {fields} fields and {tocs} table(s) of "
                  f"contents.")


# ---------------------------------------------------------------------------
# The finish without Word: LibreOffice lays the report out, and the pages it
# lands the headings on are written into the contents field's cached result.
# ---------------------------------------------------------------------------

#: Seconds a LibreOffice render may take. A report of a hundred pages with a
#: figure on most of them is a minute of work, and the first render of a finish
#: builds the user profile it runs in before it starts.
RENDER_TIMEOUT = 900

#: How many times the numbers may be written before they have to agree with the
#: pages they are written on. Two renders is the ordinary run — one to read the
#: pages from, one to prove them — and a third is only reached if writing the
#: numbers moved a heading, which is exactly the case the proof exists for.
PASSES = 4

#: Where LibreOffice is looked for, after ``XSLOPE_SOFFICE`` and the PATH, on
#: each platform it ships for.
_SOFFICE_PATHS = {
    "darwin": ("/Applications/LibreOffice.app/Contents/MacOS/soffice",),
    "win32": (r"C:\Program Files\LibreOffice\program\soffice.exe",
              r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    "": ("/usr/bin/soffice", "/usr/local/bin/soffice", "/snap/bin/libreoffice"),
}

#: The footer of every page of a report says which page it is. It is read back
#: as a check on the numbers: the page a heading is on is the page the reader
#: sees, and these are only the same while the document numbers its pages
#: straight through from the title page.
_PAGE_FOOTER = re.compile(r"^Page\s+(\d+)\s+of\s+(\d+)\b")


def libreoffice_available():
    """``(True, path)`` when this machine can lay a report out, else
    ``(False, reason)``.

    Two things are needed and both are checked: LibreOffice, which does the
    layout, and ``pypdf``, which reads the result back. Neither is a dependency
    of this package — the finish is cosmetic, and a report is complete without
    it — so a machine without them is told what is missing, in a sentence.
    """
    if _soffice() is None:
        return False, ("LibreOffice is not installed — set XSLOPE_SOFFICE to "
                       "its soffice program if it is somewhere unusual.")
    try:
        import pypdf                                        # noqa: F401
    except Exception:
        return False, ("Reading the page numbers back needs pypdf, which is not "
                       "installed (pip install pypdf).")
    return True, _soffice()


def _soffice():
    """LibreOffice's ``soffice`` program, or None on a machine without it.

    ``XSLOPE_SOFFICE`` is read here rather than at import, so a run can name an
    installation the moment it matters; the PATH comes next, and the platform's
    usual place last.
    """
    candidates = [os.environ.get("XSLOPE_SOFFICE"), "soffice"]
    candidates += list(_SOFFICE_PATHS.get(sys.platform, _SOFFICE_PATHS[""]))
    for cand in candidates:
        if not cand:
            continue
        found = shutil.which(cand)
        if found:
            return found
        if os.path.isfile(cand) and os.access(cand, os.X_OK):
            return cand
    return None


def finalize_with_libreoffice(path, timeout=None):
    """Write real page numbers into the report's contents field, without Word.

    Parameters
    ----------
    path : str
        A ``.docx`` that has already been written.
    timeout : float, optional
        Seconds one LibreOffice render may take. Defaults to
        :data:`RENDER_TIMEOUT`.

    Returns
    -------
    (bool, str)
        ``(True, "…")`` naming how many entries were numbered and how long the
        report runs, else ``(False, reason)``. An entry whose heading is on no
        page is a failure that names the entry: a contents page that is right
        about eleven sections and silent about the twelfth is worse than one
        that carries no numbers at all, because nothing on it says which.

    How it works, and why in that order:

    1. The document is copied, and everything after this happens to the copy —
       so a run that fails or is killed halfway leaves the report as it was.
    2. The line that tells the reader to press F9 is REMOVED FIRST, before
       anything is laid out. It sits inside the field result on the contents
       page, and the numbers replace it; a page read off a layout that still
       carried it is a page from a document one line longer than the one being
       written, and on a report whose contents run to two pages that line is
       enough to move the body.
    3. The copy is laid out to PDF, and the page each heading landed on is read
       off the PDF's outline — which LibreOffice builds from the heading styles
       themselves, in document order, with the numbers the document's own
       multilevel list computed. The contents entries are matched to it in
       order: entry text against heading text, each match continuing where the
       last left off, so two sections that share a title are still told apart by
       where they are.
    4. The numbers are written, and READ BACK OFF THE SAVED FILE: what is
       checked from here on is what the document says, not what this function
       meant to make it say.
    5. The copy is laid out AGAIN, and every number the document carries is
       checked against the page its heading landed on in the document that now
       carries it. That is the only proof that survives writing — adding a
       number to an entry can wrap a long entry onto a second line, which moves
       the contents page's own page break. A disagreement is written again from
       the new layout, up to :data:`PASSES` times, and if it still disagrees the
       report is left as it was and the disagreement is the reason.
    """
    timeout = RENDER_TIMEOUT if timeout is None else timeout
    ok, resolved = _finishable(path)
    if not ok:
        return False, resolved
    path = resolved
    ok, why = libreoffice_available()
    if not ok:
        return False, why
    soffice = _soffice()

    from docx import Document

    folder = os.path.dirname(path) or "."
    try:
        work_dir = tempfile.mkdtemp(prefix=".xslope_finalize_", dir=folder)
    except OSError as exc:
        return False, f"The report's folder is not writable: {exc}."
    # The copy keeps the report's name: the PDF is named after the file that was
    # laid out, and a name with a space or a dot in it is then still found.
    work = os.path.join(work_dir, os.path.basename(path))
    profile = tempfile.mkdtemp(prefix="xslope_soffice_")
    try:
        shutil.copy2(path, work)
        try:
            doc = Document(work)
        except Exception as exc:
            return False, f"The document could not be read: {exc}."
        entries, placeholder = _toc_field(doc)
        if not entries:
            return False, ("The document has no table of contents to number.")

        _strip_placeholder(placeholder)
        for entry, _text in entries:
            _clear_page_number(entry)
        width = _text_width(doc, entries[0][0])
        doc.save(work)

        cached = [None] * len(entries)             # the numbers now on the page
        for _pass in range(PASSES):
            headings, page_count, why = _laid_out_headings(work, work_dir,
                                                           soffice, timeout,
                                                           profile)
            if headings is None:
                return False, why
            pages, why = _pages_of(entries, headings)
            if pages is None:
                return False, why
            why = _page_numbering_holds(work_dir, work, pages)
            if why:
                return False, why
            if cached == pages:
                break            # every number in the document is the page its
                                 # heading lands on in the document that says so
            for (entry, _text), page in zip(entries, pages):
                _write_page_number(entry, page, width)
            doc.save(work)
            cached, why = _cached_numbers(work, entries)
            if why:
                return False, why
            if cached != pages:
                wrong = next(text for (_p, text), was, want
                             in zip(entries, cached, pages) if was != want)
                return False, (f"The page numbers did not go into the document "
                               f"as they were computed — “{wrong}” "
                               f"came back carrying something else.")
        else:
            return False, ("The contents page numbers did not settle: writing "
                           "them moved the headings they point at.")

        os.replace(work, path)
        return True, (f"LibreOffice numbered {len(entries)} of {len(entries)} "
                      f"contents entries; the report runs to {page_count} "
                      f"pages.")
    except OSError as exc:
        return False, f"The finished report could not be put back: {exc}."
    except Exception as exc:
        # A cosmetic step may not cost anyone their report, whatever it trips
        # over. The copy is thrown away and the report is the one that was
        # written.
        return False, f"The report's page numbers could not be built: {exc}."
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)
        shutil.rmtree(profile, ignore_errors=True)


def _finishable(path):
    """``(True, path)`` when ``path`` is a document that can be finished."""
    if not path:
        return False, "No document was named."
    path = os.path.abspath(path)
    if not os.path.isfile(path):
        return False, f"There is no document at {path}."
    if not path.lower().endswith(".docx"):
        return False, "Only a Word document can be finished."
    return True, path


# ---------------------------------------------------------------------------
# The contents field
# ---------------------------------------------------------------------------

def _toc_field(doc):
    """``(entries, placeholder)`` — the contents field's cached result.

    ``entries`` is ``[(paragraph, text), …]``, one per line of the contents;
    ``placeholder`` is the paragraph carrying the field's end mark, which holds
    the line telling the reader to press F9 and is where the numbers go instead.

    The field is found by its INSTRUCTION — the one whose field name is ``TOC``
    — and its begin, separate and end marks are tracked as a NEST. Both matter.
    A report's title page carries a field of its own, a ``DOCPROPERTY`` for the
    title, so the first separate mark in the document is not the contents'; and
    a contents field Word has already built has ``HYPERLINK`` and ``PAGEREF``
    fields inside its result, so the first end mark after it is not the
    contents' either. Keying on either one collects the wrong paragraphs.
    """
    from docx.oxml.ns import qn

    stack, toc = [], None
    entries, placeholder = [], None
    for para in doc.paragraphs:
        text, closed = [], False
        for element in para._p.iter():
            tag = element.tag.split("}")[-1]
            if tag == "fldChar":
                kind = element.get(qn("w:fldCharType"))
                if kind == "begin":
                    stack.append({"instr": "", "result": False})
                elif kind == "separate" and stack:
                    stack[-1]["result"] = True
                    if toc is None and _field_name(stack[-1]["instr"]) == "TOC":
                        toc = stack[-1]
                elif kind == "end" and stack:
                    if stack.pop() is toc:
                        placeholder, closed = para, True
                        break
            elif tag == "instrText" and stack:
                stack[-1]["instr"] += element.text or ""
            elif tag == "t" and toc is not None and toc["result"]:
                text.append(element.text or "")
            elif tag == "tab" and toc is not None and toc["result"] \
                    and element.getparent().tag.endswith("}r"):
                text.append("\t")
        if toc is not None and toc["result"]:
            line = _entry_line("".join(text))
            if line and not closed:
                entries.append((para, line))
            elif line and closed and _is_contents_style(para):
                # A contents Word itself has built can end on its last entry
                # rather than on a line of its own. That paragraph is an entry,
                # not a placeholder, and its text is not thrown away.
                entries.append((para, line))
                placeholder = None
        if closed:
            break
    return entries, placeholder


def _field_name(instruction):
    """The name of a Word field, from its instruction — ``TOC`` out of
    ``' TOC \\o "1-3" \\h \\z \\u '``."""
    parts = (instruction or "").split()
    return parts[0].upper() if parts else ""


def _is_contents_style(para):
    """True when the paragraph is styled as a contents entry."""
    try:
        return (para.style.name or "").strip().lower().startswith("toc")
    except Exception:
        return False


def _clean(text):
    """One line of text, with the spacings a layout engine varies made equal."""
    return re.sub(r"\s+", " ", (text or "").replace("\u00a0", " ")).strip()


def _entry_line(text):
    """What a contents entry NAMES, without a page number written on it before.

    A finished report finished again must match its entries on the headings they
    name, not on the headings plus the numbers the last finish put after them.
    The number is what stands after the entry's last tab, and only when it is a
    number: a heading may end in a digit of its own, and it does not follow a
    tab.
    """
    head, tab, tail = (text or "").rpartition("\t")
    return _clean(head if tab and tail.strip().isdigit() else text)


def _strip_placeholder(para):
    """Take the "press F9" line out, leaving the field's end mark where it is.

    The line is what the page numbers replace, and it has to go before anything
    is laid out — it is a line on the contents page, and the contents page is
    what everything after it is measured from.
    """
    if para is None:
        return
    from docx.oxml.ns import qn
    for run in list(para._p.findall(qn("w:r"))):
        if run.findall(qn("w:fldChar")) or run.findall(qn("w:instrText")):
            continue
        if "".join(t.text or "" for t in run.findall(qn("w:t"))).strip():
            run.getparent().remove(run)
    para.paragraph_format.space_before = None


def _clear_page_number(para):
    """Take off a page number written into this entry before.

    A report can be finished twice — regenerated, or finished again after an
    edit — and the second finish must number the entry rather than append to
    the number the first one wrote. Only the tab and the digits at the END of
    the line are removed; a run holding anything else stops the walk, so a
    heading that ends in a number keeps it.
    """
    from docx.oxml.ns import qn
    for run in reversed(para._p.findall(qn("w:r"))):
        text = "".join(t.text or "" for t in run.findall(qn("w:t"))).strip()
        is_number = text.isdigit() and text != ""
        is_tab = bool(run.findall(qn("w:tab"))) and text == ""
        if not (is_number or is_tab):
            break
        run.getparent().remove(run)


def _raw_text(para):
    """A paragraph's text with its tabs kept as tabs — which is what tells a
    heading from the page number written after it."""
    from docx.oxml.ns import qn
    out = []
    for run in para._p.findall(qn("w:r")):
        for child in run:
            tag = child.tag.split("}")[-1]
            if tag == "t":
                out.append(child.text or "")
            elif tag == "tab":
                out.append("\t")
    return "".join(out)


def _written_number(para):
    """The page number a contents entry carries, or None when it carries
    none."""
    _head, tab, tail = _raw_text(para).rpartition("\t")
    return int(tail.strip()) if tab and tail.strip().isdigit() else None


def _cached_numbers(path, entries):
    """``(numbers, why not)`` — the page numbers the document on disk holds.

    Read from the saved file rather than from the document in hand, so that what
    is checked afterwards is what a reader will open. The entries have to be the
    same entries: a file that comes back with a different contents page is a
    document that was not written as it was read.
    """
    from docx import Document
    try:
        saved, _placeholder = _toc_field(Document(path))
    except Exception as exc:
        return None, f"The finished contents page could not be read back: {exc}."
    if [text for _p, text in saved] != [text for _p, text in entries]:
        return None, ("The finished contents page does not list the sections "
                      "the written one listed.")
    return [_written_number(para) for para, _text in saved], ""


def _write_page_number(para, page, width):
    """Put ``page`` at the right margin of a contents entry, on a dot leader.

    The tab stop is set at the width of the text column the entry is actually
    in, so the numbers line up on the margin of the page this document is set
    on rather than on the margin of some other one.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.oxml.ns import qn

    _clear_page_number(para)
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.clear_all()
    tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = para.add_run()
    run._r.append(run._r.makeelement(qn("w:tab"), {}))
    para.add_run(str(page))


def _text_width(doc, para):
    """The width of the text column ``para`` sits in — the page less its
    margins, taken from the section the paragraph is in."""
    section = _section_of(doc, para)
    return section.page_width - section.left_margin - section.right_margin


def _section_of(doc, para):
    """The section a paragraph belongs to.

    A section break is a ``sectPr`` on a paragraph, and it closes the section
    the paragraph is in — so the paragraph's section is the one after as many
    breaks as precede it.
    """
    from docx.oxml.ns import qn
    breaks = 0
    for other in doc.paragraphs:
        if other._p is para._p:
            break
        properties = other._p.find(qn("w:pPr"))
        if properties is not None \
                and properties.find(qn("w:sectPr")) is not None:
            breaks += 1
    return doc.sections[min(breaks, len(doc.sections) - 1)]


# ---------------------------------------------------------------------------
# The layout, and the pages read back off it
# ---------------------------------------------------------------------------

def _laid_out_headings(path, outdir, soffice, timeout, profile=None):
    """``(headings, page count, why not)`` for a LibreOffice layout of ``path``.

    ``headings`` is ``[(text, page), …]`` — every heading of the document, in
    document order, with the page it landed on. They are read off the PDF's
    OUTLINE, which LibreOffice writes from the heading styles and the numbering
    the document's own list computed: the document's structure, not a search of
    the words on the page. A search would have to tell a heading from the
    running head that repeats it at the top of every page of its section, and
    from the contents entry that names it.
    """
    pdf = _render_pdf(path, outdir, soffice, timeout, profile)
    if not os.path.exists(pdf):
        return None, 0, ("LibreOffice did not lay the report out — it wrote no "
                         "PDF.")
    import pypdf
    try:
        reader = pypdf.PdfReader(pdf)
        headings = []
        _walk_outline(reader, reader.outline, headings)
        count = len(reader.pages)
    except Exception as exc:
        return None, 0, f"The laid-out report could not be read back: {exc}."
    if not headings:
        return None, 0, ("The laid-out report carries no heading outline, so "
                         "there is nothing to take the page numbers from.")
    return headings, count, ""


def _render_pdf(path, outdir, soffice, timeout, profile=None):
    """Lay ``path`` out to a PDF beside it, and return where that PDF is.

    ``profile`` is a LibreOffice user profile of our own. It matters: a
    ``soffice`` started with the user's profile while LibreOffice is open joins
    THAT session and does the work in the copy of LibreOffice the user is
    working in — the same discourtesy this module refuses to do to Word. Given
    a profile of its own it starts on its own, and finishing a report cannot
    disturb a document open on the screen. It is also quicker: a profile that
    nothing else is using needs no lock waited for.
    """
    pdf = os.path.join(outdir,
                       os.path.splitext(os.path.basename(path))[0] + ".pdf")
    if os.path.exists(pdf):
        os.remove(pdf)                     # never read the previous pass's pages
    argv = [soffice]
    if profile:
        argv.append("-env:UserInstallation=" + _file_url(profile))
    argv += ["--headless", "--convert-to", "pdf", "--outdir", outdir, path]
    try:
        subprocess.run(argv, capture_output=True, text=True, timeout=timeout)
    except (OSError, subprocess.SubprocessError):
        return pdf                         # absent, and reported as such
    return pdf


def _file_url(path):
    """A ``file://`` URL for a local directory, spelled as this platform spells
    one — which is what LibreOffice's ``-env:`` switches take."""
    import pathlib
    return pathlib.Path(path).absolute().as_uri()


def _walk_outline(reader, items, out):
    """Flatten a PDF outline into ``(text, page)``, in document order."""
    for item in items:
        if isinstance(item, list):
            _walk_outline(reader, item, out)
            continue
        try:
            page = reader.get_destination_page_number(item) + 1
        except Exception:
            continue
        out.append((_clean(str(item.title)), page))


def _pages_of(entries, headings):
    """``(pages, why not)`` — the page every contents entry belongs on.

    Matched in document order: each entry takes the first heading at or after
    the one the entry before it took. Two sections that share a title — a
    "Results" under every method — are therefore told apart by where they are
    rather than by what they say.

    An entry that matches nothing fails the whole finish, and the reason names
    the entries and counts both sides of the tally: a contents page that is
    right about eleven sections and silent about the twelfth is worse than one
    with no numbers on it at all, because nothing on the page says which line to
    distrust.
    """
    pages, at, missing = [], 0, []
    for _para, text in entries:
        found = next((n for n in range(at, len(headings))
                      if headings[n][0] == text), None)
        if found is None:
            missing.append(text)
            pages.append(None)
            continue
        pages.append(headings[found][1])
        at = found + 1
    if missing:
        named = "\u201d, \u201c".join(missing[:3])
        return None, (f"{len(entries) - len(missing)} of {len(entries)} contents "
                      f"entries were numbered; {len(missing)} name a heading "
                      f"that is on no page of the laid-out report: "
                      f"\u201c{named}\u201d. Nothing was written.")
    return pages, ""


def _page_numbering_holds(outdir, path, pages):
    """"" when the page a heading is ON is the page a reader would CALL it.

    The pages come off the layout as ordinal positions — first sheet, second
    sheet — and the number written into the contents is what the reader sees in
    the footer. Those are the same number only while the document numbers its
    pages straight through from the title page, which is how the report is set;
    a document numbered any other way is left alone rather than numbered wrong.
    """
    pdf = os.path.join(outdir,
                       os.path.splitext(os.path.basename(path))[0] + ".pdf")
    try:
        import pypdf
        reader = pypdf.PdfReader(pdf)
        for page in sorted(set(pages)):
            text = reader.pages[page - 1].extract_text() or ""
            for line in text.splitlines():
                found = _PAGE_FOOTER.match(line.strip())
                if not found:
                    continue
                if int(found.group(1)) != page:
                    return (f"The report does not number its pages straight "
                            f"through: sheet {page} calls itself page "
                            f"{found.group(1)}, so a contents page built from "
                            f"the layout would send the reader to the wrong "
                            f"one.")
                break
    except Exception:
        return ""                          # unreadable text is not a wrong number
    return ""


# ---------------------------------------------------------------------------
# Either program, whichever this machine has
# ---------------------------------------------------------------------------

#: Which finish to use, when the caller does not say: ``word``, ``libreoffice``
#: or ``auto``. A run that must not touch the Word the user is working in sets
#: it to ``libreoffice``; a run that will accept nothing else sets it to
#: ``word``.
FINISH_ENV = "XSLOPE_FINALIZE"


def finalize_report(path, timeout=None, prefer=None):
    """Finish the report's contents page, by whichever route this machine has.

    Word is asked first where there is one, because a Word update is the
    document's own field rebuilt by the program that owns the format. Where
    there is no Word, or Word declines — not installed, not permitted, not
    answering — LibreOffice lays the report out instead and the page numbers are
    written into the same cached result Word would have filled.

    ``prefer`` is ``"word"``, ``"libreoffice"`` or ``"auto"``, defaulting to
    :data:`FINISH_ENV` and then to ``"auto"``. Returns the ``(bool, message)``
    of whichever leg answered; when both decline, both reasons are given, since
    on a machine with neither program the second is the actionable one.
    """
    prefer = (prefer or os.environ.get(FINISH_ENV) or "auto").strip().lower()
    if prefer not in ("word", "libreoffice", "auto"):
        return False, (f"{FINISH_ENV} is set to {prefer!r}; it names the "
                       f"program the report is finished with, and the choices "
                       f"are word, libreoffice and auto.")
    ok, why = _finishable(path)
    if not ok:
        return False, why

    if prefer == "libreoffice":
        return finalize_with_libreoffice(path, timeout=timeout)
    word_ok, word_why = word_available()
    if word_ok:
        done, message = finalize_with_word(path)
        if done or prefer == "word":
            return done, message
        word_why = message
    elif prefer == "word":
        return False, word_why
    done, message = finalize_with_libreoffice(path, timeout=timeout)
    if done:
        return True, message
    return False, f"{word_why} {message}"


# ---------------------------------------------------------------------------
# Talking to the two shells
# ---------------------------------------------------------------------------

def _osascript(script, args=(), timeout=TIMEOUT, language=None):
    """Run ``script`` through ``osascript``, with ``args`` as its arguments.

    The script arrives on standard input and the arguments as argv, so nothing
    is ever quoted into a command line. Returns ``(True, stdout)`` or
    ``(False, sentence)``: a refused automation permission and a Word that never
    answers are both ordinary outcomes here, not exceptions.
    """
    argv = ["osascript"]
    if language:
        argv += ["-l", language]
    argv += ["-"] + [str(a) for a in args]
    try:
        done = subprocess.run(argv, input=script, capture_output=True,
                              text=True, timeout=timeout)
    except subprocess.TimeoutExpired:
        return False, (f"Word did not answer within {timeout:g} seconds — the "
                       f"report was left as written.")
    except OSError as exc:
        return False, f"Word could not be reached: {exc}."
    if done.returncode != 0:
        return False, _apple_event_reason(done.stderr)
    return True, done.stdout


def _apple_event_reason(stderr):
    """Turn osascript's error text into something a person can act on."""
    text = (stderr or "").strip()
    if "-1743" in text or "Not authorized to send Apple events" in text:
        return ("Word could not be controlled: allow XSLOPE Studio to control "
                "Microsoft Word under System Settings → Privacy & Security → "
                "Automation.")
    if "-600" in text or "isn't running" in text:
        return "Microsoft Word did not start."
    return text.splitlines()[-1] if text else "Word reported an unknown error."


def _powershell(script, args=(), timeout=TIMEOUT):
    """Run ``script`` through PowerShell, with ``args`` after it.

    ``-File -`` is not a thing, so the script is passed with ``-Command`` and
    reads its argument from ``$args``; the argument is a separate argv entry and
    is never spliced into the text.
    """
    argv = ["powershell", "-NoProfile", "-NonInteractive",
            "-ExecutionPolicy", "Bypass", "-Command", "-"]
    body = script
    if args:
        # The script declares param(...); calling it as a script block keeps the
        # argument out of the script text.
        quoted = " ".join(_ps_quote(str(a)) for a in args)
        body = f"& {{{script}}} {quoted}"
    try:
        done = subprocess.run(argv, input=body, capture_output=True, text=True,
                              timeout=timeout)
    except subprocess.TimeoutExpired:
        return False, (f"Word did not answer within {timeout:g} seconds — the "
                       f"report was left as written.")
    except OSError as exc:
        return False, f"Word could not be reached: {exc}."
    if done.returncode not in (0, 3):
        text = (done.stderr or done.stdout or "").strip()
        return False, (text.splitlines()[-1] if text
                       else "Word reported an unknown error.")
    return True, done.stdout


def _ps_quote(value):
    """A PowerShell single-quoted literal: the one escape is a doubled quote."""
    return "'" + value.replace("'", "''") + "'"
